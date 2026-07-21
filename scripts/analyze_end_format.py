"""Compare completed batch outputs with the correct template, excluding letterhead."""

from __future__ import annotations

from collections import Counter, defaultdict
from difflib import SequenceMatcher
import hashlib
import json
from pathlib import Path
import re
import unicodedata

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "test_docx" / "end_docx"
TEMPLATE = ROOT / "test_docx" / "correct_docx" / "005班子对照检查材料_正确格式.docx"

PARAGRAPH_FIELDS = (
    "style", "font", "size_pt", "bold", "italic", "underline", "alignment",
    "first_indent", "left_indent", "right_indent", "line_spacing", "space_before",
    "space_after", "page_break_before", "keep_with_next", "keep_together", "page_breaks",
)


def _length(value) -> int:
    return int(value) if value else 0


def _first_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def _style_value(style, container: str, field: str):
    current = style
    while current is not None:
        source = getattr(current, container)
        value = getattr(source, field)
        if value is not None:
            return value
        current = current.base_style
    return None


def _paragraph_value(paragraph, field: str):
    value = getattr(paragraph.paragraph_format, field)
    if value is not None:
        return value
    return _style_value(paragraph.style, "paragraph_format", field) if paragraph.style else None


def _font_value(paragraph, run, field: str):
    value = getattr(run.font, field) if run else None
    if value is not None:
        return value
    return _style_value(paragraph.style, "font", field) if paragraph.style else None


def _normalize(text: str) -> str:
    value = unicodedata.normalize("NFKC", text or "")
    value = re.sub(r"\s+", "", value)
    return value


def _text_id(text: str) -> str:
    return hashlib.sha256(_normalize(text).encode("utf-8")).hexdigest()[:16]


def _paragraph_snapshot(paragraph, index: int) -> dict:
    run = _first_run(paragraph)
    page_breaks = len(paragraph._element.findall(f".//{qn('w:br')}[@{qn('w:type')}='page']"))
    font_name = _font_value(paragraph, run, "name")
    font_size = _font_value(paragraph, run, "size")
    alignment = paragraph.alignment
    if alignment is None:
        alignment = _style_value(paragraph.style, "paragraph_format", "alignment") if paragraph.style else None
    return {
        "index": index,
        "text_id": _text_id(paragraph.text),
        "length": len(paragraph.text),
        "style": paragraph.style.name if paragraph.style else "",
        "font": font_name or "",
        "size_pt": round(font_size.pt, 3) if font_size else None,
        "bold": bool(_font_value(paragraph, run, "bold")),
        "italic": bool(_font_value(paragraph, run, "italic")),
        "underline": bool(_font_value(paragraph, run, "underline")),
        "alignment": str(alignment) if alignment is not None else "",
        "first_indent": _length(_paragraph_value(paragraph, "first_line_indent")),
        "left_indent": _length(_paragraph_value(paragraph, "left_indent")),
        "right_indent": _length(_paragraph_value(paragraph, "right_indent")),
        "line_spacing": str(_paragraph_value(paragraph, "line_spacing")) if _paragraph_value(paragraph, "line_spacing") is not None else "",
        "space_before": _length(_paragraph_value(paragraph, "space_before")),
        "space_after": _length(_paragraph_value(paragraph, "space_after")),
        "page_break_before": bool(_paragraph_value(paragraph, "page_break_before")),
        "keep_with_next": bool(_paragraph_value(paragraph, "keep_with_next")),
        "keep_together": bool(_paragraph_value(paragraph, "keep_together")),
        "page_breaks": page_breaks,
    }


def _section_snapshot(section) -> dict:
    return {
        "page_width": _length(section.page_width),
        "page_height": _length(section.page_height),
        "top_margin": _length(section.top_margin),
        "bottom_margin": _length(section.bottom_margin),
        "left_margin": _length(section.left_margin),
        "right_margin": _length(section.right_margin),
        "header_distance": _length(section.header_distance),
        "footer_distance": _length(section.footer_distance),
        "orientation": str(section.orientation),
    }


def snapshot(path: Path, template: bool = False) -> dict:
    document = Document(path)
    paragraphs = [_paragraph_snapshot(paragraph, index) for index, paragraph in enumerate(document.paragraphs)]
    if template:
        paragraphs = [
            item for item in paragraphs
            if not item["style"].startswith("Docxtool Letterhead")
            and not (item["length"] == 0 and item["style"] == "Docxtool Role Name")
        ]
    return {
        "paragraphs": paragraphs,
        "sections": [_section_snapshot(section) for section in document.sections],
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "headers": sum(len(section.header.paragraphs) for section in document.sections),
        "footers": sum(len(section.footer.paragraphs) for section in document.sections),
    }


def compare_file(actual: dict, expected: dict) -> dict:
    actual_ids = [item["text_id"] for item in actual["paragraphs"]]
    expected_ids = [item["text_id"] for item in expected["paragraphs"]]
    matcher = SequenceMatcher(a=expected_ids, b=actual_ids, autojunk=False)
    field_differences = []
    unmatched_actual = 0
    unmatched_expected = 0
    matched = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for expected_item, actual_item in zip(expected["paragraphs"][i1:i2], actual["paragraphs"][j1:j2]):
                matched += 1
                for field in PARAGRAPH_FIELDS:
                    if actual_item[field] != expected_item[field]:
                        field_differences.append({
                            "field": field,
                            "actual_index": actual_item["index"],
                            "expected_index": expected_item["index"],
                            "actual": actual_item[field],
                            "expected": expected_item[field],
                        })
        else:
            unmatched_expected += i2 - i1
            unmatched_actual += j2 - j1
    document_differences = []
    for field in ("tables", "inline_shapes", "headers", "footers"):
        if actual[field] != expected[field]:
            document_differences.append({"field": field, "actual": actual[field], "expected": expected[field]})
    if len(actual["sections"]) != len(expected["sections"]):
        document_differences.append({"field": "section_count", "actual": len(actual["sections"]), "expected": len(expected["sections"])})
    for index, (actual_section, expected_section) in enumerate(zip(actual["sections"], expected["sections"])):
        for field in actual_section:
            if actual_section[field] != expected_section[field]:
                document_differences.append({"field": f"section_{index}_{field}", "actual": actual_section[field], "expected": expected_section[field]})
    return {
        "matched_paragraphs": matched,
        "unmatched_actual_paragraphs": unmatched_actual,
        "unmatched_expected_paragraphs": unmatched_expected,
        "format_differences": field_differences,
        "document_differences": document_differences,
    }


def main() -> int:
    expected = snapshot(TEMPLATE, template=True)
    records = []
    field_counts: Counter[str] = Counter()
    transition_counts: dict[str, Counter[str]] = defaultdict(Counter)
    file_counts = {}
    for path in sorted(OUTPUT.glob("*_排版结果.docx")):
        result = compare_file(snapshot(path), expected)
        differences = result["format_differences"] + result["document_differences"]
        for item in result["format_differences"]:
            field_counts[item["field"]] += 1
            transition_counts[item["field"]][f"{item['actual']!r} -> {item['expected']!r}"] += 1
        for item in result["document_differences"]:
            field_counts[item["field"]] += 1
        file_counts[path.name[:3]] = len(differences)
        records.append({"编号": path.name[:3], "文件名": path.name, **result})
    report = {
        "文件总数": len(records),
        "存在格式差异文件数": sum(bool(item["format_differences"] or item["document_differences"]) for item in records),
        "格式字段差异统计": dict(field_counts),
        "常见实际值到模板值": {field: dict(counter.most_common(8)) for field, counter in transition_counts.items()},
        "各文件差异数量": file_counts,
        "说明": "模板版头及空职务段已排除；段落使用规范化文本对齐，只对文本相同的对应段落比较格式。未匹配段落单独统计，不作为格式差异。",
        "结果": records,
        "视觉渲染检查": "未执行",
    }
    (OUTPUT / "排版结果格式差异分析.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        f"文件总数：{report['文件总数']}",
        f"存在格式差异文件数：{report['存在格式差异文件数']}",
        "格式字段差异统计：" + json.dumps(report["格式字段差异统计"], ensure_ascii=False),
        "说明：" + report["说明"],
        "",
        "各文件：",
    ]
    lines.extend(
        f"{item['编号']} | 格式差异={len(item['format_differences'])} | 文档级差异={len(item['document_differences'])} | 未匹配输出段={item['unmatched_actual_paragraphs']} | 未匹配模板段={item['unmatched_expected_paragraphs']}"
        for item in records
    )
    (OUTPUT / "排版结果格式差异分析.txt").write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps({"files": len(records), "with_format_differences": report["存在格式差异文件数"], "field_counts": report["格式字段差异统计"]}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
