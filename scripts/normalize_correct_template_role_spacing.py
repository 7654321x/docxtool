"""Normalize the retained correct template to use role paragraph spacing."""

from pathlib import Path
import tempfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "test_docx" / "correct_docx" / "005班子对照检查材料_正确格式.docx"


def main() -> int:
    document = Document(TARGET)
    role = next(paragraph for paragraph in document.paragraphs if paragraph.style and paragraph.style.name == "Docxtool Role Name" and paragraph.text.strip())
    role.paragraph_format.space_before = Pt(28)
    role.paragraph_format.space_after = Pt(28)
    properties = role._p.get_or_add_pPr()
    spacing = properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        properties.append(spacing)
    spacing.set(qn("w:before"), "560")
    spacing.set(qn("w:after"), "560")
    spacing.set(qn("w:beforeLines"), "100")
    spacing.set(qn("w:afterLines"), "100")
    following = role._p.getnext()
    if following is not None and following.tag.endswith("}p"):
        empty = next((paragraph for paragraph in document.paragraphs if paragraph._p is following), None)
        if empty is not None and not empty.text.strip() and empty.style and empty.style.name == "Docxtool Role Name":
            following.getparent().remove(following)
    with tempfile.NamedTemporaryFile(suffix=".docx", dir=TARGET.parent, delete=False) as handle:
        temporary = Path(handle.name)
    document.save(temporary)
    temporary.replace(TARGET)
    print(TARGET)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
