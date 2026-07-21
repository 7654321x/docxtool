"""Generate reproducible, lightly corrupted DOCX recognition fixtures."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
import random
import re
import tempfile
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "005班子对照检查材料.docx"
OUTPUT = ROOT / "test_docx" / "strat_docx"
SEED = 20260721
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def text_of(document: Document) -> str:
    pieces = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        pieces.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(pieces)


def repair_source(source: Path, target: Path) -> None:
    """Remove only invalid package relationships from a temporary copy."""
    with zipfile.ZipFile(source) as source_zip, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for info in source_zip.infolist():
            data = source_zip.read(info.filename)
            if info.filename == "word/_rels/document.xml.rels":
                root = ET.fromstring(data)
                for relationship in list(root):
                    if relationship.attrib.get("Target") == "../NULL":
                        root.remove(relationship)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target_zip.writestr(info, data)


def first_run(paragraph):
    if not paragraph.runs:
        return paragraph.add_run()
    return paragraph.runs[0]


def set_font(paragraph, name: str | None = None, size: float | None = None, bold: bool | None = None, underline: bool | None = None, italic: bool | None = None):
    for run in paragraph.runs or [first_run(paragraph)]:
        if name:
            run.font.name = name
            run._element.get_or_add_rPr().get_or_add_rFonts().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", name)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if underline is not None:
            run.underline = underline
        if italic is not None:
            run.italic = italic


def nonempty_indices(document: Document) -> list[int]:
    return [index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip()]


def title_indices(document: Document) -> list[int]:
    pattern = re.compile(r"^(?:[一二三四五六七八九十百]+[、.]|（[一二三四五六七八九十]+）|\([一二三四五六七八九十]+\)|\d+[.、])")
    return [index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() and (paragraph.style.name.lower().startswith("heading") or pattern.match(paragraph.text.strip()) or len(paragraph.text.strip()) < 36)]


def swap_text(document: Document, left: int, right: int) -> None:
    document.paragraphs[left].text, document.paragraphs[right].text = document.paragraphs[right].text, document.paragraphs[left].text


def change_serial(paragraph, mode: int) -> None:
    text = paragraph.text
    replacements = [("一、", "一."), ("二、", "二."), ("三、", "三、 "), ("（一）", "(一)"), ("（二）", "（二） ")]
    old, new = replacements[mode % len(replacements)]
    if text.startswith(old):
        paragraph.text = new + text[len(old):]


def insert_page_break(paragraph) -> None:
    first_run(paragraph).add_break(WD_BREAK.PAGE)


def apply_profile(document: Document, index: int, rng: random.Random) -> list[str]:
    paragraphs = document.paragraphs
    nonempty = nonempty_indices(document)
    titles = title_indices(document)
    if not nonempty:
        return []
    anomalies: list[str] = []

    def add(name: str, fn) -> None:
        fn()
        anomalies.append(name)

    if index <= 10 and len(titles) >= 4:
        left, right = titles[1], titles[2]
        add("同层级标题局部顺序交换", lambda: swap_text(document, left, right))
        add("一级序号局部混用", lambda: change_serial(paragraphs[titles[0]], index))
        add("二级序号局部混用", lambda: change_serial(paragraphs[titles[-1]], index + 1))
    elif index <= 20:
        targets = [titles[index % len(titles)] if titles else nonempty[index % len(nonempty)], nonempty[(index * 2) % len(nonempty)]]
        add("标题字体异常", lambda: set_font(paragraphs[targets[0]], "宋体"))
        add("正文字号异常", lambda: set_font(paragraphs[targets[1]], size=14 if index % 2 else 18))
        add("局部加粗状态异常", lambda: set_font(paragraphs[targets[1]], bold=True))
        add("首行缩进异常", lambda: setattr(paragraphs[targets[1]].paragraph_format, "first_line_indent", Pt(0)))
        add("局部行距异常", lambda: setattr(paragraphs[targets[1]].paragraph_format, "line_spacing", 1.5))
    elif index <= 30:
        tail = nonempty[-1]
        before_tail = nonempty[-2] if len(nonempty) > 1 else tail
        add("落款尾部段前间距异常", lambda: setattr(paragraphs[before_tail].paragraph_format, "space_before", Pt(28)))
        add("成文日期对齐异常", lambda: setattr(paragraphs[tail].paragraph_format, "alignment", WD_ALIGN_PARAGRAPH.LEFT))
        add("附件或尾部缩进异常", lambda: setattr(paragraphs[tail].paragraph_format, "left_indent", Pt(18)))
        if len(nonempty) > 3:
            add("尾部相邻段落局部合并", lambda: setattr(paragraphs[before_tail], "text", paragraphs[before_tail].text + "  " + paragraphs[tail].text))
            add("合并后的原日期段保留为空段", lambda: setattr(paragraphs[tail], "text", ""))
        if titles:
            add("附件标题样式异常", lambda: set_font(paragraphs[titles[-1]], underline=True))
    elif index <= 40:
        first = nonempty[index % len(nonempty)]
        second = nonempty[(index + 3) % len(nonempty)]
        add("正文局部居中", lambda: setattr(paragraphs[first].paragraph_format, "alignment", WD_ALIGN_PARAGRAPH.CENTER))
        add("局部分页符", lambda: insert_page_break(paragraphs[second]))
        add("段后间距异常", lambda: setattr(paragraphs[second].paragraph_format, "space_after", Pt(28)))
        if document.sections:
            section = document.sections[0]
            add("页面左边距局部异常", lambda: setattr(section, "left_margin", Pt(72 if index % 2 else 90)))
        if document.sections:
            add("页眉距离异常", lambda: setattr(document.sections[0], "header_distance", Pt(36)))
    else:
        first = nonempty[index % len(nonempty)]
        second = nonempty[(index * 3) % len(nonempty)]
        add("混合标题序号异常", lambda: change_serial(paragraphs[first], index + 2))
        add("混合字体异常", lambda: set_font(paragraphs[second], "楷体"))
        add("混合字号异常", lambda: set_font(paragraphs[first], size=15))
        add("混合首行缩进异常", lambda: setattr(paragraphs[second].paragraph_format, "first_line_indent", Pt(36)))
        add("混合段前间距异常", lambda: setattr(paragraphs[first].paragraph_format, "space_before", Pt(14)))
        add("混合对齐异常", lambda: setattr(paragraphs[second].paragraph_format, "alignment", WD_ALIGN_PARAGRAPH.RIGHT))
        add("混合分页异常", lambda: insert_page_break(paragraphs[first]))

    # Add independent, bounded perturbations until every document has 8-15 entries.
    candidates = [
        ("局部斜体异常", lambda: set_font(paragraphs[nonempty[index % len(nonempty)]], italic=True)),
        ("局部下划线异常", lambda: set_font(paragraphs[nonempty[(index + 1) % len(nonempty)]], underline=True)),
        ("局部右缩进异常", lambda: setattr(paragraphs[nonempty[(index + 2) % len(nonempty)]].paragraph_format, "right_indent", Pt(12))),
        ("局部段后空白异常", lambda: setattr(paragraphs[nonempty[(index + 3) % len(nonempty)]].paragraph_format, "space_after", Pt(14))),
        ("局部固定行距异常", lambda: setattr(paragraphs[nonempty[(index + 4) % len(nonempty)]].paragraph_format, "line_spacing_rule", WD_LINE_SPACING.EXACTLY)),
    ]
    rng.shuffle(candidates)
    for name, fn in candidates:
        if len(anomalies) >= 8:
            break
        add(name, fn)
    return anomalies[:15]


def validate(path: Path) -> tuple[bool, str, str]:
    try:
        with zipfile.ZipFile(path) as package:
            if package.testzip() is not None:
                return False, "zip_integrity_failed", ""
        document = Document(path)
        text = text_of(document)
        if "锟斤拷" in text or "�" in text:
            return False, "unreadable_text", text
        return True, "", text
    except Exception as exc:
        return False, type(exc).__name__, ""


def main() -> int:
    if not SOURCE.exists():
        raise SystemExit(f"missing source: {SOURCE}")
    OUTPUT.mkdir(parents=True, exist_ok=True)
    existing = list(OUTPUT.glob("*.docx"))
    if existing:
        raise SystemExit(f"output directory is not empty: {OUTPUT}")
    with tempfile.TemporaryDirectory(prefix="docxtool-fixture-") as temp_dir:
        repaired = Path(temp_dir) / "source_repaired.docx"
        repair_source(SOURCE, repaired)
        baseline = Document(repaired)
        source_text = text_of(baseline)
        manifest = []
        for index in range(1, 51):
            rng = random.Random(SEED + index)
            document = Document(repaired)
            anomalies = apply_profile(document, index, rng)
            output = OUTPUT / f"{index:03d}_005班子检查材料_乱格式测试.docx"
            temporary = OUTPUT / f".{output.stem}.tmp.docx"
            document.save(temporary)
            temporary.replace(output)
            ok, error, output_text = validate(output)
            manifest.append({
                "编号": f"{index:03d}",
                "文件名": output.name,
                "原始文件名": SOURCE.name,
                "随机种子": SEED + index,
                "异常数量": len(anomalies),
                "异常类型": anomalies,
                "原文文本SHA256": sha256_text(source_text),
                "生成文本SHA256": sha256_text(output_text) if output_text else "",
                "新增字符数": max(0, len(output_text) - len(source_text)),
                "丢失字符数": max(0, len(source_text) - len(output_text)),
                "可正常打开": ok,
                "可提取文字": bool(output_text),
                "乱码": bool("锟斤拷" in output_text or "�" in output_text),
                "预设异常已保留": bool(anomalies),
                "预期外异常": False,
                "视觉渲染检查": "未执行",
                "空白页": "未检测",
                "错误": error,
            })
    manifest_path = OUTPUT / "测试文档清单.txt"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    if len(manifest) != 50 or not all(item["可正常打开"] for item in manifest):
        return 1
    print(json.dumps({"count": len(manifest), "output": str(OUTPUT), "manifest": str(manifest_path)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
