"""Capture privacy-preserving Phase A import and export equivalence snapshots.

This utility deliberately observes the existing importer through its public
``DocxImporter.load`` entry point and its long-lived importer compatibility
aliases.  It is a migration test tool, not part of the formatting runtime.
Snapshots retain hashes, lengths, structure and locator facts only; source
paragraph text never reaches the JSON artifact.
"""

from __future__ import annotations

import argparse
from dataclasses import replace
import hashlib
import json
import posixpath
from pathlib import Path
import sys
from typing import Any, Dict, Iterable, List, Mapping, Optional
from urllib.parse import unquote
import zipfile

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from docxtool.document.engine import export_doc  # noqa: E402
from docxtool.document.importer import DocxImporter  # noqa: E402
from docxtool.document.style_config import PageSettings, StyleRule  # noqa: E402


STANDARD_ROOT = ROOT / "test_docx" / "tset1" / "test1"
SPECIAL_ROOT = ROOT / "test_docx" / "test2" / "test2"
MODES = ("strict", "structural", "normalize")
_CONTENT_TYPES_PART = "[Content_Types].xml"
_CONTENT_TYPES_CONTENT_TYPE = "application/vnd.openxmlformats-package.content-types+xml"
_CORE_PROPERTIES_PART = "docProps/core.xml"
_CORE_TIME_FIELDS = frozenset(("created", "modified", "lastPrinted"))
_NORMALIZED_TIME_TEXT = "__DOCXTOOL_NORMALIZED_CORE_TIME__"


def _sha256(value: bytes) -> str:
    """Return the stable SHA-256 digest used by redacted snapshots."""
    return hashlib.sha256(value).hexdigest()


def _text_summary(value: Any) -> Dict[str, Any]:
    """Represent text without exposing source content in a snapshot file."""
    text = str(value or "")
    return {"length": len(text), "sha256": _sha256(text.encode("utf-8"))}


def _xml_digest(value: Any) -> str:
    """Hash an OOXML element without persisting its potentially private XML."""
    element = getattr(value, "_element", value)
    xml = getattr(element, "xml", "")
    return _sha256(str(xml).encode("utf-8"))


def _json_digest(value: Any) -> str:
    """Hash arbitrary diagnostic data after deterministic JSON serialization."""
    encoded = json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
    return _sha256(encoded.encode("utf-8"))


def _parse_xml(value: bytes) -> Any:
    """Parse package XML without network or entity resolution."""
    parser = etree.XMLParser(resolve_entities=False, no_network=True, recover=False)
    return etree.fromstring(value, parser=parser)


def _content_type_maps(value: Optional[bytes]) -> tuple[Dict[str, str], Dict[str, str], List[str]]:
    """Return OPC default/override content types and stable parse errors."""
    defaults: Dict[str, str] = {}
    overrides: Dict[str, str] = {}
    errors: List[str] = []
    if value is None:
        return defaults, overrides, ["MISSING_CONTENT_TYPES"]
    try:
        root = _parse_xml(value)
    except (etree.XMLSyntaxError, ValueError):
        return defaults, overrides, ["MALFORMED_CONTENT_TYPES"]
    for child in root:
        local_name = etree.QName(child).localname
        if local_name == "Default":
            extension = str(child.get("Extension", "") or "").lower()
            content_type = str(child.get("ContentType", "") or "")
            if extension:
                defaults[extension] = content_type
        elif local_name == "Override":
            part_name = str(child.get("PartName", "") or "").lstrip("/")
            content_type = str(child.get("ContentType", "") or "")
            if part_name:
                overrides[part_name] = content_type
    return defaults, overrides, errors


def _part_content_type(
    part_name: str,
    defaults: Mapping[str, str],
    overrides: Mapping[str, str],
) -> str:
    """Resolve one OPC part content type without guessing unknown extensions."""
    if part_name == _CONTENT_TYPES_PART:
        return _CONTENT_TYPES_CONTENT_TYPE
    if part_name in overrides:
        return str(overrides[part_name])
    extension = posixpath.splitext(part_name)[1].lstrip(".").lower()
    return str(defaults.get(extension, "") or "")


def _normalize_part(part_name: str, value: bytes) -> tuple[bytes, List[str], Optional[str]]:
    """Normalize only documented non-deterministic core-property time text."""
    if part_name != _CORE_PROPERTIES_PART:
        return value, [], None
    try:
        root = _parse_xml(value)
    except (etree.XMLSyntaxError, ValueError):
        return value, [], "MALFORMED_CORE_PROPERTIES"
    normalized_fields: List[str] = []
    for element in root.iter():
        local_name = etree.QName(element).localname
        if local_name in _CORE_TIME_FIELDS:
            normalized_fields.append(local_name)
            element.text = _NORMALIZED_TIME_TEXT
    if not normalized_fields:
        return value, [], None
    canonical = etree.tostring(root, method="c14n", with_comments=True)
    return canonical, sorted(set(normalized_fields)), None


def _relationship_source_part(relationships_part: str) -> str:
    """Map an OPC .rels member to its source part, using '/' for package root."""
    if relationships_part == "_rels/.rels":
        return "/"
    marker = "/_rels/"
    if marker not in relationships_part or not relationships_part.endswith(".rels"):
        return ""
    prefix, filename = relationships_part.rsplit(marker, 1)
    return posixpath.join(prefix, filename[:-5])


def _resolve_relationship_target(source_part: str, target: str) -> str:
    """Resolve one internal relationship target to a normalized package part name."""
    target_path = unquote(str(target or "").split("#", 1)[0].split("?", 1)[0])
    target_path = target_path.replace("\\", "/")
    if target_path.startswith("/"):
        return posixpath.normpath(target_path).lstrip("/")
    base = "" if source_part == "/" else posixpath.dirname(source_part)
    return posixpath.normpath(posixpath.join(base, target_path)).lstrip("/")


def _package_manifest(path: Path) -> Dict[str, Any]:
    """Build a deterministic manifest for every package part and relationship."""
    with zipfile.ZipFile(path) as package:
        raw_entries = []
        for info in package.infolist():
            if info.is_dir():
                continue
            part_name = info.filename.replace("\\", "/").lstrip("/")
            raw_entries.append((part_name, package.read(info)))
        raw_entries.sort(key=lambda item: (item[0], _sha256(item[1])))
        package_names = {name for name, _value in raw_entries}
        content_types_value = next(
            (value for name, value in raw_entries if name == _CONTENT_TYPES_PART),
            None,
        )
        defaults, overrides, manifest_errors = _content_type_maps(content_types_value)
        duplicate_counts: Dict[str, int] = {}
        for name, _value in raw_entries:
            duplicate_counts[name] = duplicate_counts.get(name, 0) + 1
        duplicate_ordinals: Dict[str, int] = {}
        parts: Dict[str, Dict[str, Any]] = {}
        normalized_metadata: List[Dict[str, Any]] = []
        for part_name, raw_value in raw_entries:
            duplicate_ordinals[part_name] = duplicate_ordinals.get(part_name, 0) + 1
            ordinal = duplicate_ordinals[part_name]
            record_key = (
                part_name
                if duplicate_counts[part_name] == 1
                else "{0}#{1}".format(part_name, ordinal)
            )
            normalized_value, normalized_fields, normalization_error = _normalize_part(
                part_name,
                raw_value,
            )
            if normalization_error:
                manifest_errors.append(normalization_error)
            if normalized_fields:
                normalized_metadata.append({
                    "part_name": part_name,
                    "fields": normalized_fields,
                    "strategy": "replace-core-property-time-text",
                })
            parts[record_key] = {
                "part_name": part_name,
                "size": len(normalized_value),
                "sha256": _sha256(normalized_value),
                "content_type": _part_content_type(part_name, defaults, overrides),
            }
        duplicate_parts = sorted(name for name, count in duplicate_counts.items() if count > 1)
        if duplicate_parts:
            manifest_errors.append("DUPLICATE_PACKAGE_PART")

        relationships: Dict[str, Dict[str, Any]] = {}
        relationship_parse_errors: List[Dict[str, str]] = []
        missing_targets: List[Dict[str, str]] = []
        for relationships_part, raw_value in raw_entries:
            if not relationships_part.endswith(".rels"):
                continue
            source_part = _relationship_source_part(relationships_part)
            try:
                root = _parse_xml(raw_value)
            except (etree.XMLSyntaxError, ValueError):
                relationship_parse_errors.append({
                    "part_name": relationships_part,
                    "error_code": "MALFORMED_RELATIONSHIPS_XML",
                })
                continue
            duplicate_relationship_ids: Dict[str, int] = {}
            for element in root:
                if etree.QName(element).localname != "Relationship":
                    continue
                relationship_id = str(element.get("Id", "") or "")
                relationship_type = str(element.get("Type", "") or "")
                target = str(element.get("Target", "") or "")
                target_mode = str(element.get("TargetMode", "") or "Internal")
                duplicate_relationship_ids[relationship_id] = (
                    duplicate_relationship_ids.get(relationship_id, 0) + 1
                )
                ordinal = duplicate_relationship_ids[relationship_id]
                key = "{0}|{1}".format(source_part, relationship_id)
                if ordinal > 1:
                    key = "{0}#{1}".format(key, ordinal)
                external = target_mode.lower() == "external"
                resolved_target = "" if external else _resolve_relationship_target(source_part, target)
                target_exists: Optional[bool] = None if external else resolved_target in package_names
                relationships[key] = {
                    "source_part": source_part,
                    "relationship_id": relationship_id,
                    "relationship_type": relationship_type,
                    "target": target,
                    "target_mode": target_mode,
                    "target_exists": target_exists,
                    "resolved_target": resolved_target,
                }
                if target_exists is False:
                    missing_targets.append({
                        "source_part": source_part,
                        "relationship_id": relationship_id,
                        "target": target,
                        "resolved_target": resolved_target,
                    })
        if relationship_parse_errors:
            manifest_errors.append("RELATIONSHIP_PARSE_ERROR")
        bad_member = package.testzip()

    normalized_metadata.sort(key=lambda item: item["part_name"])
    missing_targets.sort(
        key=lambda item: (item["source_part"], item["relationship_id"], item["target"])
    )
    relationship_parse_errors.sort(key=lambda item: item["part_name"])
    errors = sorted(set(manifest_errors))
    return {
        "schema": "docx-package-manifest-v1",
        "part_count": len(parts),
        "parts": parts,
        "relationship_count": len(relationships),
        "relationships": relationships,
        "normalized_metadata": normalized_metadata,
        "missing_relationship_target_count": len(missing_targets),
        "missing_relationship_targets": missing_targets,
        "relationship_parse_errors": relationship_parse_errors,
        "duplicate_parts": duplicate_parts,
        "zip_integrity_member": bad_member or "",
        "errors": errors,
        "valid": not errors and not missing_targets and not bad_member,
    }


def _source_run_snapshot(source_run: Any) -> Dict[str, Any]:
    """Capture non-text source-run facts used by segment-format verification."""
    keys = (
        "start", "end", "font_name", "east_asia_font_name", "ascii_font_name",
        "font_size_pt", "bold", "italic", "underline", "explicit", "inherited",
        "known", "format_sources",
    )
    return {key: getattr(source_run, key, None) for key in keys}


def _feature_snapshot(features: Any) -> Dict[str, Any]:
    """Capture physical/segment facts without retaining visible paragraph text."""
    if features is None:
        return {"present": False}
    scalar_keys = (
        "style_name", "font_name", "font_size_pt", "bold", "alignment",
        "first_line_indent", "numbering_prefix", "paragraph_index",
        "source_physical_paragraph_index", "source_start_utf16", "source_end_utf16",
        "source_canonical_start_utf16", "source_canonical_end_utf16",
        "source_locator_status", "segment_font_name", "segment_dominant_font_name",
        "segment_font_name_east_asia", "segment_font_name_ascii", "segment_font_size_pt",
        "segment_weighted_font_size_pt", "segment_bold_char_ratio",
        "segment_italic_char_ratio", "segment_underline_char_ratio",
        "segment_explicit_format_ratio", "segment_inherited_format_ratio",
        "segment_run_count", "segment_visible_char_count", "segment_mapped_format_char_count",
        "segment_format_coverage_ratio", "segment_format_status", "segment_style_name",
        "segment_has_mixed_fonts", "segment_has_mixed_sizes", "segment_numbering_features",
        "segment_alignment", "segment_position_in_physical_paragraph", "segment_index",
        "segment_count", "is_in_table", "contains_image", "is_new_line",
        "first_run_font_name", "first_run_font_size_pt", "first_run_bold",
        "dominant_font_name", "weighted_font_size", "max_font_size", "min_font_size",
        "bold_char_ratio", "italic_char_ratio", "explicitly_formatted_char_ratio",
        "inline_lead_bold",
    )
    result = {"present": True}
    result.update({key: getattr(features, key, None) for key in scalar_keys})
    result["source_physical_text"] = _text_summary(
        getattr(features, "source_physical_text", "")
    )
    result["source_canonical_text"] = _text_summary(
        getattr(features, "source_canonical_text", "")
    )
    result["source_fragment_text"] = _text_summary(
        getattr(features, "source_fragment_text", "")
    )
    result["source_canonical_fragment_text"] = _text_summary(
        getattr(features, "source_canonical_fragment_text", "")
    )
    result["source_locator_evidence"] = list(
        getattr(features, "source_locator_evidence", ()) or ()
    )
    result["source_locator_warnings"] = list(
        getattr(features, "source_locator_warnings", ()) or ()
    )
    result["segment_format_warnings"] = list(
        getattr(features, "segment_format_warnings", ()) or ()
    )
    result["segment_format_sources"] = list(
        getattr(features, "segment_format_sources", ()) or ()
    )
    result["source_run_spans"] = [
        _source_run_snapshot(item)
        for item in (getattr(features, "source_run_spans", ()) or ())
    ]
    return result


def _token_snapshot(tokens: Iterable[Any]) -> List[Dict[str, Any]]:
    """Capture inline-token ordering and text identity without source text."""
    return [
        {"kind": getattr(token, "kind", ""), "text": _text_summary(getattr(token, "text", ""))}
        for token in tokens or ()
    ]


def _legacy_type(meta: Mapping[str, Any]) -> str:
    """Read both legacy flat and structured metadata forms unchanged."""
    value = meta.get("legacy_type_id", "")
    if isinstance(value, Mapping):
        return str(value.get("value", "") or "")
    return str(value or "")


def _paragraph_snapshot(paragraph: Any, index: int) -> Dict[str, Any]:
    """Capture the type, text identity, locator and review contract of one segment."""
    meta = dict(getattr(paragraph, "meta", None) or {})
    return {
        "index": index,
        "type_id": str(getattr(paragraph, "type_id", "") or ""),
        "text": _text_summary(getattr(paragraph, "text", "")),
        "original_text": _text_summary(getattr(paragraph, "original_text", "")),
        "features": _feature_snapshot(getattr(paragraph, "features", None)),
        "inline_tokens": _token_snapshot(getattr(paragraph, "inline_tokens", ()) or ()),
        "legacy_type_id": _legacy_type(meta),
        "classification_kind": meta.get("classification_kind"),
        "classification_confidence": meta.get("classification_confidence"),
        "classification_auto_format": meta.get("classification_auto_format"),
        "recognition_type": meta.get("recognition_type"),
        "recognized_type": meta.get("recognized_type"),
        "final_type": meta.get("final_type"),
        "recognition_block_index": meta.get("recognition_block_index"),
        "review_level": meta.get("review_level"),
        "review_confidence": meta.get("review_confidence"),
        "review_reasons": list(meta.get("review_reasons", ()) or ()),
        "heading_level": meta.get("heading_level"),
        "numbering": meta.get("numbering"),
        "attachment_single": meta.get("attachment_single"),
        "attachment_multi": meta.get("attachment_multi"),
        "source_locator_status": getattr(getattr(paragraph, "features", None), "source_locator_status", ""),
    }


def _diagnostics_snapshot(data: Any) -> Dict[str, Any]:
    """Retain review diagnostics while avoiding previews and raw candidate text."""
    report = getattr(data, "recognition_diagnostics", {}) or {}
    source_items = report.get("paragraphs", ()) if isinstance(report, Mapping) else ()
    allowed = (
        "paragraph_index", "block_index", "legacy_type", "legacy_type_id",
        "recognized_type", "final_type", "review_confidence", "review_level",
        "review_reasons", "needs_review", "mapping_applied", "mapping_failed",
        "final_score", "candidate_margin", "candidate_count", "winner_provider",
    )
    paragraphs = []
    for item in source_items if isinstance(source_items, list) else ():
        if not isinstance(item, Mapping):
            continue
        snapshot = {key: item.get(key) for key in allowed}
        snapshot["candidate_digest"] = _json_digest(item.get("candidates", ()))
        snapshot["evidence_digest"] = _json_digest(item.get("evidence_summary", ()))
        paragraphs.append(snapshot)
    return {
        "summary": dict(report.get("summary", {}) or {}) if isinstance(report, Mapping) else {},
        "validation": dict(report.get("validation", {}) or {}) if isinstance(report, Mapping) else {},
        "paragraphs": paragraphs,
    }


def _raw_block_snapshot(raw_blocks: Iterable[tuple]) -> List[Dict[str, Any]]:
    """Capture physical reader output before logical segmentation occurs."""
    result: List[Dict[str, Any]] = []
    for index, block in enumerate(raw_blocks):
        kind = block[0]
        item: Dict[str, Any] = {"index": index, "kind": kind}
        if kind == "paragraph":
            _kind, paragraph, features, tokens, sect_pr = block
            item.update({
                "text": _text_summary(getattr(paragraph, "text", "")),
                "features": _feature_snapshot(features),
                "tokens": _token_snapshot(tokens),
                "section_xml_sha256": _xml_digest(sect_pr) if sect_pr is not None else "",
            })
        elif kind in {"paragraph_xml", "letterhead_paragraph_xml"}:
            item.update({
                "text": _text_summary(getattr(block[1], "text", "")),
                "xml_sha256": _xml_digest(block[1]),
            })
        elif kind == "protected_paragraph_xml":
            item.update({
                "text": _text_summary(getattr(block[1], "text", "")),
                "features": _feature_snapshot(block[2]),
                "xml_sha256": _xml_digest(block[1]),
            })
        elif kind == "table":
            item["xml_sha256"] = _xml_digest(block[1])
        else:
            item["value_sha256"] = _json_digest(block[1:])
        result.append(item)
    return result


def _logical_line_snapshot(flat_lines: Iterable[tuple]) -> List[Dict[str, Any]]:
    """Capture logical segmentation output before importer classification begins."""
    result: List[Dict[str, Any]] = []
    for index, line in enumerate(flat_lines):
        kind = line[0]
        item: Dict[str, Any] = {"index": index, "kind": kind}
        if kind == "text":
            _kind, text, features, tokens, sect_pr = line
            item.update({
                "text": _text_summary(text),
                "features": _feature_snapshot(features),
                "tokens": _token_snapshot(tokens),
                "section_xml_sha256": _xml_digest(sect_pr) if sect_pr is not None else "",
            })
        elif kind == "table":
            item["xml_sha256"] = _xml_digest(line[1])
        elif kind in {"paragraph_xml", "letterhead_paragraph_xml"}:
            item.update({
                "text": _text_summary(getattr(line[1], "text", "")),
                "xml_sha256": _xml_digest(line[1]),
            })
        elif kind == "protected_paragraph_xml":
            item.update({
                "text": _text_summary(getattr(line[1], "text", "")),
                "features": _feature_snapshot(line[2]),
                "xml_sha256": _xml_digest(line[1]),
            })
        else:
            item["value_sha256"] = _json_digest(line[1:])
        result.append(item)
    return result


def _document_snapshot(data: Any) -> Dict[str, Any]:
    """Capture importer output after recognition and existing normalization order."""
    paragraphs = list(getattr(data, "paragraphs", ()) or ())
    return {
        "paragraph_count": len(paragraphs),
        "paragraphs": [_paragraph_snapshot(item, index) for index, item in enumerate(paragraphs)],
        "table_count": len(getattr(data, "tables", ()) or ()),
        "image_count": sum(getattr(item, "type_id", "") == "__image__" for item in paragraphs),
        "section_relationship_count": len(getattr(data, "section_relationship_parts", {}) or {}),
        "has_body_section": bool(getattr(data, "body_sectPr", None) is not None),
        "doc_mode": getattr(data, "doc_mode", ""),
        "processing_strategy": getattr(data, "processing_strategy", ""),
        "recognition_mode": getattr(data, "recognition_mode", ""),
        "strict_preservation": bool(getattr(data, "strict_preservation", False)),
        "normalization_change_count": len(getattr(data, "normalization_changes", ()) or ()),
        "diagnostics": _diagnostics_snapshot(data),
    }


def _output_docx_snapshot(path: Path) -> Dict[str, Any]:
    """Capture the complete OPC package manifest and visible document structure."""
    package_manifest = _package_manifest(path)
    document = Document(path)
    return {
        "package_manifest": package_manifest,
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "paragraphs": [_text_summary(item.text) for item in document.paragraphs],
    }


def _load_with_capture(
    source: Path,
    mode: str,
    rules: list,
    *,
    legacy_candidates_enabled: Optional[bool] = None,
) -> tuple[Dict[str, Any], Any]:
    """Load one DOCX while observing existing importer compatibility aliases."""
    import docxtool.document.importer as importer_module

    captured: Dict[str, Any] = {}
    real_read_blocks = importer_module._read_body_blocks
    real_build_lines = importer_module._build_logical_lines
    real_apply_recognition = importer_module.apply_recognition

    def capture_blocks(*args: Any, **kwargs: Any) -> Any:
        result = real_read_blocks(*args, **kwargs)
        captured["physical_blocks"] = _raw_block_snapshot(result)
        return result

    def capture_lines(*args: Any, **kwargs: Any) -> Any:
        result = real_build_lines(*args, **kwargs)
        captured["logical_lines"] = _logical_line_snapshot(result)
        return result

    def capture_recognition(data: Any, config: Any = None) -> None:
        captured["pre_recognition"] = _document_snapshot(data)
        effective_config = config
        if legacy_candidates_enabled is not None and config is not None:
            effective_config = replace(config, enable_legacy_candidates=legacy_candidates_enabled)
        real_apply_recognition(data, effective_config)

    importer_module._read_body_blocks = capture_blocks
    importer_module._build_logical_lines = capture_lines
    importer_module.apply_recognition = capture_recognition
    try:
        data = DocxImporter().load(
            str(source),
            rules,
            features={
                "processing": {"strategy": mode},
                "numbering": {"enabled": True, "mode": "safe"},
            },
        )
    finally:
        importer_module._read_body_blocks = real_read_blocks
        importer_module._build_logical_lines = real_build_lines
        importer_module.apply_recognition = real_apply_recognition

    captured["result"] = _document_snapshot(data)
    return captured, data


def _document_sources() -> List[tuple[str, Path]]:
    """Return the required 50 standard and 5 special regression sources."""
    standard = sorted(STANDARD_ROOT.glob("*.docx"))
    special = sorted(SPECIAL_ROOT.glob("*.docx"))
    if len(standard) != 50:
        raise ValueError("expected 50 standard DOCX files, found {0}".format(len(standard)))
    if len(special) != 5:
        raise ValueError("expected 5 special DOCX files, found {0}".format(len(special)))
    return [("standard", item) for item in standard] + [("special", item) for item in special]


def capture_snapshot(output: Path, artifact_dir: Optional[Path]) -> Dict[str, Any]:
    """Capture strict, structural and normalize snapshots for all 55 fixtures."""
    rules = StyleRule.from_config()
    settings = PageSettings.from_config()
    documents: List[Dict[str, Any]] = []
    if artifact_dir is not None:
        artifact_dir.mkdir(parents=True, exist_ok=True)
    for group, source in _document_sources():
        document_record: Dict[str, Any] = {
            "group": group,
            "source_name": source.name,
            "source_sha256": _sha256(source.read_bytes()),
            "modes": {},
        }
        for mode in MODES:
            captured, data = _load_with_capture(source, mode, rules)
            if artifact_dir is not None:
                artifact = artifact_dir / "{0}-{1}-{2}.docx".format(
                    group, source.stem, mode
                )
                export_doc(data, rules, settings, str(artifact))
                captured["export"] = _output_docx_snapshot(artifact)
            document_record["modes"][mode] = captured
        documents.append(document_record)
    payload = {
        "schema": "phase-a-equivalence-v2",
        "text_storage": "hash-and-length-only",
        "document_count": len(documents),
        "mode_count": len(MODES),
        "documents": documents,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")
    return payload


def capture_legacy_provider_comparison(output: Path) -> Dict[str, Any]:
    """Compare provider-toggle input invariance separately from output drift."""
    rules = StyleRule.from_config()
    input_differences: List[Dict[str, Any]] = []
    output_differences: List[Dict[str, Any]] = []
    checked = 0
    for group, source in _document_sources():
        for mode in MODES:
            enabled, _enabled_data = _load_with_capture(
                source, mode, rules, legacy_candidates_enabled=True
            )
            disabled, _disabled_data = _load_with_capture(
                source, mode, rules, legacy_candidates_enabled=False
            )
            checked += 1
            input_paths = _difference_paths(
                enabled.get("pre_recognition", {}),
                disabled.get("pre_recognition", {}),
            )
            if input_paths:
                input_differences.append({
                    "group": group,
                    "source_name": source.name,
                    "mode": mode,
                    "different_field_count": len(input_paths),
                    "fields": input_paths,
                })
            output_paths = _difference_paths(
                enabled.get("result", {}),
                disabled.get("result", {}),
            )
            if output_paths:
                output_differences.append({
                    "group": group,
                    "source_name": source.name,
                    "mode": mode,
                    "different_field_count": len(output_paths),
                    "fields": output_paths,
                })
    payload = {
        "schema": "phase-a-legacy-provider-comparison-v2",
        "comparison": "legacy-candidate-provider-input-invariance",
        "scope_warning": (
            "This experiment toggles only LegacyCandidateProvider and does not disable "
            "importer Legacy preprocessing."
        ),
        "legacy_candidate_provider": {
            "status": "toggled",
            "enabled_run": True,
            "disabled_run": False,
        },
        "importer_legacy_preprocessing": {
            "status": "blocked",
            "enabled_in_both_runs": True,
            "tested": False,
            "reason": "No production-neutral importer preprocessing bypass is available.",
        },
        "checked_cases": checked,
        "input_comparison": {
            "difference_count": len(input_differences),
            "differences": input_differences,
        },
        "output_comparison": {
            "difference_count": len(output_differences),
            "differences": output_differences,
        },
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")
    return payload


def capture_legacy_input_comparison(output: Path) -> Dict[str, Any]:
    """Compatibility wrapper for the renamed provider-only experiment."""
    return capture_legacy_provider_comparison(output)


def _difference_paths(left: Any, right: Any, path: str = "$") -> List[str]:
    """Return deterministic field paths for two JSON-compatible snapshots."""
    if type(left) is not type(right):
        return [path]
    if isinstance(left, Mapping):
        paths: List[str] = []
        for key in sorted(set(left) | set(right), key=str):
            key_path = "{0}.{1}".format(path, key)
            if key not in left or key not in right:
                paths.append(key_path)
            else:
                paths.extend(_difference_paths(left[key], right[key], key_path))
        return paths
    if isinstance(left, list):
        paths = []
        if len(left) != len(right):
            paths.append("{0}.length".format(path))
        for index, (left_item, right_item) in enumerate(zip(left, right)):
            paths.extend(_difference_paths(left_item, right_item, "{0}[{1}]".format(path, index)))
        return paths
    return [] if left == right else [path]


def _package_validation_paths(value: Any, path: str = "$") -> List[str]:
    """Return paths for invalid package manifests even when both sides match."""
    issues: List[str] = []
    if isinstance(value, Mapping):
        manifest = value.get("package_manifest")
        if isinstance(manifest, Mapping):
            for key, relationship in (manifest.get("relationships", {}) or {}).items():
                if isinstance(relationship, Mapping) and relationship.get("target_exists") is False:
                    issues.append(
                        "{0}.package_manifest.relationships.{1}.target_exists".format(path, key)
                    )
            for index, _item in enumerate(manifest.get("relationship_parse_errors", ()) or ()):
                issues.append(
                    "{0}.package_manifest.relationship_parse_errors[{1}]".format(path, index)
                )
            if manifest.get("zip_integrity_member"):
                issues.append("{0}.package_manifest.zip_integrity_member".format(path))
        for key, item in value.items():
            issues.extend(_package_validation_paths(item, "{0}.{1}".format(path, key)))
    elif isinstance(value, list):
        for index, item in enumerate(value):
            issues.extend(_package_validation_paths(item, "{0}[{1}]".format(path, index)))
    return issues


def _categorized_difference_paths(paths: Iterable[str]) -> Dict[str, List[str]]:
    """Separate package, relationship, normalized metadata and structure drift."""
    categories = {
        "package_part_differences": [],
        "relationship_differences": [],
        "normalized_metadata_differences": [],
        "document_structure_differences": [],
    }
    for path in paths:
        if ".package_manifest.normalized_metadata" in path:
            categories["normalized_metadata_differences"].append(path)
        elif (
            ".package_manifest.relationships" in path
            or ".package_manifest.relationship_parse_errors" in path
            or ".package_manifest.missing_relationship" in path
        ):
            categories["relationship_differences"].append(path)
        elif ".package_manifest" in path:
            categories["package_part_differences"].append(path)
        else:
            categories["document_structure_differences"].append(path)
    return categories


def compare_snapshots(before: Path, after: Path, output: Path) -> Dict[str, Any]:
    """Compare two captured payloads and retain only redacted structural paths."""
    before_payload = json.loads(before.read_text(encoding="utf-8"))
    after_payload = json.loads(after.read_text(encoding="utf-8"))
    paths = sorted(set(
        _difference_paths(before_payload, after_payload)
        + _package_validation_paths(before_payload)
        + _package_validation_paths(after_payload)
    ))
    categories = _categorized_difference_paths(paths)
    payload = {
        "schema": "phase-a-equivalence-comparison-v2",
        "before_sha256": _sha256(before.read_bytes()),
        "after_sha256": _sha256(after.read_bytes()),
        "difference_count": len(paths),
        "difference_paths": paths,
        **categories,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")
    return payload


def _parse_args(argv: Optional[List[str]] = None) -> argparse.Namespace:
    """Parse the small CLI used by Phase A migration verification."""
    parser = argparse.ArgumentParser(description=__doc__)
    subparsers = parser.add_subparsers(dest="command", required=True)
    capture = subparsers.add_parser("capture")
    capture.add_argument("--output", type=Path, required=True)
    capture.add_argument("--artifact-dir", type=Path)
    legacy_provider = subparsers.add_parser("legacy-provider-input-invariance")
    legacy_provider.add_argument("--output", type=Path, required=True)
    legacy_alias = subparsers.add_parser("legacy-input")
    legacy_alias.add_argument("--output", type=Path, required=True)
    compare = subparsers.add_parser("compare")
    compare.add_argument("--before", type=Path, required=True)
    compare.add_argument("--after", type=Path, required=True)
    compare.add_argument("--output", type=Path, required=True)
    return parser.parse_args(argv)


def main(argv: Optional[List[str]] = None) -> int:
    """Run one snapshot command and return a standard process status."""
    args = _parse_args(argv)
    if args.command == "capture":
        capture_snapshot(args.output, args.artifact_dir)
    elif args.command in {"legacy-provider-input-invariance", "legacy-input"}:
        if args.command == "legacy-input":
            print(
                "warning: 'legacy-input' is deprecated; use "
                "'legacy-provider-input-invariance'.",
                file=sys.stderr,
            )
        comparison = capture_legacy_provider_comparison(args.output)
        if comparison["input_comparison"]["difference_count"]:
            return 1
    else:
        comparison = compare_snapshots(args.before, args.after, args.output)
        if comparison["difference_count"]:
            return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
