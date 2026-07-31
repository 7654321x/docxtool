"""Generate sanitized DOCX files for manual WPS integration validation.

The fixtures deliberately contain no user material.  They are not committed
and are written under ``test_docx`` so they can be opened by WPS directly.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _east_asia(run, value: str) -> None:
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), value)


def _save(document: Document, directory: Path, name: str) -> None:
    document.save(directory / name)


def generate(directory: Path) -> tuple[Path, ...]:
    directory.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_paragraph("DocxTool WPS 验收：基础标题与正文")
    document.add_paragraph("一、总体要求")
    document.add_paragraph("这是可直接判断为正文的脱敏示例内容，用于验证批注和格式应用。")
    _save(document, directory, "01-basic-heading-body.docx")

    document = Document()
    paragraph = document.add_paragraph()
    heading = paragraph.add_run("一、同段标题。")
    heading.font.bold = True
    heading.font.size = Pt(16)
    body = paragraph.add_run("同一物理段落中的正文必须只允许预览，不得自动整段排版。")
    body.font.size = Pt(12)
    _save(document, directory, "02-inline-mixed-segments.docx")

    document = Document()
    for value in ("唯一前缀", "重复测试段落", "中间锚点", "重复测试段落", "唯一后缀"):
        document.add_paragraph(value)
    _save(document, directory, "03-duplicate-paragraphs.docx")

    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("一、软换行标题")
    body = paragraph.add_run()
    body.add_break(WD_BREAK.LINE)
    body.add_text("正文含有全角空格　、NBSP\u00a0、Tab\t和 Emoji 😀。")
    _save(document, directory, "04-soft-break-and-spaces.docx")

    document = Document()
    base = document.styles.add_style("WpsValidationBase", WD_STYLE_TYPE.PARAGRAPH)
    base.font.size = Pt(14)
    base.font.bold = True
    derived = document.styles.add_style("WpsValidationDerived", WD_STYLE_TYPE.PARAGRAPH)
    derived.base_style = base
    paragraph = document.add_paragraph(style=derived)
    run = paragraph.add_run("东亚字体与样式继承示例")
    _east_asia(run, "SimHei")
    _save(document, directory, "05-eastasia-style-inheritance.docx")

    document = Document()
    paragraph = document.add_paragraph("用户原有批注锚点：此文字不得被插件删除。")
    if hasattr(document, "add_comment"):
        document.add_comment(paragraph.runs[0], text="这是用户原有批注，应始终保留。", author="User", initials="U")
    document.add_paragraph("一、插件预览与清理只能操作自己的批注。")
    _save(document, directory, "06-user-comment-protection.docx")

    document = Document()
    document.add_paragraph("确认段落：独立正文可进入一键排版。")
    document.add_paragraph("重复定位段落")
    document.add_paragraph("重复定位段落")
    document.add_paragraph("复核段落：重复文字发生局部歧义时只能预览或跳过。")
    _save(document, directory, "07-review-and-unresolved.docx")

    document = Document()
    document.add_paragraph("分节与页面格式验收")
    document.add_paragraph("纵向节中的正文。")
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    section.orientation = 1
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    document.add_paragraph("横向节中的正文。")
    _save(document, directory, "08-section-page-format.docx")

    return tuple(directory / name for name in (
        "01-basic-heading-body.docx", "02-inline-mixed-segments.docx",
        "03-duplicate-paragraphs.docx", "04-soft-break-and-spaces.docx",
        "05-eastasia-style-inheritance.docx", "06-user-comment-protection.docx",
        "07-review-and-unresolved.docx", "08-section-page-format.docx",
    ))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, default=Path("test_docx") / "wps_validation")
    args = parser.parse_args()
    for path in generate(args.output_dir.resolve()):
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
