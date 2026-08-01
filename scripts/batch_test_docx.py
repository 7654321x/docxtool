"""Run DOCX regression fixtures with structural comparison and optional visual QA."""

from __future__ import annotations

import argparse
from collections import Counter
from difflib import SequenceMatcher
import hashlib
import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import sys
import tempfile
import time
from typing import Any, Iterable
import zipfile

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from docxtool.document.engine import export_doc  # noqa: E402
from docxtool.document.importer import DocxImporter  # noqa: E402
from docxtool.document.letterhead_config import default_letterhead_config  # noqa: E402
from docxtool.document.style_config import PageSettings, StyleRule  # noqa: E402

INPUT_DIR = ROOT / "test_docx" / "tset1" / "test1"
TEMPLATE_DIR = ROOT / "test_docx" / "tset1" / "test1正确格式"
OUTPUT_DIR = ROOT / "test_docx" / "tset1" / "test1测试结果"
SPECIAL_INPUT_DIR = ROOT / "test_docx" / "test2" / "test2"
SPECIAL_TEMPLATE_DIR = ROOT / "test_docx" / "test2" / "test2正确格式"
SPECIAL_OUTPUT_DIR = ROOT / "test_docx" / "test2" / "测试结果"
COMPARISON_VERSION = "structural-alignment-v3"
_DISPATCH_NUMBER_RE = re.compile(r"^(?P<agency_code>.+?)〔(?P<year>\d{4})〕(?P<sequence>\d+)号$")
_PAGE_NUMBER_RE = re.compile(r"^[—–－\-\s]*\d{1,4}[—–－\-\s]*$")
_DATE_RE = re.compile(
    r"^(?P<year>[0-9〇零一二三四五六七八九]{4})年"
    r"(?P<month>[0-9一二三四五六七八九十]{1,3})月"
    r"(?P<day>[0-9一二三四五六七八九十]{1,3})日$"
)
_SIGNATURE_ORG_SUFFIX_RE = re.compile(
    r"(?:委员会|工作委员会|人民政府|人民法院|人民检察院|代表大会|"
    r"办公室|街道办事处|领导小组|工作组|党组|党委|政府|政协|人大|"
    r"总工会|局|厅|部|院|处|科|办|镇|乡)$"
)
_CJK_DIGITS = {"〇": 0, "零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
_REGIONS = ("letterhead", "front", "body", "tail", "attachment")
_RECOGNITION_FIELDS = ("type", "recognition_type", "section", "heading_level", "region")
_PHYSICAL_FIELDS = (
    "style", "font", "size_pt", "bold", "italic", "underline", "alignment",
    "first_indent", "left_indent", "right_indent", "line_spacing", "space_before",
    "space_after", "page_break_before", "keep_with_next", "keep_together",
)
_HEADING_PREFIX_RE = re.compile(
    r"^(?:"
    r"[一二三四五六七八九十百千零〇]+[、.．。]+|"
    r"[（(][一二三四五六七八九十百千零〇0-9]+[）)]|"
    r"[0-9]+[.．、]+|"
    r"[（(][0-9]+[）)]"
    r")\s*"
)
_OBJECT_CAPTION_RE = re.compile(r"^[表图]\s*[一二三四五六七八九十百千零〇0-9]+(?:\s|[：:、.．]|$)")
_INDENT_FIELDS = frozenset(("first_indent", "left_indent", "right_indent"))
_PROTECTED_FORMAT_FIELDS = frozenset((
    "style", "font", "size_pt", "bold", "italic", "underline", "alignment",
))
_INDENT_TOLERANCE_CHARS = 0.15


class RenderUnavailable(RuntimeError):
    """Raised when optional local visual QA dependencies are unavailable."""


def text_hash(text: str) -> str:
    return hashlib.sha256((text or "").encode("utf-8")).hexdigest()[:16]


def first_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def _normalized_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u3000", " ").strip())


def _cjk_number(value: str) -> int | None:
    if value.isdigit():
        return int(value)
    if value and all(char in _CJK_DIGITS for char in value):
        return int("".join(str(_CJK_DIGITS[char]) for char in value))
    if value == "十":
        return 10
    if len(value) == 2 and value.startswith("十") and value[1] in _CJK_DIGITS:
        return 10 + _CJK_DIGITS[value[1]]
    if len(value) == 2 and value.endswith("十") and value[0] in _CJK_DIGITS:
        return _CJK_DIGITS[value[0]] * 10
    if len(value) == 3 and value[1] == "十" and value[0] in _CJK_DIGITS and value[2] in _CJK_DIGITS:
        return _CJK_DIGITS[value[0]] * 10 + _CJK_DIGITS[value[2]]
    return None


def _date_identity(text: str) -> str | None:
    match = _DATE_RE.fullmatch(text)
    if match is None:
        return None
    year = _cjk_number(match.group("year"))
    month = _cjk_number(match.group("month"))
    day = _cjk_number(match.group("day"))
    if year is None or month is None or day is None or not 1 <= month <= 12 or not 1 <= day <= 31:
        return None
    return f"date:{year:04d}-{month:02d}-{day:02d}"


def _is_heading(item: dict[str, Any]) -> bool:
    return bool(item.get("heading_level")) or str(item.get("type", "")).startswith("heading")


def _heading_content(text: str) -> tuple[str, bool]:
    match = _HEADING_PREFIX_RE.match(text)
    if match is None:
        return text, False
    return text[match.end():].strip(), True


def _identity_for(item: dict[str, Any]) -> tuple[str, str]:
    """Return a private in-memory alignment key and its explicit normalization reason."""
    text = _normalized_text(str(item.get("_text", "")))
    if not text:
        return "", "empty"
    date_identity = _date_identity(text)
    if date_identity:
        return date_identity, "date_normalization"
    heading_content, had_prefix = _heading_content(text)
    if _is_heading(item) or had_prefix:
        had_period = heading_content.endswith("。")
        if had_period:
            heading_content = heading_content[:-1].rstrip()
        if heading_content:
            if had_prefix:
                return "heading:" + heading_content, "heading_number_normalization"
            if had_period:
                return "heading:" + heading_content, "heading_terminal_period"
            return "heading:" + heading_content, "heading_content"
    return text, "exact_text"


def _identity_hash(item: dict[str, Any]) -> str:
    identity, _reason = _identity_for(item)
    return text_hash(identity) if identity else ""


def _recognition_region(item: dict[str, Any], index: int) -> str:
    type_id = str(item.get("type", ""))
    section = str(item.get("section", ""))
    if type_id in {"__letterhead__", "dispatch_number", "signer", "letterhead_separator"}:
        return "letterhead"
    if type_id.startswith("attachment") or "attachment" in section:
        return "attachment"
    if type_id in {"sign_org", "sign_date", "responsibility_line"} or "tail" in section:
        return "tail"
    if type_id in {"title", "title_cont", "role_name", "date_line", "salutation", "meeting_meta"}:
        return "front"
    return "body"


def _physical_region(item: dict[str, Any], index: int, total: int) -> str:
    style = str(item.get("style", ""))
    text = _normalized_text(str(item.get("_text", "")))
    if style.startswith("DCT-Letterhead") or style == "DCT-DocumentNumber":
        return "letterhead"
    if style.startswith("DCT-Attachment") or text.startswith("附件"):
        return "attachment"
    if style.startswith("DCT-Sign"):
        return "tail"
    if style in {"DCT-Title", "DCT-TitleCont", "DCT-RoleName", "DCT-DateLine"}:
        return "front"
    if index > max(4, total * 3 // 4) and _date_identity(text):
        return "tail"
    return "body"


def _paragraph_property_chain(paragraph) -> Iterable[Any]:
    direct = getattr(getattr(paragraph, "_p", None), "pPr", None)
    if direct is not None:
        yield direct
    seen: set[str] = set()
    style = getattr(paragraph, "style", None)
    while style is not None:
        style_id = str(getattr(style, "style_id", "") or id(style))
        if style_id in seen:
            break
        seen.add(style_id)
        properties = getattr(getattr(style, "element", None), "pPr", None)
        if properties is not None:
            yield properties
        style = getattr(style, "base_style", None)
    styles_element = getattr(getattr(getattr(paragraph, "part", None), "styles", None), "element", None)
    defaults = styles_element.find(qn("w:docDefaults")) if styles_element is not None else None
    default = defaults.find(qn("w:pPrDefault")) if defaults is not None else None
    properties = default.find(qn("w:pPr")) if default is not None else None
    if properties is not None:
        yield properties


def _indent_chars(paragraph, name: str, fallback_size_pt: float = 16.0) -> float:
    char_attribute = {
        "first_indent": "w:firstLineChars",
        "left_indent": "w:leftChars",
        "right_indent": "w:rightChars",
    }[name]
    twip_attribute = {
        "first_indent": "w:firstLine",
        "left_indent": "w:left",
        "right_indent": "w:right",
    }[name]
    for properties in _paragraph_property_chain(paragraph):
        indent = properties.find(qn("w:ind"))
        if indent is None:
            continue
        value = indent.get(qn(char_attribute))
        if value is not None:
            try:
                return round(float(value) / 100.0, 3)
            except ValueError:
                pass
        value = indent.get(qn(twip_attribute))
        if value is not None:
            try:
                return round(float(value) / max(fallback_size_pt * 20.0, 1.0), 3)
            except ValueError:
                pass
    return 0.0


def _protected_object_kind(paragraph) -> str:
    text = _normalized_text(paragraph.text)
    if not _OBJECT_CAPTION_RE.match(text):
        return ""
    previous = paragraph._p.getprevious()
    if previous is None:
        return ""
    if previous.tag == qn("w:tbl"):
        return "table_caption"
    if previous.tag == qn("w:p"):
        has_drawing = any(previous.find(".//" + qn(tag)) is not None for tag in ("w:drawing", "w:pict", "w:object"))
        visible_text = "".join(previous.itertext()).strip()
        if has_drawing and not visible_text:
            return "image_caption"
    return ""


def _physical_item(paragraph, index: int, total: int) -> dict[str, Any]:
    run = first_run(paragraph)
    fmt = paragraph.paragraph_format
    size_pt = round(run.font.size.pt, 3) if run and run.font.size else 16.0
    protected_kind = _protected_object_kind(paragraph)
    item = {
        "index": index,
        "_text": paragraph.text,
        "text_hash": text_hash(paragraph.text),
        "length": len(paragraph.text),
        "style": paragraph.style.name if paragraph.style else "",
        "style_id": paragraph.style.style_id if paragraph.style else "",
        "font": run.font.name if run else "",
        "size_pt": round(run.font.size.pt, 3) if run and run.font.size else None,
        "bold": bool(run.bold) if run else False,
        "italic": bool(run.italic) if run else False,
        "underline": bool(run.underline) if run else False,
        "alignment": str(paragraph.alignment) if paragraph.alignment is not None else "",
        "first_indent": _indent_chars(paragraph, "first_indent", size_pt),
        "left_indent": _indent_chars(paragraph, "left_indent", size_pt),
        "right_indent": _indent_chars(paragraph, "right_indent", size_pt),
        "first_indent_raw": int(fmt.first_line_indent) if fmt.first_line_indent else 0,
        "left_indent_raw": int(fmt.left_indent) if fmt.left_indent else 0,
        "right_indent_raw": int(fmt.right_indent) if fmt.right_indent else 0,
        "line_spacing": str(fmt.line_spacing) if fmt.line_spacing is not None else "",
        "space_before": int(fmt.space_before) if fmt.space_before else 0,
        "space_after": int(fmt.space_after) if fmt.space_after else 0,
        "page_break_before": bool(fmt.page_break_before),
        "keep_with_next": bool(fmt.keep_with_next),
        "keep_together": bool(fmt.keep_together),
        "protected_object": bool(protected_kind),
        "protected_kind": protected_kind,
    }
    item["region"] = _physical_region(item, index, total)
    return item


def physical_snapshot(path: Path) -> dict[str, Any]:
    document = Document(path)
    total = len(document.paragraphs)
    return {
        "paragraphs": [_physical_item(paragraph, index, total) for index, paragraph in enumerate(document.paragraphs)],
        "tables": len(document.tables),
        "sections": len(document.sections),
        "inline_shapes": len(document.inline_shapes),
        "headers": sum(len(section.header.paragraphs) for section in document.sections),
        "footers": sum(len(section.footer.paragraphs) for section in document.sections),
    }


def recognition_snapshot(path: Path, rules: list[StyleRule], *, processing_strategy: str) -> dict[str, Any]:
    data = DocxImporter().load(
        str(path),
        rules,
        features={
            "classification": {"enabled": True},
            "processing": {"strategy": processing_strategy},
        },
    )
    paragraphs = []
    for index, item in enumerate(data.paragraphs):
        features = item.features
        snapshot_item = {
            "index": index,
            "_text": item.original_text or item.text,
            "text_hash": text_hash(item.original_text or item.text),
            "length": len(item.original_text or item.text),
            "type": item.type_id,
            "recognition_type": item.meta.get("recognition_type", ""),
            "section": item.meta.get("recognition_section", ""),
            "heading_level": _heading_level(item.type_id),
            "font": features.font_name if features else "",
            "size_pt": features.font_size_pt if features else None,
            "bold": bool(features.bold) if features else False,
            "alignment": features.alignment if features else "",
            "first_indent": features.first_line_indent if features else 0,
            "style": features.style_name if features else "",
        }
        snapshot_item["region"] = _recognition_region(snapshot_item, index)
        paragraphs.append(snapshot_item)
    diagnostics = getattr(data, "recognition_diagnostics", {}) or {}
    return {
        "mode": data.doc_mode,
        "paragraphs": paragraphs,
        "tables": len(data.tables),
        "blocks": len(getattr(data, "body_blocks", []) or []),
        "images": sum(1 for item in paragraphs if item.get("contains_image")),
        "diagnostic_summary": diagnostics.get("summary", {}),
    }


def _heading_level(type_id: str) -> int | None:
    match = re.fullmatch(r"heading([1-4])(?:_report)?", type_id or "")
    return int(match.group(1)) if match else None


def template_for(source: Path) -> tuple[Path | None, str]:
    templates = sorted(TEMPLATE_DIR.glob("*.docx"))
    if not templates:
        return None, "无模板"
    match = re.match(r"^(\d{3})_", source.name)
    if match:
        numbered = [path for path in templates if path.name.startswith(match.group(1) + "_")]
        if len(numbered) == 1:
            return numbered[0], "编号匹配"
        if len(numbered) > 1:
            return None, "编号模板不唯一"
    if len(templates) == 1:
        return templates[0], "单模板统一对照"
    return None, "模板匹配不明确"


def letterhead_options_from_template(template: Path | None) -> dict | None:
    """Build managed letterhead options from a correct-format reference file."""
    if template is None:
        return None
    paragraphs = [paragraph.text.strip() for paragraph in Document(template).paragraphs]
    mark = next((text for text in paragraphs[:20] if text.endswith("文件") and len(text) > 2), "")
    number_match = next((
        _DISPATCH_NUMBER_RE.fullmatch(text)
        for text in paragraphs[:20]
        if _DISPATCH_NUMBER_RE.fullmatch(text)
    ), None)
    if not mark or number_match is None:
        return None
    options = default_letterhead_config()
    options.update({
        "enabled": True,
        "agencies": [{
            "id": "template-agency-1", "name": mark[:-2], "short_name": "",
            "role": "sponsor", "order": 1,
        }],
        "document_number": {
            "agency_code": number_match.group("agency_code"),
            "year": int(number_match.group("year")),
            "sequence": int(number_match.group("sequence")),
        },
    })
    return options


def _alignment_keys(items: Iterable[dict[str, Any]]) -> list[str]:
    keys = []
    for item in items:
        identity, reason = _identity_for(item)
        # Empty paragraphs are intentionally not cross-matched: inserted spacers must not shift text anchors.
        keys.append(identity if identity else f"__empty_{item['index']}_{reason}")
    return keys


def align_paragraphs(
    actual_items: list[dict[str, Any]], expected_items: list[dict[str, Any]],
) -> dict[str, Any]:
    """Align per semantic region while preserving duplicate paragraph order."""
    pairs: list[tuple[dict[str, Any], dict[str, Any], str]] = []
    actual_unmatched: list[dict[str, Any]] = []
    expected_unmatched: list[dict[str, Any]] = []
    for region in _REGIONS:
        actual_region = [item for item in actual_items if item.get("region") == region]
        expected_region = [item for item in expected_items if item.get("region") == region]
        matcher = SequenceMatcher(
            None, _alignment_keys(actual_region), _alignment_keys(expected_region), autojunk=False,
        )
        for tag, actual_start, actual_end, expected_start, expected_end in matcher.get_opcodes():
            if tag == "equal":
                for offset in range(actual_end - actual_start):
                    actual = actual_region[actual_start + offset]
                    expected = expected_region[expected_start + offset]
                    actual_identity, actual_reason = _identity_for(actual)
                    expected_identity, expected_reason = _identity_for(expected)
                    reason = "exact_text"
                    if _normalized_text(actual.get("_text", "")) != _normalized_text(expected.get("_text", "")):
                        reason = actual_reason if actual_reason != "exact_text" else expected_reason
                    pairs.append((actual, expected, reason))
            else:
                actual_unmatched.extend(actual_region[actual_start:actual_end])
                expected_unmatched.extend(expected_region[expected_start:expected_end])
    pairs.sort(key=lambda pair: pair[0]["index"])
    actual_unmatched.sort(key=lambda item: item["index"])
    expected_unmatched.sort(key=lambda item: item["index"])
    return {"pairs": pairs, "actual_unmatched": actual_unmatched, "expected_unmatched": expected_unmatched}


def _is_generated_letterhead(item: dict[str, Any], source: str) -> bool:
    if source == "recognition":
        return item.get("type") in {"__letterhead__", "dispatch_number", "signer", "letterhead_separator"}
    style = str(item.get("style", ""))
    return style.startswith("DCT-Letterhead") or style == "DCT-DocumentNumber"


def _difference(
    category: str,
    actual: Any,
    expected: Any,
    *,
    source: str,
    actual_index: int | None,
    expected_index: int | None,
    match_reason: str,
    expected_change: bool = False,
    expected_reason: str = "",
    difference_origin: str = "output_regression",
    rule_basis: str = "template_reference",
) -> dict[str, Any]:
    return {
        "category": category,
        "actual": actual,
        "expected": expected,
        "paragraph_index": actual_index,
        "actual_paragraph_index": actual_index,
        "expected_paragraph_index": expected_index,
        "comparison_source": source,
        "match_reason": match_reason,
        "expected_change": expected_change,
        "expected_reason": expected_reason,
        "difference_origin": difference_origin,
        "rule_basis": rule_basis,
    }


def severity(category: str, *, expected_change: bool = False) -> str:
    if expected_change:
        return "P3"
    if category in {
        "output_addition", "template_missing", "source_text_addition", "source_text_loss",
        "text_change", "mode", "tables", "images",
        "type", "recognition_type", "section", "heading_level", "region",
    }:
        return "P1"
    if category in {
        "font", "size_pt", "bold", "italic", "underline", "alignment", "first_indent",
        "left_indent", "right_indent", "line_spacing", "space_before", "space_after",
        "style", "page_break_before", "keep_with_next", "keep_together", "sections",
        "inline_shapes", "headers", "footers",
    }:
        return "P2"
    return "P3"


def _field_values_equal(field: str, actual: Any, expected: Any) -> bool:
    if field in _INDENT_FIELDS:
        try:
            return abs(float(actual or 0.0) - float(expected or 0.0)) <= _INDENT_TOLERANCE_CHARS
        except (TypeError, ValueError):
            return actual == expected
    if field == "size_pt" and actual is not None and expected is not None:
        try:
            return abs(float(actual) - float(expected)) <= 0.05
        except (TypeError, ValueError):
            return actual == expected
    return actual == expected


def _mark_expected(item: dict[str, Any], *, reason: str, origin: str, rule_basis: str) -> None:
    item["expected_change"] = True
    item["expected_reason"] = reason
    item["difference_origin"] = origin
    item["rule_basis"] = rule_basis


def compare_snapshots(
    actual: dict[str, Any],
    expected: dict[str, Any],
    *,
    source: str,
    fields: tuple[str, ...],
    document_fields: tuple[str, ...],
    report_unmatched: bool,
) -> tuple[list[dict[str, Any]], dict[str, int]]:
    alignment = align_paragraphs(actual.get("paragraphs", []), expected.get("paragraphs", []))
    differences: list[dict[str, Any]] = []
    expected_normalizations = 0
    for actual_item, expected_item, reason in alignment["pairs"]:
        raw_actual = _normalized_text(str(actual_item.get("_text", "")))
        raw_expected = _normalized_text(str(expected_item.get("_text", "")))
        if raw_actual != raw_expected:
            expected_normalizations += 1
            differences.append(_difference(
                "expected_normalization", actual_item.get("text_hash"), expected_item.get("text_hash"),
                source=source, actual_index=actual_item["index"], expected_index=expected_item["index"],
                match_reason=reason, expected_change=True,
                expected_reason=reason,
                difference_origin="expected_normalization",
                rule_basis="configured_normalization",
            ))
        for field in fields:
            if not _field_values_equal(field, actual_item.get(field), expected_item.get(field)):
                difference = _difference(
                    field, actual_item.get(field), expected_item.get(field),
                    source=source, actual_index=actual_item["index"], expected_index=expected_item["index"],
                    match_reason=reason,
                )
                difference.update({
                    "actual_identity_hash": _identity_hash(actual_item),
                    "expected_identity_hash": _identity_hash(expected_item),
                    "actual_style_id": actual_item.get("style_id", ""),
                    "expected_style_id": expected_item.get("style_id", ""),
                    "actual_protected_object": bool(actual_item.get("protected_object")),
                    "expected_protected_object": bool(expected_item.get("protected_object")),
                    "protected_kind": actual_item.get("protected_kind", "") or expected_item.get("protected_kind", ""),
                })
                differences.append(difference)
    if report_unmatched:
        for item in alignment["actual_unmatched"]:
            generated = _is_generated_letterhead(item, source)
            empty = not _normalized_text(str(item.get("_text", "")))
            differences.append(_difference(
                "output_addition", item.get("text_hash"), None,
                source=source, actual_index=item["index"], expected_index=None,
                match_reason="generated_letterhead" if generated else "unmatched_output",
                expected_change=generated or empty,
                expected_reason="generated_letterhead" if generated else ("empty_paragraph" if empty else ""),
                difference_origin="expected_normalization" if generated or empty else "output_regression",
                rule_basis="configured_letterhead" if generated else "template_reference",
            ))
            differences[-1]["actual_identity_hash"] = _identity_hash(item)
        for item in alignment["expected_unmatched"]:
            empty = not _normalized_text(str(item.get("_text", "")))
            differences.append(_difference(
                "template_missing", None, item.get("text_hash"),
                source=source, actual_index=None, expected_index=item["index"],
                match_reason="unmatched_template",
                expected_change=empty,
                expected_reason="empty_paragraph" if empty else "",
                difference_origin="expected_normalization" if empty else "output_regression",
            ))
            differences[-1]["expected_identity_hash"] = _identity_hash(item)
    for field in document_fields:
        if actual.get(field) != expected.get(field):
            differences.append(_difference(
                field, actual.get(field), expected.get(field), source=source,
                actual_index=None, expected_index=None, match_reason="document_property",
            ))
    stats = {
        "matched_pairs": len(alignment["pairs"]),
        "output_additions": len(alignment["actual_unmatched"]),
        "template_missing": len(alignment["expected_unmatched"]),
        "expected_normalizations": expected_normalizations,
    }
    return differences, stats


def _mark_input_fixture_differences(
    differences: list[dict[str, Any]], source_snapshot: dict[str, Any], expected_snapshot: dict[str, Any],
) -> Counter[str]:
    """Separate source-fixture content variants from output-introduced changes."""
    baseline, _ = compare_snapshots(
        source_snapshot, expected_snapshot, source="recognition", fields=(), document_fields=(), report_unmatched=True,
    )
    source_additions = Counter(
        item.get("actual_identity_hash", "") for item in baseline
        if item["category"] == "output_addition" and not item["expected_change"]
    )
    source_missing = Counter(
        item.get("expected_identity_hash", "") for item in baseline
        if item["category"] == "template_missing" and not item["expected_change"]
    )
    source_identities = Counter(_identity_hash(item) for item in source_snapshot.get("paragraphs", []))
    expected_identities = Counter(_identity_hash(item) for item in expected_snapshot.get("paragraphs", []))
    marked: Counter[str] = Counter()
    for item in differences:
        if item["comparison_source"] != "recognition" or item["expected_change"]:
            continue
        identity = ""
        if item["category"] == "output_addition":
            identity = item.get("actual_identity_hash", "")
            if source_additions[identity] <= 0:
                continue
            source_additions[identity] -= 1
        elif item["category"] == "template_missing":
            identity = item.get("expected_identity_hash", "")
            if source_missing[identity] <= 0:
                continue
            source_missing[identity] -= 1
        else:
            continue
        reason = (
            "source_order_difference"
            if identity and source_identities[identity] and expected_identities[identity]
            else "input_fixture_difference"
        )
        _mark_expected(
            item,
            reason=reason,
            origin=reason,
            rule_basis="source_fixture",
        )
        item["match_reason"] = reason
        marked[reason] += 1
    return marked


def _source_preservation_differences(
    actual_snapshot: dict[str, Any], source_snapshot: dict[str, Any],
    existing_differences: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], dict[str, int]]:
    """Report output-only text changes that a template comparison cannot see."""
    preservation, stats = compare_snapshots(
        actual_snapshot, source_snapshot, source="source_preservation",
        fields=(), document_fields=(), report_unmatched=True,
    )
    existing = {
        (item.get("actual_identity_hash", ""), item.get("expected_identity_hash", ""))
        for item in existing_differences
        if not item.get("expected_change") and item.get("comparison_source") == "recognition"
    }
    results: list[dict[str, Any]] = []
    for item in preservation:
        if item.get("expected_change"):
            continue
        if item["category"] == "output_addition":
            key = (item.get("actual_identity_hash", ""), "")
            category = "source_text_addition"
        elif item["category"] == "template_missing":
            key = ("", item.get("expected_identity_hash", ""))
            category = "source_text_loss"
        else:
            continue
        if key in existing:
            continue
        item["category"] = category
        item["difference_origin"] = "output_regression"
        item["rule_basis"] = "source_text_preservation"
        item["match_reason"] = "output_source_mismatch"
        results.append(item)
    return results, stats


def _mark_project_config_differences(differences: list[dict[str, Any]]) -> Counter[str]:
    marked: Counter[str] = Counter()
    for item in differences:
        if item.get("expected_change") or item.get("comparison_source") != "physical":
            continue
        style_id = str(item.get("actual_style_id", ""))
        category = str(item.get("category", ""))
        actual = item.get("actual")
        configured = False
        if category == "size_pt" and style_id == "DCT-LetterheadMark":
            configured = _field_values_equal(category, actual, 32.0)
        elif category == "right_indent" and style_id == "DCT-Signature":
            configured = _field_values_equal(category, actual, 2.0)
        elif category == "right_indent" and style_id == "DCT-Date":
            configured = _field_values_equal(category, actual, 4.0)
        elif category == "first_indent" and style_id in {"DCT-Body", "DCT-Responsibility"}:
            configured = _field_values_equal(category, actual, 2.0)
        elif category == "left_indent" and style_id == "DCT-AttachmentNoteItem":
            configured = _field_values_equal(category, actual, 5.0)
        if not configured:
            continue
        _mark_expected(
            item,
            reason="project_config_difference",
            origin="project_config_difference",
            rule_basis="GB/T 9704-2012 + versioned_project_config",
        )
        marked[category] += 1
    return marked


def _mark_protected_source_differences(
    differences: list[dict[str, Any]], actual_snapshot: dict[str, Any], source_snapshot: dict[str, Any],
) -> int:
    alignment = align_paragraphs(
        actual_snapshot.get("paragraphs", []), source_snapshot.get("paragraphs", []),
    )
    source_by_actual = {actual["index"]: source for actual, source, _reason in alignment["pairs"]}
    marked = 0
    for item in differences:
        if (
            item.get("expected_change")
            or item.get("comparison_source") != "physical"
            or item.get("category") not in _PROTECTED_FORMAT_FIELDS
            or not item.get("actual_protected_object")
        ):
            continue
        source_item = source_by_actual.get(item.get("actual_paragraph_index"))
        if source_item is None or not source_item.get("protected_object"):
            continue
        field = str(item["category"])
        if not _field_values_equal(field, item.get("actual"), source_item.get(field)):
            continue
        _mark_expected(
            item,
            reason="protected_source_difference",
            origin="protected_source_difference",
            rule_basis="protected_object_preservation_contract",
        )
        marked += 1
    return marked


def compare_documents(
    actual: dict[str, dict[str, Any]], expected: dict[str, dict[str, Any]],
    *, source_recognition: dict[str, Any] | None = None,
    source_physical: dict[str, Any] | None = None,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    recognition_diffs, recognition_stats = compare_snapshots(
        actual["recognition"], expected["recognition"], source="recognition",
        fields=_RECOGNITION_FIELDS, document_fields=("mode", "tables", "images"), report_unmatched=True,
    )
    physical_diffs, physical_stats = compare_snapshots(
        actual["physical"], expected["physical"], source="physical",
        fields=_PHYSICAL_FIELDS, document_fields=("tables", "sections", "inline_shapes", "headers", "footers"),
        report_unmatched=False,
    )
    differences = recognition_diffs + physical_diffs
    fixture_reasons: Counter[str] = Counter()
    preservation_stats: dict[str, int] = {}
    if source_recognition is not None:
        fixture_reasons = _mark_input_fixture_differences(
            differences, source_recognition, expected["recognition"],
        )
        preservation, preservation_stats = _source_preservation_differences(
            actual["recognition"], source_recognition, differences,
        )
        differences.extend(preservation)
    protected_source_differences = 0
    if source_physical is not None:
        protected_source_differences = _mark_protected_source_differences(
            differences, actual["physical"], source_physical,
        )
    project_config_differences = _mark_project_config_differences(differences)
    for item in differences:
        item["severity"] = severity(item["category"], expected_change=bool(item["expected_change"]))
    real = [item for item in differences if not item["expected_change"]]
    expected_changes = [item for item in differences if item["expected_change"]]
    return differences, {
        "version": COMPARISON_VERSION,
        "alignment_method": "区域分段 + 去序号标题正文锚点 + 保序三方归因",
        "recognition": recognition_stats,
        "physical": physical_stats,
        "matched_pairs": recognition_stats["matched_pairs"],
        "output_additions": recognition_stats["output_additions"],
        "template_missing": recognition_stats["template_missing"],
        "expected_normalizations": sum(item["category"] == "expected_normalization" for item in expected_changes),
        "input_fixture_differences": sum(fixture_reasons.values()),
        "source_order_differences": fixture_reasons.get("source_order_difference", 0),
        "source_preservation": preservation_stats,
        "protected_source_differences": protected_source_differences,
        "project_config_differences": dict(project_config_differences),
        "unexpected_differences": len(real),
        "real_issue_counts": dict(Counter(item["severity"] for item in real)),
        "difference_origin_counts": dict(Counter(item["difference_origin"] for item in differences)),
        "expected_difference_counts": dict(Counter(item["expected_reason"] or item["category"] for item in expected_changes)),
    }


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", type=Path, default=OUTPUT_DIR, help="Empty directory under test_docx/tset1/test1测试结果 for this batch run.")
    parser.add_argument(
        "--special-output-dir",
        type=Path,
        default=SPECIAL_OUTPUT_DIR,
        help="Directory under test_docx/test2/测试结果 for the special regression set.",
    )
    parser.add_argument("--without-template-letterhead", action="store_true", help="Do not derive managed letterhead options from the matched correct template.")
    parser.add_argument("--strict-preservation", action="store_true", help="Use strict preservation instead of the default smart/structural mode.")
    parser.add_argument("--mode", choices=("structural", "strict", "normalize"), default="structural", help="Processing strategy. The default matches the frontend smart mode.")
    parser.add_argument("--render-review", action="store_true", help="Render the deterministic visual QA sample to PNG after processing.")
    parser.add_argument("--require-render", action="store_true", help="Fail the batch when requested visual rendering cannot complete.")
    parser.add_argument("--render-sample-size", type=int, default=10, help="Number of standard fixtures in the visual sample (default: 10).")
    args = parser.parse_args(argv)
    if args.require_render:
        args.render_review = True
    if args.render_sample_size < 1:
        parser.error("--render-sample-size must be positive")
    return args


def _validated_output_dir(path: Path) -> Path:
    output_dir = path.resolve()
    output_root = OUTPUT_DIR.resolve()
    try:
        output_dir.relative_to(output_root)
    except ValueError as exc:
        raise SystemExit(f"output directory must be inside {output_root}: {output_dir}") from exc
    if output_dir.exists() and any(output_dir.iterdir()):
        raise SystemExit(f"output directory is not empty: {output_dir}")
    return output_dir


def visual_rendering_not_run(reason: str = "未执行视觉渲染检查：未请求。") -> dict[str, Any]:
    return {"executed": False, "reason": reason, "documents": [], "suspected_pages": []}


def _validate_special_output_dir(path: Path) -> Path:
    output_dir = path.resolve()
    expected_root = SPECIAL_OUTPUT_DIR.resolve()
    try:
        output_dir.relative_to(expected_root)
    except ValueError as exc:
        raise SystemExit(
            f"special output directory must be inside {expected_root}: {output_dir}"
        ) from exc
    if output_dir != expected_root and output_dir.exists() and any(output_dir.iterdir()):
        raise SystemExit(f"special output directory is not empty: {output_dir}")
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _export_and_validate(
    importer: DocxImporter, source: Path, output: Path, rules: list[StyleRule], settings: PageSettings,
    processing_features: dict[str, Any], *, letterhead_options: dict[str, Any] | None = None,
) -> tuple[object, object]:
    """Write through a temporary file and accept it only after DOCX checks pass."""
    temporary = output.with_name(f".{output.stem}.tmp.docx")
    try:
        source_data = importer.load(str(source), rules, features=processing_features)
        export_doc(source_data, rules, settings, str(temporary), letterhead_options=letterhead_options)
        with zipfile.ZipFile(temporary) as package:
            bad_member = package.testzip()
            if bad_member:
                raise ValueError(f"ZIP_INTEGRITY:{bad_member}")
        output_data = importer.load(str(temporary), rules, features=processing_features)
        Document(temporary)
        temporary.replace(output)
        return source_data, output_data
    except Exception:
        if temporary.exists():
            temporary.unlink()
        raise


def source_heading_cue_audit(source_data: Any, output_data: Any) -> dict[str, Any]:
    """Check source-backed Word list headings against round-tripped output.

    The report intentionally stores hashes and indexes instead of paragraph
    text.  It is useful for special fixtures that have no canonical template:
    a short bold ``@lvl_N`` source item must remain the corresponding heading
    after export and re-import.
    """
    def heading_content_key(value: str) -> str:
        normalized = _normalized_text(value)
        return re.sub(
            r"^(?:[一二三四五六七八九十百千零〇]+、|"
            r"[（(][一二三四五六七八九十百千零〇0-9]+[）)]|"
            r"\d+[.．、])\s*",
            "",
            normalized,
            count=1,
        )

    output_by_text: dict[str, list[tuple[int, str]]] = {}
    for index, paragraph in enumerate(getattr(output_data, "paragraphs", ())):
        key = heading_content_key(str(getattr(paragraph, "text", "") or ""))
        if key:
            output_by_text.setdefault(key, []).append((index, str(getattr(paragraph, "type_id", ""))))

    cue_count = 0
    mismatches: list[dict[str, Any]] = []
    for index, paragraph in enumerate(getattr(source_data, "paragraphs", ())):
        features = getattr(paragraph, "features", None)
        marker = str(
            getattr(features, "segment_numbering_features", "")
            or getattr(features, "numbering_prefix", "")
            or ""
        )
        match = re.fullmatch(r"@lvl_(\d+)", marker)
        text = heading_content_key(str(getattr(paragraph, "text", "") or ""))
        bold_ratio = float(
            getattr(features, "segment_bold_char_ratio", 0.0)
            or getattr(features, "bold_char_ratio", 0.0)
            or 0.0
        )
        if match is None or not text or len(text) > 40 or bold_ratio < 0.65:
            continue
        level = min(int(match.group(1)) + 2, 4)
        expected_type = f"heading{level}"
        cue_count += 1
        matches = output_by_text.get(text, [])
        if any(type_id == expected_type for _output_index, type_id in matches):
            continue
        mismatches.append({
            "source_paragraph_index": index,
            "source_physical_paragraph_index": getattr(features, "source_physical_paragraph_index", None),
            "text_hash": text_hash(text),
            "expected_type": expected_type,
            "output_paragraph_indexes": [output_index for output_index, _type_id in matches],
            "output_types": sorted({type_id for _output_index, type_id in matches}),
            "evidence": [marker, "short-bold-source-list-item"],
        })
    return {"cue_count": cue_count, "mismatch_count": len(mismatches), "mismatches": mismatches}


def signature_continuity_audit(output_data: Any) -> dict[str, Any]:
    """Verify that recognized signature organizations directly precede dates."""
    visible = [
        (index, paragraph)
        for index, paragraph in enumerate(getattr(output_data, "paragraphs", ()))
        if str(getattr(paragraph, "text", "") or "").strip()
    ]
    issues: list[dict[str, Any]] = []
    for position, (output_index, paragraph) in enumerate(visible):
        if str(getattr(paragraph, "type_id", "")) != "sign_date":
            continue
        prior_orgs = [
            prior_position
            for prior_position, (_prior_index, prior) in enumerate(visible[:position])
            if str(getattr(prior, "type_id", "")) == "sign_org"
        ]
        if not prior_orgs:
            for prior_index, prior in reversed(visible[max(0, position - 12):position]):
                if str(getattr(prior, "type_id", "")) != "body":
                    continue
                source_text = str(
                    getattr(prior, "original_text", "")
                    or getattr(prior, "text", "")
                    or ""
                )
                source_lines = [line.strip() for line in source_text.splitlines() if line.strip()]
                candidate = source_lines[-1] if len(source_lines) >= 2 else ""
                if not (
                    2 <= len(candidate) <= 40
                    and not any(mark in candidate for mark in "。；;：:")
                    and _SIGNATURE_ORG_SUFFIX_RE.search(candidate)
                ):
                    continue
                issues.append({
                    "issue": "signature_org_embedded_in_body",
                    "organization_paragraph_index": prior_index,
                    "date_paragraph_index": output_index,
                    "organization_text_hash": text_hash(candidate),
                    "date_text_hash": text_hash(str(getattr(paragraph, "text", ""))),
                })
                break
            continue
        org_position = prior_orgs[-1]
        if org_position == position - 1:
            continue
        org_index, organization = visible[org_position]
        between = visible[org_position + 1:position]
        issues.append({
            "organization_paragraph_index": org_index,
            "date_paragraph_index": output_index,
            "organization_text_hash": text_hash(str(getattr(organization, "text", ""))),
            "date_text_hash": text_hash(str(getattr(paragraph, "text", ""))),
            "intervening_types": [str(getattr(item, "type_id", "")) for _index, item in between],
        })
    return {"issue_count": len(issues), "issues": issues}


def _special_record(
    importer: DocxImporter, source: Path, output_dir: Path, rules: list[StyleRule], settings: PageSettings,
    processing_features: dict[str, Any], processing_strategy: str,
) -> tuple[dict[str, Any], Path]:
    start = time.perf_counter()
    output = output_dir / f"{source.stem}_排版结果.docx"
    record: dict[str, Any] = {"文件名": source.name, "输出文件名": output.name, "处理策略": processing_strategy, "成功": False, "错误": "", "耗时_ms": 0}
    try:
        source_data, output_data = _export_and_validate(importer, source, output, rules, settings, processing_features)
        diagnostics = getattr(output_data, "recognition_diagnostics", {}) or {}
        summary = diagnostics.get("summary", {}) if isinstance(diagnostics, dict) else {}
        title_cue_audit = source_heading_cue_audit(source_data, output_data)
        signature_audit = signature_continuity_audit(output_data)
        record.update({
            "成功": True, "原文段落数": len(source_data.paragraphs), "输出段落数": len(output_data.paragraphs),
            "原文文字哈希": text_hash("\n".join(item.original_text for item in source_data.paragraphs)),
            "输出文字哈希": text_hash("\n".join(item.original_text for item in output_data.paragraphs)),
            "结构复核数": int(summary.get("needs_review_count", 0) or 0),
            "关键结构复核数": int(summary.get("critical_review_count", 0) or 0), "表格数": len(output_data.tables),
            "源标题线索数": title_cue_audit["cue_count"],
            "源标题线索未保留数": title_cue_audit["mismatch_count"],
            "源标题线索检查": title_cue_audit["mismatches"],
            "落款连续性问题数": signature_audit["issue_count"],
            "落款连续性检查": signature_audit["issues"],
        })
    except Exception as exc:
        record["错误"] = f"{type(exc).__name__}: {exc}"
    record["耗时_ms"] = round((time.perf_counter() - start) * 1000, 2)
    return record, output


def _rendering_dependencies() -> tuple[str, Any]:
    soffice = os.environ.get("DOCXTOOL_SOFFICE", "") or shutil.which("soffice")
    if not soffice:
        raise RenderUnavailable("未找到 soffice；请安装 LibreOffice 或设置 DOCXTOOL_SOFFICE。")
    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RenderUnavailable("缺少 PyMuPDF 开发依赖；请安装 requirements-dev.lock。") from exc
    return soffice, fitz


def _ink_ratio(pixmap: Any) -> float:
    samples = memoryview(pixmap.samples)
    components = max(int(pixmap.n), 3)
    pixels = max(1, len(samples) // components)
    step = 16
    inspected = 0
    ink = 0
    for pixel in range(0, pixels, step):
        offset = pixel * components
        if offset + 2 >= len(samples):
            break
        inspected += 1
        if samples[offset] < 245 or samples[offset + 1] < 245 or samples[offset + 2] < 245:
            ink += 1
    return round(ink / max(1, inspected), 6)


def _page_visual_status(text: str, ink_ratio: float) -> tuple[str, bool, int]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    visible_lines = [line for line in lines if not _PAGE_NUMBER_RE.fullmatch(line)]
    visible_characters = sum(len(re.sub(r"\s+", "", line)) for line in visible_lines)
    joined = " ".join(visible_lines)
    attachment_title = "附件" in joined and visible_characters <= 80
    if visible_characters == 0 and ink_ratio < 0.012:
        return "empty_page", True, visible_characters
    if visible_characters <= 35 or ink_ratio < 0.004:
        if attachment_title:
            return "attachment_title_page", False, visible_characters
        return "sparse_page", True, visible_characters
    return "normal", False, visible_characters


def _cleanup_render_temporary(path: Path) -> None:
    """Best-effort cleanup for LibreOffice profile files released late on Windows."""
    for _ in range(5):
        try:
            shutil.rmtree(path)
            return
        except FileNotFoundError:
            return
        except OSError:
            time.sleep(0.2)


def _render_document(doc_path: Path, artifact_dir: Path, *, label: str) -> dict[str, Any]:
    soffice, fitz = _rendering_dependencies()
    artifact_dir.mkdir(parents=True, exist_ok=True)
    runtime_root = ROOT / "var" / "runtime"
    runtime_root.mkdir(parents=True, exist_ok=True)
    temporary = Path(tempfile.mkdtemp(prefix="docxtool-render-", dir=str(runtime_root)))
    try:
        profile = temporary / "profile"
        converted = temporary / "converted"
        profile.mkdir()
        converted.mkdir()
        command = [
            soffice, f"-env:UserInstallation={profile.resolve().as_uri()}", "--headless", "--nologo",
            "--nodefault", "--nolockcheck", "--norestore", "--convert-to", "pdf:writer_pdf_Export",
            "--outdir", str(converted), str(doc_path.resolve()),
        ]
        completed = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, check=False)
        pdf_candidates = sorted(converted.glob("*.pdf"))
        if not pdf_candidates:
            details = (completed.stderr or completed.stdout or "LibreOffice 未产生 PDF").strip()
            raise RuntimeError(f"渲染失败：{details[:300]}")
        pdf_path = artifact_dir / f"{label}.pdf"
        shutil.copy2(pdf_candidates[0], pdf_path)
    finally:
        _cleanup_render_temporary(temporary)
    pages = []
    document = fitz.open(str(pdf_path))
    try:
        for page_number, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            page_path = artifact_dir / f"page-{page_number}.png"
            pixmap.save(str(page_path))
            text = page.get_text("text")
            ink_ratio = _ink_ratio(pixmap)
            status, suspected, visible_characters = _page_visual_status(text, ink_ratio)
            pages.append({
                "page": page_number, "image": page_path.name, "visible_characters": visible_characters,
                "ink_ratio": ink_ratio, "status": status, "suspected": suspected,
            })
    finally:
        document.close()
    return {"file": doc_path.name, "artifact_dir": str(artifact_dir), "pdf": pdf_path.name, "page_count": len(pages), "pages": pages}


def select_render_samples(records: list[dict[str, Any]], sample_size: int) -> list[dict[str, Any]]:
    """Select a deterministic, risk-first sample without relying on report order."""
    successful = [record for record in records if record.get("成功")]
    if len(successful) <= sample_size:
        return successful
    def score(record: dict[str, Any]) -> tuple[int, str]:
        alignment = record.get("模板对齐", {}) or {}
        real = alignment.get("real_issue_counts", {}) or {}
        expected = alignment.get("expected_difference_counts", {}) or {}
        diagnostics = record.get("结构诊断", {}) or {}
        value = int(real.get("P1", 0)) * 100 + int(real.get("P2", 0)) * 10
        value += int(diagnostics.get("critical_review_count", 0) or 0) * 80
        value += int(diagnostics.get("needs_review_count", 0) or 0) * 5
        value += int(expected.get("protected_source_difference", 0) or 0) * 30
        value += int(record.get("表格数", 0) or 0) * 2
        return value, str(record.get("编号", ""))
    selected: list[dict[str, Any]] = []
    for record in sorted(successful, key=lambda item: (-score(item)[0], score(item)[1])):
        if score(record)[0] <= 0:
            break
        selected.append(record)
        if len(selected) == sample_size:
            return selected
    remaining = [record for record in successful if record not in selected]
    needed = sample_size - len(selected)
    if needed == 1:
        selected.append(remaining[len(remaining) // 2])
    else:
        for offset in range(needed):
            index = round(offset * (len(remaining) - 1) / (needed - 1))
            candidate = remaining[index]
            if candidate not in selected:
                selected.append(candidate)
    return sorted(selected, key=lambda item: str(item.get("编号", "")))


def run_visual_review(
    standard_records: list[dict[str, Any]], standard_outputs: dict[str, Path], special_records: list[dict[str, Any]],
    special_outputs: dict[str, Path], template_paths: Iterable[Path], *, standard_root: Path,
    special_root: Path, sample_size: int,
) -> dict[str, Any]:
    try:
        _rendering_dependencies()
    except RenderUnavailable as exc:
        return visual_rendering_not_run(str(exc))
    documents: list[dict[str, Any]] = []
    failures: list[dict[str, str]] = []
    selected_standard = select_render_samples(standard_records, sample_size)
    jobs = [("standard", record["编号"], standard_outputs[record["编号"]], standard_root / str(record["编号"])) for record in selected_standard]
    jobs.extend(("special", record["文件名"], special_outputs[record["文件名"]], special_root / Path(record["文件名"]).stem) for record in special_records if record.get("成功"))
    rendered_templates: set[Path] = set()
    for template in template_paths:
        if template in rendered_templates:
            continue
        rendered_templates.add(template)
        jobs.append(("template", template.name, template, standard_root / "模板" / template.stem))
    for kind, record_id, path, artifact_dir in jobs:
        try:
            rendered = _render_document(path, artifact_dir, label="document")
            rendered["kind"] = kind
            rendered["record_id"] = record_id
            rendered["artifact_dir"] = str(artifact_dir.relative_to(ROOT))
            for page in rendered["pages"]:
                page["image"] = str((artifact_dir / page["image"]).relative_to(ROOT))
            documents.append(rendered)
        except Exception as exc:
            failures.append({"kind": kind, "record_id": record_id, "error": f"{type(exc).__name__}: {exc}"})
    suspected_pages = [
        {"kind": document["kind"], "record_id": document["record_id"], "page": page["page"], "status": page["status"], "image": page["image"]}
        for document in documents for page in document["pages"] if page["suspected"]
    ]
    return {
        "executed": True, "renderer": "LibreOffice + PyMuPDF", "standard_sample_size": len(selected_standard),
        "standard_sample_ids": [record["编号"] for record in selected_standard], "documents": documents,
        "suspected_pages": suspected_pages, "failures": failures,
        "reason": "已完成视觉渲染抽查。" if not failures else "部分文档渲染失败；详见 failures。",
    }


def _report_lines(report: dict[str, Any]) -> list[str]:
    return [
        f"总数：{report['总数']}", f"成功：{report['成功数']}", f"失败：{report['失败数']}",
        f"存在真实模板问题的文件：{report['真实问题文件数']}",
        f"仅存在预期差异的文件：{report['预期差异文件数']}",
        f"含输入固有差异的文件：{report['输入固有差异文件数']}",
        "问题统计：" + json.dumps(report["问题统计"], ensure_ascii=False),
        "输入固有差异统计：" + json.dumps(report["输入固有差异统计"], ensure_ascii=False),
        "差异归因统计：" + json.dumps(report.get("差异归因统计", {}), ensure_ascii=False),
        "项目配置差异统计：" + json.dumps(report.get("项目配置差异统计", {}), ensure_ascii=False),
        "受保护对象差异统计：" + json.dumps(report.get("受保护对象差异统计", {}), ensure_ascii=False),
        f"源标题线索未保留：{report.get('源标题线索未保留数', 0)}",
        f"落款连续性问题：{report.get('落款连续性问题数', 0)}",
        "视觉渲染检查：" + str(report["视觉渲染检查"]["reason"]),
    ]


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    processing_strategy = "strict" if args.strict_preservation else args.mode
    sources = sorted(INPUT_DIR.glob("*.docx"))
    if len(sources) != 50:
        raise SystemExit(f"expected 50 input DOCX files, found {len(sources)}")
    output_dir = _validated_output_dir(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    rules = StyleRule.from_config()
    settings = PageSettings.from_config()
    importer = DocxImporter()
    template_letterheads: dict[Path, dict[str, Any] | None] = {}
    # The user-facing formatter enables safe heading numbering by default.
    # Keep the batch regression harness on the same explicit configuration so
    # level 1-4 source numbering is rebuilt consistently during comparison.
    processing_features: dict[str, Any] = {
        "processing": {"strategy": processing_strategy},
        "numbering": {"enabled": True, "mode": "safe"},
    }
    results: list[dict[str, Any]] = []
    standard_outputs: dict[str, Path] = {}
    templates_used: set[Path] = set()
    for source in sources:
        start = time.perf_counter()
        output = output_dir / source.name.replace("_乱格式测试.docx", "_排版结果.docx")
        template, template_status = template_for(source)
        letterhead_options = None
        if template is not None and not args.without_template_letterhead:
            templates_used.add(template)
            letterhead_options = template_letterheads.setdefault(template, letterhead_options_from_template(template))
        record: dict[str, Any] = {
            "编号": source.name[:3], "原始文件名": source.name, "输出文件名": output.name,
            "模板": template.name if template else "", "模板匹配": template_status,
            "模板版头": bool(letterhead_options), "处理策略": processing_strategy, "成功": False,
            "错误": "", "差异": [], "耗时_ms": 0,
        }
        try:
            source_data, output_data = _export_and_validate(
                importer, source, output, rules, settings, processing_features, letterhead_options=letterhead_options,
            )
            title_cue_audit = source_heading_cue_audit(source_data, output_data)
            signature_audit = signature_continuity_audit(output_data)
            standard_outputs[record["编号"]] = output
            record.update({
                "成功": True, "原文段落数": len(source_data.paragraphs), "输出段落数": len(output_data.paragraphs),
                "输出模式": output_data.doc_mode,
                "版头状态": output_data.letterhead_detection.status if output_data.letterhead_detection else "none",
                "版头详情": list(output_data.letterhead_detection.details) if output_data.letterhead_detection else [],
                "版头段落数": sum(item.type_id == "__letterhead__" for item in output_data.paragraphs),
                "发文字号数": sum(item.type_id == "dispatch_number" for item in output_data.paragraphs),
                "表格数": len(output_data.tables),
                "结构诊断": (getattr(output_data, "recognition_diagnostics", {}) or {}).get("summary", {}),
                "源标题线索数": title_cue_audit["cue_count"],
                "源标题线索未保留数": title_cue_audit["mismatch_count"],
                "源标题线索检查": title_cue_audit["mismatches"],
                "落款连续性问题数": signature_audit["issue_count"],
                "落款连续性检查": signature_audit["issues"],
                "原文文字哈希": text_hash("\n".join(item.original_text for item in source_data.paragraphs)),
                "输出文字哈希": text_hash("\n".join(item.original_text for item in output_data.paragraphs)),
            })
            if template:
                actual = {"recognition": recognition_snapshot(output, rules, processing_strategy=processing_strategy), "physical": physical_snapshot(output)}
                expected = {"recognition": recognition_snapshot(template, rules, processing_strategy=processing_strategy), "physical": physical_snapshot(template)}
                source_recognition = recognition_snapshot(source, rules, processing_strategy=processing_strategy)
                source_physical = physical_snapshot(source)
                record["差异"], record["模板对齐"] = compare_documents(
                    actual, expected,
                    source_recognition=source_recognition,
                    source_physical=source_physical,
                )
        except Exception as exc:
            record["错误"] = f"{type(exc).__name__}: {exc}"
        record["耗时_ms"] = round((time.perf_counter() - start) * 1000, 2)
        results.append(record)
    special_output_dir = _validate_special_output_dir(args.special_output_dir)
    special_results: list[dict[str, Any]] = []
    special_outputs: dict[str, Path] = {}
    for source in sorted(SPECIAL_INPUT_DIR.glob("*.docx")):
        record, output = _special_record(importer, source, special_output_dir, rules, settings, processing_features, processing_strategy)
        special_results.append(record)
        if record.get("成功"):
            special_outputs[record["文件名"]] = output
    visual_status = visual_rendering_not_run()
    if args.render_review:
        visual_status = run_visual_review(
            results, standard_outputs, special_results, special_outputs, templates_used,
            standard_root=output_dir / "视觉抽查", special_root=special_output_dir / "视觉抽查",
            sample_size=args.render_sample_size,
        )
    special_failures = [item for item in visual_status.get("failures", []) if item.get("kind") == "special"]
    special_visual = {
        **visual_status,
        "documents": [item for item in visual_status.get("documents", []) if item.get("kind") == "special"],
        "suspected_pages": [item for item in visual_status.get("suspected_pages", []) if item.get("kind") == "special"],
        "failures": special_failures,
        "reason": "已完成专项集视觉渲染抽查。" if visual_status.get("executed") and not special_failures else visual_status["reason"],
    }
    special_report = {
        "总数": len(special_results), "成功数": sum(bool(item["成功"]) for item in special_results),
        "失败数": sum(not item["成功"] for item in special_results), "处理策略": processing_strategy,
        "模板对比": f"专项集模板目录：{SPECIAL_TEMPLATE_DIR.relative_to(ROOT)}；当前专项报告仍以结构审计和视觉抽查为主。",
        "源标题线索未保留数": sum(int(item.get("源标题线索未保留数", 0) or 0) for item in special_results),
        "落款连续性问题数": sum(int(item.get("落款连续性问题数", 0) or 0) for item in special_results),
        "视觉渲染检查": special_visual, "结果": special_results,
    }
    (special_output_dir / "批量测试报告.json").write_text(json.dumps(special_report, ensure_ascii=False, indent=2), encoding="utf-8")
    (special_output_dir / "批量测试报告.txt").write_text("\n".join([
        f"总数：{special_report['总数']}", f"成功：{special_report['成功数']}", f"失败：{special_report['失败数']}",
        f"处理策略：{processing_strategy}", "模板对比：" + special_report["模板对比"],
        f"源标题线索未保留：{special_report['源标题线索未保留数']}",
        f"落款连续性问题：{special_report['落款连续性问题数']}",
        "视觉渲染检查：" + str(special_visual["reason"]),
        *[f"{item['文件名']} | 成功={item['成功']} | 复核={item.get('结构复核数', 0)} | 源标题未保留={item.get('源标题线索未保留数', 0)} | 落款连续性={item.get('落款连续性问题数', 0)} | 错误={item['错误']}" for item in special_results],
    ]), encoding="utf-8")
    real_issue_files = sum(bool((item.get("模板对齐", {}).get("unexpected_differences", 0))) for item in results if item.get("成功"))
    expected_only_files = sum(
        bool(item.get("差异")) and not bool(item.get("模板对齐", {}).get("unexpected_differences", 0))
        for item in results if item.get("成功")
    )
    fixture_difference_files = sum(
        bool(item.get("模板对齐", {}).get("input_fixture_differences", 0))
        for item in results if item.get("成功")
    )
    report = {
        "总数": len(results), "成功数": sum(bool(item["成功"]) for item in results),
        "失败数": sum(not item["成功"] for item in results), "差异文件数": sum(bool(item["差异"]) for item in results),
        "真实问题文件数": real_issue_files, "预期差异文件数": expected_only_files,
        "输入固有差异文件数": fixture_difference_files,
        "模板对齐方式": COMPARISON_VERSION,
        "问题统计": {level: sum(1 for item in results for diff in item["差异"] if diff["severity"] == level and not diff["expected_change"]) for level in ("P0", "P1", "P2", "P3")},
        "预期差异统计": dict(Counter(diff["category"] for item in results for diff in item["差异"] if diff["expected_change"])),
        "输入固有差异统计": dict(Counter(
            diff["category"] for item in results for diff in item["差异"]
            if diff.get("expected_reason") in {"input_fixture_difference", "source_order_difference"}
        )),
        "差异归因统计": dict(Counter(
            diff.get("difference_origin", "unclassified")
            for item in results for diff in item["差异"]
        )),
        "项目配置差异统计": dict(Counter(
            diff["category"] for item in results for diff in item["差异"]
            if diff.get("expected_reason") == "project_config_difference"
        )),
        "受保护对象差异统计": dict(Counter(
            diff["category"] for item in results for diff in item["差异"]
            if diff.get("expected_reason") == "protected_source_difference"
        )),
        "源标题线索未保留数": sum(int(item.get("源标题线索未保留数", 0) or 0) for item in results),
        "落款连续性问题数": sum(int(item.get("落款连续性问题数", 0) or 0) for item in results),
        "视觉渲染检查": visual_status,
        "专项集": {"目录": str(SPECIAL_INPUT_DIR.relative_to(ROOT)), "结果目录": str(special_output_dir.relative_to(ROOT)), "总数": special_report["总数"], "成功数": special_report["成功数"], "失败数": special_report["失败数"]},
        "结果": results,
    }
    (output_dir / "批量测试报告.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = _report_lines(report)
    for item in results:
        alignment = item.get("模板对齐", {})
        lines.append(f"{item['编号']} {item['原始文件名']} | 成功={item['成功']} | 真实差异={alignment.get('unexpected_differences', 0)} | 预期差异={alignment.get('expected_normalizations', 0)} | 错误={item['错误']}")
    (output_dir / "批量测试报告.txt").write_text("\n".join(lines), encoding="utf-8")
    render_failed = args.require_render and (not visual_status.get("executed") or bool(visual_status.get("failures")))
    return 0 if report["失败数"] == 0 and special_report["失败数"] == 0 and not render_failed else 1


if __name__ == "__main__":
    raise SystemExit(main())
