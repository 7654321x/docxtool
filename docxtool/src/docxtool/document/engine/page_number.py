"""Safe page-number footer helpers."""

from __future__ import annotations

import re
import copy
from collections.abc import Iterator, Mapping, Sequence
from itertools import count
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


_PAGE_INSTRUCTION_RE = re.compile(r"\b(?:PAGE|NUMPAGES)\b", re.IGNORECASE)
_PAGE_ONLY_TEXT_RE = re.compile(r"[\s\d/\\\-–—第页共总计：:PagePAGEofNUMPAGES]+", re.IGNORECASE)
_PAGE_DECORATION_TEXT_RE = re.compile(r"[\s\d/\\\-–—第页共总计：:Pageof]+", re.IGNORECASE)
_STATIC_PAGE_NUMBER_RE = re.compile(
    r"(?:"
    r"[—–-]\s*\d+\s*[—–-]"
    r"|\d+"
    r"|第\s*\d+\s*页(?:\s*(?:共|/)\s*\d+\s*页?)?"
    r"|\d+\s*/\s*\d+"
    r")"
)
_DASH_STYLES = {"dash", "numberindash", "number_in_dash"}
_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "o": "urn:schemas-microsoft-com:office:office",
    "v": "urn:schemas-microsoft-com:vml",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
}


def apply_page_number(document, options: Mapping[str, Any] | None = None):
    """Apply clean Word PAGE fields to document footers."""

    return apply_page_numbers(document, options)


def apply_page_numbers(document, options: Mapping[str, Any] | None = None):
    """Apply Word PAGE/NUMPAGES fields to section footers without erasing other content."""

    opts = dict(options or {})
    if opts.get("enabled", True) is False:
        return document

    style = str(opts.get("style", opts.get("format", "dash"))).lower()
    position = str(opts.get("position", opts.get("alignment", "outside"))).lower()
    first_page_policy = _first_page_policy(opts)
    section_numbering = str(opts.get("section_numbering", "continue")).lower()
    restart_at = int(opts.get("restart_at", 1))

    if position == "outside":
        _set_even_and_odd_headers(document, True)
    has_even_footer = position == "outside" or _even_and_odd_headers_enabled(document)
    _set_update_fields_on_open(document)
    shape_ids = count(_next_shape_id(document))

    for section_index, section in enumerate(document.sections):
        _set_footer_distance(section, opts)
        _set_section_numbering(section, section_numbering, restart_at, opts, section_index)
        _set_page_number_format(section, style)
        first_page_snapshot = None
        first_page_owned = False
        if first_page_policy in {"hide", "hidden", "no", "skip"}:
            section.different_first_page_header_footer = True
            first_page_footer = section.first_page_footer
            first_page_owned = not first_page_footer.is_linked_to_previous
            if first_page_owned:
                first_page_snapshot = copy.deepcopy(first_page_footer._element)

        _apply_to_footer(
            section.footer,
            _alignment_for(position, "default"),
            style,
            opts,
            position=position,
            footer_kind="default",
            shape_id=next(shape_ids),
        )
        if has_even_footer:
            _apply_to_footer(
                section.even_page_footer,
                _alignment_for(position, "even"),
                style,
                opts,
                position=position,
                footer_kind="even",
                shape_id=next(shape_ids),
            )

        if first_page_policy in {"hide", "hidden", "no", "skip"}:
            if first_page_owned and first_page_snapshot is not None:
                _replace_footer_content(section.first_page_footer, first_page_snapshot)
                _remove_existing_page_numbers(section.first_page_footer)
        elif first_page_policy in {"show", "same", "display"} or section.different_first_page_header_footer:
            section.different_first_page_header_footer = True
            _apply_to_footer(
                section.first_page_footer,
                _alignment_for(position, "first"),
                style,
                opts,
                position=position,
                footer_kind="first",
                shape_id=next(shape_ids),
            )

    return document


def _first_page_policy(options: Mapping[str, Any]) -> str:
    if "first_page" in options:
        raw = options["first_page"]
    elif "first_page_policy" in options:
        raw = options["first_page_policy"]
    else:
        return "default"
    if isinstance(raw, bool):
        return "show" if raw else "hide"
    normalized = str(raw).strip().lower()
    if normalized in {"true", "yes", "1"}:
        return "show"
    if normalized in {"false", "no", "0"}:
        return "hide"
    return normalized or "default"


def _set_footer_distance(section, options: Mapping[str, Any]) -> None:
    if options.get("offset_from_text_mm") is None:
        return
    try:
        offset_cm = float(options.get("offset_from_text_mm")) / 10
    except (TypeError, ValueError):
        return
    bottom_margin_cm = section.bottom_margin.cm if section.bottom_margin else 3.5
    section.footer_distance = Cm(max(0.3, bottom_margin_cm - offset_cm))


def _replace_footer_content(footer, snapshot) -> None:
    footer.is_linked_to_previous = False
    element = footer._element
    for child in list(element):
        element.remove(child)
    for child in snapshot:
        element.append(copy.deepcopy(child))


def _apply_to_footer(
    footer,
    alignment,
    style: str,
    options: Mapping[str, Any],
    *,
    position: str,
    footer_kind: str,
    shape_id: int,
) -> None:
    _materialize_linked_footer(footer)
    footer.is_linked_to_previous = False
    _remove_existing_page_numbers(footer)
    if style in _DASH_STYLES:
        _append_official_dash_page_number(
            footer,
            alignment,
            options,
            position=position,
            footer_kind=footer_kind,
            shape_id=shape_id,
        )
        return
    paragraph = next(
        (candidate for candidate in footer.paragraphs if _is_reusable_empty_paragraph(candidate)),
        None,
    )
    if paragraph is None:
        _append_canonical_page_number_paragraph(
            footer,
            alignment,
            style,
            options,
            position=position,
            footer_kind=footer_kind,
        )
    else:
        paragraph.clear()
        paragraph.alignment = alignment
        _set_page_number_indent(paragraph, position, footer_kind, options)
        _write_page_number(paragraph, style, options)


def _append_canonical_page_number_paragraph(
    footer,
    alignment,
    style: str,
    options: Mapping[str, Any],
    *,
    position: str,
    footer_kind: str,
) -> None:
    """Append one clean page-number paragraph with canonical formatting."""
    paragraph = footer.add_paragraph()
    paragraph.alignment = alignment
    _set_page_number_indent(paragraph, position, footer_kind, options)
    _write_page_number(paragraph, style, options)


def _next_shape_id(document) -> int:
    highest = 0
    doc_pr_tag = f"{{{_NS['wp']}}}docPr"
    for part in document.part.package.parts:
        element = getattr(part, "element", None)
        if element is None:
            element = getattr(part, "_element", None)
        if element is None:
            continue
        for doc_pr in element.iter(doc_pr_tag):
            try:
                highest = max(highest, int(doc_pr.get("id", "0")))
            except ValueError:
                continue
    return highest + 1


def _append_official_dash_page_number(
    footer,
    alignment,
    options: Mapping[str, Any],
    *,
    position: str,
    footer_kind: str,
    shape_id: int,
) -> None:
    paragraph = next(
        (candidate for candidate in footer.paragraphs if _is_reusable_empty_paragraph(candidate)),
        None,
    )
    if paragraph is None:
        paragraph = footer.add_paragraph()
    else:
        paragraph.clear()
    paragraph.alignment = alignment
    _set_page_number_indent(paragraph, position, footer_kind, options)
    run = paragraph.add_run()
    run._r.append(
        _build_official_page_number_alternate_content(
            alignment,
            options,
            position=position,
            footer_kind=footer_kind,
            shape_id=shape_id,
        )
    )


def _build_official_page_number_alternate_content(
    alignment,
    options: Mapping[str, Any],
    *,
    position: str,
    footer_kind: str,
    shape_id: int,
):
    alternate = etree.Element(_tag("mc", "AlternateContent"), nsmap=_NS)
    choice = etree.SubElement(alternate, _tag("mc", "Choice"), Requires="wps")
    drawing = etree.SubElement(choice, _tag("w", "drawing"))
    anchor = etree.SubElement(
        drawing,
        _tag("wp", "anchor"),
        distT="0",
        distB="0",
        distL="114300",
        distR="114300",
        simplePos="0",
        relativeHeight=str(251658240 + shape_id),
        behindDoc="0",
        locked="0",
        layoutInCell="1",
        allowOverlap="1",
    )
    etree.SubElement(anchor, _tag("wp", "simplePos"), x="0", y="0")
    horizontal = etree.SubElement(anchor, _tag("wp", "positionH"), relativeFrom="margin")
    horizontal_alignment = _shape_horizontal_alignment(position)
    etree.SubElement(horizontal, _tag("wp", "align")).text = horizontal_alignment
    vertical = etree.SubElement(anchor, _tag("wp", "positionV"), relativeFrom="paragraph")
    etree.SubElement(vertical, _tag("wp", "posOffset")).text = "0"
    etree.SubElement(anchor, _tag("wp", "extent"), cx="1828800", cy="1828800")
    etree.SubElement(anchor, _tag("wp", "effectExtent"), l="0", t="0", r="0", b="0")
    etree.SubElement(anchor, _tag("wp", "wrapNone"))
    etree.SubElement(anchor, _tag("wp", "docPr"), id=str(shape_id), name="DocxTool Page Number")
    etree.SubElement(anchor, _tag("wp", "cNvGraphicFramePr"))
    graphic = etree.SubElement(anchor, _tag("a", "graphic"))
    graphic_data = etree.SubElement(
        graphic,
        _tag("a", "graphicData"),
        uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    )
    shape = etree.SubElement(graphic_data, _tag("wps", "wsp"))
    etree.SubElement(shape, _tag("wps", "cNvSpPr"), txBox="1")
    shape_properties = etree.SubElement(shape, _tag("wps", "spPr"))
    transform = etree.SubElement(shape_properties, _tag("a", "xfrm"))
    etree.SubElement(transform, _tag("a", "off"), x="0", y="0")
    etree.SubElement(transform, _tag("a", "ext"), cx="1828800", cy="1828800")
    geometry = etree.SubElement(shape_properties, _tag("a", "prstGeom"), prst="rect")
    etree.SubElement(geometry, _tag("a", "avLst"))
    etree.SubElement(shape_properties, _tag("a", "noFill"))
    line = etree.SubElement(shape_properties, _tag("a", "ln"), w="6350")
    etree.SubElement(line, _tag("a", "noFill"))
    textbox = etree.SubElement(shape, _tag("wps", "txbx"))
    textbox.append(
        _build_page_number_textbox_content(
            alignment, options, position=position, footer_kind=footer_kind
        )
    )
    body = etree.SubElement(
        shape,
        _tag("wps", "bodyPr"),
        rot="0",
        spcFirstLastPara="0",
        vertOverflow="overflow",
        horzOverflow="overflow",
        vert="horz",
        wrap="none",
        lIns="0",
        tIns="0",
        rIns="0",
        bIns="0",
        numCol="1",
        spcCol="0",
        rtlCol="0",
        fromWordArt="0",
        anchor="t",
        anchorCtr="0",
        forceAA="0",
        upright="0",
        compatLnSpc="1",
    )
    etree.SubElement(body, _tag("a", "spAutoFit"))

    fallback = etree.SubElement(alternate, _tag("mc", "Fallback"))
    pict = etree.SubElement(fallback, _tag("w", "pict"))
    vml_shape = etree.SubElement(
        pict,
        _tag("v", "shape"),
        id=f"_x0000_s{1024 + shape_id}",
        style=(
            "position:absolute;left:0pt;margin-top:0pt;height:144pt;width:144pt;"
            f"mso-position-horizontal:{horizontal_alignment};"
            "mso-position-horizontal-relative:margin;"
            "mso-wrap-style:none;"
        ),
        filled="f",
        stroked="f",
        coordsize="21600,21600",
    )
    vml_textbox = etree.SubElement(
        vml_shape,
        _tag("v", "textbox"),
        inset="0mm,0mm,0mm,0mm",
        style="mso-fit-shape-to-text:t;",
    )
    vml_textbox.append(
        _build_page_number_textbox_content(
            alignment, options, position=position, footer_kind=footer_kind
        )
    )
    return alternate


def _build_page_number_textbox_content(
    alignment,
    options: Mapping[str, Any],
    *,
    position: str,
    footer_kind: str,
):
    content = etree.Element(_tag("w", "txbxContent"))
    paragraph = etree.SubElement(content, _tag("w", "p"))
    _set_element_paragraph_format(paragraph, alignment, position, footer_kind)
    paragraph.append(_build_styled_text_run("— ", options))
    paragraph.extend(_build_styled_field_runs("PAGE", options))
    paragraph.append(_build_styled_text_run(" —", options))
    return content


def _set_element_paragraph_format(paragraph, alignment, position: str, footer_kind: str) -> None:
    properties = etree.SubElement(paragraph, _tag("w", "pPr"))
    etree.SubElement(properties, _tag("w", "jc"), {qn("w:val"): _alignment_value(alignment)})
    indent = etree.SubElement(
        properties,
        _tag("w", "ind"),
        {qn("w:firstLine"): "0", qn("w:firstLineChars"): "0"},
    )
    if position == "outside":
        indent.set(qn("w:leftChars" if footer_kind == "even" else "w:rightChars"), "100")


def _build_styled_text_run(text: str, options: Mapping[str, Any]):
    run = _build_styled_run(options)
    text_element = etree.SubElement(run, _tag("w", "t"))
    text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    return run


def _build_styled_field_runs(instruction: str, options: Mapping[str, Any]) -> list:
    begin = _build_styled_run(options)
    etree.SubElement(begin, _tag("w", "fldChar"), {qn("w:fldCharType"): "begin", qn("w:dirty"): "true"})
    instr = _build_styled_run(options)
    instr_text = etree.SubElement(instr, _tag("w", "instrText"))
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = f" {instruction} "
    separate = _build_styled_run(options)
    etree.SubElement(separate, _tag("w", "fldChar"), {qn("w:fldCharType"): "separate"})
    end = _build_styled_run(options)
    etree.SubElement(end, _tag("w", "fldChar"), {qn("w:fldCharType"): "end"})
    return [begin, instr, separate, end]


def _build_styled_run(options: Mapping[str, Any]):
    font_name = str(options.get("font_name") or options.get("font") or "宋体")
    font_size_pt = float(options.get("font_size_pt", 14))
    run = etree.Element(_tag("w", "r"))
    properties = etree.SubElement(run, _tag("w", "rPr"))
    fonts = etree.SubElement(properties, _tag("w", "rFonts"))
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), font_name)
    etree.SubElement(properties, _tag("w", "b"), {qn("w:val"): "1" if bool(options.get("bold", False)) else "0"})
    half_points = str(round(font_size_pt * 2))
    etree.SubElement(properties, _tag("w", "sz"), {qn("w:val"): half_points})
    etree.SubElement(properties, _tag("w", "szCs"), {qn("w:val"): half_points})
    return run


def _tag(prefix: str, local_name: str) -> str:
    return f"{{{_NS[prefix]}}}{local_name}"


def _alignment_value(alignment) -> str:
    if alignment == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    return "right"


def _shape_horizontal_alignment(position: str) -> str:
    normalized = {"centre": "center"}.get(position, position)
    if normalized in {"left", "center", "right", "outside"}:
        return normalized
    return "right"


def _materialize_linked_footer(footer) -> None:
    """Copy the effective inherited footer before breaking its section link."""
    if not footer.is_linked_to_previous:
        return
    source_part = footer.part
    inherited = copy.deepcopy(footer._element)
    footer.is_linked_to_previous = False
    _remap_footer_relationships(inherited, source_part, footer.part)
    _replace_footer_content(footer, inherited)


def _remap_footer_relationships(element, source_part, target_part) -> None:
    """Recreate relationships referenced by copied footer XML in its new part."""
    relationship_attributes = {qn("r:id"), qn("r:embed"), qn("r:link")}
    remapped: dict[str, str] = {}
    for node in element.iter():
        for attribute in relationship_attributes:
            old_rid = node.get(attribute)
            if not old_rid:
                continue
            if old_rid not in source_part.rels:
                raise ValueError(f"FOOTER_RELATIONSHIP_NOT_FOUND: {old_rid}")
            if old_rid not in remapped:
                relationship = source_part.rels[old_rid]
                if relationship.is_external:
                    new_rid = target_part.relate_to(
                        relationship.target_ref,
                        relationship.reltype,
                        is_external=True,
                    )
                else:
                    new_rid = target_part.relate_to(
                        relationship.target_part,
                        relationship.reltype,
                    )
                remapped[old_rid] = new_rid
            node.set(attribute, remapped[old_rid])


def _set_page_number_indent(
    paragraph, position: str, footer_kind: str, options: Mapping[str, Any]
) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    indent = ppr.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        ppr.append(indent)
    for name in (
        "w:left",
        "w:right",
        "w:leftChars",
        "w:rightChars",
        "w:firstLine",
        "w:firstLineChars",
        "w:hanging",
        "w:hangingChars",
    ):
        indent.attrib.pop(qn(name), None)

    if position != "outside":
        if not indent.attrib:
            ppr.remove(indent)
        return

    indent.set(qn("w:firstLine"), "0")
    indent.set(qn("w:firstLineChars"), "0")
    if footer_kind == "even":
        indent.set(qn("w:leftChars"), "100")
    else:
        indent.set(qn("w:rightChars"), "100")


def _is_reusable_empty_paragraph(paragraph) -> bool:
    element = paragraph._element
    if paragraph.text.strip():
        return False
    if element.find(qn("w:r")) is not None:
        return False
    for polluted_format in ("w:pBdr", "w:tabs", "w:shd"):
        if element.find(".//" + qn(polluted_format)) is not None:
            return False
    protected_content = (
        "w:fldSimple",
        "w:instrText",
        "w:drawing",
        "w:pict",
        "w:object",
        "w:hyperlink",
    )
    return not any(element.find(".//" + qn(tag)) is not None for tag in protected_content)


def _write_page_number(paragraph, style: str, options: Mapping[str, Any]) -> None:
    if style in {"dash", "numberindash", "number_in_dash", "plain", "number", "page"}:
        _add_field(paragraph, "PAGE", options)
        return
    if style in {"cn", "chinese", "第page页"}:
        _add_text(paragraph, "第 ", options)
        _add_field(paragraph, "PAGE", options)
        _add_text(paragraph, " 页", options)
        return
    if style in {"cn_total", "chinese_total", "page_numpages", "第page页共numpages页"}:
        _add_text(paragraph, "第 ", options)
        _add_field(paragraph, "PAGE", options)
        _add_text(paragraph, " 页 共 ", options)
        _add_field(paragraph, "NUMPAGES", options)
        _add_text(paragraph, " 页", options)
        return
    _add_field(paragraph, "PAGE", options)


def _add_text(paragraph, text: str, options: Mapping[str, Any]):
    run = paragraph.add_run(text)
    _style_run(run, options)
    return run


def _add_field(paragraph, instruction: str, options: Mapping[str, Any]) -> None:
    begin = paragraph.add_run()
    _style_run(begin, options)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_begin.set(qn("w:dirty"), "true")
    begin._r.append(field_begin)

    instr = paragraph.add_run()
    _style_run(instr, options)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    _style_run(separate, options)
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    separate._r.append(field_separate)

    end = paragraph.add_run()
    _style_run(end, options)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    end._r.append(field_end)


def _style_run(run, options: Mapping[str, Any]) -> None:
    font_name = str(options.get("font_name") or options.get("font") or "宋体")
    font_size_pt = float(options.get("font_size_pt", 14))
    bold = bool(options.get("bold", False))
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.insert(0, r_fonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(attr), font_name)
    run.font.size = Pt(font_size_pt)
    half_points = str(round(font_size_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        element = rpr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            rpr.append(element)
        element.set(qn("w:val"), half_points)
    run.font.bold = bold


def _alignment_for(position: str, footer_kind: str):
    normalized = {"right": "right", "left": "left", "center": "center", "centre": "center"}.get(position, position)
    if normalized == "outside":
        return WD_ALIGN_PARAGRAPH.LEFT if footer_kind == "even" else WD_ALIGN_PARAGRAPH.RIGHT
    if normalized == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if normalized == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.RIGHT


def _set_even_and_odd_headers(document, enabled: bool) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:evenAndOddHeaders"))
    if not enabled:
        if existing is not None:
            settings.remove(existing)
        return
    if existing is None:
        existing = OxmlElement("w:evenAndOddHeaders")
        settings.append(existing)
    existing.set(qn("w:val"), "1")


def _even_and_odd_headers_enabled(document) -> bool:
    existing = document.settings._element.find(qn("w:evenAndOddHeaders"))
    if existing is None:
        return False
    return existing.get(qn("w:val"), "1") not in {"0", "false", "False"}


def _set_update_fields_on_open(document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def _set_section_numbering(section, mode: str, restart_at: int, options: Mapping[str, Any], section_index: int) -> None:
    starts = options.get("section_starts")
    if isinstance(starts, Sequence) and not isinstance(starts, (str, bytes)):
        if section_index < len(starts) and starts[section_index] is not None:
            _set_pg_num_start(section, int(starts[section_index]))
        else:
            _clear_pg_num_start(section)
        return
    if mode in {"restart", "restart_each_section", "new"}:
        _set_pg_num_start(section, restart_at)
    else:
        _clear_pg_num_start(section)


def _set_page_number_format(section, style: str) -> None:
    """Set the native PAGE number format independently from restart state."""
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), "decimal")


def _set_pg_num_start(section, start: int) -> None:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def _clear_pg_num_start(section) -> None:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        return
    pg_num_type.attrib.pop(qn("w:start"), None)
    if not pg_num_type.attrib and len(pg_num_type) == 0:
        section._sectPr.remove(pg_num_type)


def _remove_existing_page_numbers(footer) -> None:
    _remove_page_number_shapes(footer._element)
    for paragraph in list(footer.paragraphs):
        if _is_page_number_only_paragraph(paragraph):
            paragraph._element.getparent().remove(paragraph._element)
            continue
        if not _paragraph_has_page_number(paragraph):
            continue
        _remove_page_field_runs(paragraph)
        _remove_literal_page_tokens(paragraph)
        if _is_structurally_empty_after_page_cleanup(paragraph._element):
            paragraph._element.getparent().remove(paragraph._element)


def _remove_page_number_shapes(footer_element) -> None:
    for alternate in list(footer_element.iter(_tag("mc", "AlternateContent"))):
        if not _is_page_number_shape(alternate):
            continue
        paragraph = _ancestor_with_tag(alternate, qn("w:p"))
        parent = alternate.getparent()
        parent.remove(alternate)
        if parent.tag == qn("w:r") and _is_empty_run(parent):
            run_parent = parent.getparent()
            run_parent.remove(parent)
        if (
            paragraph is not None
            and paragraph.getparent() is not None
            and _is_structurally_empty_after_page_cleanup(paragraph)
        ):
            paragraph.getparent().remove(paragraph)


def _ancestor_with_tag(element, tag):
    current = element.getparent()
    while current is not None:
        if current.tag == tag:
            return current
        current = current.getparent()
    return None


def _is_empty_run(run) -> bool:
    for child in run:
        if child.tag == qn("w:rPr"):
            continue
        if child.tag == qn("w:t") and not (child.text or "").strip():
            continue
        return False
    return True


def _is_structurally_empty_after_page_cleanup(paragraph) -> bool:
    """Return True only for a touched page-number paragraph with no business content."""
    if any((text.text or "").strip() for text in paragraph.iter(qn("w:t"))):
        return False
    if any((text.text or "").strip() for text in paragraph.iter(qn("w:instrText"))):
        return False
    protected = (
        qn("w:fldSimple"),
        qn("w:drawing"),
        qn("w:pict"),
        qn("w:object"),
        qn("w:hyperlink"),
        qn("w:txbxContent"),
    )
    if any(any(True for _ in paragraph.iter(tag)) for tag in protected):
        return False
    allowed_top_level = {
        qn("w:pPr"),
        qn("w:r"),
        qn("w:bookmarkStart"),
        qn("w:bookmarkEnd"),
        qn("w:proofErr"),
    }
    for child in paragraph:
        if child.tag not in allowed_top_level:
            return False
        if child.tag == qn("w:r") and not _is_empty_run(child):
            return False
    return True


def _is_page_number_shape(alternate) -> bool:
    contents = list(alternate.iter(qn("w:txbxContent")))
    if not contents:
        return False
    for content in contents:
        instructions = [
            "".join(node.itertext())
            for node in content.iter(qn("w:instrText"))
        ]
        if not any(_PAGE_INSTRUCTION_RE.search(item) for item in instructions):
            return False
        if any(
            item.strip() and not _PAGE_INSTRUCTION_RE.search(item)
            for item in instructions
        ):
            return False
        visible = "".join(node.text or "" for node in content.iter(qn("w:t")))
        if re.sub(r"[\s\d—–-]", "", visible):
            return False
    return True


def _paragraph_has_page_number(paragraph) -> bool:
    return any(
        _PAGE_INSTRUCTION_RE.search(instruction)
        for instruction in _iter_field_instructions(paragraph)
    )


def _iter_field_instructions(paragraph) -> Iterator[str]:
    """Yield complex and simple Word field instructions from one paragraph."""
    for _start, _end, instruction in _iter_complex_fields(paragraph):
        yield instruction
    for field in paragraph._element.findall(qn("w:fldSimple")):
        yield field.get(qn("w:instr"), "")


def _iter_complex_fields(paragraph) -> Iterator[tuple[int, int, str]]:
    """Yield direct-run bounds and joined instructions for complex fields."""
    runs = list(paragraph._element.findall(qn("w:r")))
    field_start: int | None = None
    instruction_parts: list[str] = []
    for index, run in enumerate(runs):
        if field_start is None and _has_field_char(run, "begin"):
            field_start = index
            instruction_parts = []
        if field_start is None:
            continue
        instruction_parts.extend(
            element.text or ""
            for element in run.findall(".//" + qn("w:instrText"))
        )
        if _has_field_char(run, "end"):
            yield field_start, index, "".join(instruction_parts)
            field_start = None
            instruction_parts = []


def _is_page_number_only_paragraph(paragraph) -> bool:
    instructions = tuple(_iter_field_instructions(paragraph))
    if any(
        instruction.strip() and not _PAGE_INSTRUCTION_RE.search(instruction)
        for instruction in instructions
    ):
        return False
    visible_text = "".join(text.text or "" for text in paragraph._element.findall(".//" + qn("w:t")))
    normalized = re.sub(r"\s+", "", visible_text)
    if not normalized:
        return any(_PAGE_INSTRUCTION_RE.search(item) for item in instructions)
    if _STATIC_PAGE_NUMBER_RE.fullmatch(normalized):
        return True
    if _paragraph_has_page_number(paragraph):
        return not _PAGE_ONLY_TEXT_RE.sub("", visible_text).strip()
    return False


def _remove_page_field_runs(paragraph) -> None:
    runs = list(paragraph._element.findall(qn("w:r")))
    remove_indexes: set[int] = set()
    for start, end, instruction in _iter_complex_fields(paragraph):
        if not _PAGE_INSTRUCTION_RE.search(instruction):
            continue
        if start > 0:
            _trim_attached_dash_decoration(runs[start - 1], trailing=True)
        if end < len(runs) - 1:
            _trim_attached_dash_decoration(runs[end + 1], trailing=False)
        while start > 0 and _is_page_decoration_run(runs[start - 1]):
            start -= 1
        while end < len(runs) - 1 and _is_page_decoration_run(runs[end + 1]):
            end += 1
        remove_indexes.update(range(start, end + 1))
    for element in paragraph._element.findall(qn("w:fldSimple")):
        instruction = element.get(qn("w:instr"), "")
        if _PAGE_INSTRUCTION_RE.search(instruction):
            paragraph._element.remove(element)
    for index in sorted(remove_indexes, reverse=True):
        paragraph._element.remove(runs[index])


def _trim_attached_dash_decoration(run, *, trailing: bool) -> None:
    """Remove dash decoration attached to business text beside a PAGE field."""
    text_elements = run.findall(".//" + qn("w:t"))
    if not text_elements:
        return
    target = text_elements[-1] if trailing else text_elements[0]
    value = target.text or ""
    if trailing:
        target.text = re.sub(r"(?:\s*[—–-]\s*)+$", "", value)
    else:
        target.text = re.sub(r"^(?:\s*[—–-]\s*)+", "", value)


def _has_field_char(run, field_char_type: str) -> bool:
    return any(
        field_char.get(qn("w:fldCharType")) == field_char_type
        for field_char in run.findall(".//" + qn("w:fldChar"))
    )


def _is_page_decoration_run(run) -> bool:
    if run.find(".//" + qn("w:drawing")) is not None or run.find(".//" + qn("w:pict")) is not None:
        return False
    text_elements = run.findall(".//" + qn("w:t"))
    text = "".join(text_element.text or "" for text_element in text_elements)
    if not text:
        return bool(text_elements)
    return not _PAGE_DECORATION_TEXT_RE.sub("", text).strip()


def _remove_literal_page_tokens(paragraph) -> None:
    for run in list(paragraph.runs):
        if _PAGE_INSTRUCTION_RE.search(run.text or ""):
            run._element.getparent().remove(run._element)
