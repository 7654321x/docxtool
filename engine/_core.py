"""engine — 排版引擎。

职责：纯排版执行器，不做段落识别业务。
  - 输入 DocumentData + StyleRules → 输出 .docx
  - apply_style 为纯执行器，不处理 meta 判断
  - 所有 meta 映射在 export_doc 主循环中完成
"""

import copy
import io
import logging
import re
from typing import Dict, List, Optional, Tuple, Union

from style_config import (
    StyleRule, PageSettings, cn_size_to_pt, chinese_number,
    arabic_number, parse_indent, parse_alignment,
    logger, ExportError, StyleError,
)
import math

from importer import DocumentData, ParagraphData
from engine.normal import resolve as _resolve_rule

# ── python-docx 模块级导入 ──
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── XML 安全写入（防重复节点）──
def _set_unique(pPr, tag, element):
    """替换 pPr 中已存在的同名标签，避免重复 append 导致 Word 行为不一致。"""
    old = pPr.find(tag)
    if old is not None:
        pPr.remove(old)
    pPr.append(element)


def _copy_run_style(src_run, dst_run) -> None:
    """Copy the resolved run properties so rewritten inline segments keep the paragraph font."""
    src_rpr = getattr(src_run._element, "rPr", None)
    if src_rpr is None:
        return
    dst_rpr = getattr(dst_run._element, "rPr", None)
    if dst_rpr is not None:
        dst_run._element.remove(dst_rpr)
    dst_run._element.insert(0, copy.deepcopy(src_rpr))


def _segment_writer(para):
    """Return a writer that replaces all existing run text without losing base styling."""
    if not para.runs:
        para.add_run("")
    base_run = para.runs[0]
    for run in para.runs:
        run.text = ""
    used_base = False

    def write(text: str, *, bold=None, cn_font=None, size_pt=None):
        nonlocal used_base
        if not text:
            return None
        run = base_run if not used_base else para.add_run("")
        if used_base:
            _copy_run_style(base_run, run)
        used_base = True
        run.text = text
        if cn_font:
            _set_run_fonts(run, cn_font=cn_font, en_font="Times New Roman")
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        return run

    return write


def _line_spacing_twips(settings: PageSettings) -> int:
    """PageSettings 行距 pt → Word twips。"""
    try:
        value = float(settings.line_spacing_value)
    except (TypeError, ValueError):
        value = 28.0
    if value <= 0:
        value = 28.0
    return int(round(value * 20))


# ── 落款排版辅助 ──
def _apply_right_indent(para, n=2):
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None: ind = OxmlElement('w:ind'); pPr.append(ind)
    ind.set(qn('w:right'), str(int(n * 560)))
    ind.set(qn('w:rightChars'), str(n * 100))

def _apply_first_line_indent_chars(para, chars: int):
    """只设置首行缩进，不设置悬挂缩进。"""
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    for attr in (qn('w:firstLine'), qn('w:firstLineChars'),
                 qn('w:left'), qn('w:leftChars'),
                 qn('w:hanging'), qn('w:hangingChars')):
        if attr in ind.attrib:
            del ind.attrib[attr]
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    ind.set(qn('w:firstLine'), str(chars * 320))


def _apply_left_indent_chars(para, chars: float):
    """只设置整段左缩进，不设置首行或悬挂缩进。"""
    pPr = para._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    for attr in (qn('w:firstLine'), qn('w:firstLineChars'),
                 qn('w:left'), qn('w:leftChars'),
                 qn('w:hanging'), qn('w:hangingChars')):
        if attr in ind.attrib:
            del ind.attrib[attr]
    ind.set(qn('w:leftChars'), str(int(chars * 100)))
    ind.set(qn('w:left'), str(int(chars * 320)))

def _attachment_note_wrap_start_chars(text: str) -> int:
    """附件说明首行回行列。

    无编号：`附件：正文` → 回行对齐到正文首字（5 字列）
    有编号：`附件：1. 正文` → 回行对齐到编号后的正文首字
    """
    m = re.match(r'^\s*附件\s*[:：]\s*(\d+)[.．、]\s*', text or "")
    if not m:
        return 5
    return 2 + 3 + len(m.group(1)) + 2

def _attachment_item_wrap_start_chars(text: str) -> int:
    """附件续项回行列：对齐到编号后的正文首字。"""
    m = re.match(r'^\s*(\d+)[.．、]\s*', text or "")
    if not m:
        return 8
    return 5 + len(m.group(1)) + 2

# ── 清理旧编号 ──
_LEADING_NUM_RE = re.compile(
    r'^\s*(?:[（(]?[一二三四五六七八九十百千零〇0-9]+[)）]?[、\.．]\s*)+'
)

def _strip_heading_numbering(text: str) -> str:
    """去掉段首已有的编号，避免再次插入时重复。"""
    return _LEADING_NUM_RE.sub("", text, count=1).lstrip()


def _apply_special_bold(para, text: str) -> None:
    """特殊加粗入口：按匹配拆 run，每句独立加粗。"""
    if not para.runs:
        return
    from style_config import NB_SUFFIXES as _NBS, NB_FIXED as _NBF
    import re as _re
    _suffixes = "|".join(_NBS)
    _fixed = "|".join(map(_re.escape, _NBF))
    XSHI = rf'(?:[一二三四五六七八九十]+(?:{_suffixes})|{_fixed})'
    parts = _re.split(f'(?={XSHI})', text)
    write = _segment_writer(para)
    for pi, part in enumerate(parts):
        if not part:
            continue
        if pi == 0:
            write(part, bold=False)
        elif _NBF and any(part.startswith(f) for f in _NBF):
            colon = part.find('：')
            if colon < 0:
                colon = part.find(':')
            if colon >= 0:
                write(part[:colon + 1], bold=True)
                write(part[colon + 1:], bold=False)
            else:
                write(part, bold=True)
        else:
            m = _re.match(rf'([一二三四五六七八九十]+(?:{_suffixes}).*?。)(.*)', part)
            if m:
                write(m.group(1), bold=True)
                write(m.group(2), bold=False)
            else:
                write(part, bold=True)

def _apply_fixed_bold(para, part: str) -> None:
    """固定词组：加粗到冒号止（如'比如：xxx' → 加粗'比如：'）。"""
    colon = part.find('：')
    if colon < 0:
        colon = part.find(':')
    if colon >= 0:
        br = para.add_run(part[:colon+1])
        br.font.bold = True
        if part[colon+1:]:
            nr = para.add_run(part[colon+1:])
            nr.font.bold = False
    else:
        r = para.add_run(part)
        r.font.bold = True

def _apply_numbered_bold(para, part: str) -> None:
    """X是/X要：加粗到句号止（如'一是…。xxx' → 加粗'一是…。'）。"""
    from style_config import NB_SUFFIXES
    import re as _re
    _suffixes = "|".join(NB_SUFFIXES)
    XSHI = rf'[一二三四五六七八九十]+(?:{_suffixes})'
    m = _re.match(f'({XSHI}.*?。)(.*)', part)
    if m:
        br = para.add_run(m.group(1))
        br.font.bold = True
        if m.group(2):
            nr = para.add_run(m.group(2))
            nr.font.bold = False
    else:
        r = para.add_run(part)
        r.font.bold = True

def _apply_colon_bold(para, text: str) -> None:
    """冒号关键词加粗：从行首到第一个冒号（含）加粗，后面正常。"""
    if not para.runs:
        return
    for colon in ('：', ':'):
        pos = text.find(colon)
        if pos > 0 and pos <= 10:
            para.runs[0].text = ""
            br = para.add_run(text[:pos + 1])
            br.font.bold = True
            if text[pos + 1:]:
                nr = para.add_run(text[pos + 1:])
                nr.font.bold = False
            return


def _apply_heading1_report_split(para, text: str, rule) -> None:
    """heading1_report 句号后内容换行：标题部分黑体，句号后另起一段 body。"""
    period = text.find('。')
    if period <= 0 or period >= len(text) - 1:
        return
    heading_text = text[:period + 1]  # 含句号
    body_text = text[period + 1:].lstrip()
    # 截断当前段落为标题部分
    for run in para.runs:
        run.text = ""
    para.runs[0].text = heading_text
    # 在当前段落后插入 body 段落
    body_para = OxmlElement('w:p')
    body_run = OxmlElement('w:r')
    body_t = OxmlElement('w:t')
    body_t.set(qn('xml:space'), 'preserve')
    body_t.text = body_text
    body_run.append(body_t)
    body_para.append(body_run)
    para._element.addnext(body_para)


def _apply_glossary_item(para, text: str, rule) -> None:
    """名词解释条目：编号不加粗 + 关键词（冒号前）黑体，正文（冒号后）仿宋。"""
    if len(para.runs) < 2:
        return
    cp = -1
    for c in ('：', ':'):
        cp = text.find(c)
        if cp > 0:
            break
    if cp <= 0:
        return
    kw = text[:cp + 1]   # 关键词（含冒号）
    bd = text[cp + 1:]   # 正文
    # 保留编号 run（runs[0]），清除内容 run（runs[-1]）
    para.runs[-1].text = ""
    # 关键词 → 黑体
    kr = para.add_run(kw)
    kr.font.name = "黑体"
    _set_run_fonts(kr, cn_font="黑体", en_font="Times New Roman")
    kr.font.size = Pt(rule.font_size_pt)
    kr.font.bold = False
    # 正文 → 仿宋
    if bd:
        br = para.add_run(bd)
        br.font.name = "仿宋_GB2312"
        _set_run_fonts(br, cn_font="仿宋_GB2312", en_font="Times New Roman")
        br.font.size = Pt(rule.font_size_pt)
        br.font.bold = False


def _apply_report_first_sentence(para, text: str, rule) -> None:
    """报告首句加粗：首句（到第一个 。）楷体加粗，剩余仿宋正文。"""
    if not para.runs:
        return
    period = text.find('。')
    if period <= 0:
        return
    # 首句（含句号）→ 楷体加粗
    first = text[:period + 1]
    rest = text[period + 1:]
    write = _segment_writer(para)
    write(first, bold=True, cn_font="楷体_GB2312", size_pt=rule.font_size_pt)
    # 剩余 → 仿宋
    if rest:
        write(rest, bold=False, cn_font=rule.font, size_pt=rule.font_size_pt)


def _set_para_spacing(para, before_lines: float = 0, after_lines: float = 0,
                      line_twips: int = 560) -> None:
    """设置段前段后间距（单位：行）。"""
    pPr = para._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    if before_lines > 0:
        spacing.set(qn('w:before'), str(int(round(before_lines * line_twips))))
        spacing.set(qn('w:beforeLines'), str(int(round(before_lines * 100))))
    if after_lines > 0:
        spacing.set(qn('w:after'), str(int(round(after_lines * line_twips))))
        spacing.set(qn('w:afterLines'), str(int(round(after_lines * 100))))
    if line_twips > 0:
        spacing.set(qn('w:line'), str(line_twips))
        spacing.set(qn('w:lineRule'), 'exact')
    _set_unique(pPr, qn('w:spacing'), spacing)


def _apply_rule_paragraph_format(para, rule: StyleRule, line_twips: int) -> None:
    """应用 JSON 中的段落级扩展配置。"""
    _set_para_spacing(
        para,
        before_lines=getattr(rule, "spacing_before", 0.0) or 0.0,
        after_lines=getattr(rule, "spacing_after", 0.0) or 0.0,
        line_twips=line_twips,
    )
    left_indent = getattr(rule, "left_indent", 0.0) or 0.0
    right_indent = getattr(rule, "right_indent", 0.0) or 0.0
    if left_indent > 0:
        _apply_left_indent_chars(para, left_indent)
    if right_indent > 0:
        _apply_right_indent(para, right_indent)
    if getattr(rule, "page_break_before", False):
        pPr = para._element.get_or_add_pPr()
        pb = OxmlElement('w:pageBreakBefore')
        _set_unique(pPr, qn('w:pageBreakBefore'), pb)


def _handle_heading_period(text: str) -> str:
    """处理标题句号（heading2/heading3）。

    类型 A（独立）："（一）坚持党的领导。" → 去掉句号
    类型 B（行内）："1.深入学习…统一行动。我们举办…" → 保留整段
                      后续在 run 层拆分为标题 run + 正文 run（仿宋）
    """
    # 找第一个中文句号
    period_pos = text.find("。")
    if period_pos < 0:
        return text  # 无句号，类型 A，直接返回

    after_period = text[period_pos + 1:].strip()
    # 句号后 ≥5 字 → 类型 B（行内标题），保留整段
    if len(after_period) >= 5:
        return text
    # 句号后无内容或很短 → 类型 A（独立标题），去掉句号
    return text[:period_pos] + text[period_pos + 1:]


# ═══════════════════════════════════════════════════════════════
# 四级层级计数器（分级控制）
# ═══════════════════════════════════════════════════════════════

class NumberingCounter:
    """四级层级计数器：高级递增 → 低级清零。

    heading1 → a+=1, b=0, c=0, d=0
    heading2 → b+=1, c=0, d=0
    heading3 → c+=1, d=0
    heading4 → d+=1
    """
    a: int = 0
    b: int = 0
    c: int = 0
    d: int = 0

    def advance(self, type_id: str) -> None:
        if type_id == "heading1":
            self.a += 1
            self.b = 0; self.c = 0; self.d = 0
        elif type_id == "heading2":
            self.b += 1
            self.c = 0; self.d = 0
        elif type_id == "heading3":
            self.c += 1
            self.d = 0
        elif type_id == "heading4":
            self.d += 1

    def render(self, pattern: str, type_id: str) -> str:
        """精确替换模板中的 {a}/{b}/{c}/{d}。只替换花括号变量，不动字面字母。"""
        if not pattern:
            return ""
        is_chinese = type_id in ("heading1", "heading2")
        num_fn = chinese_number if is_chinese else arabic_number
        result = pattern
        result = result.replace("{a}", num_fn(self.a))
        result = result.replace("{b}", num_fn(self.b))
        result = result.replace("{c}", arabic_number(self.c))
        result = result.replace("{d}", arabic_number(self.d))
        return result


# ═══════════════════════════════════════════════════════════════
# 中英分离字体设置
# ═══════════════════════════════════════════════════════════════



def _set_run_fonts(run, cn_font="宋体", en_font="Times New Roman"):
    """中英混排字体：中文=ea，英文=ascii/hAnsi。"""
    run.font.name = en_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)


# ═══════════════════════════════════════════════════════════════
# 纯样式执行器
# ═══════════════════════════════════════════════════════════════

def apply_style(paragraph, rule: StyleRule) -> None:
    """在 Run 级别设置字体。纯执行器，不处理 meta，不捕获异常。"""



    # 确保段落至少有一个 Run
    if not paragraph.runs:
        paragraph.add_run("")

    # 遍历所有 Run，中英分离字体
    for run in paragraph.runs:
        _set_run_fonts(run, cn_font=rule.font, en_font="Times New Roman")
        run.font.size = Pt(rule.font_size_pt)
        if rule.bold is not None:
            run.font.bold = rule.bold

    # 对齐
    align_id = _align_to_enum(rule.alignment, WD_ALIGN_PARAGRAPH)
    if align_id is not None:
        paragraph.alignment = align_id

    # 正文：启用网格对齐（row_index≥5 为正文/附件）
    if rule.row_index >= 5:
        pPr = paragraph._element.get_or_add_pPr()
        snap = OxmlElement('w:snapToGrid')
        snap.set(qn('w:val'), '1')
        _set_unique(pPr, qn('w:snapToGrid'), snap)

    # 标题：禁用网格对齐 + 孤行控制（row_index≤4 为标题）
    if rule.row_index < 5:
        pPr = paragraph._element.get_or_add_pPr()
        snap = OxmlElement('w:snapToGrid')
        snap.set(qn('w:val'), '0')
        _set_unique(pPr, qn('w:snapToGrid'), snap)
        widowControl = OxmlElement('w:widowControl')
        widowControl.set(qn('w:val'), '0')
        _set_unique(pPr, qn('w:widowControl'), widowControl)
        # 关闭字符间距（标题自然排列，不受网格拉宽）
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:val'), '0')
            rPr.append(sp)

    # 首行缩进：字符单位 w:firstLineChars（2字符 = 200）
    if rule.first_line_indent > 0:
        pPr = paragraph._element.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), str(int(rule.first_line_indent * 100)))
        # 清除 twip 单位避免冲突
        if qn('w:firstLine') in ind.attrib:
            del ind.attrib[qn('w:firstLine')]


def _align_to_enum(align_str: str, WD_ALIGN_PARAGRAPH) -> Optional[int]:
    """中文对齐字符串 → python-docx 对齐枚举。"""
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    # 传入的可能是中文，也有可能是英文
    zh_to_en = {
        "左对齐": "left", "居中": "center", "右对齐": "right",
        "两端对齐": "justify", "分散对齐": "distribute",
    }
    key = zh_to_en.get(align_str, align_str)
    return mapping.get(key)


def apply_style_safe(paragraph, rule: StyleRule) -> bool:
    """带降级的 apply_style 包装。"""
    try:
        apply_style(paragraph, rule)
        return True
    except Exception as e:
        logger.warning(f"[引擎] 字体 '{rule.font}' 应用失败: {e}，回退宋体")
        try:
            fallback = copy.copy(rule)
            fallback.font = "宋体"
            fallback.font_size_pt = 12
            fallback.bold = False
            apply_style(paragraph, fallback)
        except Exception as e2:
            logger.error(f"[引擎] 完全降级失败: {e2}")
        return False


# ═══════════════════════════════════════════════════════════════
# 编号
# ═══════════════════════════════════════════════════════════════

def apply_numbering(paragraph, rule: StyleRule, counter: NumberingCounter) -> None:
    """在段落前插入纯文本编号。"""
    numbering = counter.render(rule.numbering_pattern, f"heading{rule.row_index + 1}" if rule.row_index < 4 else "body")
    if numbering:
        logger.debug(f"[编号] {rule.level_name} → \"{numbering}\" (a={counter.a} b={counter.b} c={counter.c} d={counter.d})")
    if not numbering:
        return
    existing = paragraph.text.strip()
    if existing.startswith(numbering):
        return
    # XML 层：先创建 run，补齐字体，再移到最前
    if paragraph.runs:
        first_run = paragraph.runs[0]
        new_run = paragraph.add_run(numbering)
        _set_run_fonts(new_run, cn_font=rule.font, en_font="Times New Roman")
        new_run.font.size = Pt(rule.font_size_pt)
        if rule.bold is not None:
            new_run.font.bold = rule.bold
        first_run._element.addprevious(new_run._element)
    else:
        new_run = paragraph.add_run(numbering)
        _set_run_fonts(new_run, cn_font=rule.font, en_font="Times New Roman")
        new_run.font.size = Pt(rule.font_size_pt)
        if rule.bold is not None:
            new_run.font.bold = rule.bold


# ═══════════════════════════════════════════════════════════════
# 上标拆分
# ═══════════════════════════════════════════════════════════════

def apply_superscript_split(paragraph) -> None:
    """拆分 Run：匹配 [n] 或 〔n〕模式，设为上标并统一为 [n] 格式。"""

    # 匹配 [1] [12] 〔1〕 〔12〕，跳过 4 位数字（如 [2025] 是文号）
    superscript_pattern = re.compile(r'(\[\d{1,2}\]|〔\d{1,2}〕)')
    to_remove = []

    for run in list(paragraph.runs):
        text = run.text
        parts = superscript_pattern.split(text)
        if len(parts) == 1:
            continue

        parent = run._element.getparent()
        insert_before = run._element

        for part in parts:
            if not part:
                continue
            new_run = paragraph.add_run(part)
            if superscript_pattern.match(part):
                # 统一为 [n] 格式
                num = re.search(r'\d+', part).group()
                new_run.text = f'[{num}]'
                new_run.font.superscript = True
                new_run.font.name = "Times New Roman"
                new_run.font.size = Pt(16)  # 三号
            else:
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
            parent.insert(parent.index(insert_before), new_run._element)
            # 上标与前一个 run 保持同行
            if superscript_pattern.match(part):
                prev_elem = new_run._element.getprevious()
                if prev_elem is not None:
                    prev_rPr = prev_elem.find(qn('w:rPr'))
                    if prev_rPr is None:
                        prev_rPr = OxmlElement('w:rPr')
                        prev_elem.insert(0, prev_rPr)
                    keep = OxmlElement('w:keepNext')
                    prev_rPr.append(keep)

        to_remove.append(run._element)

    for elem in to_remove:
        elem.getparent().remove(elem)


# ═══════════════════════════════════════════════════════════════
# 页面设置（底层 XML 注入 w:docGrid）
# ═══════════════════════════════════════════════════════════════

def _write_doc_grid(section, settings: PageSettings, doc_mode: str = "") -> bool:
    """写入 docGrid：强制每行28字符，每页22行，精确计算字符间距。

    关键：charSpace = 列宽 - 字体固有宽度（16pt=320twip），Word 用"字体宽 + charSpace"算列宽。
    返回 True 表示网格已写入，False 表示跳过（chars/lines 为 0）。
    """
    if settings.chars_per_line <= 0 or settings.lines_per_page <= 0:
        return False  # 自然排版模式：不写网格
    if doc_mode == "SCHEME":
        return False  # 方案模式：不写网格
    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:docGrid')):
        sectPr.remove(old)

    # ═══ 字符网格 ═══
    TWIP = 567
    page_w = 21.0 * TWIP
    left   = settings.margin_left_cm * TWIP
    right  = settings.margin_right_cm * TWIP
    cw     = page_w - left - right
    col    = cw / settings.chars_per_line
    csp    = math.floor(col - 320)  # floor(315.9-320) = -5
    line_pitch = _line_spacing_twips(settings)

    logger.info(f"[网格] charSpace={csp}  版心={cw:.0f}twip  期望{settings.chars_per_line}字/行")
    logger.info(f"[网格] XML: charsPerLine={settings.chars_per_line} linesPerPage={settings.lines_per_page} charSpace={csp} linePitch={line_pitch}")

    docGrid = OxmlElement('w:docGrid')
    docGrid.set(qn('w:type'), 'linesAndChars')
    docGrid.set(qn('w:charsPerLine'), str(settings.chars_per_line))
    docGrid.set(qn('w:linesPerPage'), str(settings.lines_per_page))
    docGrid.set(qn('w:charSpace'), str(csp))
    docGrid.set(qn('w:linePitch'), str(line_pitch))
    sectPr.append(docGrid)
    return True


def apply_page_settings(doc, settings: PageSettings, doc_mode: str = "") -> None:
    """设置页面布局：页边距 + docGrid 网格（22行×28字） + 行距。"""

    # ═══ 方案：写 documentDefaults 控制全局默认字体（比 Normal 样式更可靠） ═══
    styles_element = doc.styles._element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = OxmlElement('w:docDefaults')
        styles_element.insert(0, docDefaults)
    # Run 默认属性（字体 + 字号）
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        docDefaults.append(rPrDefault)
    rPrDef = rPrDefault.find(qn('w:rPr'))
    if rPrDef is None:
        rPrDef = OxmlElement('w:rPr')
        rPrDefault.append(rPrDef)
    # 清旧建新
    for old in rPrDef.findall(qn('w:rFonts')):
        rPrDef.remove(old)
    for old in rPrDef.findall(qn('w:sz')):
        rPrDef.remove(old)
    for old in rPrDef.findall(qn('w:szCs')):
        rPrDef.remove(old)
    # 覆盖 w:lang，强制东亚语言 = 中文（否则 CJK 网格不生效）
    for old_lang in rPrDef.findall(qn('w:lang')):
        rPrDef.remove(old_lang)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:eastAsia'), 'zh-CN')
    lang.set(qn('w:bidi'), 'ar-SA')
    rPrDef.append(lang)
    df = OxmlElement('w:rFonts')
    df.set(qn('w:eastAsia'), '仿宋_GB2312')
    df.set(qn('w:ascii'), 'Times New Roman')
    df.set(qn('w:hAnsi'), 'Times New Roman')
    rPrDef.append(df)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '32')  # 16pt = 32 half-pt
    rPrDef.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '32')
    rPrDef.append(szCs)

    # 同时也要设 Normal 样式（兼容旧版 Word）
    style = doc.styles['Normal']
    style.font.size = Pt(16)
    style.font.name = "Times New Roman"
    rPr = style.element.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.insert(0, rFonts)

    for section in doc.sections:
        # 写入网格（字体已就绪）
        _write_doc_grid(section, settings, doc_mode)

        # A4 页面尺寸
        section.page_width = Cm(settings.page_width_cm)
        section.page_height = Cm(settings.page_height_cm)

        # 页边距
        section.top_margin = Cm(settings.margin_top_cm)
        section.bottom_margin = Cm(settings.margin_bottom_cm)
        section.left_margin = Cm(settings.margin_left_cm)
        section.right_margin = Cm(settings.margin_right_cm)

        # 页码距版心下边缘 7mm → 页脚距页底 = 下边距 - 0.7cm = 2.8cm（GB/T 9704-2012 7.5）
        section.header_distance = Cm(0)
        section.footer_distance = Cm(settings.margin_bottom_cm - 0.7)

    logger.info(f"[页面] 边距 上{settings.margin_top_cm} 下{settings.margin_bottom_cm} 左{settings.margin_left_cm} 右{settings.margin_right_cm} cm")

    # 强制激活 Word 网格模式（否则 docGrid 只有数据不生效）
    settings_element = doc.settings._element
    compat = settings_element.find(qn('w:compat'))
    if compat is None:
        compat = OxmlElement('w:compat')
        settings_element.append(compat)
    for name, val in [
        ("compatibilityMode", "15"),
        ("overrideTableStyleFontSizeAndJustification", "1"),
        ("noExtraLineSpacing", "1"),
        ("useFELayout", "1"),              # 触发东亚字符网格排版模式
        ("balanceSingleByteDoubleByteWidth", "1"),  # 标点压缩
        ("doNotExpand", "1"),              # 不扩展字符间距
        ("doNotLeaveBackslashAlone", "1"), # CJK 禁则
    ]:
        el = OxmlElement('w:compatSetting')
        el.set(qn('w:name'), name)
        el.set(qn('w:val'), val)
        compat.append(el)

    # Normal 样式已在 docGrid 写入前设置字号/字体
    from docx.enum.text import WD_LINE_SPACING
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = Pt(settings.line_spacing_value)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    # 禁用 Normal 样式的自动上下文间距
    style_pPr = style.element.get_or_add_pPr()
    for old in style_pPr.findall(qn('w:contextualSpacing')):
        style_pPr.remove(old)
    ctxSpc = OxmlElement('w:contextualSpacing')
    ctxSpc.set(qn('w:val'), '0')
    style_pPr.append(ctxSpc)


# ═══════════════════════════════════════════════════════════════
# 页码（Union[奇右, 偶左]）
# ═══════════════════════════════════════════════════════════════

def apply_header_footer(doc, page_rule: StyleRule) -> None:
    """设置页码：开启奇偶页不同，奇右偶左，格式 "- 1 -"。"""

    # 开启奇偶页不同
    settings_element = doc.settings._element
    even_and_odd = settings_element.find(qn('w:evenAndOddHeaders'))
    if even_and_odd is None:
        even_and_odd = OxmlElement('w:evenAndOddHeaders')
        settings_element.append(even_and_odd)
    even_and_odd.set(qn('w:val'), '1')  # 必须设置值才生效

    for section in doc.sections:
        # 奇数页页脚（默认页脚）— 右对齐
        footer_odd = section.footer
        footer_odd.is_linked_to_previous = False
        _setup_page_number_paragraph(footer_odd, WD_ALIGN_PARAGRAPH.RIGHT,
                                     page_rule)

        # 偶数页页脚 — 左对齐
        try:
            footer_even = section.even_page_footer
            footer_even.is_linked_to_previous = False
            _setup_page_number_paragraph(footer_even, WD_ALIGN_PARAGRAPH.LEFT,
                                         page_rule)
        except Exception as e:
            logger.warning(f"[引擎] 偶数页页脚设置失败: {e}（部分 Word 版本不支持）")


def _setup_page_number_paragraph(footer, alignment, page_rule: StyleRule) -> None:
    """在一个页脚中设置页码。使用 Word PAGE 域代码，格式由 page_rule 定义。"""

    # 清空页脚
    for p in footer.paragraphs:
        p.clear()

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = alignment

    # 居右空 1 字 / 居左空 1 字（GB/T 9704-2012 7.5）
    pPr = para._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        ind.set(qn('w:rightChars'), '100')  # 右缩进 1 字符
    else:
        ind.set(qn('w:leftChars'), '100')   # 左缩进 1 字符
    pPr.append(ind)

    # 添加文本 "— " + PAGE 域 + " —"
    for text, is_field in [("— ", False), ("PAGE", True), (" —", False)]:
        run = para.add_run(text)
        run.font.name = page_rule.font
        run.font.size = Pt(page_rule.font_size_pt)
        run.font.bold = page_rule.bold

        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), page_rule.font)
        rFonts.set(qn('w:ascii'), page_rule.font)
        rFonts.set(qn('w:hAnsi'), page_rule.font)
        rPr.append(rFonts)

        if is_field:
            # 完整 PAGE 域：begin → instrText → separate → 默认文本 → end
            run.text = ""
            fldChar_begin = OxmlElement('w:fldChar')
            fldChar_begin.set(qn('w:fldCharType'), 'begin')
            run._element.append(fldChar_begin)

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' PAGE '
            run._element.append(instrText)

            fldChar_sep = OxmlElement('w:fldChar')
            fldChar_sep.set(qn('w:fldCharType'), 'separate')
            run._element.append(fldChar_sep)

            t = OxmlElement('w:t')
            t.text = '1'
            run._element.append(t)

            fldChar_end = OxmlElement('w:fldChar')
            fldChar_end.set(qn('w:fldCharType'), 'end')
            run._element.append(fldChar_end)


# ═══════════════════════════════════════════════════════════════
# 后处理：数字/字母 → Times New Roman
# ═══════════════════════════════════════════════════════════════

_DIGIT_LATIN_RE = re.compile(r'([0-9]+|[a-zA-Z]+)')

def _apply_digit_latin_font(paragraph) -> None:
    """遍历段落所有 run，把数字和拉丁字母拆分，设为 Times New Roman。"""

    # 跳过含换行符的段落（如标题中的 <w:br/>），避免破坏换行
    for run in paragraph.runs:
        if run._element.find(qn('w:br')) is not None:
            return

    # 先收集所有需要处理的 run（避免迭代中修改）
    all_runs = list(paragraph.runs)

    for run in all_runs:
        text = run.text
        if not text or not _DIGIT_LATIN_RE.search(text):
            continue
        if run.font.superscript:  # 上标 run 不拆
            continue

        rPr_xml = run._element.find(qn('w:rPr'))
        rPr_clone = copy.deepcopy(rPr_xml) if rPr_xml is not None else None

        # 按数字/字母拆分，在原 run 后插入新 run（正向，保持顺序）
        parts = _DIGIT_LATIN_RE.split(text)
        insert_after = run._element
        for part in parts:
            if not part:
                continue
            new_r = OxmlElement('w:r')
            if rPr_clone is not None:
                new_r.append(copy.deepcopy(rPr_clone))
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = part
            new_r.append(t)
            if _DIGIT_LATIN_RE.fullmatch(part):
                nrPr = new_r.find(qn('w:rPr'))
                if nrPr is None:
                    nrPr = OxmlElement('w:rPr')
                    new_r.insert(0, nrPr)
                rf = nrPr.find(qn('w:rFonts'))
                if rf is None:
                    rf = OxmlElement('w:rFonts')
                    nrPr.append(rf)
                rf.set(qn('w:ascii'), 'Times New Roman')
                rf.set(qn('w:hAnsi'), 'Times New Roman')
            insert_after.addnext(new_r)
            insert_after = new_r

        run._element.getparent().remove(run._element)

# ═══════════════════════════════════════════════════════════════
# 后处理：上标 → [n] 格式
# ═══════════════════════════════════════════════════════════════

def _apply_universal_superscript(paragraph) -> None:
    """扫描段落中的所有 run，将已经是上标的文本统一为 [n] 格式。"""
    for run in paragraph.runs:
        if run.font.superscript:
            # 已经是上标 → 确保格式一致
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)  # 三号
            # 仅纯数字加 [n] 括号，非数字（如注/※）原样保留
            text = run.text.strip()
            if re.match(r'^\d+$', text):
                run.text = f"[{text}]"


# ═══════════════════════════════════════════════════════════════
# 表格/图片复制
# ═══════════════════════════════════════════════════════════════

def _copy_table(doc, table) -> None:
    """跨文档复制表格（底层 XML 操作）。"""
    import copy as pycopy

    tbl_element = pycopy.deepcopy(table._tbl)
    _remap_image_relationships(tbl_element, table.part, doc.part)
    _append_body_element(doc, tbl_element)


def _append_body_element(doc, element) -> None:
    """Append a raw body child before sectPr, keeping Word body XML valid."""
    body = doc.element.body
    sectPr = body.sectPr
    if sectPr is not None:
        sectPr.addprevious(element)
    else:
        body.append(element)


def _remap_image_relationships(element, source_part, target_part) -> None:
    """Copy image relationships referenced by a copied paragraph/table XML."""
    for blip in element.findall('.//' + qn('a:blip')):
        for attr in (qn('r:embed'), qn('r:link')):
            old_rid = blip.get(attr)
            if not old_rid:
                continue
            related = source_part.related_parts.get(old_rid)
            if related is None or not hasattr(related, "blob"):
                continue
            new_rid, _ = target_part.get_or_add_image(io.BytesIO(related.blob))
            blip.set(attr, new_rid)

    try:
        legacy_images = element.findall('.//' + qn('v:imagedata'))
    except KeyError:
        legacy_images = []
    for imagedata in legacy_images:
        old_rid = imagedata.get(qn('r:id'))
        if not old_rid:
            continue
        related = source_part.related_parts.get(old_rid)
        if related is None or not hasattr(related, "blob"):
            continue
        new_rid, _ = target_part.get_or_add_image(io.BytesIO(related.blob))
        imagedata.set(qn('r:id'), new_rid)


def _copy_image(doc, source_para) -> None:
    """跨文档复制图片（二进制 blob 路线）。

    注意：此功能在 V1 为简化实现，仅保留原段落文本，
    图片部分暂不做完全复制。
    """
    import copy as pycopy

    p_element = pycopy.deepcopy(source_para._p)
    _remap_image_relationships(p_element, source_para.part, doc.part)
    _append_body_element(doc, p_element)
    logger.debug(f"[引擎] 图片段落已按原位复制: '{source_para.text[:30]}'")


# ═══════════════════════════════════════════════════════════════
# Type → StyleRule 索引映射
# ═══════════════════════════════════════════════════════════════

TYPE_TO_RULE_INDEX: Dict[str, int] = {
    "title": 0, "title_cont": 0,     # 主标题 + 续行 → row 0
    "heading1": 1, "heading1_report": 1,  # 报告 heading1 同 row 1，但无编号
    "heading2": 2, "heading3": 3, "heading4": 4,
    "body": 5, "attachment": 5,
    "addressing": 10, "date_line": 11, "author_line": 12, "role_name": 13,
    "title2": 14, "sign_off": 15,
    "glossary_title": 0, "glossary_item": 16,
    "attachment_note": 17, "attachment_note_item": 18,
    "attachment_page_mark": 19, "attachment_title": 20, "attachment_body": 21,
    "sign_org": 22, "sign_date": 23,
    "number": 6, "letter": 7,
    "page_number": 8, "superscript": 9,
}

HEAD_TYPES_REQUIRING_GAP = ("title", "title_cont", "date_line", "author_line", "role_name", "attachment_title")
HEAD_GAP_FOLLOW_TYPES = ("body", "attachment_body", "heading1")


# ═══════════════════════════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════════════════════════

def export_doc(doc_data: DocumentData, rules: List[StyleRule],
               settings: PageSettings, output_path: str,
               numbered_bold_enabled: bool = True) -> dict:
    """排版引擎主入口。DocumentData → .docx 文件。

    Returns:
        dict: 排版统计信息。
    """

    logger.info(f"[引擎] 排版开始: {doc_data.filepath} → {output_path}")
    logger.info(f"[引擎] 共 {len(doc_data.paragraphs)} 段, {len(doc_data.tables)} 表格")

    doc = Document()

    stats = {
        "total": len(doc_data.paragraphs),
        "heading1": 0, "heading1_report": 0, "heading2": 0, "heading3": 0, "heading4": 0,
        "body": 0, "output_path": output_path,
    }

    # 查找页码规则（row 7）
    page_rule = rules[8] if len(rules) > 8 else StyleRule.default_for_row(8)

    prev_was_title = False
    prev_type_id = ""
    line_twips = _line_spacing_twips(settings)

    render_items = doc_data.paragraphs
    paragraph_i = 0
    _deferred_body_log = []  # heading1 拆出的 body 日志

    for i, pd in enumerate(render_items):
        # 表格占位符 → 原位复制
        if pd.type_id == "__table__":
            try: _copy_table(doc, pd.meta.get("table"))
            except Exception as e: logger.warning(f"[引擎] 表格复制失败: {e}")
            continue
        # 图片占位符 → 原位复制
        if pd.type_id == "__image__":
            try: _copy_image(doc, pd.meta.get("image_xml"))
            except Exception as e: logger.warning(f"[引擎] 图片段落复制失败: {e}")
            continue

        para_no = paragraph_i
        paragraph_i += 1
        try:
            logger.debug(f"[引擎] 段落 {para_no}: type={pd.type_id} text='{pd.text[:30]}'")

            # 确定对应的 StyleRule 索引
            rule_index = TYPE_TO_RULE_INDEX.get(pd.type_id, 5)  # fallback → 正文
            raw_rule = rules[rule_index] if rule_index < len(rules) else StyleRule.default_for_row(rule_index)

            # meta → 重写规则（不脏 apply_style）
            meta = pd.meta or {}
            # glossary_title 特殊处理（内联段落）
            if pd.type_id == "glossary_title":
                resolved = copy.copy(raw_rule)
                resolved.row_index = 0; resolved.font = "方正小标宋简体"
                resolved.font_size_pt = 22; resolved.bold = False
                resolved.alignment = "居中"; resolved.first_line_indent = 0.0
                para = doc.add_paragraph(pd.text)
                run = para.runs[0] if para.runs else para.add_run(pd.text)
                apply_style(para, resolved)
                _set_run_fonts(run, cn_font=resolved.font)
                run.font.size = Pt(resolved.font_size_pt)
                run.font.bold = resolved.bold
                _set_para_spacing(para, before_lines=1, after_lines=1)
                pPr = para._element.get_or_add_pPr()
                page_break = OxmlElement('w:pageBreakBefore')
                pPr.append(page_break)
                prev_was_title = True
                continue

            # meta → 重写规则（按文种分派）
            resolved = _resolve_rule(pd, raw_rule, rules)

            # 空行插入（三号 16pt 行距 28 磅）
            # date_line 后的留白由段后间距控制，避免生成真实空段落。
            need_gap = (prev_was_title and pd.type_id in HEAD_GAP_FOLLOW_TYPES
                        and prev_type_id != "date_line" and pd.text.strip())

            if need_gap:
                spacer = doc.add_paragraph("")
                spPPr = spacer._element.get_or_add_pPr()
                spSpacing = OxmlElement('w:spacing')
                spSpacing.set(qn('w:line'), str(line_twips))
                spSpacing.set(qn('w:lineRule'), 'exact')
                spPPr.append(spSpacing)
                spacer.add_run("").font.size = Pt(16)

            # 标题先清理旧编号，再建段落
            text = pd.text
            if pd.type_id.startswith("heading"):
                text = _strip_heading_numbering(text)
                # 一/二级标题特殊处理：句号分割的行内标题（政协报告体例）
                if pd.type_id in ("heading1", "heading2"):
                    text = _handle_heading_period(text)

            para = doc.add_paragraph(text)

            # 逐段写入 XML 属性
            pPr = para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            # 段前/段后间距（自然模式下生效）
            before_twip = int(settings.space_before_line * line_twips)
            after_twip = int(settings.space_after_line * line_twips)
            spacing.set(qn('w:before'), str(before_twip))
            spacing.set(qn('w:after'), str(after_twip))
            spacing.set(qn('w:beforeLines'), str(int(settings.space_before_line * 100)))
            spacing.set(qn('w:afterLines'), str(int(settings.space_after_line * 100)))
            # 网格模式 → 固定行距；自然模式(每行=0) → 不设，让段间距生效
            if settings.chars_per_line > 0:
                spacing.set(qn('w:line'), str(line_twips))
                spacing.set(qn('w:lineRule'), 'exact')
            pPr.append(spacing)
            if i < 3:
                logger.info(f"[段落{i}] XML: grid={'on' if settings.chars_per_line>0 else 'off'}")
            # 禁用 Word 自动上下文间距
            ctxSpc = OxmlElement('w:contextualSpacing')
            ctxSpc.set(qn('w:val'), '0')
            _set_unique(pPr, qn('w:contextualSpacing'), ctxSpc)
            # 关闭孤行控制
            wc = OxmlElement('w:widowControl')
            wc.set(qn('w:val'), '0')
            _set_unique(pPr, qn('w:widowControl'), wc)

            # 应用样式（带降级）
            if not apply_style_safe(para, resolved):
                logger.warning(f"[引擎] 段落 {i} 样式降级，继续后续")
                num, htext = _contact_wrap_data
                first_start = getattr(resolved, 'first_line_indent', 2.0) or 2.0
                follow_start = first_start + len(num + htext)
                _apply_hanging_indent_chars(para, first_start, follow_start)

            _apply_rule_paragraph_format(para, resolved, line_twips)

            # 头部署名/日期的相邻间距：
            # “职务姓名 + 日期”连续出现时，中间段后为 0；日期之后用段后 1 行。
            if pd.type_id in ("role_name", "author_line"):
                _set_para_spacing(para, before_lines=0, after_lines=0, line_twips=line_twips)
            elif pd.type_id == "date_line":
                _set_para_spacing(para, before_lines=0, after_lines=1, line_twips=line_twips)

            # (colon_inline_body removed — scheme mode deleted)            # date_line 强制适应一行（自动计算压缩量）
            if getattr(resolved, 'date_line_compress', False):
                pPr = para._element.get_or_add_pPr()
                csc = OxmlElement('w:characterSpacingControl')
                csc.set(qn('w:val'), 'compressPunctuation')
                pPr.append(csc)
                # 自动计算字符间距压缩
                text_len = len(pd.text)
                if text_len > 27:
                    shrink = min(int((text_len - 27) * 5), 40)  # 每超出1字收紧0.25pt，最多2pt
                    for run in para.runs:
                        rPr = run._element.get_or_add_rPr()
                        sp = OxmlElement('w:spacing')
                        sp.set(qn('w:val'), str(-shrink))
                        rPr.append(sp)

            # glossary_title 独立分支保留；其他段前段后由 JSON StyleRule 控制。
            if pd.type_id == "glossary_title":
                _set_para_spacing(para, before_lines=1, after_lines=1, line_twips=line_twips)

            # heading1/heading2 行内标题：句号分割，标题样式 + 正文仿宋
            if pd.type_id == "heading1" and "。" in para.text:
                period_pos = para.text.find("。")  # 用 rendered text
                full_text = para.text  # 保存全文
                after = full_text[period_pos + 1:].strip()
                if len(after) >= 5 and para.runs:
                    heading_text = full_text[:period_pos + 1]
                    body_text = full_text[period_pos + 1:]
                    # 标题文字去掉句号
                    para.runs[-1].text = heading_text.rstrip("。")
                    # 正文另起一段
                    body_para = doc.add_paragraph(body_text)
                    body_font = rules[5].font if len(rules) > 5 else "仿宋_GB2312"
                    # 预写入 spacing
                    bpPr = body_para._element.get_or_add_pPr()
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:line'), str(line_twips))
                    spacing.set(qn('w:lineRule'), 'exact')
                    bpPr.append(spacing)
                    # 应用正文样式
                    body_rule = rules[5] if len(rules) > 5 else StyleRule.default_for_row(5)
                    apply_style(body_para, body_rule)
                    _apply_rule_paragraph_format(body_para, body_rule, line_twips)
                    # 关闭孤行控制
                    wc = OxmlElement('w:widowControl')
                    wc.set(qn('w:val'), '0')
                    _set_unique(bpPr, qn('w:widowControl'), wc)
                    ctxSpc = OxmlElement('w:contextualSpacing')
                    ctxSpc.set(qn('w:val'), '0')
                    _set_unique(bpPr, qn('w:contextualSpacing'), ctxSpc)
                    # 编号 + 后续处理
                    apply_superscript_split(body_para)
                    if numbered_bold_enabled:
                        _apply_special_bold(body_para, body_text)
                    _deferred_body_log.append(
                        f"[排版] #{i}h1→body | "
                        f"\"{body_text[:28]}\" | "
                        f"字体={body_rule.font} | 字号={body_rule.font_size_pt}pt | 加粗={body_rule.bold} | "
                        f"对齐={body_rule.alignment} | 首行缩进={body_rule.first_line_indent}字符 | "
                        f"行距={settings.line_spacing_value}pt固定 | 对网=1"
                    )

            # heading2 句号分割，标题+正文同段（方案模式不拆分）
            if pd.type_id == "heading2" and "。" in para.text:
                period_pos = para.text.find("。")
                full_text = para.text
                after = full_text[period_pos + 1:].strip()
                if len(after) >= 15 and para.runs:
                    para.runs[-1].text = full_text[:period_pos + 1]
                    body_run = para.add_run(full_text[period_pos + 1:])
                    _set_run_fonts(body_run, cn_font=rules[5].font, en_font="Times New Roman")
                    body_run.font.size = Pt(resolved.font_size_pt)
                    body_run.font.bold = False
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # X是/固定词组 特殊加粗
            if pd.meta.get("numbered_bold") and para.runs:
                _apply_special_bold(para, pd.text)

            # 冒号关键词加粗（如"责任单位：区政府" → "责任单位："加粗）
            if pd.meta.get("colon_bold") and para.runs:
                _apply_colon_bold(para, pd.text)

            # heading1_report 句号后换行
            if pd.meta.get("heading1_report_split") and para.runs:
                _apply_heading1_report_split(para, pd.text, resolved)

            # 报告首句加粗（首句楷体_GB2312 加粗，剩余仿宋正文）
            if pd.meta.get("report_first_sentence_bold") and para.runs:
                _apply_report_first_sentence(para, pd.text, resolved)

            # 编号：从 meta 读取预计算编号，插入到第一个 run 前
            numbering = "" if pd.meta.get("colon_inline_body") else pd.meta.get("numbering", "")
            logger.debug(f"[编号] meta={numbering!r} type={pd.type_id}")
            if numbering and para.runs:
                new_r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = numbering
                new_r.append(t)
                # 设置字体
                rPr = OxmlElement('w:rPr')
                rF = OxmlElement('w:rFonts')
                rF.set(qn('w:eastAsia'), resolved.font)
                rF.set(qn('w:ascii'), 'Times New Roman')
                rF.set(qn('w:hAnsi'), 'Times New Roman')
                rPr.append(rF)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(resolved.font_size_pt * 2)))  # half-pt → 1/2pt
                rPr.append(sz)
                if resolved.bold:
                    b = OxmlElement('w:b')
                    rPr.append(b)
                new_r.insert(0, rPr)
                para.runs[0]._element.addprevious(new_r)

            # 名词解释条目（编号后执行：关键词黑体，正文仿宋）
            if pd.meta.get("glossary_item") and para.runs:
                _apply_glossary_item(para, pd.text, resolved)

            # 上标拆分
            apply_superscript_split(para)

            # 最终强制 snapToGrid
            pPr_final = para._element.get_or_add_pPr()
            snap_val = '0' if pd.type_id in ('title', 'heading1', 'heading2', 'heading3', 'heading4') else '1'
            snap = OxmlElement('w:snapToGrid')
            snap.set(qn('w:val'), snap_val)
            _set_unique(pPr_final, qn('w:snapToGrid'), snap)
            # 段落排版日志（每段汇总格式信息）
            text_preview = pd.text[:28].replace('\n', ' ')
            indent = getattr(resolved, 'first_line_indent', 0)
            logger.info(
                f"[排版] #{i} {pd.type_id} | "
                f"\"{text_preview}\" | "
                f"字体={resolved.font} | "
                f"字号={resolved.font_size_pt}pt | "
                f"加粗={resolved.bold} | "
                f"对齐={resolved.alignment} | "
                f"首行缩进={int(indent)}字符 | "
                    f"行距={settings.line_spacing_value}pt固定 | "
                f"对网={snap_val}"
            )
            # 输出 heading1 拆出的 body 日志
            for log_line in _deferred_body_log:
                logger.info(log_line)
            _deferred_body_log.clear()

            # 记录头部区域；后接正文或一级标题时需要空一行。
            prev_was_title = (pd.type_id in HEAD_TYPES_REQUIRING_GAP)
            prev_type_id = pd.type_id

            # 统计
            if pd.type_id in stats:
                stats[pd.type_id] += 1
            elif pd.type_id.startswith("heading"):
                k = pd.type_id
                stats[k] = stats.get(k, 0) + 1
            else:
                stats["body"] += 1

        except Exception as e:
            logger.error(f"[引擎] 段落 {i} 异常: {e}，降级为纯文本")
            # 降级兜底：不跳过，用原始文本 + 正文格式写入
            try:
                fallback = doc.add_paragraph(pd.text)
                fallback.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                stats["body"] += 1
            except Exception:
                pass  # 最终兜底：实在写不了才跳过
            continue

    # 后处理：上标统一
    for para in doc.paragraphs:
        _apply_universal_superscript(para)

    # 后处理：数字/字母 → Times New Roman
    for para in doc.paragraphs:
        _apply_digit_latin_font(para)

    # 页面设置（边距 + compat + Normal 样式）
    apply_page_settings(doc, settings, doc_data.doc_mode)

    # 页码
    apply_header_footer(doc, page_rule)

    # 页面行数诊断
    page_h_cm = 29.7 - settings.margin_top_cm - settings.margin_bottom_cm
    max_lines = page_h_cm / (settings.line_spacing_value * 0.0353)  # pt→cm
    logger.info(f"[页面] 版心高度={page_h_cm:.1f}cm 行距={settings.line_spacing_value}pt → 理论最大={max_lines:.1f}行 设定={settings.lines_per_page}行")

    # 保存
    try:
        doc.save(output_path)
        logger.info(f"[引擎] 排版完成: {output_path}")
    except Exception as e:
        raise ExportError(f"保存失败 {output_path}: {e}")

    return stats


# ═══════════════════════════════════════════════════════════════
# 验证（依赖 python-docx）
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # NumberingCounter 验证
    nc = NumberingCounter()
    nc.advance("heading1")
    assert nc.a == 1 and nc.b == 0 and nc.c == 0 and nc.d == 0
    nc.advance("heading2")
    assert nc.a == 1 and nc.b == 1 and nc.c == 0 and nc.d == 0
    nc.advance("heading3")
    assert nc.a == 1 and nc.b == 1 and nc.c == 1 and nc.d == 0
    nc.advance("heading4")
    assert nc.a == 1 and nc.b == 1 and nc.c == 1 and nc.d == 1
    nc.advance("heading2")
    assert nc.b == 2 and nc.c == 0 and nc.d == 0, f"级联清零失败: {nc.b},{nc.c},{nc.d}"
    nc.advance("heading1")
    assert nc.a == 2 and nc.b == 0 and nc.c == 0 and nc.d == 0, f"heading1 级联清零失败: {nc}"

    # render 验证
    nc2 = NumberingCounter()
    nc2.a = 1; nc2.b = 2; nc2.c = 3; nc2.d = 4
    assert nc2.render("{a}、", "heading1") == "一、", f"render 中文失败: {nc2.render('{a}、', 'heading1')}"
    assert nc2.render("（{b}）", "heading2") == "（二）", f"render 括号中文失败"
    assert nc2.render("{c}.", "heading3") == "3.", f"render 阿拉伯失败"
    assert nc2.render("({d})", "heading4") == "(4)", f"render 括号阿拉伯失败"
    assert nc2.render("- 1 -", "page_number") == "- 1 -", "render 固定值不应变动"

    # TYPE_TO_RULE_INDEX
    assert TYPE_TO_RULE_INDEX["body"] == 5
    assert TYPE_TO_RULE_INDEX["title"] == 0
    assert TYPE_TO_RULE_INDEX["heading1"] == 1
    assert TYPE_TO_RULE_INDEX["heading2"] == 2
    assert TYPE_TO_RULE_INDEX["heading3"] == 3
    assert TYPE_TO_RULE_INDEX["heading4"] == 4

    # Regression: attachment continuation items should use whole-paragraph
    # left indent of 5 chars, with no hanging indent.
    from pathlib import Path
    from tempfile import TemporaryDirectory
    from zipfile import ZipFile
    from xml.etree import ElementTree as ET

    with TemporaryDirectory() as tmpdir:
        out = Path(tmpdir) / "attachment_indent.docx"
        data = DocumentData(paragraphs=[
            ParagraphData("附件：1. 基本情况", "attachment_note", "附件：1. 基本情况", None),
            ParagraphData("2. 具体情况", "attachment_note_item", "2. 具体情况", None),
        ])
        export_doc(data, [StyleRule.default_for_row(i) for i in range(10)], PageSettings(), str(out))
        with ZipFile(out) as zf:
            root = ET.fromstring(zf.read("word/document.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = root.findall(".//w:body/w:p", ns)
        item_ind = paragraphs[1].find("w:pPr/w:ind", ns)
        assert item_ind is not None
        assert item_ind.get(qn("w:leftChars")) == "500", item_ind.attrib
        assert qn("w:hangingChars") not in item_ind.attrib and qn("w:hanging") not in item_ind.attrib

    print("✅ engine.py 纯函数验证全部通过")
