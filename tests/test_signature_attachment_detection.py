import logging
import tempfile
import unittest
from pathlib import Path

from docx import Document

from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import StyleRule, logger


def _rules():
    return [StyleRule.default_for_row(i) for i in range(10)]


class SignatureAttachmentDetectionTest(unittest.TestCase):
    def setUp(self):
        logger.setLevel(logging.ERROR)
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)

    def tearDown(self):
        self.tmp.cleanup()

    def _load_lines(self, lines):
        doc = Document()
        for line in lines:
            doc.add_paragraph(line)
        path = self.root / "input.docx"
        doc.save(path)
        return DocxImporter().load(str(path), _rules())

    def test_sign_org_date_and_attachment_page(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "区政府人才保障工作组",
            "2025年10月15日",
            "附件1",
            "基本情况",
            "测试正文测试正文。",
        ])

        self.assertEqual(
            [p.type_id for p in data.paragraphs[-5:]],
            ["sign_org", "sign_date", "attachment_page_mark", "attachment_title", "attachment_body"],
        )

    def test_tail_date_without_sign_org_before_attachment_page(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "2025年十月15日",
            "附件 1",
            "基本情况",
            "测试正文测试正文。",
        ])

        self.assertEqual(
            [p.type_id for p in data.paragraphs[-4:]],
            ["sign_date", "attachment_page_mark", "attachment_title", "attachment_body"],
        )
        self.assertEqual(data.paragraphs[-4].text, "2025年10月15日")

    def test_tail_date_after_responsibility_line_before_attachment_page(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "责任单位：区政府",
            "2025年十月15日",
            "附件1",
            "基本情况",
            "测试正文测试正文。",
        ])

        self.assertEqual(data.paragraphs[-4].type_id, "sign_date")
        self.assertEqual(data.paragraphs[-4].text, "2025年10月15日")
        self.assertEqual(
            [p.type_id for p in data.paragraphs[-3:]],
            ["attachment_page_mark", "attachment_title", "attachment_body"],
        )

    def test_attachment_note_can_start_attachment_page_without_date(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "附件：基本情况",
            "附件一",
            "基本情况",
            "测试正文测试正文。",
        ])

        self.assertEqual(
            [p.type_id for p in data.paragraphs[-4:]],
            ["attachment_note", "attachment_page_mark", "attachment_title", "attachment_body"],
        )
        self.assertEqual(data.paragraphs[-3].text, "附件 一")

    def test_multiple_attachment_pages_continue_after_attachment_body(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "附件：1. 基本情况",
            "2. 具体情况",
            "2025年10月15日",
            "附件1",
            "基本情况",
            "第一项正文。",
            "附件2",
            "具体情况",
            "第二项正文。",
        ])

        tail_types = [p.type_id for p in data.paragraphs[-9:]]
        self.assertEqual(
            tail_types,
            [
                "attachment_note",
                "attachment_note_item",
                "sign_date",
                "attachment_page_mark",
                "attachment_title",
                "attachment_body",
                "attachment_page_mark",
                "attachment_title",
                "attachment_body",
            ],
        )

    def test_middle_date_is_not_signature_date(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "2025年10月15日",
            "这里是后续正文内容这里是后续正文内容。",
            "附件：基本情况",
            "附件1",
            "基本情况",
            "测试正文测试正文。",
        ])

        middle_date = next(p for p in data.paragraphs if p.original_text == "2025年10月15日")
        self.assertNotEqual(middle_date.type_id, "sign_date")

    def test_inline_attachment_reference_is_not_attachment_page_mark(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "区政府人才保障工作组",
            "2025年10月15日",
            "附件1所示",
        ])

        self.assertNotEqual(data.paragraphs[-1].type_id, "attachment_page_mark")

    def test_colon_attachment_line_remains_attachment_note(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "2025年10月15日",
            "附件：基本情况",
            "附件",
            "基本情况",
            "测试正文测试正文。",
        ])

        self.assertEqual(data.paragraphs[-4].type_id, "attachment_note")
        self.assertEqual(data.paragraphs[-3].type_id, "attachment_page_mark")

    def test_chinese_month_day_tail_date_is_normalized(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "2025年十月15日",
        ])

        self.assertEqual(data.paragraphs[-1].type_id, "sign_date")
        self.assertEqual(data.paragraphs[-1].text, "2025年10月15日")


if __name__ == "__main__":
    unittest.main()
