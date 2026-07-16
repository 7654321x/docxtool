from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxtool.document.engine.core import export_doc
from docxtool.document.engine.signature_block import apply_signature_block
from docxtool.document.engine.style_catalog import ensure_document_styles
from docxtool.document.importer import DocumentData, ParagraphData, ParagraphFeatures
from docxtool.document.style_config import PageSettings, StyleRule
from docxtool.security.docx_integrity import validate_docx_integrity


def _rules() -> list[StyleRule]:
    return [StyleRule.default_for_row(index) for index in range(24)]


def _document_with_styles():
    document = Document()
    ensure_document_styles(document, _rules(), PageSettings())
    return document


def _style(document, style_id: str):
    return next(style for style in document.styles if style.style_id == style_id)


def _paragraph(document, text: str, style_id: str):
    paragraph = document.add_paragraph(text)
    paragraph.style = _style(document, style_id)
    return paragraph


def _right_chars(paragraph) -> str | None:
    properties = paragraph._element.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    return indent.get(qn("w:rightChars")) if indent is not None else None


def _set_right_chars(paragraph, value: str) -> None:
    properties = paragraph._element.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    indent.set(qn("w:rightChars"), value)


def test_preserve_mode_does_not_change_existing_signature_indents() -> None:
    document = _document_with_styles()
    signature = _paragraph(document, "内江市东兴区人民政府", "DCT-Signature")
    date = _paragraph(document, "2026年7月16日", "DCT-Date")
    _set_right_chars(signature, "700")
    _set_right_chars(date, "300")

    assert apply_signature_block(document, {"mode": "preserve"}) == 0

    assert _right_chars(signature) == "700"
    assert _right_chars(date) == "300"


def test_without_seal_places_date_two_characters_right_of_long_signature() -> None:
    document = _document_with_styles()
    signature_text = "内江市东兴区人民政府"
    date_text = "2026年7月16日"
    signature = _paragraph(document, signature_text, "DCT-Signature")
    date = _paragraph(document, date_text, "DCT-Date")

    assert apply_signature_block(document, {"mode": "without_seal"}) == 1

    assert _right_chars(signature) == "200"
    assert _right_chars(date) == "350"
    assert signature.text == signature_text
    assert date.text == date_text


def test_without_seal_keeps_two_character_minimum_for_long_date() -> None:
    document = _document_with_styles()
    signature = _paragraph(document, "区政府", "DCT-Signature")
    date = _paragraph(document, "二〇二六年七月十六日", "DCT-Date")

    apply_signature_block(document, {"mode": "without_seal"})

    assert _right_chars(signature) == "200"
    assert _right_chars(date) == "200"


def test_with_seal_centers_signature_on_date_text_area() -> None:
    document = _document_with_styles()
    signature = _paragraph(document, "内江市东兴区人民政府", "DCT-Signature")
    date = _paragraph(document, "2026年7月16日", "DCT-Date")

    apply_signature_block(document, {"mode": "with_seal"})

    assert _right_chars(date) == "400"
    assert _right_chars(signature) == "225"
    signature_center_from_right = 2.25 + 10 / 2
    date_center_from_right = 4 + 6.5 / 2
    assert signature_center_from_right == date_center_from_right


def test_unreliable_and_joint_signature_blocks_are_preserved() -> None:
    document = _document_with_styles()
    nonadjacent_signature = _paragraph(document, "甲机关", "DCT-Signature")
    document.add_paragraph("间隔内容")
    nonadjacent_date = _paragraph(document, "2026年7月16日", "DCT-Date")
    first_joint = _paragraph(document, "甲机关", "DCT-Signature")
    second_joint = _paragraph(document, "乙机关", "DCT-Signature")
    joint_date = _paragraph(document, "2026年7月16日", "DCT-Date")

    assert apply_signature_block(document, {"mode": "with_seal"}) == 0

    for paragraph in (
        nonadjacent_signature,
        nonadjacent_date,
        first_joint,
        second_joint,
        joint_date,
    ):
        assert _right_chars(paragraph) is None


def test_table_signature_and_complex_signature_are_preserved() -> None:
    document = _document_with_styles()
    table = document.add_table(rows=2, cols=1)
    table_signature = table.cell(0, 0).paragraphs[0]
    table_signature.text = "表格机关"
    table_signature.style = _style(document, "DCT-Signature")
    table_date = table.cell(1, 0).paragraphs[0]
    table_date.text = "2026年7月16日"
    table_date.style = _style(document, "DCT-Date")
    complex_signature = _paragraph(document, "甲机关\n乙机关", "DCT-Signature")
    complex_date = _paragraph(document, "2026年7月16日", "DCT-Date")

    assert apply_signature_block(document, {"mode": "without_seal"}) == 0

    assert _right_chars(table_signature) is None
    assert _right_chars(table_date) is None
    assert _right_chars(complex_signature) is None
    assert _right_chars(complex_date) is None


def test_signature_layout_is_idempotent_and_output_reopens(tmp_path: Path) -> None:
    document = _document_with_styles()
    _paragraph(document, "内江市东兴区人民政府", "DCT-Signature")
    _paragraph(document, "2026年7月16日", "DCT-Date")

    apply_signature_block(document, {"mode": "with_seal"})
    first_xml = document._body._element.xml
    apply_signature_block(document, {"mode": "with_seal"})
    assert document._body._element.xml == first_xml
    output = tmp_path / "signature-block.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    assert [paragraph.text for paragraph in Document(output).paragraphs[-2:]] == [
        "内江市东兴区人民政府",
        "2026年7月16日",
    ]


def test_export_doc_invokes_signature_block_engine(tmp_path: Path) -> None:
    output = tmp_path / "exported-signature.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData("内江市东兴区人民政府", "sign_org", "内江市东兴区人民政府", ParagraphFeatures()),
            ParagraphData("2026年7月16日", "sign_date", "2026年7月16日", ParagraphFeatures()),
        ],
        filepath=str(tmp_path / "input.docx"),
    )

    stats = export_doc(
        data,
        _rules(),
        PageSettings(),
        str(output),
        page_number_options={"enabled": False},
        signature_block_options={"mode": "without_seal"},
    )

    reopened = Document(output)
    assert stats["signature_blocks_adjusted"] == 1
    assert _right_chars(reopened.paragraphs[-2]) == "200"
    assert _right_chars(reopened.paragraphs[-1]) == "350"
    assert validate_docx_integrity(output).ok is True
