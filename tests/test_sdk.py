from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from docxtool import __version__ as package_version
from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import load_rules_and_settings
from docxtool.paths import default_format_config_path
from docxtool.sdk import (
    InvalidRequestError,
    RecognitionInputError,
    recognize_docx,
    validate_recognition_plan,
)
from docxtool.sdk.cli import main as sdk_main, recognize_main


def _source_document(path: Path) -> None:
    document = Document()
    document.add_paragraph("关于推进基层治理工作的通知")
    document.add_paragraph("一、总体要求")
    document.add_paragraph("要坚持问题导向，完善基层治理工作机制。")
    document.save(path)


def _default_config() -> dict:
    return json.loads(default_format_config_path().read_text(encoding="utf-8"))


def test_sdk_matches_existing_authoritative_recognition_and_redacts_text(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _source_document(source)
    config = _default_config()
    plan = recognize_docx(source, format_config=config)

    rules, _settings, features = load_rules_and_settings(config)
    features["processing"] = {"strategy": "structural"}
    features["recognition"] = {"mode": "authoritative"}
    imported = DocxImporter().load(str(source), rules, features=features)

    assert plan.processing_mode == "structural"
    assert plan.recognition_mode == "authoritative"
    assert plan.package_version == package_version
    assert plan.schema_version == "recognition-plan-v1"
    assert plan.integration_contract_version == "integration-contract-v1"
    assert plan.plan_id
    assert all(block.block_id and block.physical_group_id for block in plan.blocks)
    assert plan.locator_version == "source-locator-v2"
    assert plan.host_text_contract_version == "host-text-v1"
    assert [block.type_id for block in plan.blocks] == [item.type_id for item in imported.paragraphs]
    assert all(len(block.text_sha256) in {0, 64} for block in plan.blocks)
    payload = json.dumps(plan.to_dict(), ensure_ascii=False)
    assert "基层治理" not in payload
    assert "总体要求" not in payload


def test_sdk_assigns_unique_ids_to_unlocated_letterhead_and_image_blocks(
    tmp_path: Path,
) -> None:
    source = tmp_path / "letterhead-objects.docx"
    document = Document()
    document.add_paragraph()
    mark = document.add_paragraph("测试机关文件")
    mark.alignment = 1
    mark.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    document.add_paragraph("测发〔2026〕1号")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)
    document.add_paragraph()
    document.add_paragraph("关于测试工作的通知")
    document.add_paragraph("正文内容。")
    for _ in range(2):
        image = document.add_paragraph()
        drawing = OxmlElement("w:drawing")
        extent = OxmlElement("wp:extent")
        extent.set("cx", "100")
        extent.set("cy", "100")
        drawing.append(extent)
        image._p.append(drawing)
    document.save(source)

    plan = recognize_docx(source)
    special = [block for block in plan.blocks if block.physical_paragraph_index is None]

    assert len(special) >= 7
    assert len({block.physical_group_id for block in special}) == len(special)
    assert len({block.block_id for block in plan.blocks}) == len(plan.blocks)
    assert validate_recognition_plan(plan).valid


def test_sdk_emits_verified_physical_utf16_ranges_and_local_text_is_opt_in(tmp_path: Path) -> None:
    source = tmp_path / "mixed.docx"
    document = Document()
    document.add_paragraph("一、测试标题。这是同一物理段落中的测试正文内容。")
    document.save(source)

    redacted = recognize_docx(source, recognition_mode="legacy")
    local = recognize_docx(source, recognition_mode="legacy", include_text=True)
    assert len(local.blocks) >= 2
    assert "recognized_text" not in redacted.blocks[0].to_dict()
    for block in local.blocks:
        if block.kind != "paragraph":
            continue
        assert block.source_paragraph_index == 0
        assert block.physical_paragraph_index == 0
        assert block.locator_verified is True
        assert len(block.physical_text_sha256) == 64
        assert block.range_start_utf16 is not None
        assert block.range_end_utf16 is not None
        assert block.text_length_utf16 == len(block.recognized_text.encode("utf-16-le")) // 2
        physical = "一、测试标题。这是同一物理段落中的测试正文内容。".encode("utf-16-le")
        selected = physical[block.range_start_utf16 * 2:block.range_end_utf16 * 2].decode("utf-16-le")
        assert selected == block.recognized_text


def test_sdk_keeps_numbered_heading2_colon_body_in_one_verified_range(
    tmp_path: Path,
) -> None:
    source = tmp_path / "heading2-colon-inline-body.docx"
    text = "（一）工作安排： 正文内容"
    document = Document()
    document.add_paragraph(text)
    document.save(source)

    plan = recognize_docx(source, include_text=True)

    paragraph_blocks = [block for block in plan.blocks if block.kind == "paragraph"]
    assert len(paragraph_blocks) == 1
    block = paragraph_blocks[0]
    assert block.type_id == "heading2"
    assert block.recognized_text == text
    assert block.locator_verified is True
    assert block.range_start_utf16 == 0
    assert block.range_end_utf16 == len(text.encode("utf-16-le")) // 2


def test_sdk_locates_front_role_and_date_split_from_one_physical_paragraph(tmp_path: Path) -> None:
    source = tmp_path / "front-metadata.docx"
    document = Document()
    document.add_paragraph("在某地区委员会会议闭幕大会上的讲话", style="Heading 1")
    byline = document.add_paragraph()
    role = byline.add_run("某区委员会党组书记、主席　张测试")
    role.bold = True
    role.font.size = Pt(22)
    byline.add_run().add_break(WD_BREAK.LINE)
    date = byline.add_run("（2026年8月27日11:00，某地区委员会会议中心）")
    date.bold = True
    date.font.size = Pt(22)
    document.add_paragraph("各位代表、同志们：")
    document.add_paragraph("现将有关事项说明如下，请认真抓好落实。")
    document.save(source)

    plan = recognize_docx(source, include_text=True, include_raw_text=True)
    front = [block for block in plan.blocks if block.physical_paragraph_index == 1]

    assert [block.type_id for block in front] == ["role_name", "date_line"]
    assert all(block.locator_verified for block in front)
    assert front[0].raw_end_utf16 <= front[1].raw_start_utf16
    physical = document.paragraphs[1].text.encode("utf-16-le")
    for block in front:
        selected = physical[
            block.raw_start_utf16 * 2:block.raw_end_utf16 * 2
        ].decode("utf-16-le")
        assert selected == block.raw_fragment_text


def test_sdk_cli_writes_json_plan(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "plan.json"
    _source_document(source)

    assert sdk_main([str(source), "--output", str(output)]) == 0
    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload["ok"] is True
    assert payload["data"]["schema_version"] == "recognition-plan-v1"
    assert payload["data"]["blocks"]


def test_sdk_legacy_entry_point_writes_json_plan(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "legacy-plan.json"
    _source_document(source)

    assert recognize_main([str(source), "--output", str(output)]) == 0
    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload["ok"] is True
    assert payload["data"]["schema_version"] == "recognition-plan-v1"
    assert payload["data"]["blocks"]


def test_sdk_rejects_invalid_input_and_modes(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _source_document(source)

    with pytest.raises(InvalidRequestError, match="processing_mode"):
        recognize_docx(source, processing_mode="unexpected")
    with pytest.raises(RecognitionInputError, match="可读取"):
        recognize_docx(tmp_path / "missing.docx")
