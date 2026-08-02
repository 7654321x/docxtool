from __future__ import annotations

import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from lxml import etree

from docxtool.document.models import DocumentData
from docxtool.document.importing.images import (
    contains_visible_image,
    is_object_caption,
    is_object_caption_text,
    is_standalone_image_paragraph,
)
from docxtool.document.importing.inline_tokens import (
    extract_inline_tokens,
    inline_tokens_text,
    normalize_inline_tokens,
)
from docxtool.document.importing.numbering import (
    detect_numbering_prefix,
    heading_style_prefix,
    is_auto_numbered_item,
    word_list_level_prefix,
)
from docxtool.document.models import ParagraphFeatures
from docxtool.document.importing.relationships import repair_broken_rels
from docxtool.document.importing.sections import (
    collect_section_header_footer_parts,
    extract_paragraph_sectPr,
)


class _Style:
    """保存样式名称，模拟 python-docx 段落样式对象。"""

    def __init__(self, name: str) -> None:
        self.name = name


class _Paragraph:
    """保存段落文本和样式，模拟题注判断需要的最小段落对象。"""

    def __init__(self, text: str, style_name: str = "") -> None:
        self.text = text
        self.style = _Style(style_name) if style_name else None


def test_image_caption_helpers_detect_caption_and_empty_image_paragraph() -> None:
    assert is_object_caption(_Paragraph("图1 测试图片"))
    assert is_object_caption_text("表2 测试表格")
    assert is_object_caption(_Paragraph("普通说明", "Caption"))
    assert is_object_caption(_Paragraph("普通说明", "题注"))
    assert not is_object_caption(_Paragraph("普通正文"))
    assert is_standalone_image_paragraph(_Paragraph(""))
    assert not is_standalone_image_paragraph(_Paragraph("正文"))


def test_contains_visible_image_ignores_zero_extent_drawing() -> None:
    element = etree.fromstring(
        b"""
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
          <w:r><w:drawing><wp:inline><wp:extent cx="1" cy="0"/></wp:inline></w:drawing></w:r>
        </w:p>
        """
    )

    assert contains_visible_image(element) is False


def test_contains_visible_image_accepts_positive_extent_and_legacy_pict() -> None:
    drawing = etree.fromstring(
        b"""
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
             xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
          <w:r><w:drawing><wp:inline><wp:extent cx="1" cy="2"/></wp:inline></w:drawing></w:r>
        </w:p>
        """
    )
    pict = etree.fromstring(
        b"""
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:r><w:pict/></w:r>
        </w:p>
        """
    )

    assert contains_visible_image(drawing) is True
    assert contains_visible_image(pict) is True


def test_inline_tokens_extract_text_tabs_line_breaks_and_page_breaks() -> None:
    paragraph = Document().add_paragraph()
    run = paragraph.add_run("甲")
    run.add_tab()
    run.add_text("乙")
    run.add_break(WD_BREAK.LINE)
    run.add_text("丙")
    run.add_break(WD_BREAK.PAGE)
    run.add_text("丁")

    tokens = extract_inline_tokens(paragraph)

    assert [token.kind for token in tokens] == ["text", "tab", "text", "line_break", "text", "page_break", "text"]
    assert inline_tokens_text(tokens) == "甲\t乙\n丙丁"


def test_normalize_inline_tokens_only_changes_text_tokens_when_enabled() -> None:
    paragraph = Document().add_paragraph()
    paragraph.add_run("A")
    paragraph.add_run().add_tab()
    tokens = extract_inline_tokens(paragraph)

    disabled = normalize_inline_tokens(tokens, enabled=False, normalize_text=lambda text: text.lower())
    enabled = normalize_inline_tokens(tokens, enabled=True, normalize_text=lambda text: text.lower())

    assert [token.text for token in disabled] == ["A", ""]
    assert [token.text for token in enabled] == ["a", ""]
    assert [token.kind for token in enabled] == ["text", "tab"]


def test_extract_paragraph_sectPr_returns_section_break_copy() -> None:
    paragraph = Document().add_paragraph("第一节")
    original = OxmlElement("w:sectPr")
    paragraph._p.get_or_add_pPr().append(original)

    sect_pr = extract_paragraph_sectPr(paragraph)

    assert sect_pr is not None
    assert sect_pr is not original


def test_collect_section_header_footer_parts_records_referenced_parts() -> None:
    doc = Document()
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "测试页眉"
    data = DocumentData()

    collect_section_header_footer_parts(doc, section._sectPr, data)

    assert data.section_relationship_parts


def test_detect_numbering_prefix_keeps_existing_importer_shapes() -> None:
    assert detect_numbering_prefix("一、一级标题") == "一、"
    assert detect_numbering_prefix("二.损坏一级标题") == "二."
    assert detect_numbering_prefix("（一）二级标题") == "（一）"
    assert detect_numbering_prefix("1.三级标题") == "1."
    assert detect_numbering_prefix("（1）四级标题") == "（1）"
    assert detect_numbering_prefix("一是普通正文") == ""


def test_word_list_level_prefix_keeps_native_list_evidence_without_literal_numbering() -> None:
    paragraph = Document().add_paragraph("自动列表标题")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "1")
    num_pr.append(ilvl)
    p_pr.append(num_pr)

    assert word_list_level_prefix(paragraph._p, paragraph.text) == "@lvl_1"
    assert word_list_level_prefix(paragraph._p, "1.已有字面编号") == ""


def test_heading_style_prefix_supports_chinese_and_english_heading_names() -> None:
    assert heading_style_prefix("Heading 1") == "@style_heading1"
    assert heading_style_prefix("标题 2") == "@style_heading2"
    assert heading_style_prefix("正文") == ""


def test_is_auto_numbered_item_reads_word_numbering_facts_only() -> None:
    """验证 Word 自动编号事实判断不把字面编号误作原生编号。"""
    assert is_auto_numbered_item(ParagraphFeatures(numbering_prefix="@lvl_0"))
    assert is_auto_numbered_item(ParagraphFeatures(numbering_prefix="@style_heading1"))
    assert not is_auto_numbered_item(ParagraphFeatures(numbering_prefix="一、"))
    assert not is_auto_numbered_item(None)


def test_repair_broken_rels_removes_null_relationship_from_temporary_copy(tmp_path) -> None:
    """损坏关系修复应返回临时副本，并只删除 Target=../NULL 的关系。"""
    source = tmp_path / "source.docx"
    broken = tmp_path / "broken.docx"
    document = Document()
    document.add_paragraph("测试正文")
    document.save(source)

    rel_name = "word/_rels/document.xml.rels"
    namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    with zipfile.ZipFile(source) as input_archive, zipfile.ZipFile(
        broken, "w", zipfile.ZIP_DEFLATED
    ) as output_archive:
        for item in input_archive.infolist():
            data = input_archive.read(item)
            if item.filename == rel_name:
                root = ET.fromstring(data)
                relationship = ET.SubElement(root, f"{{{namespace}}}Relationship")
                relationship.set("Id", "rIdBrokenNull")
                relationship.set(
                    "Type",
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
                )
                relationship.set("Target", "../NULL")
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            output_archive.writestr(item, data)

    repaired_path = repair_broken_rels(str(broken))

    assert repaired_path != str(broken)
    assert broken.exists()
    with zipfile.ZipFile(repaired_path) as archive:
        repaired_rels = ET.fromstring(archive.read(rel_name))
    assert all(
        (relationship.get("Target") or "").replace("\\", "/") != "../NULL"
        for relationship in repaired_rels
    )
