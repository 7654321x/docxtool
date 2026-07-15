import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

from docxtool.document.engine.style_catalog import ensure_document_styles
from docxtool.document.style_config import PageSettings, StyleRule


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

STYLE_IDS = [
    "DCT-Title",
    "DCT-DocumentNumber",
    "DCT-Author",
    "DCT-RoleName",
    "DCT-Recipient",
    "DCT-Heading1",
    "DCT-Heading2",
    "DCT-Heading3",
    "DCT-Heading4",
    "DCT-Body",
    "DCT-Responsibility",
    "DCT-Signature",
    "DCT-Date",
    "DCT-AttachmentNote",
    "DCT-AttachmentNoteItem",
    "DCT-AttachmentMark",
    "DCT-AttachmentTitle",
    "DCT-AttachmentBody",
]


def _styles_xml(path: Path) -> ET.Element:
    with zipfile.ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/styles.xml"))


def _style(root: ET.Element, style_id: str) -> ET.Element:
    matches = [
        style
        for style in root.findall("w:style", NS)
        if style.get(qn("w:styleId")) == style_id
    ]
    assert len(matches) == 1
    return matches[0]


def _ppr(root: ET.Element, style_id: str) -> ET.Element:
    ppr = _style(root, style_id).find("w:pPr", NS)
    assert ppr is not None
    return ppr


def _save_with_structural_styles(tmp_path: Path) -> tuple[Path, Document]:
    document = Document()
    document.add_paragraph("原文不应被样式目录修改")
    rules = [StyleRule.default_for_row(index) for index in range(24)]
    rules[5].spacing_before = 0.5
    rules[5].spacing_after = 0.25

    ensure_document_styles(document, rules, PageSettings(line_spacing_value=30.0))
    ensure_document_styles(document, rules, PageSettings(line_spacing_value=30.0))

    output = tmp_path / "structural-styles.docx"
    document.save(output)
    return output, document


def test_structural_style_ids_are_stable_and_idempotent(tmp_path: Path) -> None:
    output, document = _save_with_structural_styles(tmp_path)

    assert [paragraph.text for paragraph in document.paragraphs] == ["原文不应被样式目录修改"]

    root = _styles_xml(output)
    for style_id in STYLE_IDS:
        style = _style(root, style_id)
        assert style.get(qn("w:type")) == "paragraph"
        assert style.find("w:pPr", NS) is not None
        assert style.findall(".//w:numPr", NS) == []


def test_heading_styles_have_outline_without_catalog_keep_chain(tmp_path: Path) -> None:
    output, _document = _save_with_structural_styles(tmp_path)
    root = _styles_xml(output)

    for level, style_id in enumerate(["DCT-Heading1", "DCT-Heading2", "DCT-Heading3", "DCT-Heading4"]):
        ppr = _ppr(root, style_id)
        assert ppr.find("w:keepNext", NS) is None
        assert ppr.find("w:keepLines", NS) is None
        assert ppr.find("w:outlineLvl", NS).get(qn("w:val")) == str(level)

    title_ppr = _ppr(root, "DCT-Title")
    assert title_ppr.find("w:keepNext", NS) is not None
    assert title_ppr.find("w:keepLines", NS) is not None


def test_body_style_uses_justified_indent_and_exact_spacing(tmp_path: Path) -> None:
    output, _document = _save_with_structural_styles(tmp_path)
    body_ppr = _ppr(_styles_xml(output), "DCT-Body")

    assert body_ppr.find("w:jc", NS).get(qn("w:val")) == "both"

    indent = body_ppr.find("w:ind", NS)
    assert indent.get(qn("w:firstLineChars")) == "200"
    assert indent.get(qn("w:firstLine")) == "640"

    spacing = body_ppr.find("w:spacing", NS)
    assert spacing.get(qn("w:line")) == "600"
    assert spacing.get(qn("w:lineRule")) == "exact"
    assert spacing.get(qn("w:before")) == "300"
    assert spacing.get(qn("w:after")) == "150"
    assert spacing.get(qn("w:beforeLines")) == "50"
    assert spacing.get(qn("w:afterLines")) == "25"


def test_explicit_zero_spacing_does_not_fall_back_to_page_defaults(tmp_path: Path) -> None:
    document = Document()
    rules = [StyleRule.default_for_row(index) for index in range(24)]
    ensure_document_styles(
        document,
        rules,
        PageSettings(line_spacing_value=30.0, space_before_line=1.0, space_after_line=1.0),
    )

    output = tmp_path / "zero-spacing.docx"
    document.save(output)
    body_ppr = _ppr(_styles_xml(output), "DCT-Body")
    spacing = body_ppr.find("w:spacing", NS)

    assert spacing.get(qn("w:beforeLines")) == "0"
    assert spacing.get(qn("w:afterLines")) == "0"


def test_responsibility_style_is_left_aligned_without_distribution_or_indent(tmp_path: Path) -> None:
    output, _document = _save_with_structural_styles(tmp_path)
    ppr = _ppr(_styles_xml(output), "DCT-Responsibility")

    assert ppr.find("w:jc", NS).get(qn("w:val")) == "left"
    assert ppr.find("w:jc", NS).get(qn("w:val")) not in {"both", "distribute"}

    indent = ppr.find("w:ind", NS)
    assert indent.get(qn("w:firstLineChars")) == "0"
    assert indent.get(qn("w:firstLine")) == "0"
    assert ppr.find("w:numPr", NS) is None


def test_attachment_mark_and_title_keep_with_following_content(tmp_path: Path) -> None:
    output, _document = _save_with_structural_styles(tmp_path)
    root = _styles_xml(output)

    for style_id in ("DCT-AttachmentMark", "DCT-AttachmentTitle"):
        ppr = _ppr(root, style_id)
        assert ppr.find("w:keepNext", NS) is not None
        assert ppr.find("w:keepLines", NS) is not None
        assert ppr.find("w:numPr", NS) is None
