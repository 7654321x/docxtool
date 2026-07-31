from __future__ import annotations

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docxtool.document.importer import (
    ParagraphFeatures,
    SourceRun,
    _apply_segment_format_features,
    extract_features,
)


def _set_east_asia_font(run, name: str, ascii_name: str = "Arial") -> None:
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), ascii_name)
    fonts.set(qn("w:hAnsi"), ascii_name)


def test_effective_format_uses_direct_east_asia_font_and_preserves_boolean_unknown() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("中文 Title")
    _set_east_asia_font(run, "SimHei", "Calibri")
    run.font.size = Pt(16)

    features = extract_features(paragraph, 0)
    source_run = features.source_run_spans[0]

    assert source_run.east_asia_font_name == "SimHei"
    assert source_run.ascii_font_name == "Calibri"
    assert source_run.font_size_pt == 16.0
    assert source_run.bold is None
    assert source_run.explicit is True
    assert "direct_run" in source_run.format_sources
    assert features.segment_font_name_east_asia == "SimHei"
    assert features.segment_font_name_ascii == "Calibri"


def test_effective_format_resolves_paragraph_based_on_chain() -> None:
    document = Document()
    base = document.styles.add_style("Base Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    base.font.name = "Arial"
    base.font.size = Pt(14)
    base.font.bold = True
    derived = document.styles.add_style("Derived Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    derived.base_style = base
    paragraph = document.add_paragraph("继承格式正文", style=derived)

    source_run = extract_features(paragraph, 0).source_run_spans[0]

    assert source_run.font_size_pt == 14.0
    assert source_run.bold is True
    assert source_run.explicit is False
    assert source_run.inherited is True
    assert "based_on" in source_run.format_sources


def test_direct_run_override_wins_over_character_and_paragraph_style() -> None:
    document = Document()
    paragraph_style = document.styles.add_style("Bold Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    paragraph_style.font.bold = True
    character_style = document.styles.add_style("Unbold Character", WD_STYLE_TYPE.CHARACTER)
    character_style.font.bold = False
    paragraph = document.add_paragraph(style=paragraph_style)
    run = paragraph.add_run("直接覆盖")
    run.style = character_style
    run.font.bold = True

    source_run = extract_features(paragraph, 0).source_run_spans[0]

    assert source_run.bold is True
    assert source_run.explicit is True
    assert source_run.inherited is False
    assert "direct_run" in source_run.format_sources


def test_effective_format_reads_doc_defaults_and_theme_font_aliases() -> None:
    document = Document()
    defaults = document.styles.element.find(qn("w:docDefaults"))
    default = defaults.find(qn("w:rPrDefault"))
    rpr = default.find(qn("w:rPr"))
    bold = OxmlElement("w:b")
    rpr.append(bold)
    run = document.add_paragraph().add_run("默认格式")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:asciiTheme"), "minorHAnsi")
    run._element.get_or_add_rPr().append(fonts)

    source_run = extract_features(run._parent, 0).source_run_spans[0]

    assert source_run.bold is True
    assert "doc_defaults" in source_run.format_sources
    # Office's default theme has a minor Latin face; the resolver must use it
    # when no direct ASCII/HAnsi name is present.
    assert source_run.ascii_font_name


def test_segment_format_coverage_marks_unmapped_run_as_review() -> None:
    parent = ParagraphFeatures(source_physical_text="标题正文")
    parent.source_run_spans = (
        SourceRun(
            start=0,
            end=2,
            font_name="SimHei",
            east_asia_font_name="SimHei",
            ascii_font_name="Arial",
            font_size_pt=16.0,
            bold=True,
            italic=None,
            underline=None,
            explicit=True,
            inherited=False,
            known=True,
            format_sources=("direct_run",),
        ),
        SourceRun(
            start=2,
            end=4,
            font_name="",
            east_asia_font_name=None,
            ascii_font_name=None,
            font_size_pt=None,
            bold=None,
            italic=None,
            underline=None,
            explicit=False,
            inherited=False,
            known=False,
            format_warnings=("FORMAT_UNKNOWN",),
        ),
    )
    child = ParagraphFeatures(source_physical_text="标题正文")

    _apply_segment_format_features(child, parent, 0, 4)

    assert child.segment_visible_char_count == 4
    assert child.segment_mapped_format_char_count == 2
    assert child.segment_format_coverage_ratio == 0.5
    assert child.segment_format_status == "review"
    assert "PARTIAL_RUN_FORMAT_COVERAGE" in child.segment_format_warnings


def test_one_run_crossing_two_segments_is_weighted_by_intersection() -> None:
    parent = ParagraphFeatures(source_physical_text="一、标题。正文内容")
    parent.source_run_spans = (
        SourceRun(
            start=0,
            end=len(parent.source_physical_text),
            font_name="SimSun",
            east_asia_font_name="SimSun",
            ascii_font_name="Times New Roman",
            font_size_pt=12.0,
            bold=False,
            italic=False,
            underline=False,
            explicit=False,
            inherited=True,
            known=True,
            format_sources=("paragraph_style",),
        ),
    )
    heading = ParagraphFeatures(source_physical_text=parent.source_physical_text)
    body = ParagraphFeatures(source_physical_text=parent.source_physical_text)

    _apply_segment_format_features(heading, parent, 0, 5)
    _apply_segment_format_features(body, parent, 5, len(parent.source_physical_text))

    assert heading.segment_visible_char_count == 5
    assert body.segment_visible_char_count == 4
    assert heading.segment_format_coverage_ratio == 1.0
    assert body.segment_format_coverage_ratio == 1.0
    assert heading.segment_font_name_east_asia == body.segment_font_name_east_asia == "SimSun"
