from docx import Document
from docx.oxml.ns import qn

from docxtool.document.engine.header_footer import apply_header_footer
from docxtool.document.style_config import StyleRule


def test_apply_header_footer_writes_odd_and_even_page_fields() -> None:
    """页脚兼容入口应写入奇偶页不同和 PAGE 字段。"""
    document = Document()
    rule = StyleRule.default_for_row(8)

    apply_header_footer(document, rule)

    even_and_odd = document.settings._element.find(qn("w:evenAndOddHeaders"))
    assert even_and_odd is not None
    assert even_and_odd.get(qn("w:val")) == "1"

    section = document.sections[0]
    odd_xml = section.footer.paragraphs[0]._element
    even_xml = section.even_page_footer.paragraphs[0]._element
    assert odd_xml.find(".//" + qn("w:instrText")).text == " PAGE "
    assert even_xml.find(".//" + qn("w:instrText")).text == " PAGE "
    assert odd_xml.find(".//" + qn("w:ind")).get(qn("w:rightChars")) == "100"
    assert even_xml.find(".//" + qn("w:ind")).get(qn("w:leftChars")) == "100"
