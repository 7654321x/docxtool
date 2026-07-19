import copy
import hashlib
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

from docxtool.document.engine.core import export_doc
from docxtool.document.engine.letterhead import (
    apply_letterhead,
    detect_letterhead,
)
from docxtool.document.importer import DocumentData, ParagraphData, ParagraphFeatures, DocxImporter
from docxtool.document.letterhead_config import default_letterhead_config
from docxtool.document.style_config import PageSettings, StyleRule
from docxtool.security import validate_docx_integrity


def rules():
    return [StyleRule.default_for_row(index) for index in range(24)]


def config(**changes):
    value = default_letterhead_config()
    value.update(
        {
            "enabled": True,
            "agencies": [
                {"id": "agency-1", "name": "测试机关", "short_name": "", "role": "sponsor", "order": 1}
            ],
            "document_number": {"agency_code": "测发", "year": 2026, "sequence": 12},
        }
    )
    value.update(changes)
    return value


def data(source: Path | None = None):
    return DocumentData(
        filepath=str(source or "generated.docx"),
        paragraphs=[
            ParagraphData("公文标题", "title", "公文标题", ParagraphFeatures()),
            ParagraphData("正文内容。", "body", "正文内容。", ParagraphFeatures()),
        ],
    )


def export(tmp_path, letterhead, name="output.docx", source=None, page_settings=None):
    output = tmp_path / name
    stats = export_doc(
        data(source), rules(), page_settings or PageSettings(), str(output),
        page_number_enabled=False, letterhead_options=letterhead,
    )
    validate_docx_integrity(output)
    return output, stats


def style_ids(document):
    return [paragraph.style.style_id for paragraph in document.paragraphs]


def spacing_value(paragraph, name):
    spacing = paragraph._p.get_or_add_pPr().find(qn("w:spacing"))
    return spacing.get(qn(f"w:{name}")) if spacing is not None else None


def run_font_value(run, name):
    fonts = run._r.get_or_add_rPr().find(qn("w:rFonts"))
    return fonts.get(qn(f"w:{name}")) if fonts is not None else None


def test_missing_null_and_disabled_do_not_generate(tmp_path):
    for index, value in enumerate((None, {**default_letterhead_config(), "enabled": False})):
        output, stats = export(tmp_path, value, f"disabled-{index}.docx")
        document = Document(output)
        assert "DCT-LetterheadMark" not in style_ids(document)
        assert [paragraph.text for paragraph in document.paragraphs if paragraph.text] == ["公文标题", "正文内容。"]
        assert stats["letterhead_action"] == "preserved-disabled"


def test_single_mark_document_number_separator_and_title_spacing(tmp_path):
    output, stats = export(tmp_path, config())
    document = Document(output)
    assert [paragraph.text for paragraph in document.paragraphs[:9]] == [
        "", "", "", "测试机关文件", "", "", "测发〔2026〕12号", "", "公文标题"
    ]
    assert style_ids(document)[:9] == [
        "DCT-LetterheadSpacer", "DCT-LetterheadSpacer", "DCT-LetterheadSpacer",
        "DCT-LetterheadMark",
        "DCT-LetterheadSpacer", "DCT-LetterheadSpacer",
        "DCT-DocumentNumber", "DCT-LetterheadSeparator", "DCT-Title",
    ]
    assert stats["letterhead_action"] == "generated"
    mark = document.paragraphs[3]
    assert mark.runs[0].font.size.pt == 32
    assert mark.runs[0]._r.rPr.find(qn("w:w")) is None
    assert spacing_value(mark, "beforeLines") == "0"
    assert spacing_value(mark, "afterLines") == "0"
    assert spacing_value(mark, "before") is None
    assert spacing_value(mark, "after") is None
    for spacer in document.paragraphs[:3] + document.paragraphs[4:6]:
        assert spacer.style.style_id == "DCT-LetterheadSpacer"
        assert spacing_value(spacer, "beforeLines") == "0"
        assert spacing_value(spacer, "afterLines") == "0"
    assert spacing_value(document.paragraphs[6], "beforeLines") == "0"
    assert spacing_value(document.paragraphs[6], "afterLines") == "0"
    assert spacing_value(document.paragraphs[7], "afterLines") == "200"
    assert spacing_value(document.paragraphs[7], "after") is None
    assert document.paragraphs[7].paragraph_format.left_indent.pt == 0
    assert document.paragraphs[7].paragraph_format.right_indent.pt == 0
    empty_styles = [p.style.style_id for p in document.paragraphs if not p.text]
    assert empty_styles == [
        "DCT-LetterheadSpacer", "DCT-LetterheadSpacer", "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer", "DCT-LetterheadSpacer", "DCT-LetterheadSeparator",
    ]
    assert document.paragraphs[9].text == "正文内容。"
    assert document.paragraphs[9].paragraph_format.page_break_before is not True
    with ZipFile(output) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        styles_xml = etree.fromstring(archive.read("word/styles.xml"))
        custom_xml = archive.read("docProps/custom.xml").decode("utf-8")
        bottom = document_xml.find(".//" + qn("w:pBdr") + "/" + qn("w:bottom"))
        assert bottom.get(qn("w:val")) == "single"
        assert bottom.get(qn("w:color")) == "FF0000"
        assert bottom.get(qn("w:sz")) == "12"
        assert bottom.get(qn("w:space")) == "0"
        assert len(document_xml.findall(".//" + qn("w:pBdr") + "/" + qn("w:bottom"))) == 1
        assert not document_xml.findall(".//" + qn("w:drawing"))
        assert not document_xml.findall(".//" + qn("w:pict"))
        assert not document_xml.findall(".//" + qn("w:object"))
        assert "DocxtoolLetterheadVersion" in custom_xml
        assert not any("header" in name and name.endswith(".xml") for name in archive.namelist())
        for style_id in (
            "DCT-LetterheadSpacer", "DCT-LetterheadMark", "DCT-DocumentNumber",
            "DCT-SignerLine", "DCT-LetterheadSeparator",
        ):
            assert styles_xml.find(f".//{qn('w:style')}[@{qn('w:styleId')}='{style_id}']") is not None
        assert b"------" not in archive.read("word/document.xml")
    assert round(document.paragraphs[7].paragraph_format.space_before.cm, 1) == 0.4
    assert spacing_value(document.paragraphs[7], "beforeLines") is None


def test_enabled_letterhead_does_not_override_document_page_layout(tmp_path):
    custom = PageSettings(
        page_width_cm=20,
        page_height_cm=28,
        margin_top_cm=1,
        margin_bottom_cm=1.2,
        margin_left_cm=1.4,
        margin_right_cm=1.6,
    )
    enabled, _ = export(tmp_path, config(), "custom-enabled-page.docx", page_settings=custom)
    section = Document(enabled).sections[0]
    assert round(section.page_width.mm) == 200
    assert round(section.page_height.mm) == 280
    assert round(section.top_margin.mm) == 10
    assert round(section.bottom_margin.mm) == 12
    assert round(section.left_margin.mm) == 14
    assert round(section.right_margin.mm) == 16

    disabled_config = {**default_letterhead_config(), "enabled": False}
    disabled, _ = export(
        tmp_path,
        disabled_config,
        "custom-page.docx",
        page_settings=custom,
    )
    disabled_section = Document(disabled).sections[0]
    assert round(disabled_section.page_width.mm) == 200
    assert round(disabled_section.top_margin.mm) == 10


def test_agency_only_and_name_ending_in_document_are_not_duplicated(tmp_path):
    output, _ = export(tmp_path, config(mark_display_mode="agency_only"), "agency-only.docx")
    assert next(
        p.text for p in Document(output).paragraphs if p.style.style_id == "DCT-LetterheadMark"
    ) == "测试机关"
    ending = config(agencies=[{"id": "agency-1", "name": "测试机关文件", "short_name": "", "role": "sponsor", "order": 1}])
    output2, _ = export(tmp_path, ending, "ending.docx")
    assert next(
        p.text for p in Document(output2).paragraphs if p.style.style_id == "DCT-LetterheadMark"
    ) == "测试机关文件"


def test_upward_multiple_signers_use_separate_runs_and_tabs(tmp_path):
    signers = [
        {"id": "signer-1", "agency_id": "agency-1", "name": "张三", "label": "签发人", "order": 1},
        {"id": "signer-2", "agency_id": "agency-1", "name": "李四", "label": "签发人", "order": 2},
        {"id": "signer-3", "agency_id": "agency-1", "name": "王五", "label": "签发人", "order": 3},
    ]
    output, _ = export(tmp_path, config(document_direction="upward", signers=signers), "upward.docx")
    document = Document(output)
    signer_paragraphs = [p for p in document.paragraphs if p.style.style_id == "DCT-SignerLine"]
    assert len(signer_paragraphs) == 1
    assert [p.text for p in signer_paragraphs] == ["\t签发人：张三\t签发人：李四"]
    assert signer_paragraphs[0].runs[1].text == "签发人："
    assert signer_paragraphs[0].runs[2].text == "张三"
    assert run_font_value(signer_paragraphs[0].runs[2], "eastAsia") == "楷体_GB2312"
    expected_right = Cm(15.6) - Pt(16)
    expected_positions = [
        expected_right - Cm(4.6),
        expected_right,
    ]
    actual_positions = [
        stop.position for stop in signer_paragraphs[0].paragraph_format.tab_stops
    ]
    assert all(
        abs(actual - expected) <= 635
        for actual, expected in zip(actual_positions, expected_positions, strict=True)
    )
    assert all(
        stop.alignment == WD_TAB_ALIGNMENT.RIGHT
        for stop in signer_paragraphs[0].paragraph_format.tab_stops
    )
    number_paragraph = next(p for p in document.paragraphs if p.style.style_id == "DCT-DocumentNumber")
    assert number_paragraph.text == "测发〔2026〕12号\t签发人：王五"
    assert number_paragraph.paragraph_format.left_indent.pt == 16
    final_tab_stop = list(number_paragraph.paragraph_format.tab_stops)[0]
    assert abs(final_tab_stop.position - expected_positions[0]) <= 635
    assert final_tab_stop.alignment == WD_TAB_ALIGNMENT.RIGHT
    assert spacing_value(signer_paragraphs[0], "beforeLines") == "0"
    assert spacing_value(signer_paragraphs[0], "afterLines") == "0"
    assert spacing_value(number_paragraph, "beforeLines") == "0"
    assert spacing_value(number_paragraph, "afterLines") == "0"


def test_upward_single_signer_shares_document_number_line(tmp_path):
    signers = [
        {"id": "signer-1", "agency_id": "agency-1", "name": "张三", "label": "签发人", "order": 1}
    ]
    output, _ = export(
        tmp_path,
        config(document_direction="upward", signers=signers),
        "upward-single.docx",
    )
    document = Document(output)
    number_paragraph = next(
        p for p in document.paragraphs if p.style.style_id == "DCT-DocumentNumber"
    )
    assert number_paragraph.text == "测发〔2026〕12号\t签发人：张三"
    assert number_paragraph.paragraph_format.left_indent.pt == 16
    assert spacing_value(number_paragraph, "beforeLines") == "0"
    assert spacing_value(number_paragraph, "afterLines") == "0"
    signer_label = next(run for run in number_paragraph.runs if run.text == "签发人：")
    signer_name = next(run for run in number_paragraph.runs if run.text == "张三")
    assert run_font_value(signer_label, "eastAsia") == "仿宋_GB2312"
    assert run_font_value(signer_name, "eastAsia") == "楷体_GB2312"
    number_runs = [run for run in number_paragraph.runs if run.text in {"2026", "12"}]
    assert len(number_runs) == 2
    assert all(run_font_value(run, "ascii") == "Times New Roman" for run in number_runs)
    assert all(run_font_value(run, "hAnsi") == "Times New Roman" for run in number_runs)
    assert "  " not in number_paragraph.text
    tab_stop = list(number_paragraph.paragraph_format.tab_stops)[0]
    assert abs(tab_stop.position - (Cm(15.6) - Pt(16))) <= 635
    assert tab_stop.alignment == WD_TAB_ALIGNMENT.RIGHT


def test_title_recipient_and_body_follow_letterhead_without_page_breaks(tmp_path):
    document_data = DocumentData(
        filepath="generated.docx",
        paragraphs=[
            ParagraphData("公文标题", "title", "公文标题", ParagraphFeatures()),
            ParagraphData("测试机关：", "addressing", "测试机关：", ParagraphFeatures()),
            ParagraphData("正文内容。", "body", "正文内容。", ParagraphFeatures()),
        ],
    )
    output = tmp_path / "title-recipient-body.docx"
    export_doc(
        document_data,
        rules(),
        PageSettings(),
        str(output),
        page_number_enabled=False,
        letterhead_options=config(),
    )
    validate_docx_integrity(output)
    document = Document(output)
    assert style_ids(document) == [
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadMark",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-DocumentNumber",
        "DCT-LetterheadSeparator",
        "DCT-Title",
        "DCT-Recipient",
        "DCT-Body",
    ]
    assert [p.text for p in document.paragraphs[8:]] == ["公文标题", "测试机关：", "正文内容。"]
    recipient = document.paragraphs[9]
    assert recipient._p.pPr.find(qn("w:ind")) is None
    recipient_style_indent = recipient.style.element.pPr.find(qn("w:ind"))
    assert recipient_style_indent.get(qn("w:firstLineChars")) == "0"
    assert recipient_style_indent.get(qn("w:firstLine")) == "0"
    assert all(p.paragraph_format.page_break_before is not True for p in document.paragraphs[8:])


def test_downward_and_parallel_preserve_configured_signers_without_rendering_them(tmp_path):
    signers = [
        {"id": "signer-1", "agency_id": "agency-1", "name": "张三", "label": "签发人", "order": 1}
    ]
    for direction in ("downward", "parallel"):
        output, _ = export(
            tmp_path,
            config(document_direction=direction, signers=signers),
            f"{direction}.docx",
        )
        document = Document(output)
        assert "DCT-SignerLine" not in style_ids(document)
        number_paragraph = next(
            p for p in document.paragraphs if p.style.style_id == "DCT-DocumentNumber"
        )
        assert number_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_joint_all_and_sponsor_only_preserve_sponsor_order_and_number(tmp_path):
    agencies = [
        {"id": "agency-2", "name": "联合机关乙", "short_name": "", "role": "joint", "order": 1},
        {"id": "agency-1", "name": "主办机关甲", "short_name": "", "role": "sponsor", "order": 2},
        {"id": "agency-3", "name": "联合机关丙", "short_name": "", "role": "joint", "order": 3},
    ]
    joint = config(issuance_mode="joint", agencies=agencies)
    output, _ = export(tmp_path, joint, "joint.docx")
    document = Document(output)
    marks = [p.text for p in document.paragraphs if p.style.style_id == "DCT-LetterheadMark" and p.text]
    assert marks == ["主办机关甲", "联合机关乙\t文件", "联合机关丙"]
    assert "测发〔2026〕12号" in [p.text for p in document.paragraphs]

    sponsor_only = copy.deepcopy(joint)
    sponsor_only["joint_mark_scope"] = "sponsor_only"
    output2, _ = export(tmp_path, sponsor_only, "sponsor-only.docx")
    marks2 = [
        p.text for p in Document(output2).paragraphs
        if p.style.style_id == "DCT-LetterheadMark" and p.text
    ]
    assert marks2 == ["主办机关甲文件"]


def test_managed_output_is_detected_and_reprocessing_is_idempotent(tmp_path):
    first, _ = export(tmp_path, config(), "first.docx")
    first_data = DocxImporter().load(str(first), rules(), features={})
    assert first_data.letterhead_detection.status == "managed"
    second = tmp_path / "second.docx"
    second_stats = export_doc(
        first_data, rules(), PageSettings(), str(second),
        page_number_enabled=False, letterhead_options=config(),
    )
    assert second_stats["letterhead_action"] == "replaced"
    assert style_ids(Document(second)).count("DCT-LetterheadMark") == 1
    assert detect_letterhead(Document(second)).status == "managed"

    replace = config(replace_managed=True, document_number={"agency_code": "测发", "year": 2026, "sequence": 99})
    replaced = tmp_path / "replaced.docx"
    replaced_stats = export_doc(
        first_data, rules(), PageSettings(), str(replaced),
        page_number_enabled=False, letterhead_options=replace,
    )
    assert replaced_stats["letterhead_action"] == "replaced"
    assert [p.text for p in Document(replaced).paragraphs].count("测发〔2026〕99号") == 1


def _external_document(path: Path):
    document = Document()
    mark = document.add_paragraph()
    mark.alignment = 1
    mark_run = mark.add_run("测试机关文件")
    mark_run.font.color.rgb = None
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    mark_run._element.get_or_add_rPr().append(color)
    document.add_paragraph("测发〔2026〕3号")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)
    title = document.add_paragraph("外部公文标题")
    # Random title formatting must not widen the protected letterhead prefix.
    title.alignment = 1
    title.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    title.runs[0].font.size = Pt(32)
    title_borders = OxmlElement("w:pBdr")
    title_bottom = OxmlElement("w:bottom")
    title_bottom.set(qn("w:val"), "single")
    title_bottom.set(qn("w:color"), "FF0000")
    title_borders.append(title_bottom)
    title._p.get_or_add_pPr().append(title_borders)
    document.add_paragraph("正文内容。")
    document.save(path)


def test_enabled_letterhead_replaces_external_letterhead(tmp_path):
    source = tmp_path / "external.docx"
    _external_document(source)
    before_hash = hashlib.sha256(source.read_bytes()).hexdigest()
    imported = DocxImporter().load(str(source), rules(), features={})
    assert imported.letterhead_detection.status == "recognized_external"
    ordinary = [item for item in imported.paragraphs if not item.type_id.startswith("__")]
    assert [(item.type_id, item.text) for item in ordinary[:2]] == [
        ("title", "外部公文标题"),
        ("body", "正文内容。"),
    ]
    output = tmp_path / "external-output.docx"
    stats = export_doc(
        imported, rules(), PageSettings(), str(output),
        page_number_enabled=False, letterhead_options=config(),
    )
    assert stats["letterhead_action"] == "replaced"
    assert stats["compatibility_warnings"] == []
    assert [p.text for p in Document(output).paragraphs[:8]] == [
        "", "", "", "测试机关文件", "", "", "测发〔2026〕12号", "",
    ]
    assert hashlib.sha256(source.read_bytes()).hexdigest() == before_hash


def test_single_leading_drawing_is_not_enough_to_replace_content():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph._p.append(OxmlElement("w:drawing"))
    document.add_paragraph("公文标题")
    detection = detect_letterhead(document)
    assert detection.status == "none"
    result = apply_letterhead(document, config(), detection=detection, rules=rules(), settings=PageSettings())
    assert result.action == "generated"
    assert result.warnings == []
    assert "DCT-LetterheadMark" in style_ids(document)
    assert any(
        paragraph._p.find(".//" + qn("w:drawing")) is not None
        for paragraph in document.paragraphs
    )


def test_zero_size_drawing_and_captioned_image_are_not_letterhead_signals():
    zero_size = Document()
    paragraph = zero_size.add_paragraph("（三）普通段落。")
    drawing = OxmlElement("w:drawing")
    extent = OxmlElement("wp:extent")
    extent.set("cx", "100")
    extent.set("cy", "0")
    drawing.append(extent)
    paragraph._p.append(drawing)
    assert detect_letterhead(zero_size).status == "none"

    captioned = Document()
    image_paragraph = captioned.add_paragraph()
    drawing = OxmlElement("w:drawing")
    extent = OxmlElement("wp:extent")
    extent.set("cx", "100")
    extent.set("cy", "100")
    drawing.append(extent)
    image_paragraph._p.append(drawing)
    captioned.add_paragraph("图2结构示意图")
    assert detect_letterhead(captioned).status == "none"


def test_document_number_reference_in_body_does_not_block_letterhead_generation():
    document = Document()
    document.add_paragraph("公文标题")
    reference = (
        "按照《中共四川省纪委机关、中共四川省委组织部关于开好2025年度县以上党和国家机关"
        "党员领导干部民主生活会的通知》（川组通〔2025〕51号）要求，形成如下材料。"
    )
    document.add_paragraph(reference)
    original_text = [paragraph.text for paragraph in document.paragraphs]

    detection = detect_letterhead(document)
    assert detection.status == "none"

    result = apply_letterhead(
        document,
        config(),
        detection=detection,
        rules=rules(),
        settings=PageSettings(),
    )

    assert result.action == "generated"
    assert result.warnings == []
    assert [paragraph.text for paragraph in document.paragraphs[-2:]] == original_text
    assert style_ids(document)[:8] == [
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadMark",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-DocumentNumber",
        "DCT-LetterheadSeparator",
    ]


def test_enabled_letterhead_replaces_unknown_complex_prefix_and_keeps_input_unchanged(tmp_path):
    source = tmp_path / "unknown.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph._p.append(OxmlElement("w:drawing"))
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)
    document.add_paragraph("公文标题")
    document.add_paragraph("正文内容。")
    document.save(source)
    before_hash = hashlib.sha256(source.read_bytes()).hexdigest()

    imported = DocxImporter().load(str(source), rules(), features={})
    assert imported.letterhead_detection.status == "unknown"
    output = tmp_path / "unknown-output.docx"
    stats = export_doc(
        imported,
        rules(),
        PageSettings(),
        str(output),
        page_number_enabled=False,
        letterhead_options=config(),
    )

    assert stats["letterhead_action"] == "replaced"
    assert stats["compatibility_warnings"] == []
    assert [p.style.style_id for p in Document(output).paragraphs[:4]] == [
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadSpacer",
        "DCT-LetterheadMark",
    ]
    assert not any(
        paragraph._p.find(".//" + qn("w:drawing")) is not None
        for paragraph in Document(output).paragraphs
    )
    assert hashlib.sha256(source.read_bytes()).hexdigest() == before_hash


def test_disabled_letterhead_preserves_external_and_unknown_blocks(tmp_path):
    disabled = {**default_letterhead_config(), "enabled": False}

    external = tmp_path / "external-disabled.docx"
    _external_document(external)
    external_data = DocxImporter().load(str(external), rules(), features={})
    external_output = tmp_path / "external-disabled-output.docx"
    external_stats = export_doc(
        external_data,
        rules(),
        PageSettings(),
        str(external_output),
        page_number_enabled=False,
        letterhead_options=disabled,
    )
    assert external_stats["letterhead_action"] == "preserved-disabled"
    assert [p.text for p in Document(external_output).paragraphs[:3]] == [
        "测试机关文件",
        "测发〔2026〕3号",
        "",
    ]

    unknown = tmp_path / "unknown-disabled.docx"
    unknown_document = Document()
    drawing_paragraph = unknown_document.add_paragraph()
    drawing_paragraph._p.append(OxmlElement("w:drawing"))
    unknown_separator = unknown_document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    unknown_separator._p.get_or_add_pPr().append(borders)
    unknown_document.add_paragraph("公文标题")
    unknown_document.save(unknown)
    unknown_data = DocxImporter().load(str(unknown), rules(), features={})
    unknown_output = tmp_path / "unknown-disabled-output.docx"
    unknown_stats = export_doc(
        unknown_data,
        rules(),
        PageSettings(),
        str(unknown_output),
        page_number_enabled=False,
        letterhead_options=disabled,
    )
    assert unknown_stats["letterhead_action"] == "preserved-disabled"
    assert Document(unknown_output).paragraphs[0]._p.find(".//" + qn("w:drawing")) is not None


def test_random_red_text_and_body_drawing_do_not_create_letterhead_block():
    document = Document()
    red = document.add_paragraph("测试")
    red.alignment = WD_ALIGN_PARAGRAPH.CENTER
    red_run = red.runs[0]
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    red_run._element.get_or_add_rPr().append(color)
    document.add_paragraph("一、一级标题")
    image = document.add_paragraph()
    drawing = OxmlElement("w:drawing")
    extent = OxmlElement("wp:extent")
    extent.set("cx", "100")
    extent.set("cy", "100")
    drawing.append(extent)
    image._p.append(drawing)

    detection = detect_letterhead(document)

    assert detection.status == "none"
    assert detection.protected_body_indexes == ()


def test_external_detection_is_bounded_at_separator_and_trailing_blanks():
    document = Document()
    document.add_paragraph()
    mark = document.add_paragraph()
    mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mark.add_run("测试机关文件")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    run._element.get_or_add_rPr().append(color)
    document.add_paragraph("测发〔2026〕3号")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)
    document.add_paragraph()
    document.add_paragraph("公文标题")
    later_red = document.add_paragraph("正文红字")
    later_color = OxmlElement("w:color")
    later_color.set(qn("w:val"), "FF0000")
    later_red.runs[0]._element.get_or_add_rPr().append(later_color)

    detection = detect_letterhead(document)

    assert detection.status == "recognized_external"
    assert detection.protected_body_indexes == (0, 1, 2, 3, 4)


def test_document_number_and_red_separator_are_sufficient_without_mark():
    document = Document()
    document.add_paragraph("市委办〔2026〕1号")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)
    document.add_paragraph("关于推进重点工作的通知")

    detection = detect_letterhead(document)

    assert detection.status == "recognized_external"
    assert detection.details == ("document-number", "separator", "bounded-prefix")
    assert detection.protected_body_indexes == (0, 1)


def test_document_number_without_red_separator_is_incomplete_letterhead():
    document = Document()
    document.add_paragraph("市委办〔2026〕1号")
    document.add_paragraph("关于推进重点工作的通知")

    detection = detect_letterhead(document)

    assert detection.status == "unknown"
    assert detection.details == ("incomplete-document-number", "bounded-prefix")
    assert detection.protected_body_indexes == (0,)


def test_leading_metadata_is_kept_with_complete_letterhead():
    document = Document()
    document.add_paragraph("000123")
    document.add_paragraph("机密★5年")
    document.add_paragraph("加急")
    document.add_paragraph("测发〔2026〕1号")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:color"), "E60012")
    borders.append(top)
    separator._p.get_or_add_pPr().append(borders)

    detection = detect_letterhead(document)

    assert detection.status == "recognized_external"
    assert detection.protected_body_indexes == (0, 1, 2, 3, 4)


def test_compatible_document_number_is_incomplete_not_standard():
    document = Document()
    document.add_paragraph("测发[2026]第1号")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "E60012")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)

    detection = detect_letterhead(document)

    assert detection.status == "unknown"
    assert detection.details == ("compatible-document-number", "bounded-prefix")


def test_joint_agency_prefix_and_signer_continuation_are_bounded():
    document = Document()
    document.add_paragraph("中共测试市委")
    document.add_paragraph("测试市人民政府文件")
    document.add_paragraph("测发〔2026〕2号")
    document.add_paragraph("签发人：张三")
    document.add_paragraph("李四")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)
    document.add_paragraph("关于推进重点工作的通知")

    detection = detect_letterhead(document)

    assert detection.status == "recognized_external"
    assert detection.protected_body_indexes == (0, 1, 2, 3, 4, 5)


def test_agency_file_mark_without_red_formatting_is_incomplete_letterhead():
    document = Document()
    document.add_paragraph("中共测试市委办公室文件")
    document.add_paragraph("关于推进重点工作的通知")

    detection = detect_letterhead(document)

    assert detection.status == "unknown"
    assert detection.details == ("incomplete-letterhead-mark", "following-title", "bounded-prefix")
    assert detection.protected_body_indexes == (0,)


def test_agency_file_text_without_following_title_is_not_letterhead():
    document = Document()
    document.add_paragraph("中共测试市委办公室文件")
    document.add_paragraph("各有关单位：")
    document.add_paragraph("现将有关事项安排如下。")

    detection = detect_letterhead(document)

    assert detection.status == "none"
    assert detection.protected_body_indexes == ()


def test_red_agency_file_text_without_following_title_is_not_letterhead():
    document = Document()
    paragraph = document.add_paragraph("中共测试市委办公室文件")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    document.add_paragraph("各有关单位：")
    document.add_paragraph("现将有关事项安排如下。")

    detection = detect_letterhead(document)

    assert detection.status == "none"


def test_agency_file_text_does_not_use_later_separator_past_body():
    document = Document()
    document.add_paragraph("中共测试市委办公室文件")
    body = document.add_paragraph("这是一行尚未添加句号的正文内容")
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)

    detection = detect_letterhead(document)

    assert detection.status == "none"


def test_agency_file_mark_accepts_visually_identifiable_material_title():
    document = Document()
    document.add_paragraph("中共测试市委办公室文件")
    title = document.add_paragraph("2026年度重点任务清单")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("一、总体安排")

    detection = detect_letterhead(document)

    assert detection.status == "unknown"
    assert detection.protected_body_indexes == (0,)


def test_title_ending_with_file_is_not_semantic_letterhead_mark():
    document = Document()
    document.add_paragraph("关于报送规范性文件")
    document.add_paragraph("正文内容")

    detection = detect_letterhead(document)

    assert detection.status == "none"


def test_red_title_starting_with_about_is_not_visual_letterhead_mark():
    document = Document()
    paragraph = document.add_paragraph("关于报送规范性文件")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    document.add_paragraph("正文内容。")

    detection = detect_letterhead(document)

    assert detection.status == "none"
