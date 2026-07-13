import os
import tempfile
import time
import unittest
import uuid
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse
from unittest.mock import patch

from docx import Document

import server


class ServerProductionControlsTest(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        root = Path(self.tmp.name)
        self.old_db = server._DB_PATH
        self.old_log_dir = server.LOG_DIR
        self.old_output_dir = server.OUTPUT_DIR
        self.old_admin_token = server.ADMIN_TOKEN
        self.old_proxy_secret = server.PROXY_SECRET
        self.old_allow_local_file_api = server.ALLOW_LOCAL_FILE_API
        self.old_bind_host = server.BIND_HOST
        self.old_frontend_origin = server.FRONTEND_ORIGIN
        server._DB_PATH = str(root / "stats.db")
        server.LOG_DIR = str(root / "logs")
        server.OUTPUT_DIR = str(root / "outputs")
        server.ADMIN_TOKEN = ""
        server.PROXY_SECRET = ""
        server.ALLOW_LOCAL_FILE_API = False
        server.BIND_HOST = "127.0.0.1"
        server.FRONTEND_ORIGIN = ""
        os.makedirs(server.LOG_DIR, exist_ok=True)
        os.makedirs(server.OUTPUT_DIR, exist_ok=True)
        server._sql_init()
        with server.TASKS_LOCK:
            server.TASKS.clear()
        with server.QUEUE_COND:
            server.TASK_QUEUE.clear()

    def tearDown(self):
        server._DB_PATH = self.old_db
        server.LOG_DIR = self.old_log_dir
        server.OUTPUT_DIR = self.old_output_dir
        server.ADMIN_TOKEN = self.old_admin_token
        server.PROXY_SECRET = self.old_proxy_secret
        server.ALLOW_LOCAL_FILE_API = self.old_allow_local_file_api
        server.BIND_HOST = self.old_bind_host
        server.FRONTEND_ORIGIN = self.old_frontend_origin
        with server.TASKS_LOCK:
            server.TASKS.clear()
        with server.QUEUE_COND:
            server.TASK_QUEUE.clear()
        self.tmp.cleanup()

    def test_file_ttl_is_24_hours(self):
        self.assertEqual(server.FILE_TTL, 86400)

    def test_default_secrets_are_not_embedded_in_source(self):
        self.assertEqual(server.DEFAULT_ADMIN_TOKEN, "")
        self.assertEqual(server.DEFAULT_PROXY_SECRET, "")

    def test_startup_urls_use_clean_monitor_url(self):
        urls = server._startup_urls()

        self.assertEqual(urls["tool"], "http://127.0.0.1:9527")
        self.assertEqual(urls["monitor"], "http://127.0.0.1:9527/monitor")
        self.assertEqual(urls["tunnel_command"], "cloudflared tunnel --url http://127.0.0.1:9527")

    def test_backend_binds_to_loopback_by_default(self):
        self.assertEqual(server.BIND_HOST, "127.0.0.1")
        self.assertEqual(server._server_bind_address(), ("127.0.0.1", server.PORT))

    def test_queue_position_reports_people_ahead(self):
        server._enqueue_task("task-a", "a.docx", "a.docx", "203.0.113.1", "ua")
        server._enqueue_task("task-b", "b.docx", "b.docx", "203.0.113.1", "ua")
        server._enqueue_task("task-c", "c.docx", "c.docx", "203.0.113.1", "ua")

        self.assertEqual(server._task_queue_info("task-a"), {
            "queue_position": 1,
            "queue_ahead": 0,
            "message": "排队中，前方还有 0 个任务",
        })
        self.assertEqual(server._task_queue_info("task-c"), {
            "queue_position": 3,
            "queue_ahead": 2,
            "message": "排队中，前方还有 2 个任务",
        })

    def test_task_temp_paths_live_under_project_runtime_dir(self):
        task_tmp = server._task_tmp_input_path("task-a", "测试 文件.docx")
        self.assertIn(os.path.join("runtime", "tmp"), task_tmp)
        self.assertTrue(task_tmp.endswith(os.path.join("task-a", "input.docx")))

    def test_spawn_subprocess_keeps_runtime_tmp_input_until_loaded(self):
        task_id = f"spawn-{uuid.uuid4()}"
        task_dir = Path(server.RUNTIME_TMP_DIR) / task_id
        input_path = task_dir / "input.docx"
        result = {}
        task_dir.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_paragraph("测试标题")
        doc.add_paragraph("这是一段用于真实 spawn 回归测试的正文。")
        doc.save(input_path)

        try:
            result = server._task_process_subprocess(
                task_id,
                str(input_path),
                "spawn-regression.docx",
                "127.0.0.1",
                "unittest",
            )

            self.assertTrue(input_path.exists())
            self.assertEqual(result.get("status"), "done", result.get("error"))
            self.assertTrue(Path(result["output_path"]).exists())
        finally:
            server._cleanup_task_tmp(task_id)
            for key in ("output_dir", "log_path"):
                path = result.get(key, "")
                if path:
                    server._cleanup_output_path(path)

    def test_enqueue_task_is_visible_in_monitor_immediately(self):
        queued_path = Path(server.OUTPUT_DIR) / "queued-input.docx"
        queued_path.write_bytes(b"PK queued")

        server._enqueue_task("task-a", str(queued_path), "queued.docx", "203.0.113.7", "ua")

        stats = server.get_sql_stats()
        self.assertEqual(stats["total"], 1)
        self.assertEqual(stats["recent"][0]["status"], "queued")
        self.assertEqual(stats["recent"][0]["filename"], "queued.docx")
        self.assertEqual(stats["top_ips"][0]["ip"], "203.0.113.7")

    def test_enqueue_task_full_does_not_leave_persisted_row(self):
        queued_path = Path(server.OUTPUT_DIR) / "queued-input.docx"
        queued_path.write_bytes(b"PK queued")
        old_max_queue = server.MAX_QUEUE
        server.MAX_QUEUE = 0
        try:
            with self.assertRaises(OverflowError):
                server._enqueue_task("task-a", str(queued_path), "queued.docx", "203.0.113.7", "ua")
        finally:
            server.MAX_QUEUE = old_max_queue

        with server._SQL_LOCK:
            conn = server._sql()
            row = conn.execute("SELECT 1 FROM tasks WHERE id=?", ("task-a",)).fetchone()
            conn.close()
        self.assertIsNone(row)
        self.assertNotIn("task-a", server.TASK_QUEUE)

    def test_completed_task_updates_queued_monitor_record(self):
        queued_path = Path(server.OUTPUT_DIR) / "queued-input.docx"
        queued_path.write_bytes(b"PK queued")
        server._enqueue_task("task-a", str(queued_path), "queued.docx", "203.0.113.7", "ua")

        server.log_sql(
            "task-a", "203.0.113.7", "ua", "queued.docx", 9,
            "NORMAL", 5, 1, 4, 1200, "done",
            log_filename="task-a.log", log_path=str(Path(server.LOG_DIR) / "task-a.log"),
        )

        stats = server.get_sql_stats()
        self.assertEqual(stats["total"], 1)
        self.assertEqual(stats["done"], 1)
        self.assertEqual(stats["recent"][0]["status"], "done")
        self.assertEqual(stats["recent"][0]["log_filename"], "task-a.log")

    def test_task_times_use_python_local_24_hour_clock(self):
        with patch.object(server, "_now_local", return_value="2026-06-02 13:24:00", create=True):
            server.log_sql(
                "task-a", "203.0.113.7", "ua", "done.docx", 9,
                "NORMAL", 5, 1, 4, 1200, "done",
            )

        with server._SQL_LOCK:
            conn = server._sql()
            row = conn.execute("SELECT created_at, done_at FROM tasks WHERE id=?", ("task-a",)).fetchone()
            conn.close()

        self.assertEqual(row["created_at"], "2026-06-02 13:24:00")
        self.assertEqual(row["done_at"], "2026-06-02 13:24:00")

    def test_now_local_uses_system_local_24_hour_clock(self):
        local_time = time.struct_time((2026, 6, 2, 13, 24, 0, 1, 153, -1))
        with patch.object(server.time, "localtime", return_value=local_time):
            self.assertEqual(server._now_local(), "2026-06-02 13:24:00")

    def test_parse_network_beijing_time_from_http_date_header(self):
        dt = server._parse_http_date_to_beijing("Tue, 02 Jun 2026 05:24:33 GMT")

        self.assertEqual(dt.strftime("%Y-%m-%d %H:%M:%S"), "2026-06-02 13:24:33")

    def test_startup_time_check_warns_when_system_time_differs_from_beijing_minute(self):
        network_time = datetime(2026, 6, 2, 13, 24, 33, tzinfo=timezone.utc)

        with patch.object(server, "_now_local", return_value="2026-06-02 01:24:05"), \
             patch.object(server, "_fetch_beijing_network_time", return_value=network_time):
            lines = server._startup_time_check_lines()

        self.assertIn("系统时间与北京网络时间不一致", lines[0])
        self.assertIn("系统时间为：2026-06-02 01:24", lines[1])
        self.assertIn("北京时间为：2026-06-02 13:24", lines[2])
        self.assertIn("sudo timedatectl set-timezone Asia/Shanghai", lines[3])
        self.assertIn("sudo timedatectl set-ntp true", lines[4])

    def test_queued_task_time_uses_python_local_24_hour_clock(self):
        with patch.object(server, "_now_local", return_value="2026-06-02 13:24:00", create=True):
            server.record_task_queued("task-a", "203.0.113.7", "ua", "queued.docx", 9)

        with server._SQL_LOCK:
            conn = server._sql()
            row = conn.execute("SELECT created_at, done_at FROM tasks WHERE id=?", ("task-a",)).fetchone()
            conn.close()

        self.assertEqual(row["created_at"], "2026-06-02 13:24:00")
        self.assertEqual(row["done_at"], "")

    def test_task_is_not_marked_done_before_stats_are_written(self):
        input_path = Path(server.OUTPUT_DIR) / "input.docx"
        input_path.write_bytes(b"PK input")
        observed_statuses = []

        class FakeParagraph:
            type_id = "body"

        class FakeDocData:
            doc_mode = "NORMAL"
            paragraphs = [FakeParagraph()]

        class FakeImporter:
            def load(self, _input_path, _rules):
                return FakeDocData()

        def fake_export_doc(_doc_data, _rules, _settings, output_path, numbered_bold_enabled=True):
            Path(output_path).write_bytes(b"PK output")

        def fake_log_sql(*_args, **_kwargs):
            with server.TASKS_LOCK:
                observed_statuses.append(server.TASKS["task-a"].get("status"))

        with server.TASKS_LOCK:
            server.TASKS["task-a"] = {"status": "processing"}

        with patch.object(server, "DocxImporter", FakeImporter), \
             patch.object(server, "export_doc", fake_export_doc), \
             patch.object(server, "log_sql", fake_log_sql), \
             patch.object(server.StyleRule, "from_config", return_value=None), \
             patch.object(server.PageSettings, "from_config", return_value=None):
            server._process_task("task-a", str(input_path), "input.docx", "203.0.113.7", "ua")

        self.assertEqual(observed_statuses, ["processing"])
        self.assertEqual(server._public_task_state("task-a")["status"], "done")

    def test_processing_task_has_no_queue_ahead(self):
        with server.TASKS_LOCK:
            server.TASKS["task-a"] = {"status": "processing"}

        self.assertEqual(server._public_task_state("task-a")["queue_ahead"], 0)
        self.assertEqual(server._public_task_state("task-a")["message"], "正在排版")

    def test_startup_recovery_marks_inflight_tasks_interrupted(self):
        with server._SQL_LOCK:
            conn = server._sql()
            conn.execute(
                "INSERT INTO tasks (id, ip, ua, filename, status, created_at, done_at) VALUES (?,?,?,?,?,?,?)",
                ("task-a", "203.0.113.7", "ua", "queued.docx", "processing", "2026-06-02 13:24:00", ""),
            )
            conn.commit()
            conn.close()

        recovered = server._recover_inflight_tasks_on_startup()
        self.assertEqual(recovered, 1)
        with server._SQL_LOCK:
            conn = server._sql()
            row = conn.execute("SELECT status, error, done_at FROM tasks WHERE id=?", ("task-a",)).fetchone()
            conn.close()
        self.assertEqual(row["status"], "interrupted")
        self.assertEqual(row["error"], "服务重启后任务中断")
        self.assertNotEqual(row["done_at"], "")

    def test_admin_token_accepts_query_header_or_cookie(self):
        self.assertFalse(server._admin_authorized(urlparse("/monitor"), {}, ""))

        server.ADMIN_TOKEN = "secret"

        self.assertFalse(server._admin_authorized(urlparse("/monitor"), {}, ""))
        self.assertTrue(server._admin_authorized(urlparse("/monitor?token=secret"), {}, ""))
        self.assertTrue(server._admin_authorized(urlparse("/monitor"), {"X-Admin-Token": "secret"}, ""))
        self.assertTrue(server._admin_authorized(urlparse("/monitor"), {}, "admin_token=secret"))

    def test_admin_session_cookie_round_trip(self):
        session = server._create_admin_session("ua", "203.0.113.7")
        cookie = server._admin_cookie_header(session["session_id"])
        headers = {"Cookie": cookie}

        from_headers = server._admin_session_from_headers(headers)
        self.assertEqual(from_headers["session_id"], session["session_id"])
        self.assertEqual(from_headers["csrf_token"], session["csrf_token"])
        self.assertIn(server.ADMIN_SESSION_COOKIE, cookie)

    def test_file_api_rejects_direct_requests_when_proxy_secret_is_unset(self):
        server.PROXY_SECRET = ""

        self.assertFalse(server._file_api_authorized({}))

    def test_file_api_requires_proxy_secret_when_configured(self):
        server.PROXY_SECRET = "proxy-secret"

        self.assertFalse(server._file_api_authorized({}))
        self.assertFalse(server._file_api_authorized({"X-Proxy-Secret": "wrong"}))
        self.assertTrue(server._file_api_authorized({"X-Proxy-Secret": "proxy-secret"}))

    def test_file_api_rejects_client_controlled_localhost_host(self):
        server.PROXY_SECRET = "proxy-secret"

        self.assertFalse(server._file_api_authorized({"Host": "127.0.0.1:9527"}))
        self.assertFalse(server._file_api_authorized({"Host": "localhost:9527"}, ("203.0.113.10", 41234)))
        self.assertFalse(server._file_api_authorized({"Host": "[::1]:9527"}))
        self.assertFalse(server._file_api_authorized({"Host": "example.trycloudflare.com"}))

    def test_file_api_rejects_loopback_proxy_without_secret_in_production(self):
        server.PROXY_SECRET = "proxy-secret"
        server.ALLOW_LOCAL_FILE_API = True
        server.FRONTEND_ORIGIN = "https://example.pages.dev"

        self.assertFalse(server._file_api_authorized({}, ("127.0.0.1", 51234)))

    def test_file_api_allows_loopback_only_when_local_development_enabled(self):
        server.PROXY_SECRET = "proxy-secret"
        server.ALLOW_LOCAL_FILE_API = True
        server.FRONTEND_ORIGIN = ""
        server.BIND_HOST = "127.0.0.1"

        self.assertTrue(server._file_api_authorized({}, ("127.0.0.1", 51234)))
        self.assertTrue(server._file_api_authorized({}, ("::1", 51234)))
        self.assertFalse(server._file_api_authorized({}, ("203.0.113.10", 51234)))

    def test_cleanup_expired_outputs_deletes_only_old_files(self):
        old_file = Path(server.OUTPUT_DIR) / "old.docx"
        new_file = Path(server.OUTPUT_DIR) / "new.docx"
        old_file.write_text("old", encoding="utf-8")
        new_file.write_text("new", encoding="utf-8")
        old_mtime = time.time() - server.FILE_TTL - 60
        os.utime(old_file, (old_mtime, old_mtime))

        result = server._cleanup_expired_outputs()

        self.assertEqual(result["removed"], 1)
        self.assertFalse(old_file.exists())
        self.assertTrue(new_file.exists())

    def test_health_ready_version_payloads(self):
        self.assertEqual(server._health_payload()["status"], "ok")

        ready = server._ready_payload()
        self.assertTrue(ready["ok"], ready)
        self.assertTrue(ready["checks"]["database"])
        self.assertTrue(ready["checks"]["output_dir"])
        self.assertTrue(ready["checks"]["log_dir"])

        version = server._version_payload()
        self.assertEqual(version["file_ttl_seconds"], 86400)
        self.assertEqual(version["max_upload_mb"], 10)
        self.assertIn("started_at", version)

        server.PROXY_SECRET = "proxy-secret"
        self.assertTrue(server._version_payload()["proxy_secret_required"])
        server.PROXY_SECRET = ""
        self.assertTrue(server._version_payload()["proxy_secret_required"])
        self.assertFalse(server._version_payload()["proxy_secret_configured"])

    def test_error_payload_has_stable_code(self):
        self.assertEqual(
            server._error_payload("FILE_TOO_LARGE", "文件过大")["code"],
            "FILE_TOO_LARGE",
        )

    def test_monitor_links_use_session_forms_not_query_tokens(self):
        server.log_sql(
            "aaaaaaaa-aaaa-aaaa-aaaa-aaaaaaaaaaaa",
            "203.0.113.8", "ua", "a.docx", 100, "NORMAL",
            3, 1, 2, 1200, "done",
        )
        session = server._create_admin_session("ua", "203.0.113.8")

        html = server._monitor_html(server.get_sql_stats(), session["csrf_token"])

        self.assertIn('/stats"', html)
        self.assertIn('method="post" action="/cleanup"', html.replace("\n", " "))
        self.assertIn('name="csrf_token" value="' + session["csrf_token"], html)
        self.assertNotIn('token=', html)

    def test_monitor_query_keeps_pagination_and_ignores_time_filters(self):
        query = server._monitor_query_from(
            urlparse("/monitor?start=2026-06-01T09:30&end=2026-06-02&recent_page=2&ip_page=3&recent_size=50&ip_size=999")
        )

        self.assertNotIn("start", query)
        self.assertNotIn("end", query)
        self.assertEqual(query["recent_page"], 2)
        self.assertEqual(query["ip_page"], 3)
        self.assertEqual(query["recent_size"], 50)
        self.assertEqual(query["ip_size"], 100)

    def test_monitor_query_defaults_to_fifty_rows_per_page(self):
        query = server._monitor_query_from(urlparse("/monitor"))

        self.assertEqual(query["recent_size"], 50)
        self.assertEqual(query["ip_size"], 50)

    def test_sql_stats_support_pagination_without_time_filtering(self):
        for i in range(25):
            server.log_sql(
                f"task-{i}", f"203.0.113.{i % 3}", "ua", f"file-{i}.docx",
                100 + i, "NORMAL", 3, 1, 2, 1000, "done",
            )
            created = f"2026-06-01 10:{i:02d}:00" if i < 22 else "2026-05-01 10:00:00"
            with server._SQL_LOCK:
                conn = server._sql()
                conn.execute("UPDATE tasks SET created_at=? WHERE id=?", (created, f"task-{i}"))
                conn.commit()
                conn.close()

        stats = server.get_sql_stats({
            "start": "2026-06-01 00:00:00",
            "end": "2026-06-01 23:59:59",
            "recent_page": 2,
            "recent_size": 10,
            "ip_page": 1,
            "ip_size": 2,
        })

        self.assertEqual(stats["total"], 25)
        self.assertEqual(stats["recent_total"], 25)
        self.assertEqual(stats["recent_page"], 2)
        self.assertEqual(stats["recent_pages"], 3)
        self.assertEqual(len(stats["recent"]), 10)
        self.assertEqual(stats["recent"][0]["filename"], "file-14.docx")
        self.assertEqual(stats["ip_total"], 3)
        self.assertEqual(stats["ip_pages"], 2)
        self.assertEqual(len(stats["top_ips"]), 2)

    def test_sql_stats_clamps_out_of_range_pages_to_visible_data(self):
        for i in range(3):
            server.log_sql(
                f"task-{i}", f"203.0.113.{i}", "ua", f"file-{i}.docx",
                100, "NORMAL", 3, 1, 2, 1000, "done",
            )

        stats = server.get_sql_stats({
            "recent_page": 99,
            "recent_size": 20,
            "ip_page": 99,
            "ip_size": 20,
        })

        self.assertEqual(stats["total"], 3)
        self.assertEqual(stats["recent_page"], 1)
        self.assertEqual(stats["recent_pages"], 1)
        self.assertEqual(len(stats["recent"]), 3)
        self.assertEqual(stats["ip_page"], 1)
        self.assertEqual(stats["ip_pages"], 1)
        self.assertEqual(len(stats["top_ips"]), 3)

    def test_recent_tasks_use_insert_order_when_created_at_is_stale(self):
        server.log_sql(
            "old-task", "203.0.113.1", "ua", "old.docx",
            100, "NORMAL", 3, 1, 2, 1000, "done",
        )
        server.log_sql(
            "new-task", "203.0.113.2", "ua", "new.docx",
            100, "NORMAL", 3, 1, 2, 1000, "done",
        )
        with server._SQL_LOCK:
            conn = server._sql()
            conn.execute("UPDATE tasks SET created_at=? WHERE id=?", ("2026-01-01 00:00:00", "new-task"))
            conn.commit()
            conn.close()

        stats = server.get_sql_stats()

        self.assertEqual(stats["recent"][0]["filename"], "new.docx")

    def test_active_ips_use_latest_insert_order_before_upload_count(self):
        server.log_sql(
            "old-task-1", "203.0.113.10", "ua", "old-1.docx",
            100, "NORMAL", 3, 1, 2, 1000, "done",
        )
        server.log_sql(
            "old-task-2", "203.0.113.10", "ua", "old-2.docx",
            100, "NORMAL", 3, 1, 2, 1000, "done",
        )
        server.log_sql(
            "new-task", "203.0.113.99", "ua", "new.docx",
            100, "NORMAL", 3, 1, 2, 1000, "done",
        )

        stats = server.get_sql_stats()

        self.assertEqual(stats["top_ips"][0]["ip"], "203.0.113.99")
        self.assertEqual(stats["top_ips"][0]["last_filename"], "new.docx")

    def test_monitor_html_excludes_time_filter_and_keeps_pagination_links(self):
        stats = server.get_sql_stats({
            "recent_page": 2,
            "recent_size": 20,
            "ip_page": 1,
            "ip_size": 20,
        })

        html = server._monitor_html(stats, "secret")

        self.assertNotIn('name="start"', html)
        self.assertNotIn('name="end"', html)
        self.assertNotIn('type="date"', html)
        self.assertNotIn('时间查询', html)
        self.assertNotIn('id="filterForm"', html)
        self.assertIn('显示设置', html)
        self.assertIn('name="recent_size"', html)
        self.assertIn('name="ip_size"', html)
        self.assertIn('Noto Sans CJK SC', html)
        self.assertIn('WenQuanYi Micro Hei', html)
        self.assertIn('recent_page=1', html)
        self.assertNotIn('start=', html)

    def test_monitor_auto_refresh_is_kept_without_filter_form(self):
        html = server._monitor_html(server.get_sql_stats(), "secret")

        self.assertNotIn('http-equiv="refresh"', html)
        self.assertNotIn('id="filterForm"', html)
        self.assertNotIn('monitorAutoRefreshPaused', html)
        self.assertIn('setInterval', html)


if __name__ == "__main__":
    unittest.main()
