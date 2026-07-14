import hashlib
import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION

from docxtool.security.docx_integrity import DocxIntegrityError, validate_docx_integrity


REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\xf8\x0f"
    b"\x00\x01\x01\x01\x00\x1b\xb6\xeeV\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _content_types(extra: str = "") -> str:
    return (
        "<?xml version='1.0' encoding='UTF-8'?>"
        "<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'>"
        "<Default Extension='rels' ContentType='application/vnd.openxmlformats-package.relationships+xml'/>"
        "<Default Extension='xml' ContentType='application/xml'/>"
        "<Default Extension='png' ContentType='image/png'/>"
        "<Override PartName='/word/document.xml' "
        "ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'/>"
        f"{extra}</Types>"
    )


def _root_rels() -> str:
    return (
        "<?xml version='1.0' encoding='UTF-8'?>"
        f"<Relationships xmlns='{REL_NS}'>"
        "<Relationship Id='rId1' "
        "Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument' "
        "Target='word/document.xml'/>"
        "</Relationships>"
    )


def _document_xml(body: str = "") -> str:
    return (
        "<?xml version='1.0' encoding='UTF-8'?>"
        f"<w:document xmlns:w='{W_NS}' xmlns:r='{R_NS}'>"
        f"<w:body>{body}<w:sectPr/></w:body></w:document>"
    )


def _rels_xml(*relationships: str) -> str:
    return (
        "<?xml version='1.0' encoding='UTF-8'?>"
        f"<Relationships xmlns='{REL_NS}'>"
        f"{''.join(relationships)}"
        "</Relationships>"
    )


def _relationship(rel_id: str, target: str, *, target_mode: str | None = None) -> str:
    mode = f" TargetMode='{target_mode}'" if target_mode else ""
    return (
        f"<Relationship Id='{rel_id}' "
        "Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink' "
        f"Target='{target}'{mode}/>"
    )


def _minimal_docx_bytes(extra_members: dict[str, bytes | str] | None = None, document_xml: str | None = None) -> bytes:
    archive = io.BytesIO()
    with zipfile.ZipFile(archive, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _content_types())
        zf.writestr("_rels/.rels", _root_rels())
        zf.writestr("word/document.xml", document_xml or _document_xml())
        for name, content in (extra_members or {}).items():
            zf.writestr(name, content)
    return archive.getvalue()


def _replace_zip_member(source: bytes, replacements: dict[str, bytes | str]) -> bytes:
    source_archive = io.BytesIO(source)
    output = io.BytesIO()
    with zipfile.ZipFile(source_archive, "r") as src, zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename in replacements:
                continue
            dst.writestr(item, src.read(item.filename))
        for name, content in replacements.items():
            dst.writestr(name, content)
    return output.getvalue()


def test_minimal_docx_passes() -> None:
    report = validate_docx_integrity(_minimal_docx_bytes())

    assert report.ok is True
    assert report.part_count == 3
    assert report.relationship_count == 1


def test_multi_section_docx_passes(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("first section")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("second section")
    path = tmp_path / "multi-section.docx"
    document.save(path)

    report = validate_docx_integrity(path)

    assert report.ok is True
    assert report.part_count > 3


def test_missing_relationship_target_rejected() -> None:
    archive = _minimal_docx_bytes(
        {
            "word/_rels/document.xml.rels": _rels_xml(_relationship("rId7", "missing.xml")),
        },
        _document_xml("<w:p><w:hyperlink r:id='rId7'/></w:p>"),
    )

    with pytest.raises(DocxIntegrityError) as error:
        validate_docx_integrity(archive)

    assert error.value.code == "MISSING_REL_TARGET"


def test_duplicate_relationship_id_rejected() -> None:
    archive = _minimal_docx_bytes(
        {
            "word/_rels/document.xml.rels": _rels_xml(
                _relationship("rId1", "target-a.xml"),
                _relationship("rId1", "target-b.xml"),
            ),
            "word/target-a.xml": f"<w:root xmlns:w='{W_NS}'/>",
            "word/target-b.xml": f"<w:root xmlns:w='{W_NS}'/>",
        }
    )

    with pytest.raises(DocxIntegrityError) as error:
        validate_docx_integrity(archive)

    assert error.value.code == "DUPLICATE_REL_ID"


def test_external_relationship_not_reported_as_missing_package_part() -> None:
    archive = _minimal_docx_bytes(
        {
            "word/_rels/document.xml.rels": _rels_xml(
                _relationship("rId9", "https://example.test/doc", target_mode="External")
            ),
        },
        _document_xml("<w:p><w:hyperlink r:id='rId9'><w:r><w:t>link</w:t></w:r></w:hyperlink></w:p>"),
    )

    report = validate_docx_integrity(archive)

    assert report.ok is True
    assert report.relationship_count == 2


def test_header_footer_image_relationships_validate_recursively() -> None:
    document = _document_xml(
        "<w:sectPr><w:headerReference w:type='default' r:id='rIdHeader'/>"
        "<w:footerReference w:type='default' r:id='rIdFooter'/></w:sectPr>"
    )
    header = (
        f"<w:hdr xmlns:w='{W_NS}' xmlns:r='{R_NS}'>"
        "<w:p><w:r><w:drawing><a:blip xmlns:a='http://schemas.openxmlformats.org/drawingml/2006/main' "
        "r:embed='rIdImage'/></w:drawing></w:r></w:p></w:hdr>"
    )
    footer = f"<w:ftr xmlns:w='{W_NS}'><w:p><w:r><w:t>footer</w:t></w:r></w:p></w:ftr>"
    archive = _minimal_docx_bytes(
        {
            "word/_rels/document.xml.rels": _rels_xml(
                _relationship("rIdHeader", "header1.xml"),
                _relationship("rIdFooter", "footer1.xml"),
            ),
            "word/header1.xml": header,
            "word/footer1.xml": footer,
            "word/_rels/header1.xml.rels": _rels_xml(_relationship("rIdImage", "media/image1.png")),
            "word/media/image1.png": PNG_1X1,
        },
        document,
    )

    report = validate_docx_integrity(archive)

    assert report.ok is True
    assert report.checked_part_count >= 4


def test_input_sha256_unchanged(tmp_path: Path) -> None:
    path = tmp_path / "readonly.docx"
    path.write_bytes(_minimal_docx_bytes())
    before = hashlib.sha256(path.read_bytes()).hexdigest()

    validate_docx_integrity(path)

    after = hashlib.sha256(path.read_bytes()).hexdigest()
    assert after == before


def test_error_messages_do_not_include_local_absolute_paths(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    archive = _replace_zip_member(_minimal_docx_bytes(), {"word/document.xml": "<w:document>"})
    path.write_bytes(archive)

    with pytest.raises(DocxIntegrityError) as error:
        validate_docx_integrity(path)

    message = str(error.value)
    assert str(path) not in message
    assert str(tmp_path) not in message
    assert "word/document.xml" in message


def test_relationship_target_cannot_escape_package_boundary() -> None:
    archive = _minimal_docx_bytes(
        {
            "word/_rels/document.xml.rels": _rels_xml(_relationship("rId7", "../../evil.xml")),
        },
        _document_xml("<w:p><w:hyperlink r:id='rId7'/></w:p>"),
    )

    with pytest.raises(DocxIntegrityError) as error:
        validate_docx_integrity(archive)

    assert error.value.code == "REL_TARGET_ESCAPE"


def test_missing_content_type_coverage_rejected() -> None:
    archive = _minimal_docx_bytes({"word/media/image1.gif": b"GIF89a"})

    with pytest.raises(DocxIntegrityError) as error:
        validate_docx_integrity(archive)

    assert error.value.code == "MISSING_CONTENT_TYPE"
