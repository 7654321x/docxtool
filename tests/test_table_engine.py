import base64
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from docxtool.document.engine.table import format_tables
from docxtool.security.docx_integrity import validate_docx_integrity


def _tiny_png(path: Path) -> None:
    path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQ"
            "VR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )


def _document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def test_safe_table_format_preserves_merged_cells_nested_tables_and_images(tmp_path: Path) -> None:
    image = tmp_path / "tiny.png"
    _tiny_png(image)
    document = Document()
    table = document.add_table(rows=4, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Amount"
    table.cell(0, 2).text = "Description"
    table.cell(1, 0).merge(table.cell(1, 1)).text = "Merged"
    table.cell(1, 2).paragraphs[0].add_run().add_picture(str(image), width=Cm(0.2))
    table.cell(2, 0).text = "A"
    table.cell(2, 1).text = "123.45"
    table.cell(2, 2).text = "Long text that should remain left aligned"
    nested = table.cell(3, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested"
    before_grid_span_count = table._tbl.xml.count("gridSpan")
    before_drawing_count = table._tbl.xml.count("drawing")

    format_tables(
        document,
        {
            "width_cm": 16,
            "indent_cm": 0.5,
            "borders": {"val": "single", "size": 6, "color": "444444"},
            "cell_margin_cm": 0.12,
            "vertical_align": "center",
            "header": {"repeat": True, "bold": True, "shading": "D9EAF7", "alignment": "center"},
            "auto_align": True,
        },
    )
    output = tmp_path / "tables.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    xml = _document_xml(output)
    assert xml.count("<w:tblW") >= 2
    assert "<w:tblInd" in xml
    assert "<w:tblBorders>" in xml
    assert "<w:tblCellMar>" in xml
    assert "<w:tblHeader" in xml
    assert 'w:fill="D9EAF7"' in xml
    assert "gridSpan" in xml
    assert "drawing" in xml
    assert xml.count("gridSpan") == before_grid_span_count
    assert xml.count("drawing") == before_drawing_count
    assert 'w:val="right"' in xml
    assert 'w:val="center"' in xml
    assert 'w:val="left"' in xml


def test_table_format_disabled_does_not_change_table_xml() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "unchanged"
    before = table._tbl.xml

    format_tables(document, {"enabled": False, "width_cm": 10})

    assert table._tbl.xml == before
    assert table._tbl.tblPr.find(qn("w:tblW")) is not None
