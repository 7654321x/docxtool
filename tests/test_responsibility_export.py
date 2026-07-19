import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.oxml.ns import qn

from docxtool.document.engine import export_doc
from docxtool.document.engine.core import _apply_responsibility_line
from docxtool.document.importer import DocxImporter, DocumentData, ParagraphData, ParagraphFeatures
from docxtool.document.style_config import PageSettings, StyleRule
from docxtool.security.docx_integrity import validate_docx_integrity


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _rules() -> list[StyleRule]:
    return [StyleRule.default_for_row(index) for index in range(24)]


def _document_xml(path: Path) -> ET.Element:
    with zipfile.ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/document.xml"))


def _paragraphs(root: ET.Element) -> list[ET.Element]:
    return root.findall(".//w:body/w:p", NS)


def _text(paragraph: ET.Element) -> str:
    return "".join(element.text or "" for element in paragraph.findall(".//w:t", NS))


def _lines(paragraph: ET.Element) -> list[str]:
    lines = [""]
    for element in paragraph.iter():
        if element.tag == qn("w:t"):
            lines[-1] += element.text or ""
        elif element.tag == qn("w:br"):
            lines.append("")
    return lines


def _pstyle(paragraph: ET.Element) -> str:
    p_style = paragraph.find("w:pPr/w:pStyle", NS)
    return p_style.get(qn("w:val")) if p_style is not None else ""


def _run_text_and_bold(paragraph: ET.Element) -> list[tuple[str, bool]]:
    result = []
    for run in paragraph.findall("w:r", NS):
        text = "".join(element.text or "" for element in run.findall("w:t", NS))
        if not text:
            continue
        bold = run.find("w:rPr/w:b", NS)
        bold_value = bold.get(qn("w:val")) if bold is not None else None
        result.append((text, bold is not None and bold_value not in {"0", "false", "False"}))
    return result


def _export_responsibility(tmp_path: Path, text: str) -> tuple[Path, dict]:
    output = tmp_path / "responsibility.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData("主标题", "title", "主标题", ParagraphFeatures()),
            ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
            ParagraphData(text, "responsibility_line", text, ParagraphFeatures(), meta={"colon_bold": True}),
            ParagraphData("附件 1", "attachment_page_mark", "附件 1", ParagraphFeatures()),
        ],
        filepath="generated.docx",
    )
    stats = export_doc(data, _rules(), PageSettings(), str(output))
    return output, stats


@pytest.mark.parametrize(
    ("source_text", "expected_lines"),
    [
        ("责任单位：区政府", ["责任单位：区政府"]),
        ("责任单位：区政府\n责任单位：商务局", ["责任单位：区政府", "责任单位：商务局"]),
        ("责任单位：区政府责任单位：商务局责任单位：区政府", ["责任单位：区政府", "责任单位：商务局", "责任单位：区政府"]),
        ("责 任 单 位 ：区政府", ["责任单位：区政府"]),
    ],
)
def test_responsibility_export_splits_and_normalizes_labels(
    tmp_path: Path,
    source_text: str,
    expected_lines: list[str],
) -> None:
    output, stats = _export_responsibility(tmp_path, source_text)

    assert stats["fallback_count"] == 0
    assert validate_docx_integrity(output).ok is True
    root = _document_xml(output)
    responsibility = next(paragraph for paragraph in _paragraphs(root) if "责任单位" in _text(paragraph))

    assert _pstyle(responsibility) == "DCT-Responsibility"
    assert _lines(responsibility) == expected_lines
    assert len(responsibility.findall(".//w:br", NS)) == len(expected_lines) - 1
    assert all(line.count("责任单位：") == 1 for line in expected_lines)
    assert root.findall(".//w:numPr", NS) == []


def test_responsibility_runs_format_label_only_and_direct_paragraph_format(tmp_path: Path) -> None:
    output, _stats = _export_responsibility(tmp_path, "责 任 单 位 ：区政府责任单位：商务局")
    root = _document_xml(output)
    responsibility = next(paragraph for paragraph in _paragraphs(root) if "责任单位" in _text(paragraph))

    assert _run_text_and_bold(responsibility) == [
        ("责任单位：", True),
        ("区政府", False),
        ("责任单位：", True),
        ("商务局", False),
    ]
    assert _lines(responsibility) == ["责任单位：区政府", "责任单位：商务局"]
    assert responsibility.find("w:pPr/w:jc", NS).get(qn("w:val")) == "left"
    indent = responsibility.find("w:pPr/w:ind", NS)
    assert indent is not None
    assert indent.get(qn("w:leftChars")) == "200"
    assert indent.get(qn("w:left")) == "640"
    assert indent.get(qn("w:firstLineChars")) == "0"
    assert indent.get(qn("w:firstLine")) == "0"


def test_generic_key_value_line_uses_fixed_layout_and_label_only_bold(tmp_path: Path) -> None:
    output = tmp_path / "key-value.docx"
    rules = _rules()
    rules[5].font_size_pt = 12
    rules[5].spacing_before = 2
    rules[5].spacing_after = 2
    data = DocumentData(
        paragraphs=[
            ParagraphData(
                "联系人：张三", "body", "联系人：张三", ParagraphFeatures(),
                meta={"colon_bold": True},
            ),
        ],
        filepath="generated.docx",
    )

    export_doc(data, rules, PageSettings(line_spacing_value=20), str(output))
    paragraph = _paragraphs(_document_xml(output))[0]

    assert _run_text_and_bold(paragraph) == [("联系人：", True), ("张三", False)]
    indent = paragraph.find("w:pPr/w:ind", NS)
    assert indent.get(qn("w:leftChars")) == "0"
    assert indent.get(qn("w:firstLineChars")) == "200"
    spacing = paragraph.find("w:pPr/w:spacing", NS)
    assert spacing.get(qn("w:beforeLines")) == "0"
    assert spacing.get(qn("w:afterLines")) == "0"
    assert spacing.get(qn("w:line")) == "560"
    assert spacing.get(qn("w:lineRule")) == "exact"
    for run in paragraph.findall("w:r", NS):
        if run.find("w:t", NS) is not None:
            assert run.find("w:rPr/w:sz", NS).get(qn("w:val")) == "32"


def test_generic_key_value_line_is_detected_from_source_document(tmp_path: Path) -> None:
    source = tmp_path / "key-value-source.docx"
    document = Document()
    document.add_paragraph("主标题")
    document.add_paragraph("一、工作情况")
    document.add_paragraph("联系人：张三")
    document.save(source)

    data = DocxImporter().load(
        str(source),
        _rules(),
        features={"punctuation": {"enabled": True, "mode": "safe"}},
    )
    key_value = next(paragraph for paragraph in data.paragraphs if paragraph.original_text == "联系人：张三")

    assert key_value.type_id == "body"
    assert key_value.meta["colon_bold"] is True


def test_soft_broken_key_value_lines_are_split_and_detected(tmp_path: Path) -> None:
    source = tmp_path / "soft-key-values.docx"
    document = Document()
    document.add_paragraph("主标题")
    document.add_paragraph("一、工作情况")
    paragraph = document.add_paragraph("联系人：张三")
    paragraph.add_run().add_break()
    paragraph.add_run("网址：https://example.test/a:b?q=1.2")
    paragraph.add_run().add_break()
    paragraph.add_run("空值：")
    paragraph.add_run().add_break()
    paragraph.add_run("：只有值没有键名")
    paragraph.add_run().add_break()
    paragraph.add_run("超长键名边界测试字段：这是一个会自动换行的长值，用于验证标签和值格式。")
    document.save(source)

    data = DocxImporter().load(str(source), _rules())
    values = [item for item in data.paragraphs if item.meta.get("colon_bold")]

    assert [item.text for item in values] == [
        "联系人：张三",
        "网址：https://example.test/a:b?q=1.2",
        "超长键名边界测试字段：这是一个会自动换行的长值，用于验证标签和值格式。",
    ]


def test_responsibility_renderer_is_idempotent() -> None:
    document = Document()
    paragraph = document.add_paragraph("责任单位：区政府责任单位：区政府")

    _apply_responsibility_line(paragraph, paragraph.text)
    first = paragraph.text
    _apply_responsibility_line(paragraph, paragraph.text)

    assert paragraph.text == first
    assert paragraph.text == "责任单位：区政府\n责任单位：区政府"


def test_heading1_report_split_body_has_dct_body_and_structural_invariants(tmp_path: Path) -> None:
    output = tmp_path / "heading1-report.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData(
                "政协报告标题。正文内容正文内容正文内容",
                "heading1_report",
                "政协报告标题。正文内容正文内容正文内容",
                ParagraphFeatures(),
                meta={"heading1_report_split": True},
            ),
            ParagraphData("责任单位：区政府责任单位：商务局", "responsibility_line", "责任单位：区政府责任单位：商务局", ParagraphFeatures()),
            ParagraphData("附件 1", "attachment_page_mark", "附件 1", ParagraphFeatures()),
        ],
        filepath="generated.docx",
    )

    stats = export_doc(data, _rules(), PageSettings(), str(output))

    assert stats["fallback_count"] == 0
    assert validate_docx_integrity(output).ok is True
    root = _document_xml(output)
    paragraphs = [paragraph for paragraph in _paragraphs(root) if _text(paragraph).strip()]
    style_by_text = {_text(paragraph): _pstyle(paragraph) for paragraph in paragraphs}

    assert style_by_text["政协报告标题。"] == "DCT-Heading1"
    assert style_by_text["正文内容正文内容正文内容"] == "DCT-Body"
    assert style_by_text["责任单位：区政府责任单位：商务局"] == "DCT-Responsibility"
    assert style_by_text["附件 1"] == "DCT-AttachmentMark"
    assert all(_pstyle(paragraph).startswith("DCT-") for paragraph in paragraphs)
    assert root.findall(".//w:numPr", NS) == []
