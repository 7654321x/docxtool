import logging
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn

from docxtool.document.engine import export_doc
from docxtool.document.engine import core as engine_core
from docxtool.document.importer import DocumentData, ParagraphData, ParagraphFeatures
from docxtool.document.style_config import ExportError, PageSettings, StyleRule, logger


def _rules():
    return [StyleRule.default_for_row(i) for i in range(10)]


def _body_font(run):
    rPr = run._element.rPr
    rFonts = rPr.rFonts if rPr is not None else None
    return rFonts.get(qn("w:eastAsia")) if rFonts is not None else None


def _font_size_half_points(run):
    rPr = run._element.rPr
    size = rPr.find(qn("w:sz")) if rPr is not None else None
    return size.get(qn("w:val")) if size is not None else None


def _has_bold(run):
    rPr = run._element.rPr
    return rPr.find(qn("w:b")) is not None if rPr is not None else False


def _spacing_after_lines(paragraph):
    spacing = paragraph._element.get_or_add_pPr().find(qn("w:spacing"))
    return spacing.get(qn("w:afterLines")) if spacing is not None else None


def _spacing_before_lines(paragraph):
    spacing = paragraph._element.get_or_add_pPr().find(qn("w:spacing"))
    return spacing.get(qn("w:beforeLines")) if spacing is not None else None


def _snap_to_grid(paragraph):
    snap = paragraph._element.get_or_add_pPr().find(qn("w:snapToGrid"))
    return snap.get(qn("w:val")) if snap is not None else None


class EngineHeadingSpacingTest(unittest.TestCase):
    def setUp(self):
        logger.setLevel(logging.ERROR)
        self.tmp = tempfile.TemporaryDirectory()
        self.out = str(Path(self.tmp.name) / "out.docx")

    def tearDown(self):
        self.tmp.cleanup()

    def _export(
        self,
        paragraphs,
        *,
        processing_strategy="normalize",
        strict_preservation=False,
    ):
        doc_data = DocumentData(
            paragraphs=paragraphs,
            filepath="input.docx",
            processing_strategy=processing_strategy,
            strict_preservation=strict_preservation,
        )
        export_doc(doc_data, _rules(), PageSettings(), self.out)
        return Document(self.out)

    def test_title2_has_explicit_zero_first_line_indent(self):
        paragraph = ParagraphData(
            text="今后五年工作建议",
            type_id="title2",
            original_text="今后五年工作建议",
            features=ParagraphFeatures(),
            meta={},
        )

        for strategy in ("structural", "normalize"):
            document = self._export([paragraph], processing_strategy=strategy)
            properties = document.paragraphs[0]._element.get_or_add_pPr()
            indent = properties.find(qn("w:ind"))
            self.assertIsNotNone(indent)
            self.assertEqual(indent.get(qn("w:firstLineChars")), "0")
            self.assertEqual(indent.get(qn("w:firstLine")), "0")

    def test_numbered_heading2_colon_inline_body_formats_editable_modes_only(self):
        def paragraph():
            return ParagraphData(
                text="工作安排： 正文内容",
                type_id="heading2",
                original_text="（一）工作安排： 正文内容",
                features=ParagraphFeatures(),
                meta={
                    "numbered_heading2_colon_inline_body": True,
                    "numbering": "（一）",
                },
            )

        for strategy in ("structural", "normalize"):
            doc = self._export([paragraph()], processing_strategy=strategy)
            output = doc.paragraphs[0]
            self.assertEqual(len(doc.paragraphs), 1)
            self.assertEqual(output.text, "（一）工作安排： 正文内容")
            colon_position = output.text.index("：")
            cursor = 0
            for run in output.runs:
                if not run.text:
                    continue
                cursor += len(run.text)
                expected_font = (
                    "楷体_GB2312"
                    if cursor <= colon_position + 1
                    else "仿宋_GB2312"
                )
                self.assertEqual(_body_font(run), expected_font)

        strict_paragraph = ParagraphData(
            text="（一）工作安排： 正文内容",
            type_id="heading2",
            original_text="（一）工作安排： 正文内容",
            features=ParagraphFeatures(),
            meta={"numbered_heading2_colon_inline_body": True},
        )
        strict = self._export(
            [strict_paragraph],
            processing_strategy="strict",
            strict_preservation=True,
        )
        self.assertEqual(strict.paragraphs[0].text, "（一）工作安排： 正文内容")
        self.assertEqual(
            [run.text for run in strict.paragraphs[0].runs if run.text],
            ["（一）工作安排： 正文内容"],
        )

    def test_heading2_inline_body_participates_in_document_grid(self):
        paragraphs = [
            ParagraphData(
                text=f"标题{i}。这是需要继续按正文网格排列的较长正文内容",
                type_id="heading2",
                original_text=f"（{label}）标题{i}。这是需要继续按正文网格排列的较长正文内容",
                features=ParagraphFeatures(),
                meta={
                    "numbered_heading2_period_inline_body": True,
                    "numbering": f"（{label}）",
                },
            )
            for i, label in enumerate(("一", "二", "三"), start=1)
        ]

        document = self._export(paragraphs, processing_strategy="structural")

        self.assertEqual(len(document.paragraphs), 3)
        self.assertEqual([_snap_to_grid(item) for item in document.paragraphs], ["1"] * 3)
        for paragraph in document.paragraphs:
            self.assertIn("。", paragraph.text)
            period = paragraph.text.index("。")
            cursor = 0
            for run in paragraph.runs:
                cursor += len(run.text)
                expected_font = "楷体_GB2312" if cursor <= period + 1 else "仿宋_GB2312"
                self.assertEqual(_body_font(run), expected_font)

    def test_standalone_heading_and_body_keep_existing_grid_behavior(self):
        document = self._export(
            [
                ParagraphData(
                    "独立标题",
                    "heading2",
                    "（一）独立标题",
                    ParagraphFeatures(),
                    meta={"numbering": "（一）"},
                ),
                ParagraphData(
                    "这是普通正文。",
                    "body",
                    "这是普通正文。",
                    ParagraphFeatures(),
                ),
            ],
            processing_strategy="structural",
        )

        self.assertEqual(_snap_to_grid(document.paragraphs[0]), "0")
        self.assertEqual(_snap_to_grid(document.paragraphs[1]), "1")

    def test_numbered_heading2_colon_skips_competing_inline_effects(self):
        paragraph = ParagraphData(
            text="工作安排： 正文内容",
            type_id="heading2",
            original_text="（一）工作安排： 正文内容",
            features=ParagraphFeatures(),
            meta={
                "numbered_heading2_colon_inline_body": True,
                "numbering": "（一）",
                "numbered_bold": True,
                "colon_bold": True,
                "inline_lead_bold": True,
            },
        )

        with patch.object(engine_core, "_apply_special_bold") as special, \
                patch.object(engine_core, "_apply_colon_bold") as colon, \
                patch.object(engine_core, "_apply_key_value_line_format") as key_value, \
                patch.object(engine_core, "_apply_inline_lead_bold") as lead:
            doc = self._export([paragraph], processing_strategy="structural")

        self.assertEqual(doc.paragraphs[0].text, "（一）工作安排： 正文内容")
        special.assert_not_called()
        colon.assert_not_called()
        key_value.assert_not_called()
        lead.assert_not_called()

    def test_heading1_period_splits_one_complete_body_paragraph(self):
        doc = self._export([
            ParagraphData(
                text="一级标题。这里是正文内容这里是正文内容",
                type_id="heading1",
                original_text="一、一级标题。这里是正文内容这里是正文内容",
                features=ParagraphFeatures(),
                meta={"numbering": "一、"},
            )
        ])

        heading = doc.paragraphs[0]
        body = doc.paragraphs[1]

        self.assertEqual(len(doc.paragraphs), 2)
        self.assertEqual(heading.text, "一、一级标题")
        self.assertNotIn("。", heading.text)
        self.assertEqual(body.text, "这里是正文内容这里是正文内容")
        self.assertFalse(body.runs[-1].bold)
        self.assertEqual(_body_font(body.runs[-1]), "仿宋_GB2312")

    def test_structural_heading1_period_splits_one_complete_body_paragraph(self):
        doc = self._export([
            ParagraphData(
                text="一级标题。这里是正文内容这里是正文内容",
                type_id="heading1",
                original_text="一、一级标题。这里是正文内容这里是正文内容",
                features=ParagraphFeatures(),
                meta={"numbering": "一、", "heading_inline_body": True},
            )
        ], processing_strategy="structural")

        self.assertEqual(len(doc.paragraphs), 2)
        self.assertEqual(doc.paragraphs[0].text, "一、一级标题")
        self.assertEqual(doc.paragraphs[1].text, "这里是正文内容这里是正文内容")
        self.assertEqual(doc.paragraphs[1].style.style_id, "DCT-Body")
        self.assertFalse(doc.paragraphs[1].runs[-1].bold)

    def test_inline_heading_body_verification_blocks_truncated_output(self):
        original = engine_core._apply_inline_heading_body_split

        def corrupt_body(*args, **kwargs):
            body = original(*args, **kwargs)
            if body is not None:
                body.add_run("错误附加内容")
            return body

        paragraphs = [
            ParagraphData(
                text="一级标题。这里是完整正文内容这里是完整正文内容",
                type_id="heading1",
                original_text="一、一级标题。这里是完整正文内容这里是完整正文内容",
                features=ParagraphFeatures(),
                meta={"numbering": "一、"},
            )
        ]
        with patch.object(engine_core, "_apply_inline_heading_body_split", corrupt_body):
            with self.assertRaisesRegex(ExportError, "正文未完整保留"):
                self._export(paragraphs)

    def test_structural_render_failure_is_not_downgraded_to_body(self):
        paragraph = ParagraphData(
            text="一级标题",
            type_id="heading1",
            original_text="一、一级标题",
            features=ParagraphFeatures(),
        )
        with patch.object(
            engine_core,
            "_resolve_rule",
            side_effect=RuntimeError("injected structural render failure"),
        ):
            with self.assertRaisesRegex(
                ExportError,
                "结构段落排版失败: type=heading1 index=0",
            ):
                self._export([paragraph])

    def test_terminal_body_uses_widow_control_without_changing_earlier_body(self):
        doc = self._export([
            ParagraphData("第一段正文内容。", "body", "第一段正文内容。", ParagraphFeatures()),
            ParagraphData("最后一段正文内容。", "body", "最后一段正文内容。", ParagraphFeatures()),
        ])

        first = doc.paragraphs[0]._p.get_or_add_pPr().find(qn("w:widowControl"))
        last = doc.paragraphs[1]._p.get_or_add_pPr().find(qn("w:widowControl"))
        self.assertEqual(first.get(qn("w:val")), "0")
        self.assertEqual(last.get(qn("w:val")), "1")

    def test_standalone_heading_terminal_period_is_removed_at_every_level(self):
        cases = (
            ("heading1", "一、一级标题。", "一、一级标题"),
            ("heading2", "（一）二级标题。", "（一）二级标题"),
            ("heading3", "1.三级标题。", "1.三级标题"),
            ("heading4", "（1）四级标题。", "（1）四级标题"),
        )
        for type_id, source, expected in cases:
            with self.subTest(type_id=type_id):
                doc = self._export([
                    ParagraphData(
                        text=source,
                        type_id=type_id,
                        original_text=source,
                        features=ParagraphFeatures(),
                    )
                ], processing_strategy="structural")

                self.assertEqual(doc.paragraphs[0].text, expected)

    def test_heading_period_before_body_is_not_removed_as_terminal_punctuation(self):
        source = "（一）二级标题。这里是标题后的正文内容。"
        doc = self._export([
            ParagraphData(
                text=source,
                type_id="heading2",
                original_text=source,
                features=ParagraphFeatures(),
            )
        ], processing_strategy="structural")

        self.assertEqual(doc.paragraphs[0].text, source)

    def test_heading2_period_body_keeps_one_paragraph_with_body_font(self):
        source = "（一）会议安排。 后续正文内容完整说明有关工作要求。"
        doc = self._export([
            ParagraphData(
                text=source,
                type_id="heading2",
                original_text=source,
                features=ParagraphFeatures(),
                meta={"numbered_heading2_period_inline_body": True},
            )
        ], processing_strategy="structural")

        self.assertEqual(len(doc.paragraphs), 1)
        paragraph = doc.paragraphs[0]
        self.assertEqual(paragraph.text, source)
        self.assertEqual(
            [run.text for run in paragraph.runs if run.text],
            ["（一）会议安排。", " 后续正文内容完整说明有关工作要求。"],
        )
        self.assertEqual(_body_font(paragraph.runs[0]), "楷体_GB2312")
        self.assertEqual(_body_font(paragraph.runs[1]), "仿宋_GB2312")

    def test_head_area_inserts_blank_line_before_body_or_heading1(self):
        cases = [
            ("title", "正文内容正文内容。"),
            ("title_cont", "一、一级标题"),
            ("author_line", "一、一级标题"),
            ("role_name", "正文内容正文内容。"),
        ]
        for head_type, next_text in cases:
            with self.subTest(head_type=head_type, next_text=next_text):
                next_type = "heading1" if next_text.startswith("一、") else "body"
                doc = self._export([
                    ParagraphData(
                        text="总题目" if head_type != "author_line" else "姓名",
                        type_id=head_type,
                        original_text="总题目",
                        features=ParagraphFeatures(),
                        meta={"is_title": head_type == "title"},
                    ),
                    ParagraphData(
                        text=next_text.replace("一、", ""),
                        type_id=next_type,
                        original_text=next_text,
                        features=ParagraphFeatures(),
                        meta={"numbering": "一、"} if next_type == "heading1" else {},
                    ),
                ])

                if head_type == "role_name":
                    self.assertEqual(doc.paragraphs[1].text, next_text)
                    self.assertEqual(_spacing_after_lines(doc.paragraphs[0]), "100")
                    self.assertIn(_spacing_before_lines(doc.paragraphs[1]), (None, "0"))
                else:
                    self.assertEqual(doc.paragraphs[1].text, "")

    def test_date_line_uses_spacing_after_without_blank_paragraph(self):
        doc = self._export([
            ParagraphData(
                text="（2026年7月  日）",
                type_id="date_line",
                original_text="（2026年7月  日）",
                features=ParagraphFeatures(),
                meta={},
            ),
            ParagraphData(
                text="正文内容正文内容。",
                type_id="body",
                original_text="正文内容正文内容。",
                features=ParagraphFeatures(),
                meta={},
            ),
        ])

        self.assertEqual([p.text for p in doc.paragraphs[:2]], ["（2026年7月  日）", "正文内容正文内容。"])
        self.assertEqual(_spacing_after_lines(doc.paragraphs[0]), "100")

    def test_opening_addressing_keeps_configured_one_line_before_spacing(self):
        doc = self._export([
            ParagraphData("讲话标题", "title", "讲话标题", ParagraphFeatures(), meta={"is_title": True}),
            ParagraphData("（2026年8月27日）", "date_line", "（2026年8月27日）", ParagraphFeatures()),
            ParagraphData("各位委员、同志们：", "addressing", "各位委员、同志们：", ParagraphFeatures()),
            ParagraphData("正文内容完整保留。", "body", "正文内容完整保留。", ParagraphFeatures()),
        ], processing_strategy="structural")

        addressing = next(p for p in doc.paragraphs if p.text == "各位委员、同志们：")
        self.assertEqual(_spacing_before_lines(addressing), "100")
        self.assertIn(_spacing_after_lines(addressing), (None, "0"))

    def test_addressing_after_body_has_no_before_spacing(self):
        doc = self._export([
            ParagraphData("讲话标题", "title", "讲话标题", ParagraphFeatures(), meta={"is_title": True}),
            ParagraphData("各位委员、同志们：", "addressing", "各位委员、同志们：", ParagraphFeatures()),
            ParagraphData("正文内容完整保留。", "body", "正文内容完整保留。", ParagraphFeatures()),
            ParagraphData("各位委员、同志们！", "addressing", "各位委员、同志们！", ParagraphFeatures()),
            ParagraphData("结尾正文内容完整保留。", "body", "结尾正文内容完整保留。", ParagraphFeatures()),
        ], processing_strategy="structural")

        opening = next(p for p in doc.paragraphs if p.text == "各位委员、同志们：")
        closing = next(p for p in doc.paragraphs if p.text == "各位委员、同志们！")
        self.assertEqual(_spacing_before_lines(opening), "100")
        self.assertEqual(_spacing_before_lines(closing), "0")
        self.assertEqual(_spacing_after_lines(closing), "0")

    def test_role_name_and_date_line_are_adjacent(self):
        doc = self._export([
            ParagraphData(
                text="区政协副主席   杨明远",
                type_id="role_name",
                original_text="区政协副主席   杨明远",
                features=ParagraphFeatures(),
                meta={},
            ),
            ParagraphData(
                text="（2026年7月  日）",
                type_id="date_line",
                original_text="（2026年7月  日）",
                features=ParagraphFeatures(),
                meta={},
            ),
            ParagraphData(
                text="正文内容正文内容。",
                type_id="body",
                original_text="正文内容正文内容。",
                features=ParagraphFeatures(),
                meta={},
            ),
        ])

        self.assertEqual(
            [p.text for p in doc.paragraphs[:3]],
            ["区政协副主席   杨明远", "（2026年7月  日）", "正文内容正文内容。"],
        )
        self.assertIn(_spacing_after_lines(doc.paragraphs[0]), (None, "0"))
        self.assertEqual(_spacing_after_lines(doc.paragraphs[1]), "100")

    def test_role_name_after_multiline_title_has_one_line_before_spacing(self):
        doc = self._export([
            ParagraphData(
                text="中共内江市东兴区政协党组班子",
                type_id="title",
                original_text="中共内江市东兴区政协党组班子",
                features=ParagraphFeatures(),
                meta={"is_title": True},
            ),
            ParagraphData(
                text="2025年度民主生活会对照检查材料",
                type_id="title_cont",
                original_text="2025年度民主生活会对照检查材料",
                features=ParagraphFeatures(),
                meta={},
            ),
            ParagraphData(
                text="区政协办公室主任  李弟弟",
                type_id="role_name",
                original_text="区政协办公室主任  李弟弟",
                features=ParagraphFeatures(),
                meta={},
            ),
            ParagraphData(
                text="一级标题",
                type_id="heading1",
                original_text="一、一级标题",
                features=ParagraphFeatures(),
                meta={"numbering": "一、"},
            ),
        ])

        role = next(p for p in doc.paragraphs if p.style.style_id == "DCT-RoleName")
        self.assertEqual(_spacing_before_lines(role), "100")
        self.assertEqual(_spacing_after_lines(role), "100")
        self.assertEqual(doc.paragraphs[1].text, "2025年度民主生活会对照检查材料")
        self.assertEqual(doc.paragraphs[2].text, "区政协办公室主任  李弟弟")
        self.assertEqual(doc.paragraphs[3].text, "一、一级标题")
        self.assertIn(_spacing_before_lines(doc.paragraphs[3]), (None, "0"))

    def test_attachment_note_has_one_line_gap_after_body(self):
        doc = self._export([
            ParagraphData("正文内容。", "body", "正文内容。", ParagraphFeatures()),
            ParagraphData("附件：1. 基本情况", "attachment_note", "附件：1. 基本情况", ParagraphFeatures()),
            ParagraphData("2. 具体情况", "attachment_note_item", "2. 具体情况", ParagraphFeatures()),
            ParagraphData("区政协办", "sign_org", "区政协办", ParagraphFeatures()),
            ParagraphData("2025年10月15日", "sign_date", "2025年10月15日", ParagraphFeatures()),
        ])

        self.assertEqual(_spacing_before_lines(doc.paragraphs[1]), "100")
        self.assertIn(_spacing_before_lines(doc.paragraphs[2]), (None, "0"))
        self.assertEqual(doc.paragraphs[1].style.style_id, "DCT-AttachmentNote")
        self.assertEqual(doc.paragraphs[2].style.style_id, "DCT-AttachmentNoteItem")
        self.assertEqual(_spacing_before_lines(doc.paragraphs[3]), "300")
        self.assertIn(_spacing_before_lines(doc.paragraphs[4]), (None, "0"))
        note_indent = doc.paragraphs[1]._p.get_or_add_pPr().find(qn("w:ind"))
        item_indent = doc.paragraphs[2]._p.get_or_add_pPr().find(qn("w:ind"))
        signature_indent = doc.paragraphs[3]._p.get_or_add_pPr().find(qn("w:ind"))
        date_indent = doc.paragraphs[4]._p.get_or_add_pPr().find(qn("w:ind"))
        self.assertEqual(note_indent.get(qn("w:leftChars")), "200")
        self.assertEqual(item_indent.get(qn("w:leftChars")), "500")
        self.assertEqual(signature_indent.get(qn("w:rightChars")), "200")
        self.assertEqual(date_indent.get(qn("w:rightChars")), "400")

    def test_export_normalizes_attachment_note_before_signature_block(self):
        doc = self._export([
            ParagraphData("正文内容。", "body", "正文内容。", ParagraphFeatures()),
            ParagraphData("区政协办", "sign_org", "区政协办", ParagraphFeatures()),
            ParagraphData("附件：1. 基本情况", "attachment_note", "附件：1. 基本情况", ParagraphFeatures()),
            ParagraphData("2025年10月15日", "sign_date", "2025年10月15日", ParagraphFeatures()),
            ParagraphData("2. 具体情况", "attachment_note_item", "2. 具体情况", ParagraphFeatures()),
            ParagraphData("3. 超级情况", "attachment_note_item", "3. 超级情况", ParagraphFeatures()),
        ])

        self.assertEqual(
            [paragraph.text for paragraph in doc.paragraphs],
            [
                "正文内容。",
                "附件：1. 基本情况",
                "2. 具体情况",
                "3. 超级情况",
                "区政协办",
                "2025年10月15日",
            ],
        )

    def test_author_role_and_meeting_title_metadata_use_kaiti_gb2312_16pt_bold(self):
        doc = self._export([
            ParagraphData(
                text="张三",
                type_id="author_line",
                original_text="张三",
                features=ParagraphFeatures(),
                meta={},
            ),
            ParagraphData(
                text="区政协办公室主任  李弟弟",
                type_id="role_name",
                original_text="区政协办公室主任  李弟弟",
                features=ParagraphFeatures(),
                meta={},
            ),
            ParagraphData(
                text="在全市重点工作会议上",
                type_id="meeting_title_meta",
                original_text="在全市重点工作会议上",
                features=ParagraphFeatures(),
                meta={},
            ),
        ])

        for paragraph in doc.paragraphs:
            with self.subTest(text=paragraph.text):
                run = paragraph.runs[0]
                self.assertEqual(_body_font(run), "楷体_GB2312")
                self.assertEqual(_font_size_half_points(run), "32")
                self.assertTrue(_has_bold(run))

    def test_numbered_bold_does_not_duplicate_text(self):
        text = (
            "一是加强理论武装，把牢正确履职方向。"
            "坚持把学习贯彻习近平总书记关于树立和践行正确政绩观的重要论述作为重要政治任务。"
        )

        doc = self._export([
            ParagraphData(
                text=text,
                type_id="body",
                original_text=text,
                features=ParagraphFeatures(),
                meta={"numbered_bold": True},
            )
        ])

        self.assertEqual(doc.paragraphs[0].text, text)
        self.assertEqual(doc.paragraphs[0].text.count("一是加强理论武装"), 1)

    def test_inline_lead_bold_stays_in_one_body_paragraph(self):
        text = "推动重点工作走深走实。各单位结合实际持续抓好任务落实。"
        doc = self._export([
            ParagraphData(
                text=text,
                type_id="body",
                original_text=text,
                features=ParagraphFeatures(),
                meta={"inline_lead_bold": True},
            )
        ], processing_strategy="structural")

        self.assertEqual(len(doc.paragraphs), 1)
        self.assertEqual(doc.paragraphs[0].text, text)
        self.assertTrue(_has_bold(doc.paragraphs[0].runs[0]))
        self.assertFalse(doc.paragraphs[0].runs[-1].bold)

    def test_repeated_numbered_leads_only_bold_each_lead_sentence(self):
        text = (
            "一要坚持统筹推进。第一部分正文保持普通格式。"
            "二要强化协同配合。第二部分正文保持普通格式。"
            "三要完善长效机制。第三部分正文保持普通格式。"
        )
        doc = self._export([
            ParagraphData(
                text=text,
                type_id="body",
                original_text=text,
                features=ParagraphFeatures(),
                meta={"numbered_bold": True},
            )
        ], processing_strategy="structural")

        paragraph = doc.paragraphs[0]
        assert paragraph.text == text
        assert len(paragraph.runs) >= 6
        for index, run in enumerate(paragraph.runs[:6]):
            assert bool(run.bold) is (index % 2 == 0)


if __name__ == "__main__":
    unittest.main()
