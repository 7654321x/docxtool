from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from docxtool.document.engine.cleanup import cleanup_styles
from docxtool.security.docx_integrity import validate_docx_integrity


def _add_run_shading(run, fill: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    r_pr.append(shading)


def test_cleanup_default_off_does_not_change_document_xml() -> None:
    document = Document()
    run = document.add_paragraph().add_run("keep highlighted")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.color.rgb = RGBColor(255, 0, 0)
    before = document._element.xml

    cleanup_styles(document)

    assert document._element.xml == before


def test_safe_cleanup_removes_only_explicitly_configured_run_styles(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    abnormal = paragraph.add_run("configured anomaly")
    abnormal.font.highlight_color = WD_COLOR_INDEX.YELLOW
    abnormal.font.color.rgb = RGBColor(0, 255, 0)
    abnormal.font.underline = True
    abnormal.font.italic = True
    _add_run_shading(abnormal, "FFFF00")
    redhead = paragraph.add_run(" redhead")
    redhead.font.color.rgb = RGBColor(255, 0, 0)
    signature = document.add_paragraph("电子签名有效").add_run(" keep signature italic")
    signature.font.italic = True
    table = document.add_table(rows=1, cols=1)
    cell_shading = OxmlElement("w:shd")
    cell_shading.set(qn("w:fill"), "D9EAF7")
    table.cell(0, 0)._tc.get_or_add_tcPr().append(cell_shading)

    cleanup_styles(
        document,
        {
            "mode": "safe",
            "clear_highlight_colors": ["YELLOW"],
            "clear_shading_fills": ["FFFF00"],
            "clear_font_colors": ["00FF00"],
            "clear_underline_values": [True],
            "clear_italic_values": [True],
        },
    )
    output = tmp_path / "cleanup.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    xml = document._element.xml
    assert "highlight" not in abnormal._element.xml
    assert "FFFF00" not in abnormal._element.xml
    assert "00FF00" not in abnormal._element.xml
    assert "w:u" not in abnormal._element.xml
    assert "w:i" not in abnormal._element.xml
    assert "FF0000" in redhead._element.xml
    assert "w:i" in signature._element.xml
    assert "D9EAF7" in xml
