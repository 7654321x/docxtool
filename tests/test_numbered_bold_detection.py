import logging
import tempfile
import unittest
from pathlib import Path

from docx import Document

from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import StyleRule, logger


def _rules():
    return [StyleRule.default_for_row(i) for i in range(10)]


class NumberedBoldDetectionTest(unittest.TestCase):
    def setUp(self):
        logger.setLevel(logging.ERROR)

    def test_yishi_body_does_not_also_use_report_first_sentence_bold(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "source.docx"
            doc = Document()
            doc.add_paragraph("一、一级标题")
            doc.add_paragraph(
                "一是加强理论武装，把牢正确履职方向。"
                "坚持把学习贯彻习近平总书记关于树立和践行正确政绩观的重要论述作为重要政治任务。"
            )
            doc.save(src)

            data = DocxImporter().load(str(src), _rules())
            body = next(p for p in data.paragraphs if p.text.startswith("一是加强理论武装"))

            self.assertTrue(body.meta.get("numbered_bold"))
            self.assertNotIn("report_first_sentence_bold", body.meta)

    def test_date_after_role_name_is_detected_as_date_line(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "source.docx"
            doc = Document()
            doc.add_paragraph("在理论学习中心组学习会上的发言")
            doc.add_paragraph("杨明远")
            doc.add_paragraph("（2026年7月  日）")
            doc.add_paragraph("各位同志：")
            doc.save(src)

            data = DocxImporter().load(str(src), _rules())
            types_by_text = {p.text: p.type_id for p in data.paragraphs}

            self.assertEqual(types_by_text["杨明远"], "role_name")
            self.assertEqual(types_by_text["（2026年7月  日）"], "date_line")


if __name__ == "__main__":
    unittest.main()
