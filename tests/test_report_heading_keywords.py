from __future__ import annotations

from docx import Document
import pytest

from docxtool.document.importer import DocxImporter
from docxtool.document.recognition import RecognitionConfig, apply_recognition
from docxtool.document.style_config import StyleRule
from docxtool.sdk import recognize_docx
from tests.support.recognition_helpers import _document, _paragraph


def test_report_numbered_headings_use_the_ordinary_numbering_provider() -> None:
    first = _paragraph("一、过去五年工作回顾", "body", 1)
    second = _paragraph("二、今后五年工作安排", "body", 2)
    data = _document(
        _paragraph("政府工作报告", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        first,
        second,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert [first.type_id, second.type_id] == ["heading1", "heading1"]
    assert all(
        item.meta["recognition_provider"].startswith("numbering:")
        for item in (first, second)
    )


def test_report_salutations_use_the_ordinary_structural_provider() -> None:
    first = _paragraph("各位委员、同志们：", "body", 1)
    second = _paragraph("各位同志：", "body", 2)
    data = _document(
        _paragraph("政府工作报告", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        first,
        second,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert [first.type_id, second.type_id] == ["addressing", "addressing"]
    assert all(
        "standalone-addressing" in item.meta["recognition_evidence"]
        for item in (first, second)
    )


def test_annual_review_inline_heading_is_segmented_into_heading1_and_body(tmp_path) -> None:
    source = tmp_path / "report-inline-heading.docx"
    document = Document()
    title = document.add_paragraph()
    title.alignment = 1
    title.add_run("政府工作报告").bold = True
    document.add_paragraph("五年来。我们持续推进重点工作并取得新的明显成效。")
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative", include_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 1]

    assert [block.type_id for block in blocks] == ["heading1", "body"]
    assert [block.recognized_text for block in blocks] == [
        "五年来。", "我们持续推进重点工作并取得新的明显成效。",
    ]


def test_annual_review_is_segmented_from_local_evidence_in_any_mode(tmp_path) -> None:
    report_source = tmp_path / "report-one-year.docx"
    report = Document()
    title = report.add_paragraph()
    title.alignment = 1
    title.add_run("政府工作报告").bold = True
    report.add_paragraph("一年来。我们持续推进重点工作并取得新的明显成效。")
    report.save(report_source)

    report_plan = recognize_docx(report_source, recognition_mode="authoritative", include_text=True)
    report_blocks = [block for block in report_plan.blocks if block.physical_paragraph_index == 1]
    assert [block.type_id for block in report_blocks] == ["heading1", "body"]

    normal_source = tmp_path / "normal-one-year.docx"
    normal = Document()
    normal.add_paragraph("普通材料")
    normal.add_paragraph("一年来。我们持续推进重点工作并取得新的明显成效。")
    normal.add_paragraph("五年来。我们持续推进重点工作并取得新的明显成效。")
    normal.save(normal_source)

    normal_plan = recognize_docx(normal_source, recognition_mode="authoritative", include_text=True)
    for paragraph_index in (1, 2):
        normal_blocks = [
            block for block in normal_plan.blocks
            if block.physical_paragraph_index == paragraph_index
        ]
        assert [block.type_id for block in normal_blocks] == ["heading1", "body"]


def test_report_front_short_titles_restore_title2_and_glossary_title() -> None:
    title2 = _paragraph("工作安排", "body", 1)
    glossary_title = _paragraph("名词解释", "body", 2)
    data = _document(
        _paragraph("政府工作报告", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        title2,
        glossary_title,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert [title2.type_id, glossary_title.type_id] == ["title2", "glossary_title"]
    title2_trace = data.recognition_diagnostics["candidate_trace"][1]
    title2_candidate = next(
        item for item in title2_trace["candidates"] if item["type"] == "title2"
    )
    assert "report-mode-prior" in title2_candidate["evidence"]


def test_normal_front_short_title_gets_a_soft_title2_candidate() -> None:
    title2 = _paragraph("工作安排", "body", 1)
    data = _document(
        _paragraph("普通材料", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        title2,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    trace = data.recognition_diagnostics["candidate_trace"][1]
    assert any(
        item["type"] == "title2" and item["source"] == "body-title"
        for item in trace["candidates"]
    )


@pytest.mark.parametrize("glossary_text", ["名词解释", "名词解释与说明", "名词解释及说明"])
def test_normal_front_glossary_title_uses_fuzzy_marker(glossary_text: str) -> None:
    glossary_title = _paragraph(glossary_text, "body", 1)
    data = _document(
        _paragraph("普通材料", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        glossary_title,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert glossary_title.type_id == "glossary_title"


def test_body_title2_between_body_paragraphs_uses_centered_subtitle_semantics() -> None:
    title2 = _paragraph("今后五年工作建议", "body", 1, bold_char_ratio=1.0)
    data = _document(
        _paragraph("前一段正文已经完整说明工作情况和阶段性成效。", "body", 0),
        title2,
        _paragraph("后一段正文继续说明下一阶段的重点任务和工作安排。", "body", 2),
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert title2.type_id == "title2"
    assert title2.meta["recognition_provider"].startswith("body-title:")


def test_report_short_section_titles_do_not_require_source_bold() -> None:
    first_title = _paragraph("过去五年主要工作", "body", 2)
    second_title = _paragraph("今后五年工作建议", "body", 4)
    data = _document(
        _paragraph("人民检察院工作报告", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        _paragraph("正文开头已经完整说明报告背景和总体工作情况。", "body", 1),
        first_title,
        _paragraph("过去五年工作正文已经完整说明主要成绩和具体做法。", "body", 3),
        second_title,
        _paragraph("今后五年工作正文已经完整说明总体思路和重点任务。", "body", 5),
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert [first_title.type_id, second_title.type_id] == ["title2", "title2"]
    assert all(
        "report-middle-body-position" in item.meta["recognition_evidence"]
        for item in (first_title, second_title)
    )


def test_production_recognition_keeps_report_short_title_over_core_body_candidate(tmp_path) -> None:
    source = tmp_path / "report-production-short-title.docx"
    document = Document()
    title = document.add_paragraph("人民检察院工作报告")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("前一段正文已经完整说明工作情况和阶段性成效。")
    document.add_paragraph("今后五年工作建议")
    document.add_paragraph("后一段正文继续说明下一阶段的重点任务和工作安排，内容足够长以形成正文上下文。")
    document.save(source)

    data = DocxImporter().load(
        str(source),
        [StyleRule.default_for_row(index) for index in range(10)],
        strict_preservation=False,
    )

    short_title = data.paragraphs[2]
    assert short_title.type_id == "title2"
    assert short_title.meta["recognition_provider"].startswith("body-title:")


@pytest.mark.parametrize("title_text", ["名词解释", "名词解释与说明", "注释及说明"])
def test_production_recognition_keeps_fuzzy_glossary_title_over_core_body_candidate(
    tmp_path, title_text: str,
) -> None:
    source = tmp_path / "report-production-glossary-title.docx"
    document = Document()
    title = document.add_paragraph("人民检察院工作报告")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("正文已经开始并完整说明工作背景和基本情况。")
    glossary_title = document.add_paragraph(title_text)
    glossary_title.paragraph_format.page_break_before = True
    document.add_paragraph("1.术语：这是名词解释条目的正文说明。")
    document.save(source)

    data = DocxImporter().load(
        str(source),
        [StyleRule.default_for_row(index) for index in range(10)],
        strict_preservation=False,
    )

    assert data.paragraphs[2].type_id == "glossary_title"
    assert data.paragraphs[3].type_id == "glossary_item"


@pytest.mark.parametrize("numbered_text, expected_type", [
    ("一、总体情况", "heading1"),
    ("（一）总体情况", "heading2"),
])
def test_glossary_does_not_claim_non_level3_numbered_headings(
    numbered_text: str,
    expected_type: str,
) -> None:
    glossary_item = _paragraph(numbered_text, "body", 2, numbering_prefix=numbered_text.split("总体", 1)[0])
    data = _document(
        _paragraph("普通材料", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        _paragraph("正文已经开始并完整说明工作背景和基本情况。", "body", 1),
        _paragraph("名词解释", "body", 2),
        glossary_item,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert glossary_item.type_id == expected_type
    assert glossary_item.type_id != "glossary_item"


@pytest.mark.parametrize(
    "text",
    [
        "五年来，我们持续推进重点工作并取得新的明显成效。",
        "过去五年来。我们持续推进重点工作并取得新的明显成效。",
        "五年来。",
    ],
)
def test_annual_review_requires_exact_lead_period_and_body(text, tmp_path) -> None:
    source = tmp_path / "normal-annual-review-negative.docx"
    document = Document()
    document.add_paragraph("普通材料")
    document.add_paragraph(text)
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative", include_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 1]

    assert len(blocks) == 1


def test_numbered_glossary_item_without_colon_preserves_missing_colon_metadata() -> None:
    glossary_title = _paragraph("名词解释", "body", 0)
    glossary_item = _paragraph("1.术语", "body", 1, numbering_prefix="1.")
    data = _document(
        _paragraph("普通公文", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        _paragraph("正文已经开始并完整说明工作背景和基本情况。", "body", 1),
        glossary_title,
        glossary_item,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert glossary_item.type_id == "glossary_item"
    assert glossary_item.meta["colon_pos"] == -1


def test_title2_and_glossary_are_mode_independent() -> None:
    title2 = _paragraph("工作安排", "body", 2, bold_char_ratio=1.0)
    glossary_title = _paragraph("名词解释", "body", 3, bold_char_ratio=1.0)
    glossary_item = _paragraph("术语：说明内容", "body", 4)
    data = _document(
        _paragraph("普通公文", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        _paragraph("正文已经开始并完整说明工作背景和基本情况。", "body", 1),
        title2,
        glossary_title,
        glossary_item,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert [title2.type_id, glossary_title.type_id, glossary_item.type_id] == [
        "title2", "glossary_title", "glossary_item",
    ]
    assert title2.meta["recognition_provider"].startswith("body-title:")
    assert glossary_title.meta["recognition_provider"].startswith("glossary:")
    assert glossary_item.meta["recognition_provider"].startswith("glossary:")


def test_long_colon_glossary_item_is_not_claimed_by_body_structure() -> None:
    glossary_title = _paragraph("名词解释与说明", "body", 2)
    glossary_item = _paragraph(
        "1.术语：这是一个较长的解释内容，用于验证名词解释条目在冒号解释结构下不会被普通正文候选抢占。",
        "body",
        3,
        numbering_prefix="1.",
    )
    data = _document(
        _paragraph("普通公文", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        _paragraph("正文已经开始并完整说明工作背景和基本情况。", "body", 1),
        glossary_title,
        glossary_item,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert glossary_item.type_id == "glossary_item"
    assert glossary_item.meta["recognition_provider"].startswith("glossary:")


def test_glossary_items_continue_after_one_body_continuation() -> None:
    first_item = _paragraph("三个规定：解释正文内容。", "body", 3)
    continuation = _paragraph(
        "上一条名词解释的续行内容继续说明适用范围和具体工作要求。",
        "body",
        4,
    )
    second_item = _paragraph(
        "帮教模式：解释正文内容继续说明具体组成和工作方式。",
        "body",
        5,
    )
    data = _document(
        _paragraph("人民检察院工作报告", "title", 0, alignment="CENTER", bold_char_ratio=1.0),
        _paragraph("正文已经开始并完整说明工作背景和基本情况。", "body", 1),
        _paragraph("名词解释与说明", "body", 2),
        first_item,
        continuation,
        second_item,
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_core_candidates=False, enable_legacy_candidates=False),
    )

    assert [first_item.type_id, continuation.type_id, second_item.type_id] == [
        "glossary_item", "body", "glossary_item",
    ]
