import pytest
from docx import Document

from docxtool.document.engine.heading_body_split import (
    apply_heading1_report_split,
    insert_paragraph_after,
    verify_inline_heading_body_pair,
)
from docxtool.document.engine.style_catalog import ensure_document_styles
from docxtool.document.style_config import ExportError, PageSettings, StyleRule


def test_insert_paragraph_after_places_new_paragraph_next_to_source() -> None:
    doc = Document()
    first = doc.add_paragraph("第一段")
    second = insert_paragraph_after(first)
    second.add_run("第二段")

    assert [paragraph.text for paragraph in doc.paragraphs] == ["第一段", "第二段"]


def test_apply_heading1_report_split_outputs_heading_and_single_body_paragraph() -> None:
    doc = Document()
    ensure_document_styles(doc, [StyleRule.default_for_row(i) for i in range(24)], PageSettings())
    paragraph = doc.add_paragraph("一、提高认识。这里是完整正文")
    heading_rule = StyleRule.default_for_row(1)
    body_rule = StyleRule.default_for_row(5)

    body_paragraph = apply_heading1_report_split(
        paragraph,
        paragraph.text,
        heading_rule,
        body_rule,
        560,
        remove_heading_period=True,
    )

    assert body_paragraph is not None
    assert [item.text for item in doc.paragraphs] == ["一、提高认识", "这里是完整正文"]
    assert body_paragraph.style.style_id == "DCT-Body"


def test_apply_heading1_report_split_ignores_short_or_missing_body() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("一、提高认识。短")

    result = apply_heading1_report_split(
        paragraph,
        paragraph.text,
        StyleRule.default_for_row(1),
        StyleRule.default_for_row(5),
        560,
    )

    assert result is None
    assert [item.text for item in doc.paragraphs] == ["一、提高认识。短"]


def test_verify_inline_heading_body_pair_rejects_non_adjacent_or_changed_body() -> None:
    doc = Document()
    ensure_document_styles(doc, [StyleRule.default_for_row(i) for i in range(24)], PageSettings())
    heading = doc.add_paragraph("一、标题")
    body = doc.add_paragraph("正文")
    body.style = "DCT-Body"

    verify_inline_heading_body_pair(heading, body, "正文")

    doc.add_paragraph("尾段")
    with pytest.raises(ExportError):
        verify_inline_heading_body_pair(heading, doc.paragraphs[-1], "正文")

    with pytest.raises(ExportError):
        verify_inline_heading_body_pair(heading, body, "被改变")
