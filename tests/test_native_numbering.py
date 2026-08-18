from __future__ import annotations

from pathlib import Path
import zipfile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from docxtool.document.engine import export_doc
from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import load_rules_and_settings
from docxtool.security import validate_docx_integrity


W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _append_numbering_definition(
    document,
    *,
    num_id: int,
    abstract_num_id: int,
    levels: tuple[tuple[int, str, str, int], ...],
    start_override: tuple[int, int] | None = None,
) -> None:
    numbering = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_num_id))
    for ilvl, num_fmt, level_text, start_value in levels:
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), str(start_value))
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), level_text)
        level.extend((start, fmt, text))
        abstract.append(level)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    if start_override is not None:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(start_override[0]))
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), str(start_override[1]))
        override.append(start)
        num.append(override)
    first_num = numbering.find(qn("w:num"))
    first_num.addprevious(abstract)
    numbering.append(num)


def _apply_numbering(paragraph, num_id: int, ilvl: int = 0) -> None:
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), str(ilvl))
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend((level, number))
    paragraph._p.get_or_add_pPr().append(num_pr)


def _native_numbering_doc(path: Path) -> None:
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    templates = (
        ("chineseCounting", "%1、"),
        ("chineseCounting", "（%1）"),
        ("decimal", "%1."),
        ("decimal", "（%1）"),
    )
    for offset, (num_fmt, level_text) in enumerate(templates, 20):
        _append_numbering_definition(
            document,
            num_id=offset,
            abstract_num_id=offset,
            levels=((0, num_fmt, level_text, 1),),
        )

        paragraph = document.add_paragraph(f"第{offset}项工作")
        _apply_numbering(paragraph, offset)
        document.add_paragraph("后续正文对该标题展开具体说明。")
    document.save(path)


def _load(path: Path, enabled: bool):
    rules, settings, features = load_rules_and_settings(
        {"mode": "smart", "numbering": {"enabled": enabled}}
    )
    data = DocxImporter().load(str(path), rules, features=features)
    return rules, settings, features, data


def test_native_numbering_templates_define_heading_levels_without_bold(tmp_path: Path) -> None:
    source = tmp_path / "native-headings.docx"
    _native_numbering_doc(source)

    _rules, _settings, _features, data = _load(source, False)
    native = [item for item in data.paragraphs if item.features.native_numbering]

    assert [item.type_id for item in native] == [
        "heading1", "heading2", "heading3", "heading4"
    ]
    assert [item.features.native_numbering.ordinal for item in native] == [1, 1, 1, 1]
    assert all(item.features.bold_char_ratio == 0 for item in native)
    assert data.recognition_diagnostics["summary"]["source_native_heading_clue_count"] == 4
    assert data.recognition_diagnostics["summary"]["source_native_heading_unpreserved_count"] == 0


def test_native_numbering_is_preserved_when_heading_normalization_is_disabled(tmp_path: Path) -> None:
    source = tmp_path / "native-headings.docx"
    output = tmp_path / "preserved.docx"
    _native_numbering_doc(source)
    rules, settings, features, data = _load(source, False)

    stats = export_doc(
        data, rules, settings, str(output), numbering_options=features["numbering"]
    )

    with zipfile.ZipFile(output) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        numbering_xml = etree.fromstring(archive.read("word/numbering.xml"))
    assert len(document_xml.findall(".//w:numPr", W_NS)) == 4
    referenced = {
        item.get(qn("w:val"))
        for item in document_xml.findall(".//w:numPr/w:numId", W_NS)
    }
    defined = {
        item.get(qn("w:numId")) for item in numbering_xml.findall("w:num", W_NS)
    }
    assert referenced <= defined
    assert stats["native_numbering_preserved"] == 4
    assert validate_docx_integrity(output).ok is True


def test_native_heading_numbering_is_rebuilt_once_when_enabled(tmp_path: Path) -> None:
    source = tmp_path / "native-headings.docx"
    output = tmp_path / "rebuilt.docx"
    _native_numbering_doc(source)
    rules, settings, features, data = _load(source, True)

    stats = export_doc(
        data, rules, settings, str(output), numbering_options=features["numbering"]
    )

    document = Document(output)
    heading_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.style.style_id.startswith("DCT-Heading")
    ]
    headings = [paragraph.text for paragraph in heading_paragraphs]
    with zipfile.ZipFile(output) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
    assert headings == ["一、第20项工作", "（一）第21项工作", "1.第22项工作", "（1）第23项工作"]
    prefixes = ("一、", "（一）", "1.", "（1）")
    expected_bold = (False, True, True, False)
    for paragraph, prefix, bold in zip(
        heading_paragraphs, prefixes, expected_bold
    ):
        prefix_runs = []
        prefix_text = ""
        for run in paragraph.runs:
            prefix_runs.append(run)
            prefix_text += run.text
            if len(prefix_text) >= len(prefix):
                break
        assert prefix_text == prefix
        if bold:
            assert all(run.bold is True for run in prefix_runs)
        else:
            assert all(run.bold is not True for run in prefix_runs)
        body_runs = [run for run in paragraph.runs if run.text and not run.text.startswith(prefix)]
        assert body_runs
        if bold:
            assert all(run.bold is True for run in body_runs)
        else:
            assert all(run.bold is not True for run in body_runs)
    assert document_xml.findall(".//w:numPr", W_NS) == []
    assert stats["native_numbering_preserved"] == 0


def test_missing_native_numbering_definition_fails_import(tmp_path: Path) -> None:
    source = tmp_path / "broken-native-numbering.docx"
    document = Document()
    paragraph = document.add_paragraph("标题")
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "9999")
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)
    document.save(source)
    rules, _settings, features = load_rules_and_settings({"mode": "smart"})

    with pytest.raises(ValueError, match="WPS_NATIVE_NUMBERING_NUM_MISSING"):
        DocxImporter().load(str(source), rules, features=features)


def test_native_numbering_is_inherited_from_paragraph_style(tmp_path: Path) -> None:
    source = tmp_path / "style-numbering.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    _append_numbering_definition(
        document,
        num_id=60,
        abstract_num_id=60,
        levels=((0, "chineseCounting", "（%1）", 1),),
    )
    style = document.styles.add_style("NativeHeadingStyle", 1)
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "60")
    num_pr.extend((level, num_id))
    style.element.get_or_add_pPr().append(num_pr)
    document.add_paragraph("样式继承的标题", style="NativeHeadingStyle")
    document.save(source)

    _rules, _settings, _features, data = _load(source, False)
    paragraph = data.paragraphs[-1]

    assert paragraph.type_id == "heading2"
    assert paragraph.features.native_numbering is not None
    assert paragraph.features.native_numbering.num_id == 60


def test_start_override_sets_actual_ordinal_without_false_conflict(tmp_path: Path) -> None:
    source = tmp_path / "start-override.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    _append_numbering_definition(
        document,
        num_id=61,
        abstract_num_id=61,
        levels=((0, "chineseCounting", "（%1）", 1),),
        start_override=(0, 5),
    )
    for text in ("第五项标题", "第六项标题"):
        paragraph = document.add_paragraph(text)
        _apply_numbering(paragraph, 61)
        document.add_paragraph("后续正文对该标题展开具体说明。")
    document.save(source)

    _rules, _settings, _features, data = _load(source, False)
    headings = [item for item in data.paragraphs if item.type_id == "heading2"]

    assert [item.features.native_numbering.ordinal for item in headings] == [5, 6]
    for heading in headings:
        evidence = heading.meta["recognition_evidence"]
        assert "numbering-starts-after-one" not in evidence


def test_native_child_numbering_restarts_after_parent_changes(tmp_path: Path) -> None:
    source = tmp_path / "nested-numbering.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    _append_numbering_definition(
        document,
        num_id=62,
        abstract_num_id=62,
        levels=(
            (0, "chineseCounting", "%1、", 1),
            (1, "chineseCounting", "（%2）", 1),
        ),
    )
    for level, text in (
        (0, "父标题一"),
        (1, "子标题一"),
        (1, "子标题二"),
        (0, "父标题二"),
        (1, "新父级下的子标题一"),
    ):
        paragraph = document.add_paragraph(text)
        _apply_numbering(paragraph, 62, level)
    document.save(source)

    _rules, _settings, _features, data = _load(source, False)
    numbered = [item for item in data.paragraphs if item.features.native_numbering]

    assert [item.features.native_numbering.ordinal for item in numbered] == [1, 1, 2, 2, 1]


def test_native_numbering_family_supports_unformatted_custom_template(tmp_path: Path) -> None:
    source = tmp_path / "custom-family.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    _append_numbering_definition(
        document,
        num_id=63,
        abstract_num_id=63,
        levels=((0, "decimal", "第%1节", 1),),
    )
    first = document.add_paragraph("首个自定义标题", style="Heading 3")
    _apply_numbering(first, 63)
    second = document.add_paragraph("后续无格式标题")
    _apply_numbering(second, 63)
    document.save(source)

    _rules, _settings, _features, data = _load(source, False)
    numbered = [item for item in data.paragraphs if item.features.native_numbering]

    assert [item.type_id for item in numbered] == ["heading3", "heading3"]
    second_evidence = numbered[1].meta["recognition_evidence"]
    assert "native-numbering-family-sibling" in second_evidence


def test_numbering_normalization_keeps_native_body_list(tmp_path: Path) -> None:
    source = tmp_path / "body-list.docx"
    output = tmp_path / "body-list-output.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    _append_numbering_definition(
        document,
        num_id=64,
        abstract_num_id=64,
        levels=((0, "decimal", "%1.", 1),),
    )
    body = document.add_paragraph(
        "这是一个内容完整且长度明显超过四十个可见字符的自动列表正文，用于验证开启序号规范后仍保留原生列表编号而不升级为标题。"
    )
    _apply_numbering(body, 64)
    document.save(source)
    rules, settings, features, data = _load(source, True)

    assert data.paragraphs[-1].type_id == "body"
    stats = export_doc(
        data, rules, settings, str(output), numbering_options=features["numbering"]
    )

    with zipfile.ZipFile(output) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
    assert len(document_xml.findall(".//w:numPr", W_NS)) == 1
    assert stats["native_numbering_preserved"] == 1
    assert validate_docx_integrity(output).ok is True


def test_colon_introduced_native_list_stays_body(tmp_path: Path) -> None:
    source = tmp_path / "colon-list.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    document.add_paragraph("具体事项如下：")
    _append_numbering_definition(
        document,
        num_id=65,
        abstract_num_id=65,
        levels=((0, "decimal", "%1.", 1),),
    )
    for text in ("第一项", "第二项"):
        paragraph = document.add_paragraph(text)
        _apply_numbering(paragraph, 65)
    document.save(source)

    _rules, _settings, _features, data = _load(source, False)
    numbered = [item for item in data.paragraphs if item.features.native_numbering]

    assert [item.type_id for item in numbered] == ["body", "body"]


def test_native_heading2_colon_inline_body_is_recognized_and_rebuilt_once(
    tmp_path: Path,
) -> None:
    source = tmp_path / "native-heading2-colon.docx"
    preserved_output = tmp_path / "native-heading2-colon-preserved.docx"
    output = tmp_path / "native-heading2-colon-output.docx"
    document = Document()
    _append_numbering_definition(
        document,
        num_id=66,
        abstract_num_id=66,
        levels=((0, "chineseCounting", "（%1）", 1),),
    )
    paragraph = document.add_paragraph("工作安排：具体内容")
    _apply_numbering(paragraph, 66)
    document.save(source)

    preserved_rules, preserved_settings, preserved_features, preserved_data = _load(
        source, False
    )
    assert preserved_data.paragraphs[0].type_id == "heading2"
    assert preserved_data.paragraphs[0].meta[
        "numbered_heading2_colon_inline_body"
    ] is True
    export_doc(
        preserved_data,
        preserved_rules,
        preserved_settings,
        str(preserved_output),
        numbering_options=preserved_features["numbering"],
    )
    with zipfile.ZipFile(preserved_output) as archive:
        preserved_xml = etree.fromstring(archive.read("word/document.xml"))
    assert len(preserved_xml.findall(".//w:numPr", W_NS)) == 1
    assert validate_docx_integrity(preserved_output).ok is True

    rules, settings, features, data = _load(source, True)

    recognized = data.paragraphs[0]
    assert recognized.type_id == "heading2"
    assert recognized.meta["numbered_heading2_colon_inline_body"] is True
    assert "numbered-heading2-colon-inline-body" in recognized.meta[
        "recognition_evidence"
    ]

    export_doc(
        data, rules, settings, str(output), numbering_options=features["numbering"]
    )

    rendered = Document(output).paragraphs[0]
    assert rendered.text == "（一）工作安排：具体内容"
    colon_position = rendered.text.index("：")
    cursor = 0
    for run in rendered.runs:
        if not run.text:
            continue
        cursor += len(run.text)
        expected_font = rules[2].font if cursor <= colon_position + 1 else rules[5].font
        assert run._r.rPr.rFonts.get(qn("w:eastAsia")) == expected_font
    with zipfile.ZipFile(output) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
    assert document_xml.findall(".//w:numPr", W_NS) == []
    assert validate_docx_integrity(output).ok is True


def test_native_heading2_period_body_stays_one_paragraph(tmp_path: Path) -> None:
    source = tmp_path / "native-heading2-period.docx"
    output = tmp_path / "native-heading2-period-output.docx"
    document = Document()
    _append_numbering_definition(
        document,
        num_id=67,
        abstract_num_id=67,
        levels=((0, "chineseCounting", "（%1）", 1),),
    )
    paragraph = document.add_paragraph(
        "会议安排。后续正文内容完整说明有关工作要求。"
    )
    _apply_numbering(paragraph, 67)
    document.save(source)
    rules, settings, features, data = _load(source, False)

    assert len(data.paragraphs) == 1
    assert data.paragraphs[0].type_id == "heading2"
    assert data.paragraphs[0].meta["numbered_heading2_period_inline_body"] is True
    assert "numbered-heading2-period-inline-body" in data.paragraphs[0].meta[
        "recognition_evidence"
    ]

    export_doc(
        data, rules, settings, str(output), numbering_options=features["numbering"]
    )
    rendered = Document(output)
    assert len(rendered.paragraphs) == 1
    assert rendered.paragraphs[0].text == "会议安排。后续正文内容完整说明有关工作要求。"
    with zipfile.ZipFile(output) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
    assert len(document_xml.findall(".//w:numPr", W_NS)) == 1
    assert validate_docx_integrity(output).ok is True
