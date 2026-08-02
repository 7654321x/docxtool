from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from docxtool.document.importer import DocumentData, ParagraphData, ParagraphFeatures
from docxtool.document.source_tape import (
    HOST_TEXT_CONTRACT_VERSION,
    SourceTape,
    UnknownTextContractVersion,
    canonicalize_text,
    canonicalize_host_paragraph_text,
)
from docxtool.sdk import bind_recognition_plan, recognize_docx
from docxtool.sdk.cli import main as sdk_main
from docxtool.sdk.recognition import _build_plan


def _write_document(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for value in paragraphs:
        document.add_paragraph(value)
    document.save(path)


def _snapshot(paragraphs: list[dict] | list[str], *, snapshot_id: str = "snap-1") -> dict:
    items = []
    for ordinal, value in enumerate(paragraphs):
        if isinstance(value, dict):
            raw_text = value["raw_text"]
            host_index = value.get("host_paragraph_index", ordinal)
        else:
            raw_text = value
            host_index = ordinal
        items.append({
            "host_paragraph_id": "main:{0:06d}".format(host_index),
            "host_paragraph_index": host_index,
            "story_id": "main",
            "story_type": "main",
            "story_paragraph_index": ordinal,
            "section_index": 0,
            "is_in_table": False,
            "raw_text": raw_text,
        })
    return {
        "schema_version": "host-snapshot-v1",
        "integration_contract_version": "integration-contract-v1",
        "snapshot_id": snapshot_id,
        "document_identity": "local-only",
        "document_revision": "rev-1",
        "host": {"kind": "test-host"},
        "text_contract_version": "host-text-v1",
        "offset_encoding": "utf16_code_unit",
        "paragraphs": items,
    }


def test_source_tape_keeps_utf16_and_canonical_boundaries_reversible() -> None:
    raw = "甲\u00a0😀\r\nＢ"
    tape = SourceTape.from_text(raw)

    assert tape.canonical_text == "甲 😀\nB"
    raw_start = tape.raw_offset_utf16(1)
    raw_end = tape.raw_offset_utf16(3)
    assert raw_start is not None and raw_end is not None
    canonical_range = tape.canonical_range_for_raw_span(1, 3)
    assert canonical_range is not None
    assert tape.raw_span_for_canonical_range(*canonical_range) == (1, 3)
    assert tape.raw_slice_utf16(raw_start, raw_end) == "\u00a0😀"


def test_host_text_contract_normalizes_controls_with_a_versioned_mapping() -> None:
    result = canonicalize_host_paragraph_text("甲\r\n乙\r丙\n丁\v戊\f己\t庚\u00a0\u3000辛😀\x07")

    assert result.contract_version == HOST_TEXT_CONTRACT_VERSION
    assert result.canonical_text == "甲\n乙\n丙\n丁\n戊\f己\t庚  辛😀"
    assert len(result.raw_to_canonical_utf16) == len(result.raw_text) + 1
    assert "TRAILING_TABLE_CELL_MARKER_PRESENT" in result.warnings
    with pytest.raises(UnknownTextContractVersion, match="不支持"):
        canonicalize_host_paragraph_text("正文", contract_version="unknown-v9")


def test_segment_counts_report_incomplete_source_locator_group() -> None:
    source = "甲乙丙"

    def feature(start: int | None, end: int | None, status: str) -> ParagraphFeatures:
        raw = source[start:end] if start is not None and end is not None else ""
        return ParagraphFeatures(
            source_physical_paragraph_index=0,
            source_physical_text=source,
            source_start_utf16=start,
            source_end_utf16=end,
            source_canonical_text=source,
            source_canonical_start_utf16=start,
            source_canonical_end_utf16=end,
            source_fragment_text=raw,
            source_canonical_fragment_text=raw,
            source_locator_status=status,
        )

    data = DocumentData(paragraphs=[
        ParagraphData("甲", "body", "甲", feature(0, 1, "confirmed")),
        ParagraphData("乙", "body", "乙", feature(None, None, "unresolved")),
        ParagraphData("丙", "body", "丙", feature(2, 3, "confirmed")),
    ])
    plan = _build_plan(data, "0" * 64, include_text=True)

    assert [block.segment_index for block in plan.blocks] == [0, 2, 1]
    assert [block.block_index for block in plan.blocks] == [0, 1, 2]
    assert {block.segment_count_total for block in plan.blocks} == {3}
    assert {block.segment_count_located for block in plan.blocks} == {2}
    assert {block.segment_count_confirmed for block in plan.blocks} == {2}
    assert [block.source_locator_status for block in plan.blocks] == ["confirmed", "unresolved", "confirmed"]


def test_partial_locator_group_keeps_located_source_order_and_stable_unresolved_order() -> None:
    source = "甲乙丙丁"

    def feature(start: int | None, end: int | None, status: str) -> ParagraphFeatures:
        raw = source[start:end] if start is not None and end is not None else ""
        return ParagraphFeatures(
            source_physical_paragraph_index=0,
            source_physical_text=source,
            source_start_utf16=start,
            source_end_utf16=end,
            source_canonical_text=source,
            source_canonical_start_utf16=start,
            source_canonical_end_utf16=end,
            source_fragment_text=raw,
            source_canonical_fragment_text=raw,
            source_locator_status=status,
        )

    data = DocumentData(paragraphs=[
        ParagraphData("丙", "attachment_note", "丙", feature(2, 3, "confirmed")),
        ParagraphData("未定位一", "body", "未定位一", feature(None, None, "unresolved")),
        ParagraphData("甲", "sign_org", "甲", feature(0, 1, "confirmed")),
        ParagraphData("未定位二", "body", "未定位二", feature(None, None, "unresolved")),
    ])
    plan = _build_plan(data, "0" * 64, include_text=True)

    assert [block.block_index for block in plan.blocks] == [0, 1, 2, 3]
    assert [block.segment_index for block in plan.blocks] == [1, 2, 0, 3]
    assert [block.source_locator_status for block in plan.blocks] == [
        "confirmed",
        "unresolved",
        "confirmed",
        "unresolved",
    ]
    assert {block.segment_count_total for block in plan.blocks} == {4}
    assert {block.segment_count_located for block in plan.blocks} == {2}
    assert {block.segment_count_confirmed for block in plan.blocks} == {2}
    assert type(plan).from_dict(plan.to_dict()).to_dict() == plan.to_dict()


def test_partial_locator_group_detects_overlap_only_between_located_ranges() -> None:
    source = "甲乙丙丁"

    def feature(start: int | None, end: int | None, status: str) -> ParagraphFeatures:
        raw = source[start:end] if start is not None and end is not None else ""
        return ParagraphFeatures(
            source_physical_paragraph_index=0,
            source_physical_text=source,
            source_start_utf16=start,
            source_end_utf16=end,
            source_canonical_text=source,
            source_canonical_start_utf16=start,
            source_canonical_end_utf16=end,
            source_fragment_text=raw,
            source_canonical_fragment_text=raw,
            source_locator_status=status,
        )

    data = DocumentData(paragraphs=[
        ParagraphData("乙丙", "body", "乙丙", feature(1, 3, "confirmed")),
        ParagraphData("未定位", "body", "未定位", feature(None, None, "unresolved")),
        ParagraphData("甲乙", "body", "甲乙", feature(0, 2, "confirmed")),
    ])
    plan = _build_plan(data, "source-hash", include_text=True)

    assert [block.segment_index for block in plan.blocks] == [1, 2, 0]
    assert plan.blocks[2].source_locator_status == "confirmed"
    assert "SOURCE_RANGE_OVERLAP" not in plan.blocks[2].source_locator_warnings
    assert plan.blocks[0].source_locator_status == "unresolved"
    assert "SOURCE_RANGE_OVERLAP" in plan.blocks[0].source_locator_warnings
    assert plan.blocks[1].source_locator_status == "unresolved"
    assert "SOURCE_RANGE_OVERLAP" not in plan.blocks[1].source_locator_warnings


def test_sdk_uses_source_range_order_when_normalization_reorders_blocks() -> None:
    source = "甲乙丙"

    def feature(start: int, end: int) -> ParagraphFeatures:
        raw = source[start:end]
        return ParagraphFeatures(
            source_physical_paragraph_index=0,
            source_physical_text=source,
            source_start_utf16=start,
            source_end_utf16=end,
            source_canonical_text=source,
            source_canonical_start_utf16=start,
            source_canonical_end_utf16=end,
            source_fragment_text=raw,
            source_canonical_fragment_text=raw,
            source_locator_status="confirmed",
        )

    data = DocumentData(paragraphs=[
        ParagraphData("丙", "attachment_note", "丙", feature(2, 3)),
        ParagraphData("甲", "sign_org", "甲", feature(0, 1)),
        ParagraphData("乙", "sign_date", "乙", feature(1, 2)),
    ])
    plan = _build_plan(data, "source-hash", include_text=True)

    assert [block.type_id for block in plan.blocks] == [
        "attachment_note",
        "sign_org",
        "sign_date",
    ]
    assert [block.segment_index for block in plan.blocks] == [2, 0, 1]
    assert all(block.locator_verified for block in plan.blocks)
    assert all("SOURCE_RANGE_OVERLAP" not in block.source_locator_warnings for block in plan.blocks)


def test_sdk_splits_one_physical_paragraph_into_ordered_verified_segments(tmp_path: Path) -> None:
    source = tmp_path / "mixed.docx"
    raw = "一、总体要求。总体要求应覆盖😀、\t空格和重复文字。"
    _write_document(source, [raw])

    plan = recognize_docx(source, recognition_mode="legacy", include_text=True, include_raw_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert len(blocks) == 2
    assert [block.segment_index for block in blocks] == [0, 1]
    assert {block.segment_count for block in blocks} == {2}
    assert [block.type_id for block in blocks] == ["heading1", "body"]
    assert all(block.source_locator_status == "confirmed" for block in blocks)
    assert all(block.locator_verified for block in blocks)
    assert blocks[0].raw_end_utf16 <= blocks[1].raw_start_utf16
    physical = raw.encode("utf-16-le")
    for block in blocks:
        selected = physical[block.raw_start_utf16 * 2:block.raw_end_utf16 * 2].decode("utf-16-le")
        assert selected == block.raw_fragment_text
        assert canonicalize_text(selected) == block.canonical_fragment_text
        assert block.raw_fragment_sha256
        assert block.canonical_fragment_sha256


def test_mixed_physical_paragraph_uses_run_intersection_format_features(tmp_path: Path) -> None:
    source = tmp_path / "mixed-format.docx"
    document = Document()
    paragraph = document.add_paragraph()
    heading = paragraph.add_run("一、总体要求。")
    heading.font.name = "SimHei"
    heading.font.size = Pt(16)
    heading.font.bold = True
    body = paragraph.add_run("各单位要认真贯彻执行，形成长效机制。")
    body.font.name = "FangSong"
    body.font.size = Pt(12)
    body.font.bold = False
    document.save(source)

    plan = recognize_docx(source, recognition_mode="legacy")
    heading_block, body_block = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert [heading_block.type_id, body_block.type_id] == ["heading1", "body"]
    assert heading_block.segment_format["run_count"] == 1
    assert body_block.segment_format["run_count"] == 1
    assert heading_block.segment_format["bold_char_ratio"] == 1.0
    assert body_block.segment_format["bold_char_ratio"] == 0.0
    assert heading_block.segment_format["weighted_font_size_pt"] == 16.0
    assert body_block.segment_format["weighted_font_size_pt"] == 12.0
    assert heading_block.segment_format != body_block.segment_format


def test_host_binding_uses_monotonic_order_and_handles_shifted_paragraphs(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_document(source, ["重复段落", "一、总体要求。正文内容不少于五个字。", "重复段落"])
    plan = recognize_docx(source, recognition_mode="legacy")

    binding = bind_recognition_plan(plan, _snapshot([
        {"host_paragraph_index": 9, "raw_text": "额外段落"},
        {"host_paragraph_index": 10, "raw_text": "重复段落"},
        {"host_paragraph_index": 11, "raw_text": "一、总体要求。正文内容不少于五个字。"},
        {"host_paragraph_index": 12, "raw_text": "重复段落"},
    ]))

    assert all(item.binding_status == "confirmed" for item in binding.blocks)
    assert [item.host_paragraph_index for item in binding.blocks] == [10, 11, 11, 12]
    assert all(item.host_raw_end_utf16 is not None for item in binding.blocks)
    assert all(item.host_canonical_end_utf16 is not None for item in binding.blocks)
    assert any("DUPLICATE_TEXT_DISAMBIGUATED" in item.binding_evidence for item in binding.blocks)


def test_host_binding_maps_canonical_text_to_host_raw_text_without_reusing_offsets(tmp_path: Path) -> None:
    source = tmp_path / "canonical.docx"
    _write_document(source, ["正文\u00a0内容😀"])
    plan = recognize_docx(source, recognition_mode="legacy")

    binding = bind_recognition_plan(plan, _snapshot(["正文 内容😀"]))

    assert binding.blocks[0].binding_status == "review"
    assert binding.host_text_contract_version == HOST_TEXT_CONTRACT_VERSION
    assert binding.blocks[0].host_canonical_start_utf16 == 0
    assert binding.blocks[0].host_canonical_end_utf16 == len("正文 内容😀".encode("utf-16-le")) // 2
    assert "PHYSICAL_CANONICAL_TEXT_MATCH" in binding.blocks[0].binding_evidence
    assert "RAW_TEXT_NORMALIZED" in binding.blocks[0].binding_warnings


def test_ambiguous_repeat_only_blocks_its_own_physical_paragraph_group(tmp_path: Path) -> None:
    source = tmp_path / "local-ambiguity.docx"
    _write_document(source, ["唯一开头", "重复内容", "唯一结尾"])
    plan = recognize_docx(source, recognition_mode="legacy")

    binding = bind_recognition_plan(plan, _snapshot([
        {"host_paragraph_index": 10, "raw_text": "唯一开头"},
        {"host_paragraph_index": 11, "raw_text": "重复内容"},
        {"host_paragraph_index": 12, "raw_text": "重复内容"},
        {"host_paragraph_index": 13, "raw_text": "唯一结尾"},
    ]))

    assert [item.binding_status for item in binding.blocks] == ["confirmed", "unresolved", "confirmed"]
    assert [item.host_paragraph_index for item in binding.blocks] == [10, None, 13]
    assert [item.status for item in binding.physical_paragraphs] == [
        "matched_unique", "ambiguous", "matched_unique",
    ]
    ambiguous = binding.physical_paragraphs[1]
    assert ambiguous.candidate_host_paragraph_indexes == (11, 12)
    assert "SOURCE_OCCURRENCE_AMBIGUOUS" in ambiguous.warnings


def test_repeated_paragraphs_with_distinct_contexts_remain_confirmed(tmp_path: Path) -> None:
    source = tmp_path / "context-repeat.docx"
    _write_document(source, ["文首", "重复内容", "中间锚点", "重复内容", "文末"])
    plan = recognize_docx(source, recognition_mode="legacy")

    binding = bind_recognition_plan(plan, _snapshot(["文首", "重复内容", "中间锚点", "重复内容", "文末"]))

    assert all(item.binding_status == "confirmed" for item in binding.blocks)
    assert all(item.status == "matched_unique" for item in binding.physical_paragraphs)


def test_host_binding_refuses_ambiguous_duplicate_occurrences(tmp_path: Path) -> None:
    source = tmp_path / "ambiguous.docx"
    _write_document(source, ["相同内容"])
    plan = recognize_docx(source, recognition_mode="legacy")

    binding = bind_recognition_plan(plan, _snapshot(["相同内容", "相同内容"]))

    assert binding.blocks[0].binding_status == "unresolved"
    assert "SOURCE_OCCURRENCE_AMBIGUOUS" in binding.blocks[0].binding_warnings


def test_sdk_cli_can_bind_a_local_snapshot_without_default_text_output(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "plan.json"
    snapshot = tmp_path / "snapshot.json"
    _write_document(source, ["普通正文内容"])
    snapshot.write_text(json.dumps(_snapshot(["普通正文内容"]), ensure_ascii=False), encoding="utf-8")

    assert sdk_main([str(source), "--host-snapshot", str(snapshot), "--output", str(output)]) == 0
    payload = json.loads(output.read_text(encoding="utf-8"))
    assert payload["binding"]["blocks"][0]["binding_status"] == "confirmed"
    assert "recognized_text" not in payload["data"]["blocks"][0]
