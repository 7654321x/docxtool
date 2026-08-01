from __future__ import annotations

import json
from pathlib import Path
from dataclasses import replace

import pytest
from docx import Document

from docxtool.sdk import (
    DocxToolSdkError,
    InvalidHostSnapshotError,
    InvalidRequestError,
    RecognitionRequest,
    bind_recognition_plan,
    get_json_schema,
    get_sdk_manifest,
    host_snapshot_from_dict,
    load_json_schema,
    recognition_request_from_dict,
    recognition_plan_from_dict,
    summarize_host_snapshot,
    recognize_docx,
    validate_host_snapshot,
    validate_recognition_binding,
    validate_recognition_plan,
    validate_recognition_request,
    validate_sdk_error,
    validate_sdk_manifest,
)
from docxtool.sdk.cli import main as sdk_cli
from docxtool.sdk.constants import SCHEMA_FILENAMES


def _write_document(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for value in paragraphs:
        document.add_paragraph(value)
    document.save(path)


def _snapshot(raw_text: str, *, kind: str = "test-host", snapshot_id: str = "snap-1") -> dict:
    return {
        "schema_version": "host-snapshot-v1",
        "integration_contract_version": "integration-contract-v1",
        "snapshot_id": snapshot_id,
        "document_identity": "local-doc",
        "document_revision": "rev-1",
        "host": {"kind": kind},
        "text_contract_version": "host-text-v1",
        "offset_encoding": "utf16_code_unit",
        "paragraphs": [
            {
                "host_paragraph_id": "main:000000",
                "host_paragraph_index": 0,
                "story_id": "main",
                "story_type": "main",
                "story_paragraph_index": 0,
                "section_index": 0,
                "is_in_table": False,
                "raw_text": raw_text,
            }
        ],
    }


def test_manifest_and_schema_files_are_public_and_valid() -> None:
    manifest = get_sdk_manifest().to_dict()

    assert manifest["schema_version"] == "sdk-manifest-v1"
    assert manifest["integration_contract_versions"] == ["integration-contract-v1"]
    assert "host_snapshot_binding" in manifest["capabilities"]
    assert manifest["binding_scope"]["supported_story_types"] == ["main"]
    assert manifest["binding_scope"]["tables"] == "excluded"
    assert validate_sdk_manifest(manifest).valid
    assert get_json_schema("recognition-request")["title"]
    for schema_version in SCHEMA_FILENAMES:
        schema = load_json_schema(schema_version)
        assert schema["$schema"].startswith("https://json-schema.org/")
        assert schema["$id"].endswith(SCHEMA_FILENAMES[schema_version])
        assert schema["title"]


def test_recognition_request_supports_round_trip_and_raw_text_implies_text() -> None:
    request = RecognitionRequest.from_dict({
        "schema_version": "recognition-request-v1",
        "processing_mode": "structural",
        "recognition_mode": "authoritative",
        "include_text": False,
        "include_raw_text": True,
        "format_config": None,
        "feature_overrides": None,
    })

    assert request.include_text is True
    assert request.include_raw_text is True
    assert validate_recognition_request(request).valid
    assert recognition_request_from_dict(request.to_dict()).to_dict() == request.to_dict()


def test_recognition_plan_round_trip_schema_validation_and_privacy(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_document(source, ["一、总体要求。正文内容不少于五个字。"])

    plan = recognize_docx(source, recognition_mode="legacy")
    assert validate_recognition_plan(plan).valid
    payload = plan.to_dict()
    restored = recognition_plan_from_dict(json.loads(json.dumps(payload, ensure_ascii=False)))

    assert restored.plan_id == plan.plan_id
    assert [block.block_id for block in restored.blocks] == [block.block_id for block in plan.blocks]
    assert "正文内容" not in json.dumps(payload, ensure_ascii=False)
    assert payload["blocks"][0]["source_locator"]["raw_span"]["encoding"] == "utf16_code_unit"
    assert payload["blocks"][0]["range_coordinate_system"] == "source_raw_text_utf16"


def test_host_snapshot_v1_validation_rejects_missing_and_duplicate_ids() -> None:
    valid = _snapshot("普通正文")
    report = validate_host_snapshot(valid)
    snapshot = host_snapshot_from_dict(valid)

    assert report.valid
    assert snapshot.snapshot_id == "snap-1"
    assert snapshot.paragraphs[0].host_paragraph_id == "main:000000"

    missing_snapshot_id = dict(valid)
    missing_snapshot_id.pop("snapshot_id")
    report = validate_host_snapshot(missing_snapshot_id)
    assert not report.valid
    assert report.errors[0].code == "INVALID_HOST_SNAPSHOT"
    with pytest.raises(InvalidHostSnapshotError, match="snapshot_id"):
        host_snapshot_from_dict(missing_snapshot_id)

    duplicate = dict(valid)
    duplicate["paragraphs"] = [dict(valid["paragraphs"][0]), dict(valid["paragraphs"][0])]
    duplicate["paragraphs"][1]["host_paragraph_index"] = 1
    report = validate_host_snapshot(duplicate)
    assert not report.valid
    assert report.errors[0].code == "INVALID_HOST_SNAPSHOT"


def test_legacy_host_snapshot_is_rejected_by_public_binding_contract() -> None:
    legacy = {"host_type": "wps", "paragraphs": [{"raw_text": "普通正文"}]}

    report = validate_host_snapshot(legacy)
    assert not report.valid
    with pytest.raises(InvalidHostSnapshotError):
        host_snapshot_from_dict(legacy)


def test_binding_is_host_neutral_and_contains_preconditions(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    raw = "普通正文内容"
    _write_document(source, [raw])
    plan = recognize_docx(source, recognition_mode="legacy")

    bindings = []
    for kind in ("wps", "microsoft-word", "test-host"):
        binding = bind_recognition_plan(plan.to_dict(), _snapshot(raw, kind=kind, snapshot_id="same-snap"))
        assert validate_recognition_binding(binding).valid
        bindings.append(binding.to_dict())

    first_blocks = bindings[0]["blocks"]
    for payload in bindings[1:]:
        assert [block["block_id"] for block in payload["blocks"]] == [block["block_id"] for block in first_blocks]
        assert [
            block["binding"]["status"] for block in payload["blocks"]
        ] == [block["binding"]["status"] for block in first_blocks]
        assert [
            block["host_target"]["raw_span"] for block in payload["blocks"]
        ] == [block["host_target"]["raw_span"] for block in first_blocks]
        assert [
            block["preconditions"]["raw_fragment_sha256"] for block in payload["blocks"]
        ] == [block["preconditions"]["raw_fragment_sha256"] for block in first_blocks]
    confirmed = bindings[0]["blocks"][0]
    assert confirmed["binding"]["recommended_action"] == "verify_host_range"
    assert confirmed["preconditions"]["snapshot_id"] == "same-snap"
    assert confirmed["preconditions"]["host_paragraph_id"] == "main:000000"


def test_different_snapshots_produce_different_binding_ids(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_document(source, ["普通正文内容"])
    plan = recognize_docx(source, recognition_mode="legacy")

    first = bind_recognition_plan(plan, _snapshot("普通正文内容", snapshot_id="snap-a"))
    second = bind_recognition_plan(plan, _snapshot("普通正文内容", snapshot_id="snap-b"))

    assert first.binding_id != second.binding_id


def test_binding_id_changes_when_target_span_status_or_preconditions_change(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_document(source, ["普通正文内容"])
    plan = recognize_docx(source, recognition_mode="legacy")
    binding = bind_recognition_plan(plan, _snapshot("普通正文内容", snapshot_id="snap-a"))
    block = binding.blocks[0]

    shifted = replace(binding, binding_id="", blocks=(replace(
        block,
        host_raw_end_utf16=(block.host_raw_end_utf16 or 0) + 1,
    ),))
    reviewed = replace(binding, binding_id="", blocks=(replace(
        block,
        binding_status="review",
        recommended_action="preview_only",
    ),))
    changed_hash = replace(binding, binding_id="", blocks=(replace(
        block,
        preconditions={**dict(block.preconditions), "raw_fragment_sha256": "a" * 64},
    ),))

    assert shifted.binding_id != binding.binding_id
    assert reviewed.binding_id != binding.binding_id
    assert changed_hash.binding_id != binding.binding_id


def test_request_boolean_contract_rejects_strings_and_cli_uses_same_code(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    request_path = tmp_path / "request.json"
    output_path = tmp_path / "output.json"
    _write_document(source, ["普通正文内容"])
    request = RecognitionRequest().to_dict()
    request["include_text"] = "false"
    request_path.write_text(json.dumps(request, ensure_ascii=False), encoding="utf-8")

    report = validate_recognition_request(request)
    assert not report.valid
    assert report.errors[0].code == "INVALID_RECOGNITION_REQUEST"
    with pytest.raises(InvalidRequestError):
        recognition_request_from_dict(request)
    with pytest.raises(InvalidRequestError):
        recognize_docx(source, include_text="false")

    assert sdk_cli(["recognize", "--source", str(source), "--request", str(request_path), "--output", str(output_path)]) == 2
    payload = json.loads(output_path.read_text(encoding="utf-8"))
    assert payload["ok"] is False
    assert payload["error"]["code"] == report.errors[0].code


def test_strict_mode_rejects_unknown_fields_but_non_strict_warns() -> None:
    request = RecognitionRequest().to_dict()
    request["extension_note"] = "ignored"

    non_strict = validate_recognition_request(request)
    strict = validate_recognition_request(request, strict=True)

    assert non_strict.valid
    assert non_strict.warnings[0].path == "$.extension_note"
    assert not strict.valid
    assert strict.errors[0].path == "$.extension_note"


def test_review_items_follow_block_id_when_blocks_are_reordered(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_document(source, ["第一段", "第二段"])
    plan = recognize_docx(source, recognition_mode="legacy")
    payload = plan.to_dict()
    payload["blocks"] = list(reversed(payload["blocks"]))
    target = payload["blocks"][0]
    payload["review_items"] = [{
        "block_id": target["block_id"],
        "block_index": 0,
        "source_paragraph_index": target.get("source_paragraph_index"),
        "final_type": target["semantic"]["type_id"],
        "level": "review",
        "confidence": 0.5,
        "reasons": ["TEST_REORDER"],
        "evidence": [],
    }]

    restored = recognition_plan_from_dict(payload)

    assert restored.review_items[0].block_id == target["block_id"]
    assert restored.review_items[0].block_index == target["block_index"]


def test_binding_state_invariants_are_validated(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_document(source, ["普通正文内容"])
    plan = recognize_docx(source, recognition_mode="legacy")
    payload = bind_recognition_plan(plan, _snapshot("普通正文内容")).to_dict()

    invalid_action = json.loads(json.dumps(payload, ensure_ascii=False))
    invalid_action["blocks"][0]["binding"]["recommended_action"] = "skip"
    assert not validate_recognition_binding(invalid_action).valid

    unresolved = json.loads(json.dumps(payload, ensure_ascii=False))
    unresolved["blocks"][0]["binding"]["status"] = "unresolved"
    unresolved["blocks"][0]["binding"]["recommended_action"] = "skip"
    assert not validate_recognition_binding(unresolved).valid

    review = json.loads(json.dumps(payload, ensure_ascii=False))
    review["blocks"][0]["binding"]["status"] = "review"
    review["blocks"][0]["binding"]["recommended_action"] = "verify_host_range"
    assert not validate_recognition_binding(review).valid


def test_host_snapshot_summary_is_text_free_and_not_bindable(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _write_document(source, ["普通正文内容"])
    plan = recognize_docx(source, recognition_mode="legacy")
    summary = summarize_host_snapshot(host_snapshot_from_dict(_snapshot("普通正文内容"))).to_dict()

    assert "raw_text" not in json.dumps(summary, ensure_ascii=False)
    assert summary["bindable"] is False
    assert not validate_host_snapshot(summary).valid
    with pytest.raises(DocxToolSdkError):
        bind_recognition_plan(plan, summary)


def test_cli_manifest_recognize_bind_and_validate_round_trip(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    manifest_output = tmp_path / "manifest.json"
    plan_output = tmp_path / "plan.json"
    binding_output = tmp_path / "binding.json"
    snapshot_output = tmp_path / "snapshot.json"
    validate_output = tmp_path / "valid.json"
    _write_document(source, ["普通正文内容"])
    snapshot_output.write_text(json.dumps(_snapshot("普通正文内容"), ensure_ascii=False), encoding="utf-8")

    assert sdk_cli(["manifest", "--output", str(manifest_output)]) == 0
    assert json.loads(manifest_output.read_text(encoding="utf-8"))["data"]["schema_version"] == "sdk-manifest-v1"

    assert sdk_cli(["recognize", "--source", str(source), "--output", str(plan_output)]) == 0
    plan_payload = json.loads(plan_output.read_text(encoding="utf-8"))
    assert plan_payload["data"]["schema_version"] == "recognition-plan-v1"

    assert sdk_cli([
        "bind",
        "--plan",
        str(plan_output),
        "--snapshot",
        str(snapshot_output),
        "--output",
        str(binding_output),
    ]) == 0
    binding_payload = json.loads(binding_output.read_text(encoding="utf-8"))
    assert binding_payload["data"]["schema_version"] == "recognition-binding-v1"

    assert sdk_cli([
        "validate",
        "--kind",
        "recognition-binding",
        "--input",
        str(binding_output),
        "--output",
        str(validate_output),
    ]) == 0
    assert json.loads(validate_output.read_text(encoding="utf-8"))["data"]["valid"] is True

    assert sdk_cli([
        "summarize-snapshot",
        "--snapshot",
        str(snapshot_output),
        "--output",
        str(validate_output),
    ]) == 0
    summary = json.loads(validate_output.read_text(encoding="utf-8"))["data"]
    assert summary["summary_type"] == "host-snapshot-summary-v1"
    assert "raw_text" not in json.dumps(summary, ensure_ascii=False)


def test_error_payload_matches_public_schema() -> None:
    error = DocxToolSdkError(
        "测试错误",
        code="SOURCE_TEXT_HASH_MISMATCH",
        details={"path": "blocks[0]", "raw_text": "不得输出"},
    )
    payload = error.to_dict()

    assert "raw_text" not in payload["details"]
    assert payload["code"] == "SOURCE_TEXT_HASH_MISMATCH"
    assert validate_sdk_error(payload).valid
