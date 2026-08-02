from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docxtool.document.engine.paragraph_format import (
    apply_first_line_indent_chars,
    apply_left_indent_chars,
    apply_right_indent,
    apply_rule_paragraph_format,
    apply_style,
    set_para_spacing,
    set_widow_control,
)
from docxtool.document.style_config import StyleRule


def _paragraph_properties(paragraph):
    return paragraph._element.get_or_add_pPr()


def test_apply_style_sets_body_snap_to_grid_and_first_line_indent() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("正文")
    rule = StyleRule.default_for_row(5)

    apply_style(paragraph, rule)

    properties = _paragraph_properties(paragraph)
    snap = properties.find(qn("w:snapToGrid"))
    indent = properties.find(qn("w:ind"))
    assert snap is not None and snap.get(qn("w:val")) == "1"
    assert indent is not None and indent.get(qn("w:firstLineChars")) == "200"
    assert paragraph.runs[0].font.size.pt == 16


def test_apply_style_sets_heading_snap_and_widow_control_off() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("标题")
    rule = StyleRule.default_for_row(1)

    apply_style(paragraph, rule)

    properties = _paragraph_properties(paragraph)
    snap = properties.find(qn("w:snapToGrid"))
    widow = properties.find(qn("w:widowControl"))
    assert snap is not None and snap.get(qn("w:val")) == "0"
    assert widow is not None and widow.get(qn("w:val")) == "0"


def test_indent_helpers_clear_conflicting_indent_attributes() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("测试")

    apply_left_indent_chars(paragraph, 5)
    apply_first_line_indent_chars(paragraph, 2)

    indent = _paragraph_properties(paragraph).find(qn("w:ind"))
    assert indent is not None
    assert indent.get(qn("w:firstLineChars")) == "200"
    assert qn("w:leftChars") not in indent.attrib
    assert qn("w:hangingChars") not in indent.attrib


def test_right_indent_clears_left_and_first_line_indent() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("落款")

    apply_first_line_indent_chars(paragraph, 2)
    apply_right_indent(paragraph, 4)

    indent = _paragraph_properties(paragraph).find(qn("w:ind"))
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert indent is not None and indent.get(qn("w:rightChars")) == "400"
    assert qn("w:firstLineChars") not in indent.attrib


def test_spacing_and_rule_paragraph_format_write_expected_values() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("附件标题")
    rule = StyleRule.default_for_row(20)

    apply_rule_paragraph_format(paragraph, rule, 560)
    set_para_spacing(paragraph, before_lines=0, after_lines=0, line_twips=560, explicit_zero=True)
    set_widow_control(paragraph, True)

    properties = _paragraph_properties(paragraph)
    spacing = properties.find(qn("w:spacing"))
    widow = properties.find(qn("w:widowControl"))
    assert spacing is not None
    assert spacing.get(qn("w:beforeLines")) == "0"
    assert spacing.get(qn("w:afterLines")) == "0"
    assert spacing.get(qn("w:line")) == "560"
    assert widow is not None and widow.get(qn("w:val")) == "1"
