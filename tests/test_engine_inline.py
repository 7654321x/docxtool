from dataclasses import dataclass

from docx import Document
from docx.shared import Pt

from docxtool.document.engine.inline import (
    segment_writer,
    without_redundant_trailing_body_page_breaks,
    write_inline_tokens,
)


@dataclass
class _Token:
    """测试用行内 token，提供 kind 和 text 两个字段。"""

    kind: str
    text: str = ""


@dataclass
class _ParagraphData:
    """测试用段落数据，提供分页清理函数需要读取的字段。"""

    text: str
    type_id: str


def test_segment_writer_reuses_base_style_and_applies_overrides() -> None:
    """行内片段写入器应复用首个 run 样式，并允许覆盖加粗和字号。"""
    document = Document()
    paragraph = document.add_paragraph()
    base = paragraph.add_run("旧文本")
    base.font.bold = True
    base.font.size = Pt(14)

    writer = segment_writer(paragraph)
    writer("第一段", bold=False)
    writer("第二段", size_pt=16)

    assert [run.text for run in paragraph.runs] == ["第一段", "第二段"]
    assert paragraph.runs[0].font.bold is False
    assert paragraph.runs[1].font.bold is False
    assert paragraph.runs[1].font.size == Pt(16)


def test_segment_writer_uses_injected_font_callback() -> None:
    """传入字体回调时，写入器应把中文字体设置交给调用方处理。"""
    document = Document()
    paragraph = document.add_paragraph("旧文本")
    calls: list[tuple[str, str]] = []

    def set_run_fonts(_run, *, cn_font: str, en_font: str) -> None:
        calls.append((cn_font, en_font))

    writer = segment_writer(paragraph, set_run_fonts=set_run_fonts)
    writer("新文本", cn_font="黑体")

    assert calls == [("黑体", "Times New Roman")]


def test_write_inline_tokens_restores_text_tabs_breaks_and_page_breaks() -> None:
    """行内 token 写回应保留文本、制表符、软换行和分页符顺序。"""
    document = Document()
    paragraph = document.add_paragraph()

    write_inline_tokens(
        paragraph,
        [
            _Token("text", "甲"),
            _Token("tab"),
            _Token("text", "乙"),
            _Token("line_break"),
            _Token("text", "丙"),
            _Token("page_break"),
        ],
    )

    assert paragraph.text == "甲\t乙\n丙"
    assert len(paragraph.runs) == 6


def test_without_redundant_trailing_body_page_breaks_keeps_protected_boundaries() -> None:
    """分页清理遇到附件、表格、图片等边界时应原样返回。"""
    tokens = [_Token("text", "正文"), _Token("page_break")]

    result = without_redundant_trailing_body_page_breaks(
        _ParagraphData("正文", "body"),
        _ParagraphData("附件标题", "attachment_title"),
        tokens,
    )

    assert result == tokens


def test_without_redundant_trailing_body_page_breaks_removes_mid_sentence_and_trailing_breaks() -> None:
    """普通正文中的句中分页和末尾分页应被清理。"""
    tokens = [
        _Token("text", "前半句，"),
        _Token("page_break"),
        _Token("text", "后半句"),
        _Token("page_break"),
    ]

    result = without_redundant_trailing_body_page_breaks(
        _ParagraphData("前半句，后半句", "body"),
        _ParagraphData("下一段", "body"),
        tokens,
    )

    assert result == [_Token("text", "前半句，"), _Token("text", "后半句")]
