from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

from docxtool.document.engine import export_doc
from docxtool.document.importer import DocxImporter, DocumentData, ParagraphData, ParagraphFeatures
from docxtool.document.style_config import PageSettings, StyleRule
from docxtool.security.docx_integrity import validate_docx_integrity

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _rules() -> list[StyleRule]:
    return [StyleRule.default_for_row(index) for index in range(24)]


def _document_xml(path: Path) -> ET.Element:
    with zipfile.ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/document.xml"))


def _styles_xml(path: Path) -> ET.Element:
    with zipfile.ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/styles.xml"))


def _footer_xmls(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        return [
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        ]


def _paragraphs(root: ET.Element) -> list[ET.Element]:
    return root.findall(".//w:body/w:p", NS)


def _pstyle(paragraph: ET.Element) -> str:
    p_style = paragraph.find("w:pPr/w:pStyle", NS)
    return p_style.get(qn("w:val")) if p_style is not None else ""


def _cm_twips(value: str | None) -> float:
    return int(value or 0) / 567


def _text(paragraph: ET.Element) -> str:
    return "".join(element.text or "" for element in paragraph.findall(".//w:t", NS))


def _run_text_and_bold(paragraph: ET.Element) -> list[tuple[str, bool]]:
    result = []
    for run in paragraph.findall("w:r", NS):
        text = "".join(element.text or "" for element in run.findall("w:t", NS))
        bold = run.find("w:rPr/w:b", NS)
        bold_value = bold.get(qn("w:val")) if bold is not None else None
        result.append((text, bold is not None and bold_value not in {"0", "false", "False"}))
    return result


def _run_east_asia_font(paragraph: ET.Element, text: str) -> str | None:
    for run in paragraph.findall("w:r", NS):
        run_text = "".join(element.text or "" for element in run.findall("w:t", NS))
        if run_text == text:
            fonts = run.find("w:rPr/w:rFonts", NS)
            return fonts.get(qn("w:eastAsia")) if fonts is not None else None
    return None


def _spacing_after_lines(paragraph: ET.Element) -> str | None:
    spacing = paragraph.find("w:pPr/w:spacing", NS)
    return spacing.get(qn("w:afterLines")) if spacing is not None else None


def _style(root: ET.Element, style_id: str) -> ET.Element:
    for style in root.findall("w:style", NS):
        if style.get(qn("w:styleId")) == style_id:
            return style
    raise AssertionError(f"missing style {style_id}")


def _export_sample(output: Path) -> None:
    data = DocumentData(
        paragraphs=[
            ParagraphData("主标题", "title", "主标题", ParagraphFeatures()),
            ParagraphData("一级标题", "heading1", "一、一级标题", ParagraphFeatures()),
            ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
            ParagraphData("责任单位：区政府\n责任单位：商务局", "responsibility_line", "责 任 单 位：区政府\n责任单位：商务局", ParagraphFeatures(), meta={"colon_bold": True}),
            ParagraphData("附件：测试材料", "attachment_note", "附件：测试材料", ParagraphFeatures()),
            ParagraphData("附件 1", "attachment_page_mark", "附件 1", ParagraphFeatures()),
            ParagraphData("附件标题", "attachment_title", "附件标题", ParagraphFeatures()),
            ParagraphData("附件正文", "attachment_body", "附件正文", ParagraphFeatures()),
        ],
        filepath="generated.docx",
    )
    export_doc(data, _rules(), PageSettings(), str(output))


def test_exported_nonempty_paragraphs_have_structural_pstyle_and_no_numpr(tmp_path: Path) -> None:
    output = tmp_path / "structured.docx"
    _export_sample(output)

    assert validate_docx_integrity(output).ok is True
    root = _document_xml(output)
    paragraphs = [paragraph for paragraph in _paragraphs(root) if _text(paragraph).strip()]

    assert paragraphs
    assert all(_pstyle(paragraph).startswith("DCT-") for paragraph in paragraphs)
    assert root.findall(".//w:numPr", NS) == []
    assert _pstyle(paragraphs[1]) == "DCT-Heading1"
    assert _pstyle(paragraphs[2]) == "DCT-Body"


def test_heading_body_responsibility_and_attachment_styles_are_applied(tmp_path: Path) -> None:
    output = tmp_path / "structured.docx"
    _export_sample(output)
    doc_root = _document_xml(output)
    style_root = _styles_xml(output)

    heading = next(paragraph for paragraph in _paragraphs(doc_root) if _text(paragraph) == "一级标题")
    assert heading.find("w:pPr/w:keepNext", NS) is None
    assert heading.find("w:pPr/w:keepLines", NS) is None

    body_style = _style(style_root, "DCT-Body")
    body_indent = body_style.find("w:pPr/w:ind", NS)
    assert body_indent.get(qn("w:firstLineChars")) == "200"

    responsibility = next(paragraph for paragraph in _paragraphs(doc_root) if "责任单位" in _text(paragraph))
    assert _pstyle(responsibility) == "DCT-Responsibility"
    assert _text(responsibility).startswith("责任单位：区政府")
    assert responsibility.find("w:pPr/w:jc", NS).get(qn("w:val")) == "left"
    assert responsibility.find("w:pPr/w:jc", NS).get(qn("w:val")) not in {"both", "distribute"}
    responsibility_indent = responsibility.find("w:pPr/w:ind", NS)
    assert responsibility_indent.get(qn("w:leftChars")) == "200"
    assert responsibility_indent.get(qn("w:firstLineChars")) == "0"
    assert responsibility.findall(".//w:br", NS)

    style_by_text = {_text(paragraph): _pstyle(paragraph) for paragraph in _paragraphs(doc_root)}
    assert style_by_text["附件 1"] == "DCT-AttachmentMark"
    assert style_by_text["附件标题"] == "DCT-AttachmentTitle"
    assert style_by_text["附件正文"] == "DCT-AttachmentBody"
    attachment_title = next(
        paragraph for paragraph in _paragraphs(doc_root) if _text(paragraph) == "附件标题"
    )
    title_spacing = attachment_title.find("w:pPr/w:spacing", NS)
    assert title_spacing.get(qn("w:beforeLines")) == "100"
    assert title_spacing.get(qn("w:afterLines")) == "100"
    attachment_title_style = _style(style_root, "DCT-AttachmentTitle")
    style_spacing = attachment_title_style.find("w:pPr/w:spacing", NS)
    assert style_spacing.get(qn("w:beforeLines")) == "100"
    assert style_spacing.get(qn("w:afterLines")) == "100"
    paragraph_texts = [_text(paragraph) for paragraph in _paragraphs(doc_root)]
    title_index = paragraph_texts.index("附件标题")
    assert paragraph_texts[title_index + 1] == "附件正文"
    for text in ("附件 1", "附件标题"):
        paragraph = next(paragraph for paragraph in _paragraphs(doc_root) if _text(paragraph) == text)
        assert paragraph.find("w:pPr/w:keepNext", NS) is not None
        assert paragraph.find("w:pPr/w:keepLines", NS) is not None


def test_structural_styles_have_heading_outline_levels(tmp_path: Path) -> None:
    output = tmp_path / "structured.docx"
    _export_sample(output)
    style_root = _styles_xml(output)

    for index, style_id in enumerate(("DCT-Heading1", "DCT-Heading2", "DCT-Heading3", "DCT-Heading4")):
        outline = _style(style_root, style_id).find("w:pPr/w:outlineLvl", NS)
        assert outline.get(qn("w:val")) == str(index)


def test_page_margins_and_clean_footer_page_field(tmp_path: Path) -> None:
    output = tmp_path / "structured.docx"
    _export_sample(output)

    sect = _document_xml(output).find(".//w:body/w:sectPr", NS)
    margins = sect.find("w:pgMar", NS)
    assert abs(_cm_twips(margins.get(qn("w:top"))) - 3.7) < 0.01
    assert abs(_cm_twips(margins.get(qn("w:bottom"))) - 3.5) < 0.01
    assert abs(_cm_twips(margins.get(qn("w:left"))) - 2.8) < 0.01
    assert abs(_cm_twips(margins.get(qn("w:right"))) - 2.6) < 0.01

    footer_xmls = _footer_xmls(output)
    assert len(footer_xmls) == 3
    for footer_xml in footer_xmls:
        assert footer_xml.count("PAGE") == 1
        assert "NUMPAGES" not in footer_xml
        assert "AlternateContent" not in footer_xml
        assert "txbxContent" not in footer_xml
        assert "textbox" not in footer_xml


def test_document_grid_char_space_uses_ooxml_4096_point_units(tmp_path: Path) -> None:
    output = tmp_path / "grid.docx"
    _export_sample(output)
    grid = _document_xml(output).find(".//w:body/w:sectPr/w:docGrid", NS)

    assert grid is not None
    assert grid.get(qn("w:charsPerLine")) == "28"
    assert int(grid.get(qn("w:charSpace"))) == -842


def test_import_export_keeps_input_sha256_unchanged(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("一、一级标题")
    document.add_paragraph("正文内容")
    document.add_paragraph("责 任 单 位：区政府")
    document.save(source)
    before = hashlib.sha256(source.read_bytes()).hexdigest()

    data = DocxImporter().load(str(source), _rules())
    export_doc(data, _rules(), PageSettings(), str(output))

    assert hashlib.sha256(source.read_bytes()).hexdigest() == before
    assert validate_docx_integrity(output).ok is True
    grid = _document_xml(output).find(".//w:body/w:sectPr/w:docGrid", NS)
    assert grid is not None
    assert grid.get(qn("w:charsPerLine")) == "28"
    assert int(grid.get(qn("w:charSpace"))) == -842


def test_responsibility_line_normalizes_quotes_and_repeated_labels(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("总题目")
    document.add_paragraph("一、一级标题")
    document.add_paragraph("正文内容正文内容正文内容。")
    document.add_paragraph("“责任单位：区政府责任单位：商务局”")
    document.save(source)

    data = DocxImporter().load(str(source), _rules())
    responsibility = next(paragraph for paragraph in data.paragraphs if paragraph.type_id == "responsibility_line")
    assert responsibility.text == "责任单位：区政府\n责任单位：商务局"

    export_doc(data, _rules(), PageSettings(), str(output))
    root = _document_xml(output)
    output_responsibility = next(paragraph for paragraph in _paragraphs(root) if "责任单位" in _text(paragraph))
    assert _pstyle(output_responsibility) == "DCT-Responsibility"
    assert output_responsibility.findall(".//w:br", NS)


def test_heading4_does_not_insert_blank_paragraph_after_it(tmp_path: Path) -> None:
    output = tmp_path / "output.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData("一级标题", "heading1", "一、一级标题", ParagraphFeatures(), meta={"numbering": "一、"}),
            ParagraphData("二级标题", "heading2", "（一）二级标题", ParagraphFeatures(), meta={"numbering": "（一）"}),
            ParagraphData("三级标题", "heading3", "1.三级标题", ParagraphFeatures(), meta={"numbering": "1."}),
            ParagraphData("四级标题", "heading4", "（1）四级标题", ParagraphFeatures(), meta={"numbering": "（1）"}),
            ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
        ],
        filepath="generated.docx",
    )

    export_doc(data, _rules(), PageSettings(), str(output))
    texts = [_text(paragraph) for paragraph in _paragraphs(_document_xml(output))]

    assert "（1）四级标题" in texts
    heading_index = texts.index("（1）四级标题")
    assert texts[heading_index + 1] == "正文内容"


def test_heading2_title_runs_stay_bold_when_inline_body_is_split(tmp_path: Path) -> None:
    output = tmp_path / "output.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData(
                "强化理论武装，筑牢思想政治根基。坚持学习提升履职能力和服务水平。",
                "heading2",
                "（一）强化理论武装，筑牢思想政治根基。坚持学习提升履职能力和服务水平。",
                ParagraphFeatures(),
                meta={"numbering": "（一）"},
            ),
        ],
        filepath="generated.docx",
    )

    export_doc(data, _rules(), PageSettings(), str(output))
    paragraph = _paragraphs(_document_xml(output))[0]

    assert _pstyle(paragraph) == "DCT-Heading2"
    assert _run_text_and_bold(paragraph) == [
        ("（一）", True),
        ("强化理论武装，筑牢思想政治根基。", True),
        ("坚持学习提升履职能力和服务水平。", False),
    ]


def test_short_inline_heading2_body_uses_the_body_rule(tmp_path: Path) -> None:
    output = tmp_path / "output.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData(
                "标题。短正文测试",
                "heading2",
                "（一）标题。短正文测试",
                ParagraphFeatures(),
                meta={"numbering": "（一）"},
            ),
        ],
        filepath="generated.docx",
    )

    export_doc(data, _rules(), PageSettings(), str(output))
    paragraph = _paragraphs(_document_xml(output))[0]

    assert _run_text_and_bold(paragraph) == [
        ("（一）", True),
        ("标题。", True),
        ("短正文测试", False),
    ]
    assert _run_east_asia_font(paragraph, "短正文测试") == "仿宋_GB2312"


def test_heading_rules_respect_configured_boldness(tmp_path: Path) -> None:
    output = tmp_path / "output.docx"
    rules = _rules()
    rules[2].bold = False
    rules[3].bold = False
    data = DocumentData(
        paragraphs=[
            ParagraphData("二级标题", "heading2", "（一）二级标题", ParagraphFeatures(), meta={"numbering": "（一）"}),
            ParagraphData("三级标题", "heading3", "1.三级标题", ParagraphFeatures(), meta={"numbering": "1."}),
        ],
        filepath="generated.docx",
    )

    export_doc(data, rules, PageSettings(), str(output))
    paragraphs = _paragraphs(_document_xml(output))

    assert all(not bold for _text_value, bold in _run_text_and_bold(paragraphs[0]))
    assert all(not bold for _text_value, bold in _run_text_and_bold(paragraphs[1]))


def test_heading3_followed_by_body_has_no_blank_gap_or_spacing(tmp_path: Path) -> None:
    output = tmp_path / "output.docx"
    rules = _rules()
    rules[3].spacing_after = 1.0
    data = DocumentData(
        paragraphs=[
            ParagraphData("测试测试", "heading3", "1.测试测试", ParagraphFeatures(), meta={"numbering": "1."}),
            ParagraphData("正文内容正文内容正文内容。", "body", "正文内容正文内容正文内容。", ParagraphFeatures()),
        ],
        filepath="generated.docx",
    )

    export_doc(data, rules, PageSettings(), str(output))
    paragraphs = _paragraphs(_document_xml(output))
    texts = [_text(paragraph) for paragraph in paragraphs]

    assert texts[:2] == ["1.测试测试", "正文内容正文内容正文内容。"]
    assert _spacing_after_lines(paragraphs[0]) in {None, "0"}


def test_heading1_before_nested_heading_does_not_create_keep_chain(tmp_path: Path) -> None:
    output = tmp_path / "output.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData("一级标题", "heading1", "一、一级标题", ParagraphFeatures(), meta={"numbering": "一、"}),
            ParagraphData("二级标题", "heading2", "（一）二级标题", ParagraphFeatures(), meta={"numbering": "（一）"}),
            ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
        ],
        filepath="generated.docx",
    )

    export_doc(data, _rules(), PageSettings(), str(output))
    heading1 = _paragraphs(_document_xml(output))[0]

    assert heading1.find("w:pPr/w:keepNext", NS) is None
    assert heading1.find("w:pPr/w:keepLines", NS) is None


def test_heading1_before_body_does_not_require_same_page(tmp_path: Path) -> None:
    output = tmp_path / "output.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData("一级标题", "heading1", "一、一级标题", ParagraphFeatures(), meta={"numbering": "一、"}),
            ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
        ],
        filepath="generated.docx",
    )

    export_doc(data, _rules(), PageSettings(), str(output))
    heading1 = _paragraphs(_document_xml(output))[0]

    assert heading1.find("w:pPr/w:keepNext", NS) is None
    assert heading1.find("w:pPr/w:keepLines", NS) is None


def test_imported_heading3_heading4_and_responsibility_are_exported_without_blank_gap(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    for text in (
        "总题目",
        "一、一级标题",
        "（一）二级标题",
        "1.测试",
        "（1）测试",
        "正文内容正文内容正文内容。",
        "责任单位：区政府责任单位：商务局",
    ):
        document.add_paragraph(text)
    document.save(source)

    data = DocxImporter().load(str(source), _rules())
    type_by_original = {paragraph.original_text: paragraph.type_id for paragraph in data.paragraphs}

    assert type_by_original["1.测试"] == "heading3"
    assert type_by_original["（1）测试"] == "heading4"
    responsibility = next(paragraph for paragraph in data.paragraphs if paragraph.type_id == "responsibility_line")
    assert responsibility.text == "责任单位：区政府\n责任单位：商务局"

    export_doc(data, _rules(), PageSettings(), str(output))
    texts = [_text(paragraph) for paragraph in _paragraphs(_document_xml(output))]

    heading4_index = texts.index("（1）测试")
    assert texts[heading4_index + 1] == "正文内容正文内容正文内容。"
