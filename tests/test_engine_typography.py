from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docxtool.document.engine.typography import (
    apply_digit_latin_font,
    apply_superscript_split,
    apply_universal_superscript,
    set_run_fonts,
)


def _run_fonts(run):
    """读取 run 字体 XML，返回 rFonts 元素用于断言。"""
    return run._element.get_or_add_rPr().find(qn("w:rFonts"))


def test_set_run_fonts_writes_east_asia_and_latin_fonts() -> None:
    """字体辅助应分别写入中文和英文字体。"""
    document = Document()
    run = document.add_paragraph().add_run("测试ABC")

    set_run_fonts(run, cn_font="黑体", en_font="Times New Roman")

    fonts = _run_fonts(run)
    assert fonts.get(qn("w:eastAsia")) == "黑体"
    assert fonts.get(qn("w:ascii")) == "Times New Roman"
    assert fonts.get(qn("w:hAnsi")) == "Times New Roman"


def test_apply_superscript_split_converts_bracket_markers() -> None:
    """上标拆分应把 [n]/〔n〕 拆成上标 run，并统一为 [n]。"""
    document = Document()
    paragraph = document.add_paragraph("正文[1]继续〔2〕")

    apply_superscript_split(paragraph)

    assert paragraph.text == "正文[1]继续[2]"
    superscripts = [run.text for run in paragraph.runs if run.font.superscript]
    assert superscripts == ["[1]", "[2]"]


def test_apply_superscript_split_preserves_direct_chinese_font() -> None:
    """拆分脚注标记时不得丢失正文 run 的中文字体。"""
    document = Document()
    paragraph = document.add_paragraph("正文[1]继续")
    set_run_fonts(paragraph.runs[0], cn_font="仿宋_GB2312")

    apply_superscript_split(paragraph)

    fonts = _run_fonts(paragraph.runs[0])
    assert fonts.get(qn("w:eastAsia")) == "仿宋_GB2312"
    assert paragraph.text == "正文[1]继续"


def test_apply_digit_latin_font_splits_digits_and_letters_only() -> None:
    """数字和拉丁字母应拆成独立 run 并使用 Times New Roman。"""
    document = Document()
    paragraph = document.add_paragraph("第2026年ABC测试")

    apply_digit_latin_font(paragraph)

    texts = [run.text for run in paragraph.runs]
    assert texts == ["第", "2026", "年", "ABC", "测试"]
    latin_runs = [run for run in paragraph.runs if run.text in {"2026", "ABC"}]
    assert all(_run_fonts(run).get(qn("w:ascii")) == "Times New Roman" for run in latin_runs)


def test_apply_digit_latin_font_skips_paragraph_with_line_break() -> None:
    """含软换行段落应跳过数字字体拆分，避免破坏换行结构。"""
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("第2026年")
    run._element.append(OxmlElement("w:br"))

    apply_digit_latin_font(paragraph)

    assert len(paragraph.runs) == 1
    assert paragraph.runs[0].text == "第2026年\n"


def test_apply_universal_superscript_formats_existing_superscript_run() -> None:
    """已有纯数字上标应格式化为 [n]，并设置三号 Times New Roman。"""
    document = Document()
    run = document.add_paragraph().add_run("3")
    run.font.superscript = True

    apply_universal_superscript(document.paragraphs[0])

    assert run.text == "[3]"
    assert run.font.name == "Times New Roman"
    assert run.font.size == Pt(16)
