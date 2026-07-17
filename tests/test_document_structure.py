from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn

from docxtool.document.engine.document_structure import (
    BOUNDARY_RULES,
    ElementKind,
    analyze_document_structure,
)
from docxtool.document.importer import (
    DocumentData,
    InlineToken,
    ParagraphData,
    ParagraphFeatures,
)


def paragraph(text, type_id, *, tokens=(), meta=None):
    return ParagraphData(
        text,
        type_id,
        text,
        ParagraphFeatures(text=text),
        meta or {},
        list(tokens),
    )


def letterhead(style_id):
    document = Document()
    style = document.styles.add_style(style_id, 1)
    style.style_id = style_id
    value = document.add_paragraph("版头", style=style)
    return paragraph("", "__letterhead__", meta={"paragraph_xml": value})


def complete_data():
    page = [InlineToken("page_break")]
    return DocumentData(paragraphs=[
        letterhead("DCT-LetterheadMark"),
        letterhead("DCT-DocumentNumber"),
        letterhead("DCT-SignerLine"),
        letterhead("DCT-LetterheadSeparator"),
        paragraph("关于测试工作的通知", "title"),
        paragraph("测试机关", "author_line"),
        paragraph("2026年7月17日", "date_line"),
        paragraph("内江会议中心", "location_line"),
        paragraph("一、一级标题", "heading1"),
        paragraph("（一）二级标题", "heading2"),
        paragraph("1.三级标题", "heading3"),
        paragraph("（1）四级标题", "heading4"),
        paragraph("一是列表内容", "list"),
        paragraph("引用内容", "quote"),
        paragraph("正文中的2025年10月15日不是落款。", "body"),
        paragraph("", "__table__"),
        paragraph("", "__image__"),
        paragraph("图1 测试图片", "__object_caption__"),
        paragraph("附件：1. 测试清单", "attachment_note"),
        paragraph("2. 测试表格", "attachment_note_item"),
        paragraph("测试市人民政府", "sign_org"),
        paragraph("2026年7月17日", "sign_date"),
        paragraph("（此件公开发布）", "note"),
        paragraph("附件1", "attachment_page_mark", tokens=page),
        paragraph("附件一标题", "attachment_title"),
        paragraph("附件一正文", "attachment_body"),
        paragraph("附件2", "attachment_page_mark", tokens=page),
        paragraph("附件二正文", "attachment_body"),
    ])


def test_complete_standard_structure_and_boundaries():
    structure = analyze_document_structure(complete_data())

    assert structure.front_matter is not None
    assert [item.kind for item in structure.front_matter.elements] == [
        ElementKind.LETTERHEAD_MARK,
        ElementKind.DOCUMENT_NUMBER,
        ElementKind.SIGNER,
        ElementKind.LETTERHEAD_SEPARATOR,
    ]
    assert structure.title is not None
    assert len(structure.title.title_elements) == 1
    assert len(structure.title.metadata_elements) == 3
    assert structure.body is not None
    assert {item.kind for item in structure.body.elements} >= {
        ElementKind.HEADING_1, ElementKind.HEADING_2, ElementKind.HEADING_3,
        ElementKind.HEADING_4, ElementKind.LIST, ElementKind.QUOTE,
        ElementKind.TABLE, ElementKind.FIGURE, ElementKind.CAPTION,
    }
    assert structure.attachment_note is not None
    assert structure.signature is not None
    assert [item.kind for item in structure.signature.elements] == [
        ElementKind.SIGNATURE_AGENCY, ElementKind.SIGNATURE_DATE, ElementKind.NOTE,
    ]
    assert [item.ordinal for item in structure.attachments] == [1, 2]
    assert all(item.starts_after_page_break for item in structure.attachments)
    assert all(item.span.elements[0].kind == ElementKind.PAGE_BREAK for item in structure.attachments)
    assert structure.boundary_rules == BOUNDARY_RULES
    assert all(span.confidence >= 0.85 for span in (
        structure.front_matter, structure.title.span, structure.attachment_note, structure.signature,
    ))
    assert all(span.evidence for span in (
        structure.front_matter, structure.title.span, structure.body.span,
        structure.attachment_note, structure.signature,
    ))


def test_no_letterhead_and_title_date_is_not_signature():
    data = DocumentData(paragraphs=[
        paragraph("测试标题", "title"),
        paragraph("测试作者", "author_line"),
        paragraph("2026年7月17日", "date_line"),
        paragraph("正文", "body"),
    ])
    structure = analyze_document_structure(data)

    assert structure.front_matter is None
    assert structure.title is not None
    assert [item.kind for item in structure.title.metadata_elements] == [
        ElementKind.TITLE_METADATA, ElementKind.TITLE_METADATA,
    ]
    assert structure.signature is None


def test_body_date_does_not_form_signature_and_unknown_is_preserved():
    data = DocumentData(paragraphs=[
        paragraph("标题", "title"),
        paragraph("正文", "body"),
        paragraph("2026年7月17日", "body"),
        paragraph("无法识别", "mystery"),
    ])
    structure = analyze_document_structure(data)

    assert structure.signature is None
    assert structure.body is not None
    unknown = [item for item in structure.body.elements if item.kind == ElementKind.UNKNOWN]
    assert len(unknown) == 1
    assert unknown[0].confidence < 0.6


def test_attachment_note_and_content_are_separate():
    structure = analyze_document_structure(complete_data())

    note_indexes = {item.index for item in structure.attachment_note.elements}
    content_indexes = {item.index for block in structure.attachments for item in block.span.elements}
    assert note_indexes.isdisjoint(content_indexes)
    assert all(item.kind == ElementKind.ATTACHMENT_NOTE_ITEM for item in structure.attachment_note.elements)


def test_attachment_without_real_page_break_is_not_attachment_block():
    data = DocumentData(paragraphs=[
        paragraph("标题", "title"),
        paragraph("正文", "body"),
        paragraph("测试机关", "sign_org"),
        paragraph("2026年7月17日", "sign_date"),
        paragraph("附件1", "attachment_page_mark"),
        paragraph("附件正文", "attachment_body"),
    ])
    structure = analyze_document_structure(data)

    assert structure.attachments == ()
    assert any(
        item.kind == ElementKind.ATTACHMENT_TITLE
        for span in structure.unknown for item in span.elements
    )


def test_page_break_before_and_section_break_are_detected():
    data = DocumentData(paragraphs=[
        paragraph("标题", "title"),
        paragraph("正文", "body"),
        paragraph("附件1", "attachment_page_mark", meta={"page_break_before": True}),
        paragraph("附件正文", "attachment_body", meta={"sectPr": object()}),
    ])
    structure = analyze_document_structure(data)

    assert len(structure.attachments) == 1
    assert [item.evidence[0] for item in structure.elements if item.kind == ElementKind.PAGE_BREAK] == [
        "ooxml:pageBreakBefore", "ooxml:sectPr",
    ]


def test_empty_spacing_does_not_change_structure():
    base = complete_data()
    with_unknown_blanks = deepcopy(base)
    # Importer normally removes empty paragraphs; retained unknowns must not act as boundaries.
    with_unknown_blanks.paragraphs.insert(8, paragraph("", "blank"))
    structure = analyze_document_structure(with_unknown_blanks)

    assert structure.title is not None
    assert structure.body is not None
    assert len(structure.attachments) == 2
    assert any(item.kind == ElementKind.UNKNOWN for item in structure.body.elements)


def test_signature_requires_agency_date_pair_and_keeps_note():
    data = DocumentData(paragraphs=[
        paragraph("标题", "title"), paragraph("正文", "body"),
        paragraph("测试机关", "sign_org"), paragraph("附注", "note"),
        paragraph("2026年7月17日", "sign_date"),
    ])
    structure = analyze_document_structure(data)
    assert structure.signature is not None
    assert structure.signature.start_index == 2
    assert structure.signature.end_index == 5


def test_analysis_is_deterministic_and_does_not_change_xml():
    data = complete_data()
    xml_nodes = [
        item.meta["paragraph_xml"]._p
        for item in data.paragraphs if "paragraph_xml" in item.meta
    ]
    before = [node.xml for node in xml_nodes]

    first = analyze_document_structure(data)
    second = analyze_document_structure(data)

    assert first == second
    assert before == [node.xml for node in xml_nodes]
    accounted = sorted(
        item.index
        for span in [
            first.front_matter, first.title.span, first.body.span,
            first.attachment_note, first.signature, *[block.span for block in first.attachments],
            *first.unknown,
        ] if span is not None
        for item in span.elements
    )
    assert accounted == list(range(len(first.elements)))
    assert len(accounted) == len(set(accounted))


def test_real_page_break_xml_is_detected_without_text_spacing():
    document = Document()
    value = document.add_paragraph("附件1")
    value.paragraph_format.page_break_before = True
    data = DocumentData(paragraphs=[
        paragraph("标题", "title"), paragraph("正文", "body"),
        paragraph("", "attachment_page_mark", meta={"paragraph_xml": value}),
        paragraph("附件正文", "attachment_body"),
    ])
    structure = analyze_document_structure(data)
    assert len(structure.attachments) == 1
    assert any(node.tag == qn("w:pageBreakBefore") for node in value._p.iter())
