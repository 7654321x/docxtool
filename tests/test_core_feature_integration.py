from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document

from docxtool.document.importer import DocumentData, DocxImporter, ParagraphData, ParagraphFeatures
from docxtool.document.engine.core import export_doc
from docxtool.document.style_config import (
    ConfigValidationError,
    PageSettings,
    StyleRule,
    load_rules_and_settings,
    validate_format_config,
)
from docxtool.security.docx_integrity import validate_docx_integrity


def _rules() -> list[StyleRule]:
    return [StyleRule.default_for_row(index) for index in range(24)]


def _document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def _footer_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )


def test_core_feature_options_default_to_safe_off_without_changing_legacy_features() -> None:
    _, _, features = load_rules_and_settings({})

    assert features["punctuation_enabled"] is True
    assert features["page_number_enabled"] is True
    assert features["punctuation"] == {
        "enabled": False,
        "mode": "safe",
        "scope": {"body": True, "tables": False, "headers": False, "footers": False},
    }
    assert features["classification"]["enabled"] is True
    assert features["numbering"]["enabled"] is False
    assert features["page_number"]["enabled"] is False
    assert features["page_number"]["position"] == "outside"
    assert features["page_number"]["first_page"] is True
    assert features["page_number"]["offset_from_text_mm"] == 7
    assert features["table_format"]["enabled"] is False
    assert features["cleanup"]["enabled"] is False


def test_core_feature_options_validate_modes() -> None:
    with pytest.raises(ConfigValidationError, match="punctuation.mode"):
        validate_format_config({"punctuation": {"mode": "unsafe"}})


def test_core_feature_options_accept_legacy_boolean_flags() -> None:
    _, _, features = load_rules_and_settings(
        {
            "styles": [],
            "page": {},
            "punctuation": True,
            "numbering": "false",
            "page_number": True,
            "table_format": 0,
            "cleanup": "on",
        }
    )

    assert features["punctuation"]["enabled"] is True
    assert features["numbering"]["enabled"] is False
    assert features["page_number"]["enabled"] is True
    assert features["table_format"]["enabled"] is False
    assert features["cleanup"]["enabled"] is True


def test_importer_uses_safe_punctuation_engine_and_records_classification(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("请访问 https://example.com/a,b?x=1.2, 并说明:可以吗?")
    document.save(source)

    data = DocxImporter().load(
        str(source),
        _rules(),
        features={
            "punctuation_enabled": False,
            "punctuation": {"enabled": True, "mode": "safe"},
            "classification": {"enabled": True, "minimum_auto_format_confidence": 0.85},
        },
    )

    paragraph = data.paragraphs[0]
    assert "https://example.com/a,b?x=1.2" in paragraph.original_text
    assert "说明：可以吗？" in paragraph.original_text
    assert paragraph.meta["classification_kind"]
    assert 0 <= paragraph.meta["classification_confidence"] <= 1


def test_export_doc_applies_explicit_numbering_and_field_page_number_options(tmp_path: Path) -> None:
    output = tmp_path / "formatted.docx"
    data = DocumentData(
        paragraphs=[
            ParagraphData("1、重点任务", "body", "1、重点任务", ParagraphFeatures()),
            ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
        ],
        filepath=str(tmp_path / "input.docx"),
    )

    export_doc(
        data,
        _rules(),
        PageSettings(),
        str(output),
        page_number_enabled=False,
        numbering_options={"enabled": True, "mode": "safe"},
        page_number_options={"enabled": True, "style": "cn_total", "position": "center", "offset_from_text_mm": 7},
    )

    assert validate_docx_integrity(output).ok is True
    assert Document(output).paragraphs[0].text == "1.重点任务"
    footer_xml = _footer_xml(output)
    assert "PAGE" in footer_xml
    assert "NUMPAGES" in footer_xml
    assert 'w:val="center"' in footer_xml
    assert abs(Document(output).sections[0].footer_distance.cm - 2.8) < 0.02
