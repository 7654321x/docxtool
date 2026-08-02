from __future__ import annotations

from pathlib import Path

from docx import Document

import docxtool.document.importer as importer_module
from docxtool.document.models import DocumentData, ParagraphData, ParagraphFeatures
from docxtool.document.recognition.legacy import DetectionContext


def test_importer_load_is_a_thin_document_pipeline_facade(monkeypatch) -> None:
    expected = object()
    captured = {}

    def fake_pipeline(filepath, rules, features, **kwargs):
        captured.update(
            filepath=filepath,
            rules=rules,
            features=features,
            kwargs=kwargs,
        )
        return expected

    monkeypatch.setattr(importer_module, "run_document_pipeline", fake_pipeline)
    importer = importer_module.DocxImporter()

    result = importer.load(
        "sample.docx",
        [],
        {"processing": {"strategy": "structural"}},
        strict_preservation=False,
        recognition_mode="shadow",
    )

    assert result is expected
    assert captured["filepath"] == "sample.docx"
    assert captured["rules"] == []
    assert captured["features"] == {"processing": {"strategy": "structural"}}
    assert captured["kwargs"]["importer"] is importer
    assert captured["kwargs"]["compatibility"] is importer_module
    assert captured["kwargs"]["strict_preservation"] is False
    assert captured["kwargs"]["recognition_mode"] == "shadow"


def test_importer_facade_reexports_stable_models_and_legacy_context() -> None:
    assert importer_module.DocumentData is DocumentData
    assert importer_module.ParagraphData is ParagraphData
    assert importer_module.ParagraphFeatures is ParagraphFeatures
    assert importer_module.DetectionContext is DetectionContext


def test_importer_pipeline_uses_legacy_reader_and_segmentation_patch_points(
    tmp_path: Path,
    monkeypatch,
) -> None:
    source = tmp_path / "facade.docx"
    document = Document()
    document.add_paragraph("一、测试标题")
    document.add_paragraph("测试正文。")
    document.save(source)

    real_read = importer_module._read_body_blocks
    real_build = importer_module._build_logical_lines
    calls = []

    def capture_read(*args, **kwargs):
        calls.append("read")
        return real_read(*args, **kwargs)

    def capture_build(*args, **kwargs):
        calls.append("build")
        return real_build(*args, **kwargs)

    monkeypatch.setattr(importer_module, "_read_body_blocks", capture_read)
    monkeypatch.setattr(importer_module, "_build_logical_lines", capture_build)

    data = importer_module.DocxImporter().load(str(source), [], strict_preservation=True)

    assert calls == ["read", "build"]
    assert [paragraph.original_text for paragraph in data.paragraphs] == [
        "一、测试标题",
        "测试正文。",
    ]
