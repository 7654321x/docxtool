import base64
import posixpath
import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from docxtool.document.engine import export_doc
from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import PageSettings, StyleRule

REL_NS = "{http://schemas.openxmlformats.org/package/2006/relationships}"
IMAGE_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def _rules():
    return [StyleRule.default_for_row(i) for i in range(10)]


def _tiny_png(path: Path) -> None:
    path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQ"
            "VR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )


def _set_part_text(part, text: str) -> None:
    part.paragraphs[0].text = text


def _add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f" {instruction} "
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    paragraph.add_run("1")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def _set_footer_with_page_fields(footer, label: str) -> None:
    paragraph = footer.paragraphs[0]
    paragraph.text = f"{label} footer "
    _add_field(paragraph, "PAGE")
    paragraph.add_run("/")
    _add_field(paragraph, "NUMPAGES")


def _configure_section(section, label: str, image_path: Path | None = None) -> None:
    section.different_first_page_header_footer = True
    for header in (section.header, section.first_page_header, section.even_page_header):
        header.is_linked_to_previous = False
    for footer in (section.footer, section.first_page_footer, section.even_page_footer):
        footer.is_linked_to_previous = False

    _set_part_text(section.header, f"{label} default header")
    if image_path is not None:
        section.header.paragraphs[0].add_run().add_picture(str(image_path), width=Cm(0.2))
    _set_part_text(section.first_page_header, f"{label} first header")
    _set_part_text(section.even_page_header, f"{label} even header")
    _set_footer_with_page_fields(section.footer, f"{label} default")
    _set_footer_with_page_fields(section.first_page_footer, f"{label} first")
    _set_footer_with_page_fields(section.even_page_footer, f"{label} even")


def _export_round_trip(source: Path, output: Path) -> None:
    data = DocxImporter().load(str(source), _rules())
    export_doc(data, _rules(), PageSettings(), str(output))
    Document(output)


def _part_name(target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    return posixpath.normpath(target if target.startswith("word/") else "word/" + target)


def _child_attr(element, child_tag: str, attr_name: str, default: str) -> str:
    child = element.find(child_tag)
    if child is None:
        return default
    return child.get(attr_name, default)


def _rels(zf: zipfile.ZipFile, rels_name: str) -> dict[str, dict[str, str]]:
    root = ET.fromstring(zf.read(rels_name))
    return {
        rel.get("Id"): {
            "target": rel.get("Target"),
            "type": rel.get("Type"),
        }
        for rel in root.findall(REL_NS + "Relationship")
    }


def _document_sect_prs(zf: zipfile.ZipFile):
    root = ET.fromstring(zf.read("word/document.xml"))
    return root.findall(".//" + qn("w:sectPr"))


def _refs_for_section(section):
    refs = []
    for tag, kind in (
        (qn("w:headerReference"), "header"),
        (qn("w:footerReference"), "footer"),
    ):
        for ref in section.findall(tag):
            refs.append(
                {
                    "kind": kind,
                    "type": ref.get(qn("w:type")),
                    "rid": ref.get(qn("r:id")),
                }
            )
    return refs


def _texts_from_part(zf: zipfile.ZipFile, part_name: str) -> str:
    root = ET.fromstring(zf.read(part_name))
    return "".join(text.text or "" for text in root.findall(".//" + qn("w:t")))


def _field_instructions_from_part(zf: zipfile.ZipFile, part_name: str) -> list[str]:
    root = ET.fromstring(zf.read(part_name))
    return [
        "".join(element.itertext()).strip()
        for element in root.findall(".//" + qn("w:instrText"))
    ]


class SectionHeaderFooterTest(unittest.TestCase):
    def test_preserves_explicit_section_headers_footers_and_media(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp = Path(tmp_dir)
            image = tmp / "tiny.png"
            source = tmp / "source.docx"
            output = tmp / "output.docx"
            _tiny_png(image)

            doc = Document()
            doc.settings.odd_and_even_pages_header_footer = True
            first = doc.sections[0]
            first.orientation = WD_ORIENT.PORTRAIT
            first.page_width = Cm(21.0)
            first.page_height = Cm(29.7)
            _configure_section(first, "S1", image)
            doc.add_paragraph("S1 body")

            second = doc.add_section(WD_SECTION.NEW_PAGE)
            second.orientation = WD_ORIENT.LANDSCAPE
            second.page_width = Cm(29.7)
            second.page_height = Cm(21.0)
            _configure_section(second, "S2")
            doc.add_paragraph("S2 body")

            third = doc.add_section(WD_SECTION.CONTINUOUS)
            third.orientation = WD_ORIENT.PORTRAIT
            third.page_width = Cm(21.0)
            third.page_height = Cm(29.7)
            _configure_section(third, "S3")
            doc.add_paragraph("S3 body")
            doc.save(source)

            _export_round_trip(source, output)

            with zipfile.ZipFile(output) as zf:
                names = set(zf.namelist())
                document_rels = _rels(zf, "word/_rels/document.xml.rels")
                sections = _document_sect_prs(zf)
                self.assertGreaterEqual(len(sections), 3)

                orientations = [
                    _child_attr(sect, qn("w:pgSz"), qn("w:orient"), "portrait")
                    for sect in sections[:3]
                ]
                self.assertIn("landscape", orientations)
                section_types = [
                    _child_attr(sect, qn("w:type"), qn("w:val"), "nextPage")
                    for sect in sections[:3]
                ]
                self.assertIn("nextPage", section_types)
                self.assertIn("continuous", section_types)
                self.assertTrue(any(sect.find(qn("w:titlePg")) is not None for sect in sections))

                settings_root = ET.fromstring(zf.read("word/settings.xml"))
                self.assertIsNotNone(settings_root.find(".//" + qn("w:evenAndOddHeaders")))

                labels_by_section = ["S1", "S2", "S3"]
                for section, label in zip(sections[:3], labels_by_section, strict=True):
                    refs = _refs_for_section(section)
                    header_types = {ref["type"] for ref in refs if ref["kind"] == "header"}
                    footer_types = {ref["type"] for ref in refs if ref["kind"] == "footer"}
                    self.assertEqual(header_types, {"default", "first", "even"})
                    self.assertEqual(footer_types, {"default", "first", "even"})

                    default_header = next(
                        ref for ref in refs if ref["kind"] == "header" and ref["type"] == "default"
                    )
                    header_target = _part_name(document_rels[default_header["rid"]]["target"])
                    header_text = _texts_from_part(zf, header_target)
                    self.assertIn(f"{label} default header", header_text)
                    for other in set(labels_by_section) - {label}:
                        self.assertNotIn(f"{other} default header", header_text)

                    for ref in refs:
                        rel = document_rels[ref["rid"]]
                        part_name = _part_name(rel["target"])
                        self.assertIn(part_name, names)
                        self.assertTrue(part_name.startswith(f"word/{ref['kind']}"))
                        if ref["kind"] == "footer":
                            footer_text = _texts_from_part(zf, part_name)
                            self.assertIn(f"{label} {ref['type']} footer", footer_text)

                footer_xml = "\n".join(
                    zf.read(name).decode("utf-8")
                    for name in names
                    if name.startswith("word/footer") and name.endswith(".xml")
                )
                self.assertIn("PAGE", footer_xml)
                self.assertNotIn("NUMPAGES", footer_xml)
                self.assertIn("<w:instrText", footer_xml)
                self.assertNotIn("<w:t>PAGE</w:t>", footer_xml)
                self.assertNotIn("<w:t>NUMPAGES</w:t>", footer_xml)
                for name in names:
                    if not name.startswith("word/footer") or not name.endswith(".xml"):
                        continue
                    instructions = _field_instructions_from_part(zf, name)
                    self.assertLessEqual(instructions.count("PAGE"), 1)
                    self.assertNotIn("NUMPAGES", instructions)

                image_header = next(
                    name
                    for name in names
                    if name.startswith("word/header")
                    and "S1 default header" in zf.read(name).decode("utf-8")
                )
                header_rels_name = f"word/_rels/{Path(image_header).name}.rels"
                self.assertIn(header_rels_name, names)
                header_rels = _rels(zf, header_rels_name)
                image_rels = [rel for rel in header_rels.values() if rel["type"] == IMAGE_REL]
                self.assertTrue(image_rels)
                for rel in image_rels:
                    self.assertIn(_part_name(rel["target"]), names)

    def test_inherited_section_header_does_not_create_duplicate_part(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp = Path(tmp_dir)
            source = tmp / "source.docx"
            output = tmp / "output.docx"

            doc = Document()
            first = doc.sections[0]
            first.header.is_linked_to_previous = False
            _set_part_text(first.header, "Shared inherited header")
            doc.add_paragraph("First section")

            second = doc.add_section(WD_SECTION.NEW_PAGE)
            second.orientation = WD_ORIENT.LANDSCAPE
            second.page_width = Cm(29.7)
            second.page_height = Cm(21.0)
            doc.add_paragraph("Second section inherits")

            third = doc.add_section(WD_SECTION.NEW_PAGE)
            third.header.is_linked_to_previous = False
            _set_part_text(third.header, "Third independent header")
            doc.add_paragraph("Third section")
            doc.save(source)

            _export_round_trip(source, output)

            with zipfile.ZipFile(output) as zf:
                names = set(zf.namelist())
                document_rels = _rels(zf, "word/_rels/document.xml.rels")
                sections = _document_sect_prs(zf)
                self.assertGreaterEqual(len(sections), 3)

                default_refs = []
                effective_headers = []
                current_header = None
                for section in sections[:3]:
                    refs = [
                        ref
                        for ref in _refs_for_section(section)
                        if ref["kind"] == "header" and ref["type"] == "default"
                    ]
                    default_refs.extend(refs)
                    if refs:
                        target = _part_name(document_rels[refs[0]["rid"]]["target"])
                        current_header = _texts_from_part(zf, target)
                    effective_headers.append(current_header)

                second_refs = [
                    ref
                    for ref in _refs_for_section(sections[1])
                    if ref["kind"] == "header" and ref["type"] == "default"
                ]
                self.assertEqual(second_refs, [])
                self.assertEqual(effective_headers[0], "Shared inherited header")
                self.assertEqual(effective_headers[1], "Shared inherited header")
                self.assertEqual(effective_headers[2], "Third independent header")

                header_targets = {
                    _part_name(document_rels[ref["rid"]]["target"])
                    for ref in default_refs
                }
                self.assertEqual(len(header_targets), 2)
                self.assertTrue(all(target in names for target in header_targets))


if __name__ == "__main__":
    unittest.main()
