import pytest
from docx import Document
from docx.oxml.ns import qn

from docxtool.document.engine.inline_effects import (
    apply_colon_bold,
    apply_inline_lead_bold,
    apply_key_value_line_format,
    apply_responsibility_line,
    apply_special_bold,
    enforce_inline_heading2_format,
    handle_heading_period,
    normalize_responsibility_lines,
)
from docxtool.document.style_config import StyleRule


def test_special_bold_keeps_body_after_lead_sentence_not_bold() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("一是提高认识。后续正文")

    apply_special_bold(paragraph, paragraph.text)

    runs = [(run.text, run.font.bold) for run in paragraph.runs if run.text]
    assert runs == [("一是提高认识。", True), ("后续正文", False)]


def test_colon_bold_only_bolds_short_label() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("联系人：张三")

    apply_colon_bold(paragraph, paragraph.text)

    runs = [(run.text, run.font.bold) for run in paragraph.runs if run.text]
    assert runs == [("联系人：", True), ("张三", False)]


def test_colon_bold_ignores_numeric_colon_and_uses_later_semantic_colon() -> None:
    numeric_doc = Document()
    numeric = numeric_doc.add_paragraph("11:00")
    apply_colon_bold(numeric, numeric.text)
    assert [(run.text, run.font.bold) for run in numeric.runs if run.text] == [("11:00", None)]

    labeled_doc = Document()
    labeled = labeled_doc.add_paragraph("时间：11:00")
    apply_colon_bold(labeled, labeled.text)
    assert [(run.text, run.font.bold) for run in labeled.runs if run.text] == [
        ("时间：", True),
        ("11:00", False),
    ]


def test_colon_bold_uses_raw_offset_with_whitespace_around_colons() -> None:
    numeric_doc = Document()
    numeric = numeric_doc.add_paragraph("11　：　00")
    apply_colon_bold(numeric, numeric.text)
    assert [(run.text, run.font.bold) for run in numeric.runs if run.text] == [
        ("11　：　00", None),
    ]

    labeled_doc = Document()
    labeled = labeled_doc.add_paragraph("时间 ： 11:00")
    apply_colon_bold(labeled, labeled.text)
    assert [(run.text, run.font.bold) for run in labeled.runs if run.text] == [
        ("时间 ：", True),
        (" 11:00", False),
    ]

    prefixed_doc = Document()
    prefixed = prefixed_doc.add_paragraph("1 : 2 标签：内容")
    apply_colon_bold(prefixed, prefixed.text)
    assert [(run.text, run.font.bold) for run in prefixed.runs if run.text] == [
        ("1 : 2 ", False),
        ("标签：", True),
        ("内容", False),
    ]


@pytest.mark.parametrize(
    ("value", "prefix"),
    [
        ("时间11:00 标签：内容", "时间11:00 "),
        ("版本1:2 标签：内容", "版本1:2 "),
        ("会议时间 11 : 00 标签：内容", "会议时间 11 : 00 "),
        ("序号1：2 标签：内容", "序号1：2 "),
    ],
)
def test_colon_bold_keeps_prefixed_numeric_expression_normal(
    value: str,
    prefix: str,
) -> None:
    doc = Document()
    paragraph = doc.add_paragraph(value)

    apply_colon_bold(paragraph, paragraph.text)

    assert [(run.text, run.font.bold) for run in paragraph.runs if run.text] == [
        (prefix, False),
        ("标签：", True),
        ("内容", False),
    ]


def test_responsibility_line_splits_repeated_labels_and_clears_indent() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("责任单位：区政府责任单位：商务局")

    apply_responsibility_line(paragraph, paragraph.text)

    assert paragraph.text == "责任单位：区政府\n责任单位：商务局"
    assert [run.font.bold for run in paragraph.runs if run.text.startswith("责任单位：")] == [True, True]
    indent = paragraph._element.get_or_add_pPr().find(qn("w:ind"))
    assert indent is not None and indent.get(qn("w:firstLineChars")) == "0"


def test_key_value_line_format_uses_left_indent_when_manual_break_exists() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("责任单位：区政府")
    run.add_break()
    paragraph.add_run("联系人：张三")

    apply_key_value_line_format(paragraph)

    indent = paragraph._element.get_or_add_pPr().find(qn("w:ind"))
    assert indent is not None
    assert indent.get(qn("w:leftChars")) == "200"
    assert indent.get(qn("w:firstLineChars")) == "0"


def test_inline_lead_bold_and_heading2_format_split_run_styles() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("推动落实。后续正文")
    rule = StyleRule.default_for_row(5)

    apply_inline_lead_bold(paragraph, paragraph.text, rule)
    assert [(run.text, run.font.bold) for run in paragraph.runs if run.text] == [
        ("推动落实。", True),
        ("后续正文", False),
    ]

    heading = doc.add_paragraph("（一）标题内容。这里是正文")
    enforce_inline_heading2_format(heading, True, False)
    assert [(run.text, run.font.bold) for run in heading.runs if run.text] == [
        ("（一）标题内容。", True),
        ("这里是正文", False),
    ]


def test_heading_period_handler_removes_only_standalone_heading_period() -> None:
    assert handle_heading_period("（一）提高认识。") == "（一）提高认识"
    assert handle_heading_period("（一）提高认识。这里是正文") == "（一）提高认识。这里是正文"
    assert normalize_responsibility_lines("“责 任 单 位：区政府”") == ["责任单位：区政府"]
