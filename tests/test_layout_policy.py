from __future__ import annotations

from docx import Document

from docxtool.document.analysis.layout_policy import (
    LayoutPolicy,
    validate_layout_preservation,
)
from docxtool.document.analysis.document_structure import ElementKind
from docxtool.document.models import DocumentData, ParagraphData, ParagraphFeatures
from docxtool.document.engine import export_doc
from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import load_rules_and_settings


def _source(path, content_rows: tuple[str, ...], *, title: str) -> None:
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    document.add_paragraph("2026年8月13日")
    marker = document.add_paragraph("附件1")
    marker.paragraph_format.page_break_before = True
    document.add_paragraph(title)
    for row in content_rows:
        document.add_paragraph(row)
    document.save(path)


def _load(path, mode: str):
    rules, settings, features = load_rules_and_settings({"mode": mode})
    return rules, settings, DocxImporter().load(str(path), rules, features=features)


def test_plain_attachment_body_uses_normalize_policy(tmp_path) -> None:
    source = tmp_path / "plain-attachment.docx"
    _source(
        source,
        ("一、工作目标", "正文内容按照正常公文规则进行排版。"),
        title="工作方案",
    )

    _rules, _settings, data = _load(source, "smart")

    attachment_bodies = [
        paragraph for paragraph in data.paragraphs
        if paragraph.type_id == "attachment_body"
    ]
    assert attachment_bodies
    assert all(
        paragraph.meta["layout_policy"] == LayoutPolicy.NORMALIZE.value
        for paragraph in attachment_bodies
    )


def test_attachment_manual_columns_preserve_layout_end_to_end(tmp_path) -> None:
    source = tmp_path / "manual-columns.docx"
    output = tmp_path / "manual-columns-output.docx"
    rows = (
        "部门甲：人员甲\t部门乙：人员乙",
        "部门丙：人员丙      部门丁：人员丁",
        "部门戊：人员戊\u3000\u3000部门己：人员己",
        "部门庚：人员庚\u00a0\u00a0部门辛：人员辛",
    )
    _source(source, rows, title="参会人员名单")
    rules, settings, data = _load(source, "normalize")

    protected = [
        paragraph for paragraph in data.paragraphs
        if paragraph.meta["layout_policy"] == LayoutPolicy.PRESERVE_LAYOUT.value
    ]
    assert [paragraph.type_id for paragraph in protected] == ["attachment_body"] * 4
    assert [paragraph.text for paragraph in protected] == list(rows)

    export_doc(data, rules, settings, str(output))
    exported = Document(output)
    exported_rows = [paragraph for paragraph in exported.paragraphs if paragraph.text in rows]
    assert [paragraph.text for paragraph in exported_rows] == list(rows)
    assert all(paragraph.style.style_id == "DCT-AttachmentBody" for paragraph in exported_rows)
    assert all(paragraph.runs[0].font.size.pt == 16 for paragraph in exported_rows)


def test_body_key_value_lines_are_not_layout_protected(tmp_path) -> None:
    source = tmp_path / "key-values.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("责任单位：部门甲")
    document.add_paragraph("主要任务：完成有关工作。")
    document.save(source)

    _rules, _settings, data = _load(source, "smart")

    assert all(
        paragraph.meta["layout_policy"] == LayoutPolicy.NORMALIZE.value
        for paragraph in data.paragraphs
    )


def test_layout_policy_is_stable_across_processing_modes(tmp_path) -> None:
    source = tmp_path / "mode-columns.docx"
    rows = (
        "部门甲：人员甲\t部门乙：人员乙",
        "部门丙：人员丙      部门丁：人员丁",
    )
    _source(source, rows, title="参会人员名单")

    for mode in ("strict", "smart", "normalize"):
        _rules, _settings, data = _load(source, mode)
        protected = [
            paragraph for paragraph in data.paragraphs
            if paragraph.meta["layout_policy"] == LayoutPolicy.PRESERVE_LAYOUT.value
        ]
        assert [paragraph.type_id for paragraph in protected] == ["attachment_body"] * 2
        assert [paragraph.text for paragraph in protected] == list(rows)


def test_preserved_objects_use_preserve_object_policy() -> None:
    data = DocumentData(paragraphs=[
        ParagraphData("", "__table__", "", ParagraphFeatures()),
        ParagraphData("", "__image__", "", ParagraphFeatures()),
        ParagraphData("图1", "__object_caption__", "图1", ParagraphFeatures()),
    ])
    from docxtool.document.analysis.document_structure import analyze_document_structure
    from docxtool.document.analysis.layout_policy import assign_layout_policies

    structure = analyze_document_structure(data)
    assign_layout_policies(data, structure)

    assert [paragraph.meta["layout_policy"] for paragraph in data.paragraphs] == [
        LayoutPolicy.PRESERVE_OBJECT.value,
        LayoutPolicy.PRESERVE_OBJECT.value,
        LayoutPolicy.PRESERVE_OBJECT.value,
    ]
    assert [element.kind for element in structure.elements] == [
        ElementKind.TABLE,
        ElementKind.FIGURE,
        ElementKind.CAPTION,
    ]


def test_layout_preservation_invariant_rejects_changed_text() -> None:
    paragraph = ParagraphData(
        "已改变文本",
        "attachment_body",
        "原始\t文本",
        ParagraphFeatures(),
        {"layout_policy": LayoutPolicy.PRESERVE_LAYOUT.value},
    )

    import pytest

    with pytest.raises(ValueError, match="preserve-layout paragraph text changed"):
        validate_layout_preservation(DocumentData(paragraphs=[paragraph]))
