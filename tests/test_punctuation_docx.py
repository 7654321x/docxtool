import hashlib
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

from docxtool.document.engine.punctuation_docx import PunctuationScope, normalize_docx_punctuation
from docxtool.security.docx_integrity import validate_docx_integrity


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"w": W_NS, "r": R_NS, "rel": REL_NS}


def _document_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    return "".join(root.xpath(".//w:t/text()", namespaces=NS))


def _xml_root(path: Path, name: str) -> etree._Element:
    with zipfile.ZipFile(path) as archive:
        return etree.fromstring(archive.read(name))


def _rels_root(path: Path, name: str) -> etree._Element:
    with zipfile.ZipFile(path) as archive:
        return etree.fromstring(archive.read(name))


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_field(paragraph) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    paragraph.add_run("1,页.")
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def test_cross_run_normalization_preserves_run_format_and_integrity(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    run1 = paragraph.add_run('他说"')
    run1.bold = True
    run2 = paragraph.add_run("你好,世界")
    run2.italic = True
    run3 = paragraph.add_run('..."')
    run3.underline = True
    doc.save(source)
    before_hash = hashlib.sha256(source.read_bytes()).hexdigest()

    report = normalize_docx_punctuation(source, output)

    assert report.integrity_ok is True
    assert report.replacements == 4
    assert hashlib.sha256(source.read_bytes()).hexdigest() == before_hash
    result = Document(output)
    out_paragraph = result.paragraphs[0]
    assert out_paragraph.text == "他说“你好，世界……”"
    assert [run.text for run in out_paragraph.runs] == ["他说“", "你好，世界", "……”"]
    assert out_paragraph.runs[0].bold is True
    assert out_paragraph.runs[1].italic is True
    assert out_paragraph.runs[2].underline is True
    validate_docx_integrity(output)


def test_preserves_fields_hyperlink_targets_breaks_tabs_and_bookmarks(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    normal = doc.add_paragraph("正文,要处理.")
    normal.add_run().add_tab()
    normal.add_run("继续!")
    break_run = normal.add_run()
    break_run.add_break(WD_BREAK.PAGE)
    normal.add_run("分页后?")
    field_para = doc.add_paragraph("页码字段,不处理.")
    _add_page_field(field_para)
    link_para = doc.add_paragraph("链接前,处理.")
    _add_hyperlink(link_para, "https://example.com/a,b?x=1.2", "链接,不处理.")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "42")
    bookmark_start.set(qn("w:name"), "mark")
    normal._p.insert(0, bookmark_start)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "42")
    normal._p.append(bookmark_end)
    doc.save(source)

    normalize_docx_punctuation(source, output)

    text = _document_text(output)
    assert "正文，要处理。继续！分页后？" in text
    assert "页码字段,不处理.1,页." in text
    assert "链接前，处理。链接,不处理." in text
    root = _xml_root(output, "word/document.xml")
    assert len(root.xpath(".//w:br[@w:type='page']", namespaces=NS)) == 1
    assert len(root.xpath(".//w:tab", namespaces=NS)) == 1
    assert len(root.xpath(".//w:bookmarkStart[@w:name='mark']", namespaces=NS)) == 1
    assert len(root.xpath(".//w:bookmarkEnd[@w:id='42']", namespaces=NS)) == 1
    assert len(root.xpath(".//w:instrText[text()=' PAGE ']", namespaces=NS)) == 1
    rels = _rels_root(output, "word/_rels/document.xml.rels")
    targets = rels.xpath(".//rel:Relationship/@Target", namespaces=NS)
    assert "https://example.com/a,b?x=1.2" in targets


def test_table_scope_is_opt_in(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    default_output = tmp_path / "default.docx"
    scoped_output = tmp_path / "scoped.docx"
    doc = Document()
    doc.add_paragraph("正文,处理.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格,默认不处理."
    doc.save(source)

    normalize_docx_punctuation(source, default_output)
    normalize_docx_punctuation(source, scoped_output, scope=PunctuationScope(tables=True))

    assert "正文，处理。表格,默认不处理." in _document_text(default_output)
    assert "正文，处理。表格，默认不处理。" in _document_text(scoped_output)


def test_header_footer_scope_is_opt_in(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    default_output = tmp_path / "default.docx"
    scoped_output = tmp_path / "scoped.docx"
    doc = Document()
    doc.add_paragraph("正文,处理.")
    section = doc.sections[0]
    section.header.paragraphs[0].text = "页眉,默认不处理."
    section.footer.paragraphs[0].text = "页脚,默认不处理."
    doc.save(source)

    normalize_docx_punctuation(source, default_output)
    normalize_docx_punctuation(source, scoped_output, scope=PunctuationScope(headers=True, footers=True))

    assert "正文，处理。" in _document_text(default_output)
    assert "页眉,默认不处理." in "".join(_xml_root(default_output, "word/header1.xml").xpath(".//w:t/text()", namespaces=NS))
    assert "页脚,默认不处理." in "".join(_xml_root(default_output, "word/footer1.xml").xpath(".//w:t/text()", namespaces=NS))
    assert "页眉，默认不处理。" in "".join(_xml_root(scoped_output, "word/header1.xml").xpath(".//w:t/text()", namespaces=NS))
    assert "页脚，默认不处理。" in "".join(_xml_root(scoped_output, "word/footer1.xml").xpath(".//w:t/text()", namespaces=NS))
