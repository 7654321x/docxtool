from __future__ import annotations

import importlib

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import docxtool.document.importer as importer_module
from docxtool.document.importer import extract_features
from docxtool.document.importing import numbering as numbering_module
from docxtool.document.importing.features import extract_paragraph_features
from docxtool.document.importing.numbering import detect_numbering_prefix
from docxtool.document.models import ParagraphFeatures
from docxtool.document.segmentation import source_locator as source_locator_module


def _set_east_asia_font(run, name: str, ascii_name: str = "Arial") -> None:
    """给测试 run 写入中英文字体，传入 run 和字体名，无返回值。"""
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), ascii_name)
    fonts.set(qn("w:hAnsi"), ascii_name)


def test_importing_features_extracts_physical_run_and_locator_facts() -> None:
    """验证物理导入层只提取格式和定位事实，不依赖 importer 分类链。"""
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("一、测试标题")
    _set_east_asia_font(run, "SimHei", "Calibri")
    run.font.size = Pt(16)
    run.bold = True

    features = extract_paragraph_features(
        paragraph,
        3,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )

    assert features.text == "一、测试标题"
    assert features.paragraph_index == 3
    assert features.source_locator_status == "confirmed"
    assert features.numbering_prefix == "一、"
    assert features.source_run_spans[0].east_asia_font_name == "SimHei"
    assert features.source_run_spans[0].ascii_font_name == "Calibri"
    assert features.source_run_spans[0].font_size_pt == 16.0
    assert features.source_run_spans[0].bold is True


def test_importer_extract_features_keeps_legacy_facade() -> None:
    """验证旧 `document.importer.extract_features` 入口仍转发并返回同等事实。"""
    document = Document()
    paragraph = document.add_paragraph("正文内容")

    direct = extract_paragraph_features(
        paragraph,
        0,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )
    legacy = extract_features(paragraph, 0)

    assert legacy.text == direct.text
    assert legacy.source_physical_text == direct.source_physical_text
    assert legacy.source_locator_status == direct.source_locator_status
    assert legacy.source_run_spans == direct.source_run_spans


def test_importing_features_calls_source_locator_builder(monkeypatch) -> None:
    """物理特征提取应调用 source locator 模块中的初始特征构建器。"""
    document = Document()
    paragraph = document.add_paragraph()
    sentinel = ParagraphFeatures()
    calls = []

    def fake_builder(raw_text: str, paragraph_index: int) -> ParagraphFeatures:
        calls.append((raw_text, paragraph_index))
        return sentinel

    monkeypatch.setattr(
        source_locator_module,
        "build_physical_source_features",
        fake_builder,
    )

    result = extract_paragraph_features(
        paragraph,
        4,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )

    assert result is sentinel
    assert calls == [("", 4)]


def test_importer_extract_features_keeps_legacy_monkeypatch_path(monkeypatch) -> None:
    """旧 importer 私有依赖被 patch 后，facade 仍应调用该替身。"""
    document = Document()
    paragraph = document.add_paragraph("正文")
    sentinel = ParagraphFeatures(text="sentinel")
    calls = []

    def fake_extract(paragraph_value, index_value, *, detect_numbering_prefix_func):
        calls.append((paragraph_value, index_value, detect_numbering_prefix_func))
        return sentinel

    monkeypatch.setattr(
        importer_module,
        "_importing_extract_paragraph_features",
        fake_extract,
    )

    result = importer_module.extract_features(paragraph, 9)

    assert result is sentinel
    assert calls == [(paragraph, 9, importer_module._detect_numbering_prefix)]


def test_importing_features_calls_physical_numbering_application(monkeypatch) -> None:
    """物理特征提取应通过 numbering 模块应用已有 Word 编号事实。"""
    document = Document()
    paragraph = document.add_paragraph("自动列表标题")
    calls = []

    def fake_apply(features, paragraph_value, text, *, debug_logger):
        calls.append((features, paragraph_value, text, debug_logger))
        features.numbering_prefix = "@lvl_7"

    monkeypatch.setattr(
        numbering_module,
        "apply_physical_numbering_features",
        fake_apply,
    )

    result = extract_paragraph_features(
        paragraph,
        2,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )

    assert result.numbering_prefix == "@lvl_7"
    assert calls == [(result, paragraph, paragraph.text.strip(), importer_module.logger)]


def test_physical_numbering_application_keeps_native_and_literal_precedence() -> None:
    """原生列表只补空前缀，已有手写编号继续优先。"""
    document = Document()
    native = document.add_paragraph("自动列表标题")
    p_pr = native._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "2")
    num_pr.append(ilvl)
    p_pr.append(num_pr)

    native_features = extract_paragraph_features(
        native,
        0,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )
    literal = document.add_paragraph("1.手写编号标题")
    literal._p.get_or_add_pPr().append(num_pr)
    literal_features = extract_paragraph_features(
        literal,
        1,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )

    assert native_features.numbering_prefix == "@lvl_2"
    assert literal_features.numbering_prefix == "1."


def test_physical_numbering_application_exposes_numbering_read_failure(monkeypatch) -> None:
    """编号读取失败必须明确暴露，不能靠 Heading 样式静默降级。"""
    document = Document()
    paragraph = document.add_paragraph("样式标题", style="Heading 1")

    def fail_word_numbering(paragraph_element, text):
        raise ValueError("fixture failure")

    monkeypatch.setattr(numbering_module, "word_list_level_prefix", fail_word_numbering)

    with pytest.raises(ValueError, match="fixture failure"):
        extract_paragraph_features(
            paragraph,
            0,
            detect_numbering_prefix_func=detect_numbering_prefix,
        )


def test_importing_features_calls_physical_format_application(monkeypatch) -> None:
    """物理特征编排应调用独立格式读取模块，并使用同一特征对象。"""
    physical_format_module = importlib.import_module(
        "docxtool.document.importing.physical_format"
    )
    document = Document()
    paragraph = document.add_paragraph("格式正文")
    calls = []

    def fake_apply(paragraph_value, features_value):
        calls.append((paragraph_value, features_value))
        features_value.first_run_font_name = "sentinel"

    monkeypatch.setattr(
        physical_format_module,
        "apply_physical_format_features",
        fake_apply,
    )

    result = extract_paragraph_features(
        paragraph,
        5,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )

    assert result.first_run_font_name == "sentinel"
    assert calls == [(paragraph, result)]


def test_importing_features_keeps_paragraph_alignment_and_indent_facts() -> None:
    """格式迁移后仍按原单位保存段落对齐和首行缩进。"""
    document = Document()
    paragraph = document.add_paragraph("格式正文")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(18)

    features = extract_paragraph_features(
        paragraph,
        0,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )

    assert features.alignment == "center (1)"
    assert features.first_line_indent == 0.635
