import logging
import tempfile
import unittest
from pathlib import Path

from docx import Document

from importer import DocxImporter
from style_config import StyleRule, logger


def _rules():
    return StyleRule.from_config()


class ImporterHeadingFlowTest(unittest.TestCase):
    def setUp(self):
        logger.setLevel(logging.ERROR)
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)

    def tearDown(self):
        self.tmp.cleanup()

    def _load_lines(self, lines):
        doc = Document()
        for line in lines:
            doc.add_paragraph(line)
        path = self.root / "input.docx"
        doc.save(path)
        return DocxImporter().load(str(path), _rules())

    def _find_original(self, data, prefix):
        for pd in data.paragraphs:
            if pd.original_text.startswith(prefix):
                return pd
        self.fail(f"missing paragraph starting with {prefix!r}")

    def test_heading4_can_return_to_heading3_and_child_numbering_resets(self):
        data = self._load_lines([
            "总题目报告",
            "区委：",
            "一、2026年上半年工作总结",
            "（一）工作开展情况",
            "1.坚持党的全面领导",
            "（1）持续强化思想理论武装。正文内容。",
            "（2）始终坚持党的全面领导。正文内容。",
            "2.聚焦主责主业，充分发挥专门协商机构作用",
            "（1）有序推进各类协商议政活动。正文内容。",
        ])

        second_heading3 = self._find_original(data, "2.聚焦主责主业")
        first_child_after_second_heading3 = self._find_original(data, "（1）有序推进")

        self.assertEqual(second_heading3.type_id, "heading3")
        self.assertEqual(second_heading3.meta["numbering"], "2.")
        self.assertEqual(first_child_after_second_heading3.type_id, "heading4")
        self.assertEqual(first_child_after_second_heading3.meta["numbering"], "（1）")

    def test_heading4_can_return_to_heading2(self):
        data = self._load_lines([
            "总题目报告",
            "区委：",
            "一、2026年上半年工作总结",
            "（一）工作开展情况",
            "1.坚持党的全面领导",
            "（1）持续强化思想理论武装。正文内容。",
            "（二）特色亮点工作情况",
        ])

        heading2 = self._find_original(data, "（二）特色亮点")

        self.assertEqual(heading2.type_id, "heading2")
        self.assertEqual(heading2.meta["numbering"], "（二）")

    def test_heading3_can_return_to_heading2(self):
        data = self._load_lines([
            "总题目报告",
            "区委：",
            "一、2026年上半年工作总结",
            "（一）工作开展情况",
            "1.坚持党的全面领导",
            "（二）特色亮点工作情况",
            "1.基层协商平台建设持续深化",
        ])

        heading2 = self._find_original(data, "（二）特色亮点")
        first_heading3_under_second_heading2 = self._find_original(data, "1.基层协商")

        self.assertEqual(heading2.type_id, "heading2")
        self.assertEqual(heading2.meta["numbering"], "（二）")
        self.assertEqual(first_heading3_under_second_heading2.type_id, "heading3")
        self.assertEqual(first_heading3_under_second_heading2.meta["numbering"], "1.")

    def test_inline_newline_keeps_split_heading4_sequence(self):
        data = self._load_lines([
            "总题目报告",
            "区委：",
            "一、2026年上半年工作总结",
            "（一）工作开展情况",
            "1.坚持党的全面领导",
            "（1）持续强化思想理论武装。正文内容。\n      (2)推动两大园区建设。正文内容。",
            "（3）积极服务产业发展。正文内容。",
        ])

        second_child = self._find_original(data, "(2)推动两大园区")
        third_child = self._find_original(data, "（3）积极服务")

        self.assertEqual(second_child.type_id, "heading4")
        self.assertEqual(second_child.meta["numbering"], "（2）")
        self.assertEqual(third_child.type_id, "heading4")
        self.assertEqual(third_child.meta["numbering"], "（3）")

    def test_heading1_does_not_jump_directly_to_heading4(self):
        data = self._load_lines([
            "总题目报告",
            "区委：",
            "一、2026年上半年工作总结",
            "（1）不应直接成为四级标题。正文内容。",
        ])

        illegal_child = self._find_original(data, "（1）不应直接")

        self.assertNotEqual(illegal_child.type_id, "heading4")


if __name__ == "__main__":
    unittest.main()
