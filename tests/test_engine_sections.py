from copy import deepcopy

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree

from docxtool.document.engine.sections import (
    OFFICIAL_GRID_FONT_SIZE_PT,
    apply_page_settings,
    copy_paragraph_sectPr,
    doc_grid_char_space,
    line_spacing_twips,
    replace_body_sectPr,
    section_margins_cm,
    sectPr_is_landscape,
    set_sectPr_page_layout,
)
from docxtool.document.style_config import PageSettings


def _settings() -> PageSettings:
    """构造测试页面设置，返回带标准公文页边距和网格的 PageSettings。"""
    return PageSettings()


def _style_element(document: Document, style_id: str):
    return next(
        style
        for style in document.styles._element.findall(qn("w:style"))
        if style.get(qn("w:styleId")) == style_id
    )


def _style_without_font_sizes_xml(style_element) -> str:
    copied = deepcopy(style_element)
    run_properties = copied.find(qn("w:rPr"))
    if run_properties is not None:
        for tag in ("w:sz", "w:szCs"):
            for element in run_properties.findall(qn(tag)):
                run_properties.remove(element)
    return copied.xml


def test_section_margins_rotate_physical_edges_for_landscape() -> None:
    """横向分节应按物理边旋转页边距，而不是只交换页面宽高。"""
    settings = _settings()

    assert section_margins_cm(settings, False) == (3.7, 3.5, 2.8, 2.6)
    assert section_margins_cm(settings, True) == (2.8, 2.6, 3.5, 3.7)


def test_doc_grid_char_space_uses_ooxml_point_delta_units() -> None:
    """docGrid 字距应使用磅差值乘 4096，保证 WPS 28 字版心稳定。"""
    settings = _settings()
    content_width_twips = (
        int(round(settings.page_width_cm * 567))
        - int(round(settings.margin_left_cm * 567))
        - int(round(settings.margin_right_cm * 567))
    )

    assert (
        doc_grid_char_space(
            content_width_twips,
            settings.chars_per_line,
            OFFICIAL_GRID_FONT_SIZE_PT,
        )
        == -842
    )
    assert line_spacing_twips(settings) == 560


def test_set_sectPr_page_layout_writes_portrait_page_and_grid() -> None:
    """纵向分节布局应写入页面尺寸、页边距和文档网格。"""
    sect_pr = OxmlElement("w:sectPr")

    set_sectPr_page_layout(sect_pr, _settings())

    pg_sz = sect_pr.find(qn("w:pgSz"))
    pg_mar = sect_pr.find(qn("w:pgMar"))
    doc_grid = sect_pr.find(qn("w:docGrid"))
    assert pg_sz.get(qn("w:orient")) is None
    assert pg_mar.get(qn("w:top")) == str(int(round(3.7 * 567)))
    assert pg_mar.get(qn("w:left")) == str(int(round(2.8 * 567)))
    assert doc_grid.get(qn("w:charsPerLine")) == "28"
    assert doc_grid.get(qn("w:linePitch")) == "560"


def test_set_sectPr_page_layout_preserves_landscape_size_and_rotates_margins() -> None:
    """横向分节应保留源页面尺寸，设置横向标记并旋转页边距。"""
    sect_pr = OxmlElement("w:sectPr")
    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "16838")
    pg_sz.set(qn("w:h"), "11906")
    sect_pr.append(pg_sz)

    assert sectPr_is_landscape(sect_pr) is True
    set_sectPr_page_layout(sect_pr, _settings())

    pg_mar = sect_pr.find(qn("w:pgMar"))
    assert pg_sz.get(qn("w:orient")) == "landscape"
    assert pg_sz.get(qn("w:w")) == "16838"
    assert pg_sz.get(qn("w:h")) == "11906"
    assert pg_mar.get(qn("w:top")) == str(int(round(2.8 * 567)))
    assert pg_mar.get(qn("w:right")) == str(int(round(3.7 * 567)))
    assert sect_pr.find(qn("w:docGrid")).get(qn("w:charsPerLine")) == "28"


def test_copy_and_replace_section_properties_keep_single_sectPr() -> None:
    """段落和正文替换分节属性时应先移除旧 sectPr，再写入新副本。"""
    document = Document()
    paragraph = document.add_paragraph("正文")
    source_sect_pr = OxmlElement("w:sectPr")
    source_pg_mar = OxmlElement("w:pgMar")
    source_pg_mar.set(qn("w:top"), "123")
    source_sect_pr.append(source_pg_mar)

    copy_paragraph_sectPr(paragraph, source_sect_pr)
    copy_paragraph_sectPr(paragraph, source_sect_pr)

    paragraph_sections = paragraph._element.findall(".//" + qn("w:sectPr"))
    assert len(paragraph_sections) == 1
    assert paragraph_sections[0].find(qn("w:pgMar")).get(qn("w:top")) == "123"

    replace_body_sectPr(document, source_sect_pr)
    replace_body_sectPr(document, source_sect_pr)
    body_sections = document._body._element.findall(qn("w:sectPr"))
    assert len(body_sections) == 1
    assert body_sections[0].find(qn("w:pgMar")).get(qn("w:top")) == "123"


def test_apply_page_settings_writes_defaults_compat_and_grid() -> None:
    """页面设置入口应写入默认字体、兼容网格设置和正文 sectPr。"""
    document = Document()

    apply_page_settings(document, _settings())

    styles = document.styles._element
    defaults = styles.find(qn("w:docDefaults"))
    r_fonts = defaults.find(".//" + qn("w:rFonts"))
    assert r_fonts.get(qn("w:eastAsia")) == "仿宋_GB2312"
    assert r_fonts.get(qn("w:ascii")) == "Times New Roman"

    compat = document.settings._element.find(qn("w:compat"))
    names = {item.get(qn("w:name")) for item in compat.findall(qn("w:compatSetting"))}
    assert "useFELayout" in names
    assert "doNotExpand" in names

    sect_pr = document.sections[0]._sectPr
    assert sect_pr.find(qn("w:docGrid")).get(qn("w:charsPerLine")) == "28"


def test_wps_docxtool_normalizes_only_normal_size_and_preserves_headings() -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Source Normal Font"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    normal.paragraph_format.left_indent = Pt(12)
    normal.paragraph_format.line_spacing = Pt(19)
    heading1_properties = _style_element(document, "Heading1").get_or_add_rPr()
    heading1_fonts = OxmlElement("w:rFonts")
    heading1_fonts.set(qn("w:eastAsia"), "Source Heading One")
    heading1_properties.append(heading1_fonts)
    heading2_size = OxmlElement("w:sz")
    heading2_size.set(qn("w:val"), "18")
    _style_element(document, "Heading2").get_or_add_rPr().append(heading2_size)
    heading3_color = OxmlElement("w:color")
    heading3_color.set(qn("w:val"), "654321")
    _style_element(document, "Heading3").get_or_add_rPr().append(heading3_color)
    heading4_bold = OxmlElement("w:b")
    heading4_bold.set(qn("w:val"), "1")
    _style_element(document, "Heading4").get_or_add_rPr().append(heading4_bold)

    normal_before = _style_without_font_sizes_xml(_style_element(document, "Normal"))
    headings_before = {
        style_id: _style_element(document, style_id).xml
        for style_id in ("Heading1", "Heading2", "Heading3", "Heading4")
    }

    apply_page_settings(document, _settings(), style_profile="wps_docxtool")

    normal_element = _style_element(document, "Normal")
    normal_run_properties = normal_element.find(qn("w:rPr"))
    headings_after = {
        style_id: _style_element(document, style_id).xml
        for style_id in ("Heading1", "Heading2", "Heading3", "Heading4")
    }
    assert normal_run_properties.find(qn("w:sz")).get(qn("w:val")) == "32"
    assert normal_run_properties.find(qn("w:szCs")).get(qn("w:val")) == "32"
    assert _style_without_font_sizes_xml(normal_element) == normal_before
    assert headings_after == headings_before


@pytest.mark.parametrize("normal_font_size_pt", [10.5, 11.0, 12.0])
def test_wps_docxtool_official_grid_uses_fixed_16pt_baseline(
    normal_font_size_pt: float,
) -> None:
    document = Document()
    document.styles["Normal"].font.size = Pt(normal_font_size_pt)
    settings = _settings()

    apply_page_settings(document, settings, style_profile="wps_docxtool")

    normal_element = _style_element(document, "Normal")
    normal_run_properties = normal_element.find(qn("w:rPr"))
    grid = document.sections[0]._sectPr.find(qn("w:docGrid"))
    char_space = int(grid.get(qn("w:charSpace")))
    margins = document.sections[0]._sectPr.find(qn("w:pgMar"))
    page_size = document.sections[0]._sectPr.find(qn("w:pgSz"))
    content_width = (
        int(page_size.get(qn("w:w")))
        - int(margins.get(qn("w:left")))
        - int(margins.get(qn("w:right")))
    )
    assert content_width == 8845
    assert normal_run_properties.find(qn("w:sz")).get(qn("w:val")) == "32"
    assert normal_run_properties.find(qn("w:szCs")).get(qn("w:val")) == "32"
    assert grid.get(qn("w:type")) == "linesAndChars"
    assert grid.get(qn("w:charsPerLine")) == "28"
    assert grid.get(qn("w:linesPerPage")) == "22"
    assert char_space == -842
    assert char_space != 19638
    assert grid.get(qn("w:linePitch")) == "560"


def test_wps_docxtool_normalizes_missing_normal_size_over_11pt_doc_defaults() -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.size = None
    normal.base_style = None
    defaults = document.styles._element.find(qn("w:docDefaults"))
    default_properties = defaults.find(f"{qn('w:rPrDefault')}/{qn('w:rPr')}")
    size = default_properties.find(qn("w:sz"))
    size.set(qn("w:val"), "22")
    size_cs = default_properties.find(qn("w:szCs"))
    if size_cs is None:
        size_cs = OxmlElement("w:szCs")
        default_properties.append(size_cs)
    size_cs.set(qn("w:val"), "22")
    defaults_before = etree.tostring(defaults)

    apply_page_settings(document, _settings(), style_profile="wps_docxtool")

    normal_run_properties = _style_element(document, "Normal").find(qn("w:rPr"))
    grid = document.sections[0]._sectPr.find(qn("w:docGrid"))
    assert normal_run_properties.find(qn("w:sz")).get(qn("w:val")) == "32"
    assert normal_run_properties.find(qn("w:szCs")).get(qn("w:val")) == "32"
    assert etree.tostring(defaults) == defaults_before
    assert grid.get(qn("w:charSpace")) == "-842"
    assert grid.get(qn("w:linePitch")) == "560"

@pytest.mark.parametrize(
    ("settings", "doc_mode"),
    [
        (PageSettings(chars_per_line=0), ""),
        (PageSettings(lines_per_page=0), ""),
        (PageSettings(), "SCHEME"),
    ],
)
def test_wps_docxtool_does_not_normalize_normal_without_official_grid(
    settings: PageSettings,
    doc_mode: str,
) -> None:
    document = Document()
    document.styles["Normal"].font.size = Pt(11)

    apply_page_settings(
        document,
        settings,
        doc_mode,
        style_profile="wps_docxtool",
    )

    assert document.styles["Normal"].font.size.pt == 11
