from concurrent.futures import ThreadPoolExecutor
from types import SimpleNamespace

from docx import Document
import pytest
from docxtool.document.engine import export_doc
import docxtool.document.recognition.decoder as decoder
from docxtool.document.recognition.candidates import Candidate
from docxtool.document.recognition import (
    DocumentMode,
    ParagraphType,
    RecognitionConfig,
    SectionKind,
    apply_recognition,
    diagnostics_to_json,
    extract_blocks,
    extract_features,
    resolve_render_mapping,
)
from docxtool.document.recognition.features import BlockKind
from docxtool.document.style_config import PageSettings, StyleRule


def _paragraph(text, type_id="body", index=0, **meta):
    numbering_prefix = meta.pop("numbering_prefix", "")
    segment_numbering_features = meta.pop("segment_numbering_features", numbering_prefix)
    bold_char_ratio = meta.pop("bold_char_ratio", 0.0)
    return SimpleNamespace(
        text=text,
        original_text=text,
        type_id=type_id,
        features=SimpleNamespace(
            paragraph_index=index,
            alignment=meta.pop("alignment", ""),
            style_name=meta.pop("style_name", ""),
            bold=bold_char_ratio >= 0.65,
            bold_char_ratio=bold_char_ratio,
            font_size_pt=None,
            numbering_prefix=numbering_prefix,
            segment_numbering_features=segment_numbering_features,
        ),
        meta=meta,
    )


def _document(*paragraphs, mode="NORMAL"):
    return SimpleNamespace(paragraphs=list(paragraphs), doc_mode=mode)


def _rules():
    return [StyleRule.default_for_row(index) for index in range(10)]


def test_dispatch_number_vetoes_title_continuation():
    data = _document(
        _paragraph("关于推进公共服务工作的通知", "title", 0),
        _paragraph("内政发〔2026〕23号", "title_cont", 1),
    )

    apply_recognition(data)

    assert data.paragraphs[1].type_id == "dispatch_number"
    assert data.paragraphs[1].meta["recognition_provider"].startswith("structural:")


def test_numbered_meeting_label_is_metadata_not_heading():
    data = _document(
        _paragraph("2026年第一次党委会会议纪要", "title", 0),
        _paragraph("（一）缺席：无", "heading2", 1),
    )

    apply_recognition(data)

    assert data.doc_mode == "MEETING_MINUTES"
    assert data.paragraphs[1].type_id == "meeting_meta"


def test_embedded_document_title_after_signature_note():
    data = _document(
        _paragraph("2026年7月21日", "sign_date", 0),
        _paragraph("（本文有删减）", "body", 1),
        _paragraph("公共服务提升规划", "body", 2),
        _paragraph("第一章 总则", "heading1", 3),
    )

    apply_recognition(data)

    assert data.paragraphs[2].type_id == "embedded_document_title"
    trace = data.recognition_diagnostics["candidate_trace"][2]
    assert "embedded_document_title" in [item["type"] for item in trace["candidates"]]
    assert data.recognition_diagnostics["paragraphs"][2]["provider"].startswith("embedded-document:")


def test_report_bold_metadata_removed_outside_report_mode():
    paragraph = _paragraph("这是正文。后续内容。", "body", 0, report_first_sentence_bold=True)
    data = _document(paragraph, mode="NORMAL")

    apply_recognition(data)

    assert "report_first_sentence_bold" not in paragraph.meta


def test_shared_features_preserve_raw_text_and_extract_numbered_key_value():
    paragraph = _paragraph("（一）缺  席：无", "heading2", 0)
    data = _document(paragraph)
    block = extract_blocks(data)[0]

    features = extract_features(block)

    assert features.raw_text == "（一）缺  席：无"
    assert features.normalized_text == "(一)缺 席:无"
    assert features.compact_text == "(一)缺席:无"
    assert features.optional_numbering_before_label == "(一)"
    assert features.key_value_label == "缺席"


def test_multiline_date_and_attachment_is_not_a_fabricated_key_value():
    paragraph = _paragraph("2026年7月21日\n附件：1.材料", "body", 0)
    block = extract_blocks(_document(paragraph))[0]

    features = extract_features(block)

    assert features.key_value_label is None
    assert features.key_value_value is None


def test_non_text_blocks_remain_in_original_sequence():
    table = _paragraph("", "__table__", 1)
    image = _paragraph("", "__image__", 2)
    data = _document(_paragraph("标题", "title", 0), table, image, _paragraph("正文", "body", 3))

    blocks = extract_blocks(data)

    assert [block.kind for block in blocks] == [BlockKind.PARAGRAPH, BlockKind.TABLE, BlockKind.IMAGE, BlockKind.PARAGRAPH]
    assert blocks[1].paragraph_index is None
    assert blocks[1].table_index == 0


def test_mode_and_result_are_deterministic_and_idempotent():
    data = _document(
        _paragraph("2026年第一次党委会会议纪要", "title", 0),
        _paragraph("出席：甲、乙", "body", 1),
        _paragraph("（一）缺席：无", "heading2", 2),
    )

    apply_recognition(data)
    first = [(item.type_id, dict(item.meta)) for item in data.paragraphs]
    first_diagnostics = dict(data.recognition_diagnostics)
    apply_recognition(data)

    assert data.doc_mode == "MEETING_MINUTES"
    assert first == [(item.type_id, dict(item.meta)) for item in data.paragraphs]
    assert first_diagnostics == data.recognition_diagnostics
    assert data.recognition_diagnostics["validation"]["ok"] is True


def test_public_model_vocabulary_is_stable():
    assert DocumentMode.MEETING_MINUTES.value == "meeting_minutes"
    assert ParagraphType.DISPATCH_NUMBER.value == "dispatch_number"


def test_diagnostics_json_is_safe_and_configurable():
    data = _document(_paragraph("标题", "title", 0), _paragraph("正文", "body", 1))
    apply_recognition(data)
    serialized = diagnostics_to_json(data.recognition_diagnostics)

    assert "标题" not in serialized
    assert '"beam_width": 12' in serialized
    assert data.recognition_diagnostics["structure_tree"] in {"built", "unavailable"}


def test_dispatch_has_multiple_candidates_before_hard_veto():
    data = _document(
        _paragraph("关于推进公共服务工作的通知", "title", 0),
        _paragraph("内政发〔2026〕23号", "title_cont", 1),
    )
    apply_recognition(data)

    trace = data.recognition_diagnostics["candidate_trace"][1]
    assert trace["candidate_count"] >= 2
    candidate_types = [item["type"] for item in trace["candidates"]]
    assert "dispatch_number" in candidate_types
    assert "title_continuation" in candidate_types


def test_authoritative_can_disable_core_and_legacy_candidates() -> None:
    paragraph = _paragraph(
        "国发〔2026〕23号",
        "title",
        0,
        classification_kind="body",
        classification_confidence=0.95,
    )
    data = _document(paragraph)

    apply_recognition(
        data,
        RecognitionConfig(
            mode="authoritative",
            enable_core_candidates=False,
            enable_legacy_candidates=False,
        ),
    )

    assert paragraph.type_id == "dispatch_number"
    trace = data.recognition_diagnostics["candidate_trace"][0]["candidates"]
    assert {item["source"] for item in trace}.isdisjoint({"core", "legacy"})


def test_unverified_legacy_attachment_note_does_not_force_final_type() -> None:
    paragraph = _paragraph("相关附件另行发送，后续仍继续说明正文事项。", "attachment_note", 1)
    data = _document(
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 0),
        paragraph,
        _paragraph("后一段正文继续说明工作安排。", "body", 2),
    )

    apply_recognition(data)

    assert paragraph.type_id == "body"
    diagnostic = data.recognition_diagnostics["paragraphs"][1]
    assert diagnostic["legacy_type"] == "attachment_note"
    assert diagnostic["final_type"] == "body"


def test_previous_title_changes_ambiguous_centered_line_decision():
    after_title = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        _paragraph("补充说明", "body", 1, alignment="CENTER", bold_char_ratio=1.0),
    )
    after_body = _document(
        _paragraph("正文开头", "body", 0),
        _paragraph("补充说明", "body", 1, alignment="CENTER"),
    )

    apply_recognition(after_title)
    apply_recognition(after_body)

    assert after_title.paragraphs[1].meta["recognition_type"] == "title_continuation"
    assert after_body.paragraphs[1].meta["recognition_type"] == "body"


def test_front_matter_context_overrides_misused_heading_style_without_keywords():
    data = _document(
        _paragraph("关于推进基层治理重点工作的通知", "heading1", 0, alignment="CENTER", style_name="Heading 1"),
        _paragraph("区政协办公室主任  张三", "role_name", 1),
        _paragraph("2026年7月27日", "date_line", 2),
        _paragraph("各有关单位：", "addressing", 3),
        _paragraph("现将有关事项通知如下。", "body", 4),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:4]] == [
        "title", "role_name", "date_line", "addressing",
    ]
    context = data.recognition_diagnostics["document_context"]
    assert context["front_matter_positions"] == [0, 1, 2, 3]
    assert context["body_start"] == 4
    assert context["body_start_reason"] == "recipient-following-body"


def test_wrong_legacy_front_metadata_cannot_veto_structural_title() -> None:
    title = _paragraph(
        "关于推进基层治理重点工作的通知",
        "role_name",
        0,
        alignment="CENTER",
        style_name="Heading 1",
        legacy_type_id="role_name",
    )
    data = _document(
        title,
        _paragraph("各有关单位：", "body", 1),
        _paragraph("现将有关事项通知如下，请结合实际抓好落实。", "body", 2),
    )

    apply_recognition(data)

    assert title.type_id == "title"
    diagnostic = data.recognition_diagnostics["paragraphs"][0]
    assert diagnostic["final_type"] == "title"
    assert "legacy-reclassified" in diagnostic["evidence_summary"]


def test_disabling_legacy_ignores_legacy_context_and_candidates() -> None:
    title = _paragraph(
        "关于推进基层治理重点工作的通知",
        "role_name",
        0,
        alignment="CENTER",
        style_name="Heading 1",
        legacy_type_id="role_name",
    )
    data = _document(
        title,
        _paragraph("各有关单位：", "body", 1),
        _paragraph("现将有关事项通知如下，请结合实际抓好落实。", "body", 2),
    )

    apply_recognition(
        data,
        RecognitionConfig(enable_legacy_candidates=False),
    )

    assert title.type_id == "title"
    diagnostic = data.recognition_diagnostics["paragraphs"][0]
    assert all(
        "legacy" not in evidence
        for evidence in [*diagnostic["evidence_summary"], *diagnostic["title_context_evidence"]]
    )
    assert {item["source"] for item in data.recognition_diagnostics["candidate_trace"][0]["candidates"]}.isdisjoint({"legacy"})


def test_body_empty_colon_label_is_not_a_recipient_or_key_value() -> None:
    label = _paragraph("某某学院：", "body", 2, no_indent=True)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        label,
        _paragraph("调研发现有关工作正在有序推进。", "body", 3),
    )

    apply_recognition(data)

    assert label.type_id == "body"
    assert label.meta["no_indent"] is True
    assert label.meta["recognition_section"] == "body"


def test_standalone_salutation_remains_addressing_after_body_started() -> None:
    salutation = _paragraph("各位委员、同志们！", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        salutation,
        _paragraph("后续正文继续正常排版。", "body", 3),
    )

    apply_recognition(data)

    assert salutation.type_id == "addressing"
    assert "standalone-addressing" in salutation.meta["recognition_evidence"]


def test_personal_title_salutation_remains_addressing_after_body_started() -> None:
    salutation = _paragraph("余书记：", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        salutation,
        _paragraph("后续正文继续正常排版。", "body", 3),
    )

    apply_recognition(data)

    assert salutation.type_id == "addressing"
    assert "standalone-addressing" in salutation.meta["recognition_evidence"]


def test_organization_label_is_not_a_personal_title_salutation() -> None:
    label = _paragraph("某某学院：", "body", 2, no_indent=True)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        label,
        _paragraph("后续正文继续正常排版。", "body", 3),
    )

    apply_recognition(data)

    assert label.type_id == "body"


def test_organization_label_with_inline_content_is_not_a_key_value() -> None:
    paragraph = _paragraph("某某职业学院：调研工作正在有序推进。", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
    )

    apply_recognition(data)

    assert paragraph.type_id == "body"
    assert paragraph.meta["recognition_section"] == "body"


def test_front_recipient_is_structural_without_legacy_metadata() -> None:
    recipient = _paragraph("某某学院：", "body", 1)
    data = _document(
        _paragraph("关于开展专项调研的通知", "body", 0, alignment="CENTER"),
        recipient,
        _paragraph("现将有关事项通知如下，请结合实际抓好落实。", "body", 2),
    )

    apply_recognition(data)

    assert recipient.type_id == "addressing"
    assert data.recognition_diagnostics["document_context"]["body_start"] == 2
    assert "front-recipient" in recipient.meta["recognition_evidence"]


def test_explanatory_colon_sentence_remains_body_not_key_value() -> None:
    paragraph = _paragraph("主要原因如下：第一阶段已经完成。", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
    )

    apply_recognition(data)

    assert paragraph.type_id == "body"
    assert paragraph.meta["recognition_section"] == "body"
    assert "colon-explanatory-body" in paragraph.meta["recognition_evidence"]


def test_structural_key_value_generalizes_after_label_rewording() -> None:
    first = _paragraph("责任单位：综合管理部门", "body", 2)
    second = _paragraph("联系人：张某", "body", 3)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        first,
        second,
    )

    apply_recognition(data)

    assert [first.type_id, second.type_id] == ["responsibility_line", "responsibility_line"]
    assert "explicit-label" in first.meta["recognition_evidence"]
    assert "explicit-label" in second.meta["recognition_evidence"]


def test_mid_body_attachment_keyword_is_body_not_hard_attachment_note() -> None:
    paragraph = _paragraph("附件：材料清单", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
        _paragraph("后续正文继续说明材料将另行发送。", "body", 3),
    )

    apply_recognition(data)

    assert paragraph.type_id == "body"
    candidates = data.recognition_diagnostics["candidate_trace"][2]["candidates"]
    attachment = next(item for item in candidates if item["type"] == "attachment_note")
    assert attachment["hard"] is False
    assert "attachment-keyword-without-tail-context" in paragraph.meta["recognition_evidence"]


def test_empty_attachment_keyword_without_items_stays_body() -> None:
    paragraph = _paragraph("附件：", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
    )

    apply_recognition(data)

    assert paragraph.type_id == "body"


def test_tail_attachment_note_and_items_are_confirmed_by_context() -> None:
    note = _paragraph("附件：1.材料清单", "body", 2)
    item = _paragraph("2.补充材料", "body", 3)
    sign_org = _paragraph("星河治理委员会", "body", 4)
    sign_date = _paragraph("2026年7月20日", "body", 5)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        note,
        item,
        sign_org,
        sign_date,
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[2:]] == [
        "attachment_note", "attachment_note_item", "sign_org", "sign_date",
    ]
    note_candidates = data.recognition_diagnostics["candidate_trace"][2]["candidates"]
    attachment = next(item for item in note_candidates if item["type"] == "attachment_note")
    assert attachment["hard"] is True
    assert "attachment-tail-context" in attachment["evidence"]


def test_signature_org_uses_generic_tail_context_without_name_list() -> None:
    sign_org = _paragraph("星河治理委员会", "body", 2)
    sign_date = _paragraph("2026年7月20日", "body", 3)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        sign_org,
        sign_date,
    )

    apply_recognition(data)

    assert [sign_org.type_id, sign_date.type_id] == ["sign_org", "sign_date"]
    assert sign_org.text == "星河治理委员会"
    assert "signature-tail-context" in sign_org.meta["recognition_evidence"]


def test_body_unit_mention_is_not_signature_org_without_tail_context() -> None:
    paragraph = _paragraph("星河治理委员会持续推进相关工作，阶段性成效已经形成。", "body", 1)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        paragraph,
        _paragraph("后续正文继续说明工作安排。", "body", 2),
    )

    apply_recognition(data)

    assert paragraph.type_id == "body"


def test_duplicate_heading_sequence_is_applied_but_marked_for_review() -> None:
    duplicate = _paragraph("一、重复编号标题", "body", 4)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、持续夯实基层基础", "body", 2),
        _paragraph("持续完善责任体系，确保各项部署落实到位。", "body", 3),
        duplicate,
        _paragraph("后续正文对该标题展开具体说明，并列举工作安排、责任分工和落实要求，确保形成完整支撑。", "body", 5),
    )

    apply_recognition(data)

    assert duplicate.type_id == "heading1"
    diagnostic = data.recognition_diagnostics["paragraphs"][4]
    assert diagnostic["review_level"] == "review"
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]
    assert "numbering-duplicate" in diagnostic["heading_context_evidence"]


def test_child_heading_numbering_resets_under_new_parent_scope() -> None:
    first_child = _paragraph("（一）第一章下的子标题", "body", 3)
    second_child = _paragraph("（一）第二章下的子标题", "body", 6)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、第一部分", "body", 2),
        first_child,
        _paragraph("正文对第一章子标题展开说明。", "body", 4),
        _paragraph("二、第二部分", "body", 5),
        second_child,
        _paragraph("正文对第二章子标题展开说明。", "body", 7),
    )

    apply_recognition(data)

    assert [first_child.type_id, second_child.type_id] == ["heading2", "heading2"]
    second_diagnostic = data.recognition_diagnostics["paragraphs"][6]
    assert "numbering-duplicate" not in second_diagnostic["heading_context_evidence"]
    assert second_diagnostic["review_level"] not in {"review", "critical_review"}
    h2_families = [
        item for item in data.recognition_diagnostics["document_context"]["heading_families"]
        if item["level"] == 2
    ]
    assert [item["parent_scope"] for item in h2_families] == [[2], [5]]


def test_duplicate_child_heading_in_same_parent_scope_is_reviewed() -> None:
    duplicate = _paragraph("（一）同一父标题下重复的子标题", "body", 5)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、第一部分", "body", 2),
        _paragraph("（一）第一个子标题", "body", 3),
        _paragraph("正文对第一个子标题展开说明。", "body", 4),
        duplicate,
        _paragraph("正文对重复子标题展开说明。", "body", 6),
    )

    apply_recognition(data)

    assert duplicate.type_id == "heading2"
    diagnostic = data.recognition_diagnostics["paragraphs"][5]
    assert "numbering-duplicate" in diagnostic["heading_context_evidence"]
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]


def test_deep_heading_cannot_borrow_parent_from_previous_h1_scope() -> None:
    orphan = _paragraph("1.缺少当前父级的三级标题", "body", 7)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、第一部分", "body", 2),
        _paragraph("（一）第一章子标题", "body", 3),
        _paragraph("1.第一章三级标题", "body", 4),
        _paragraph("正文对第一章三级标题展开说明。", "body", 5),
        _paragraph("二、第二部分", "body", 6),
        orphan,
        _paragraph("正文对孤立三级标题展开说明。", "body", 8),
    )

    apply_recognition(data)

    diagnostic = data.recognition_diagnostics["paragraphs"][7]
    assert "missing-parent-heading" in diagnostic["heading_context_evidence"]
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]


def test_orphan_deep_heading_is_not_confirmed_without_parent_context() -> None:
    orphan = _paragraph("1.缺少父级的三级标题", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        orphan,
        _paragraph("后续正文对该段展开具体说明，并形成足够的正文支撑材料。", "body", 3),
    )

    apply_recognition(data)

    diagnostic = data.recognition_diagnostics["paragraphs"][2]
    assert "heading3" in diagnostic["candidate_types"]
    assert "missing-parent-heading" in diagnostic["heading_context_evidence"]
    assert diagnostic["review_level"] == "review"
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]


def test_short_bold_source_list_heading_survives_core_body_candidate() -> None:
    paragraph = _paragraph(
        "企业在订单班的重视度和投入上主动性不足",
        "heading2",
        2,
        numbering_prefix="@lvl_0",
        bold_char_ratio=1.0,
        classification_kind="body",
        classification_confidence=0.78,
    )
    data = _document(
        _paragraph("三、存在的主要问题", "heading1", 0),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
        _paragraph("后续正文对该标题展开具体说明。", "body", 3),
    )

    apply_recognition(data)

    assert paragraph.type_id == "heading2"
    assert paragraph.meta["recognition_provider"].startswith("source-list-numbering:")


def test_long_source_list_prose_is_not_promoted_to_heading() -> None:
    paragraph = _paragraph(
        "这是继承了Word列表属性的较长正文段落，其中包含完整叙述和多项具体情况，不应因为隐藏列表属性被识别成标题。",
        "body",
        1,
        numbering_prefix="@lvl_0",
        bold_char_ratio=1.0,
        classification_kind="body",
        classification_confidence=0.9,
    )
    data = _document(_paragraph("工作情况", "title", 0), paragraph)

    apply_recognition(data)

    assert paragraph.type_id == "body"


def test_title_formatting_cannot_override_front_role_and_placeholder_date():
    data = _document(
        _paragraph("在区政协第九届委员会第一次会议", "title", 0, alignment="CENTER"),
        _paragraph("召集人会议上的讲话", "title_cont", 1, alignment="CENTER"),
        _paragraph("区委副书记  XXX", "role_name", 2, alignment="CENTER"),
        _paragraph("2026年8月X日", "date_line", 3, alignment="CENTER"),
        _paragraph("同志们：", "addressing", 4),
        _paragraph("现将会议有关事项说明如下，确保各项工作衔接有序。", "body", 5),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:5]] == [
        "title", "title_cont", "role_name", "date_line", "addressing",
    ]
    assert data.recognition_diagnostics["paragraphs"][2]["review_level"] != "critical_review"
    assert data.recognition_diagnostics["paragraphs"][3]["review_level"] != "critical_review"


def test_front_role_metadata_vetoes_title_continuation_transition_bonus() -> None:
    role = _paragraph(
        "某区委员会党组书记、主席张测试",
        "role_name",
        1,
        alignment="CENTER",
        bold_char_ratio=1.0,
        legacy_type_id="role_name",
    )
    date = _paragraph(
        "（2026年8月27日11:00，某地区委员会会议中心）",
        "date_line",
        2,
        alignment="CENTER",
        bold_char_ratio=1.0,
        legacy_type_id="date_line",
    )
    data = _document(
        _paragraph("在某地区委员会会议闭幕大会上的讲话", "title", 0, alignment="CENTER"),
        role,
        date,
        _paragraph("各位代表、同志们：", "addressing", 3),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:5]] == [
        "title", "role_name", "date_line", "addressing", "body",
    ]
    assert role.meta["review_level"] != "critical_review"
    assert date.meta["review_level"] != "critical_review"
    assert "colon-explanatory-body" not in date.meta["recognition_evidence"]


def test_front_role_metadata_supports_generic_role_and_name_shapes() -> None:
    cases = (
        ("某办公室主任 张某", "2026年8月27日"),
        ("某委员会书记、主席　张某某", "2026年8月27日"),
        ("某区总工程师赵某某", "2026年8月27日11:00"),
        ("某工作组负责人  张某", "各位代表、同志们："),
    )

    for role_text, following_text in cases:
        following_type = "date_line" if "年" in following_text else "addressing"
        data = _document(
            _paragraph("年度工作报告", "title", 0, alignment="CENTER"),
            _paragraph(role_text, "body", 1, alignment="CENTER"),
            _paragraph(following_text, following_type, 2, alignment="CENTER"),
            _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
        )

        apply_recognition(data)

        assert data.paragraphs[1].type_id == "role_name", role_text
        assert data.paragraphs[1].meta["recognition_provider"].startswith("front-metadata:")


def test_role_word_inside_real_title_continuation_is_not_a_person_name() -> None:
    continuation = _paragraph(
        "主席会议制度建设情况",
        "title_cont",
        1,
        alignment="CENTER",
    )
    data = _document(
        _paragraph("年度工作报告", "title", 0, alignment="CENTER"),
        continuation,
        _paragraph("2026年8月27日", "date_line", 2, alignment="CENTER"),
        _paragraph("各有关单位：", "addressing", 3),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert continuation.type_id == "title_cont"
    assert data.recognition_diagnostics["document_context"]["front_metadata"] == [
        {"position": 2, "kind": "date_line"},
    ]


def test_front_metadata_shape_recognizes_unseen_role_and_placeholder_date_without_legacy_type():
    data = _document(
        _paragraph("在推进重点项目专题会议上的讲话", "body", 0, alignment="CENTER"),
        _paragraph("某区总工程师张某", "body", 1, alignment="CENTER"),
        _paragraph("2026年8月X日", "body", 2, alignment="CENTER"),
        _paragraph("同志们：", "body", 3),
        _paragraph("现将会议有关事项说明如下，确保各项工作衔接有序。", "body", 4),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:5]] == [
        "title", "role_name", "date_line", "addressing", "body",
    ]
    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert diagnostics[1]["provider"].startswith("front-metadata:")
    assert diagnostics[2]["provider"].startswith("front-metadata:")
    assert diagnostics[1]["review_level"] in {"confirmed", "info"}
    assert diagnostics[2]["review_level"] in {"confirmed", "info"}


def test_numbered_heading_family_confirms_parallel_body_headings():
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、持续夯实基层基础", "body", 2),
        _paragraph("持续完善责任体系，确保各项部署落实到位，并将责任细化到岗、任务落实到人、过程监督到位。", "body", 3),
        _paragraph("二、稳步提升服务质效", "body", 4),
        _paragraph("围绕群众需求优化流程，持续提升服务效率和群众办事体验，确保事项办理规范有序。", "body", 5),
    )

    apply_recognition(data)

    assert [data.paragraphs[index].type_id for index in (2, 4)] == ["heading1", "heading1"]
    context = data.recognition_diagnostics["document_context"]
    assert context["heading_families"] == [
        {"level": 1, "count": 2, "positions": [2, 4], "supported_count": 2},
    ]
    assert "parallel-heading-family" in data.recognition_diagnostics["paragraphs"][2]["heading_context_evidence"]


def test_single_numbered_heading_is_applied_but_marked_for_review():
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("一、持续夯实基层基础", "body", 1),
        _paragraph("持续完善责任体系，确保各项部署落实到位，并将责任细化到岗、任务落实到人、过程监督到位。", "body", 2),
    )

    apply_recognition(data)

    heading = data.recognition_diagnostics["paragraphs"][1]
    assert data.paragraphs[1].type_id == "heading1"
    assert heading["review_level"] == "review"
    assert "LEGACY_TYPE_CONFLICT" in heading["review_reasons"]
    assert data.paragraphs[1].original_text == "一、持续夯实基层基础"


def test_table_boundary_blocks_title_continuation():
    data = _document(
        _paragraph("主标题", "title", 0),
        _paragraph("", "__table__", 1),
        _paragraph("表后说明", "body", 2, alignment="CENTER"),
    )

    apply_recognition(data)

    assert data.recognition_diagnostics["candidate_trace"][1]["boundary_before"] is True
    assert data.paragraphs[2].meta["recognition_type"] != "title_continuation"


def test_wrong_legacy_and_docxtool_style_do_not_override_dispatch():
    paragraph = _paragraph("国发〔2026〕23号", "title", 0, style_name="DCT-Title", legacy_type_id="title")
    data = _document(paragraph)

    apply_recognition(data)

    assert paragraph.type_id == "dispatch_number"
    assert paragraph.meta["legacy_type_id"]["value"] == "title"


def test_wrong_heading_legacy_does_not_override_meeting_metadata():
    paragraph = _paragraph("（一）缺席：李四", "heading2", 1, style_name="DCT-Heading2", legacy_type_id="heading2")
    data = _document(_paragraph("党委会会议纪要", "title", 0), paragraph)

    apply_recognition(data)

    assert paragraph.type_id == "meeting_meta"


def test_every_paragraph_type_has_explicit_render_mapping():
    for paragraph_type in ParagraphType:
        type_id, rule_index, style_id = resolve_render_mapping(paragraph_type)
        if paragraph_type is ParagraphType.UNKNOWN:
            assert type_id is None and rule_index is None and style_id is None
        else:
            assert type_id
            assert isinstance(rule_index, int)
            assert style_id and style_id.startswith("DCT-")


def test_real_docx_round_trip_preserves_semantics_and_layout(tmp_path):
    source = tmp_path / "source.docx"
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    document = Document()
    document.add_paragraph("国务院关于印发规划的通知")
    document.add_paragraph("国发〔2026〕23号")
    document.add_paragraph("各有关单位：")
    document.add_paragraph("正文内容。")
    document.add_paragraph("国务院")
    document.add_paragraph("2026年3月18日")
    document.add_paragraph("（本文有删减）")
    document.add_paragraph("公共服务提升规划")
    document.add_paragraph("第一章 总则")
    document.save(source)

    importer = __import__("docxtool.document.importer", fromlist=["DocxImporter"]).DocxImporter()
    first_data = importer.load(str(source), _rules())
    export_doc(first_data, _rules(), PageSettings(), str(first))
    second_data = importer.load(str(first), _rules())
    export_doc(second_data, _rules(), PageSettings(), str(second))
    third_data = importer.load(str(second), _rules())

    def signature(data):
        return [(item.type_id, item.text, item.meta.get("recognition_type"), item.meta.get("recognition_section")) for item in data.paragraphs]

    assert signature(second_data) == signature(third_data)
    assert [item.type_id for item in second_data.paragraphs].count("dispatch_number") == 1
    assert [item.type_id for item in second_data.paragraphs].count("embedded_document_title") == 1
    assert len(second_data.paragraphs) == len(third_data.paragraphs)
    assert second_data.recognition_diagnostics["validation"]["ok"] is True


def test_recognition_config_rejects_invalid_values():
    for kwargs in (
        {"beam_width": 1},
        {"max_candidates_per_paragraph": 1},
        {"legacy_score": 1.1},
        {"text_preview_length": -1},
        {"unknown_render_type": "silent"},
    ):
        try:
            RecognitionConfig(**kwargs)
        except ValueError:
            continue
        raise AssertionError(f"invalid config accepted: {kwargs}")


def test_disabling_diagnostics_does_not_change_decisions():
    enabled = _document(_paragraph("主标题", "title", 0), _paragraph("国发〔2026〕23号", "title_cont", 1))
    disabled = _document(_paragraph("主标题", "title", 0), _paragraph("国发〔2026〕23号", "title_cont", 1))

    apply_recognition(enabled, RecognitionConfig(enable_diagnostics=True))
    apply_recognition(disabled, RecognitionConfig(enable_diagnostics=False))

    assert [item.type_id for item in enabled.paragraphs] == [item.type_id for item in disabled.paragraphs]
    assert enabled.recognition_diagnostics["candidate_trace"]
    assert disabled.recognition_diagnostics["candidate_trace"] == []
    assert disabled.recognition_diagnostics["engine_version"] == "3.0"
    assert disabled.recognition_diagnostics["schema_version"] == "1.0"


def test_empty_and_table_only_documents_do_not_create_fake_paragraphs():
    empty = _document()
    table_only = _document(_paragraph("", "__table__", 0))

    apply_recognition(empty)
    apply_recognition(table_only)

    assert empty.recognition_diagnostics["paragraphs"] == []
    assert table_only.recognition_diagnostics["paragraphs"] == []
    assert table_only.recognition_diagnostics["blocks"][0]["kind"] == "table"


def test_long_and_unusual_unicode_text_is_bounded_and_preserved():
    raw = "Ａ：" + "甲\u00a0\u200b" * 25000
    paragraph = _paragraph(raw, "body", 0)
    data = _document(paragraph)

    apply_recognition(data, RecognitionConfig(text_preview_length=10))

    assert paragraph.original_text == raw
    assert len(data.recognition_diagnostics["paragraphs"][0]["text_preview"]) == 10
    assert len(data.recognition_diagnostics["candidate_trace"]) == 1


def test_incomplete_heading_level_is_diagnosed_without_text_loss():
    paragraph = _paragraph("（三）直接出现", "heading2", 0)
    data = _document(paragraph)

    apply_recognition(data)

    assert paragraph.original_text == "（三）直接出现"
    assert data.recognition_diagnostics["paragraphs"][0]["candidate_count"] >= 1


def test_dispatch_variants_are_stable_after_nfkc():
    for index, text in enumerate(("国发〔2026〕23号", "市府[2027]1号", "ＡＢ〔２０２８〕１２３号", "国发〔 2026 〕 23 号")):
        paragraph = _paragraph(text, "title_cont", index)
        data = _document(paragraph)
        apply_recognition(data)
        assert paragraph.type_id == "dispatch_number", text


def test_front_matter_scan_skips_empty_and_caption_placeholders() -> None:
    prefix = [_paragraph("", "body", index) for index in range(14)]
    caption = _paragraph("图1 工作流程", "__object_caption__", 14)
    title = _paragraph("关于推进基层治理工作的通知", "body", 15, alignment="CENTER")
    recipient = _paragraph("各有关单位：", "body", 16)
    body = _paragraph("现将有关事项通知如下，请结合实际抓好落实。", "body", 17)
    data = _document(*prefix, caption, title, recipient, body)

    apply_recognition(data)

    assert title.type_id == "title"
    assert recipient.type_id == "addressing"
    context = data.recognition_diagnostics["document_context"]
    assert 15 in context["front_matter_positions"]
    assert context["front_scan_reason"] == "body-boundary"
    assert context["body_start"] == 17


def test_front_scan_uses_soft_threshold_until_real_structure_boundary() -> None:
    prefix = [
        _paragraph(f"联合发文机关{index}", "body", index)
        for index in range(13)
    ]
    title = _paragraph("关于推进基层治理工作的通知", "body", 13, alignment="CENTER")
    recipient = _paragraph("各有关单位：", "body", 14)
    body = _paragraph("现将有关事项通知如下，请结合实际抓好落实。", "body", 15)
    data = _document(*prefix, title, recipient, body)

    apply_recognition(data)

    assert title.type_id == "title"
    assert recipient.type_id == "addressing"
    context = data.recognition_diagnostics["document_context"]
    assert 13 in context["front_matter_positions"]
    assert context["front_scan_reason"] == "body-boundary"
    assert context["front_soft_threshold_exceeded"] is True
    assert context["front_scan_soft_threshold"] == 12
    assert context["body_start"] == 15


def test_key_value_and_source_variants_do_not_promote_numbering():
    for index, text in enumerate(("（一）缺席：李四", "（二）出 席:张三", "来源：国家卫生健康委员会")):
        paragraph = _paragraph(text, "heading2", index)
        data = _document(_paragraph("党委会会议纪要", "title", 0), paragraph)
        apply_recognition(data)
        if text.startswith("来源"):
            assert paragraph.type_id == "note"
        else:
            assert paragraph.type_id == "meeting_meta"


def test_review_flags_and_safe_summary_do_not_change_final_types():
    clear = _document(_paragraph("国发〔2026〕23号", "title_cont", 0))
    ambiguous = _document(_paragraph("补充说明", "body", 0, alignment="CENTER"))

    apply_recognition(clear)
    apply_recognition(ambiguous, RecognitionConfig(review_low_score=0.9))

    clear_diagnostic = clear.recognition_diagnostics["paragraphs"][0]
    ambiguous_diagnostic = ambiguous.recognition_diagnostics["paragraphs"][0]
    assert clear_diagnostic["needs_review"] is False
    assert clear_diagnostic["review_level"] == "info"
    assert "STRUCTURE_CONFIRMED_RECLASSIFICATION" in clear_diagnostic["review_reasons"]
    assert ambiguous_diagnostic["needs_review"] is True
    assert ambiguous_diagnostic["review_reasons"]
    assert clear.paragraphs[0].type_id == "dispatch_number"
    summary = ambiguous.recognition_diagnostics["summary"]
    assert summary["needs_review_count"] == 1
    assert "补充说明" not in diagnostics_to_json(ambiguous.recognition_diagnostics)


def test_explicit_numbering_is_confirmed_even_when_raw_candidate_softmax_is_low():
    data = _document(
        _paragraph("一、工作开展情况", "heading1", 0),
        _paragraph("（一）落实重点任务", "heading2", 1),
        _paragraph("1. 压实责任", "heading3", 2),
        _paragraph("（1）明确时限", "heading4", 3),
    )

    apply_recognition(data)

    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert [item["final_type"] for item in diagnostics] == ["heading1", "heading2", "heading3", "heading4"]
    assert all(item["review_level"] == "confirmed" for item in diagnostics)
    assert not any(item["needs_review"] for item in diagnostics)
    assert all("explicit-numbering" in item["evidence_summary"] for item in diagnostics)


def test_same_input_is_thread_safe_across_twenty_independent_documents():
    def recognize(_):
        data = _document(
            _paragraph("2026年第一次党委会会议纪要", "title", 0),
            _paragraph("出席：甲、乙", "body", 1),
            _paragraph("（一）缺席：无", "heading2", 2),
        )
        apply_recognition(data)
        return (
            data.doc_mode,
            tuple((item.type_id, item.meta["recognition_section"]) for item in data.paragraphs),
            data.recognition_diagnostics["summary"],
        )

    with ThreadPoolExecutor(max_workers=8) as executor:
        results = list(executor.map(recognize, range(20)))

    assert len(results) == 20
    assert all(result == results[0] for result in results)


def test_diagnostics_use_candidates_from_final_winning_beam(monkeypatch) -> None:
    class BranchingProvider:
        name = "branching"

        def propose(self, block, features, context):
            if context.index == 0:
                return [
                    Candidate(ParagraphType.BODY, 0.60, self.name, ("start-body",), section_hint=SectionKind.BODY),
                    Candidate(ParagraphType.ADDRESSING, 0.55, self.name, ("start-addressing",), section_hint=SectionKind.RECIPIENT),
                ]
            if context.index == 1 and context.previous_type == ParagraphType.ADDRESSING:
                return [
                    Candidate(ParagraphType.HEADING_1, 1.00, self.name, ("after-addressing",), section_hint=SectionKind.BODY),
                ]
            if context.index == 1:
                return [
                    Candidate(ParagraphType.BODY, 0.10, self.name, ("after-body",), section_hint=SectionKind.BODY),
                ]
            return [
                Candidate(ParagraphType.BODY, 0.10, self.name, ("tail",), section_hint=SectionKind.BODY),
            ]

    monkeypatch.setattr(decoder, "DEFAULT_PROVIDERS", (BranchingProvider(),))
    data = _document(
        _paragraph("起始行", "body", 0),
        _paragraph("路径相关行", "body", 1),
    )

    apply_recognition(data, RecognitionConfig(beam_width=2, max_candidates_per_paragraph=4))

    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert [item["final_type"] for item in diagnostics] == ["addressing", "heading1"]
    assert diagnostics[1]["candidate_types"] == ["heading1"]
    assert diagnostics[1]["provider"].startswith("branching:after-addressing")
    assert diagnostics[1]["selected_candidate_score"] == 1.0


@pytest.mark.parametrize(
    "title_text",
    [
        "代表委员履职报告",
        "项目负责人 工作方案",
        "客户经理年度总结",
        "优秀同志 调研报告",
    ],
)
def test_front_title_with_broad_role_word_is_not_role_name(title_text):
    data = _document(
        _paragraph(title_text, "title", 0, alignment="CENTER"),
        _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
        _paragraph("各位代表、同志们：", "addressing", 2),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
    )

    apply_recognition(data)

    assert data.paragraphs[0].type_id == "title"
    context = data.recognition_diagnostics["document_context"]
    assert not any(item["position"] == 0 for item in context["front_metadata"])
    trace = data.recognition_diagnostics["candidate_trace"][0]
    assert "role_name" not in [item["type"] for item in trace["candidates"]]


@pytest.mark.parametrize(
    "value",
    [
        "人大代表履职情况 专题调研",
        "项目负责人履职情况 工作部署",
        "委员联系群众情况 调研分析",
        "主席会议制度建设 经验交流",
    ],
)
@pytest.mark.parametrize(
    "context_kind",
    ["previous_title", "following_date", "following_addressing", "heading_style", "centered"],
)
def test_front_role_words_not_adjacent_to_name_do_not_form_role_name(value, context_kind):
    candidate = _paragraph(value, "body", 1)
    paragraphs = []
    if context_kind == "previous_title":
        paragraphs.append(_paragraph("年度专题材料", "title", 0, alignment="CENTER"))
        candidate.features.alignment = "CENTER"
        paragraphs.extend([candidate, _paragraph("正文内容已经开始。", "body", 2)])
    elif context_kind == "following_date":
        candidate.features.alignment = "CENTER"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 2),
        ])
    elif context_kind == "following_addressing":
        candidate.features.alignment = "CENTER"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("各位代表、同志们：", "addressing", 1),
            _paragraph("正文内容已经开始。", "body", 2),
        ])
    elif context_kind == "heading_style":
        candidate.features.style_name = "Heading 1"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 2),
        ])
    else:
        candidate.features.alignment = "CENTER"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 2),
        ])

    data = _document(*paragraphs)
    apply_recognition(data)

    candidate_index = data.paragraphs.index(candidate)
    assert candidate.type_id != "role_name"
    context = data.recognition_diagnostics["document_context"]
    assert not any(
        item["position"] == candidate_index and item["kind"] == "role_name"
        for item in context["front_metadata"]
    )
    trace = data.recognition_diagnostics["candidate_trace"][candidate_index]
    assert "role_name" not in [item["type"] for item in trace["candidates"]]


@pytest.mark.parametrize(
    "role_text",
    [
        "办公室主任 张三",
        "办公室主任 李测试",
        "办公室主任 欧阳测试",
        "办公室主任 阿·明",
        "办公室主任　李测试",
        "党组书记、主席 欧阳测试",
        "某某机构办公室主任 张三",
        "办公室主任王测试",
        "会议代表 张三",
        "项目负责人 李测试",
    ],
)
def test_front_role_name_shape_survives_changed_names_and_roles(role_text):
    data = _document(
        _paragraph("年度重点工作会议讲话", "title", 0, alignment="CENTER"),
        _paragraph(role_text, "body", 1, alignment="CENTER"),
        _paragraph("2026年8月27日", "date_line", 2, alignment="CENTER"),
        _paragraph("各位代表、同志们：", "addressing", 3),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert data.paragraphs[1].type_id == "role_name"
    context = data.recognition_diagnostics["document_context"]
    assert {item["position"]: item["kind"] for item in context["front_metadata"]}[1] == "role_name"
    trace = data.recognition_diagnostics["candidate_trace"][1]
    assert "role_name" in [item["type"] for item in trace["candidates"]]


def test_spaced_role_name_can_use_previous_title_anchor_without_date():
    data = _document(
        _paragraph("年度重点工作会议讲话", "title", 0, alignment="CENTER"),
        _paragraph("办公室主任 李测试", "body", 1, alignment="CENTER"),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 2),
    )

    apply_recognition(data)

    assert data.paragraphs[1].type_id == "role_name"
