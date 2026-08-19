from __future__ import annotations

import base64
import io
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm

from docxtool.document.engine.page_number import apply_page_number, apply_page_numbers
from docxtool.security.docx_integrity import validate_docx_integrity

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def _add_field(paragraph, instruction: str, display_text: str = "1") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    paragraph.add_run()._r.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    paragraph.add_run(display_text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def _add_simple_field(paragraph, instruction: str, display_text: str = "1") -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = display_text
    run.append(text)
    field.append(run)
    paragraph._element.append(field)


def _add_split_field(
    paragraph, instruction_parts: tuple[str, ...], display_text: str = "1"
) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)
    for part in instruction_parts:
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = part
        paragraph.add_run()._r.append(instruction)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    paragraph.add_run(display_text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    visible = OxmlElement("w:t")
    visible.text = text
    run.append(visible)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def _footer_xml(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {
            name: archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        }


def _document_xml(path: Path) -> ET.Element:
    with zipfile.ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/document.xml"))


def _footer_roots(path: Path) -> dict[str, ET.Element]:
    with zipfile.ZipFile(path) as archive:
        return {
            name: ET.fromstring(archive.read(name))
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        }


def _field_instructions(root: ET.Element) -> list[str]:
    return ["".join(element.itertext()).strip() for element in root.findall(f".//{{{W_NS}}}instrText")]


def _field_char_types(root: ET.Element) -> list[str]:
    return [
        element.get(qn("w:fldCharType"), "")
        for element in root.findall(f".//{{{W_NS}}}fldChar")
    ]


def _assert_complex_fields_are_paired(root: ET.Element) -> None:
    open_fields = 0
    waiting_for_end = 0
    for field_type in _field_char_types(root):
        if field_type == "begin":
            open_fields += 1
        elif field_type == "separate":
            assert open_fields > 0
            waiting_for_end += 1
        elif field_type == "end":
            assert open_fields > 0
            assert waiting_for_end > 0
            open_fields -= 1
            waiting_for_end -= 1
    assert open_fields == 0
    assert waiting_for_end == 0


def _visible_text(root: ET.Element) -> str:
    return "".join(element.text or "" for element in root.findall(f".//{{{W_NS}}}t"))


def _page_paragraph(root: ET.Element) -> ET.Element:
    return next(
        paragraph
        for paragraph in root.findall(f".//{{{W_NS}}}p")
        if "PAGE" in " ".join(_field_instructions(paragraph))
    )


def _page_paragraphs(root: ET.Element) -> list[ET.Element]:
    return [
        paragraph
        for paragraph in root.findall(f".//{{{W_NS}}}p")
        if "PAGE" in "".join(_field_instructions(paragraph))
    ]


def _assert_native_dash_paragraph(paragraph: ET.Element) -> None:
    assert _field_instructions(paragraph) == ["PAGE"]
    texts = [
        element.text or ""
        for element in paragraph.findall(f".//{{{W_NS}}}t")
    ]
    assert texts == []
    assert paragraph.find(".//" + qn("w:pBdr")) is None
    assert paragraph.find(".//" + qn("w:tabs")) is None
    assert paragraph.find(".//" + qn("w:u")) is None
    assert paragraph.find(".//" + qn("w:bdr")) is None
    assert paragraph.find(".//" + qn("w:shd")) is None


def _paragraph_alignment_and_indent(paragraph: ET.Element) -> tuple[str, ET.Element | None]:
    properties = paragraph.find(qn("w:pPr"))
    alignment = properties.find(qn("w:jc")) if properties is not None else None
    indent = properties.find(qn("w:ind")) if properties is not None else None
    return (alignment.get(qn("w:val"), "") if alignment is not None else "", indent)


def test_page_number_fields_outside_position_and_first_page_hidden(tmp_path: Path) -> None:
    document = Document()
    first = document.sections[0]
    first.bottom_margin = Cm(3.5)
    first.footer.is_linked_to_previous = False
    first.footer.paragraphs[0].text = "Confidential footer"
    old_page = first.footer.add_paragraph("old ")
    _add_field(old_page, "PAGE")
    first.even_page_footer.is_linked_to_previous = False
    first.even_page_footer.paragraphs[0].text = "Even footer note"
    document.add_paragraph("section one")
    second = document.add_section(WD_SECTION.NEW_PAGE)
    second.footer.is_linked_to_previous = False
    second.footer.paragraphs[0].text = "Section two footer note"
    document.add_paragraph("section two")
    first = document.sections[0]
    first.different_first_page_header_footer = True
    first.first_page_footer.is_linked_to_previous = False
    first.first_page_footer.paragraphs[0].text = "First page legal notice"

    apply_page_number(
        document,
        {
            "style": "cn_total",
            "position": "outside",
            "first_page": False,
            "section_starts": [1, None],
            "offset_from_text_mm": 7,
            "font_name": "SimSun",
            "font_size_pt": 10.5,
        },
    )
    output = tmp_path / "page-numbered.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    footers = _footer_xml(output)
    combined_footer_xml = "\n".join(footers.values())
    assert "Confidential footer" in combined_footer_xml
    assert "Section two footer note" in combined_footer_xml
    assert "First page legal notice" in combined_footer_xml
    assert "PAGE" in combined_footer_xml
    assert "NUMPAGES" in combined_footer_xml
    assert "<w:t>PAGE</w:t>" not in combined_footer_xml
    assert "<w:t>NUMPAGES</w:t>" not in combined_footer_xml
    assert "<w:t>1</w:t>" not in combined_footer_xml
    assert "AlternateContent" not in combined_footer_xml
    assert "txbxContent" not in combined_footer_xml
    assert "textbox" not in combined_footer_xml
    first_page_footer_xml = next(xml for xml in footers.values() if "First page legal notice" in xml)
    assert "PAGE" not in first_page_footer_xml
    assert "NUMPAGES" not in first_page_footer_xml
    for root in _footer_roots(output).values():
        instructions = _field_instructions(root)
        assert instructions.count("PAGE") <= 1
        assert instructions.count("NUMPAGES") <= 1
        _assert_complex_fields_are_paired(root)
    assert any('w:val="right"' in xml for xml in footers.values())
    assert any('w:val="left"' in xml for xml in footers.values())
    with zipfile.ZipFile(output) as archive:
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
    assert "evenAndOddHeaders" in settings_xml
    assert "updateFields" in settings_xml
    assert abs(first.footer_distance.cm - 2.8) < 0.02
    sections = _document_xml(output).findall(".//" + qn("w:sectPr"))
    assert any(section.find(qn("w:titlePg")) is not None for section in sections)
    starts = [
        start
        for section in sections
        if (pg_num_type := section.find(qn("w:pgNumType"))) is not None
        if (start := pg_num_type.get(qn("w:start"))) is not None
    ]
    assert starts == ["1"]


def test_page_number_styles_and_section_restart_policy(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("first")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("second")

    apply_page_numbers(document, {"style": "dash", "position": "center", "section_numbering": "restart"})
    output = tmp_path / "restart.docx"
    document.save(output)

    validate_docx_integrity(output)
    footer_xml = "\n".join(_footer_xml(output).values())
    assert "—" not in footer_xml
    assert footer_xml.count("PAGE") == 2
    assert "NUMPAGES" not in footer_xml
    assert 'w:val="center"' in footer_xml
    sections = _document_xml(output).findall(".//" + qn("w:sectPr"))
    starts = [section.find(qn("w:pgNumType")).get(qn("w:start")) for section in sections]
    assert starts == ["1", "1"]
    assert all(
        section.find(qn("w:pgNumType")).get(qn("w:fmt")) == "numberInDash"
        for section in sections
    )


def test_outside_page_numbers_reuse_empty_even_footer_paragraph(tmp_path: Path) -> None:
    document = Document()

    apply_page_numbers(document, {"style": "dash", "position": "outside"})
    output = tmp_path / "outside-page-number.docx"
    document.save(output)

    roots = _footer_roots(output)
    assert len(roots) == 2
    for root in roots.values():
        paragraphs = root.findall(f".//{{{W_NS}}}p")
        assert len(paragraphs) == 1
        assert _field_instructions(root) == ["PAGE"]


def test_page_number_formats_create_standard_fields(tmp_path: Path) -> None:
    cases = [
        ("dash", "", "", ["PAGE"]),
        ("plain", "", "", ["PAGE"]),
        ("cn", "第 ", " 页", ["PAGE"]),
        ("cn_total", "第 ", " 页 共 ", ["PAGE", "NUMPAGES"]),
    ]

    for style, prefix, suffix, instructions in cases:
        document = Document()
        document.add_paragraph(style)
        apply_page_number(document, {"style": style, "position": "center"})
        output = tmp_path / f"{style}.docx"
        document.save(output)

        assert validate_docx_integrity(output).ok is True
        roots = _footer_roots(output)
        assert len(roots) == 1
        root = next(iter(roots.values()))
        assert _field_instructions(root) == instructions
        _assert_complex_fields_are_paired(root)
        text = _visible_text(root)
        assert prefix in text
        assert suffix in text
        assert "1" not in text
        pg_num_type = document.sections[0]._sectPr.find(qn("w:pgNumType"))
        expected_format = "numberInDash" if style == "dash" else "decimal"
        assert pg_num_type.get(qn("w:fmt")) == expected_format
        if style == "dash":
            _assert_native_dash_paragraph(_page_paragraph(root))
        footer_xml = next(iter(_footer_xml(output).values()))
        assert "AlternateContent" not in footer_xml
        assert "txbxContent" not in footer_xml
        assert "textbox" not in footer_xml


def test_decimal_source_is_normalized_to_native_dash_without_footer_dashes(
    tmp_path: Path,
) -> None:
    document = Document()
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), "decimal")
    document.sections[0]._sectPr.append(pg_num_type)
    _add_field(document.sections[0].footer.paragraphs[0], "PAGE", "7")

    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "decimal-to-native-dash.docx"
    document.save(output)

    section = _document_xml(output).find(".//" + qn("w:sectPr"))
    assert section.find(qn("w:pgNumType")).get(qn("w:fmt")) == "numberInDash"
    root = next(iter(_footer_roots(output).values()))
    _assert_native_dash_paragraph(_page_paragraph(root))
    assert "—" not in _visible_text(root)
    assert "-" not in _visible_text(root)


@pytest.mark.parametrize(
    ("style", "expected_instructions", "expected_text"),
    [
        ("plain", ["PAGE"], ""),
        ("cn", ["PAGE"], "第  页"),
        ("cn_total", ["PAGE", "NUMPAGES"], "第  页 共  页"),
    ],
)
def test_non_dash_styles_clear_historical_number_in_dash(
    tmp_path: Path,
    style: str,
    expected_instructions: list[str],
    expected_text: str,
) -> None:
    document = Document()
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), "numberInDash")
    document.sections[0]._sectPr.append(pg_num_type)

    apply_page_number(document, {"style": style, "position": "center"})
    output = tmp_path / f"number-in-dash-to-{style}.docx"
    document.save(output)

    section = _document_xml(output).find(".//" + qn("w:sectPr"))
    assert section.find(qn("w:pgNumType")).get(qn("w:fmt")) == "decimal"
    root = next(iter(_footer_roots(output).values()))
    assert _field_instructions(root) == expected_instructions
    assert _visible_text(root) == expected_text


def test_native_dash_normalizes_all_sections_without_losing_start_values(
    tmp_path: Path,
) -> None:
    document = Document()
    document.add_paragraph("section one")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("section two")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("section three")
    formats = ("decimal", "numberInDash", None)
    for section, page_format in zip(document.sections, formats):
        if page_format is None:
            continue
        pg_num_type = OxmlElement("w:pgNumType")
        pg_num_type.set(qn("w:fmt"), page_format)
        section._sectPr.append(pg_num_type)

    apply_page_number(
        document,
        {
            "style": "dash",
            "position": "center",
            "section_starts": [1, None, 5],
        },
    )
    output = tmp_path / "multi-section-native-dash.docx"
    document.save(output)

    sections = _document_xml(output).findall(".//" + qn("w:sectPr"))
    assert len(sections) == 3
    page_number_types = [section.find(qn("w:pgNumType")) for section in sections]
    assert all(item.get(qn("w:fmt")) == "numberInDash" for item in page_number_types)
    assert [item.get(qn("w:start")) for item in page_number_types] == ["1", None, "5"]
    for root in _footer_roots(output).values():
        _assert_native_dash_paragraph(_page_paragraph(root))


def test_page_number_preserves_non_page_footer_content_when_replacing_old_field(tmp_path: Path) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = "Confidential"
    paragraph.add_run(" — ")
    _add_field(paragraph, "PAGE")
    paragraph.add_run(" — ")
    footer.add_paragraph("Prepared by office")
    document.add_paragraph("body")

    apply_page_number(document, {"style": "plain", "position": "center"})
    output = tmp_path / "preserve-footer.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    root = next(iter(_footer_roots(output).values()))
    assert _field_instructions(root) == ["PAGE"]
    assert "Confidential" in _visible_text(root)
    assert "Prepared by office" in _visible_text(root)
    assert "Confidential —" not in _visible_text(root)
    _assert_complex_fields_are_paired(root)


def test_dirty_duplicate_dash_page_paragraph_is_canonically_rebuilt(tmp_path: Path) -> None:
    document = Document()
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), "numberInDash")
    document.sections[0]._sectPr.append(pg_num_type)
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.add_run("— ")
    paragraph.add_run("— ")
    _add_field(paragraph, "PAGE", "7")
    paragraph.add_run(" —")

    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "dirty-duplicate-dash.docx"
    document.save(output)

    root = next(iter(_footer_roots(output).values()))
    page_paragraphs = _page_paragraphs(root)
    assert len(page_paragraphs) == 1
    _assert_native_dash_paragraph(page_paragraphs[0])
    section = _document_xml(output).find(".//" + qn("w:sectPr"))
    assert section.find(qn("w:pgNumType")).get(qn("w:fmt")) == "numberInDash"


def test_mixed_business_run_drops_attached_old_dash_and_writes_separate_page_paragraph(
    tmp_path: Path,
) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.add_run("Confidential — ")
    _add_field(paragraph, "PAGE", "7")
    paragraph.add_run(" —")

    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "mixed-business-page.docx"
    document.save(output)

    root = next(iter(_footer_roots(output).values()))
    page_paragraphs = _page_paragraphs(root)
    assert len(page_paragraphs) == 1
    _assert_native_dash_paragraph(page_paragraphs[0])
    section = _document_xml(output).find(".//" + qn("w:sectPr"))
    page_number_types = section.findall(qn("w:pgNumType"))
    assert len(page_number_types) == 1
    assert page_number_types[0].get(qn("w:fmt")) == "numberInDash"
    assert "Confidential" in _visible_text(root)
    assert "Confidential —" not in _visible_text(root)
    assert page_paragraphs[0] not in [
        paragraph
        for paragraph in root.findall(f".//{{{W_NS}}}p")
        if "Confidential" in _visible_text(paragraph)
    ]


def test_polluted_empty_footer_paragraph_is_not_reused_for_canonical_page_number(
    tmp_path: Path,
) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    polluted = footer.paragraphs[0]
    properties = polluted._element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    borders.append(bottom)
    properties.append(borders)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "hyphen")
    tabs.append(tab)
    properties.append(tabs)
    _add_field(footer.add_paragraph(), "PAGE", "7")

    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "polluted-empty-footer.docx"
    document.save(output)

    root = next(iter(_footer_roots(output).values()))
    page_paragraphs = _page_paragraphs(root)
    assert len(page_paragraphs) == 1
    _assert_native_dash_paragraph(page_paragraphs[0])


def test_dash_page_number_triple_apply_is_structurally_idempotent(tmp_path: Path) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    footer.paragraphs[0].text = "Business footer"
    _add_field(footer.add_paragraph("— "), "PAGE", "7")

    paragraph_counts: list[int] = []
    for _ in range(3):
        apply_page_number(document, {"style": "dash", "position": "center"})
        paragraph_counts.append(len(footer.paragraphs))

    output = tmp_path / "triple-apply.docx"
    document.save(output)
    root = next(iter(_footer_roots(output).values()))

    assert paragraph_counts[0] == paragraph_counts[1] == paragraph_counts[2]
    page_paragraphs = _page_paragraphs(root)
    assert len(page_paragraphs) == 1
    _assert_native_dash_paragraph(page_paragraphs[0])
    section = _document_xml(output).find(".//" + qn("w:sectPr"))
    page_number_types = section.findall(qn("w:pgNumType"))
    assert len(page_number_types) == 1
    assert page_number_types[0].get(qn("w:fmt")) == "numberInDash"


def test_page_number_preserves_complex_and_simple_non_page_fields(tmp_path: Path) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    date_paragraph = footer.paragraphs[0]
    date_paragraph.add_run("Prepared by office ")
    _add_field(date_paragraph, "DATE \\@ yyyy-MM-dd")
    style_paragraph = footer.add_paragraph("Section title: ")
    _add_field(style_paragraph, 'STYLEREF "Heading 1"')
    ref_paragraph = footer.add_paragraph("Reference: ")
    _add_simple_field(ref_paragraph, "REF bookmark1", "Referenced text")
    old_page = footer.add_paragraph("— ")
    _add_simple_field(old_page, "PAGE", "7")
    old_page.add_run(" —")

    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "preserve-non-page-fields.docx"
    document.save(output)

    root = next(iter(_footer_roots(output).values()))
    instructions = _field_instructions(root)
    simple_instructions = [
        field.get(qn("w:instr"), "").strip()
        for field in root.findall(f".//{{{W_NS}}}fldSimple")
    ]
    assert any(instruction.startswith("DATE") for instruction in instructions)
    assert any(instruction.startswith("STYLEREF") for instruction in instructions)
    assert "REF bookmark1" in simple_instructions
    assert simple_instructions.count("PAGE") == 0
    assert instructions.count("PAGE") == 1
    assert "Prepared by office" in _visible_text(root)
    assert "Referenced text" in _visible_text(root)


def test_pure_date_field_is_not_mistaken_for_a_page_number(tmp_path: Path) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    date_paragraph = footer.paragraphs[0]
    _add_field(date_paragraph, "DATE \\@ yyyy-MM-dd", "2026-08-19")

    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "pure-date-field.docx"
    document.save(output)

    root = next(iter(_footer_roots(output).values()))
    assert any(item.startswith("DATE") for item in _field_instructions(root))
    assert _field_instructions(root).count("PAGE") == 1
    assert "2026-08-19" in _visible_text(root)


def test_linked_footer_materializes_inherited_content_before_page_replacement(tmp_path: Path) -> None:
    document = Document()
    first = document.sections[0]
    first.footer.is_linked_to_previous = False
    first.footer.paragraphs[0].text = "Confidential footer"
    detail = first.footer.add_paragraph("Prepared by office ")
    _add_field(detail, "DATE \\@ yyyy-MM-dd")
    old_page = first.footer.add_paragraph("— ")
    _add_field(old_page, "PAGE")
    old_page.add_run(" —")
    document.add_paragraph("section one")
    second = document.add_section(WD_SECTION.NEW_PAGE)
    assert second.footer.is_linked_to_previous is True
    document.add_paragraph("section two")

    apply_page_number(document, {"style": "dash", "position": "center"})
    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "linked-footer.docx"
    document.save(output)

    reopened = Document(output)
    assert reopened.sections[1].footer.is_linked_to_previous is False
    second_root = reopened.sections[1].footer._element
    assert "Confidential footer" in _visible_text(second_root)
    assert "Prepared by office" in _visible_text(second_root)
    assert any(item.startswith("DATE") for item in _field_instructions(second_root))
    assert _field_instructions(second_root).count("PAGE") == 1
    _assert_native_dash_paragraph(_page_paragraph(second_root))


def test_linked_footer_preserves_image_relationship_when_materialized(tmp_path: Path) -> None:
    document = Document()
    first = document.sections[0]
    first.footer.is_linked_to_previous = False
    first.footer.paragraphs[0].text = "Confidential footer"
    image_paragraph = first.footer.add_paragraph()
    image_paragraph.add_run().add_picture(io.BytesIO(_TINY_PNG))
    old_page = first.footer.add_paragraph()
    _add_field(old_page, "PAGE")
    document.add_paragraph("section one")
    second = document.add_section(WD_SECTION.NEW_PAGE)
    assert second.footer.is_linked_to_previous is True
    document.add_paragraph("section two")

    options = {"style": "dash", "position": "center"}
    apply_page_number(document, options)
    apply_page_number(document, options)
    output = tmp_path / "linked-footer-image.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    reopened = Document(output)
    second_footer = reopened.sections[1].footer
    assert second_footer.is_linked_to_previous is False
    blips = second_footer._element.findall(".//" + qn("a:blip"))
    assert len(blips) == 1
    relationship_id = blips[0].get(qn("r:embed"))
    assert relationship_id in second_footer.part.rels
    assert second_footer.part.rels[relationship_id].target_part.blob
    assert _field_instructions(second_footer._element).count("PAGE") == 1
    _assert_native_dash_paragraph(_page_paragraph(second_footer._element))


def test_linked_footer_preserves_hyperlink_relationship_when_materialized(
    tmp_path: Path,
) -> None:
    document = Document()
    first = document.sections[0]
    first.footer.is_linked_to_previous = False
    paragraph = first.footer.paragraphs[0]
    paragraph.add_run("Reference: ")
    _add_hyperlink(paragraph, "public site", "https://example.com/footer")
    _add_field(first.footer.add_paragraph(), "PAGE")
    document.add_paragraph("section one")
    second = document.add_section(WD_SECTION.NEW_PAGE)
    assert second.footer.is_linked_to_previous is True
    document.add_paragraph("section two")

    apply_page_number(document, {"style": "dash", "position": "center"})
    output = tmp_path / "linked-footer-hyperlink.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    reopened = Document(output)
    second_footer = reopened.sections[1].footer
    hyperlinks = second_footer._element.findall(".//" + qn("w:hyperlink"))
    assert len(hyperlinks) == 1
    relationship_id = hyperlinks[0].get(qn("r:id"))
    relationship = second_footer.part.rels[relationship_id]
    assert relationship.is_external is True
    assert relationship.target_ref == "https://example.com/footer"
    _assert_native_dash_paragraph(_page_paragraph(second_footer._element))


def test_split_complex_page_field_is_replaced_without_removing_split_date(
    tmp_path: Path,
) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    _add_split_field(paragraph, (" DA", "TE \\@ yyyy-MM-dd "), "2026-08-19")
    paragraph.add_run(" ")
    _add_split_field(paragraph, (" PA", "GE "), "7")

    apply_page_number(document, {"style": "plain", "position": "center"})
    output = tmp_path / "split-complex-fields.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    root = next(iter(_footer_roots(output).values()))
    instructions = _field_instructions(root)
    assert "".join(instructions[:2]).startswith("DATE")
    assert sum("PAGE" in instruction for instruction in instructions) == 1
    assert "2026-08-19" in _visible_text(root)
    assert "7" not in _visible_text(root)
    _assert_complex_fields_are_paired(root)


def test_outside_page_numbers_preserve_default_and_even_headers(tmp_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.header.is_linked_to_previous = False
    section.header.paragraphs[0].text = "Default header"
    section.even_page_header.is_linked_to_previous = False
    section.even_page_header.paragraphs[0].text = "Even header"

    apply_page_number(document, {"style": "dash", "position": "outside"})
    output = tmp_path / "outside-preserves-headers.docx"
    document.save(output)

    reopened = Document(output)
    assert reopened.sections[0].header.paragraphs[0].text == "Default header"
    assert reopened.sections[0].even_page_header.paragraphs[0].text == "Even header"
    roots = _footer_roots(output)
    assert sorted(_field_instructions(root) for root in roots.values()) == [["PAGE"], ["PAGE"]]


def test_page_number_can_apply_first_and_even_centered_footers(tmp_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_footer.is_linked_to_previous = False
    section.first_page_footer.paragraphs[0].text = "First note"
    document.settings._element.append(OxmlElement("w:evenAndOddHeaders"))
    section.even_page_footer.is_linked_to_previous = False
    section.even_page_footer.paragraphs[0].text = "Even note"
    document.add_paragraph("body")

    apply_page_number(document, {"style": "cn", "position": "center", "first_page": True})
    output = tmp_path / "first-even.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    footers = _footer_roots(output)
    assert len(footers) == 3
    for root in footers.values():
        assert _field_instructions(root) == ["PAGE"]
        _assert_complex_fields_are_paired(root)
    footer_xml = "\n".join(_footer_xml(output).values())
    assert "First note" in footer_xml
    assert "Even note" in footer_xml
    assert footer_xml.count('w:val="center"') == 3


def test_standard_page_number_has_explicit_font_size_and_outside_indents(tmp_path: Path) -> None:
    document = Document()
    document.sections[0].bottom_margin = Cm(3.5)
    normal_ppr = document.styles["Normal"]._element.get_or_add_pPr()
    normal_indent = normal_ppr.find(qn("w:ind"))
    if normal_indent is None:
        normal_indent = OxmlElement("w:ind")
        normal_ppr.append(normal_indent)
    normal_indent.set(qn("w:firstLineChars"), "200")
    normal_indent.set(qn("w:firstLine"), "640")
    document.add_paragraph("body")

    options = {
        "style": "dash",
        "position": "outside",
        "first_page": True,
        "section_numbering": "continue",
        "offset_from_text_mm": 7,
        "font_name": "宋体",
        "font_size_pt": 14,
        "bold": False,
    }
    apply_page_number(document, options)
    apply_page_number(document, options)
    output = tmp_path / "standard-page-number.docx"
    document.save(output)

    assert validate_docx_integrity(output).ok is True
    reopened = Document(output)
    assert abs(reopened.sections[0].footer_distance.cm - 2.8) < 0.02
    roots = _footer_roots(output)
    assert len(roots) == 3
    alignments = []
    for root in roots.values():
        assert _field_instructions(root) == ["PAGE"]
        paragraph = _page_paragraph(root)
        _assert_native_dash_paragraph(paragraph)
        alignment, indent = _paragraph_alignment_and_indent(paragraph)
        alignments.append(alignment)
        assert indent is not None
        if alignment == "left":
            assert indent.get(qn("w:left")) == "280"
            assert indent.get(qn("w:right")) is None
            assert indent.get(qn("w:leftChars")) is None
        else:
            assert alignment == "right"
            assert indent.get(qn("w:right")) == "280"
            assert indent.get(qn("w:left")) is None
            assert indent.get(qn("w:rightChars")) is None
        assert indent.get(qn("w:firstLineChars")) == "0"
        assert indent.get(qn("w:firstLine")) == "0"
        runs = paragraph.findall(qn("w:r"))
        assert len(runs) == 4
        for run in runs:
            properties = run.find(qn("w:rPr"))
            fonts = properties.find(qn("w:rFonts"))
            assert fonts is not None
            for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                assert fonts.get(qn(attribute)) == "宋体"
            assert properties.find(qn("w:sz")).get(qn("w:val")) == "28"
            assert properties.find(qn("w:szCs")).get(qn("w:val")) == "28"
    assert alignments.count("right") == 2
    assert alignments.count("left") == 1


def test_page_number_removes_static_page_number_paragraphs_before_relayout(tmp_path: Path) -> None:
    document = Document()
    footer = document.sections[0].footer
    footer.paragraphs[0].text = "1"
    footer.add_paragraph("— 1 —")
    footer.add_paragraph("Confidential footer")
    document.add_paragraph("body")

    apply_page_numbers(document, {"style": "dash", "position": "center"})
    output = tmp_path / "static-page-number-reset.docx"
    document.save(output)

    root = next(iter(_footer_roots(output).values()))
    paragraphs = root.findall(f".//{{{W_NS}}}p")
    assert len(paragraphs) == 2
    assert _field_instructions(root) == ["PAGE"]
    assert _visible_text(root).count("1") == 0
    assert "Confidential footer" in _visible_text(root)


def test_non_outside_page_numbers_clear_character_indents(tmp_path: Path) -> None:
    document = Document()
    apply_page_number(document, {"position": "outside", "first_page": True})
    apply_page_number(document, {"position": "center", "first_page": True})
    output = tmp_path / "centered-page-number.docx"
    document.save(output)

    for root in _footer_roots(output).values():
        paragraph = _page_paragraph(root)
        alignment, indent = _paragraph_alignment_and_indent(paragraph)
        assert alignment == "center"
        if indent is not None:
            assert indent.get(qn("w:leftChars")) is None
            assert indent.get(qn("w:rightChars")) is None
