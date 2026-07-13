import json
import logging
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from engine import export_doc
from engine._core import TYPE_TO_RULE_INDEX
from importer import DocumentData, ParagraphData, ParagraphFeatures
from style_config import ConfigValidationError, PageSettings, StyleRule, logger


def _xml_attr(element, attr):
    return element.get(qn(attr)) if element is not None else None


def _run_east_asia_font(paragraph):
    rPr = paragraph.runs[0]._element.rPr if paragraph.runs else None
    rFonts = rPr.rFonts if rPr is not None else None
    return _xml_attr(rFonts, "w:eastAsia")


def _paragraph_spacing(paragraph):
    pPr = paragraph._element.pPr
    return pPr.find(qn("w:spacing")) if pPr is not None else None


def _paragraph_alignment(paragraph):
    pPr = paragraph._element.pPr
    jc = pPr.find(qn("w:jc")) if pPr is not None else None
    return _xml_attr(jc, "w:val")


def _paragraph_indent(paragraph):
    pPr = paragraph._element.pPr
    return pPr.find(qn("w:ind")) if pPr is not None else None


class ConfigDrivenStylesTest(unittest.TestCase):
    def setUp(self):
        logger.setLevel(logging.ERROR)
        self.tmp = tempfile.TemporaryDirectory()
        self.tmp_path = Path(self.tmp.name)

    def tearDown(self):
        self.tmp.cleanup()

    def _export(self, paragraphs, rules=None, settings=None):
        output = self.tmp_path / "out.docx"
        data = DocumentData(paragraphs=paragraphs, filepath="input.docx")
        export_doc(data, rules or StyleRule.from_config(), settings or PageSettings(), str(output))
        return Document(output)

    def test_config_declares_every_special_paragraph_style_used_by_engine(self):
        rules = StyleRule.from_config()
        names = [r.level_name for r in rules]

        for expected in [
            "附件说明",
            "附件说明续项",
            "附件正文标记",
            "附件正文标题",
            "附件正文",
            "落款署名",
            "落款日期",
        ]:
            self.assertIn(expected, names)

        self.assertEqual(TYPE_TO_RULE_INDEX["attachment_note"], names.index("附件说明"))
        self.assertEqual(TYPE_TO_RULE_INDEX["attachment_note_item"], names.index("附件说明续项"))
        self.assertEqual(TYPE_TO_RULE_INDEX["attachment_page_mark"], names.index("附件正文标记"))
        self.assertEqual(TYPE_TO_RULE_INDEX["attachment_title"], names.index("附件正文标题"))
        self.assertEqual(TYPE_TO_RULE_INDEX["attachment_body"], names.index("附件正文"))
        self.assertEqual(TYPE_TO_RULE_INDEX["sign_org"], names.index("落款署名"))
        self.assertEqual(TYPE_TO_RULE_INDEX["sign_date"], names.index("落款日期"))

    def test_from_config_reads_spacing_and_side_indent_fields(self):
        config_path = self.tmp_path / "config.json"
        styles = [
            {
                "name": f"row{i}",
                "font": "仿宋_GB2312",
                "size": "三号",
                "bold": False,
                "pattern": "",
                "lang": "",
                "indent": 0,
                "align": "左对齐",
            }
            for i in range(17)
        ]
        styles[10].update({
            "name": "附件说明",
            "spacing_before": 1.5,
            "spacing_after": 0.5,
            "left_indent": 2,
            "right_indent": 1,
        })
        config_path.write_text(
            json.dumps({"styles": styles, "page": {"line_spacing_pt": 30}}, ensure_ascii=False),
            encoding="utf-8",
        )

        rules = StyleRule.from_config(str(config_path))
        settings = PageSettings.from_config(str(config_path))

        self.assertEqual(rules[10].spacing_before, 1.5)
        self.assertEqual(rules[10].spacing_after, 0.5)
        self.assertEqual(rules[10].left_indent, 2)
        self.assertEqual(rules[10].right_indent, 1)
        self.assertEqual(settings.line_spacing_value, 30)

    def test_attachment_and_signature_styles_are_read_from_config_rows(self):
        rules = StyleRule.from_config()
        rows = {rule.level_name: i for i, rule in enumerate(rules)}
        rules[rows["附件说明"]].font = "黑体"
        rules[rows["附件说明"]].alignment = "右对齐"
        rules[rows["附件说明"]].spacing_before = 2
        rules[rows["附件说明"]].left_indent = 3
        rules[rows["落款日期"]].right_indent = 4

        doc = self._export([
            ParagraphData("附件：测试", "attachment_note", "附件：测试", ParagraphFeatures()),
            ParagraphData("2026年6月18日", "sign_date", "2026年6月18日", ParagraphFeatures()),
        ], rules=rules)

        note = doc.paragraphs[0]
        sign_date = doc.paragraphs[1]
        note_spacing = _paragraph_spacing(note)
        note_indent = _paragraph_indent(note)
        sign_indent = _paragraph_indent(sign_date)

        self.assertEqual(_run_east_asia_font(note), "黑体")
        self.assertEqual(_paragraph_alignment(note), "right")
        self.assertEqual(_xml_attr(note_spacing, "w:beforeLines"), "200")
        self.assertEqual(_xml_attr(note_indent, "w:leftChars"), "300")
        self.assertEqual(_xml_attr(sign_indent, "w:rightChars"), "400")

    def test_line_spacing_setting_controls_paragraph_and_doc_grid_pitch(self):
        doc = self._export([
            ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
        ], settings=PageSettings(line_spacing_value=30))

        spacing = _paragraph_spacing(doc.paragraphs[0])
        doc_grid = doc.sections[0]._sectPr.find(qn("w:docGrid"))

        self.assertEqual(_xml_attr(spacing, "w:line"), "600")
        self.assertEqual(_xml_attr(doc_grid, "w:linePitch"), "600")

    def test_frontend_page_config_accepts_missing_fields_and_numeric_strings(self):
        settings = PageSettings.from_config_dict({
            "styles": [],
            "page": {
                "width_cm": "21",
                "height_cm": "29.7",
                "lines_per_page": "22",
                "chars_per_line": "28",
            },
        })

        self.assertEqual(settings.page_width_cm, 21.0)
        self.assertEqual(settings.page_height_cm, 29.7)
        self.assertEqual(settings.margin_top_cm, 3.7)
        self.assertEqual(settings.lines_per_page, 22)
        self.assertEqual(settings.chars_per_line, 28)

    def test_frontend_page_config_rejects_invalid_numeric_values(self):
        invalid_values = [None, "", "abc", "NaN", "Infinity", "-Infinity", -1, 0, 1e300]

        for value in invalid_values:
            with self.subTest(value=value), self.assertRaises(ConfigValidationError) as ctx:
                PageSettings.from_config_dict({"styles": [], "page": {"width_cm": value}})
            self.assertEqual(ctx.exception.code, "FORMAT_CONFIG_INVALID")
            self.assertEqual(ctx.exception.field, "page.width_cm")

    def test_frontend_page_config_rejects_non_positive_layout_area(self):
        invalid_pages = [
            {"width_cm": 21, "margin_left_cm": 10.5, "margin_right_cm": 10.5},
            {"height_cm": 29.7, "margin_top_cm": 15, "margin_bottom_cm": 14.7},
            {"lines_per_page": 0},
            {"chars_per_line": 0},
        ]

        for page in invalid_pages:
            with self.subTest(page=page), self.assertRaises(ConfigValidationError):
                PageSettings.from_config_dict({"styles": [], "page": page})

    def test_frontend_style_config_rejects_invalid_style_numbers(self):
        invalid_styles = [
            {"size": 0},
            {"size": "不存在字号"},
            {"indent": "NaN"},
            {"spacing_before": -1},
            {"left_indent": -1},
        ]

        for style in invalid_styles:
            with self.subTest(style=style), self.assertRaises(ConfigValidationError):
                StyleRule.from_config_dict({"styles": [style], "page": {}})

    def test_frontend_style_config_keeps_legacy_valid_config(self):
        rules = StyleRule.from_config_dict({
            "styles": [{
                "name": "正文",
                "font": "仿宋_GB2312",
                "size": "三号",
                "bold": False,
                "indent": "2",
                "spacing_before": "1",
                "spacing_after": "0",
                "left_indent": "0",
                "right_indent": "0",
            }],
            "page": {},
        })

        self.assertEqual(rules[0].font_size_pt, 16.0)
        self.assertEqual(rules[0].first_line_indent, 2.0)
        self.assertEqual(rules[0].spacing_before, 1.0)


if __name__ == "__main__":
    unittest.main()
