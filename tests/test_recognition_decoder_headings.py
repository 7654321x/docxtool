"""Tests split from test_recognition_decoder.py by responsibility."""

from concurrent.futures import ThreadPoolExecutor

from docx import Document
from docxtool.document.engine import export_doc
import docxtool.document.recognition.decoder as decoder
from docxtool.document.recognition.candidates import Candidate
from docxtool.document.recognition import (
    ParagraphType,
    RecognitionConfig,
    SectionKind,
    apply_recognition,
    diagnostics_to_json,
    resolve_render_mapping,
)
from docxtool.document.style_config import PageSettings



from tests.support.recognition_helpers import _document, _paragraph, _rules


def test_duplicate_heading_sequence_is_applied_but_marked_for_review() -> None:
    duplicate = _paragraph("一、重复编号标题", "body", 4)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、持续夯实基层基础", "body", 2),
        _paragraph("持续完善责任体系，确保各项部署落实到位。", "body", 3),
        duplicate,
        _paragraph("后续正文对该标题展开具体说明，并列举工作安排、责任分工和落实要求，确保形成完整支撑。", "body", 5),
    )

    apply_recognition(data)

    assert duplicate.type_id == "heading1"
    diagnostic = data.recognition_diagnostics["paragraphs"][4]
    assert diagnostic["review_level"] == "review"
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]
    assert "numbering-duplicate" in diagnostic["heading_context_evidence"]

def test_child_heading_numbering_resets_under_new_parent_scope() -> None:
    first_child = _paragraph("（一）第一章下的子标题", "body", 3)
    second_child = _paragraph("（一）第二章下的子标题", "body", 6)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、第一部分", "body", 2),
        first_child,
        _paragraph("正文对第一章子标题展开说明。", "body", 4),
        _paragraph("二、第二部分", "body", 5),
        second_child,
        _paragraph("正文对第二章子标题展开说明。", "body", 7),
    )

    apply_recognition(data)

    assert [first_child.type_id, second_child.type_id] == ["heading2", "heading2"]
    second_diagnostic = data.recognition_diagnostics["paragraphs"][6]
    assert "numbering-duplicate" not in second_diagnostic["heading_context_evidence"]
    assert second_diagnostic["review_level"] not in {"review", "critical_review"}
    h2_families = [
        item for item in data.recognition_diagnostics["document_context"]["heading_families"]
        if item["level"] == 2
    ]
    assert [item["parent_scope"] for item in h2_families] == [[2], [5]]

def test_duplicate_child_heading_in_same_parent_scope_is_reviewed() -> None:
    duplicate = _paragraph("（一）同一父标题下重复的子标题", "body", 5)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、第一部分", "body", 2),
        _paragraph("（一）第一个子标题", "body", 3),
        _paragraph("正文对第一个子标题展开说明。", "body", 4),
        duplicate,
        _paragraph("正文对重复子标题展开说明。", "body", 6),
    )

    apply_recognition(data)

    assert duplicate.type_id == "heading2"
    diagnostic = data.recognition_diagnostics["paragraphs"][5]
    assert "numbering-duplicate" in diagnostic["heading_context_evidence"]
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]

def test_deep_heading_cannot_borrow_parent_from_previous_h1_scope() -> None:
    orphan = _paragraph("1.缺少当前父级的三级标题", "body", 7)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、第一部分", "body", 2),
        _paragraph("（一）第一章子标题", "body", 3),
        _paragraph("1.第一章三级标题", "body", 4),
        _paragraph("正文对第一章三级标题展开说明。", "body", 5),
        _paragraph("二、第二部分", "body", 6),
        orphan,
        _paragraph("正文对孤立三级标题展开说明。", "body", 8),
    )

    apply_recognition(data)

    diagnostic = data.recognition_diagnostics["paragraphs"][7]
    assert "missing-parent-heading" in diagnostic["heading_context_evidence"]
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]

def test_orphan_deep_heading_is_not_confirmed_without_parent_context() -> None:
    orphan = _paragraph("1.缺少父级的三级标题", "body", 2)
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        orphan,
        _paragraph("后续正文对该段展开具体说明，并形成足够的正文支撑材料。", "body", 3),
    )

    apply_recognition(data)

    diagnostic = data.recognition_diagnostics["paragraphs"][2]
    assert "heading3" in diagnostic["candidate_types"]
    assert "missing-parent-heading" in diagnostic["heading_context_evidence"]
    assert diagnostic["review_level"] == "review"
    assert "HEADING_SEQUENCE_CONFLICT" in diagnostic["review_reasons"]

def test_short_bold_source_list_heading_survives_core_body_candidate() -> None:
    paragraph = _paragraph(
        "企业在订单班的重视度和投入上主动性不足",
        "heading2",
        2,
        numbering_prefix="@lvl_0",
        bold_char_ratio=1.0,
        classification_kind="body",
        classification_confidence=0.78,
    )
    data = _document(
        _paragraph("三、存在的主要问题", "heading1", 0),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
        _paragraph("后续正文对该标题展开具体说明。", "body", 3),
    )

    apply_recognition(data)

    assert paragraph.type_id == "heading2"
    assert paragraph.meta["recognition_provider"].startswith("source-list-numbering:")

def test_source_list_level_zero_is_parent_of_following_literal_heading2() -> None:
    paragraph = _paragraph(
        "阶段工作安排",
        "heading2",
        2,
        numbering_prefix="@lvl_0",
        classification_kind="body",
        classification_confidence=0.78,
    )
    data = _document(
        _paragraph("工作总结", "title", 0),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
        _paragraph("（一）第一项工作", "body", 3),
        _paragraph("后续正文对该标题展开具体说明。", "body", 4),
    )

    apply_recognition(data)

    assert paragraph.type_id == "heading1"
    assert paragraph.meta["recognition_provider"].startswith("source-list-numbering:")
    trace = data.recognition_diagnostics["candidate_trace"][2]
    candidate = next(item for item in trace["candidates"] if item["source"] == "source-list-numbering")
    assert candidate["type"] == "heading1"
    assert "parent-of-following-heading2" in candidate["evidence"]

def test_unbold_source_list_level_zero_without_child_support_stays_body() -> None:
    paragraph = _paragraph(
        "阶段工作安排",
        "heading2",
        2,
        numbering_prefix="@lvl_0",
        classification_kind="body",
        classification_confidence=0.78,
    )
    data = _document(
        _paragraph("工作总结", "title", 0),
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 1),
        paragraph,
        _paragraph("后续正文继续展开具体说明。", "body", 3),
    )

    apply_recognition(data)

    assert paragraph.type_id == "body"
    trace = data.recognition_diagnostics["candidate_trace"][2]
    assert "source-list-numbering" not in {item["source"] for item in trace["candidates"]}

def test_long_source_list_prose_is_not_promoted_to_heading() -> None:
    paragraph = _paragraph(
        "这是继承了Word列表属性的较长正文段落，其中包含完整叙述和多项具体情况，不应因为隐藏列表属性被识别成标题。",
        "body",
        1,
        numbering_prefix="@lvl_0",
        bold_char_ratio=1.0,
        classification_kind="body",
        classification_confidence=0.9,
    )
    data = _document(_paragraph("工作情况", "title", 0), paragraph)

    apply_recognition(data)

    assert paragraph.type_id == "body"

def test_title_formatting_cannot_override_front_role_and_placeholder_date():
    data = _document(
        _paragraph("在区政协第九届委员会第一次会议", "title", 0, alignment="CENTER"),
        _paragraph("召集人会议上的讲话", "title_cont", 1, alignment="CENTER"),
        _paragraph("区委副书记  XXX", "role_name", 2, alignment="CENTER"),
        _paragraph("2026年8月X日", "date_line", 3, alignment="CENTER"),
        _paragraph("同志们：", "addressing", 4),
        _paragraph("现将会议有关事项说明如下，确保各项工作衔接有序。", "body", 5),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:5]] == [
        "title", "title_cont", "role_name", "date_line", "addressing",
    ]
    assert data.recognition_diagnostics["paragraphs"][2]["review_level"] != "critical_review"
    assert data.recognition_diagnostics["paragraphs"][3]["review_level"] != "critical_review"

def test_front_role_metadata_vetoes_title_continuation_transition_bonus() -> None:
    role = _paragraph(
        "某区委员会党组书记、主席张测试",
        "role_name",
        1,
        alignment="CENTER",
        bold_char_ratio=1.0,
        legacy_type_id="role_name",
    )
    date = _paragraph(
        "（2026年8月27日11:00，某地区委员会会议中心）",
        "date_line",
        2,
        alignment="CENTER",
        bold_char_ratio=1.0,
        legacy_type_id="date_line",
    )
    data = _document(
        _paragraph("在某地区委员会会议闭幕大会上的讲话", "title", 0, alignment="CENTER"),
        role,
        date,
        _paragraph("各位代表、同志们：", "addressing", 3),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:5]] == [
        "title", "role_name", "date_line", "addressing", "body",
    ]
    assert role.meta["review_level"] != "critical_review"
    assert date.meta["review_level"] != "critical_review"
    assert "colon-explanatory-body" not in date.meta["recognition_evidence"]

def test_front_role_metadata_supports_generic_role_and_name_shapes() -> None:
    cases = (
        ("某办公室主任 张某", "2026年8月27日"),
        ("某委员会书记、主席　张某某", "2026年8月27日"),
        ("某区总工程师赵某某", "2026年8月27日11:00"),
        ("某工作组负责人  张某", "各位代表、同志们："),
    )

    for role_text, following_text in cases:
        following_type = "date_line" if "年" in following_text else "addressing"
        data = _document(
            _paragraph("年度工作报告", "title", 0, alignment="CENTER"),
            _paragraph(role_text, "body", 1, alignment="CENTER"),
            _paragraph(following_text, following_type, 2, alignment="CENTER"),
            _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
        )

        apply_recognition(data)

        assert data.paragraphs[1].type_id == "role_name", role_text
        assert data.paragraphs[1].meta["recognition_provider"].startswith("front-metadata:")

def test_role_word_inside_real_title_continuation_is_not_a_person_name() -> None:
    continuation = _paragraph(
        "主席会议制度建设情况",
        "title_cont",
        1,
        alignment="CENTER",
    )
    data = _document(
        _paragraph("年度工作报告", "title", 0, alignment="CENTER"),
        continuation,
        _paragraph("2026年8月27日", "date_line", 2, alignment="CENTER"),
        _paragraph("各有关单位：", "addressing", 3),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert continuation.type_id == "title_cont"
    assert data.recognition_diagnostics["document_context"]["front_metadata"] == [
        {"position": 2, "kind": "date_line"},
    ]

def test_front_metadata_shape_recognizes_unseen_role_and_placeholder_date_without_legacy_type():
    data = _document(
        _paragraph("在推进重点项目专题会议上的讲话", "body", 0, alignment="CENTER"),
        _paragraph("某区总工程师张某", "body", 1, alignment="CENTER"),
        _paragraph("2026年8月X日", "body", 2, alignment="CENTER"),
        _paragraph("同志们：", "body", 3),
        _paragraph("现将会议有关事项说明如下，确保各项工作衔接有序。", "body", 4),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:5]] == [
        "title", "role_name", "date_line", "addressing", "body",
    ]
    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert diagnostics[1]["provider"].startswith("front-metadata:")
    assert diagnostics[2]["provider"].startswith("front-metadata:")
    assert diagnostics[1]["review_level"] in {"confirmed", "info"}
    assert diagnostics[2]["review_level"] in {"confirmed", "info"}

def test_numbered_heading_family_confirms_parallel_body_headings():
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("现将有关情况报告如下，供审阅。", "body", 1),
        _paragraph("一、持续夯实基层基础", "body", 2),
        _paragraph("持续完善责任体系，确保各项部署落实到位，并将责任细化到岗、任务落实到人、过程监督到位。", "body", 3),
        _paragraph("二、稳步提升服务质效", "body", 4),
        _paragraph("围绕群众需求优化流程，持续提升服务效率和群众办事体验，确保事项办理规范有序。", "body", 5),
    )

    apply_recognition(data)

    assert [data.paragraphs[index].type_id for index in (2, 4)] == ["heading1", "heading1"]
    context = data.recognition_diagnostics["document_context"]
    assert context["heading_families"] == [
        {"level": 1, "count": 2, "positions": [2, 4], "supported_count": 2},
    ]
    assert "parallel-heading-family" in data.recognition_diagnostics["paragraphs"][2]["heading_context_evidence"]

def test_uniform_legacy_heading_siblings_are_finalized_by_recognition() -> None:
    data = _document(
        _paragraph("一、父级标题", "heading1", 0),
        _paragraph("第一项", "heading3", 1),
        _paragraph("第二项", "heading3", 2),
        _paragraph("正文内容完整结束。", "body", 3),
    )

    apply_recognition(data)

    assert [paragraph.type_id for paragraph in data.paragraphs[:3]] == [
        "heading1", "heading2", "heading2",
    ]
    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert [item["recognized_type"] for item in diagnostics[:3]] == [
        "heading1", "heading2", "heading2",
    ]
    assert [item["final_type"] for item in diagnostics[:3]] == [
        "heading1", "heading2", "heading2",
    ]
    assert all(
        "uniform-heading-sibling-family" in item["evidence_summary"]
        for item in diagnostics[1:3]
    )

def test_single_numbered_heading_is_applied_but_marked_for_review():
    data = _document(
        _paragraph("工作情况", "title", 0, alignment="CENTER"),
        _paragraph("一、持续夯实基层基础", "body", 1),
        _paragraph("持续完善责任体系，确保各项部署落实到位，并将责任细化到岗、任务落实到人、过程监督到位。", "body", 2),
    )

    apply_recognition(data)

    heading = data.recognition_diagnostics["paragraphs"][1]
    assert data.paragraphs[1].type_id == "heading1"
    assert heading["review_level"] == "review"
    assert "LEGACY_TYPE_CONFLICT" in heading["review_reasons"]
    assert data.paragraphs[1].original_text == "一、持续夯实基层基础"

def test_table_boundary_blocks_title_continuation():
    data = _document(
        _paragraph("主标题", "title", 0),
        _paragraph("", "__table__", 1),
        _paragraph("表后说明", "body", 2, alignment="CENTER"),
    )

    apply_recognition(data)

    assert data.recognition_diagnostics["candidate_trace"][1]["boundary_before"] is True
    assert data.paragraphs[2].meta["recognition_type"] != "title_continuation"

def test_wrong_legacy_and_docxtool_style_do_not_override_dispatch():
    paragraph = _paragraph("国发〔2026〕23号", "title", 0, style_name="DCT-Title", legacy_type_id="title")
    data = _document(paragraph)

    apply_recognition(data)

    assert paragraph.type_id == "dispatch_number"
    assert paragraph.meta["legacy_type_id"]["value"] == "title"

def test_wrong_heading_legacy_preserves_numbered_heading2_colon_structure():
    paragraph = _paragraph("（一）缺席：李四", "heading2", 1, style_name="DCT-Heading2", legacy_type_id="heading2")
    data = _document(_paragraph("党委会会议纪要", "title", 0), paragraph)

    apply_recognition(data)

    assert paragraph.type_id == "heading2"

def test_every_paragraph_type_has_explicit_render_mapping():
    for paragraph_type in ParagraphType:
        type_id, rule_index, style_id = resolve_render_mapping(paragraph_type)
        if paragraph_type is ParagraphType.UNKNOWN:
            assert type_id is None and rule_index is None and style_id is None
        else:
            assert type_id
            assert isinstance(rule_index, int)
            assert style_id and style_id.startswith("DCT-")

def test_real_docx_round_trip_preserves_semantics_and_layout(tmp_path):
    source = tmp_path / "source.docx"
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    document = Document()
    document.add_paragraph("国务院关于印发规划的通知")
    document.add_paragraph("国发〔2026〕23号")
    document.add_paragraph("各有关单位：")
    document.add_paragraph("正文内容。")
    document.add_paragraph("国务院")
    document.add_paragraph("2026年3月18日")
    document.add_paragraph("（本文有删减）")
    document.add_paragraph("公共服务提升规划")
    document.add_paragraph("第一章 总则")
    document.save(source)

    importer = __import__("docxtool.document.importer", fromlist=["DocxImporter"]).DocxImporter()
    first_data = importer.load(str(source), _rules())
    export_doc(first_data, _rules(), PageSettings(), str(first))
    second_data = importer.load(str(first), _rules())
    export_doc(second_data, _rules(), PageSettings(), str(second))
    third_data = importer.load(str(second), _rules())

    def signature(data):
        return [(item.type_id, item.text, item.meta.get("recognition_type"), item.meta.get("recognition_section")) for item in data.paragraphs]

    assert signature(second_data) == signature(third_data)
    assert [item.type_id for item in second_data.paragraphs].count("dispatch_number") == 1
    assert [item.type_id for item in second_data.paragraphs].count("embedded_document_title") == 1
    assert len(second_data.paragraphs) == len(third_data.paragraphs)
    assert second_data.recognition_diagnostics["validation"]["ok"] is True

def test_recognition_config_rejects_invalid_values():
    for kwargs in (
        {"beam_width": 1},
        {"max_candidates_per_paragraph": 1},
        {"legacy_score": 1.1},
        {"text_preview_length": -1},
        {"unknown_render_type": "silent"},
    ):
        try:
            RecognitionConfig(**kwargs)
        except ValueError:
            continue
        raise AssertionError(f"invalid config accepted: {kwargs}")

def test_disabling_diagnostics_does_not_change_decisions():
    enabled = _document(_paragraph("主标题", "title", 0), _paragraph("国发〔2026〕23号", "title_cont", 1))
    disabled = _document(_paragraph("主标题", "title", 0), _paragraph("国发〔2026〕23号", "title_cont", 1))

    apply_recognition(enabled, RecognitionConfig(enable_diagnostics=True))
    apply_recognition(disabled, RecognitionConfig(enable_diagnostics=False))

    assert [item.type_id for item in enabled.paragraphs] == [item.type_id for item in disabled.paragraphs]
    assert enabled.recognition_diagnostics["candidate_trace"]
    assert disabled.recognition_diagnostics["candidate_trace"] == []
    assert disabled.recognition_diagnostics["engine_version"] == "3.0"
    assert disabled.recognition_diagnostics["schema_version"] == "1.0"

def test_empty_and_table_only_documents_do_not_create_fake_paragraphs():
    empty = _document()
    table_only = _document(_paragraph("", "__table__", 0))

    apply_recognition(empty)
    apply_recognition(table_only)

    assert empty.recognition_diagnostics["paragraphs"] == []
    assert table_only.recognition_diagnostics["paragraphs"] == []
    assert table_only.recognition_diagnostics["blocks"][0]["kind"] == "table"

def test_long_and_unusual_unicode_text_is_bounded_and_preserved():
    raw = "Ａ：" + "甲\u00a0\u200b" * 25000
    paragraph = _paragraph(raw, "body", 0)
    data = _document(paragraph)

    apply_recognition(data, RecognitionConfig(text_preview_length=10))

    assert paragraph.original_text == raw
    assert len(data.recognition_diagnostics["paragraphs"][0]["text_preview"]) == 10
    assert len(data.recognition_diagnostics["candidate_trace"]) == 1

def test_incomplete_heading_level_is_diagnosed_without_text_loss():
    paragraph = _paragraph("（三）直接出现", "heading2", 0)
    data = _document(paragraph)

    apply_recognition(data)

    assert paragraph.original_text == "（三）直接出现"
    assert data.recognition_diagnostics["paragraphs"][0]["candidate_count"] >= 1

def test_dispatch_variants_are_stable_after_nfkc():
    for index, text in enumerate(("国发〔2026〕23号", "市府[2027]1号", "ＡＢ〔２０２８〕１２３号", "国发〔 2026 〕 23 号")):
        paragraph = _paragraph(text, "title_cont", index)
        data = _document(paragraph)
        apply_recognition(data)
        assert paragraph.type_id == "dispatch_number", text

def test_key_value_and_source_variants_do_not_promote_numbering():
    for index, text in enumerate(("（一）缺席：李四", "（二）出 席:张三", "来源：国家卫生健康委员会")):
        paragraph = _paragraph(text, "heading2", index)
        data = _document(_paragraph("党委会会议纪要", "title", 0), paragraph)
        apply_recognition(data)
        if text.startswith("来源"):
            assert paragraph.type_id == "note"
        else:
            assert paragraph.type_id == "heading2"

def test_review_flags_and_safe_summary_do_not_change_final_types():
    clear = _document(_paragraph("国发〔2026〕23号", "title_cont", 0))
    ambiguous = _document(_paragraph("补充说明", "body", 0, alignment="CENTER"))

    apply_recognition(clear)
    apply_recognition(ambiguous, RecognitionConfig(review_low_score=0.9))

    clear_diagnostic = clear.recognition_diagnostics["paragraphs"][0]
    ambiguous_diagnostic = ambiguous.recognition_diagnostics["paragraphs"][0]
    assert clear_diagnostic["needs_review"] is False
    assert clear_diagnostic["review_level"] == "info"
    assert "STRUCTURE_CONFIRMED_RECLASSIFICATION" in clear_diagnostic["review_reasons"]
    assert ambiguous_diagnostic["needs_review"] is True
    assert ambiguous_diagnostic["review_reasons"]
    assert clear.paragraphs[0].type_id == "dispatch_number"
    summary = ambiguous.recognition_diagnostics["summary"]
    assert summary["needs_review_count"] == 1
    assert "补充说明" not in diagnostics_to_json(ambiguous.recognition_diagnostics)

def test_explicit_numbering_is_confirmed_even_when_raw_candidate_softmax_is_low():
    data = _document(
        _paragraph("一、工作开展情况", "heading1", 0),
        _paragraph("（一）落实重点任务", "heading2", 1),
        _paragraph("1. 压实责任", "heading3", 2),
        _paragraph("（1）明确时限", "heading4", 3),
    )

    apply_recognition(data)

    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert [item["final_type"] for item in diagnostics] == ["heading1", "heading2", "heading3", "heading4"]
    assert all(item["review_level"] == "confirmed" for item in diagnostics)
    assert not any(item["needs_review"] for item in diagnostics)
    assert all("explicit-numbering" in item["evidence_summary"] for item in diagnostics)

def test_same_input_is_thread_safe_across_twenty_independent_documents():
    def recognize(_):
        data = _document(
            _paragraph("2026年第一次党委会会议纪要", "title", 0),
            _paragraph("出席：甲、乙", "body", 1),
            _paragraph("（一）缺席：无", "heading2", 2),
        )
        apply_recognition(data)
        return (
            data.doc_mode,
            tuple((item.type_id, item.meta["recognition_section"]) for item in data.paragraphs),
            data.recognition_diagnostics["summary"],
        )

    with ThreadPoolExecutor(max_workers=8) as executor:
        results = list(executor.map(recognize, range(20)))

    assert len(results) == 20
    assert all(result == results[0] for result in results)

def test_diagnostics_use_candidates_from_final_winning_beam(monkeypatch) -> None:
    class BranchingProvider:
        name = "branching"

        def propose(self, block, features, context):
            if context.index == 0:
                return [
                    Candidate(ParagraphType.BODY, 0.60, self.name, ("start-body",), section_hint=SectionKind.BODY),
                    Candidate(ParagraphType.ADDRESSING, 0.55, self.name, ("start-addressing",), section_hint=SectionKind.RECIPIENT),
                ]
            if context.index == 1 and context.previous_type == ParagraphType.ADDRESSING:
                return [
                    Candidate(ParagraphType.HEADING_1, 1.00, self.name, ("after-addressing",), section_hint=SectionKind.BODY),
                ]
            if context.index == 1:
                return [
                    Candidate(ParagraphType.BODY, 0.10, self.name, ("after-body",), section_hint=SectionKind.BODY),
                ]
            return [
                Candidate(ParagraphType.BODY, 0.10, self.name, ("tail",), section_hint=SectionKind.BODY),
            ]

    monkeypatch.setattr(decoder, "DEFAULT_PROVIDERS", (BranchingProvider(),))
    data = _document(
        _paragraph("起始行", "body", 0),
        _paragraph("路径相关行", "body", 1),
    )

    apply_recognition(data, RecognitionConfig(beam_width=2, max_candidates_per_paragraph=4))

    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert [item["final_type"] for item in diagnostics] == ["addressing", "heading1"]
    assert diagnostics[1]["candidate_types"] == ["heading1"]
    assert diagnostics[1]["provider"].startswith("branching:after-addressing")
    assert diagnostics[1]["selected_candidate_score"] == 1.0
