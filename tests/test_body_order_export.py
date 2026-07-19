import base64
import copy
import logging
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm
from docx.shared import Pt, RGBColor

from docxtool.document.engine import export_doc
from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import PageSettings, StyleRule, logger


def _rules():
    return [StyleRule.default_for_row(i) for i in range(10)]


def _body_order(path):
    doc = Document(path)
    order = []
    for child in doc._body._element.iterchildren():
        if child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t")))
            has_image = bool(child.findall(".//" + qn("a:blip")))
            if has_image:
                order.append(("image", ""))
            elif text:
                order.append(("paragraph", text))
        elif child.tag == qn("w:tbl"):
            text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t")))
            order.append(("table", text))
    return order


def _document_xml_root(path):
    with zipfile.ZipFile(path) as zf:
        return zf.read("word/document.xml")


def _paragraph_xml_by_text(path, expected):
    doc = Document(path)
    return next(paragraph._p.xml for paragraph in doc.paragraphs if paragraph.text == expected)


def _paragraph_xml_without_spacing(path, expected):
    doc = Document(path)
    element = copy.deepcopy(next(p._p for p in doc.paragraphs if p.text == expected))
    p_pr = element.find(qn("w:pPr"))
    if p_pr is not None:
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is not None:
            p_pr.remove(spacing)
        if len(p_pr) == 0:
            element.remove(p_pr)
    return element.xml


def _table_xml_without_added_paragraph_styles(table):
    element = copy.deepcopy(table._tbl)
    for paragraph in element.findall(".//" + qn("w:p")):
        p_pr = paragraph.find(qn("w:pPr"))
        if p_pr is None:
            continue
        p_style = p_pr.find(qn("w:pStyle"))
        if p_style is not None:
            p_pr.remove(p_style)
        if len(p_pr) == 0:
            paragraph.remove(p_pr)
    return element.xml


def _replace_document_xml(source, target, transform):
    with zipfile.ZipFile(source) as src, zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/document.xml":
                data = transform(data.decode("utf-8")).encode("utf-8")
            dst.writestr(item, data)


def _inline_counts(path):
    from xml.etree import ElementTree as ET

    root = ET.fromstring(_document_xml_root(path))
    breaks = root.findall(".//" + qn("w:br"))
    page_breaks = [br for br in breaks if br.get(qn("w:type")) == "page"]
    line_breaks = [br for br in breaks if br.get(qn("w:type")) != "page"]
    tabs = root.findall(".//" + qn("w:tab"))
    rendered = root.findall(".//" + qn("w:lastRenderedPageBreak"))
    return {
        "line_breaks": len(line_breaks),
        "page_breaks": len(page_breaks),
        "rendered_page_breaks": len(rendered),
        "tabs": len(tabs),
    }


def _section_summary(path):
    from xml.etree import ElementTree as ET

    root = ET.fromstring(_document_xml_root(path))
    sections = root.findall(".//" + qn("w:sectPr"))
    summary = []
    for sect in sections:
        section_type = sect.find(qn("w:type"))
        page_size = sect.find(qn("w:pgSz"))
        margins = sect.find(qn("w:pgMar"))
        summary.append({
            "orient": page_size.get(qn("w:orient")) if page_size is not None else "",
            "type": section_type.get(qn("w:val")) if section_type is not None else "",
            "width": page_size.get(qn("w:w")) if page_size is not None else "",
            "height": page_size.get(qn("w:h")) if page_size is not None else "",
            "margin_top": margins.get(qn("w:top")) if margins is not None else "",
            "margin_bottom": margins.get(qn("w:bottom")) if margins is not None else "",
            "margin_left": margins.get(qn("w:left")) if margins is not None else "",
            "margin_right": margins.get(qn("w:right")) if margins is not None else "",
            "footer": margins.get(qn("w:footer")) if margins is not None else "",
            "grid_char_space": (
                sect.find(qn("w:docGrid")).get(qn("w:charSpace"))
                if sect.find(qn("w:docGrid")) is not None
                else ""
            ),
            "grid_chars": (
                sect.find(qn("w:docGrid")).get(qn("w:charsPerLine"))
                if sect.find(qn("w:docGrid")) is not None
                else ""
            ),
            "grid_lines": (
                sect.find(qn("w:docGrid")).get(qn("w:linesPerPage"))
                if sect.find(qn("w:docGrid")) is not None
                else ""
            ),
        })
    return summary


class BodyOrderExportTest(unittest.TestCase):
    def setUp(self):
        logger.setLevel(logging.ERROR)

    def test_keeps_table_once_at_original_body_position(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"

            doc = Document()
            doc.add_paragraph("before")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "cell"
            doc.add_paragraph("after")
            doc.save(source)

            data = DocxImporter().load(str(source), _rules())
            export_doc(data, _rules(), PageSettings(), str(output))

            self.assertEqual(
                _body_order(output),
                [("paragraph", "before"), ("table", "cell"), ("paragraph", "after")],
            )

    def test_preserves_external_hyperlink_relationship_in_table(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"

            doc = Document()
            cell_para = doc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
            rid = cell_para.part.relate_to("https://example.com/table", RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), rid)
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "link"
            run.append(text)
            hyperlink.append(run)
            cell_para._p.append(hyperlink)
            doc.save(source)

            data = DocxImporter().load(str(source), _rules())
            export_doc(data, _rules(), PageSettings(), str(output))

            exported = Document(output)
            link = exported.tables[0]._tbl.find(".//" + qn("w:hyperlink"))
            self.assertIsNotNone(link)
            exported_rel = exported.part.rels[link.get(qn("r:id"))]
            self.assertTrue(exported_rel.is_external)
            self.assertEqual(exported_rel.target_ref, "https://example.com/table")

    def test_preserves_table_and_following_caption_xml_unchanged(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"

            doc = Document()
            doc.add_paragraph("before")
            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            table.cell(0, 0).merge(table.cell(0, 1)).text = "完整表头"
            table.cell(1, 0).text = "数据A"
            table.cell(1, 1).text = "123"
            caption = doc.add_paragraph()
            caption.alignment = 2
            caption.paragraph_format.space_before = Pt(7)
            run = caption.add_run("表一完整数据表")
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(12, 34, 56)
            doc.add_paragraph("after")
            doc.save(source)

            source_doc = Document(source)
            source_table_xml = _table_xml_without_added_paragraph_styles(source_doc.tables[0])
            source_caption_xml = _paragraph_xml_without_spacing(source, "表一完整数据表")

            data = DocxImporter().load(str(source), _rules())
            caption_data = next(item for item in data.paragraphs if item.type_id == "__object_caption__")
            self.assertEqual(caption_data.meta["paragraph_xml"].text, "表一完整数据表")
            export_doc(data, _rules(), PageSettings(), str(output))

            output_doc = Document(output)
            self.assertEqual(
                _table_xml_without_added_paragraph_styles(output_doc.tables[0]),
                source_table_xml,
            )
            source_default = next(style for style in source_doc.styles if style.element.get(qn("w:default")) == "1" and style.type == 1)
            output_cell_style = output_doc.tables[0].cell(0, 0).paragraphs[0].style
            self.assertNotEqual(output_cell_style.style_id, "DCT-Body")
            self.assertTrue(output_cell_style.style_id.startswith("DCT-Preserved-"))
            self.assertNotEqual(output_cell_style.name, "Normal")
            self.assertEqual(
                [style.style_id for style in output_doc.styles if style.name == "Normal"],
                ["Normal"],
            )
            self.assertEqual(output_doc.styles["Normal"].font.size.pt, 16.0)
            self.assertEqual(
                output_cell_style.element.rPr.xml if output_cell_style.element.rPr is not None else None,
                source_default.element.rPr.xml if source_default.element.rPr is not None else None,
            )
            self.assertEqual(
                output_cell_style.element.pPr.xml if output_cell_style.element.pPr is not None else None,
                source_default.element.pPr.xml if source_default.element.pPr is not None else None,
            )
            self.assertEqual(
                _paragraph_xml_without_spacing(output, "表一完整数据表"),
                source_caption_xml,
            )
            output_caption = next(p for p in output_doc.paragraphs if p.text == "表一完整数据表")
            spacing = output_caption._p.pPr.find(qn("w:spacing"))
            self.assertEqual(spacing.get(qn("w:before")), "0")
            self.assertEqual(spacing.get(qn("w:after")), "0")
            self.assertEqual(spacing.get(qn("w:beforeLines")), "0")
            self.assertEqual(spacing.get(qn("w:afterLines")), "0")

    def test_preserves_image_and_following_caption_xml_unchanged(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"
            image = tmp / "tiny.png"
            image.write_bytes(base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
            ))

            doc = Document()
            image_paragraph = doc.add_paragraph()
            image_paragraph.alignment = 1
            image_paragraph.add_run().add_picture(str(image), width=Cm(3.21))
            caption = doc.add_paragraph("图2结构示意图")
            caption.paragraph_format.space_after = Pt(9)
            caption.runs[0].italic = True
            doc.save(source)

            source_doc = Document(source)
            source_extent = source_doc.paragraphs[0]._p.find(".//" + qn("wp:extent"))
            source_caption_xml = _paragraph_xml_without_spacing(source, "图2结构示意图")

            data = DocxImporter().load(str(source), _rules())
            export_doc(data, _rules(), PageSettings(), str(output))

            output_doc = Document(output)
            output_extent = output_doc.paragraphs[0]._p.find(".//" + qn("wp:extent"))
            self.assertEqual(dict(output_extent.attrib), dict(source_extent.attrib))
            self.assertEqual(
                _paragraph_xml_without_spacing(output, "图2结构示意图"),
                source_caption_xml,
            )

    def test_only_one_table_caption_is_protected(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"

            doc = Document()
            doc.add_table(rows=1, cols=1).cell(0, 0).text = "数据"
            table_caption = doc.add_paragraph("表1 测试表")
            table_caption.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            following = doc.add_paragraph("图2 这是普通正文")
            following.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            doc.save(source)

            data = DocxImporter().load(str(source), _rules())
            captions = [item for item in data.paragraphs if item.type_id == "__object_caption__"]

            self.assertEqual(len(captions), 1)
            self.assertEqual(captions[0].meta["paragraph_xml"].text, "表1 测试表")
            following_data = next(item for item in data.paragraphs if item.text == "图2 这是普通正文")
            self.assertEqual(following_data.type_id, "body")

            export_doc(data, _rules(), PageSettings(), str(output))
            exported = Document(output)
            exported_caption = next(p for p in exported.paragraphs if p.text == "表1 测试表")
            exported_body = next(p for p in exported.paragraphs if p.text == "图2 这是普通正文")
            self.assertEqual(str(exported_caption.runs[0].font.color.rgb), "FF0000")
            self.assertEqual(exported_body.style.style_id, "DCT-Body")
            self.assertTrue(all(run.font.color.rgb is None for run in exported_body.runs))
            self.assertTrue(all(run.font.size.pt == 16 for run in exported_body.runs if run.text))

    def test_inline_image_with_text_does_not_protect_following_body(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"
            image = tmp / "tiny.png"
            image.write_bytes(base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
            ))

            doc = Document()
            mixed = doc.add_paragraph("图片前文字")
            mixed.add_run().add_picture(str(image), width=Cm(1))
            mixed.add_run("图片后文字")
            following = doc.add_paragraph("图2 这是普通正文")
            following.runs[0].font.size = Pt(48)
            following.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            doc.save(source)

            data = DocxImporter().load(str(source), _rules())

            self.assertEqual(data.paragraphs[0].type_id, "__image__")
            following_data = next(item for item in data.paragraphs if item.text == "图2 这是普通正文")
            self.assertEqual(following_data.type_id, "body")

            export_doc(data, _rules(), PageSettings(), str(output))
            exported = Document(output)
            exported_body = next(p for p in exported.paragraphs if p.text == "图2 这是普通正文")
            self.assertEqual(exported_body.style.style_id, "DCT-Body")
            self.assertTrue(all(run.font.color.rgb is None for run in exported_body.runs))
            self.assertTrue(all(run.font.size.pt == 16 for run in exported_body.runs if run.text))

    def test_keeps_image_at_original_body_position(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"
            image = tmp / "tiny.png"
            image.write_bytes(base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
            ))

            doc = Document()
            doc.add_paragraph("before")
            doc.add_picture(str(image))
            doc.add_paragraph("after")
            doc.save(source)

            data = DocxImporter().load(str(source), _rules())
            export_doc(data, _rules(), PageSettings(), str(output))

            self.assertEqual(
                _body_order(output),
                [("paragraph", "before"), ("image", ""), ("paragraph", "after")],
            )

    def test_preserves_manual_page_breaks_line_breaks_and_tabs(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"

            doc = Document()
            p = doc.add_paragraph()
            run = p.add_run("文本前")
            run.add_break(WD_BREAK.PAGE)
            p.add_run("文本后")
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)
            run.add_break(WD_BREAK.PAGE)
            p = doc.add_paragraph()
            run = p.add_run("第一行")
            run.add_break()
            p.add_run("第二行")
            p = doc.add_paragraph()
            run = p.add_run("甲")
            run.add_tab()
            p.add_run("乙")
            doc.save(source)

            data = DocxImporter().load(str(source), _rules())
            export_doc(data, _rules(), PageSettings(), str(output))

            counts = _inline_counts(output)
            self.assertEqual(counts["page_breaks"], 4)
            self.assertEqual(counts["line_breaks"], 1)
            self.assertEqual(counts["tabs"], 1)

    def test_last_rendered_page_break_is_not_rewritten_as_manual_break(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            rendered_source = tmp / "rendered-source.docx"
            output = tmp / "output.docx"

            doc = Document()
            doc.add_paragraph("甲乙")
            doc.save(source)
            _replace_document_xml(
                source,
                rendered_source,
                lambda xml: xml.replace(
                    "<w:t>甲乙</w:t>",
                    "<w:t>甲</w:t><w:lastRenderedPageBreak/><w:t>乙</w:t>",
                ),
            )

            data = DocxImporter().load(str(rendered_source), _rules())
            export_doc(data, _rules(), PageSettings(), str(output))

            counts = _inline_counts(output)
            self.assertEqual(counts["page_breaks"], 0)
            self.assertEqual(counts["rendered_page_breaks"], 0)

    def test_preserves_multi_section_orientation_and_page_size(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            source = tmp / "source.docx"
            output = tmp / "output.docx"

            doc = Document()
            doc.add_paragraph("纵向第一页")
            landscape = doc.add_section(WD_SECTION.NEW_PAGE)
            landscape.orientation = WD_ORIENT.LANDSCAPE
            landscape.page_width = Cm(29.7)
            landscape.page_height = Cm(21.0)
            doc.add_paragraph("横向页面")
            portrait = doc.add_section(WD_SECTION.NEW_PAGE)
            portrait.orientation = WD_ORIENT.PORTRAIT
            portrait.page_width = Cm(21.0)
            portrait.page_height = Cm(29.7)
            doc.add_paragraph("纵向末页")
            doc.save(source)

            data = DocxImporter().load(str(source), _rules())
            export_doc(data, _rules(), PageSettings(), str(output))

            sections = _section_summary(output)
            self.assertGreaterEqual(len(sections), 3)
            self.assertTrue(any(section["orient"] == "landscape" for section in sections), sections)
            self.assertGreaterEqual(len({(section["width"], section["height"]) for section in sections}), 2)
            portrait_sections = [section for section in sections if section["orient"] != "landscape"]
            landscape_sections = [section for section in sections if section["orient"] == "landscape"]
            for section in portrait_sections:
                self.assertEqual(
                    (
                        section["margin_top"],
                        section["margin_bottom"],
                        section["margin_left"],
                        section["margin_right"],
                        section["footer"],
                    ),
                    ("2098", "1984", "1588", "1474", "1587"),
                )
                self.assertEqual(section["grid_char_space"], "-842")
            for section in landscape_sections:
                self.assertEqual(
                    (
                        section["margin_top"],
                        section["margin_bottom"],
                        section["margin_left"],
                        section["margin_right"],
                        section["footer"],
                    ),
                    ("1588", "1474", "1984", "2098", "1077"),
                )
                self.assertEqual(section["grid_char_space"], "27765")
            self.assertTrue(
                all(section["grid_chars"] == "28" for section in sections), sections
            )
            self.assertTrue(
                all(section["grid_lines"] == "22" for section in sections), sections
            )


if __name__ == "__main__":
    unittest.main()
