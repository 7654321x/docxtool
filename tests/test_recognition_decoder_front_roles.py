"""Tests split from test_recognition_decoder.py by responsibility."""


from docx import Document
import pytest
from docxtool.document.recognition import (
    apply_recognition,
)
from docxtool.sdk import recognize_docx



from tests.support.recognition_helpers import _document, _paragraph


def test_front_matter_scan_skips_empty_and_caption_placeholders() -> None:
    prefix = [_paragraph("", "body", index) for index in range(14)]
    caption = _paragraph("图1 工作流程", "__object_caption__", 14)
    title = _paragraph("关于推进基层治理工作的通知", "body", 15, alignment="CENTER")
    recipient = _paragraph("各有关单位：", "body", 16)
    body = _paragraph("现将有关事项通知如下，请结合实际抓好落实。", "body", 17)
    data = _document(*prefix, caption, title, recipient, body)

    apply_recognition(data)

    assert title.type_id == "title"
    assert recipient.type_id == "addressing"
    context = data.recognition_diagnostics["document_context"]
    assert 15 in context["front_matter_positions"]
    assert context["front_scan_reason"] == "body-boundary"
    assert context["body_start"] == 17

@pytest.mark.parametrize(
    "title_text",
    (
        "关于有关事项的决议",
        "关于有关事项的决定",
        "关于有关事项的命令",
        "某机关令",
        "关于有关事项的公报",
        "关于有关事项的公告",
        "关于有关事项的通告",
        "关于有关事项的意见",
        "关于有关事项的通知",
        "关于有关事项的通报",
        "关于有关事项的报告",
        "关于有关事项的请示",
        "关于有关事项的批复",
        "关于有关事项的议案",
        "关于协商有关事项的函",
        "关于有关事项的纪要",
    ),
)
def test_front_official_document_type_suffix_promotes_title_without_visual_style(
    title_text: str,
) -> None:
    title = _paragraph(title_text, "body", 0, style_name="DCT-Body")
    data = _document(
        title,
        _paragraph("各有关单位：", "body", 1),
        _paragraph("现将有关事项说明如下，请结合实际抓好落实。", "body", 2),
    )

    apply_recognition(data)

    assert title.type_id == "title"
    assert "document-type-title-suffix" in data.recognition_diagnostics["paragraphs"][0][
        "title_context_evidence"
    ]

def test_front_report_briefing_suffix_supports_multiline_title_group() -> None:
    first = _paragraph("某委员会", "body", 0, style_name="DCT-Body")
    continuation = _paragraph(
        "关于推进重点工作整改情况汇报",
        "body",
        1,
        style_name="DCT-Body",
    )
    data = _document(
        first,
        continuation,
        _paragraph("某委员会：", "body", 2),
        _paragraph("现将有关事项说明如下，请结合实际抓好落实。", "body", 3),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs[:4]] == [
        "title",
        "title_cont",
        "addressing",
        "body",
    ]
    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert "following-document-type-title" in diagnostics[0]["title_context_evidence"]
    assert "document-type-title-suffix" in diagnostics[1]["title_context_evidence"]


def test_multiline_main_title_closes_before_date_role_and_recipient() -> None:
    data = _document(
        _paragraph("关于推进重点工作的通知", "body", 0, alignment="CENTER"),
        _paragraph("阶段性安排", "body", 1, alignment="CENTER"),
        _paragraph("重点任务说明", "body", 2, alignment="CENTER"),
        _paragraph("2026年8月24日", "body", 3, alignment="CENTER"),
        _paragraph("韩双林", "body", 4, alignment="CENTER"),
        _paragraph("各位委员、各位同志：", "body", 5),
        _paragraph("现将有关情况说明如下，请结合实际抓好落实。", "body", 6),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs] == [
        "title",
        "title_cont",
        "title_cont",
        "date_line",
        "role_name",
        "addressing",
        "body",
    ]


def test_main_title_cannot_restart_after_front_metadata() -> None:
    reopened = _paragraph(
        "补充说明",
        "title",
        2,
        alignment="CENTER",
        classification_kind="main_title",
        classification_confidence=0.95,
    )
    data = _document(
        _paragraph("关于推进重点工作的通知", "title", 0, alignment="CENTER"),
        _paragraph("2026年8月24日", "date_line", 1, alignment="CENTER"),
        reopened,
        _paragraph("各有关单位：", "addressing", 3),
        _paragraph("现将有关情况说明如下，请结合实际抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert reopened.type_id != "title"
    trace = data.recognition_diagnostics["candidate_trace"][2]
    assert "main_title" in [item["type"] for item in trace["vetoed_candidates"]]


def test_title2_remains_available_after_body_and_does_not_close_title2_sequence() -> None:
    data = _document(
        _paragraph("关于推进重点工作的通知", "title", 0, alignment="CENTER"),
        _paragraph("正文已经开始并完整说明相关工作安排。", "body", 1),
        _paragraph(
            "第一部分",
            "title2",
            2,
            classification_kind="title2",
            classification_confidence=0.95,
        ),
        _paragraph("第一部分正文继续说明具体任务。", "body", 3),
        _paragraph(
            "第二部分",
            "title2",
            4,
            classification_kind="title2",
            classification_confidence=0.95,
        ),
        _paragraph("第二部分正文继续说明具体任务。", "body", 5),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs] == [
        "title",
        "body",
        "title2",
        "body",
        "title2",
        "body",
    ]


def test_front_role_name_uses_neighbors_across_empty_and_zero_width_lines() -> None:
    role = _paragraph("韩双林", "body", 5, alignment="CENTER")
    data = _document(
        _paragraph("关于推进重点工作的通知", "title", 0, alignment="CENTER"),
        _paragraph("2026年8月24日", "date_line", 1, alignment="CENTER"),
        _paragraph("", "body", 2),
        _paragraph("\u200b", "body", 3),
        role,
        _paragraph("\ufeff", "body", 6),
        _paragraph("", "body", 7),
        _paragraph("各位委员、各位同志：", "addressing", 8),
        _paragraph("现将有关情况说明如下，请结合实际抓好落实。", "body", 9),
    )

    apply_recognition(data)

    assert role.type_id == "role_name"


@pytest.mark.parametrize("date_text", ["（2026年8月27日）", "2026年8月27日"])
def test_front_meeting_date_and_context_are_distinct_from_title(date_text: str) -> None:
    data = _document(
        _paragraph("年度重点工作报告", "body", 0, alignment="CENTER"),
        _paragraph(date_text, "body", 1, alignment="CENTER"),
        _paragraph("在全市重点工作会议上", "body", 2, alignment="CENTER"),
        _paragraph("各有关单位：", "body", 3),
        _paragraph("现将有关工作情况报告如下。", "body", 4),
    )

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs] == [
        "title",
        "meeting_title_meta",
        "meeting_title_meta",
        "addressing",
        "body",
    ]
    context = data.recognition_diagnostics["document_context"]
    assert context["front_metadata"] == [
        {"position": 1, "kind": "meeting_title_meta"},
        {"position": 2, "kind": "meeting_title_meta"},
    ]


def test_front_prosecutor_role_name_uses_generic_role_shape() -> None:
    role = _paragraph(
        "人民检察院副检察长  李测试",
        "body",
        3,
        alignment="CENTER",
    )
    data = _document(
        _paragraph("人民检察院工作报告", "title", 0, alignment="CENTER"),
        _paragraph("2026年8月25日", "date_line", 1, alignment="CENTER"),
        _paragraph("在市人民代表大会会议上", "meeting_title_meta", 2, alignment="CENTER"),
        role,
        _paragraph("各位代表：", "addressing", 4),
    )

    apply_recognition(data)

    assert role.type_id == "role_name"

def test_short_body_before_recipient_is_rechecked_as_title() -> None:
    title = _paragraph("基层治理重点工作安排", "body", 0, style_name="DCT-Body")
    data = _document(
        title,
        _paragraph("各有关单位：", "body", 1),
        _paragraph("现将有关事项说明如下，请结合实际抓好落实。", "body", 2),
    )

    apply_recognition(data)

    assert title.type_id == "title"
    assert "pre-recipient-title-context" in data.recognition_diagnostics["paragraphs"][0][
        "title_context_evidence"
    ]

def test_front_document_type_words_inside_prose_do_not_promote_title() -> None:
    first = _paragraph("根据有关通知要求开展自查", "body", 0)
    second = _paragraph("现将有关情况报告如下，供审阅。", "body", 1)
    data = _document(first, second)

    apply_recognition(data)

    assert [first.type_id, second.type_id] == ["body", "body"]

def test_front_short_lines_are_rechecked_as_title_before_body_and_first_heading() -> None:
    first = _paragraph("重点工作推进情况", "body", 0, style_name="DCT-Body")
    continuation = _paragraph("阶段性安排", "body", 1, style_name="DCT-Body")
    body = _paragraph(
        "今年以来，各项重点工作有序推进，现将有关情况说明如下。",
        "body",
        2,
    )
    heading = _paragraph("一、下一阶段重点任务", "heading1", 3)
    data = _document(first, continuation, body, heading)

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs] == [
        "title",
        "title_cont",
        "body",
        "heading1",
    ]
    diagnostics = data.recognition_diagnostics["paragraphs"]
    assert "following-body-first-heading" in diagnostics[0]["title_context_evidence"]
    assert "following-body-first-heading" in diagnostics[1]["title_context_evidence"]

def test_front_short_lines_without_following_first_heading_remain_body() -> None:
    first = _paragraph("重点工作推进情况", "body", 0, style_name="DCT-Body")
    continuation = _paragraph("阶段性安排", "body", 1, style_name="DCT-Body")
    body = _paragraph(
        "今年以来，各项重点工作有序推进，现将有关情况说明如下。",
        "body",
        2,
    )
    data = _document(first, continuation, body)

    apply_recognition(data)

    assert [item.type_id for item in data.paragraphs] == ["body", "body", "body"]

def test_front_scan_uses_soft_threshold_until_real_structure_boundary() -> None:
    prefix = [
        _paragraph(f"联合发文机关{index}", "body", index)
        for index in range(13)
    ]
    title = _paragraph("关于推进基层治理工作的通知", "body", 13, alignment="CENTER")
    recipient = _paragraph("各有关单位：", "body", 14)
    body = _paragraph("现将有关事项通知如下，请结合实际抓好落实。", "body", 15)
    data = _document(*prefix, title, recipient, body)

    apply_recognition(data)

    assert title.type_id == "title"
    assert recipient.type_id == "addressing"
    context = data.recognition_diagnostics["document_context"]
    assert 13 in context["front_matter_positions"]
    assert context["front_scan_reason"] == "body-boundary"
    assert context["front_soft_threshold_exceeded"] is True
    assert context["front_scan_soft_threshold"] == 12
    assert context["body_start"] == 15

@pytest.mark.parametrize(
    "title_text",
    [
        "代表委员履职报告",
        "项目负责人 工作方案",
        "客户经理年度总结",
        "优秀同志 调研报告",
    ],
)
def test_front_title_with_broad_role_word_is_not_role_name(title_text):
    data = _document(
        _paragraph(title_text, "title", 0, alignment="CENTER"),
        _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
        _paragraph("各位代表、同志们：", "addressing", 2),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
    )

    apply_recognition(data)

    assert data.paragraphs[0].type_id == "title"
    context = data.recognition_diagnostics["document_context"]
    assert not any(item["position"] == 0 for item in context["front_metadata"])
    trace = data.recognition_diagnostics["candidate_trace"][0]
    assert "role_name" not in [item["type"] for item in trace["candidates"]]

@pytest.mark.parametrize(
    "value",
    [
        "人大代表履职情况 专题调研",
        "项目负责人履职情况 工作部署",
        "委员联系群众情况 调研分析",
        "主席会议制度建设 经验交流",
    ],
)
@pytest.mark.parametrize(
    "context_kind",
    ["previous_title", "following_date", "following_addressing", "heading_style", "centered"],
)
def test_front_role_words_not_adjacent_to_name_do_not_form_role_name(value, context_kind):
    candidate = _paragraph(value, "body", 1)
    paragraphs = []
    if context_kind == "previous_title":
        paragraphs.append(_paragraph("年度专题材料", "title", 0, alignment="CENTER"))
        candidate.features.alignment = "CENTER"
        paragraphs.extend([candidate, _paragraph("正文内容已经开始。", "body", 2)])
    elif context_kind == "following_date":
        candidate.features.alignment = "CENTER"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 2),
        ])
    elif context_kind == "following_addressing":
        candidate.features.alignment = "CENTER"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("各位代表、同志们：", "addressing", 1),
            _paragraph("正文内容已经开始。", "body", 2),
        ])
    elif context_kind == "heading_style":
        candidate.features.style_name = "Heading 1"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 2),
        ])
    else:
        candidate.features.alignment = "CENTER"
        candidate.features.paragraph_index = 0
        paragraphs.extend([
            candidate,
            _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 2),
        ])

    data = _document(*paragraphs)
    apply_recognition(data)

    candidate_index = data.paragraphs.index(candidate)
    assert candidate.type_id != "role_name"
    context = data.recognition_diagnostics["document_context"]
    assert not any(
        item["position"] == candidate_index and item["kind"] == "role_name"
        for item in context["front_metadata"]
    )
    trace = data.recognition_diagnostics["candidate_trace"][candidate_index]
    assert "role_name" not in [item["type"] for item in trace["candidates"]]

@pytest.mark.parametrize(
    "value",
    [
        "项目负责人 专题调研",
        "会议代表 经验交流",
        "办公室主任 工作部署",
        "分管领导 情况分析",
    ],
)
@pytest.mark.parametrize(
    "context_kind",
    [
        "previous_title",
        "following_date",
        "following_addressing",
        "previous_title_and_date",
        "centered",
        "heading_style",
    ],
)
def test_weak_four_character_suffix_is_not_role_name_without_complete_byline_context(
    value,
    context_kind,
):
    candidate = _paragraph(value, "body", 1)
    if context_kind == "previous_title":
        paragraphs = [
            _paragraph("年度专题材料", "title", 0, alignment="CENTER"),
            candidate,
            _paragraph("正文内容已经开始。", "body", 2),
        ]
    elif context_kind == "following_date":
        candidate.features.paragraph_index = 0
        paragraphs = [
            candidate,
            _paragraph("2026年8月27日", "date_line", 1, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 2),
        ]
    elif context_kind == "following_addressing":
        candidate.features.paragraph_index = 0
        paragraphs = [
            candidate,
            _paragraph("各位代表、同志们：", "addressing", 1),
            _paragraph("正文内容已经开始。", "body", 2),
        ]
    elif context_kind == "previous_title_and_date":
        paragraphs = [
            _paragraph("年度专题材料", "title", 0, alignment="CENTER"),
            candidate,
            _paragraph("2026年8月27日", "date_line", 2, alignment="CENTER"),
            _paragraph("正文内容已经开始。", "body", 3),
        ]
    elif context_kind == "centered":
        candidate.features.alignment = "CENTER"
        candidate.features.paragraph_index = 0
        paragraphs = [candidate, _paragraph("正文内容已经开始。", "body", 1)]
    else:
        candidate.features.style_name = "Heading 1"
        candidate.features.paragraph_index = 0
        paragraphs = [candidate, _paragraph("正文内容已经开始。", "body", 1)]

    data = _document(*paragraphs)
    apply_recognition(data)

    candidate_index = data.paragraphs.index(candidate)
    assert candidate.type_id != "role_name"
    assert not any(
        item["position"] == candidate_index and item["kind"] == "role_name"
        for item in data.recognition_diagnostics["document_context"]["front_metadata"]
    )

def test_weak_four_character_person_name_requires_complete_centered_byline_context():
    candidate = _paragraph("办公室主任 林小测试", "body", 1, alignment="CENTER")
    data = _document(
        _paragraph("年度重点工作会议讲话", "title", 0, alignment="CENTER"),
        candidate,
        _paragraph("2026年8月27日", "date_line", 2, alignment="CENTER"),
        _paragraph("各位代表、同志们：", "addressing", 3),
    )

    apply_recognition(data)

    assert candidate.type_id == "role_name"

@pytest.mark.parametrize(
    ("value", "following_text", "following_type"),
    [
        ("专题调研", "2026年8月3日", "date_line"),
        ("工作报告", "2026年8月3日", "date_line"),
        ("情况分析", "各位代表、同志们：", "addressing"),
        ("工作安排", "各位代表、同志们：", "addressing"),
    ],
)
def test_bare_four_character_title_continuation_does_not_become_role_name(
    value,
    following_text,
    following_type,
):
    candidate = _paragraph(value, "title_cont", 1, alignment="CENTER")
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        candidate,
        _paragraph(following_text, following_type, 2),
        _paragraph("正文内容已经开始。", "body", 3),
    )

    apply_recognition(data)

    assert candidate.type_id in {"title", "title_cont", "body"}
    assert candidate.type_id != "role_name"
    context = data.recognition_diagnostics["document_context"]
    assert not any(
        item["position"] == 1 and item["kind"] == "role_name"
        for item in context["front_metadata"]
    )
    trace = data.recognition_diagnostics["candidate_trace"][1]
    assert not any(
        item["type"] == "role_name"
        and (
            item.get("source") == "front-metadata"
            or item.get("provider", "").startswith("front-metadata:")
        )
        for item in trace["candidates"]
    )

@pytest.mark.parametrize(
    ("value", "centered", "expected_role"),
    [
        ("甲乙", False, True),
        ("甲乙丙", False, True),
        ("甲乙丙丁", True, True),
        ("甲乙丙丁", False, False),
        ("欧阳甲乙", False, True),
        ("甲·乙", False, False),
        ("XXX", False, True),
    ],
)
def test_bare_person_name_fallback_uses_shared_shape_strength(
    value,
    centered,
    expected_role,
):
    candidate = _paragraph(
        value,
        "body",
        1,
        alignment="CENTER" if centered else "",
    )
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        candidate,
        _paragraph("2026年8月3日", "date_line", 2, alignment="CENTER"),
        _paragraph("正文内容已经开始。", "body", 3),
    )

    apply_recognition(data)

    assert (candidate.type_id == "role_name") is expected_role
    context_roles = {
        item["position"]
        for item in data.recognition_diagnostics["document_context"]["front_metadata"]
        if item["kind"] == "role_name"
    }
    assert (1 in context_roles) is expected_role


@pytest.mark.parametrize("value", ["甲乙", "甲乙丙", "甲乙丙丁"])
def test_bare_person_name_uses_title_to_addressing_front_window(value: str) -> None:
    candidate = _paragraph(value, "body", 1)
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        candidate,
        _paragraph("各位代表、同志们：", "addressing", 2),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
    )

    apply_recognition(data)

    assert candidate.type_id == "role_name"
    assert any(
        item["position"] == 1 and item["kind"] == "role_name"
        for item in data.recognition_diagnostics["document_context"]["front_metadata"]
    )


@pytest.mark.parametrize("value", ["甲·乙", "甲，乙", "甲乙。", "甲/乙"])
def test_title_to_addressing_bare_person_name_rejects_punctuation(value: str) -> None:
    candidate = _paragraph(value, "body", 1)
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        candidate,
        _paragraph("各位代表、同志们：", "addressing", 2),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
    )

    apply_recognition(data)

    assert candidate.type_id != "role_name"


@pytest.mark.parametrize("value", ["甲乙", "甲乙丙", "甲乙丙丁"])
def test_bare_person_name_after_title_before_body_without_addressing(
    value: str,
) -> None:
    candidate = _paragraph(value, "body", 2)
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        _paragraph("2026年8月24日", "date_line", 1, alignment="CENTER"),
        candidate,
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
    )

    apply_recognition(data)

    assert candidate.type_id == "role_name"


def test_body_closes_bare_person_name_window() -> None:
    candidate = _paragraph("甲乙丙", "body", 2)
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        _paragraph("这是普通正文说明，内容已经完整结束。", "body", 1),
        candidate,
        _paragraph("各位代表、同志们：", "addressing", 3),
    )

    apply_recognition(data)

    assert candidate.type_id != "role_name"


def test_production_recognition_promotes_name_after_title_before_body(
    tmp_path,
) -> None:
    source = tmp_path / "title-date-name-body.docx"
    document = Document()
    document.add_paragraph("年度工作报告", style="Title")
    document.add_paragraph("2026年8月24日")
    document.add_paragraph("甲乙丙")
    document.add_paragraph("现将有关事项说明如下，请认真抓好落实。")
    document.save(source)

    plan = recognize_docx(
        source,
        processing_mode="structural",
        recognition_mode="authoritative",
        include_text=True,
    )

    assert [block.type_id for block in plan.blocks] == [
        "title",
        "date_line",
        "role_name",
        "body",
    ]


def test_bare_short_line_after_addressing_is_not_role_name() -> None:
    candidate = _paragraph("甲乙丙丁", "body", 2)
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        _paragraph("各位代表、同志们：", "addressing", 1),
        candidate,
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 3),
    )

    apply_recognition(data)

    assert candidate.type_id != "role_name"


def test_later_addressing_does_not_reopen_person_name_window() -> None:
    candidate = _paragraph("甲乙丙丁", "body", 2)
    data = _document(
        _paragraph("主标题", "title", 0, alignment="CENTER"),
        _paragraph("各位代表、同志们：", "addressing", 1),
        candidate,
        _paragraph("同志们：", "addressing", 3),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert candidate.type_id != "role_name"


def test_production_recognition_uses_title_to_addressing_for_four_char_name(
    tmp_path,
) -> None:
    source = tmp_path / "title-name-addressing.docx"
    document = Document()
    document.add_paragraph("年度工作报告", style="Title")
    document.add_paragraph("甲乙丙丁")
    document.add_paragraph("各位代表、同志们：")
    document.add_paragraph("现将有关事项说明如下，请认真抓好落实。")
    document.save(source)

    plan = recognize_docx(
        source,
        processing_mode="structural",
        recognition_mode="authoritative",
        include_text=True,
    )

    assert [block.type_id for block in plan.blocks] == [
        "title",
        "role_name",
        "addressing",
        "body",
    ]


def test_compound_surname_four_character_name_remains_strong_with_title_anchor():
    candidate = _paragraph("办公室主任 欧阳测试", "body", 1, alignment="CENTER")
    data = _document(
        _paragraph("年度重点工作会议讲话", "title", 0, alignment="CENTER"),
        candidate,
        _paragraph("正文内容已经开始。", "body", 2),
    )

    apply_recognition(data)

    assert candidate.type_id == "role_name"

def test_compact_role_name_prefers_strong_short_name_over_weak_long_suffix():
    candidate = _paragraph("办公室党组书记、主席张三", "body", 1)
    data = _document(
        _paragraph("年度重点工作会议讲话", "title", 0, alignment="CENTER"),
        candidate,
        _paragraph("2026年8月27日", "date_line", 2, alignment="CENTER"),
    )

    apply_recognition(data)

    assert candidate.type_id == "role_name"

@pytest.mark.parametrize(
    "role_text",
    [
        "办公室主任 张三",
        "办公室主任 李测试",
        "办公室主任 欧阳测试",
        "办公室主任 阿·明",
        "办公室主任　李测试",
        "党组书记、主席 欧阳测试",
        "某某机构办公室主任 张三",
        "办公室主任王测试",
        "会议代表 张三",
        "项目负责人 李测试",
    ],
)
def test_front_role_name_shape_survives_changed_names_and_roles(role_text):
    data = _document(
        _paragraph("年度重点工作会议讲话", "title", 0, alignment="CENTER"),
        _paragraph(role_text, "body", 1, alignment="CENTER"),
        _paragraph("2026年8月27日", "date_line", 2, alignment="CENTER"),
        _paragraph("各位代表、同志们：", "addressing", 3),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 4),
    )

    apply_recognition(data)

    assert data.paragraphs[1].type_id == "role_name"
    context = data.recognition_diagnostics["document_context"]
    assert {item["position"]: item["kind"] for item in context["front_metadata"]}[1] == "role_name"
    trace = data.recognition_diagnostics["candidate_trace"][1]
    assert "role_name" in [item["type"] for item in trace["candidates"]]

def test_spaced_role_name_can_use_previous_title_anchor_without_date():
    data = _document(
        _paragraph("年度重点工作会议讲话", "title", 0, alignment="CENTER"),
        _paragraph("办公室主任 李测试", "body", 1, alignment="CENTER"),
        _paragraph("现将有关事项说明如下，请认真抓好落实。", "body", 2),
    )

    apply_recognition(data)

    assert data.paragraphs[1].type_id == "role_name"
