from __future__ import annotations

import importlib.util
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import pytest


ROOT = Path(__file__).resolve().parents[1]
SCRIPT_PATH = ROOT / "scripts" / "batch_test_docx.py"
SPEC = importlib.util.spec_from_file_location("batch_test_docx", SCRIPT_PATH)
assert SPEC is not None and SPEC.loader is not None
batch_test_docx = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(batch_test_docx)


def test_template_letterhead_options_are_derived_from_reference(tmp_path: Path) -> None:
    template = tmp_path / "correct.docx"
    document = Document()
    document.add_paragraph("测试市政府文件")
    document.add_paragraph("测试办〔2026〕12号")
    document.save(template)

    options = batch_test_docx.letterhead_options_from_template(template)

    assert options is not None
    assert options["enabled"] is True
    assert options["agencies"][0]["name"] == "测试市政府"
    assert options["document_number"] == {"agency_code": "测试办", "year": 2026, "sequence": 12}


def test_template_letterhead_options_require_mark_and_dispatch_number(tmp_path: Path) -> None:
    template = tmp_path / "incomplete.docx"
    document = Document()
    document.add_paragraph("关于进一步推进测试工作的通知")
    document.save(template)

    assert batch_test_docx.letterhead_options_from_template(template) is None


def test_batch_defaults_to_frontend_structural_mode() -> None:
    assert batch_test_docx.parse_args([]).strict_preservation is False
    assert batch_test_docx.parse_args([]).mode == "structural"
    assert batch_test_docx.parse_args(["--strict-preservation"]).strict_preservation is True
    assert batch_test_docx.parse_args(["--mode", "normalize"]).mode == "normalize"


def test_special_output_directory_is_fixed_to_long_term_regression_location() -> None:
    assert batch_test_docx._validate_special_output_dir(batch_test_docx.SPECIAL_OUTPUT_DIR) == batch_test_docx.SPECIAL_OUTPUT_DIR.resolve()


def test_special_output_directory_accepts_a_fresh_batch_below_fixed_root(tmp_path, monkeypatch) -> None:
    special_root = tmp_path / "special"
    monkeypatch.setattr(batch_test_docx, "SPECIAL_OUTPUT_DIR", special_root)
    batch_dir = special_root / "fresh-batch"

    assert batch_test_docx._validate_special_output_dir(batch_dir) == batch_dir.resolve()


def test_batch_report_marks_visual_rendering_as_not_run() -> None:
    status = batch_test_docx.visual_rendering_not_run()

    assert status["executed"] is False
    assert "未执行视觉渲染检查" in str(status["reason"])


def test_source_heading_cue_audit_detects_heading_lost_as_body() -> None:
    features = SimpleNamespace(
        segment_numbering_features="@lvl_0",
        numbering_prefix="@lvl_0",
        segment_bold_char_ratio=1.0,
        source_physical_paragraph_index=7,
    )
    source = SimpleNamespace(paragraphs=[
        SimpleNamespace(text="缺失可见编号的短标题", type_id="heading2", features=features),
    ])
    output = SimpleNamespace(paragraphs=[
        SimpleNamespace(text="缺失可见编号的短标题", type_id="body"),
    ])

    audit = batch_test_docx.source_heading_cue_audit(source, output)

    assert audit["cue_count"] == 1
    assert audit["mismatch_count"] == 1
    assert audit["mismatches"][0]["expected_type"] == "heading2"
    assert "缺失可见编号的短标题" not in str(audit["mismatches"])


def test_source_heading_cue_audit_accepts_preserved_heading_and_ignores_long_prose() -> None:
    short_features = SimpleNamespace(
        segment_numbering_features="@lvl_0",
        numbering_prefix="@lvl_0",
        segment_bold_char_ratio=1.0,
        source_physical_paragraph_index=3,
    )
    long_features = SimpleNamespace(
        segment_numbering_features="@lvl_0",
        numbering_prefix="@lvl_0",
        segment_bold_char_ratio=1.0,
        source_physical_paragraph_index=4,
    )
    short = "短标题"
    long_text = "这是继承列表属性但长度明显超过标题边界的正文内容，不能把它纳入源标题线索审计范围并产生误报。"
    source = SimpleNamespace(paragraphs=[
        SimpleNamespace(text=short, type_id="heading2", features=short_features),
        SimpleNamespace(text=long_text, type_id="body", features=long_features),
    ])
    output = SimpleNamespace(paragraphs=[
        SimpleNamespace(text=f"（一）{short}", type_id="heading2"),
        SimpleNamespace(text=long_text, type_id="body"),
    ])

    audit = batch_test_docx.source_heading_cue_audit(source, output)

    assert audit == {"cue_count": 1, "mismatch_count": 0, "mismatches": []}


def test_native_numbering_heading_audit_is_text_free() -> None:
    text = "自动编号标题"
    features = SimpleNamespace(
        native_numbering=object(),
        source_physical_paragraph_index=8,
    )
    source = SimpleNamespace(
        paragraphs=[SimpleNamespace(text=text, type_id="heading2", features=features)],
        recognition_diagnostics={
            "paragraphs": [
                {
                    "paragraph_index": 0,
                    "evidence_summary": [
                        "numbered-heading-level-2",
                        "native-numbering-template",
                    ],
                }
            ]
        },
    )

    output = SimpleNamespace(
        paragraphs=[SimpleNamespace(text=f"（一）{text}", type_id="heading2")]
    )

    audit = batch_test_docx.native_numbering_heading_audit(source, output)

    assert audit["clue_count"] == 1
    assert audit["unpreserved_count"] == 0
    assert audit["details"][0]["type"] == "heading2"
    assert text not in str(audit)


def test_legacy_source_heading_audit_skips_parsed_native_numbering() -> None:
    features = SimpleNamespace(
        native_numbering=object(),
        segment_numbering_features="@lvl_0",
        numbering_prefix="@lvl_0",
        segment_bold_char_ratio=1.0,
        source_physical_paragraph_index=2,
    )
    source = SimpleNamespace(
        paragraphs=[SimpleNamespace(text="原生一级标题", features=features)]
    )
    output = SimpleNamespace(
        paragraphs=[SimpleNamespace(text="原生一级标题", type_id="heading1")]
    )

    assert batch_test_docx.source_heading_cue_audit(source, output) == {
        "cue_count": 0,
        "mismatch_count": 0,
        "mismatches": [],
    }


def test_signature_continuity_audit_detects_attachment_between_org_and_date() -> None:
    output = SimpleNamespace(paragraphs=[
        SimpleNamespace(text="测试单位", type_id="sign_org"),
        SimpleNamespace(text="附件：测试材料", type_id="attachment_note"),
        SimpleNamespace(text="2026年7月31日", type_id="sign_date"),
    ])

    audit = batch_test_docx.signature_continuity_audit(output)

    assert audit["issue_count"] == 1
    assert audit["issues"][0]["intervening_types"] == ["attachment_note"]
    assert "测试单位" not in str(audit["issues"])


def test_signature_continuity_audit_accepts_attachment_before_signature_pair() -> None:
    output = SimpleNamespace(paragraphs=[
        SimpleNamespace(text="附件：测试材料", type_id="attachment_note"),
        SimpleNamespace(text="测试单位", type_id="sign_org"),
        SimpleNamespace(text="2026年7月31日", type_id="sign_date"),
    ])

    assert batch_test_docx.signature_continuity_audit(output) == {
        "issue_count": 0,
        "issues": [],
    }


def test_signature_continuity_audit_detects_org_embedded_in_prior_body() -> None:
    output = SimpleNamespace(paragraphs=[
        SimpleNamespace(
            text="这是正文内容。\n\n某区工作办公室",
            original_text="这是正文内容。\n\n某区工作办公室",
            type_id="body",
        ),
        SimpleNamespace(text="附件：测试材料", type_id="attachment_note"),
        SimpleNamespace(text="2026年7月31日", type_id="sign_date"),
    ])

    audit = batch_test_docx.signature_continuity_audit(output)

    assert audit["issue_count"] == 1
    assert audit["issues"][0]["issue"] == "signature_org_embedded_in_body"
    assert "某区工作办公室" not in str(audit["issues"])


def _recognition_item(index: int, text: str, type_id: str = "body") -> dict:
    item = {
        "index": index,
        "_text": text,
        "text_hash": batch_test_docx.text_hash(text),
        "length": len(text),
        "type": type_id,
        "recognition_type": type_id,
        "section": "",
        "heading_level": batch_test_docx._heading_level(type_id),
    }
    item["region"] = batch_test_docx._recognition_region(item, index)
    return item


def _recognition_snapshot(items: list[dict]) -> dict:
    return {"mode": "document", "paragraphs": items, "tables": 0, "blocks": 0, "images": 0}


def test_structural_alignment_does_not_cascade_after_generated_letterhead() -> None:
    actual = _recognition_snapshot([
        _recognition_item(0, "测试市政府文件", "__letterhead__"),
        _recognition_item(1, "测试办〔2026〕12号", "dispatch_number"),
        _recognition_item(2, "正文甲"),
        _recognition_item(3, "正文乙"),
    ])
    expected = _recognition_snapshot([
        _recognition_item(0, "正文甲"),
        _recognition_item(1, "正文乙"),
    ])

    differences, stats = batch_test_docx.compare_snapshots(
        actual, expected, source="recognition", fields=("type",), document_fields=(), report_unmatched=True,
    )

    assert stats["matched_pairs"] == 2
    additions = [item for item in differences if item["category"] == "output_addition"]
    assert len(additions) == 2
    assert all(item["expected_change"] for item in additions)
    assert not any(item["actual_paragraph_index"] in {2, 3} for item in additions)


def test_structural_alignment_preserves_duplicate_paragraph_order() -> None:
    actual = _recognition_snapshot([
        _recognition_item(0, "重复正文"), _recognition_item(1, "重复正文"), _recognition_item(2, "结束"),
    ])
    expected = _recognition_snapshot([
        _recognition_item(0, "重复正文"), _recognition_item(1, "重复正文"), _recognition_item(2, "结束"),
    ])

    _, stats = batch_test_docx.compare_snapshots(
        actual, expected, source="recognition", fields=("type",), document_fields=(), report_unmatched=True,
    )

    assert stats["matched_pairs"] == 3
    assert stats["output_additions"] == 0
    assert stats["template_missing"] == 0


def test_expected_date_and_heading_normalization_are_not_real_structure_errors() -> None:
    actual = _recognition_snapshot([
        _recognition_item(0, "一、主要任务", "heading1"),
        _recognition_item(1, "2025年10月15日", "sign_date"),
    ])
    expected = _recognition_snapshot([
        _recognition_item(0, "一、主要任务。", "heading1"),
        _recognition_item(1, "二〇二五年十月十五日", "sign_date"),
    ])

    differences, stats = batch_test_docx.compare_snapshots(
        actual, expected, source="recognition", fields=("type",), document_fields=(), report_unmatched=True,
    )

    normalizations = [item for item in differences if item["category"] == "expected_normalization"]
    assert stats["matched_pairs"] == 2
    assert len(normalizations) == 2
    assert all(item["expected_change"] for item in normalizations)


@pytest.mark.parametrize(
    ("actual_text", "expected_text", "type_id"),
    [
        ("三、共同标题", "一、共同标题。", "heading1"),
        ("（四）共同标题", "（一）共同标题", "heading2"),
        ("7.共同标题", "1.共同标题", "heading3"),
        ("（9）共同标题", "（1）共同标题", "heading4"),
    ],
)
def test_heading_alignment_uses_content_before_comparing_numbering(
    actual_text: str, expected_text: str, type_id: str,
) -> None:
    actual = _recognition_snapshot([_recognition_item(0, actual_text, type_id)])
    expected = _recognition_snapshot([_recognition_item(0, expected_text, type_id)])

    differences, stats = batch_test_docx.compare_snapshots(
        actual, expected, source="recognition", fields=("type",), document_fields=(), report_unmatched=True,
    )

    assert stats["matched_pairs"] == 1
    assert stats["output_additions"] == 0
    assert stats["template_missing"] == 0
    normalization = next(item for item in differences if item["category"] == "expected_normalization")
    assert normalization["match_reason"] == "heading_number_normalization"
    assert normalization["difference_origin"] == "expected_normalization"


def test_unmatched_content_is_reported_as_real_p1_problem() -> None:
    actual = _recognition_snapshot([_recognition_item(0, "保留段落")])
    expected = _recognition_snapshot([_recognition_item(0, "保留段落"), _recognition_item(1, "缺失段落")])

    differences, _ = batch_test_docx.compare_snapshots(
        actual, expected, source="recognition", fields=("type",), document_fields=(), report_unmatched=True,
    )
    missing = next(item for item in differences if item["category"] == "template_missing")

    assert missing["expected_change"] is False
    assert batch_test_docx.severity(missing["category"], expected_change=missing["expected_change"]) == "P1"


def test_input_fixture_content_difference_is_not_reported_as_output_p1() -> None:
    source = _recognition_snapshot([_recognition_item(0, "测试输入中的错误序号")])
    actual = {
        "recognition": _recognition_snapshot([_recognition_item(0, "测试输入中的错误序号")]),
        "physical": {"paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0},
    }
    expected = {
        "recognition": _recognition_snapshot([_recognition_item(0, "正确模板序号")]),
        "physical": {"paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0},
    }

    differences, summary = batch_test_docx.compare_documents(actual, expected, source_recognition=source)

    fixture_differences = [item for item in differences if item["expected_reason"] == "input_fixture_difference"]
    assert summary["input_fixture_differences"] == 2
    assert len(fixture_differences) == 2
    assert summary["unexpected_differences"] == 0


def test_source_order_difference_is_separate_from_output_regression() -> None:
    source = _recognition_snapshot([
        _recognition_item(0, "二、乙标题", "heading1"),
        _recognition_item(1, "一、甲标题", "heading1"),
    ])
    actual = {
        "recognition": _recognition_snapshot([
            _recognition_item(0, "一、乙标题", "heading1"),
            _recognition_item(1, "二、甲标题", "heading1"),
        ]),
        "physical": {"paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0},
    }
    expected = {
        "recognition": _recognition_snapshot([
            _recognition_item(0, "一、甲标题", "heading1"),
            _recognition_item(1, "二、乙标题", "heading1"),
        ]),
        "physical": {"paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0},
    }

    differences, summary = batch_test_docx.compare_documents(actual, expected, source_recognition=source)

    order_differences = [item for item in differences if item["expected_reason"] == "source_order_difference"]
    assert order_differences
    assert summary["source_order_differences"] == len(order_differences)
    assert summary["unexpected_differences"] == 0


def test_output_loss_not_visible_in_template_is_still_reported() -> None:
    source = _recognition_snapshot([
        _recognition_item(0, "模板共有正文"),
        _recognition_item(1, "仅原稿存在正文"),
    ])
    actual = {
        "recognition": _recognition_snapshot([_recognition_item(0, "模板共有正文")]),
        "physical": {"paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0},
    }
    expected = {
        "recognition": _recognition_snapshot([_recognition_item(0, "模板共有正文")]),
        "physical": {"paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0},
    }

    differences, summary = batch_test_docx.compare_documents(actual, expected, source_recognition=source)

    loss = next(item for item in differences if item["category"] == "source_text_loss")
    assert loss["rule_basis"] == "source_text_preservation"
    assert loss["severity"] == "P1"
    assert summary["unexpected_differences"] == 1


def test_normalize_mode_marks_only_character_conserving_changes_as_expected() -> None:
    source = _recognition_snapshot([
        _recognition_item(0, "一、测试标题。这里是正文", "heading1"),
        _recognition_item(1, "2025年十月15日", "sign_date"),
        _recognition_item(2, "附件1", "attachment_page_mark"),
    ])
    actual_recognition = _recognition_snapshot([
        _recognition_item(0, "一、测试标题", "heading1"),
        _recognition_item(1, "这里是正文"),
        _recognition_item(2, "2025年10月15日", "sign_date"),
        _recognition_item(3, "附件 1", "attachment_page_mark"),
    ])
    empty_physical = {
        "paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0,
        "headers": 0, "footers": 0,
    }

    differences, summary = batch_test_docx.compare_documents(
        {"recognition": actual_recognition, "physical": empty_physical},
        {"recognition": actual_recognition, "physical": empty_physical},
        source_recognition=source,
        processing_strategy="normalize",
    )

    mode_differences = [
        item for item in differences
        if item.get("expected_reason") == "expected_mode_difference"
    ]
    assert mode_differences
    assert summary["expected_mode_differences"] == len(mode_differences)
    assert summary["unexpected_differences"] == 0


def test_normalize_mode_does_not_hide_real_character_loss() -> None:
    source = _recognition_snapshot([_recognition_item(0, "完整正文内容")])
    actual = _recognition_snapshot([_recognition_item(0, "完整正文")])
    empty_physical = {
        "paragraphs": [], "tables": 0, "sections": 1, "inline_shapes": 0,
        "headers": 0, "footers": 0,
    }

    differences, summary = batch_test_docx.compare_documents(
        {"recognition": actual, "physical": empty_physical},
        {"recognition": actual, "physical": empty_physical},
        source_recognition=source,
        processing_strategy="normalize",
    )

    assert any(
        item["category"] == "source_text_loss" and not item["expected_change"]
        for item in differences
    )
    assert summary["expected_mode_differences"] == 0
    assert summary["unexpected_differences"] > 0


def test_strict_inline_heading_body_alignment_is_an_expected_mode_difference() -> None:
    text = "（一）测试标题。这里是同段正文内容"
    recognition = _recognition_snapshot([_recognition_item(0, text, "heading2")])
    actual_physical = {
        "paragraphs": [{
            "index": 0, "_text": text, "text_hash": batch_test_docx.text_hash(text),
            "style": "标题 2", "style_id": "DCT-Heading2", "alignment": "LEFT (0)",
            "region": "body",
        }],
        "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0,
    }
    expected_physical = {
        "paragraphs": [{
            "index": 0, "_text": text, "text_hash": batch_test_docx.text_hash(text),
            "style": "标题 2", "style_id": "template-heading2", "alignment": "JUSTIFY (3)",
            "region": "body",
        }],
        "tables": 0, "sections": 1, "inline_shapes": 0, "headers": 0, "footers": 0,
    }

    differences, summary = batch_test_docx.compare_documents(
        {"recognition": recognition, "physical": actual_physical},
        {"recognition": recognition, "physical": expected_physical},
        processing_strategy="strict",
    )

    alignment = next(item for item in differences if item["category"] == "alignment")
    assert alignment["expected_reason"] == "expected_mode_difference"
    assert alignment["severity"] == "P3"
    assert summary["unexpected_differences"] == 0


def test_reordered_content_does_not_match_across_the_sequence() -> None:
    actual = _recognition_snapshot([_recognition_item(0, "乙"), _recognition_item(1, "甲")])
    expected = _recognition_snapshot([_recognition_item(0, "甲"), _recognition_item(1, "乙")])

    differences, _ = batch_test_docx.compare_snapshots(
        actual, expected, source="recognition", fields=("type",), document_fields=(), report_unmatched=True,
    )

    assert any(item["category"] in {"output_addition", "template_missing"} and not item["expected_change"] for item in differences)


def _set_indent(paragraph, **values: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    for name, value in values.items():
        indent.set(qn("w:" + name), value)


def test_physical_snapshot_compares_semantic_character_indent(tmp_path: Path) -> None:
    actual_path = tmp_path / "actual.docx"
    expected_path = tmp_path / "expected.docx"
    actual_document = Document()
    actual_paragraph = actual_document.add_paragraph("测试正文")
    _set_indent(actual_paragraph, firstLineChars="200", firstLine="640")
    actual_document.save(actual_path)
    expected_document = Document()
    expected_paragraph = expected_document.add_paragraph("测试正文")
    _set_indent(expected_paragraph, firstLineChars="200", firstLine="632")
    expected_document.save(expected_path)

    differences, _stats = batch_test_docx.compare_snapshots(
        batch_test_docx.physical_snapshot(actual_path),
        batch_test_docx.physical_snapshot(expected_path),
        source="physical", fields=("first_indent",), document_fields=(), report_unmatched=False,
    )

    assert differences == []


def test_protected_caption_format_is_compared_to_source_not_template(tmp_path: Path) -> None:
    paths = {name: tmp_path / f"{name}.docx" for name in ("source", "actual", "expected")}
    for name, path in paths.items():
        document = Document()
        document.add_table(rows=1, cols=1)
        caption = document.add_paragraph()
        run = caption.add_run("表1")
        run.font.name = "Source Caption Font" if name != "expected" else "Template Font"
        run.font.size = Pt(12)
        run.font.italic = name != "expected"
        document.save(path)
    empty_recognition = _recognition_snapshot([])
    actual = {"recognition": empty_recognition, "physical": batch_test_docx.physical_snapshot(paths["actual"])}
    expected = {"recognition": empty_recognition, "physical": batch_test_docx.physical_snapshot(paths["expected"])}

    differences, summary = batch_test_docx.compare_documents(
        actual,
        expected,
        source_physical=batch_test_docx.physical_snapshot(paths["source"]),
    )

    protected = [item for item in differences if item["expected_reason"] == "protected_source_difference"]
    assert {item["category"] for item in protected} == {"font", "italic"}
    assert all(item["protected_kind"] == "table_caption" for item in protected)
    assert summary["protected_source_differences"] == 2
    assert summary["unexpected_differences"] == 0


def test_protected_caption_changed_by_output_remains_a_real_problem(tmp_path: Path) -> None:
    paths = {name: tmp_path / f"{name}.docx" for name in ("source", "actual", "expected")}
    fonts = {"source": "Source Font", "actual": "Changed Font", "expected": "Template Font"}
    for name, path in paths.items():
        document = Document()
        document.add_table(rows=1, cols=1)
        run = document.add_paragraph().add_run("表1")
        run.font.name = fonts[name]
        document.save(path)
    empty_recognition = _recognition_snapshot([])
    actual = {"recognition": empty_recognition, "physical": batch_test_docx.physical_snapshot(paths["actual"])}
    expected = {"recognition": empty_recognition, "physical": batch_test_docx.physical_snapshot(paths["expected"])}

    differences, summary = batch_test_docx.compare_documents(
        actual,
        expected,
        source_physical=batch_test_docx.physical_snapshot(paths["source"]),
    )

    font_difference = next(item for item in differences if item["category"] == "font")
    assert font_difference["expected_change"] is False
    assert font_difference["severity"] == "P2"
    assert summary["protected_source_differences"] == 0
    assert summary["unexpected_differences"] == 1


def test_render_sample_selection_is_deterministic_and_risk_first() -> None:
    records = [
        {"编号": "001", "成功": True, "模板对齐": {"real_issue_counts": {"P1": 0}}, "结构诊断": {}},
        {"编号": "002", "成功": True, "模板对齐": {"real_issue_counts": {"P1": 2}}, "结构诊断": {}},
        {"编号": "003", "成功": True, "模板对齐": {"real_issue_counts": {"P2": 5}}, "结构诊断": {}},
        {"编号": "004", "成功": True, "模板对齐": {"real_issue_counts": {}}, "结构诊断": {}},
    ]

    selected = batch_test_docx.select_render_samples(records, 2)

    assert [item["编号"] for item in selected] == ["002", "003"]


def test_render_sample_selection_prioritizes_protected_source_differences() -> None:
    records = [
        {"编号": "001", "成功": True, "模板对齐": {"real_issue_counts": {}, "expected_difference_counts": {}}, "结构诊断": {}, "表格数": 1},
        {"编号": "016", "成功": True, "模板对齐": {"real_issue_counts": {}, "expected_difference_counts": {"protected_source_difference": 1}}, "结构诊断": {}, "表格数": 1},
        {"编号": "035", "成功": True, "模板对齐": {"real_issue_counts": {}, "expected_difference_counts": {"protected_source_difference": 1}}, "结构诊断": {}, "表格数": 1},
    ]

    selected = batch_test_docx.select_render_samples(records, 2)

    assert [item["编号"] for item in selected] == ["016", "035"]


@pytest.mark.parametrize(
    ("text", "ink_ratio", "expected_status", "suspected"),
    [
        ("— 17 —", 0.001, "empty_page", True),
        ("区政协办\n2025年10月15日\n— 17 —", 0.002, "sparse_page", True),
        ("附件2\n测试正文\n— 19 —", 0.001, "attachment_title_page", False),
        ("正常正文内容足够用于页面识别。" * 6, 0.04, "normal", False),
    ],
)
def test_page_visual_status_distinguishes_empty_sparse_and_attachment_pages(
    text: str, ink_ratio: float, expected_status: str, suspected: bool,
) -> None:
    status, actual_suspected, _ = batch_test_docx._page_visual_status(text, ink_ratio)

    assert status == expected_status
    assert actual_suspected is suspected


def test_require_render_implies_render_review() -> None:
    args = batch_test_docx.parse_args(["--require-render"])

    assert args.render_review is True


def test_visual_review_reports_conversion_failure_without_leaking_content(monkeypatch, tmp_path: Path) -> None:
    output = tmp_path / "sample.docx"
    output.write_bytes(b"placeholder")
    record = {"编号": "001", "成功": True, "模板对齐": {"real_issue_counts": {}}, "结构诊断": {}}
    monkeypatch.setattr(batch_test_docx, "_rendering_dependencies", lambda: ("soffice", object()))
    monkeypatch.setattr(batch_test_docx, "_render_document", lambda *args, **kwargs: (_ for _ in ()).throw(RuntimeError("conversion failed")))

    status = batch_test_docx.run_visual_review(
        [record], {"001": output}, [], {}, [],
        standard_root=tmp_path / "standard", special_root=tmp_path / "special", sample_size=1,
    )

    assert status["executed"] is True
    assert status["failures"][0]["error"] == "RuntimeError: conversion failed"
