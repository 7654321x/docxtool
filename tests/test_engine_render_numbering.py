from docx import Document

from docxtool.document.engine.render_numbering import NumberingCounter, apply_numbering
from docxtool.document.style_config import StyleRule


def test_numbering_counter_advances_and_resets_child_levels() -> None:
    counter = NumberingCounter()

    counter.advance("heading1")
    counter.advance("heading2")
    counter.advance("heading3")
    counter.advance("heading4")
    assert (counter.a, counter.b, counter.c, counter.d) == (1, 1, 1, 1)

    counter.advance("heading2")
    assert (counter.a, counter.b, counter.c, counter.d) == (1, 2, 0, 0)

    counter.advance("heading1")
    assert (counter.a, counter.b, counter.c, counter.d) == (2, 0, 0, 0)


def test_numbering_counter_renders_chinese_and_arabic_patterns() -> None:
    counter = NumberingCounter()
    counter.a = 1
    counter.b = 2
    counter.c = 3
    counter.d = 4

    assert counter.render("{a}、", "heading1") == "一、"
    assert counter.render("（{b}）", "heading2") == "（二）"
    assert counter.render("{c}.", "heading3") == "3."
    assert counter.render("（{d}）", "heading4") == "（4）"
    assert counter.render("- 1 -", "page_number") == "- 1 -"


def test_apply_numbering_inserts_formatted_prefix_before_existing_text() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("提高政治站位")
    rule = StyleRule.default_for_row(1)
    counter = NumberingCounter()
    counter.advance("heading1")

    apply_numbering(paragraph, rule, counter)

    assert paragraph.text == "一、提高政治站位"
    assert paragraph.runs[0].text == "一、"
    assert paragraph.runs[0].font.size.pt == rule.font_size_pt
    assert paragraph.runs[0].font.bold is rule.bold


def test_apply_numbering_keeps_existing_prefix_once() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("一、提高政治站位")
    rule = StyleRule.default_for_row(1)
    counter = NumberingCounter()
    counter.advance("heading1")

    apply_numbering(paragraph, rule, counter)

    assert paragraph.text == "一、提高政治站位"
    assert [run.text for run in paragraph.runs].count("一、") == 0
