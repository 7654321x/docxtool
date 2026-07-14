from __future__ import annotations

import base64
import io
import json
import os
import subprocess
import sys
import textwrap
import time
import zipfile
from pathlib import Path
from urllib.error import HTTPError
from urllib.request import Request, urlopen

from docx import Document

from docxtool.web import app as server


STRONG_ADMIN_TOKEN = "test-admin-token-20260714-worker-config"
STRONG_PROXY_SECRET = "test-proxy-secret-20260714-worker-config"


def _runtime_env(tmp_path: Path) -> dict[str, str]:
    env = os.environ.copy()
    src_path = str(Path(__file__).resolve().parents[1] / "src")
    env["PYTHONPATH"] = src_path + os.pathsep + env.get("PYTHONPATH", "")
    env["DATABASE_PATH"] = str(tmp_path / "stats.db")
    env["LOG_DIR"] = str(tmp_path / "logs")
    env["OUTPUT_DIR"] = str(tmp_path / "outputs")
    env["RUNTIME_DIR"] = str(tmp_path / "runtime")
    env["ADMIN_TOKEN"] = STRONG_ADMIN_TOKEN
    env["PROXY_SECRET"] = STRONG_PROXY_SECRET
    env["BIND_HOST"] = "127.0.0.1"
    env["PROCESS_TIMEOUT_SECONDS"] = "30"
    return env


def _format_config_headers(config: dict) -> dict[str, str]:
    raw = json.dumps(config, ensure_ascii=False).encode("utf-8")
    encoded = base64.urlsafe_b64encode(raw).decode("ascii").rstrip("=")
    return {
        "X-Format-Config": encoded,
        "X-Format-Config-Encoding": "base64url-json",
    }


def _docx_bytes() -> bytes:
    stream = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Smoke regression body text.")
    doc.save(stream)
    return stream.getvalue()


def _json_request(url: str, headers: dict[str, str] | None = None) -> dict:
    req = Request(url, headers=headers or {})
    with urlopen(req, timeout=5) as response:
        return json.loads(response.read().decode("utf-8"))


def _upload_docx(base_url: str, payload: bytes, headers: dict[str, str]) -> dict:
    req = Request(f"{base_url}/upload", data=payload, method="PUT", headers=headers)
    with urlopen(req, timeout=10) as response:
        return json.loads(response.read().decode("utf-8"))


def _read_download(base_url: str, task_id: str, headers: dict[str, str]) -> bytes:
    req = Request(f"{base_url}/download/{task_id}", headers=headers)
    with urlopen(req, timeout=10) as response:
        return response.read()


def test_importing_web_app_does_not_delete_existing_task_input(tmp_path: Path) -> None:
    runtime_dir = tmp_path / "runtime"
    input_path = runtime_dir / "tmp" / "task-id" / "input.docx"
    input_path.parent.mkdir(parents=True)
    input_path.write_bytes(b"uploaded input")

    result = subprocess.run(
        [sys.executable, "-c", "import docxtool.web.app"],
        cwd=tmp_path,
        env=_runtime_env(tmp_path),
        text=True,
        capture_output=True,
        timeout=20,
        check=False,
    )

    assert result.returncode == 0, result.stderr
    assert input_path.exists()
    assert input_path.read_bytes() == b"uploaded input"


def test_main_runs_startup_cleanup_only_after_secret_validation(monkeypatch) -> None:
    calls: list[str] = []

    class FakeSocket:
        def setsockopt(self, *_args) -> None:
            calls.append("setsockopt")

    class FakeServer:
        socket = FakeSocket()

        def __init__(self, *_args) -> None:
            calls.append("server")

        def serve_forever(self) -> None:
            calls.append("serve_forever")
            raise KeyboardInterrupt

        def server_close(self) -> None:
            calls.append("server_close")

    monkeypatch.setattr(server.sys, "argv", ["server.py"])
    monkeypatch.setattr(server, "_validate_secrets_or_exit", lambda: calls.append("validate"))
    monkeypatch.setattr(server, "_startup_cleanup", lambda: calls.append("cleanup"))
    monkeypatch.setattr(server, "_sql_init", lambda: calls.append("sql"))
    monkeypatch.setattr(server, "_recover_inflight_tasks_on_startup", lambda: calls.append("recover"))
    monkeypatch.setattr(server, "_ensure_workers_started", lambda: calls.append("workers"))
    monkeypatch.setattr(server, "ThreadingHTTPServer", FakeServer)
    monkeypatch.setattr(server, "_startup_time_check_lines", lambda: [])

    server.main()

    assert calls[:6] == ["validate", "cleanup", "sql", "recover", "workers", "server"]


def test_spawn_worker_completes_upload_with_default_format_config(tmp_path: Path) -> None:
    port_file = tmp_path / "port.txt"
    helper = tmp_path / "run_service.py"
    helper.write_text(
        textwrap.dedent(
            f"""
            import multiprocessing as mp
            import socket
            import sys
            from http.server import ThreadingHTTPServer
            from pathlib import Path

            from docxtool.web import app as server

            def main():
                server.MAX_WORKERS = 1
                server.MAX_QUEUE = 2
                server._validate_secrets_or_exit()
                server._startup_cleanup()
                server._sql_init()
                server._recover_inflight_tasks_on_startup()
                server._ensure_workers_started()
                httpd = ThreadingHTTPServer(("127.0.0.1", 0), server.Handler)
                httpd.socket.setsockopt(socket.IPPROTO_TCP, socket.TCP_NODELAY, 1)
                Path({str(port_file)!r}).write_text(str(httpd.server_address[1]), encoding="ascii")
                sys.stdout.write("ready\\n")
                sys.stdout.flush()
                httpd.serve_forever()

            if __name__ == "__main__":
                mp.freeze_support()
                main()
            """
        ),
        encoding="utf-8",
    )

    process = subprocess.Popen(
        [sys.executable, str(helper)],
        cwd=tmp_path,
        env=_runtime_env(tmp_path),
        text=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    try:
        deadline = time.monotonic() + 15
        while not port_file.exists() and time.monotonic() < deadline:
            if process.poll() is not None:
                stdout, stderr = process.communicate(timeout=2)
                raise AssertionError(f"service exited early\nSTDOUT:\n{stdout}\nSTDERR:\n{stderr}")
            time.sleep(0.05)
        assert port_file.exists(), "service did not report a listening port"
        base_url = f"http://127.0.0.1:{port_file.read_text(encoding='ascii').strip()}"

        resource = Path(server.__file__).resolve().parents[1] / "resources" / "config" / "default-format.json"
        config = json.loads(resource.read_text(encoding="utf-8"))
        headers = {
            "X-Proxy-Secret": STRONG_PROXY_SECRET,
            "X-Filename": "spawn-smoke.docx",
            **_format_config_headers(config),
        }

        upload = _upload_docx(base_url, _docx_bytes(), headers)
        task_id = upload["task_id"]
        status = {}
        deadline = time.monotonic() + 30
        while time.monotonic() < deadline:
            status = _json_request(f"{base_url}/status/{task_id}", {"X-Proxy-Secret": STRONG_PROXY_SECRET})
            if status.get("status") in {"done", "error", "timeout", "failed"}:
                break
            time.sleep(0.1)

        assert status.get("status") == "done", status
        downloaded = _read_download(base_url, task_id, {"X-Proxy-Secret": STRONG_PROXY_SECRET})
        assert zipfile.is_zipfile(io.BytesIO(downloaded))
        assert Document(io.BytesIO(downloaded)).paragraphs

        task_tmp_dir = tmp_path / "runtime" / "tmp" / task_id
        deadline = time.monotonic() + 5
        while task_tmp_dir.exists() and time.monotonic() < deadline:
            time.sleep(0.05)
        assert not task_tmp_dir.exists()
    except HTTPError as exc:
        raise AssertionError(exc.read().decode("utf-8", errors="replace")) from exc
    finally:
        process.terminate()
        try:
            process.communicate(timeout=10)
        except subprocess.TimeoutExpired:
            process.kill()
            process.communicate(timeout=10)
