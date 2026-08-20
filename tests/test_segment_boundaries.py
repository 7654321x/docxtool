from __future__ import annotations

import importlib

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
import pytest

import docxtool.document.importer as importer_module
import docxtool.document.segmentation.boundaries as boundaries_module
from docxtool.document.importer import (
    ParagraphFeatures,
    SourceRun,
    _heading_has_inline_body,
    _segment_boundary_candidates,
    _split_inline_heading_body_spans,
    _split_structural_tail_after_numbered_heading,
    _validate_source_span_partition,
    _validate_numbered_heading_body_split,
)
from docxtool.document.recognition.model import DocumentMode
from docxtool.document.segmentation.boundaries import heading_has_inline_body
from docxtool.document.segmentation.soft_breaks import (
    is_structural_key_value_line,
    should_split_structural_line_breaks,
)
from docxtool.sdk import recognize_docx


def _soft_break_decision(parts: list[str], next_text: str = "") -> bool:
    """测试辅助：传入软换行行文本，返回新 Segmenter 的拆分决策。"""
    return should_split_structural_line_breaks(
        parts,
        next_text,
        detect_numbering_prefix_func=lambda text: "一、" if (text or "").startswith("一、") else "",
        is_dispatch_number_line_func=lambda text: (text or "").startswith("发〔"),
        is_key_value_line_func=lambda text: "责任单位：" in (text or ""),
        is_sign_date_func=lambda text: bool((text or "").strip().endswith("日")),
        is_attachment_boundary_func=lambda text: (text or "").strip().startswith("附件"),
        is_tail_signature_org_func=lambda text: (text or "").strip() == "测试单位",
        is_role_name_line_func=lambda text: "  " in (text or ""),
        is_header_role_date_pair_func=lambda left, right: "主席" in (left or "") and (right or "").startswith("（"),
        is_standalone_addressing_func=importer_module._is_standalone_addressing_text,
    )


def _features_for_visual_boundary(source: str, boundary: int) -> ParagraphFeatures:
    return ParagraphFeatures(
        source_physical_text=source,
        source_run_spans=(
            SourceRun(
                start=0, end=boundary, font_name="SimHei", east_asia_font_name="SimHei",
                ascii_font_name="Arial", font_size_pt=16.0, bold=True, italic=False,
                underline=False, explicit=True, inherited=False, known=True,
                format_sources=("direct_run",),
            ),
            SourceRun(
                start=boundary, end=len(source), font_name="FangSong",
                east_asia_font_name="FangSong", ascii_font_name="Times New Roman",
                font_size_pt=12.0, bold=False, italic=False, underline=False,
                explicit=True, inherited=False, known=True, format_sources=("direct_run",),
            ),
        ),
    )


def _features_for_bold_prefix(source: str, bold_end: int) -> ParagraphFeatures:
    return ParagraphFeatures(
        source_physical_text=source,
        source_run_spans=(
            SourceRun(
                start=0, end=bold_end, font_name="FangSong", east_asia_font_name="FangSong",
                ascii_font_name="Times New Roman", font_size_pt=16.0, bold=True,
                italic=False, underline=False, explicit=True, inherited=False, known=True,
                format_sources=("direct_run",),
            ),
            SourceRun(
                start=bold_end, end=len(source), font_name="FangSong",
                east_asia_font_name="FangSong", ascii_font_name="Times New Roman",
                font_size_pt=16.0, bold=False, italic=False, underline=False,
                explicit=True, inherited=False, known=True, format_sources=("direct_run",),
            ),
        ),
    )


def test_source_span_conservation_facades_call_new_module(monkeypatch) -> None:
    """boundaries 与 importer 旧入口都应调用唯一的守恒实现。"""
    conservation_module = importlib.import_module(
        "docxtool.document.segmentation.conservation"
    )
    calls = []

    def fake_validate(source, spans):
        calls.append((source, spans))

    monkeypatch.setattr(
        conservation_module,
        "validate_source_span_partition",
        fake_validate,
    )

    boundaries_module.validate_source_span_partition("正文", [(0, 2)])
    importer_module._validate_source_span_partition("标题", [(0, 2)])

    assert calls == [("正文", [(0, 2)]), ("标题", [(0, 2)])]


@pytest.mark.parametrize("document_mode", [DocumentMode.REPORT])
def test_annual_review_boundary_is_report_only(document_mode: DocumentMode) -> None:
    source = "五年来。我们持续推进有关工作。"
    spans = _split_inline_heading_body_spans(
        source,
        0,
        len(source),
        document_mode=document_mode,
    )
    assert [source[start:end] for start, end in spans] == [
        "五年来。",
        "我们持续推进有关工作。",
    ]


@pytest.mark.parametrize("document_mode", [DocumentMode.NORMAL, DocumentMode.UNKNOWN])
def test_annual_review_boundary_is_not_used_outside_report(document_mode: DocumentMode) -> None:
    source = "五年来。我们持续推进有关工作。"
    spans = _split_inline_heading_body_spans(
        source,
        0,
        len(source),
        document_mode=document_mode,
    )
    assert spans == [(0, len(source))]


def test_annual_review_boundary_covers_one_year_prefix() -> None:
    source = "一年来。我们重点抓好了以下工作。"
    report_spans = _split_inline_heading_body_spans(
        source,
        0,
        len(source),
        document_mode=DocumentMode.REPORT,
    )
    normal_spans = _split_inline_heading_body_spans(
        source,
        0,
        len(source),
        document_mode=DocumentMode.NORMAL,
    )
    assert [source[start:end] for start, end in report_spans] == [
        "一年来。",
        "我们重点抓好了以下工作。",
    ]
    assert normal_spans == [(0, len(source))]


def test_visual_title_terminator_can_split_without_numbering() -> None:
    source = "关于推进工作的要求。各单位应当结合实际认真执行。"
    boundary = source.index("。") + 1

    candidates = _segment_boundary_candidates(
        source, 0, len(source), _features_for_visual_boundary(source, boundary)
    )
    spans = _split_inline_heading_body_spans(
        source, 0, len(source), _features_for_visual_boundary(source, boundary)
    )

    assert candidates[0].left_type_hint == "title_or_heading"
    assert "VISUAL_TITLE_TERMINATOR" in candidates[0].evidence
    assert [source[start:end] for start, end in spans] == [
        "关于推进工作的要求。", "各单位应当结合实际认真执行。",
    ]


def test_bold_glossary_prefix_does_not_split_later_sentences() -> None:
    source = "【加粗术语：】普通解释第一句。普通解释第二句。普通解释第三句。"
    features = _features_for_bold_prefix(source, source.index("】") + 1)
    first_sentence_end = source.index("。") + 1

    assert not boundaries_module.has_format_transition(
        features,
        0,
        first_sentence_end,
        len(source),
    )
    assert _split_inline_heading_body_spans(
        source,
        0,
        len(source),
        features,
    ) == [(0, len(source))]


def test_numbered_glossary_item_does_not_split_normal_explanation_sentences() -> None:
    source = "21.人工林近自然化改造：普通解释第一句。普通解释第二句。普通解释第三句。"
    features = _features_for_bold_prefix(source, source.index("：") + 1)

    assert _segment_boundary_candidates(source, 0, len(source), features) == ()
    assert _split_inline_heading_body_spans(
        source,
        0,
        len(source),
        features,
    ) == [(0, len(source))]


def test_numbered_explanatory_heading_still_splits_at_local_format_transition() -> None:
    source = "1.强化组织领导：明确工作责任。各部门要结合实际认真执行。"
    boundary = source.index("。") + 1
    features = _features_for_visual_boundary(source, boundary)

    assert [
        source[start:end]
        for start, end in _split_inline_heading_body_spans(
            source,
            0,
            len(source),
            features,
        )
    ] == ["1.强化组织领导：明确工作责任。", "各部门要结合实际认真执行。"]


def test_sdk_drops_zero_width_glossary_tail_and_zero_width_only_paragraph(
    tmp_path,
) -> None:
    source = tmp_path / "glossary-zero-width.docx"
    first_text = "21.人工林近自然化改造：普通解释第一句。普通解释第二句。普通解释第三句。"
    second_text = "22.山林季相和林相：这是下一条完整解释。"
    document = Document()
    title = document.add_paragraph("人民检察院工作报告")
    title.alignment = 1
    title.runs[0].bold = True
    document.add_paragraph("正文已经开始并完整说明工作背景和基本情况。")
    glossary_title = document.add_paragraph("名词解释")
    glossary_title.paragraph_format.page_break_before = True
    first = document.add_paragraph()
    first.add_run("21.人工林近自然化改造：").bold = True
    first.add_run("普通解释第一句。普通解释第二句。普通解释第三句。\u200c\u200c\u200c\u200c\u200c")
    document.add_paragraph("\u200c\u200c")
    second = document.add_paragraph()
    second.add_run("22.山林季相和林相：").bold = True
    second.add_run("这是下一条完整解释。")
    document.save(source)

    plan = recognize_docx(
        source,
        processing_mode="structural",
        recognition_mode="authoritative",
        include_text=True,
    )
    first_blocks = [
        block for block in plan.blocks if block.physical_paragraph_index == 3
    ]

    assert [(block.type_id, block.recognized_text) for block in first_blocks] == [
        ("glossary_item", first_text),
    ]
    assert first_blocks[0].segment_count_total == 1
    assert not any(
        block.physical_paragraph_index == 4 for block in plan.blocks
    )
    assert [
        (block.type_id, block.recognized_text)
        for block in plan.blocks
        if block.physical_paragraph_index == 5
    ] == [("glossary_item", second_text)]


def test_heading_has_inline_body_shared_boundary_helper() -> None:
    """验证标题正文粘连边界 helper 与旧 importer 入口保持一致。"""
    assert heading_has_inline_body("一、标题。正文内容不少于五字")
    assert _heading_has_inline_body("一、标题。正文内容不少于五字")
    assert not heading_has_inline_body("一、标题。短")


def test_visible_heading2_and_body_stay_in_one_physical_segment() -> None:
    source = "（一）会议安排。后续正文内容完整说明有关工作要求。"

    assert _split_inline_heading_body_spans(source, 0, len(source)) == [
        (0, len(source))
    ]


def test_soft_break_decision_uses_structural_evidence_without_final_type() -> None:
    """软换行决策传入结构证据后，只返回是否拆段，不给出最终段落类型。"""
    assert _soft_break_decision(["一、标题。", "正文继续说明。"])
    assert _soft_break_decision(["责任单位：办公室", "责任单位：研究室"])
    assert _soft_break_decision(["测试单位", "2026年5月1日"])
    assert _soft_break_decision(["正文结束", "测试单位"], "2026年5月1日")
    assert _soft_break_decision(["正文最后一段完整结束。", "各位委员、同志们！"])
    assert _soft_break_decision(["正文最后一段完整结束。", "各位委员、同志们："])
    assert _soft_break_decision(["正文最后一段完整结束。", "同志们！"])
    assert not _soft_break_decision(["普通正文第一行", "普通正文第二行"])


def test_soft_break_splits_complete_front_meeting_metadata_profile() -> None:
    """标题、括号日期和会议说明同时成立时才拆文首软换行。"""
    assert _soft_break_decision([
        "年度重点工作报告",
        "（2026年8月27日）",
        "在全市重点工作会议上",
    ])
    assert not _soft_break_decision([
        "年度重点工作报告",
        "（2026年8月27日）",
        "后续正文已经开始并说明具体工作要求。",
    ])


def test_structural_key_value_line_uses_injected_boundary_evidence() -> None:
    """键值边界判断接收责任单位和冒号标签回调，只返回软换行边界事实。"""
    assert is_structural_key_value_line(
        "责任单位：办公室",
        is_responsibility_line_func=lambda text: text.startswith("责任单位："),
        colon_bold_match_func=lambda _text: -1,
    )
    assert is_structural_key_value_line(
        "联系人：张三",
        is_responsibility_line_func=lambda _text: False,
        colon_bold_match_func=lambda text: text.find("："),
    )
    assert not is_structural_key_value_line(
        "某某学院：现将情况说明如下。",
        is_responsibility_line_func=lambda _text: False,
        colon_bold_match_func=lambda _text: -1,
    )


def test_body_visual_emphasis_does_not_create_a_paragraph_boundary() -> None:
    source = "推动重点工作走深走实。各单位应当结合实际认真执行。"
    boundary = source.index("。") + 1
    features = _features_for_visual_boundary(source, boundary)

    spans = _split_inline_heading_body_spans(
        source, 0, len(source), features, allow_visual_boundary=False
    )

    assert spans == [(0, len(source))]


def test_native_list_bold_heading_splits_inside_body_region() -> None:
    source = "阶段任务安排。后续正文内容完整保留并继续说明具体工作。"
    boundary = source.index("。") + 1
    features = _features_for_visual_boundary(source, boundary)
    features.numbering_prefix = "@lvl_0"

    candidates = _segment_boundary_candidates(source, 0, len(source), features)
    spans = _split_inline_heading_body_spans(
        source, 0, len(source), features, allow_visual_boundary=False
    )

    assert candidates[0].left_type_hint == "heading"
    assert [source[start:end] for start, end in spans] == [
        "阶段任务安排。", "后续正文内容完整保留并继续说明具体工作。",
    ]


def test_native_list_without_heading_format_transition_stays_one_segment() -> None:
    source = "这是普通自动列表正文第一句。后续内容仍然属于同一列表正文段落。"
    features = ParagraphFeatures(
        source_physical_text=source,
        numbering_prefix="@lvl_0",
        source_run_spans=(
            SourceRun(
                start=0, end=len(source), font_name="FangSong",
                east_asia_font_name="FangSong", ascii_font_name="Times New Roman",
                font_size_pt=16.0, bold=False, italic=False, underline=False,
                explicit=True, inherited=False, known=True,
                format_sources=("direct_run",),
            ),
        ),
    )

    spans = _split_inline_heading_body_spans(
        source, 0, len(source), features, allow_visual_boundary=False
    )

    assert spans == [(0, len(source))]


def test_plain_prose_sentence_is_not_split_without_structure_evidence() -> None:
    source = "这是普通正文第一句。后面仍然是同一段普通正文内容。"

    assert _segment_boundary_candidates(source, 0, len(source)) == ()
    assert _split_inline_heading_body_spans(source, 0, len(source)) == [(0, len(source))]


def test_numbered_heading_splits_to_one_complete_body_despite_later_format_change() -> None:
    source = "一、总体要求。正文首句说明。后续正文仍属于同一段。"
    heading_end = source.index("。") + 1
    body_sentence_end = source.index("。", heading_end) + 1
    features = _features_for_visual_boundary(source, body_sentence_end)

    later_candidates = _segment_boundary_candidates(source, heading_end, len(source), features)
    spans = _split_inline_heading_body_spans(source, 0, len(source), features)

    assert later_candidates[0].left_type_hint == "title_or_heading"
    assert [source[start:end] for start, end in spans] == [
        "一、总体要求。", "正文首句说明。后续正文仍属于同一段。",
    ]
    _validate_numbered_heading_body_split(source, spans, features)


def test_numbered_heading_body_validation_rejects_extra_body_segment() -> None:
    source = "一、总体要求。正文首句说明。后续正文仍属于同一段。"
    heading_end = source.index("。") + 1
    body_sentence_end = source.index("。", heading_end) + 1

    with pytest.raises(ValueError, match="一个完整正文段"):
        _validate_numbered_heading_body_split(
            source,
            [(0, heading_end), (heading_end, body_sentence_end), (body_sentence_end, len(source))],
        )


def test_numbered_heading_releases_standalone_salutation_after_one_body() -> None:
    source = "一、总体要求。正文首句说明。后续正文仍属于同一段。\n\n\n各位委员、同志们！"
    heading_body = _split_inline_heading_body_spans(source, 0, len(source))

    spans = _split_structural_tail_after_numbered_heading(source, heading_body)

    _validate_source_span_partition(source, spans)
    assert [source[start:end] for start, end in spans] == [
        "一、总体要求。",
        "正文首句说明。后续正文仍属于同一段。",
        "各位委员、同志们！",
    ]


def test_numbered_heading_releases_personal_title_salutation_after_one_body() -> None:
    source = "一、总体要求。正文首句说明。后续正文仍属于同一段。\n\n余书记："
    heading_body = _split_inline_heading_body_spans(source, 0, len(source))

    spans = _split_structural_tail_after_numbered_heading(source, heading_body)

    _validate_source_span_partition(source, spans)
    assert [source[start:end] for start, end in spans] == [
        "一、总体要求。",
        "正文首句说明。后续正文仍属于同一段。",
        "余书记：",
    ]


def test_inline_label_and_content_do_not_create_a_paragraph_boundary() -> None:
    source = "责任单位：办公室负责统筹落实相关工作。"
    candidates = _segment_boundary_candidates(source, 0, len(source))

    assert candidates == ()


def test_organization_label_and_content_stay_in_one_sdk_block(tmp_path) -> None:
    source = tmp_path / "organization-label-body.docx"
    text = "某某职业学院：调研中发现该校有关工作正在有序推进。"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("某某职业学院：")
    paragraph.add_run("调研中发现该校有关工作正在有序推进。")
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative", include_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert len(blocks) == 1
    assert blocks[0].type_id == "body"
    assert blocks[0].recognized_text == text
    assert blocks[0].locator_verified


def test_heading2_period_body_keeps_source_physical_paragraph_boundaries(
    tmp_path,
) -> None:
    one_source = tmp_path / "heading2-one-physical-paragraph.docx"
    two_source = tmp_path / "heading2-two-physical-paragraphs.docx"
    heading = "（一）会议安排。"
    body = "后续正文内容完整说明有关工作要求。"

    one = Document()
    one.add_paragraph(f"{heading}{body}")
    one.save(one_source)
    two = Document()
    two.add_paragraph(heading)
    two.add_paragraph(body)
    two.save(two_source)

    one_plan = recognize_docx(one_source, processing_mode="structural", include_text=True)
    one_blocks = [block for block in one_plan.blocks if block.kind == "paragraph"]
    assert [(block.type_id, block.recognized_text) for block in one_blocks] == [
        ("heading2", f"{heading}{body}")
    ]
    assert one_blocks[0].physical_paragraph_index == 0
    assert one_blocks[0].locator_verified is True

    two_plan = recognize_docx(two_source, processing_mode="structural", include_text=True)
    two_blocks = [block for block in two_plan.blocks if block.kind == "paragraph"]
    assert [(block.type_id, block.recognized_text) for block in two_blocks] == [
        ("heading2", heading),
        ("body", body),
    ]
    assert [block.physical_paragraph_index for block in two_blocks] == [0, 1]


def test_heading2_word_style_period_body_stays_one_physical_paragraph(
    tmp_path,
) -> None:
    source = tmp_path / "heading2-style-inline-body.docx"
    text = "会议安排。后续正文内容完整说明有关工作要求。"
    document = Document()
    document.add_paragraph(text, style="Heading 2")
    document.save(source)

    plan = recognize_docx(source, processing_mode="structural", include_text=True)
    blocks = [block for block in plan.blocks if block.kind == "paragraph"]

    assert [(block.type_id, block.recognized_text) for block in blocks] == [
        ("heading2", text)
    ]
    assert blocks[0].physical_paragraph_index == 0
    assert blocks[0].locator_verified is True


def test_inline_salutation_and_body_split_without_specific_name(tmp_path) -> None:
    source = tmp_path / "salutation-inline-body.docx"
    text = "各位代表、同志们：现在开始说明本次会议安排。"
    document = Document()
    document.add_paragraph(text)
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative", include_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert [block.type_id for block in blocks] == ["addressing", "body"]
    assert [block.recognized_text for block in blocks] == [
        "各位代表、同志们：", "现在开始说明本次会议安排。",
    ]
    assert all(block.locator_verified for block in blocks)


def test_inline_organization_label_and_body_is_not_split(tmp_path) -> None:
    source = tmp_path / "organization-inline-body.docx"
    text = "某某研究院：现将有关情况说明如下。"
    document = Document()
    document.add_paragraph(text)
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative", include_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert len(blocks) == 1
    assert blocks[0].type_id == "body"
    assert blocks[0].recognized_text == text


def test_soft_broken_front_meeting_metadata_is_not_downgraded_to_body(tmp_path) -> None:
    source = tmp_path / "front-meeting-metadata.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("年度重点工作报告")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("（2026年8月27日）")
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("在全市重点工作会议上")
    document.add_paragraph("现将有关工作情况报告如下。")
    document.save(source)

    plan = recognize_docx(
        source,
        processing_mode="structural",
        recognition_mode="authoritative",
        include_text=True,
    )
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert [block.type_id for block in blocks] == [
        "title",
        "meeting_title_meta",
        "meeting_title_meta",
    ]
    assert [block.segment_index for block in blocks] == [0, 1, 2]
    assert {block.segment_count_total for block in blocks} == {3}
    assert all(block.locator_verified for block in blocks)


def test_attachment_note_is_not_split_as_a_short_label() -> None:
    source = "附件：1.基本情况"

    assert _segment_boundary_candidates(source, 0, len(source)) == ()
    assert _split_inline_heading_body_spans(source, 0, len(source)) == [(0, len(source))]


def test_numbered_heading_with_wrapped_body_emits_one_complete_body_segment(tmp_path) -> None:
    source = tmp_path / "heading-and-complete-body.docx"
    document = Document()
    paragraph = document.add_paragraph()
    heading = paragraph.add_run("一、总体要求。")
    heading.font.bold = True
    heading.font.size = Pt(16)
    paragraph.add_run("正文第一部分属于同一正文段。")
    trailing = paragraph.add_run()
    trailing.add_break(WD_BREAK.LINE)
    trailing.add_text("正文第二部分仍属于同一正文段。")
    document.save(source)

    plan = recognize_docx(source, recognition_mode="legacy", include_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert len(blocks) == 2
    assert [block.segment_index for block in blocks] == [0, 1]
    assert {block.segment_count_total for block in blocks} == {2}
    assert all(block.source_locator_status == "confirmed" for block in blocks)
    assert [block.recognized_text for block in blocks] == [
        "一、总体要求。", "正文第一部分属于同一正文段。\n正文第二部分仍属于同一正文段。",
    ]


def test_native_list_bold_heading_and_body_emit_two_located_sdk_blocks(tmp_path) -> None:
    source = tmp_path / "native-list-heading-body.docx"
    document = Document()
    document.add_paragraph("工作总结")
    paragraph = document.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "2")
    num_pr.extend((ilvl, num_id))
    properties.append(num_pr)
    heading = paragraph.add_run("阶段任务安排。")
    heading.bold = True
    heading.font.size = Pt(16)
    paragraph.add_run("后续正文内容完整保留并继续说明具体工作。")
    document.save(source)

    plan = recognize_docx(
        source,
        processing_mode="structural",
        recognition_mode="authoritative",
        include_text=True,
    )
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 1]

    assert [block.type_id for block in blocks] == ["heading2", "body"]
    assert [block.recognized_text for block in blocks] == [
        "阶段任务安排。", "后续正文内容完整保留并继续说明具体工作。",
    ]
    assert [block.segment_index for block in blocks] == [0, 1]
    assert {block.segment_count_total for block in blocks} == {2}
    assert all(block.source_locator_status == "confirmed" for block in blocks)
    assert blocks[0].raw_end_utf16 == blocks[1].raw_start_utf16


def test_sdk_locates_heading_body_and_later_salutation_without_gaps(tmp_path) -> None:
    source = tmp_path / "heading-body-salutation.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("一、总体要求。正文内容完整保留。")
    for _unused in range(3):
        paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run("各位委员、同志们！")
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative", include_text=True)
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 0]

    assert [block.type_id for block in blocks] == ["heading1", "body", "addressing"]
    assert [block.recognized_text for block in blocks] == [
        "一、总体要求。", "正文内容完整保留。", "各位委员、同志们！",
    ]
    assert all(block.locator_verified for block in blocks)


@pytest.mark.parametrize(
    ("break_count", "addressing_text"),
    [
        (4, "各位委员、同志们！"),
        (1, "各位委员、同志们："),
    ],
)
def test_sdk_splits_body_tail_standalone_addressing_soft_lines(
    tmp_path,
    break_count: int,
    addressing_text: str,
) -> None:
    source = tmp_path / f"body-tail-addressing-{break_count}.docx"
    body_text = "正文最后一段完整结束。"
    document = Document()
    document.add_paragraph("年度工作报告", style="Title")
    document.add_paragraph("现将有关工作情况报告如下，请认真审议。")
    paragraph = document.add_paragraph()
    paragraph.add_run(body_text)
    for _unused in range(break_count):
        paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run(addressing_text)
    document.save(source)

    plan = recognize_docx(
        source,
        processing_mode="structural",
        recognition_mode="authoritative",
        include_text=True,
    )
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 2]

    assert [block.type_id for block in blocks] == ["body", "addressing"]
    assert [block.recognized_text for block in blocks] == [body_text, addressing_text]
    assert [block.segment_index for block in blocks] == [0, 1]
    assert {block.segment_count_total for block in blocks} == {2}
    assert all(block.locator_verified for block in blocks)


def test_sdk_splits_only_local_addressing_boundary_after_ordinary_soft_line(
    tmp_path,
) -> None:
    source = tmp_path / "mixed-body-tail-addressing.docx"
    first_body_line = "正文第一行"
    second_body_line = "正文第二行"
    addressing_text = "各位委员、同志们！"
    document = Document()
    document.add_paragraph("年度工作报告", style="Title")
    document.add_paragraph("现将有关工作情况报告如下，请认真审议。")
    paragraph = document.add_paragraph()
    paragraph.add_run(first_body_line)
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run(second_body_line)
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run(addressing_text)
    document.save(source)

    plan = recognize_docx(
        source,
        processing_mode="structural",
        recognition_mode="authoritative",
        include_text=True,
    )
    blocks = [block for block in plan.blocks if block.physical_paragraph_index == 2]

    assert [block.type_id for block in blocks] == ["body", "addressing"]
    assert [block.recognized_text for block in blocks] == [
        f"{first_body_line}\n{second_body_line}",
        addressing_text,
    ]
    assert [block.segment_index for block in blocks] == [0, 1]
    assert {block.segment_count_total for block in blocks} == {2}
    assert all(block.locator_verified for block in blocks)
    assert "".join(block.recognized_text.replace("\n", "") for block in blocks) == (
        first_body_line + second_body_line + addressing_text
    )
