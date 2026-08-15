import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docxtool.document.engine import export_doc
from docxtool.document.importer import (
    DocxImporter,
    DocumentData,
    InlineToken,
    ParagraphData,
    ParagraphFeatures,
    _normalize_inline_tokens,
    _normalize_quotes,
    _normalize_tail_structures,
    _normalize_text,
    _to_chinese_punctuation,
    strip_numbering,
)
from docxtool.document.pipeline.options import resolve_import_processing_options
from docxtool.document.normalization import normalize_tail_structures, strip_numbering_prefix
from docxtool.document.style_config import PageSettings, StyleRule, load_rules_and_settings


def _rules():
    return [StyleRule.default_for_row(i) for i in range(10)]


class ProcessingFlagsTest(unittest.TestCase):
    def test_tail_normalizer_importer_facade_uses_normalization_module(self):
        self.assertIs(_normalize_tail_structures, normalize_tail_structures)

    def test_numbering_strip_importer_facade_uses_normalization_module(self):
        self.assertEqual(strip_numbering("一、改革", "一、"), strip_numbering_prefix("一、改革", "一、"))

    def test_punctuation_disabled_keeps_original_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "source.docx"
            doc = Document()
            doc.add_paragraph("甲:乙,丙.丁")
            doc.save(src)

            data = DocxImporter().load(
                str(src), _rules(), features={"punctuation_enabled": False},
                strict_preservation=False,
            )

            self.assertEqual(data.paragraphs[0].text, "甲:乙,丙.丁")

    def test_punctuation_enabled_removes_fullwidth_space(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "source.docx"
            doc = Document()
            doc.add_paragraph("甲　乙,丙")
            doc.save(src)

            data = DocxImporter().load(
                str(src), _rules(), features={"punctuation_enabled": True},
                strict_preservation=False,
            )

            self.assertEqual(data.paragraphs[0].text, "甲乙，丙")

    def test_smart_mode_applies_only_enabled_punctuation_repairs(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "punctuation.docx"
            document = Document()
            document.add_paragraph("甲:乙,丙.")
            document.save(src)

            enabled = DocxImporter().load(
                str(src),
                _rules(),
                features={
                    "processing": {"strategy": "structural"},
                    "punctuation_enabled": True,
                    "punctuation": {"enabled": True, "mode": "safe"},
                },
            )
            disabled = DocxImporter().load(
                str(src),
                _rules(),
                features={
                    "processing": {"strategy": "structural"},
                    "punctuation_enabled": False,
                    "punctuation": {"enabled": False, "mode": "safe"},
                },
            )

            self.assertEqual(enabled.paragraphs[0].text, "甲：乙，丙。")
            self.assertEqual(disabled.paragraphs[0].text, "甲:乙,丙.")

    def test_page_number_disabled_skips_footer_fields(self):
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "out.docx"
            data = DocumentData(
                paragraphs=[
                    ParagraphData("正文内容", "body", "正文内容", ParagraphFeatures()),
                ],
                filepath="input.docx",
            )
            export_doc(data, _rules(), PageSettings(), str(out), page_number_enabled=False)

            with ZipFile(out) as zf:
                names = zf.namelist()
                footer_names = [name for name in names if name.startswith("word/footer")]
                self.assertEqual(footer_names, [])

    def test_smart_mode_splits_numbered_heading_from_one_complete_body_paragraph(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "fused-structure.docx"
            doc = Document()
            header = doc.add_paragraph()
            for index, text in enumerate((
                "内测发〔2026〕1号",
                "中共内江市东兴区政协党组班子",
                "2025年度民主生活会对照检查材料",
                "区政协办公室主任  李弟弟",
            )):
                header.add_run(text)
                if index < 3:
                    header.add_run().add_break(WD_BREAK.LINE)
            doc.add_paragraph("一、一级标题。这里是应当作为正文保留的完整说明文字。")
            tail = doc.add_paragraph()
            for index, text in enumerate((
                "区政协办公室",
                "2025年十月15日",
                "附件：1.基本情况",
                "2.具体情况",
            )):
                tail.add_run(text)
                if index < 3:
                    tail.add_run().add_break(WD_BREAK.LINE)
            doc.save(src)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(src), rules, features=features)

            self.assertEqual(data.processing_strategy, "structural")
            self.assertFalse(data.strict_preservation)
            self.assertEqual(
                [(item.type_id, item.text) for item in data.paragraphs[:6]],
                [
                    ("dispatch_number", "内测发〔2026〕1号"),
                    ("title_cont", "中共内江市东兴区政协党组班子"),
                    ("title_cont", "2025年度民主生活会对照检查材料"),
                    ("role_name", "区政协办公室主任  李弟弟"),
                    ("heading1", "一、一级标题。"),
                    ("body", "这里是应当作为正文保留的完整说明文字。"),
                ],
            )
            self.assertEqual(
                [(item.type_id, item.text) for item in data.paragraphs[-4:]],
                [
                    ("attachment_note", "附件：1.基本情况"),
                    ("attachment_note_item", "2.具体情况"),
                    ("sign_org", "区政协办公室"),
                    ("sign_date", "2025年10月15日"),
                ],
            )

    def test_smart_mode_splits_heading_body_despite_inline_page_break(self):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "inline-page-break.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("一、一级标题。这里是应当独立排版的正文前半部分，")
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            paragraph.add_run("后半部分仍应识别为同一段正文内容。")
            document.save(src)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(src), rules, features=features)

            self.assertEqual(
                [(item.type_id, item.text) for item in data.paragraphs],
                [
                    ("heading1", "一、一级标题。"),
                    ("body", "这里是应当独立排版的正文前半部分，后半部分仍应识别为同一段正文内容。"),
                ],
            )
            self.assertFalse(any(token.kind == "page_break" for token in data.paragraphs[0].inline_tokens))
            self.assertFalse(any(token.kind == "page_break" for token in data.paragraphs[1].inline_tokens))

    def test_smart_mode_keeps_bold_lead_sentence_in_one_body_paragraph(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "inline-emphasis.docx"
            document = Document()
            document.add_paragraph("一、工作情况")
            paragraph = document.add_paragraph()
            lead = paragraph.add_run("推动重点工作走深走实。")
            lead.bold = True
            lead.font.size = Pt(16)
            paragraph.add_run("各单位结合实际持续抓好任务落实，确保工作取得实效。")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)

            self.assertEqual(len(data.paragraphs), 2)
            self.assertEqual(data.paragraphs[1].type_id, "body")
            self.assertEqual(
                data.paragraphs[1].text,
                "推动重点工作走深走实。各单位结合实际持续抓好任务落实，确保工作取得实效。",
            )
            self.assertTrue(data.paragraphs[1].meta.get("inline_lead_bold"))

    def test_smart_mode_separates_salutation_after_heading_body_soft_lines(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "heading-body-salutation.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("一、总体要求。正文内容完整保留。")
            for _unused in range(4):
                paragraph.add_run().add_break(WD_BREAK.LINE)
            paragraph.add_run("各位委员、同志们！")
            document.add_paragraph("后续正文继续正常排版。")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)

            self.assertEqual(
                [(item.type_id, item.text) for item in data.paragraphs],
                [
                    ("heading1", "一、总体要求。"),
                    ("body", "正文内容完整保留。"),
                    ("addressing", "各位委员、同志们！"),
                    ("body", "后续正文继续正常排版。"),
                ],
            )

    def test_body_colon_label_is_not_reclassified_as_recipient(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "body-label.docx"
            document = Document()
            document.add_paragraph("工作情况", style="Title")
            document.add_paragraph("现将有关情况报告如下，供审阅。")
            document.add_paragraph("某某学院：")
            document.add_paragraph("调研发现有关工作正在有序推进。")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)
            label = next(item for item in data.paragraphs if item.text == "某某学院：")

            self.assertEqual(label.type_id, "body")
            self.assertTrue(label.meta.get("no_indent"))

    def test_smart_mode_mid_body_attachment_keyword_stays_body(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "mid-body-attachment-keyword.docx"
            document = Document()
            document.add_paragraph("工作情况", style="Title")
            document.add_paragraph("前段正文已经开始，并完整说明有关工作情况。")
            document.add_paragraph("附件：材料清单")
            document.add_paragraph("后续正文继续说明材料另行发送，不进入附件说明区。")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)
            paragraph = next(item for item in data.paragraphs if item.text == "附件：材料清单")

            self.assertEqual(paragraph.type_id, "body")
            self.assertEqual(
                paragraph.meta.get("recognition_evidence"),
                ["attachment-keyword-without-tail-context", "hard-structure", "legacy-reclassified"],
            )

    def test_tail_normalizer_does_not_reclassify_final_body_date(self):
        paragraph = ParagraphData(
            "2026年5月29日",
            "body",
            "2026年5月29日",
            ParagraphFeatures(),
            meta={
                "recognition_type": "body",
                "recognized_type": "body",
                "final_type": "body",
                "source_locator_status": "verified",
            },
            inline_tokens=[InlineToken("text", "2026年5月29日")],
        )
        paragraphs = [
            ParagraphData("前段正文已经开始，并完整说明有关工作情况。", "body", "前段正文已经开始，并完整说明有关工作情况。", ParagraphFeatures()),
            paragraph,
        ]

        _normalize_tail_structures(paragraphs, normalize_text=False)

        self.assertEqual(paragraph.type_id, "body")
        self.assertEqual(paragraph.meta["final_type"], "body")
        self.assertEqual(paragraph.meta["source_locator_status"], "verified")
        self.assertEqual(paragraph.inline_tokens, [InlineToken("text", "2026年5月29日")])

    def test_tail_reorder_keeps_recognition_diagnostics_in_output_order(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "tail-reorder-diagnostics.docx"
            document = Document()
            document.add_paragraph("工作情况", style="Title")
            document.add_paragraph("前段正文已经开始，并完整说明有关工作情况。")
            document.add_paragraph("星河治理委员会")
            document.add_paragraph("2026年5月29日")
            document.add_paragraph("附件：1.材料清单")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)

            self.assertEqual(
                [item.type_id for item in data.paragraphs[-3:]],
                ["attachment_note", "sign_org", "sign_date"],
            )
            diagnostics = data.recognition_diagnostics["paragraphs"]
            self.assertEqual(
                [item["final_type"] for item in diagnostics],
                [paragraph.type_id for paragraph in data.paragraphs],
            )
            self.assertEqual(
                [item["paragraph_index"] for item in diagnostics],
                list(range(len(data.paragraphs))),
            )
            structure = data.recognition_structure
            self.assertEqual(
                [
                    data.paragraphs[item.source_index].type_id
                    for item in structure.elements
                    if item.source_index is not None
                ],
                [paragraph.type_id for paragraph in data.paragraphs],
            )

    def test_smart_mode_recovers_speech_title_and_soft_line_body_from_heading_style(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "speech.docx"
            document = Document()
            document.add_paragraph("在区政协九届一次会议闭幕大会上的讲话", style="Heading 1")
            byline = document.add_paragraph()
            byline.add_run("区政协党组书记、主席张三")
            byline.add_run().add_break(WD_BREAK.LINE)
            byline.add_run("（2026年8月27日10:00，会议中心）")
            document.add_paragraph("各位委员、同志们：")
            fused = document.add_paragraph()
            fused.add_run("一、始终坚持党的全面领导")
            fused.add_run().add_break(WD_BREAK.LINE)
            fused.add_run().add_break(WD_BREAK.LINE)
            fused.add_run("要全面加强党的领导，切实提升履职质量。")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)

            self.assertEqual(
                [(item.type_id, item.text) for item in data.paragraphs],
                [
                    ("title", "在区政协九届一次会议闭幕大会上的讲话"),
                    ("role_name", "区政协党组书记、主席张三"),
                    ("date_line", "（2026年8月27日10:00，会议中心）"),
                    ("addressing", "各位委员、同志们："),
                    ("heading1", "一、始终坚持党的全面领导"),
                    ("body", "要全面加强党的领导，切实提升履职质量。"),
                ],
            )

    def test_smart_mode_front_role_and_date_override_inherited_title_format(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "formatted-byline.docx"
            document = Document()
            document.add_paragraph(
                "在某地区委员会会议闭幕大会上的讲话",
                style="Heading 1",
            )
            byline = document.add_paragraph()
            role = byline.add_run("某区委员会党组书记、主席　张测试")
            role.bold = True
            role.font.size = Pt(22)
            byline.add_run().add_break(WD_BREAK.LINE)
            date = byline.add_run("（2026年8月27日11:00，某地区委员会会议中心）")
            date.bold = True
            date.font.size = Pt(22)
            document.add_paragraph("各位代表、同志们：")
            document.add_paragraph("现将有关事项说明如下，请认真抓好落实。")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)

            self.assertEqual(
                [item.type_id for item in data.paragraphs[:5]],
                ["title", "role_name", "date_line", "addressing", "body"],
            )
            self.assertNotEqual(data.paragraphs[1].meta["review_level"], "critical_review")
            self.assertNotEqual(data.paragraphs[2].meta["review_level"], "critical_review")

            strict_rules, _, strict_features = load_rules_and_settings({"mode": "strict"})
            strict = DocxImporter().load(str(source), strict_rules, features=strict_features)
            self.assertIn("\n", strict.paragraphs[1].text)
            self.assertEqual(strict.paragraphs[1].features.source_physical_paragraph_index, 1)

    def test_smart_mode_removes_only_an_inferred_heading_prefix_from_speech_title(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "speech-output.docx"
            document = Document()
            document.add_paragraph("一、在区政协九届一次会议闭幕大会上的讲话", style="Heading 1")
            document.add_paragraph("各位委员、同志们：")
            document.save(source)

            rules, _, features = load_rules_and_settings({"mode": "smart"})
            data = DocxImporter().load(str(source), rules, features=features)

            self.assertEqual(data.paragraphs[0].type_id, "title")
            self.assertEqual(data.paragraphs[0].text, "在区政协九届一次会议闭幕大会上的讲话")
            self.assertEqual(data.paragraphs[0].original_text, "一、在区政协九届一次会议闭幕大会上的讲话")

    def test_smart_mode_renumbers_headings_only_when_numbering_is_enabled(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "numbering.docx"
            output = Path(tmp) / "numbering-output.docx"
            document = Document()
            for text in (
                "测试材料",
                "一、第一部分",
                "（六）第二层",
                "5..第三层",
                "（6）第四层",
                "正文内容正文内容正文内容。",
            ):
                document.add_paragraph(text)
            document.save(source)

            rules, settings, features = load_rules_and_settings(
                {"mode": "smart", "numbering": {"enabled": True}}
            )
            data = DocxImporter().load(str(source), rules, features=features)
            export_doc(
                data,
                rules,
                settings,
                str(output),
                numbering_options=features["numbering"],
            )

            headings = [
                paragraph.text
                for paragraph in Document(output).paragraphs
                if paragraph.style.style_id in {
                    "DCT-Heading1", "DCT-Heading2", "DCT-Heading3", "DCT-Heading4"
                }
            ]
            self.assertEqual(headings, ["一、第一部分", "（一）第二层", "1.第三层", "（1）第四层"])

    def test_smart_mode_restores_visible_heading1_from_wps_native_numbering(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "wps-native-numbering.docx"
            output = Path(tmp) / "wps-native-numbering-output.docx"
            document = Document()
            document.add_paragraph("工作总结", style="Title")
            document.add_paragraph("前段正文已经开始，并完整说明有关工作情况。")
            parent = document.add_paragraph()
            parent.add_run("阶段工作安排")
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "1")
            num_pr.extend((ilvl, num_id))
            parent._p.get_or_add_pPr().append(num_pr)
            child = document.add_paragraph()
            child.add_run("（一）第一项工作").bold = True
            document.add_paragraph("后续正文对该标题展开具体说明。")
            document.save(source)

            rules, settings, features = load_rules_and_settings(
                {"mode": "smart", "numbering": {"enabled": True}}
            )
            data = DocxImporter().load(str(source), rules, features=features)
            parent_data = next(item for item in data.paragraphs if item.text == "阶段工作安排")
            child_data = next(item for item in data.paragraphs if item.text == "（一）第一项工作")

            self.assertEqual(parent_data.type_id, "heading1")
            self.assertEqual(child_data.type_id, "heading2")
            export_doc(
                data,
                rules,
                settings,
                str(output),
                numbering_options=features["numbering"],
            )

            exported = Document(output)
            headings = [
                (paragraph.style.style_id, paragraph.text)
                for paragraph in exported.paragraphs
                if paragraph.style.style_id in {"DCT-Heading1", "DCT-Heading2"}
            ]
            self.assertEqual(
                headings,
                [
                    ("DCT-Heading1", "一、阶段工作安排"),
                    ("DCT-Heading2", "（一）第一项工作"),
                ],
            )
            with ZipFile(output) as package:
                self.assertIsNone(package.testzip())

    def test_smart_mode_rebuilds_malformed_chinese_dot_heading1(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "malformed-heading1.docx"
            output = Path(tmp) / "malformed-heading1-output.docx"
            document = Document()
            document.add_paragraph("测试材料")
            document.add_paragraph("一、第一部分")
            document.add_paragraph("正文内容完整保留。")
            document.add_paragraph("二.存在的问题")
            document.add_paragraph("后续正文内容完整保留。")
            document.save(source)

            rules, settings, features = load_rules_and_settings(
                {"mode": "smart", "numbering": {"enabled": True}}
            )
            data = DocxImporter().load(str(source), rules, features=features)
            export_doc(
                data,
                rules,
                settings,
                str(output),
                numbering_options=features["numbering"],
            )

            headings = [
                paragraph.text
                for paragraph in Document(output).paragraphs
                if paragraph.style.style_id == "DCT-Heading1"
            ]
            self.assertEqual(headings, ["一、第一部分", "二、存在的问题"])


    # ── canonical-first punctuation 契约（5.4.3） ─────────────────────────

    def _load_punctuation(self, text, features, *, strict_preservation=False):
        with tempfile.TemporaryDirectory() as tmp:
            src = Path(tmp) / "punctuation.docx"
            doc = Document()
            doc.add_paragraph(text)
            doc.save(src)
            data = DocxImporter().load(
                str(src),
                _rules(),
                features=features,
                strict_preservation=strict_preservation,
            )
            return data.paragraphs[0]

    def _load_with_config(self, text, config, *, strict_preservation=False):
        """通过 load_rules_and_settings 解析真实请求配置后再导入。"""
        rules, _, features = load_rules_and_settings(config)
        return self._load_punctuation(text, features, strict_preservation=strict_preservation)

    def test_punctuation_canonical_false_only_disables_everything(self):
        para = self._load_with_config(
            "甲:乙,丙.", {"punctuation": {"enabled": False}}
        )
        self.assertEqual(para.text, "甲:乙,丙.")

    def test_punctuation_canonical_true_only_uses_canonical_engine(self):
        para = self._load_with_config(
            "甲:乙,丙.", {"punctuation": {"enabled": True, "mode": "safe"}}
        )
        self.assertEqual(para.text, "甲：乙，丙。")

    def test_punctuation_legacy_false_only_disables(self):
        para = self._load_with_config(
            "甲:乙,丙.", {"features": {"punctuation_enabled": False}}
        )
        self.assertEqual(para.text, "甲:乙,丙.")

    def test_punctuation_legacy_true_only_uses_legacy_engine(self):
        para = self._load_with_config(
            "甲:乙,丙.", {"features": {"punctuation_enabled": True}}
        )
        self.assertEqual(para.text, "甲：乙，丙。")

    def test_punctuation_canonical_false_overrides_legacy_true(self):
        para = self._load_with_config(
            "甲:乙,丙.",
            {
                "punctuation": {"enabled": False},
                "features": {"punctuation_enabled": True},
            },
        )
        self.assertEqual(para.text, "甲:乙,丙.")

    def test_punctuation_canonical_true_overrides_legacy_false(self):
        para = self._load_with_config(
            "甲:乙,丙.",
            {
                "punctuation": {"enabled": True, "mode": "safe"},
                "features": {"punctuation_enabled": False},
            },
        )
        self.assertEqual(para.text, "甲：乙，丙。")

    def test_punctuation_canonical_false_disables_in_structural_mode(self):
        para = self._load_with_config(
            "甲:乙,丙.",
            {
                "processing": {"strategy": "structural"},
                "punctuation": {"enabled": False},
            },
        )
        self.assertEqual(para.text, "甲:乙,丙.")

    def test_punctuation_canonical_true_respected_in_structural_mode(self):
        para = self._load_with_config(
            "甲:乙,丙.",
            {
                "processing": {"strategy": "structural"},
                "punctuation": {"enabled": True, "mode": "safe"},
            },
        )
        self.assertEqual(para.text, "甲：乙，丙。")

    def test_punctuation_never_applies_in_strict_preservation(self):
        para = self._load_punctuation(
            "甲:乙,丙.",
            {"punctuation": {"enabled": True, "mode": "safe"}},
            strict_preservation=True,
        )
        self.assertEqual(para.text, "甲:乙,丙.")

    def _resolve_punctuation_options(self, features):
        return resolve_import_processing_options(
            features,
            strict_preservation=False,
            recognition_mode="authoritative",
            feature_bool_func=lambda value, default: bool(value),
            normalize_basic_text_func=_normalize_text,
            normalize_quotes_func=_normalize_quotes,
            to_chinese_punctuation_func=_to_chinese_punctuation,
            normalize_inline_tokens_func=_normalize_inline_tokens,
            inline_token_type=InlineToken,
            import_error_type=ValueError,
        )

    def test_punctuation_token_normalization_respects_canonical_first(self):
        tokens = [InlineToken("text", "甲:乙"), InlineToken("text", "丙.")]

        canonical_off = self._resolve_punctuation_options(
            {"punctuation": {"enabled": False}}
        )
        self.assertEqual(
            [token.text for token in canonical_off.normalize_tokens(tokens)],
            ["甲:乙", "丙."],
        )

        canonical_on = self._resolve_punctuation_options(
            {"punctuation": {"enabled": True, "mode": "safe"}}
        )
        self.assertEqual(
            [token.text for token in canonical_on.normalize_tokens(tokens)],
            ["甲：乙", "丙。"],
        )

        legacy_on = self._resolve_punctuation_options(
            {"punctuation_enabled": True}
        )
        self.assertEqual(
            [token.text for token in legacy_on.normalize_tokens(tokens)],
            ["甲：乙", "丙。"],
        )

        conflict = self._resolve_punctuation_options(
            {
                "punctuation": {"enabled": False},
                "punctuation_enabled": True,
            }
        )
        self.assertEqual(
            [token.text for token in conflict.normalize_tokens(tokens)],
            ["甲:乙", "丙."],
        )

if __name__ == "__main__":
    unittest.main()
