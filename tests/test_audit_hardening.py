from __future__ import annotations

import io
import logging
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from docx.shared import Pt

from docxtool.document.engine.numbering import chinese_integer
from docxtool.document.importer import DocxImporter, extract_features as extract_run_features
from docxtool.document.recognition import RecognitionConfig, apply_recognition, extract_blocks
from docxtool.document.recognition.candidates import StyleCandidateProvider
from docxtool.document.recognition.compatibility import TYPE_ID_MAP, to_paragraph_type
from docxtool.document.recognition.features import BlockKind
from docxtool.document.recognition.model import ParagraphType
from docxtool.document.style_config import StyleRule
from docxtool.security.docx_validator import DocxValidationError, validate_docx_upload
from docxtool.web import app as server


def _rules() -> list[StyleRule]:
    return [StyleRule.default_for_row(index) for index in range(10)]


def _paragraph(text: str, type_id: str = "body", index: int = 0, **meta):
    return SimpleNamespace(
        text=text,
        original_text=text,
        type_id=type_id,
        features=SimpleNamespace(
            paragraph_index=index,
            alignment=meta.pop("alignment", ""),
            style_name=meta.pop("style_name", ""),
            bold=False,
            font_size_pt=None,
        ),
        meta=meta,
    )


def _document(*paragraphs, mode: str = "NORMAL"):
    return SimpleNamespace(paragraphs=list(paragraphs), doc_mode=mode)


def test_recognition_modes_apply_only_authoritative_results() -> None:
    authoritative = _document(_paragraph("国发〔2026〕23号", "body"))
    shadow = _document(_paragraph("国发〔2026〕23号", "body"))
    legacy = _document(_paragraph("国发〔2026〕23号", "body"))

    apply_recognition(authoritative, RecognitionConfig(mode="authoritative"))
    apply_recognition(shadow, RecognitionConfig(mode="shadow"))
    apply_recognition(legacy, RecognitionConfig(mode="legacy"))

    assert authoritative.paragraphs[0].type_id == "dispatch_number"
    assert shadow.paragraphs[0].type_id == "body"
    assert legacy.paragraphs[0].type_id == "body"
    assert shadow.doc_mode == "NORMAL"
    assert legacy.doc_mode == "NORMAL"
    assert shadow.recognition_diagnostics["result_applied"] is False
    assert shadow.recognition_diagnostics["paragraphs"][0]["mapping_applied"] is False
    assert authoritative.recognition_diagnostics["paragraphs"][0]["mapping_applied"] is True


def test_every_supported_recognition_type_has_central_mapping_without_unknown_body_fallback() -> None:
    assert TYPE_ID_MAP
    assert all(to_paragraph_type(item) for item in TYPE_ID_MAP)
    assert to_paragraph_type(ParagraphType.UNKNOWN) is None


@pytest.mark.parametrize(
    ("classification_kind", "expected"),
    [
        ("main_title", "title"),
        ("heading_level_1", "heading1"),
        ("heading_level_2", "heading2"),
        ("heading_level_3", "heading3"),
        ("heading_level_4", "heading4"),
        ("body", "body"),
        ("attachment_title", "attachment_title"),
    ],
)
def test_authoritative_core_results_write_all_major_renderer_types(
    classification_kind: str,
    expected: str,
) -> None:
    paragraph = _paragraph(
        "中性测试文本",
        "quote" if expected == "body" else "body",
        classification_kind=classification_kind,
        classification_confidence=0.95,
    )
    data = _document(paragraph)

    apply_recognition(data, RecognitionConfig(mode="authoritative"))

    assert paragraph.type_id == expected
    diagnostic = data.recognition_diagnostics["paragraphs"][0]
    assert diagnostic["final_type"] == expected
    assert diagnostic["mapping_applied"] is True


def test_authoritative_writes_recipient_signature_date_and_attachment_note() -> None:
    sign_org = _paragraph(
        "测试市人民政府办公室",
        "body",
        1,
        classification_kind="signature_org",
        classification_confidence=0.95,
    )
    sign_date = _paragraph(
        "2026年7月20日",
        "body",
        2,
        classification_kind="signature_date",
        classification_confidence=0.95,
    )
    signature = _document(
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 0),
        sign_org,
        sign_date,
    )
    apply_recognition(signature)

    recipient = _document(_paragraph("各有关单位：", "body"))
    attachment = _document(
        _paragraph("前段正文已经开始，并完整说明有关工作情况。", "body", 0),
        _paragraph("附件：1.测试材料", "body", 1),
        _paragraph("2.补充材料", "body", 2),
    )
    apply_recognition(recipient)
    apply_recognition(attachment)

    assert [item.type_id for item in signature.paragraphs[-2:]] == ["sign_org", "sign_date"]
    assert recipient.paragraphs[0].type_id == "addressing"
    assert [item.type_id for item in attachment.paragraphs[-2:]] == ["attachment_note", "attachment_note_item"]


def test_parenthesized_arabic_number_is_heading_four_not_heading_two() -> None:
    paragraph = _paragraph("（1）四级标题", "body")
    data = _document(paragraph)

    apply_recognition(data)

    assert paragraph.type_id == "heading4"


def test_default_import_is_strict_and_preserves_text_and_order(tmp_path: Path) -> None:
    source = tmp_path / "strict.docx"
    original = [
        "工作材料,请审阅?",
        "测试市人民政府办公室",
        "2026年7月20日",
        "附件：1.测试材料",
    ]
    document = Document()
    for text in original:
        document.add_paragraph(text)
    document.save(source)

    data = DocxImporter().load(str(source), _rules())

    assert data.strict_preservation is True
    assert [item.text for item in data.paragraphs] == original
    assert [item.original_text for item in data.paragraphs] == original
    assert data.normalization_changes
    assert all(change.applied is False for change in data.normalization_changes)


def test_import_log_names_the_hashed_source_path_accurately(tmp_path: Path, caplog) -> None:
    source = tmp_path / "logged-source.docx"
    document = Document()
    document.add_paragraph("普通正文")
    document.save(source)

    caplog.set_level(logging.INFO, logger="docx_tool")
    DocxImporter().load(str(source), _rules())

    messages = [
        record.getMessage()
        for record in caplog.records
        if record.getMessage().startswith("[导入] source_path_hash=")
    ]
    assert len(messages) == 1
    assert "file_sha256=" not in messages[0]
    assert str(source) not in messages[0]


def test_non_strict_import_records_applied_changes(tmp_path: Path) -> None:
    source = tmp_path / "normalized.docx"
    document = Document()
    document.add_paragraph("工作材料,请审阅?")
    document.save(source)

    data = DocxImporter().load(
        str(source),
        _rules(),
        strict_preservation=False,
        features={"punctuation_enabled": True},
    )

    assert data.paragraphs[0].text != "工作材料,请审阅?"
    assert any(change.applied for change in data.normalization_changes)


@pytest.mark.parametrize(
    ("first_name", "second_name"),
    [
        ("word/document.xml", "word/document.xml"),
        ("word/document.xml", "WORD/DOCUMENT.XML"),
        ("word/document.xml", "word\\document.xml"),
    ],
)
def test_duplicate_zip_members_are_rejected(first_name: str, second_name: str) -> None:
    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("_rels/.rels", "<Relationships/>")
        archive.writestr(first_name, "<document/>")
        archive.writestr(second_name, "<document/>")

    with pytest.raises(DocxValidationError) as error:
        validate_docx_upload(
            payload.getvalue(),
            max_upload_bytes=1_000_000,
            max_uncompressed_bytes=1_000_000,
            max_file_count=100,
            max_xml_bytes=1_000_000,
            max_media_bytes=1_000_000,
            max_compression_ratio=10_000,
        )

    assert error.value.code == "DUPLICATE_ZIP_MEMBER"


def test_document_mode_uses_title_region_not_body_keywords() -> None:
    title = _paragraph("基层治理工作情况", "title", 0, alignment="CENTER")
    body = _paragraph("现将有关情况报告如下，并通知有关单位。", "body", 1)
    data = _document(title, body)

    apply_recognition(data)

    assert data.recognition_diagnostics["mode"] == "unknown"


def test_all_runs_are_aggregated_by_non_whitespace_character_weight() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    prefix = paragraph.add_run("1. ")
    prefix.font.name = "宋体"
    prefix.font.size = Pt(12)
    body = paragraph.add_run("这是较长的标题主体")
    body.font.name = "黑体"
    body.font.size = Pt(22)
    body.font.bold = True
    paragraph.add_run("   ").font.size = Pt(48)

    features = extract_run_features(paragraph, 0)

    assert features.first_run_font_name == "宋体"
    assert features.first_run_font_size_pt == 12
    assert features.dominant_font_name == "黑体"
    assert features.weighted_font_size > 20
    assert features.max_font_size == 22
    assert features.bold_char_ratio > 0.7


def test_run_size_falls_back_to_word_style() -> None:
    document = Document()
    style = document.styles.add_style("Audit Style", 1)
    style.font.name = "楷体"
    style.font.size = Pt(18)
    paragraph = document.add_paragraph("样式字号正文", style=style)

    features = extract_run_features(paragraph, 0)

    assert features.dominant_font_name == "楷体"
    assert features.weighted_font_size == 18
    assert features.explicitly_formatted_char_ratio == 0


def test_local_confidence_is_separate_from_document_path_score() -> None:
    data = _document(
        _paragraph("国发〔2026〕23号", "title_cont", 0),
        _paragraph("补充说明", "body", 1, alignment="CENTER"),
    )
    apply_recognition(data)
    first, second = data.recognition_diagnostics["paragraphs"]

    assert first["recognition_confidence"] != first["document_path_score"]
    assert second["recognition_confidence"] != second["document_path_score"]
    assert "document_path_score" in first
    assert "selected_candidate_score" in first
    assert "candidate_margin" in first
    assert "transition_contribution" in first


@pytest.mark.parametrize(
    ("style_name", "expected"),
    [
        ("Title", ParagraphType.MAIN_TITLE),
        ("标题", ParagraphType.MAIN_TITLE),
        ("Subtitle", ParagraphType.TITLE_CONTINUATION),
        ("副标题", ParagraphType.TITLE_CONTINUATION),
        ("Heading 1", ParagraphType.HEADING_1),
        ("标题 1", ParagraphType.HEADING_1),
        ("Heading 2", ParagraphType.HEADING_2),
        ("标题 2", ParagraphType.HEADING_2),
        ("Heading 3", ParagraphType.HEADING_3),
        ("标题 3", ParagraphType.HEADING_3),
        ("Heading 4", ParagraphType.HEADING_4),
        ("标题 4", ParagraphType.HEADING_4),
        ("Normal", ParagraphType.BODY),
        ("正文", ParagraphType.BODY),
    ],
)
def test_word_style_candidates(style_name: str, expected: ParagraphType) -> None:
    features = SimpleNamespace(style_name=style_name)
    candidates = StyleCandidateProvider().propose(None, features, None)
    assert candidates[0].paragraph_type is expected
    assert candidates[0].evidence[0].startswith("word-style-")


def test_unknown_word_style_does_not_create_body_candidate() -> None:
    features = SimpleNamespace(style_name="Custom Unknown Style")
    assert StyleCandidateProvider().propose(None, features, None) == []


def test_caption_is_a_distinct_block_and_remains_caption_after_recognition() -> None:
    image = _paragraph("", "__image__", 0)
    caption = _paragraph("图1 工作流程", "__object_caption__", 1)
    body = _paragraph("这是普通正文。", "body", 2)
    data = _document(image, caption, body)

    blocks = extract_blocks(data)
    apply_recognition(data)

    assert [item.kind for item in blocks] == [BlockKind.IMAGE, BlockKind.CAPTION, BlockKind.PARAGRAPH]
    assert caption.type_id == "__object_caption__"
    assert caption.meta["recognized_type"] == "caption"


@pytest.mark.parametrize(
    ("value", "expected"),
    [
        (0, "零"), (1, "一"), (10, "十"), (11, "十一"), (20, "二十"),
        (99, "九十九"), (100, "一百"), (101, "一百零一"),
        (110, "一百一十"), (999, "九百九十九"), (1000, "一千"),
        (1001, "一千零一"), (1010, "一千零一十"), (1100, "一千一百"),
        (9999, "九千九百九十九"),
    ],
)
def test_chinese_integer_boundaries(value: int, expected: str) -> None:
    assert chinese_integer(value) == expected


@pytest.mark.parametrize("value", [-1, 10000, True, 1.5])
def test_chinese_integer_rejects_unsupported_values(value) -> None:
    with pytest.raises(ValueError):
        chinese_integer(value)


def test_public_task_state_and_internal_detail_are_sanitized(monkeypatch) -> None:
    task_id = "audit-public-error"
    secret = "secret=top-secret-token"
    local_path = r"C:\private\document.docx"
    monkeypatch.setitem(server.TASKS, task_id, {
        "id": task_id,
        "status": "error",
        "owner_id": "owner-a",
        "error_code": "TASK_PROCESSING_ERROR",
        "error": f"{secret} {local_path}",
        "error_message": f"{secret} {local_path}",
        "internal_error_detail": f"{secret} {local_path}",
    })

    public = server._public_task_state(task_id, "owner-a")
    internal = server._sanitize_internal_error_detail(f"{secret} {local_path}")

    assert public["message"] == "排版失败"
    assert "error" not in public and "error_message" not in public
    assert "top-secret-token" not in internal
    assert "C:\\private" not in internal


def test_public_recognition_summary_contains_no_paragraph_text() -> None:
    diagnostics = {
        "recognition_mode": "shadow",
        "result_applied": False,
        "paragraphs": [{
            "paragraph_index": 4,
            "legacy_type": "body",
            "recognized_type": "heading2",
            "final_type": "body",
            "recognition_confidence": 0.51,
            "review_confidence": 0.62,
            "review_level": "review",
            "candidate_margin": 0.02,
            "needs_review": True,
            "review_reasons": ["SMALL_CANDIDATE_MARGIN"],
            "evidence_summary": ["legacy-agreement"],
            "text_preview": "sensitive-text",
        }],
    }
    summary = server._public_recognition_summary(SimpleNamespace(recognition_diagnostics=diagnostics))

    assert summary["recognition_mode"] == "shadow"
    assert summary["result_applied"] is False
    assert summary["needs_review_count"] == 1
    assert summary["review_count"] == 1
    assert summary["review_items"][0]["confidence"] == 0.62
    assert summary["review_items"][0]["review_level"] == "review"
    assert summary["review_items"][0]["evidence_summary"] == ["legacy-agreement"]
    assert summary["review_items"][0]["final_type"] == "body"
    assert "sensitive-text" not in repr(summary)


def test_runtime_dependencies_and_hashed_lock_are_complete() -> None:
    root = Path(__file__).resolve().parents[1]
    pyproject = (root / "pyproject.toml").read_text(encoding="utf-8")
    requirements = (root / "requirements.txt").read_text(encoding="utf-8")
    lock = (root / "requirements.lock").read_text(encoding="utf-8")

    assert "argon2-cffi" in pyproject
    assert "argon2-cffi" in requirements
    assert "argon2-cffi==" in lock
    assert "python-docx==" in lock
    assert "--hash=sha256:" in lock
    assert "--trusted-host" not in lock


def test_frontend_renders_sanitized_recognition_review_summary() -> None:
    source = (
        Path(__file__).resolve().parents[1]
        / "resources" / "frontend" / "pages" / "index.html"
    ).read_text(encoding="utf-8")

    assert "需要人工复核" in source
    assert "诊断结果未应用" in source
    assert "识别结果已应用" in source
    assert "renderRecognitionSummary(t.recognition_summary)" in source
    assert "li.textContent=" in source
