from __future__ import annotations

from pathlib import Path

from docx import Document
from docx import Document as DocxDocument
from docx.enum.text import WD_BREAK

from docxtool.document.importing.features import extract_paragraph_features
from docxtool.document.importing.numbering import detect_numbering_prefix
from docxtool.document.importing.reader import read_body_blocks
from docxtool.document import importer as importer_module
from docxtool.document.models import DocumentData
from docxtool.document.segmentation.pipeline import build_logical_lines


def _features(paragraph, index):
    return extract_paragraph_features(
        paragraph,
        index,
        detect_numbering_prefix_func=detect_numbering_prefix,
    )


def test_reader_keeps_body_order_and_records_physical_facts(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("第一段正文")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格内容"
    document.add_paragraph("第二段正文")
    document.save(source)

    data = DocumentData()
    raw_blocks = read_body_blocks(
        DocxDocument(source),
        data,
        strict_preservation=False,
        protected_letterhead_indexes=set(),
        extract_features_func=_features,
    )

    assert [block[0] for block in raw_blocks] == ["paragraph", "table", "paragraph"]
    assert raw_blocks[0][1].text == "第一段正文"
    assert raw_blocks[2][1].text == "第二段正文"
    assert len(data.tables) == 1
    assert data.body_sectPr is not None


def test_segmenter_preserves_visible_text_and_source_anchor_after_reader(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    source_text = "一、迁移验证标题。这是一段不少于五个字符的连续正文内容。"
    document = Document()
    document.add_paragraph(source_text)
    document.save(source)

    data = DocumentData()
    raw_blocks = read_body_blocks(
        DocxDocument(source),
        data,
        strict_preservation=False,
        protected_letterhead_indexes=set(),
        extract_features_func=_features,
    )
    logical_lines = build_logical_lines(
        raw_blocks,
        strict_preservation=False,
        structural_preservation=True,
        split_inline_heading_body_enabled=True,
        normalize_text_func=lambda text: text,
        source_starts_body_region_func=importer_module._source_starts_body_region,
        split_inline_heading_body_spans_func=importer_module._split_inline_heading_body_spans,
        validate_numbered_heading_body_split_func=(
            importer_module._validate_numbered_heading_body_split
        ),
        should_split_structural_line_breaks_func=(
            importer_module._should_split_structural_line_breaks
        ),
        split_structural_tail_after_numbered_heading_func=(
            importer_module._split_structural_tail_after_numbered_heading
        ),
        validate_source_span_partition_func=importer_module._validate_source_span_partition,
        detect_numbering_prefix_func=importer_module._detect_numbering_prefix,
        inline_lead_bold_func=importer_module._has_inline_lead_bold_transition,
    )
    text_lines = [line for line in logical_lines if line[0] == "text"]

    assert len(text_lines) == 2
    assert "".join(line[1] for line in text_lines) == source_text
    assert all(line[2].source_physical_paragraph_index == 0 for line in text_lines)
    assert text_lines[0][2].source_start_utf16 == 0
    assert text_lines[-1][2].source_end_utf16 == len(source_text)


def test_segmenter_splits_body_tail_addressing_and_drops_soft_break_tokens(
    tmp_path: Path,
) -> None:
    source = tmp_path / "body-tail-addressing.docx"
    body_text = "正文最后一段完整结束。"
    addressing_text = "各位委员、同志们！"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run(body_text)
    for _unused in range(4):
        paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run(addressing_text)
    document.save(source)

    data = DocumentData()
    raw_blocks = read_body_blocks(
        DocxDocument(source),
        data,
        strict_preservation=False,
        protected_letterhead_indexes=set(),
        extract_features_func=_features,
    )
    logical_lines = build_logical_lines(
        raw_blocks,
        strict_preservation=False,
        structural_preservation=True,
        split_inline_heading_body_enabled=True,
        normalize_text_func=lambda text: text,
        source_starts_body_region_func=importer_module._source_starts_body_region,
        split_inline_heading_body_spans_func=(
            importer_module._split_inline_heading_body_spans
        ),
        validate_numbered_heading_body_split_func=(
            importer_module._validate_numbered_heading_body_split
        ),
        should_split_structural_line_breaks_func=(
            importer_module._should_split_structural_line_breaks
        ),
        split_structural_tail_after_numbered_heading_func=(
            importer_module._split_structural_tail_after_numbered_heading
        ),
        validate_source_span_partition_func=(
            importer_module._validate_source_span_partition
        ),
        detect_numbering_prefix_func=importer_module._detect_numbering_prefix,
        inline_lead_bold_func=importer_module._has_inline_lead_bold_transition,
    )
    text_lines = [line for line in logical_lines if line[0] == "text"]

    assert [line[1] for line in text_lines] == [body_text, addressing_text]
    assert [line[3] for line in text_lines] == [[], []]
    assert [line[2].segment_index for line in text_lines] == [0, 1]
    assert {line[2].segment_count for line in text_lines} == {2}


def test_segmenter_keeps_ordinary_soft_line_before_local_addressing_boundary(
    tmp_path: Path,
) -> None:
    source = tmp_path / "mixed-body-tail-addressing.docx"
    first_body_line = "正文第一行"
    second_body_line = "正文第二行"
    addressing_text = "各位委员、同志们！"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run(first_body_line)
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run(second_body_line)
    paragraph.add_run().add_break(WD_BREAK.LINE)
    paragraph.add_run(addressing_text)
    document.save(source)

    data = DocumentData()
    raw_blocks = read_body_blocks(
        DocxDocument(source),
        data,
        strict_preservation=False,
        protected_letterhead_indexes=set(),
        extract_features_func=_features,
    )
    logical_lines = build_logical_lines(
        raw_blocks,
        strict_preservation=False,
        structural_preservation=True,
        split_inline_heading_body_enabled=True,
        normalize_text_func=lambda text: text,
        source_starts_body_region_func=importer_module._source_starts_body_region,
        split_inline_heading_body_spans_func=(
            importer_module._split_inline_heading_body_spans
        ),
        validate_numbered_heading_body_split_func=(
            importer_module._validate_numbered_heading_body_split
        ),
        should_split_structural_line_breaks_func=(
            importer_module._should_split_structural_line_breaks
        ),
        split_structural_tail_after_numbered_heading_func=(
            importer_module._split_structural_tail_after_numbered_heading
        ),
        validate_source_span_partition_func=(
            importer_module._validate_source_span_partition
        ),
        detect_numbering_prefix_func=importer_module._detect_numbering_prefix,
        inline_lead_bold_func=importer_module._has_inline_lead_bold_transition,
        split_standalone_addressing_spans_func=(
            importer_module._split_standalone_addressing_spans
        ),
    )
    text_lines = [line for line in logical_lines if line[0] == "text"]

    assert [line[1] for line in text_lines] == [
        f"{first_body_line}\n{second_body_line}",
        addressing_text,
    ]
    assert [token.kind for token in text_lines[0][3]] == [
        "text", "line_break", "text",
    ]
    assert text_lines[0][3][0].text == first_body_line
    assert text_lines[0][3][2].text == second_body_line
    assert text_lines[1][3] == []
    assert [line[2].segment_index for line in text_lines] == [0, 1]
    assert {line[2].segment_count for line in text_lines} == {2}
