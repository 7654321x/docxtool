from copy import deepcopy

from docx import Document

from docxtool.document.engine.document_structure import (
    BlockKind,
    ElementKind,
    analyze_document_structure,
)
from docxtool.document.engine.structure_context import (
    ValidationStatus,
    revalidate_element,
    validate_structure_context,
)
from docxtool.document.importer import DocumentData, InlineToken, ParagraphData, ParagraphFeatures


def paragraph(text, type_id, *, confidence=None, style="", alignment="", tokens=(), meta=None):
    values = dict(meta or {})
    if confidence is not None:
        values["classification_confidence"] = confidence
    features = ParagraphFeatures(text=text, style_name=style, alignment=alignment)
    return ParagraphData(text, type_id, text, features, values, list(tokens))


def validate(paragraphs):
    structure = analyze_document_structure(DocumentData(paragraphs=paragraphs))
    return structure, validate_structure_context(structure, paragraphs)


def result_for(validation, *, type_id=None, kind=None):
    for item in validation.elements:
        if type_id is not None and item.original_type_id == type_id:
            return item
        if kind is not None and item.context_kind == kind:
            return item
    raise AssertionError("validated element not found")


def test_title_date_and_title_signature_are_confirmed_metadata():
    paragraphs = [
        paragraph("公文标题", "title"),
        paragraph("测试机关", "author_line"),
        paragraph("2026年7月17日", "date_line"),
        paragraph("正文内容。", "body"),
    ]
    _, validation = validate(paragraphs)

    author = result_for(validation, type_id="author_line")
    date = result_for(validation, type_id="date_line")
    assert author.final_kind == ElementKind.TITLE_METADATA
    assert date.final_kind == ElementKind.TITLE_METADATA
    assert author.status == date.status == ValidationStatus.CONFIRMED
    assert any(item.detail == "title_like_before" for item in date.evidence)


def test_signature_pair_is_confirmed_at_document_tail():
    paragraphs = [
        paragraph("公文标题", "title"), paragraph("正文内容。", "body"),
        paragraph("测试市人民政府", "sign_org"), paragraph("2026年7月17日", "sign_date"),
    ]
    _, validation = validate(paragraphs)

    agency = result_for(validation, type_id="sign_org")
    date = result_for(validation, type_id="sign_date")
    assert agency.final_kind == ElementKind.SIGNATURE_AGENCY
    assert date.final_kind == ElementKind.SIGNATURE_DATE
    assert agency.status == date.status == ValidationStatus.CONFIRMED
    assert any(item.detail == "previous:signature_candidate" for item in date.evidence)


def test_body_date_and_agency_sentence_remain_body():
    paragraphs = [
        paragraph("公文标题", "title"),
        paragraph("第一段正文。", "body"),
        paragraph("2026年7月17日召开会议。", "body"),
        paragraph("由测试市人民政府负责实施。", "body"),
        paragraph("后一段正文。", "body"),
    ]
    structure, validation = validate(paragraphs)

    body_results = [item for item in validation.elements if item.original_type_id == "body"]
    assert structure.signature is None
    assert all(item.final_kind == ElementKind.BODY_PARAGRAPH for item in body_results)
    assert all(item.block_kind == BlockKind.BODY for item in body_results)


def test_body_attachment_mentions_remain_body():
    paragraphs = [
        paragraph("公文标题", "title"),
        paragraph("相关附件另行发送。", "body"),
        paragraph("详见附件内容。", "body"),
    ]
    _, validation = validate(paragraphs)
    assert all(
        item.final_kind == ElementKind.BODY_PARAGRAPH
        for item in validation.elements if item.original_type_id == "body"
    )


def test_attachment_note_and_paged_attachment_title_are_distinguished():
    page = [InlineToken("page_break")]
    paragraphs = [
        paragraph("公文标题", "title"), paragraph("正文内容。", "body"),
        paragraph("附件：测试清单", "attachment_note"),
        paragraph("测试机关", "sign_org"), paragraph("2026年7月17日", "sign_date"),
        paragraph("附件1", "attachment_page_mark", tokens=page),
        paragraph("附件正文", "attachment_body"),
    ]
    _, validation = validate(paragraphs)

    note = result_for(validation, type_id="attachment_note")
    title = result_for(validation, kind=ElementKind.ATTACHMENT_TITLE)
    assert note.final_kind == ElementKind.ATTACHMENT_NOTE_ITEM
    assert note.block_kind == BlockKind.ATTACHMENT_NOTE
    assert title.final_kind == ElementKind.ATTACHMENT_TITLE
    assert title.block_kind == BlockKind.ATTACHMENT_CONTENT
    assert title.status == ValidationStatus.CONFIRMED
    assert any(item.detail == "preceded_by:real_page_boundary" for item in title.evidence)


def test_unpaged_attachment_candidate_is_not_confirmed():
    paragraphs = [
        paragraph("公文标题", "title"), paragraph("正文内容。", "body"),
        paragraph("测试机关", "sign_org"), paragraph("2026年7月17日", "sign_date"),
        paragraph("附件1", "attachment_page_mark"), paragraph("附件正文", "attachment_body"),
    ]
    structure, validation = validate(paragraphs)
    candidate = result_for(validation, type_id="attachment_page_mark")
    assert structure.attachments == ()
    assert candidate.status != ValidationStatus.CONFIRMED
    assert candidate.final_kind == ElementKind.UNKNOWN


def test_matching_high_confidence_sources_raise_confidence():
    paragraphs = [paragraph("标题", "title"), paragraph("正文。", "body", confidence=0.91)]
    structure, validation = validate(paragraphs)
    body = result_for(validation, type_id="body")
    proposed = next(item for item in structure.elements if item.source_index == 1)
    assert body.status == ValidationStatus.CONFIRMED
    assert body.confidence > proposed.confidence


def test_high_block_low_context_is_provisional_without_forced_detail():
    paragraphs = [paragraph("标题", "title"), paragraph("正文。", "body", confidence=0.4)]
    _, validation = validate(paragraphs)
    body = result_for(validation, type_id="body")
    assert body.status == ValidationStatus.PROVISIONAL
    assert body.final_kind == ElementKind.UNKNOWN


def test_low_block_high_context_is_provisional_and_keeps_context_kind():
    paragraphs = [
        paragraph("标题", "title"),
        paragraph("未知但保留", "mystery"),
        paragraph("正文。", "body", confidence=0.95),
    ]
    _, validation = validate(paragraphs)
    body = result_for(validation, type_id="body")
    assert body.status == ValidationStatus.PROVISIONAL
    assert body.final_kind == ElementKind.BODY_PARAGRAPH


def test_high_confidence_block_context_conflict_becomes_unknown():
    paragraphs = [
        paragraph("标题", "title"), paragraph("正文。", "body"),
        paragraph("2026年7月17日", "sign_date", confidence=0.95),
    ]
    _, validation = validate(paragraphs)
    date = result_for(validation, type_id="sign_date")
    assert date.block_kind == BlockKind.BODY
    assert date.status == ValidationStatus.CONFLICT
    assert date.final_kind == ElementKind.UNKNOWN


def test_low_block_and_context_are_unknown():
    paragraphs = [paragraph("无法判断", "mystery", confidence=0.4)]
    _, validation = validate(paragraphs)
    item = validation.elements[0]
    assert item.status == ValidationStatus.UNKNOWN
    assert item.final_kind == ElementKind.UNKNOWN


def test_evidence_is_layered_and_heading_jump_is_only_recorded():
    paragraphs = [
        paragraph("标题", "title"),
        paragraph("一、一级", "heading1"),
        paragraph("1.三级", "heading3"),
        paragraph("正文。", "body"),
    ]
    _, validation = validate(paragraphs)
    heading = result_for(validation, type_id="heading3")
    sources = {item.source for item in heading.evidence}
    assert {"block", "context", "style", "text", "structural"} <= sources
    assert any(item.detail == "heading_level_jump_1_to_3" for item in heading.evidence)
    assert heading.final_kind == ElementKind.HEADING_3


def test_results_cover_input_order_and_are_deterministic():
    paragraphs = [paragraph("标题", "title"), paragraph("正文。", "body")]
    structure = analyze_document_structure(DocumentData(paragraphs=paragraphs))
    first = validate_structure_context(structure, paragraphs)
    second = validate_structure_context(structure, paragraphs)
    assert first == second
    assert [item.index for item in first.elements] == list(range(len(structure.elements)))


def test_revalidate_accepts_unchanged_and_rejects_text_style_and_neighbor_changes():
    paragraphs = [
        paragraph("标题", "title", style="标题"),
        paragraph("正文甲。", "body", style="正文", alignment="JUSTIFY"),
        paragraph("正文乙。", "body", style="正文", alignment="JUSTIFY"),
    ]
    structure, validation = validate(paragraphs)
    target = next(item for item in validation.elements if item.source_index == 1)
    assert revalidate_element(target, structure, paragraphs)

    text_changed = deepcopy(paragraphs)
    text_changed[1].original_text = "正文已变化。"
    assert not revalidate_element(target, structure, text_changed)

    style_changed = deepcopy(paragraphs)
    style_changed[1].features.style_name = "其他样式"
    assert not revalidate_element(target, structure, style_changed)

    neighbor_changed = deepcopy(paragraphs)
    neighbor_changed[2].original_text = "相邻段落已变化。"
    assert not revalidate_element(target, structure, neighbor_changed)


def test_revalidate_checks_preserved_xml_node_identity_and_analysis_is_read_only():
    document = Document()
    xml_paragraph = document.add_paragraph("图片题注")
    paragraphs = [
        paragraph("标题", "title"),
        paragraph("", "__object_caption__", meta={"paragraph_xml": xml_paragraph}),
    ]
    before = document._element.xml
    structure, validation = validate(paragraphs)
    caption = result_for(validation, type_id="__object_caption__")
    assert revalidate_element(caption, structure, paragraphs)

    replacement = deepcopy(paragraphs)
    replacement_document = Document()
    replacement[1].meta["paragraph_xml"] = replacement_document.add_paragraph("图片题注")
    assert not revalidate_element(caption, structure, replacement)
    assert document._element.xml == before
