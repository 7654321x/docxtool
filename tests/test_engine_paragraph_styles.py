from types import SimpleNamespace

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxtool.document.engine.paragraph_styles import (
    enforce_body_paragraph_invariants,
    is_standalone_keep_heading,
    paragraph_style_id,
    remove_paragraph_numbering,
    set_keep_with_next,
    set_paragraph_style_id,
    style_id_for_type,
)
from docxtool.document.engine.render_types import rule_index_for_type


def test_style_id_for_type_maps_known_types_and_falls_back_to_body() -> None:
    assert style_id_for_type("heading1") == "DCT-Heading1"
    assert style_id_for_type("attachment_note_item") == "DCT-AttachmentNoteItem"
    assert style_id_for_type("glossary_item") == "DCT-Body"
    assert style_id_for_type("meeting_title_meta") == "DCT-Author"
    assert rule_index_for_type("meeting_title_meta") == 12
    assert style_id_for_type("unknown_type") == "DCT-Body"


def test_wps_builtin_style_profile_maps_only_body_and_heading_types() -> None:
    assert style_id_for_type("body", "wps_builtin") == "Normal"
    assert style_id_for_type("meeting_meta", "wps_builtin") == "Normal"
    assert style_id_for_type("meeting_title_meta", "wps_builtin") == "DCT-Author"
    assert style_id_for_type("heading1", "wps_builtin") == "Heading1"
    assert style_id_for_type("heading2", "wps_builtin") == "Heading2"
    assert style_id_for_type("title2", "wps_builtin") == "Heading2"
    assert style_id_for_type("heading3", "wps_builtin") == "Heading3"
    assert style_id_for_type("heading4", "wps_builtin") == "Heading4"
    assert style_id_for_type("attachment_note", "wps_builtin") == "DCT-AttachmentNote"
    assert style_id_for_type("unknown_type", "wps_builtin") == "Normal"


def test_wps_docxtool_style_profile_keeps_structural_styles_out_of_builtin_slots() -> None:
    assert style_id_for_type("body", "wps_docxtool") == "DCT-Body"
    assert style_id_for_type("heading1", "wps_docxtool") == "DCT-Heading1"
    assert style_id_for_type("heading4", "wps_docxtool") == "DCT-Heading4"
    assert style_id_for_type("unknown_type", "wps_docxtool") == "DCT-Body"


def test_unknown_style_profile_fails_fast() -> None:
    with pytest.raises(ValueError, match="WPS_STYLE_PROFILE_INVALID"):
        style_id_for_type("body", "unknown-profile")


def test_set_and_read_paragraph_style_id_replaces_existing_style() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("测试")

    set_paragraph_style_id(paragraph, "DCT-Body")
    set_paragraph_style_id(paragraph, "DCT-Heading1")

    properties = paragraph._element.get_or_add_pPr()
    styles = properties.findall(qn("w:pStyle"))
    assert len(styles) == 1
    assert paragraph_style_id(paragraph) == "DCT-Heading1"


def test_set_keep_with_next_writes_single_keep_nodes() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("附件")

    set_keep_with_next(paragraph)
    set_keep_with_next(paragraph)

    properties = paragraph._element.get_or_add_pPr()
    assert len(properties.findall(qn("w:keepNext"))) == 1
    assert len(properties.findall(qn("w:keepLines"))) == 1


def test_remove_paragraph_numbering_removes_numpr_once() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("列表")
    properties = paragraph._element.get_or_add_pPr()
    properties.append(OxmlElement("w:numPr"))

    assert remove_paragraph_numbering(paragraph) is True
    assert remove_paragraph_numbering(paragraph) is False


def test_enforce_body_paragraph_invariants_skips_protected_elements() -> None:
    doc = Document()
    protected = doc.add_paragraph("受保护")
    fallback = doc.add_paragraph("普通正文")

    stats = enforce_body_paragraph_invariants(doc, {protected._p})

    assert stats == {"fallback_count": 1, "numpr_removed": 0}
    assert paragraph_style_id(protected) == ""
    assert paragraph_style_id(fallback) == "DCT-Body"


def test_wps_builtin_invariants_keep_managed_styles_and_fall_back_to_normal() -> None:
    doc = Document()
    heading = doc.add_paragraph("一级标题")
    heading.style = doc.styles["Heading 1"]
    fallback = doc.add_paragraph("普通正文")

    stats = enforce_body_paragraph_invariants(
        doc,
        style_profile="wps_builtin",
    )

    assert stats == {"fallback_count": 1, "numpr_removed": 0}
    assert paragraph_style_id(heading) == "Heading1"
    assert paragraph_style_id(fallback) == "Normal"


def test_is_standalone_keep_heading_only_allows_attachment_mark_and_title() -> None:
    assert is_standalone_keep_heading(SimpleNamespace(type_id="attachment_title"), None, "") is True
    assert is_standalone_keep_heading(SimpleNamespace(type_id="heading1"), None, "") is False
