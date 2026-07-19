import base64
import logging
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK

from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import StyleRule, logger


def _rules():
    return [StyleRule.default_for_row(i) for i in range(10)]


class SignatureDetectionTest(unittest.TestCase):
    def setUp(self):
        logger.setLevel(logging.ERROR)
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)

    def tearDown(self):
        self.tmp.cleanup()

    def _load_lines(self, lines):
        doc = Document()
        for line in lines:
            doc.add_paragraph(line)
        path = self.root / "input.docx"
        doc.save(path)
        return DocxImporter().load(str(path), _rules())

    def test_normal_document_signature_after_body(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "区政府人才保障工作组",
            "2025年十月15日",
        ])

        self.assertEqual([p.type_id for p in data.paragraphs[-2:]], ["sign_org", "sign_date"])
        self.assertEqual(data.paragraphs[-2].text, "区政府人才保障工作组")
        self.assertEqual(data.paragraphs[-1].text, "2025年10月15日")

    def test_attachment_note_signature_and_attachment_page_flow(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "附件：1. 基本情况",
            "2. 具体情况",
            "区政府人才保障工作组",
            "2025年10月15日",
            "附件1",
            "标题",
            "测试正文测试正文。",
        ])

        self.assertEqual(
            [p.type_id for p in data.paragraphs[-7:]],
            [
                "attachment_note",
                "attachment_note_item",
                "sign_org",
                "sign_date",
                "attachment_page_mark",
                "attachment_title",
                "attachment_body",
            ],
        )
        self.assertEqual(data.paragraphs[-4].text, "2025年10月15日")
        self.assertEqual(data.paragraphs[-3].text, "附件 1")

    def test_chinese_year_signature_date_is_normalized(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "内江市东兴区人民政府办公室",
            "二〇二五年十月十五日",
        ])

        self.assertEqual([p.type_id for p in data.paragraphs[-2:]], ["sign_org", "sign_date"])
        self.assertEqual(data.paragraphs[-1].text, "2025年10月15日")

    def test_short_signature_org_before_spaced_date_is_detected(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "区政协",
            "2025 年 10 月 15 日",
        ])

        self.assertEqual([p.type_id for p in data.paragraphs[-2:]], ["sign_org", "sign_date"])
        self.assertEqual(data.paragraphs[-2].text, "区政协")
        self.assertEqual(data.paragraphs[-1].text, "2025年10月15日")

    def test_long_role_and_name_line_is_detected_in_head_area(self):
        data = self._load_lines([
            "2026年度测试材料",
            "区政协办公室党组书记、主任  李弟弟",
            "（2026年7月14日）",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
        ])

        self.assertEqual(data.paragraphs[1].type_id, "role_name")
        self.assertEqual(data.paragraphs[1].text, "区政协办公室党组书记、主任  李弟弟")

    def test_office_director_and_name_line_is_detected_in_head_area(self):
        data = self._load_lines([
            "2026年度测试材料",
            "区政协办公室主任  李弟弟",
            "（2026年7月14日）",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
        ])

        self.assertEqual(data.paragraphs[1].type_id, "role_name")
        self.assertEqual(data.paragraphs[1].text, "区政协办公室主任  李弟弟")

    def test_role_name_is_split_from_soft_broken_title_block(self):
        doc = Document()
        paragraph = doc.add_paragraph("中共某区政协党组班子")
        paragraph.add_run().add_break()
        paragraph.add_run("2025年度民主生活会对照检查材料")
        paragraph.add_run().add_break()
        paragraph.add_run("区政协办公室主任  李某某")
        doc.add_paragraph("一、一级标题")
        doc.add_paragraph("这里是正文内容这里是正文内容这里是正文内容。")
        path = self.root / "soft-broken-title.docx"
        doc.save(path)

        data = DocxImporter().load(str(path), _rules())

        self.assertEqual(data.paragraphs[2].type_id, "role_name")
        self.assertEqual(data.paragraphs[2].text, "区政协办公室主任  李某某")

    def test_signature_org_is_split_from_body_before_date(self):
        doc = Document()
        doc.add_paragraph("总题目")
        doc.add_paragraph("一、一级标题")
        paragraph = doc.add_paragraph("这里是正文内容这里是正文内容这里是正文内容。")
        for _ in range(4):
            paragraph.add_run().add_break()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        paragraph.add_run().add_tab()
        paragraph.add_run("区政协办")
        trailing = doc.add_paragraph("2025年10月15日")
        trailing.add_run().add_break()
        trailing.add_run().add_break()
        trailing.add_run("附件：1. 基本情况")
        trailing.add_run().add_break()
        trailing.add_run("2. 具体情况")
        path = self.root / "soft-broken-signature.docx"
        doc.save(path)

        data = DocxImporter().load(str(path), _rules())
        tail = [(item.type_id, item.text) for item in data.paragraphs[-4:]]

        self.assertEqual(tail[0], ("attachment_note", "附件：1. 基本情况"))
        self.assertEqual(tail[1], ("attachment_note_item", "2. 具体情况"))
        self.assertEqual(tail[2], ("sign_org", "区政协办"))
        self.assertEqual(tail[3], ("sign_date", "2025年10月15日"))

    def test_soft_broken_tail_splits_signature_note_and_attachment_pages(self):
        doc = Document()
        doc.add_paragraph("总题目")
        doc.add_paragraph("一、一级标题")
        body = doc.add_paragraph("这里是正文内容这里是正文内容这里是正文内容。")
        body.add_run().add_break()
        body.add_run().add_break()
        body.add_run("区政协办")
        tail = doc.add_paragraph("2025年十月15日")
        for text in (
            "附件：1.基本情况",
            "2.具体情况",
            "3.超级情况",
            "附件1",
            "宣传材料",
            "第一份附件正文。",
            "附件2",
            "具体情况",
            "第二份附件正文。",
        ):
            tail.add_run().add_break()
            tail.add_run(text)
        path = self.root / "soft-broken-tail-with-attachments.docx"
        doc.save(path)

        data = DocxImporter().load(
            str(path),
            _rules(),
            features={"punctuation": {"enabled": True, "mode": "safe"}},
        )

        self.assertEqual(
            [item.type_id for item in data.paragraphs[-11:]],
            [
                "attachment_note",
                "attachment_note_item",
                "attachment_note_item",
                "sign_org",
                "sign_date",
                "attachment_page_mark",
                "attachment_title",
                "attachment_body",
                "attachment_page_mark",
                "attachment_title",
                "attachment_body",
            ],
        )
        self.assertEqual(data.paragraphs[-11].text, "附件：1. 基本情况")
        self.assertEqual(data.paragraphs[-10].text, "2. 具体情况")
        self.assertEqual(data.paragraphs[-8].text, "区政协办")

    def test_body_styled_numbered_paragraph_still_detects_inline_heading2(self):
        doc = Document()
        doc.add_paragraph("四、产生问题的原因")
        paragraph = doc.add_paragraph(
            "（三）宗旨意识有所不足。对群众需求掌握不够精准，服务实效仍需提升。"
        )
        paragraph.style = doc.styles["Normal"]
        path = self.root / "body-styled-heading2.docx"
        doc.save(path)

        data = DocxImporter().load(str(path), _rules())

        self.assertEqual(data.paragraphs[1].type_id, "heading2")
        self.assertTrue(data.paragraphs[1].meta.get("heading_inline_body"))

    def test_zero_height_drawing_does_not_hide_inline_heading2_text(self):
        image = self.root / "pixel.png"
        image.write_bytes(base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        ))
        doc = Document()
        doc.add_paragraph("四、产生问题的原因")
        paragraph = doc.add_paragraph("（三）宗旨意识有所不足。")
        picture = paragraph.add_run().add_picture(str(image))
        picture._inline.extent.cy = 0
        paragraph.add_run("对群众需求掌握不够精准，服务实效仍需提升。")
        path = self.root / "zero-height-drawing.docx"
        doc.save(path)

        data = DocxImporter().load(str(path), _rules())

        self.assertEqual(data.paragraphs[1].type_id, "heading2")
        self.assertIn("宗旨意识有所不足", data.paragraphs[1].text)
        self.assertIn("服务实效仍需提升", data.paragraphs[1].text)

    def test_leading_soft_break_is_not_preserved_in_body_tokens(self):
        doc = Document()
        doc.add_paragraph("五、下一步努力方向和整改措施")
        doc.add_paragraph("（一）强化理论武装，筑牢思想政治根基")
        doc.add_paragraph("1.测试测试")
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break()
        paragraph.add_run("坚持理论学习，不断提升履职能力。")
        path = self.root / "leading-soft-break.docx"
        doc.save(path)

        data = DocxImporter().load(str(path), _rules())
        body = data.paragraphs[-1]

        self.assertEqual(body.type_id, "body")
        self.assertEqual(body.text, "坚持理论学习，不断提升履职能力。")
        self.assertEqual(body.inline_tokens, [])

    def test_contact_line_before_date_is_not_signature(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "联系人：张三",
            "2025年10月15日",
        ])

        self.assertNotEqual(data.paragraphs[-2].type_id, "sign_org")
        self.assertNotEqual(data.paragraphs[-1].type_id, "sign_date")

    def test_unpunctuated_body_summary_before_date_is_not_signature(self):
        data = self._load_lines([
            "总题目",
            "一、一级标题",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "以上情况请审阅",
            "2025年10月15日",
        ])

        self.assertNotEqual(data.paragraphs[-2].type_id, "sign_org")
        self.assertNotEqual(data.paragraphs[-1].type_id, "sign_date")

    def test_signature_date_before_attachment_note_is_reordered(self):
        data = self._load_lines([
            "总题目",
            "一、存在的问题",
            "（一）带头强化政治忠诚、提高政治能力方面。",
            "这里是正文内容这里是正文内容这里是正文内容。",
            "（五）扛牢政治责任，推进全面从严治党。坚决扛起管党治党主体责任，严格落实全面从严治党相关规定。",
            "六、区政协办",
            "2025年十月15日",
            "附件：1. 基本情况",
            "2. 具体情况",
            "3. 超级情况",
            "附件1",
            "标题",
            "测试正文测试正文。",
        ])

        self.assertEqual(
            [p.type_id for p in data.paragraphs[-8:]],
            [
                "attachment_note",
                "attachment_note_item",
                "attachment_note_item",
                "sign_org",
                "sign_date",
                "attachment_page_mark",
                "attachment_title",
                "attachment_body",
            ],
        )
        self.assertEqual(data.paragraphs[-5].text, "区政协办")
        self.assertEqual(data.paragraphs[-4].text, "2025年10月15日")

    def test_interleaved_attachment_signature_tail_is_canonicalized(self):
        variants = [
            [
                "测试市人民政府办公室", "附件：1.第一项", "2026年7月19日",
                "2.第二项", "3.第三项",
            ],
            [
                "附件：1.第一项", "测试市人民政府办公室", "2.第二项",
                "3.第三项", "2026年7月19日",
            ],
            [
                "附件：1.第一项", "2026年7月19日", "2.第二项",
                "测试市人民政府办公室", "3.第三项",
            ],
            [
                "测试市人民政府办公室", "2026年7月19日", "附件：1.第一项",
                "2.第二项", "3.第三项",
            ],
        ]
        expected_types = [
            "attachment_note", "attachment_note_item", "attachment_note_item",
            "sign_org", "sign_date", "attachment_page_mark",
            "attachment_title", "attachment_body",
        ]
        for tail in variants:
            with self.subTest(tail=tail):
                data = self._load_lines([
                    "总题目",
                    "一、一级标题",
                    "这里是正文内容这里是正文内容这里是正文内容。",
                    *tail,
                    "附件一百",
                    "第一百份附件标题",
                    "第一百份附件正文。",
                ])
                self.assertEqual([item.type_id for item in data.paragraphs[-8:]], expected_types)
                self.assertEqual(data.paragraphs[-8].text, "附件：1. 第一项")
                self.assertEqual(data.paragraphs[-7].text, "2. 第二项")
                self.assertEqual(data.paragraphs[-3].text, "附件 100")

    def test_role_name_accepts_single_space_and_tab_after_role_keyword(self):
        for role_line in ("测试办公室主任 张测试", "测试办公室主任\t张测试"):
            with self.subTest(role_line=role_line):
                data = self._load_lines([
                    "测试单位领导班子",
                    "2026年度民主生活会对照检查材料",
                    role_line,
                    "一、一级标题",
                    "这里是正文内容这里是正文内容这里是正文内容。",
                ])
                role = next(item for item in data.paragraphs if "张测试" in item.text)
                self.assertEqual(role.type_id, "role_name")


if __name__ == "__main__":
    unittest.main()
