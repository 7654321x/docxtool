from __future__ import annotations

import zipfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from docxtool.document import importer as importer_module
from docxtool.document.importer import DocxImporter
from docxtool.document.style_config import StyleRule


def _rules() -> list[StyleRule]:
    return [StyleRule.default_for_row(index) for index in range(10)]


def test_importer_removes_namespaced_null_relationship(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    broken = tmp_path / "broken.docx"
    document = Document()
    document.add_paragraph("一、测试标题")
    document.add_paragraph("这是用于验证损坏关系修复的脱敏正文。")
    document.save(source)

    rel_name = "word/_rels/document.xml.rels"
    namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    with zipfile.ZipFile(source) as input_archive, zipfile.ZipFile(
        broken, "w", zipfile.ZIP_DEFLATED
    ) as output_archive:
        for item in input_archive.infolist():
            data = input_archive.read(item.filename)
            if item.filename == rel_name:
                root = ET.fromstring(data)
                relationship = ET.SubElement(root, f"{{{namespace}}}Relationship")
                relationship.set("Id", "rIdBrokenNull")
                relationship.set(
                    "Type",
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
                )
                relationship.set("Target", "../NULL")
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                assert b"ns0:Relationship" in data
            output_archive.writestr(item, data)

    before = set(tmp_path.glob("*.docx"))
    imported = DocxImporter().load(str(broken), _rules())

    assert any("损坏关系修复" in paragraph.text for paragraph in imported.paragraphs)
    assert broken.exists()
    assert set(tmp_path.glob("*.docx")) == before

    with ThreadPoolExecutor(max_workers=4) as executor:
        results = list(executor.map(lambda _: DocxImporter().load(str(broken), _rules()), range(8)))
    assert all(result.paragraphs for result in results)
    assert set(tmp_path.glob("*.docx")) == before


def test_repaired_temporary_file_is_deleted_when_parse_fails(tmp_path: Path, monkeypatch) -> None:
    original = tmp_path / "original.docx"
    original.write_bytes(b"original-user-file")
    repaired = tmp_path / "temporary-repair.docx"
    repaired.write_bytes(b"not-a-docx")
    monkeypatch.setattr(importer_module, "_repair_broken_rels", lambda _path: str(repaired))

    with pytest.raises(importer_module.ImportError):
        DocxImporter().load(str(original), _rules())

    assert original.read_bytes() == b"original-user-file"
    assert not repaired.exists()
