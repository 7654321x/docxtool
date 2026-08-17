"""渲染阶段行内 run 写入辅助。

本模块只处理输出段落内部的 run 复制、token 写回和冗余分页符清理。
它不判断段落类型，不参与识别，也不改变输入文本顺序。
"""

from __future__ import annotations

import copy
from collections.abc import Callable

from docx.enum.text import WD_BREAK
from docx.shared import Pt


def copy_run_style(src_run, dst_run) -> None:
    """复制 run 的直接样式。

    传入源 run 和目标 run；返回 None。用于拆写行内片段时保留原 run 的
    rPr 属性，避免字体、字号等直接格式丢失。
    """
    src_rpr = getattr(src_run._element, "rPr", None)
    if src_rpr is None:
        return
    dst_rpr = getattr(dst_run._element, "rPr", None)
    if dst_rpr is not None:
        dst_run._element.remove(dst_rpr)
    dst_run._element.insert(0, copy.deepcopy(src_rpr))


def segment_writer(para, *, set_run_fonts: Callable | None = None):
    """创建段落片段写入器。

    传入目标段落和可选字体设置回调；返回内部 write(text, ...) 函数。
    write 会清空原 run 文本，再逐段写入新文本，并复用首个 run 的样式。
    """
    if not para.runs:
        para.add_run("")
    base_run = para.runs[0]
    for run in para.runs:
        run.text = ""
    used_base = False

    def write(text: str, *, bold=None, cn_font=None, size_pt=None):
        """写入一个行内片段。

        传入文本和可选加粗、中文字体、字号；返回新写入的 run，空文本
        返回 None。
        """
        nonlocal used_base
        if not text:
            return None
        run = base_run if not used_base else para.add_run("")
        if used_base:
            copy_run_style(base_run, run)
        used_base = True
        run.text = text
        if cn_font and set_run_fonts is not None:
            set_run_fonts(run, cn_font=cn_font, en_font="Times New Roman")
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        return run

    return write


def write_inline_tokens(para, tokens) -> None:
    """把行内 token 写回段落。

    传入目标段落和 InlineToken 序列；返回 None。支持普通文本、制表符、
    软换行和分页符。
    """
    for token in tokens or []:
        run = para.add_run("")
        kind = getattr(token, "kind", "")
        if kind == "text":
            run.text = getattr(token, "text", "")
        elif kind == "tab":
            run.add_tab()
        elif kind == "line_break":
            run.add_break()
        elif kind == "page_break":
            run.add_break(WD_BREAK.PAGE)


def without_redundant_trailing_body_page_breaks(pd, next_pd, tokens):
    """清理普通正文中冗余的分页 token。

    传入当前段落数据、下一段落数据和 token 序列；返回新的 token list。
    只移除正文句中或末尾明显多余的分页符，不处理附件、图片、表格边界。
    """
    values = list(tokens or [])
    if pd.type_id != "body" or not pd.text.strip() or next_pd is None:
        return values
    if next_pd.type_id in {
        "attachment_page_mark", "attachment_title", "attachment_body", "glossary_title",
        "__table__", "__image__", "__object_caption__",
    }:
        return values
    last_text = max(
        (
            index for index, token in enumerate(values)
            if getattr(token, "kind", "") == "text" and getattr(token, "text", "")
        ),
        default=-1,
    )
    if last_text < 0:
        return values
    filtered = []
    for index, token in enumerate(values):
        if getattr(token, "kind", "") != "page_break":
            filtered.append(token)
            continue
        preceding = "".join(
            getattr(item, "text", "")
            for item in values[:index]
            if getattr(item, "kind", "") == "text"
        ).rstrip()
        following = "".join(
            getattr(item, "text", "")
            for item in values[index + 1:]
            if getattr(item, "kind", "") == "text"
        ).lstrip()
        if preceding.endswith(("，", ",", "、", "；", ";", "：", ":")) and following:
            continue
        filtered.append(token)
    values = filtered
    last_text = max(
        (
            index for index, token in enumerate(values)
            if getattr(token, "kind", "") == "text" and getattr(token, "text", "")
        ),
        default=-1,
    )
    trailing = values[last_text + 1:]
    if trailing and all(
        getattr(token, "kind", "") == "page_break"
        or (getattr(token, "kind", "") == "text" and not getattr(token, "text", ""))
        for token in trailing
    ):
        return values[:last_text + 1]
    return values
