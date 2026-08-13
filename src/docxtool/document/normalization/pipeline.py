"""Importer-compatible post-recognition normalization orchestration.

The importer keeps the established chain order and compatibility methods.
This module only hosts the legacy normalization operations that run after the
recognition call has completed.  Every operation is injected by the importer
so existing private monkeypatch boundaries remain effective.
"""

from __future__ import annotations

from typing import Any, Callable, Iterable, Sequence


def capture_pre_normalization_snapshot(
    data: Any,
    source_visible_texts: Iterable[str],
) -> list[tuple[str, str, str]]:
    """Build the legacy normalization ledger's pre-change snapshot.

    The function preserves the importer's former physical-source ordering and
    type fallback exactly.  It does not inspect, classify, normalize, or
    mutate paragraphs.
    """
    visible_paragraphs = [
        paragraph
        for paragraph in data.paragraphs
        if not paragraph.type_id.startswith("__")
    ]
    return [
        (
            source_text,
            source_text,
            visible_paragraphs[index].type_id
            if index < len(visible_paragraphs)
            else "",
        )
        for index, source_text in enumerate(source_visible_texts)
    ]


def strip_word_auto_numbering(
    paragraph: Any,
    *,
    log_debug: Callable[..., Any],
) -> None:
    """Remove a paragraph's native Word numbering marker when present.

    This is the original importer operation moved unchanged.  Its broad
    exception guard is retained because the import path historically treats
    malformed or non-python-docx paragraph objects as no-op cleanup cases.
    """
    try:
        from docx.oxml.ns import qn

        pPr = paragraph._element.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                pPr.remove(numPr)
                log_debug("[导入] 剥离自动编号 chars=%s", len(paragraph.text))
    except Exception:
        pass


def merge_uniform_heading_siblings(
    paragraphs: list[Any],
    *,
    log_debug: Callable[..., Any],
    log_info: Callable[..., Any],
) -> None:
    """Apply the importer's existing legacy sibling-heading merge unchanged."""
    PARENT_KEYS = {"heading1", "heading2", "heading3"}
    changed = True
    while changed:
        changed = False
        for i in range(len(paragraphs)):
            pd = paragraphs[i]
            if pd.type_id not in PARENT_KEYS:
                continue
            parent_lvl = int(pd.type_id[-1])
            parent_key = pd.type_id
            j = i + 1
            siblings = []
            while j < len(paragraphs):
                t = paragraphs[j].type_id
                if t in PARENT_KEYS and int(t[-1]) <= parent_lvl:
                    break
                if t.startswith("heading"):
                    siblings.append(j)
                j += 1
            if len(siblings) >= 2:
                levels = {int(paragraphs[s].type_id[-1]) for s in siblings}
                if len(levels) == 1:
                    target = f"heading{min(parent_lvl + 1, 4)}"
                    if any(paragraphs[s].meta.get("numbering") for s in siblings):
                        log_debug("[同级合并] 跳过：siblings 已有编号")
                        continue
                    if any(paragraphs[s].type_id != target for s in siblings):
                        for s in siblings:
                            paragraphs[s].type_id = target
                            paragraphs[s].meta["numbering"] = ""
                        log_info(f"[同级合并] {parent_key}下{len(siblings)}项L{max(levels)}→{target}")
                        changed = True


def apply_post_recognition_normalization(
    data: Any,
    rules: list[Any],
    document_paragraphs: Sequence[Any],
    *,
    strict_preservation: bool,
    structural_preservation: bool,
    processing_strategy: str,
    numbering_enabled: bool,
    before_normalization: list[tuple[str, str, str]],
    normalize_tail_structures_func: Callable[..., None],
    reorder_attachment_note_before_signature_func: Callable[[list[Any]], None],
    assign_numbering_func: Callable[..., None],
    record_applied_normalization_changes_func: Callable[[Any, list[tuple[str, str, str]]], None],
    fix_numbering_gaps_func: Callable[[list[Any]], None],
    strip_auto_numbering_func: Callable[[Any], None],
    sync_recognition_consistency_func: Callable[[Any], None],
) -> None:
    """Run text/order normalization after Recognition finalizes semantic types.

    Existing importer callbacks remain injected for compatibility, but this
    stage may not reclassify paragraphs. Heading-family decisions belong to
    Recognition and are guarded by the semantic-type invariant below.
    """
    semantic_types = {id(paragraph): paragraph.type_id for paragraph in data.paragraphs}
    local_scope = getattr(data, "format_scope", None) is not None
    if structural_preservation and not local_scope:
        normalize_tail_structures_func(data.paragraphs, normalize_text=False)
    elif not strict_preservation:
        normalize_tail_structures_func(data.paragraphs)
        reorder_attachment_note_before_signature_func(data.paragraphs)
        assign_numbering_func(data.paragraphs, rules)
        record_applied_normalization_changes_func(data, before_normalization)

    if structural_preservation and numbering_enabled:
        assign_numbering_func(data.paragraphs, rules)
        for paragraph in data.paragraphs:
            if paragraph.type_id.startswith("heading"):
                paragraph.meta["numbering_correction"] = True

    if processing_strategy == "normalize":
        fix_numbering_gaps_func(data.paragraphs)

    if processing_strategy == "normalize":
        for paragraph in document_paragraphs:
            strip_auto_numbering_func(paragraph)

    sync_recognition_consistency_func(data)
    changed_types = [
        (semantic_types[id(paragraph)], paragraph.type_id)
        for paragraph in data.paragraphs
        if id(paragraph) in semantic_types
        and semantic_types[id(paragraph)] != paragraph.type_id
    ]
    if changed_types:
        raise ValueError(
            "post-recognition normalization changed final semantic type: "
            f"{changed_types[0][0]} -> {changed_types[0][1]}"
        )
