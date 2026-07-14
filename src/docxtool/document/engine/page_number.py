"""Safe page-number footer helpers."""

from __future__ import annotations

import re
import copy
from collections.abc import Mapping, Sequence
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


_PAGE_INSTRUCTION_RE = re.compile(r"\b(?:PAGE|NUMPAGES)\b", re.IGNORECASE)
_PAGE_ONLY_TEXT_RE = re.compile(r"[\s\d/\\\-–—第页共总计：:PagePAGEofNUMPAGES]+", re.IGNORECASE)
_PAGE_DECORATION_TEXT_RE = re.compile(r"[\s\d/\\\-–—第页共总计：:Pageof]+", re.IGNORECASE)


def apply_page_number(document, options: Mapping[str, Any] | None = None):
    """Apply clean Word PAGE fields to document footers."""

    return apply_page_numbers(document, options)


def apply_page_numbers(document, options: Mapping[str, Any] | None = None):
    """Apply Word PAGE/NUMPAGES fields to section footers without erasing other content."""

    opts = dict(options or {})
    if opts.get("enabled", True) is False:
        return document

    style = str(opts.get("style", opts.get("format", "dash"))).lower()
    position = str(opts.get("position", opts.get("alignment", "outside"))).lower()
    first_page_policy = _first_page_policy(opts)
    section_numbering = str(opts.get("section_numbering", "continue")).lower()
    restart_at = int(opts.get("restart_at", 1))

    if position == "outside":
        _set_even_and_odd_headers(document, True)
    has_even_footer = position == "outside" or _even_and_odd_headers_enabled(document)
    _set_update_fields_on_open(document)

    for section_index, section in enumerate(document.sections):
        _set_footer_distance(section, opts)
        _set_section_numbering(section, section_numbering, restart_at, opts, section_index)
        first_page_snapshot = None
        first_page_owned = False
        if first_page_policy in {"hide", "hidden", "no", "skip"}:
            section.different_first_page_header_footer = True
            first_page_footer = section.first_page_footer
            first_page_owned = not first_page_footer.is_linked_to_previous
            if first_page_owned:
                first_page_snapshot = copy.deepcopy(first_page_footer._element)

        _apply_to_footer(section.footer, _alignment_for(position, "default"), style, opts)
        if has_even_footer:
            _apply_to_footer(section.even_page_footer, _alignment_for(position, "even"), style, opts)

        if first_page_policy in {"hide", "hidden", "no", "skip"}:
            if first_page_owned and first_page_snapshot is not None:
                _replace_footer_content(section.first_page_footer, first_page_snapshot)
                _remove_existing_page_numbers(section.first_page_footer)
        elif first_page_policy in {"show", "same", "display"} or section.different_first_page_header_footer:
            section.different_first_page_header_footer = True
            _apply_to_footer(section.first_page_footer, _alignment_for(position, "first"), style, opts)

    return document


def _first_page_policy(options: Mapping[str, Any]) -> str:
    if "first_page" in options:
        raw = options["first_page"]
    elif "first_page_policy" in options:
        raw = options["first_page_policy"]
    else:
        return "default"
    if isinstance(raw, bool):
        return "show" if raw else "hide"
    normalized = str(raw).strip().lower()
    if normalized in {"true", "yes", "1"}:
        return "show"
    if normalized in {"false", "no", "0"}:
        return "hide"
    return normalized or "default"


def _set_footer_distance(section, options: Mapping[str, Any]) -> None:
    if options.get("offset_from_text_mm") is None:
        return
    try:
        offset_cm = float(options.get("offset_from_text_mm")) / 10
    except (TypeError, ValueError):
        return
    bottom_margin_cm = section.bottom_margin.cm if section.bottom_margin else 3.5
    section.footer_distance = Cm(max(0.3, bottom_margin_cm - offset_cm))


def _replace_footer_content(footer, snapshot) -> None:
    footer.is_linked_to_previous = False
    element = footer._element
    for child in list(element):
        element.remove(child)
    for child in snapshot:
        element.append(copy.deepcopy(child))


def _apply_to_footer(footer, alignment, style: str, options: Mapping[str, Any]) -> None:
    footer.is_linked_to_previous = False
    _remove_existing_page_numbers(footer)
    paragraph = footer.add_paragraph()
    paragraph.alignment = alignment
    _write_page_number(paragraph, style, options)


def _write_page_number(paragraph, style: str, options: Mapping[str, Any]) -> None:
    if style in {"plain", "number", "page"}:
        _add_field(paragraph, "PAGE", options)
        return
    if style in {"cn", "chinese", "第page页"}:
        _add_text(paragraph, "第 ", options)
        _add_field(paragraph, "PAGE", options)
        _add_text(paragraph, " 页", options)
        return
    if style in {"cn_total", "chinese_total", "page_numpages", "第page页共numpages页"}:
        _add_text(paragraph, "第 ", options)
        _add_field(paragraph, "PAGE", options)
        _add_text(paragraph, " 页 共 ", options)
        _add_field(paragraph, "NUMPAGES", options)
        _add_text(paragraph, " 页", options)
        return
    _add_text(paragraph, "— ", options)
    _add_field(paragraph, "PAGE", options)
    _add_text(paragraph, " —", options)


def _add_text(paragraph, text: str, options: Mapping[str, Any]):
    run = paragraph.add_run(text)
    _style_run(run, options)
    return run


def _add_field(paragraph, instruction: str, options: Mapping[str, Any]) -> None:
    begin = paragraph.add_run()
    _style_run(begin, options)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_begin.set(qn("w:dirty"), "true")
    begin._r.append(field_begin)

    instr = paragraph.add_run()
    _style_run(instr, options)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    _style_run(separate, options)
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    separate._r.append(field_separate)

    end = paragraph.add_run()
    _style_run(end, options)
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    end._r.append(field_end)


def _style_run(run, options: Mapping[str, Any]) -> None:
    font_name = options.get("font_name") or options.get("font")
    if font_name:
        run.font.name = str(font_name)
        r_fonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            run._element.get_or_add_rPr().append(r_fonts)
        for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
            r_fonts.set(qn(attr), str(font_name))
    if options.get("font_size_pt") is not None:
        run.font.size = Pt(float(options["font_size_pt"]))
    if options.get("bold") is not None:
        run.font.bold = bool(options["bold"])


def _alignment_for(position: str, footer_kind: str):
    normalized = {"right": "right", "left": "left", "center": "center", "centre": "center"}.get(position, position)
    if normalized == "outside":
        return WD_ALIGN_PARAGRAPH.LEFT if footer_kind == "even" else WD_ALIGN_PARAGRAPH.RIGHT
    if normalized == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if normalized == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.RIGHT


def _set_even_and_odd_headers(document, enabled: bool) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:evenAndOddHeaders"))
    if not enabled:
        if existing is not None:
            settings.remove(existing)
        return
    if existing is None:
        existing = OxmlElement("w:evenAndOddHeaders")
        settings.append(existing)
    existing.set(qn("w:val"), "1")


def _even_and_odd_headers_enabled(document) -> bool:
    existing = document.settings._element.find(qn("w:evenAndOddHeaders"))
    if existing is None:
        return False
    return existing.get(qn("w:val"), "1") not in {"0", "false", "False"}


def _set_update_fields_on_open(document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def _set_section_numbering(section, mode: str, restart_at: int, options: Mapping[str, Any], section_index: int) -> None:
    starts = options.get("section_starts")
    if isinstance(starts, Sequence) and not isinstance(starts, (str, bytes)):
        if section_index < len(starts) and starts[section_index] is not None:
            _set_pg_num_start(section, int(starts[section_index]))
        else:
            _clear_pg_num_start(section)
        return
    if mode in {"restart", "restart_each_section", "new"}:
        _set_pg_num_start(section, restart_at)
    else:
        _clear_pg_num_start(section)


def _set_pg_num_start(section, start: int) -> None:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def _clear_pg_num_start(section) -> None:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        return
    pg_num_type.attrib.pop(qn("w:start"), None)
    if not pg_num_type.attrib and len(pg_num_type) == 0:
        section._sectPr.remove(pg_num_type)


def _remove_existing_page_numbers(footer) -> None:
    for paragraph in list(footer.paragraphs):
        if not _paragraph_has_page_number(paragraph):
            continue
        if _is_page_number_only_paragraph(paragraph):
            paragraph._element.getparent().remove(paragraph._element)
            continue
        _remove_page_field_runs(paragraph)
        _remove_literal_page_tokens(paragraph)


def _paragraph_has_page_number(paragraph) -> bool:
    return bool(_field_instruction_text(paragraph))


def _field_instruction_text(paragraph) -> str:
    return " ".join(instr.text or "" for instr in paragraph._element.findall(".//" + qn("w:instrText")))


def _is_page_number_only_paragraph(paragraph) -> bool:
    visible_text = "".join(text.text or "" for text in paragraph._element.findall(".//" + qn("w:t")))
    visible_text = _PAGE_ONLY_TEXT_RE.sub("", visible_text)
    return not visible_text.strip()


def _remove_page_field_runs(paragraph) -> None:
    runs = list(paragraph._element.findall(qn("w:r")))
    remove_indexes: set[int] = set()
    for index, run in enumerate(runs):
        instruction = " ".join(instr.text or "" for instr in run.findall(".//" + qn("w:instrText")))
        if not _PAGE_INSTRUCTION_RE.search(instruction):
            continue
        start = index
        while start > 0 and not _has_field_char(runs[start], "begin"):
            start -= 1
        end = index
        while end < len(runs) - 1 and not _has_field_char(runs[end], "end"):
            end += 1
        while start > 0 and _is_page_decoration_run(runs[start - 1]):
            start -= 1
        while end < len(runs) - 1 and _is_page_decoration_run(runs[end + 1]):
            end += 1
        remove_indexes.update(range(start, end + 1))
    for element in paragraph._element.findall(qn("w:fldSimple")):
        instruction = element.get(qn("w:instr"), "")
        if _PAGE_INSTRUCTION_RE.search(instruction):
            paragraph._element.remove(element)
    for index in sorted(remove_indexes, reverse=True):
        paragraph._element.remove(runs[index])


def _has_field_char(run, field_char_type: str) -> bool:
    return any(
        field_char.get(qn("w:fldCharType")) == field_char_type
        for field_char in run.findall(".//" + qn("w:fldChar"))
    )


def _is_page_decoration_run(run) -> bool:
    if run.find(".//" + qn("w:drawing")) is not None or run.find(".//" + qn("w:pict")) is not None:
        return False
    text = "".join(text_element.text or "" for text_element in run.findall(".//" + qn("w:t")))
    if not text:
        return False
    return not _PAGE_DECORATION_TEXT_RE.sub("", text).strip()


def _remove_literal_page_tokens(paragraph) -> None:
    for run in list(paragraph.runs):
        if _PAGE_INSTRUCTION_RE.search(run.text or ""):
            run._element.getparent().remove(run._element)
