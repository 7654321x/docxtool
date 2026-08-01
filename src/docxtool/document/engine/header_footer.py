"""渲染阶段页眉页脚兼容辅助。

本模块保留旧页脚 PAGE 域写入能力。当前主流程优先使用 page_number
模块，但旧导入路径仍可通过 core.py 转发到这里。
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docxtool.document.style_config import StyleRule, logger


def apply_header_footer(doc, page_rule: StyleRule) -> None:
    """设置奇偶页外侧页码。

    传入目标 Document 和页码 StyleRule；返回 None。函数会开启奇偶页
    不同，并在奇数页右侧、偶数页左侧写入 PAGE 域。
    """
    settings_element = doc.settings._element
    even_and_odd = settings_element.find(qn("w:evenAndOddHeaders"))
    if even_and_odd is None:
        even_and_odd = OxmlElement("w:evenAndOddHeaders")
        settings_element.append(even_and_odd)
    even_and_odd.set(qn("w:val"), "1")

    for section in doc.sections:
        footer_odd = section.footer
        footer_odd.is_linked_to_previous = False
        setup_page_number_paragraph(footer_odd, WD_ALIGN_PARAGRAPH.RIGHT, page_rule)

        try:
            footer_even = section.even_page_footer
            footer_even.is_linked_to_previous = False
            setup_page_number_paragraph(footer_even, WD_ALIGN_PARAGRAPH.LEFT, page_rule)
        except Exception as exc:
            logger.warning("[引擎] 偶数页页脚设置失败: %s（部分 Word 版本不支持）", exc)


def setup_page_number_paragraph(footer, alignment, page_rule: StyleRule) -> None:
    """在单个页脚中写入 PAGE 域段落。

    传入 footer、对齐方式和页码 StyleRule；返回 None。页脚会被清空后
    写入格式为“— PAGE —”的字段页码。
    """
    for paragraph in footer.paragraphs:
        paragraph.clear()

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = alignment

    pPr = para._element.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        ind.set(qn("w:rightChars"), "100")
    else:
        ind.set(qn("w:leftChars"), "100")
    pPr.append(ind)

    for text, is_field in [("— ", False), ("PAGE", True), (" —", False)]:
        run = para.add_run(text)
        run.font.name = page_rule.font
        run.font.size = Pt(page_rule.font_size_pt)
        run.font.bold = page_rule.bold

        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), page_rule.font)
        rFonts.set(qn("w:ascii"), page_rule.font)
        rFonts.set(qn("w:hAnsi"), page_rule.font)
        rPr.append(rFonts)

        if is_field:
            run.text = ""
            fldChar_begin = OxmlElement("w:fldChar")
            fldChar_begin.set(qn("w:fldCharType"), "begin")
            run._element.append(fldChar_begin)

            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = " PAGE "
            run._element.append(instrText)

            fldChar_sep = OxmlElement("w:fldChar")
            fldChar_sep.set(qn("w:fldCharType"), "separate")
            run._element.append(fldChar_sep)

            value = OxmlElement("w:t")
            value.text = "1"
            run._element.append(value)

            fldChar_end = OxmlElement("w:fldChar")
            fldChar_end.set(qn("w:fldCharType"), "end")
            run._element.append(fldChar_end)
