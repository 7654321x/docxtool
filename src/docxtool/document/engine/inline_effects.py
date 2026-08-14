"""Inline text effects used by the DOCX renderer.

本模块只负责渲染阶段的段内 run 效果：特殊加粗、冒号标签加粗、
责任单位行、名词解释、报告首句加粗和行内标题正文粗细分离。它只
处理已创建的输出段落，不决定段落类型，也不拆分逻辑段落。
"""

from __future__ import annotations

import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from docxtool.document.engine.inline import copy_run_style, segment_writer
from docxtool.document.engine.typography import set_run_fonts
from docxtool.document.recognition.colon import analyze_colon_structure, colon_bold_range
from docxtool.document.configuration.models import NB_FIXED, NB_SUFFIXES


INLINE_HEADING_BODY_MIN_CHARS = 5
_RESPONSIBILITY_LABEL_RE = re.compile(r"责\s*任\s*单\s*位\s*[:：]")


def apply_special_bold(paragraph, text: str) -> None:
    """应用特殊加粗；传入段落和全文，按规则重写 run，返回 None。"""
    if not paragraph.runs:
        return
    suffixes = "|".join(NB_SUFFIXES)
    fixed = "|".join(map(re.escape, NB_FIXED))
    lead_pattern = rf"(?:[一二三四五六七八九十]+(?:{suffixes})|{fixed})"
    parts = re.split(f"(?={lead_pattern})", text)
    write = _segment_writer(paragraph)
    for index, part in enumerate(parts):
        if not part:
            continue
        if index == 0:
            write(part, bold=False)
        elif NB_FIXED and any(part.startswith(fixed_text) for fixed_text in NB_FIXED):
            _write_fixed_bold(write, part)
        else:
            match = re.match(rf"([一二三四五六七八九十]+(?:{suffixes}).*?。)(.*)", part)
            if match:
                write(match.group(1), bold=True)
                write(match.group(2), bold=False)
            else:
                write(part, bold=True)


def apply_fixed_bold(paragraph, part: str) -> None:
    """固定词组加粗；传入段落和片段，冒号前加粗，返回 None。"""
    colon = _first_colon(part)
    if colon >= 0:
        bold_run = paragraph.add_run(part[:colon + 1])
        bold_run.font.bold = True
        if part[colon + 1:]:
            normal_run = paragraph.add_run(part[colon + 1:])
            normal_run.font.bold = False
    else:
        run = paragraph.add_run(part)
        run.font.bold = True


def apply_numbered_bold(paragraph, part: str) -> None:
    """序数引导句加粗；传入段落和片段，句号前加粗，返回 None。"""
    suffixes = "|".join(NB_SUFFIXES)
    lead_pattern = rf"[一二三四五六七八九十]+(?:{suffixes})"
    match = re.match(f"({lead_pattern}.*?。)(.*)", part)
    if match:
        bold_run = paragraph.add_run(match.group(1))
        bold_run.font.bold = True
        if match.group(2):
            normal_run = paragraph.add_run(match.group(2))
            normal_run.font.bold = False
    else:
        run = paragraph.add_run(part)
        run.font.bold = True


def apply_colon_bold(paragraph, text: str) -> None:
    """冒号标签加粗；传入段落和文本，短标签冒号前加粗，返回 None。"""
    if not paragraph.runs:
        return
    matched = colon_bold_range(text)
    if matched is not None:
        start, position = matched
        write = _segment_writer(paragraph)
        if start:
            write(text[:start], bold=False)
        write(text[start:position + 1], bold=True)
        write(text[position + 1:], bold=False)


def normalize_responsibility_lines(text: str) -> list[str]:
    """归一责任单位行；传入文本，返回拆分后的行列表。"""
    normalized = _RESPONSIBILITY_LABEL_RE.sub("责任单位：", text or "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    lines: list[str] = []
    for raw_line in normalized.split("\n"):
        line = raw_line.strip().strip("“”\"'")
        if not line:
            continue
        parts = [part.strip() for part in re.split(r"(?=责任单位：)", line) if part.strip()]
        lines.extend(parts or [line])
    return lines


def set_zero_first_line_indent(paragraph) -> None:
    """清零首行缩进；传入段落，直接修改 OOXML 并返回 None。"""
    properties = paragraph._element.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)
    indent.set(qn("w:firstLineChars"), "0")
    indent.set(qn("w:firstLine"), "0")
    for attr in ("w:hangingChars", "w:hanging"):
        indent.attrib.pop(qn(attr), None)


def force_responsibility_paragraph_format(paragraph) -> None:
    """设置责任单位段基础格式；传入段落，左对齐且清零首行缩进，返回 None。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_zero_first_line_indent(paragraph)


def apply_key_value_line_format(paragraph) -> None:
    """设置 A:B 键值行格式；传入段落，写入缩进/间距/字号，返回 None。"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    properties = paragraph._element.get_or_add_pPr()
    indent = properties.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        properties.append(indent)

    has_manual_break = paragraph._element.find(".//" + qn("w:br")) is not None
    if has_manual_break:
        indent.set(qn("w:leftChars"), "200")
        indent.set(qn("w:left"), "640")
        indent.set(qn("w:firstLineChars"), "0")
        indent.set(qn("w:firstLine"), "0")
    else:
        indent.set(qn("w:leftChars"), "0")
        indent.set(qn("w:left"), "0")
        indent.set(qn("w:firstLineChars"), "200")
        indent.set(qn("w:firstLine"), "640")
    for attr in ("w:hangingChars", "w:hanging"):
        indent.attrib.pop(qn(attr), None)

    spacing = properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        properties.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:beforeLines"), "0")
    spacing.set(qn("w:afterLines"), "0")
    spacing.set(qn("w:line"), "560")
    spacing.set(qn("w:lineRule"), "exact")

    for run in paragraph.runs:
        run.font.size = Pt(16)


def apply_responsibility_line(paragraph, text: str) -> None:
    """渲染责任单位行；传入段落和文本，按行加粗标签，返回 None。"""
    if not paragraph.runs:
        paragraph.add_run("")

    lines = normalize_responsibility_lines(text)
    if not lines:
        return

    base_run = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""
    used_base = False

    def write(text_part: str, *, bold: bool | None = None):
        """写入责任单位片段；传入文本和粗体设置，返回 run 或 None。"""
        nonlocal used_base
        if not text_part:
            return None
        run = base_run if not used_base else paragraph.add_run("")
        if used_base:
            copy_run_style(base_run, run)
        used_base = True
        run.text = text_part
        if bold is not None:
            run.font.bold = bold
        return run

    def add_break() -> None:
        """写入手动换行；不传入参数，返回 None。"""
        nonlocal used_base
        run = base_run if not used_base else paragraph.add_run("")
        if used_base:
            copy_run_style(base_run, run)
        used_base = True
        run.add_break()

    for index, line in enumerate(lines):
        if index:
            add_break()
        if line.startswith("责任单位："):
            write("责任单位：", bold=True)
            write(line[len("责任单位："):], bold=False)
        else:
            write(line, bold=False)

    force_responsibility_paragraph_format(paragraph)


def apply_glossary_item(paragraph, text: str, rule) -> None:
    """渲染名词解释条目；传入段落、文本和规则，返回 None。"""
    if len(paragraph.runs) < 2:
        return
    colon_position = _first_colon(text)
    if colon_position <= 0:
        return
    keyword = text[:colon_position + 1]
    body = text[colon_position + 1:]
    paragraph.runs[-1].text = ""

    keyword_run = paragraph.add_run(keyword)
    keyword_run.font.name = "黑体"
    set_run_fonts(keyword_run, cn_font="黑体", en_font="Times New Roman")
    keyword_run.font.size = Pt(rule.font_size_pt)
    keyword_run.font.bold = False

    if body:
        body_run = paragraph.add_run(body)
        body_run.font.name = "仿宋_GB2312"
        set_run_fonts(body_run, cn_font="仿宋_GB2312", en_font="Times New Roman")
        body_run.font.size = Pt(rule.font_size_pt)
        body_run.font.bold = False


def apply_report_first_sentence(paragraph, text: str, rule) -> None:
    """渲染报告首句强调；传入段落、文本和规则，返回 None。"""
    if not paragraph.runs:
        return
    period = text.find("。")
    if period <= 0:
        return
    first = text[:period + 1]
    rest = text[period + 1:]
    write = _segment_writer(paragraph)
    write(first, bold=True, cn_font="楷体_GB2312", size_pt=rule.font_size_pt)
    if rest:
        write(rest, bold=False, cn_font=rule.font, size_pt=rule.font_size_pt)


def apply_inline_lead_bold(paragraph, text: str, rule) -> None:
    """恢复正文首句加粗；传入段落、文本和规则，返回 None。"""
    if not paragraph.runs:
        return
    period = text.find("。")
    if period <= 0 or period >= len(text) - 1:
        return
    write = _segment_writer(paragraph)
    write(text[:period + 1], bold=True, cn_font=rule.font, size_pt=rule.font_size_pt)
    write(text[period + 1:], bold=rule.bold, cn_font=rule.font, size_pt=rule.font_size_pt)


def enforce_inline_heading2_format(paragraph, heading_bold: bool, body_bold: bool) -> None:
    """修正行内二级标题粗细；传入段落和标题/正文粗体值，返回 None。"""
    full_text = paragraph.text
    period_pos = full_text.find("。")
    if period_pos < 0 or len(full_text[period_pos + 1:].strip()) < INLINE_HEADING_BODY_MIN_CHARS:
        return
    inline_body_start = period_pos + 1

    cursor = 0
    for run in list(paragraph.runs):
        text = run.text or ""
        if not text:
            continue
        start = cursor
        end = cursor + len(text)
        cursor = end

        if end <= inline_body_start:
            run.font.bold = heading_bold
            continue
        if start >= inline_body_start:
            run.font.bold = body_bold
            continue

        heading_text = text[:inline_body_start - start]
        body_text = text[inline_body_start - start:]
        run.text = heading_text
        run.font.bold = heading_bold
        body_run = paragraph.add_run(body_text)
        copy_run_style(run, body_run)
        body_run.font.bold = body_bold
        run._element.addnext(body_run._element)


def apply_numbered_heading2_colon_format(
    paragraph,
    text: str,
    heading_rule,
    body_rule,
) -> None:
    """Format a numbered heading-2 label and its same-paragraph body."""
    analysis = analyze_colon_structure(text)
    if (
        not analysis.has_colon
        or analysis.separator_index is None
        or not analysis.label.strip()
        or not analysis.value.strip()
    ):
        return
    _apply_heading2_inline_body_format(
        paragraph,
        analysis.separator_index + 1,
        heading_rule,
        body_rule,
    )


def apply_numbered_heading2_period_format(
    paragraph,
    text: str,
    heading_rule,
    body_rule,
) -> None:
    """Format one physical heading-2 paragraph across its sentence boundary."""
    period_position = text.find("。")
    if (
        period_position < 0
        or len(text[period_position + 1:].strip()) < INLINE_HEADING_BODY_MIN_CHARS
    ):
        return
    _apply_heading2_inline_body_format(
        paragraph,
        period_position + 1,
        heading_rule,
        body_rule,
    )


def _apply_heading2_inline_body_format(
    paragraph,
    split_at: int,
    heading_rule,
    body_rule,
) -> None:
    """Apply heading/body rules in place without changing the paragraph boundary."""
    cursor = 0
    for run in list(paragraph.runs):
        run_text = run.text or ""
        if not run_text:
            continue
        start = cursor
        end = start + len(run_text)
        cursor = end
        if start < split_at < end:
            _split_text_run(run, split_at - start)
            break

    cursor = 0
    for run in paragraph.runs:
        run_text = run.text or ""
        if not run_text:
            continue
        end = cursor + len(run_text)
        rule = heading_rule if end <= split_at else body_rule
        set_run_fonts(run, cn_font=rule.font, en_font="Times New Roman")
        run.font.size = Pt(rule.font_size_pt)
        run.font.bold = rule.bold
        cursor = end


def _split_text_run(run, offset: int) -> None:
    """Split one text run while moving, rather than copying, trailing XML."""
    if offset <= 0 or offset >= len(run.text or ""):
        return
    text_tags = {
        qn("w:t"),
        qn("w:tab"),
        qn("w:br"),
        qn("w:cr"),
        qn("w:noBreakHyphen"),
        qn("w:ptab"),
    }
    consumed = 0
    children = list(run._element)
    for child_index, child in enumerate(children):
        if child.tag == qn("w:rPr"):
            continue
        value = str(child) if child.tag in text_tags else ""
        child_end = consumed + len(value)
        if offset == consumed:
            move_from = child_index
        elif offset == child_end:
            consumed = child_end
            continue
        elif consumed < offset < child_end and child.tag == qn("w:t"):
            local_offset = offset - consumed
            trailing_text = value[local_offset:]
            child.text = value[:local_offset]
            move_from = child_index + 1
        elif consumed < offset < child_end:
            raise ValueError("WPS_INLINE_HEADING2_COLON_RUN_SPLIT_FAILED")
        else:
            consumed = child_end
            continue

        new_run = run._parent.add_run("")
        copy_run_style(run, new_run)
        if consumed < offset < child_end:
            new_run.add_text(trailing_text)
        for trailing in children[move_from:]:
            new_run._element.append(trailing)
        run._element.addnext(new_run._element)
        return
    raise ValueError("WPS_INLINE_HEADING2_COLON_RUN_SPLIT_FAILED")


def handle_heading_period(text: str) -> str:
    """处理标题末尾句号；传入标题文本，返回保留或去句号后的文本。"""
    period_pos = text.find("。")
    if period_pos < 0:
        return text
    after_period = text[period_pos + 1:].strip()
    if len(after_period) >= INLINE_HEADING_BODY_MIN_CHARS:
        return text
    return text[:period_pos] + text[period_pos + 1:]


def _segment_writer(paragraph):
    """创建片段写入器；传入段落，返回 write 函数。"""
    return segment_writer(paragraph, set_run_fonts=set_run_fonts)


def _first_colon(text: str) -> int:
    """查找中英文冒号；传入文本，返回位置或 -1。"""
    positions = [position for position in (text.find("："), text.find(":")) if position >= 0]
    return min(positions) if positions else -1


def _write_fixed_bold(write, part: str) -> None:
    """用片段写入器渲染固定词；传入 write 函数和文本片段，返回 None。"""
    colon = _first_colon(part)
    if colon >= 0:
        write(part[:colon + 1], bold=True)
        write(part[colon + 1:], bold=False)
    else:
        write(part, bold=True)
