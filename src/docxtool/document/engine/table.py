"""Safe table layout helpers that only adjust formatting properties."""

from __future__ import annotations

import re
from collections.abc import Mapping
from typing import Any

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


_NUMBER_RE = re.compile(r"^[+-]?(?:\d+(?:,\d{3})*|\d*)(?:\.\d+)?%?$")


def format_tables(document, options: Mapping[str, Any] | None = None):
    """Format tables without rebuilding, merging, splitting, or deleting content."""

    opts = dict(options or {})
    if opts.get("enabled", True) is False:
        return document
    for table in _iter_tables(document):
        _format_table(table, opts)
    return document


def _iter_tables(container):
    for table in getattr(container, "tables", []):
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_tables(cell)


def _format_table(table, options: Mapping[str, Any]) -> None:
    table_properties = table._tbl.tblPr
    if table_properties is None:
        table_properties = OxmlElement("w:tblPr")
        table._tbl.insert(0, table_properties)
    if options.get("width_cm") is not None:
        _set_table_width(table_properties, _cm_to_twips(float(options["width_cm"])))
    elif options.get("width_pct") is not None:
        _set_table_width(table_properties, int(float(options["width_pct"]) * 50), width_type="pct")
    if options.get("indent_cm") is not None:
        _set_table_indent(table_properties, _cm_to_twips(float(options["indent_cm"])))
    if options.get("borders"):
        _set_borders(table_properties, options["borders"])
    margins = options.get("cell_margin_cm") or options.get("cell_margins_cm")
    if margins is not None:
        _set_cell_margins(table_properties, margins)
    if options.get("vertical_align"):
        _set_vertical_alignment(table, str(options["vertical_align"]))
    header_options = options.get("header") or {}
    if header_options:
        _format_header(table, header_options)
    if options.get("auto_align") or options.get("smart_alignment"):
        _auto_align_cell_text(table, int(options.get("long_text_threshold", 20)))


def _set_table_width(table_properties, width: int, *, width_type: str = "dxa") -> None:
    tbl_w = _first_or_new(table_properties, "w:tblW")
    tbl_w.set(qn("w:type"), width_type)
    tbl_w.set(qn("w:w"), str(width))


def _set_table_indent(table_properties, indent: int) -> None:
    tbl_ind = _first_or_new(table_properties, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))


def _set_borders(table_properties, borders: Any) -> None:
    config = borders if isinstance(borders, Mapping) else {}
    val = str(config.get("val", "single"))
    size = str(config.get("size", 4))
    color = str(config.get("color", "auto")).lstrip("#")
    space = str(config.get("space", 0))
    tbl_borders = _first_or_new(table_properties, "w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = _first_or_new(tbl_borders, f"w:{side}")
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)


def _set_cell_margins(table_properties, margins: Any) -> None:
    if isinstance(margins, Mapping):
        values = {side: _cm_to_twips(float(margins[side])) for side in ("top", "left", "bottom", "right") if side in margins}
    else:
        value = _cm_to_twips(float(margins))
        values = {side: value for side in ("top", "left", "bottom", "right")}
    tbl_cell_mar = _first_or_new(table_properties, "w:tblCellMar")
    for side, value in values.items():
        element = _first_or_new(tbl_cell_mar, f"w:{side}")
        element.set(qn("w:type"), "dxa")
        element.set(qn("w:w"), str(value))


def _set_vertical_alignment(table, value: str) -> None:
    alignment = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "middle": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }.get(value.lower(), WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = alignment


def _format_header(table, options: Mapping[str, Any]) -> None:
    if not table.rows:
        return
    row_count = int(options.get("rows", 1))
    for row in table.rows[:row_count]:
        if options.get("repeat", True):
            tr_pr = row._tr.get_or_add_trPr()
            header = _first_or_new(tr_pr, "w:tblHeader")
            header.set(qn("w:val"), "true")
        for cell in row.cells:
            if options.get("shading"):
                _set_cell_shading(cell, str(options["shading"]).lstrip("#"))
            if options.get("vertical_align"):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                if options.get("alignment"):
                    paragraph.alignment = _paragraph_alignment(str(options["alignment"]))
                for run in paragraph.runs:
                    if options.get("bold") is not None:
                        run.font.bold = bool(options["bold"])


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = _first_or_new(tc_pr, "w:shd")
    shading.set(qn("w:fill"), fill)


def _auto_align_cell_text(table, long_text_threshold: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            if _cell_has_sensitive_content(cell):
                continue
            text = " ".join(paragraph.text.strip() for paragraph in cell.paragraphs).strip()
            if not text:
                continue
            if _NUMBER_RE.match(text):
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif len(text) <= long_text_threshold:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment


def _cell_has_sensitive_content(cell) -> bool:
    return bool(
        cell._tc.findall(".//" + qn("w:drawing"))
        or cell._tc.findall(".//" + qn("w:pict"))
        or cell._tc.findall(".//" + qn("w:fldChar"))
        or cell._tc.findall(".//" + qn("w:instrText"))
        or cell._tc.findall(".//" + qn("w:hyperlink"))
    )


def _paragraph_alignment(value: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "middle": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(value.lower(), WD_ALIGN_PARAGRAPH.CENTER)


def _first_or_new(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _cm_to_twips(value: float) -> int:
    return int(round(value * 567))
