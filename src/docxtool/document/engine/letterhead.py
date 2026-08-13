"""Managed official-document letterhead rendering and compatibility facade."""

from __future__ import annotations

from dataclasses import dataclass, field
import unicodedata

from lxml import etree
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Cm, Pt, RGBColor

from docxtool.document.analysis.letterhead import (
    LETTERHEAD_STYLE_IDS,
    MANAGED_PROPERTY,
    MANAGED_VERSION,
    LetterheadDetection,
    _CUSTOM_FMTID,
    _CUSTOM_NS,
    _VT_NS,
    _has_red_separator_border,
    detect_letterhead,
    extract_letterhead_fields,
)
from docxtool.document.engine.style_catalog import ensure_letterhead_styles
from docxtool.document.letterhead_config import normalize_letterhead_config
from docxtool.document.style_config import PageSettings, StyleRule


_LETTERHEAD_STYLE_NAMES = {
    "DCT-LetterheadSpacer": "Docxtool Letterhead Spacer",
    "DCT-LetterheadMark": "Docxtool Letterhead Mark",
    "DCT-DocumentNumber": "Docxtool Document Number",
    "DCT-SignerLine": "Docxtool Signer Line",
    "DCT-LetterheadSeparator": "Docxtool Letterhead Separator",
}
WARNING_EXTERNAL = "LETTERHEAD_SKIPPED_EXISTING_EXTERNAL"
WARNING_UNKNOWN = "LETTERHEAD_SKIPPED_EXISTING_UNKNOWN"
_MARK_BASE_FONT_SIZE_PT = 32.0
_MARK_MIN_SCALE_PERCENT = 55
_MARK_TARGET_WIDTH_MAX_MM = 150.0
_LETTERHEAD_SPACER_PITCH_PT = 28.0
_VML_NS = "urn:schemas-microsoft-com:vml"


@dataclass
class LetterheadResult:
    action: str
    detection: str
    warnings: list[str] = field(default_factory=list)
    managed_paragraphs: int = 0
    protected_elements: list[object] = field(default_factory=list)


def _set_run_font(
    run,
    font_name: str,
    size_pt: float,
    *,
    color: str | None = None,
):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), font_name)
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rpr.rFonts.set(qn("w:cs"), "Times New Roman")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_run_spacing_and_scale(run, spacing_pt: float, scale_percent: int) -> None:
    rpr = run._element.get_or_add_rPr()
    spacing = rpr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        rpr.append(spacing)
    spacing.set(qn("w:val"), str(int(round(spacing_pt * 20))))
    width = rpr.find(qn("w:w"))
    if width is None:
        width = OxmlElement("w:w")
        rpr.append(width)
    width.set(qn("w:val"), str(int(scale_percent)))


def _set_paragraph_style(paragraph, style_id: str) -> None:
    paragraph.style = _LETTERHEAD_STYLE_NAMES[style_id]
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)


def _set_paragraph_line_spacing(
    paragraph,
    *,
    before_lines: int = 0,
    after_lines: int = 0,
) -> None:
    """Write paragraph spacing in OOXML line units (hundredths of a line)."""

    spacing = paragraph._p.get_or_add_pPr().get_or_add_spacing()
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:beforeLines"), str(before_lines * 100))
    spacing.set(qn("w:afterLines"), str(after_lines * 100))


def _document_number(config: dict) -> str:
    number = config["document_number"]
    return f"{number['agency_code']}〔{number['year']}〕{number['sequence']}号"


def _mark_lines(config: dict) -> list[tuple[str, bool]]:
    agencies = config["agencies"]
    if config["issuance_mode"] == "joint" and config["joint_mark_scope"] == "sponsor_only":
        agencies = agencies[:1]
    append_document = config["mark_display_mode"] == "agency_with_document"
    if len(agencies) == 1:
        name = agencies[0]["name"]
        return [(name if not append_document or name.endswith("文件") else name + "文件", False)]
    middle = (len(agencies) - 1) // 2
    return [
        (agency["name"], append_document and index == middle)
        for index, agency in enumerate(agencies)
    ]


def _estimate_mark_width_mm(
    text: str,
    font_size_pt: float,
    spacing_pt: float,
    scale_percent: int,
) -> float:
    glyph_units = sum(
        1.0 if unicodedata.east_asian_width(character) in {"W", "F"} else 0.55
        for character in text
    )
    glyph_width_pt = glyph_units * font_size_pt
    spacing_width_pt = max(0, len(text) - 1) * spacing_pt
    return (glyph_width_pt + spacing_width_pt) * scale_percent / 100 * 25.4 / 72


def _mark_typography(text: str, settings: PageSettings) -> tuple[float, float, int]:
    available_mm = (
        settings.page_width_cm - settings.margin_left_cm - settings.margin_right_cm
    ) * 10
    if available_mm <= 0:
        raise ValueError("LETTERHEAD_MARK_LAYOUT_INVALID")
    maximum_mm = min(_MARK_TARGET_WIDTH_MAX_MM, available_mm)
    font_size_pt = _MARK_BASE_FONT_SIZE_PT
    natural_mm = _estimate_mark_width_mm(text, font_size_pt, 0, 100)
    if natural_mm <= maximum_mm:
        return font_size_pt, 0.0, 100
    scale = int(maximum_mm / natural_mm * 100)
    if scale < _MARK_MIN_SCALE_PERCENT:
        raise ValueError("LETTERHEAD_MARK_TOO_LONG")
    return font_size_pt, 0.0, scale


def _add_mark_paragraphs(document, config: dict, settings: PageSettings) -> list:
    paragraphs = _add_spacer_paragraphs(
        document,
        3,
        pitch_pt=_LETTERHEAD_SPACER_PITCH_PT,
    )
    for text, add_document in _mark_lines(config):
        paragraph = document.add_paragraph()
        _set_paragraph_style(paragraph, "DCT-LetterheadMark")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        visible_text = text + ("文件" if add_document else "")
        font_size_pt, spacing_pt, scale_percent = _mark_typography(
            visible_text, settings
        )
        run = paragraph.add_run(text)
        _set_run_font(run, "方正小标宋简体", font_size_pt, color="FF0000")
        _set_run_spacing_and_scale(run, spacing_pt, scale_percent)
        if add_document:
            paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(13.5), WD_TAB_ALIGNMENT.RIGHT)
            label = paragraph.add_run("\t文件")
            _set_run_font(label, "方正小标宋简体", font_size_pt, color="FF0000")
            _set_run_spacing_and_scale(label, spacing_pt, scale_percent)
        _set_paragraph_line_spacing(paragraph)
        # The mark owns its line box: its exact line spacing follows the final
        # adaptive font size and never falls back to the 28 pt body pitch.
        paragraph.paragraph_format.line_spacing = Pt(font_size_pt)
        paragraphs.append(paragraph)
    paragraphs.extend(
        _add_spacer_paragraphs(
            document,
            2,
            pitch_pt=_effective_line_pitch(settings),
        )
    )
    return paragraphs


def _add_spacer_paragraphs(
    document,
    count: int,
    *,
    pitch_pt: float | None = None,
) -> list:
    paragraphs = []
    for _ in range(count):
        paragraph = document.add_paragraph()
        _set_paragraph_style(paragraph, "DCT-LetterheadSpacer")
        _set_paragraph_line_spacing(paragraph)
        if pitch_pt is not None:
            paragraph.paragraph_format.line_spacing = Pt(pitch_pt)
        paragraphs.append(paragraph)
    return paragraphs


def _add_number_paragraph(document, config: dict, settings: PageSettings):
    paragraph = document.add_paragraph()
    _set_paragraph_style(paragraph, "DCT-DocumentNumber")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(_document_number(config))
    _set_run_font(run, "仿宋_GB2312", 16)
    _set_paragraph_line_spacing(paragraph)
    return paragraph


def _append_signers(
    paragraph,
    signers: list[dict],
    settings: PageSettings,
    *,
    single_at_first_slot: bool = False,
) -> None:
    text_width_cm = (
        settings.page_width_cm
        - settings.margin_left_cm
        - settings.margin_right_cm
    )
    right_blank_one_char = Cm(text_width_cm) - Pt(16)
    first_signer_position = right_blank_one_char - Cm(4.6)
    if len(signers) == 2:
        tab_positions = (first_signer_position, right_blank_one_char)
    elif single_at_first_slot:
        tab_positions = (first_signer_position,)
    else:
        tab_positions = (right_blank_one_char,)
    for position in tab_positions:
        paragraph.paragraph_format.tab_stops.add_tab_stop(position, WD_TAB_ALIGNMENT.RIGHT)
    for signer in signers:
        paragraph.add_run("\t")
        label = paragraph.add_run(f"{signer['label']}：")
        _set_run_font(label, "仿宋_GB2312", 16)
        name = paragraph.add_run(signer["name"])
        _set_run_font(name, "楷体_GB2312", 16)


def _add_number_and_signer_paragraphs(
    document,
    config: dict,
    settings: PageSettings,
) -> list:
    if config["document_direction"] != "upward" or not config["signers"]:
        return [_add_number_paragraph(document, config, settings)]

    paragraphs = []
    signers = config["signers"]
    signer_rows = [signers[offset:offset + 2] for offset in range(0, len(signers), 2)]
    for index, row_signers in enumerate(signer_rows):
        final_row = index == len(signer_rows) - 1
        paragraph = document.add_paragraph()
        _set_paragraph_style(
            paragraph,
            "DCT-DocumentNumber" if final_row else "DCT-SignerLine",
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if final_row:
            paragraph.paragraph_format.left_indent = Pt(16)
            number = paragraph.add_run(_document_number(config))
            _set_run_font(number, "仿宋_GB2312", 16)
        _append_signers(
            paragraph,
            row_signers,
            settings,
            single_at_first_slot=len(signers) > 1 and len(row_signers) == 1,
        )
        _set_paragraph_line_spacing(paragraph)
        paragraphs.append(paragraph)
    return paragraphs


def _effective_line_pitch(settings: PageSettings) -> float:
    try:
        value = float(settings.line_spacing_value)
    except (TypeError, ValueError):
        value = 28.0
    return value if value > 0 else 28.0


def _add_separator(document, settings: PageSettings, style: str = "straight"):
    paragraph = document.add_paragraph()
    _set_paragraph_style(paragraph, "DCT-LetterheadSeparator")
    if style == "straight":
        ppr = paragraph._p.get_or_add_pPr()
        borders = ppr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            ppr.append(borders)
        bottom = borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "FF0000")
    elif style == "star":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_width_pt = min(
            Cm(15).pt,
            Cm(
                settings.page_width_cm
                - settings.margin_left_cm
                - settings.margin_right_cm
            ).pt,
        )
        run = paragraph.add_run()
        pict = OxmlElement("w:pict")
        group = etree.Element(f"{{{_VML_NS}}}group", nsmap={"v": _VML_NS})
        group.set("id", "DocxToolLetterheadStarSeparator")
        group.set("style", f"width:{text_width_pt:.2f}pt;height:9pt")
        group.set("coordorigin", "0,0")
        group.set("coordsize", "10000,200")
        for start, end in (("0,100", "4600,100"), ("5400,100", "10000,100")):
            line = etree.SubElement(group, f"{{{_VML_NS}}}line")
            line.set("from", start)
            line.set("to", end)
            line.set("strokecolor", "#FF0000")
            line.set("strokeweight", "1.5pt")
        # WPS renders the VML lines but can omit a filled ``v:polyline``.
        # A closed ``v:shape`` path is supported by both Word and WPS and
        # preserves the same centered five-point star geometry.
        star = etree.SubElement(group, f"{{{_VML_NS}}}shape")
        star.set("id", "DocxToolLetterheadCenterStar")
        star.set(
            "style",
            "position:absolute;left:0;top:0;width:10000;height:200",
        )
        star.set("coordorigin", "0,0")
        star.set("coordsize", "10000,200")
        star.set(
            "path",
            "m5000,10 l5022,69,5086,72,5036,112,5053,173,"
            "5000,138,4947,173,4964,112,4914,72,4978,69 x e",
        )
        star.set("filled", "t")
        star.set("stroked", "t")
        star.set("fillcolor", "#FF0000")
        star.set("strokecolor", "#FF0000")
        star.set("strokeweight", "0.5pt")
        pict.append(group)
        run._r.append(pict)
    else:
        raise ValueError("LETTERHEAD_SEPARATOR_STYLE_INVALID")
    paragraph.paragraph_format.space_before = Cm(0.4)
    _set_paragraph_line_spacing(paragraph)
    # The 4 mm distance is a physical standard, not a whole-line gap.
    spacing = paragraph._p.get_or_add_pPr().get_or_add_spacing()
    spacing.set(qn("w:before"), str(round(Cm(0.4).twips)))
    spacing.attrib.pop(qn("w:beforeLines"), None)
    spacing.set(qn("w:after"), str(int(round(_effective_line_pitch(settings) * 40))))
    spacing.attrib.pop(qn("w:afterLines"), None)
    paragraph.paragraph_format.line_spacing = Pt(1)
    return paragraph


def _first_title_element(document):
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.style_id == "DCT-Title":
            return paragraph._p
    body = document._body._element
    return next(
        (
            child
            for child in body.iterchildren()
            if child.tag != qn("w:sectPr")
            and not (
                child.tag == qn("w:p")
                and child.find("./" + qn("w:pPr") + "/" + qn("w:pStyle")) is not None
                and child.find("./" + qn("w:pPr") + "/" + qn("w:pStyle")).get(
                    qn("w:val"), ""
                )
                in LETTERHEAD_STYLE_IDS
            )
        ),
        None,
    )


def _clear_first_title_spacing(document) -> None:
    title = _first_title_element(document)
    if title is None or title.tag != qn("w:p"):
        return
    spacing = title.get_or_add_pPr().get_or_add_spacing()
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:beforeLines"), "0")


def _move_before(elements: list, anchor) -> None:
    if anchor is None:
        return
    parent = anchor.getparent()
    index = parent.index(anchor)
    for element in elements:
        parent.remove(element)
        parent.insert(index, element)
        index += 1


def _move_after_body_index(elements: list, document, body_index: int) -> bool:
    """Place newly created elements directly after a preserved header block."""

    body_children = [
        child
        for child in document._body._element.iterchildren()
        if child.tag != qn("w:sectPr")
    ]
    if not 0 <= body_index < len(body_children):
        return False
    anchor = body_children[body_index]
    parent = anchor.getparent()
    if parent is None:
        return False
    index = parent.index(anchor) + 1
    for element in elements:
        parent.remove(element)
        parent.insert(index, element)
        index += 1
    return True


def _detection_has_separator(document, detection: LetterheadDetection) -> bool:
    body_children = [
        child
        for child in document._body._element.iterchildren()
        if child.tag != qn("w:sectPr")
    ]
    return any(
        0 <= index < len(body_children)
        and body_children[index].tag == qn("w:p")
        and _has_red_separator_border(body_children[index])
        for index in detection.protected_body_indexes
    )


def _can_fill_missing_separator(document, detection: LetterheadDetection) -> bool:
    """Only complete a bounded document-number header, never an ambiguous block."""

    if detection.status != "unknown" or not detection.protected_body_indexes:
        return False
    detail = detection.details[0] if detection.details else ""
    if detail not in {
        "incomplete-document-number",
        "incomplete-document-number-signer",
        "compatible-document-number",
    }:
        return False
    return not _detection_has_separator(document, detection)


def _set_managed_property(document) -> None:
    package = document.part.package
    custom_part = next((part for part in package.parts if str(part.partname) == "/docProps/custom.xml"), None)
    if custom_part is None:
        root = etree.Element(f"{{{_CUSTOM_NS}}}Properties", nsmap={None: _CUSTOM_NS, "vt": _VT_NS})
        custom_part = Part(
            PackURI("/docProps/custom.xml"),
            CT.OFC_CUSTOM_PROPERTIES,
            etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True),
            package,
        )
        package.relate_to(custom_part, RT.CUSTOM_PROPERTIES)
    else:
        root = etree.fromstring(custom_part.blob)
    existing = None
    max_pid = 1
    for prop in root.findall(f"{{{_CUSTOM_NS}}}property"):
        max_pid = max(max_pid, int(prop.get("pid", "1") or 1))
        if prop.get("name") == MANAGED_PROPERTY:
            existing = prop
    if existing is None:
        existing = etree.SubElement(root, f"{{{_CUSTOM_NS}}}property")
        existing.set("fmtid", _CUSTOM_FMTID)
        existing.set("pid", str(max_pid + 1))
        existing.set("name", MANAGED_PROPERTY)
    for child in list(existing):
        existing.remove(child)
    value = etree.SubElement(existing, f"{{{_VT_NS}}}i4")
    value.text = MANAGED_VERSION
    custom_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def remove_managed_letterhead(document) -> int:
    removed = 0
    in_block = False
    for paragraph in list(document.paragraphs):
        style_id = paragraph.style.style_id if paragraph.style else ""
        if style_id in {"DCT-LetterheadSpacer", "DCT-LetterheadMark"}:
            in_block = True
        if in_block and style_id in LETTERHEAD_STYLE_IDS:
            paragraph._element.getparent().remove(paragraph._element)
            removed += 1
            if style_id == "DCT-LetterheadSeparator":
                break
        elif in_block:
            break
    return removed


def remove_detected_letterhead(document, detection: LetterheadDetection) -> int:
    """Remove the body-flow block identified by *detection* from this document."""

    body_children = [
        child
        for child in document._body._element.iterchildren()
        if child.tag != qn("w:sectPr")
    ]
    removed = 0
    for index in sorted(set(detection.protected_body_indexes), reverse=True):
        if 0 <= index < len(body_children):
            element = body_children[index]
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
                removed += 1
    return removed


def apply_letterhead(
    document,
    config,
    *,
    detection: LetterheadDetection | None = None,
    rules: list[StyleRule] | None = None,
    settings: PageSettings | None = None,
) -> LetterheadResult:
    normalized = normalize_letterhead_config(config)
    detection = detection or detect_letterhead(document)
    if not normalized["enabled"]:
        if _can_fill_missing_separator(document, detection):
            resolved_settings = settings or PageSettings()
            ensure_letterhead_styles(document, rules, resolved_settings)
            separator = _add_separator(document, resolved_settings)
            last_header_index = max(detection.protected_body_indexes)
            if _move_after_body_index([separator._p], document, last_header_index):
                return LetterheadResult(
                    "separator-added",
                    detection.status,
                    managed_paragraphs=1,
                    protected_elements=[separator._p],
                )
            # A failed placement must not leave an orphan separator at the end.
            separator._p.getparent().remove(separator._p)
            return LetterheadResult("ambiguous", detection.status)
        if detection.status in {"managed", "recognized_external"}:
            return LetterheadResult("preserved", detection.status)
        if detection.status == "none":
            return LetterheadResult("none", detection.status)
        return LetterheadResult("ambiguous", detection.status)
    resolved_settings = settings or PageSettings()
    if normalized.get("auto_recognize"):
        extracted = detection.fields or extract_letterhead_fields(document, detection)
        if extracted is None:
            raise ValueError("LETTERHEAD_AUTO_RECOGNITION_INCOMPLETE")
        if extracted.issuance_mode != "single":
            raise ValueError("LETTERHEAD_JOINT_SOURCE_UNSUPPORTED")
        normalized = normalize_letterhead_config(
            {
                **normalized,
                "agencies": [{
                    "id": "agency-1",
                    "name": extracted.mark_text,
                    "short_name": "",
                    "role": "sponsor",
                    "order": 1,
                }],
                "mark_display_mode": "agency_only",
                "document_number": {
                    "agency_code": extracted.agency_code,
                    "year": extracted.year,
                    "sequence": extracted.sequence,
                },
                "document_direction": "upward" if extracted.signers else "downward",
                "signers": [
                    {
                        "id": f"signer-{index}",
                        "agency_id": "agency-1",
                        "name": name,
                        "label": "签发人",
                        "order": index,
                    }
                    for index, name in enumerate(extracted.signers, 1)
                ],
                "separator_style": extracted.separator_style,
            }
        )
    # Calculate all deterministic layout values before deleting an existing
    # block. A validation or over-width failure therefore leaves it untouched.
    for text, add_document in _mark_lines(normalized):
        _mark_typography(text + ("文件" if add_document else ""), resolved_settings)
    if detection.status != "none":
        remove_detected_letterhead(document, detection)
    ensure_letterhead_styles(document, rules, resolved_settings)
    paragraphs = _add_mark_paragraphs(document, normalized, resolved_settings)
    paragraphs.extend(
        _add_number_and_signer_paragraphs(document, normalized, resolved_settings)
    )
    paragraphs.append(
        _add_separator(document, resolved_settings, normalized["separator_style"])
    )
    elements = [paragraph._p for paragraph in paragraphs]
    _move_before(elements, _first_title_element(document))
    _clear_first_title_spacing(document)
    _set_managed_property(document)
    return LetterheadResult(
        "replaced" if detection.status != "none" else "generated",
        detection.status,
        managed_paragraphs=len(paragraphs),
        protected_elements=elements,
    )


__all__ = [
    "LETTERHEAD_STYLE_IDS",
    "LetterheadDetection",
    "LetterheadResult",
    "WARNING_EXTERNAL",
    "WARNING_UNKNOWN",
    "apply_letterhead",
    "detect_letterhead",
    "remove_detected_letterhead",
    "remove_managed_letterhead",
]
