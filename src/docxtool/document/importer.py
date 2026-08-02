"""importer — .docx 导入 + 段落分类 + 结构识别。

职责边界：
  - 按 Word body XML 顺序提取段落/表格/图片段落
  - 标点规范化（引号/括号/英文标点→中文）
  - 兼容调用旧 scorer 优先级的段落分类
  - 兼容转发 Recognition 层旧 Flow 状态约束
  - 附件/落款/日期固定结构识别（正文→附件说明→落款→日期→附件页）
  - 编号分配 + 同级合并 + 连续性修复
  - 不负责排版渲染（由 engine.py 负责）
"""

from __future__ import annotations

import hashlib
import re
from types import SimpleNamespace
from typing import List, Optional, Tuple

from docxtool.document.classifier import ClassificationOptions, classify_paragraphs
from docxtool.document.models import (
    BodyBlock,
    DocumentData,
    InlineToken,
    NormalizationChange,
    ParagraphData,
    ParagraphFeatures,
    SegmentBoundaryCandidate,
    SourceRun,
)
from docxtool.document.normalization.tail import (
    normalize_tail_structures as _normalize_tail_structures,
    reorder_attachment_note_before_signature as _normalization_reorder_attachment_note_before_signature,
    sync_recognition_consistency as _sync_recognition_consistency,
)
from docxtool.document.normalization.changes import (
    record_applied_normalization_changes as _normalization_record_applied_normalization_changes,
    record_strict_normalization_suggestions as _normalization_record_strict_normalization_suggestions,
)
from docxtool.document.normalization.pipeline import (
    apply_post_recognition_normalization as _normalization_apply_post_recognition_normalization,
    capture_pre_normalization_snapshot as _normalization_capture_pre_normalization_snapshot,
    merge_uniform_heading_siblings as _normalization_merge_uniform_heading_siblings,
    strip_word_auto_numbering as _normalization_strip_word_auto_numbering,
)
from docxtool.document.normalization.dates import (
    chinese_number_to_int as _normalization_chinese_number_to_int,
    chinese_year_to_int as _normalization_chinese_year_to_int,
    is_attachment_page_mark as _normalization_is_attachment_page_mark,
    is_sign_date_text as _normalization_is_sign_date_text,
    normalize_attachment_page_mark as _normalization_normalize_attachment_page_mark,
    normalize_sign_date as _normalization_normalize_sign_date,
)
from docxtool.document.normalization.numbering import (
    assign_heading_numbering as _normalization_assign_heading_numbering,
    fix_heading_numbering_gaps as _normalization_fix_heading_numbering_gaps,
    strip_numbering_prefix as _normalization_strip_numbering_prefix,
    style_key_to_rule_row as _normalization_style_key_to_rule_row,
)
from docxtool.document.normalization.signature import normalize_sign_org as _normalization_normalize_sign_org
from docxtool.document.normalization.responsibility import (
    is_responsibility_line as _normalization_is_responsibility_line,
    normalize_responsibility_line as _normalization_normalize_responsibility_line,
)
from docxtool.document.normalization.text import (
    normalize_basic_text as _normalization_normalize_basic_text,
    normalize_quotes as _normalization_normalize_quotes,
    to_chinese_punctuation as _normalization_to_chinese_punctuation,
)
from docxtool.document.importing.images import (
    is_object_caption_text as _is_object_caption_text,
)
from docxtool.document.importing.features import (
    extract_paragraph_features as _importing_extract_paragraph_features,
)
from docxtool.document.importing.inline_tokens import (
    normalize_inline_tokens as _importing_normalize_inline_tokens,
)
from docxtool.document.importing.numbering import (
    NUMBERING_PATTERNS as _NUMBERING_PATTERNS,
    detect_numbering_prefix as _importing_detect_numbering_prefix,
    is_auto_numbered_item as _importing_is_auto_numbered_item,
)
from docxtool.document.importing.reader import (
    open_docx_document as _open_docx_document,
    read_body_blocks as _read_body_blocks,
)
from docxtool.document.importing.relationships import repair_broken_rels as _repair_broken_rels
from docxtool.document.recognition import apply_recognition
from docxtool.document.recognition.attachment import (
    can_start_attachment_note as _recognition_can_start_attachment_note,
    is_attachment_boundary_text as _recognition_is_attachment_boundary_text,
    match_attachment_item as _recognition_match_attachment_item,
    match_attachment_note as _recognition_match_attachment_note,
)
from docxtool.document.recognition.colon import (
    analyze_colon_structure,
    colon_bold_match as _recognition_colon_bold_match,
    contains_colon as _recognition_contains_colon,
    is_standalone_addressing_text,
)
from docxtool.document.recognition.document_mode import (
    detect_legacy_doc_type as _recognition_detect_legacy_doc_type,
    starts_report_heading_or_addressing as _recognition_starts_report_heading_or_addressing,
)
from docxtool.document.recognition.metadata import (
    enrich_legacy_type_metadata as _recognition_enrich_legacy_type_metadata,
)
from docxtool.document.recognition.opening_speech import (
    opening_speech_title_text as _recognition_opening_speech_title_text,
    strip_inferred_speech_numbering as _recognition_strip_inferred_speech_numbering,
)
from docxtool.document.recognition.selection import (
    build_legacy_scorer_registry as _recognition_build_legacy_scorer_registry,
    select_legacy_scored_type as _recognition_select_legacy_scored_type,
)
from docxtool.document.recognition.numbering import (
    find_numbered_bold_pos as _recognition_find_numbered_bold_pos,
    looks_like_damaged_heading as _recognition_looks_like_damaged_heading,
    match_numbering as _recognition_match_numbering,
    match_style_or_level as _recognition_match_style_or_level,
)
from docxtool.document.recognition.signature import (
    blocks_independent_sign_date as _recognition_blocks_independent_sign_date,
    has_signature_org_shape as _recognition_has_signature_org_shape,
    is_body_tail_context as _recognition_is_body_tail_context,
    is_signature_org_candidate as _recognition_is_signature_org_candidate,
)
from docxtool.document.recognition.state import (
    legacy_flow_allows as _recognition_legacy_flow_allows,
    legacy_record_structural as _recognition_legacy_record_structural,
    legacy_repair_heading2_continuation as _recognition_legacy_repair_heading2_continuation,
    legacy_repair_heading4_colon as _recognition_legacy_repair_heading4_colon,
    legacy_repair_heading_level as _recognition_legacy_repair_heading_level,
    legacy_repair_ocr_heading as _recognition_legacy_repair_ocr_heading,
    legacy_update_context_after_type as _recognition_legacy_update_context_after_type,
)
from docxtool.document.recognition.tail_structure import (
    detect_legacy_tail_structural_type as _recognition_detect_legacy_tail_structural_type,
)
from docxtool.document.recognition.version import RECOGNITION_VERSION_TAG
from docxtool.document.recognition.legacy import DetectionContext, ScoreBoard, ScoreDetail
from docxtool.document.segmentation.source_locator import (
    apply_segment_format_features as _apply_segment_format_features,
    inherit_source_locator as _inherit_source_locator,
    set_source_locator as _set_source_locator,
    source_line_spans as _source_line_spans,
    trim_source_span as _trim_source_span,
    visible_character_count as _visible_character_count,
)
from docxtool.document.segmentation.pipeline import (
    build_logical_lines as _build_logical_lines,
)
from docxtool.document.segmentation.body_tail import (
    find_last_body_candidate_index as _find_last_body_candidate_index,
)
from docxtool.document.segmentation.boundaries import (
    heading_has_inline_body as _seg_heading_has_inline_body,
    has_format_transition as _seg_has_format_transition,
    has_inline_lead_bold_transition as _seg_has_inline_lead_bold_transition,
    is_strong_soft_line_structure as _seg_is_strong_soft_line_structure,
    segment_boundary_candidates as _seg_segment_boundary_candidates,
    source_starts_body_region as _seg_source_starts_body_region,
    split_inline_heading_body_spans as _seg_split_inline_heading_body_spans,
    split_structural_tail_after_numbered_heading as _seg_split_structural_tail_after_numbered_heading,
    validate_numbered_heading_body_split as _seg_validate_numbered_heading_body_split,
    validate_source_span_partition as _seg_validate_source_span_partition,
)
from docxtool.document.segmentation.soft_breaks import (
    is_header_role_date_pair as _seg_is_header_role_date_pair,
    is_dispatch_number_line as _seg_is_dispatch_number_line,
    is_role_name_line as _seg_is_role_name_line,
    is_structural_key_value_line as _seg_is_structural_key_value_line,
    should_split_structural_line_breaks as _seg_should_split_structural_line_breaks,
)
from docxtool.document.style_config import (
    logger, ImportError,
    StyleRule,
)

# ═══════════════════════════════════════════════════════════════
# 数据模型
# ═══════════════════════════════════════════════════════════════

# Stable document models live in ``docxtool.document.models``.  They are
# imported above and re-exported from this module to keep legacy imports such
# as ``from docxtool.document.importer import ParagraphData`` compatible.
__all__ = [
    "BodyBlock",
    "DocumentData",
    "InlineToken",
    "NormalizationChange",
    "ParagraphData",
    "ParagraphFeatures",
    "SegmentBoundaryCandidate",
    "SourceRun",
    "ScoreBoard",
    "ScoreDetail",
    "_apply_segment_format_features",
    "_inherit_source_locator",
    "_set_source_locator",
    "_source_line_spans",
    "_trim_source_span",
    "_visible_character_count",
    "extract_features",
]


def _has_format_transition(
    features: Optional[ParagraphFeatures],
    start: int,
    boundary: int,
    end: int,
) -> bool:
    """兼容旧私有入口，传入段落特征和范围，返回是否存在格式切换。"""
    return _seg_has_format_transition(features, start, boundary, end)


def _segment_boundary_candidates(
    source: str,
    start: int,
    end: int,
    features: Optional[ParagraphFeatures] = None,
) -> Tuple[SegmentBoundaryCandidate, ...]:
    """兼容旧私有入口，传入源范围和特征，返回候选逻辑边界。"""
    return _seg_segment_boundary_candidates(
        source,
        start,
        end,
        features,
        analyze_colon_structure_func=analyze_colon_structure,
        detect_numbering_prefix_func=_detect_numbering_prefix,
    )


def _split_inline_heading_body_spans(
    source: str,
    start: int,
    end: int,
    features: Optional[ParagraphFeatures] = None,
    *,
    allow_visual_boundary: bool = True,
) -> List[Tuple[int, int]]:
    """兼容旧私有入口，传入源范围和特征，返回标题正文拆分范围。"""
    return _seg_split_inline_heading_body_spans(
        source,
        start,
        end,
        features,
        allow_visual_boundary=allow_visual_boundary,
        analyze_colon_structure_func=analyze_colon_structure,
        detect_numbering_prefix_func=_detect_numbering_prefix,
    )


def _is_standalone_addressing_text(text: str) -> bool:
    return is_standalone_addressing_text(text)


def _is_strong_soft_line_structure(text: str) -> bool:
    """兼容旧私有入口，传入软换行文本，返回是否为强结构行。"""
    return _seg_is_strong_soft_line_structure(
        text,
        detect_numbering_prefix_func=_detect_numbering_prefix,
        is_standalone_addressing_func=_is_standalone_addressing_text,
        is_sign_date_func=_normalization_is_sign_date_text,
        is_attachment_boundary_func=_is_attachment_boundary,
    )


def _split_structural_tail_after_numbered_heading(
    source: str,
    heading_body_spans: List[Tuple[int, int]],
    next_text: str = "",
) -> List[Tuple[int, int]]:
    """兼容旧私有入口，传入标题正文范围，返回释放尾部结构后的范围。"""
    return _seg_split_structural_tail_after_numbered_heading(
        source,
        heading_body_spans,
        next_text,
        is_strong_soft_line_structure_func=_is_strong_soft_line_structure,
        is_sign_date_func=_normalization_is_sign_date_text,
        is_tail_signature_org_func=_is_tail_signature_org_text,
    )


def _validate_source_span_partition(source: str, spans: List[Tuple[int, int]]) -> None:
    """兼容旧私有入口，传入源文本和范围列表，校验可见文字守恒。"""
    return _seg_validate_source_span_partition(source, spans)


def _source_starts_body_region(source: str) -> bool:
    """兼容旧私有入口，传入源文本，返回是否足以开启正文区域。"""
    return _seg_source_starts_body_region(
        source,
        detect_numbering_prefix_func=_detect_numbering_prefix,
    )


def _has_inline_lead_bold_transition(
    source: str,
    start: int,
    end: int,
    features: ParagraphFeatures,
) -> bool:
    """兼容旧私有入口，传入源范围和特征，返回是否存在正文首句加粗过渡。"""
    return _seg_has_inline_lead_bold_transition(
        source,
        start,
        end,
        features,
        detect_numbering_prefix_func=_detect_numbering_prefix,
    )


def _validate_numbered_heading_body_split(
    source: str,
    spans: List[Tuple[int, int]],
    features: Optional[ParagraphFeatures] = None,
) -> None:
    """兼容旧私有入口，传入源文本和范围列表，校验编号标题正文拆分契约。"""
    return _seg_validate_numbered_heading_body_split(
        source,
        spans,
        features,
        analyze_colon_structure_func=analyze_colon_structure,
        detect_numbering_prefix_func=_detect_numbering_prefix,
    )


# ═══════════════════════════════════════════════════════════════
# 特征提取
# ═══════════════════════════════════════════════════════════════

def extract_features(paragraph, index: int) -> ParagraphFeatures:
    """兼容旧入口，传入 python-docx 段落和序号，返回物理格式特征。"""
    return _importing_extract_paragraph_features(
        paragraph,
        index,
        detect_numbering_prefix_func=_detect_numbering_prefix,
    )


def _normalize_inline_tokens(tokens: List[InlineToken], punctuation_enabled: bool) -> List[InlineToken]:
    """兼容旧私有入口，传入 token 和标点开关，返回按 importer 标点策略处理后的 token。"""
    return _importing_normalize_inline_tokens(
        tokens,
        enabled=punctuation_enabled,
        normalize_text=lambda text: _to_chinese_punctuation(_normalize_quotes(text)),
    )


# ═══════════════════════════════════════════════════════════════
# 编号前缀检测（仅用于提取特征，不做业务判定）
# ═══════════════════════════════════════════════════════════════

def _detect_numbering_prefix(text: str) -> str:
    """检测文本中的编号前缀（仅特征提取，不判断类型）。"""
    return _importing_detect_numbering_prefix(text)


# ═══════════════════════════════════════════════════════════════
# V3 规则引擎 — 19 步优先级 + 独立 scorer + Flow 状态机
# ═══════════════════════════════════════════════════════════════

# ── 附件 / 落款 结构事实 ──
def _is_body_context(ctx) -> bool:
    """判断识别上下文是否已经处在正文或正文后结构中。"""
    return _recognition_is_body_tail_context(ctx.last_structural_type)

def _cn2int(s: str):
    """兼容旧私有入口，传入数字文本，返回整数或 None。"""
    return _normalization_chinese_number_to_int(s)

def _cn_year2int(s: str):
    """兼容旧私有入口，传入年份文本，返回整数年份或 None。"""
    return _normalization_chinese_year_to_int(s)

def _norm_sign_date(text: str) -> str:
    """兼容旧私有入口，传入成文日期文本，返回规范化日期。"""
    return _normalization_normalize_sign_date(text)

def _norm_attach_mark(text: str) -> str:
    """兼容旧私有入口，传入附件页标识，返回规范化标识。"""
    return _normalization_normalize_attachment_page_mark(text)

def _is_attachment_page_mark(text: str) -> bool:
    """传入段落文本，返回是否是附件页标识。"""
    return _normalization_is_attachment_page_mark(text)

def _is_attachment_boundary(text: str) -> bool:
    """传入段落文本，返回是否是附件说明或附件页边界。"""
    t = (text or "").strip()
    return _recognition_is_attachment_boundary_text(t, is_attachment_page_mark=_is_attachment_page_mark)

def _norm_sign_org(text: str) -> str:
    """兼容旧私有入口，传入落款单位文本，返回规范化单位名称。"""
    return _normalization_normalize_sign_org(text)


def _record_structural(ctx, type_id: str, text: str) -> None:
    """兼容旧私有入口，传入上下文、类型和文本，记录最后结构事实。"""
    _recognition_legacy_record_structural(ctx, type_id, text)

def _blocks_independent_sign_date(ctx) -> bool:
    """判断上一结构文本是否阻止当前日期作为独立尾部日期。"""
    return _recognition_blocks_independent_sign_date(ctx.last_structural_text)

def _looks_like_sign_org(text: str, next_text: str, ctx) -> bool:
    """综合上下文、下一段日期和机构形态判断落款单位候选。"""
    t = (text or "").strip()
    return _recognition_is_signature_org_candidate(
        t,
        next_text,
        last_structural_type=ctx.last_structural_type,
        is_attachment_note=bool(_recognition_match_attachment_note(t)),
        current_is_sign_date=_normalization_is_sign_date_text(t),
        next_is_sign_date=_normalization_is_sign_date_text(next_text or ""),
    )

def _heading_has_inline_body(text: str) -> bool:
    """兼容旧私有入口，传入标题候选文本，返回是否粘连正文。"""
    return _seg_heading_has_inline_body(text)

def _can_start_attachment_note(ctx) -> bool:
    """兼容旧私有入口，传入识别上下文，返回是否允许进入附件说明块。"""
    return _recognition_can_start_attachment_note(
        has_seen_real_body=ctx.has_seen_real_body,
        attachment_page_mode=ctx.attachment_page_mode,
        signature_complete=ctx.signature_complete,
        last_structural_type=ctx.last_structural_type,
    )

def _is_auto_numbered_item(feats: Optional[ParagraphFeatures]) -> bool:
    """兼容旧私有入口，传入段落特征，返回是否有 Word 自动编号事实。"""
    return _importing_is_auto_numbered_item(feats)


def _is_structural_key_value_line(text: str) -> bool:
    """兼容旧私有入口，传入一行文本，返回它是否可作为软换行键值边界。"""
    return _seg_is_structural_key_value_line(
        text,
        is_responsibility_line_func=_normalization_is_responsibility_line,
        colon_bold_match_func=_colon_bold_match,
    )


def _is_role_name_line(text: str) -> bool:
    """兼容旧私有入口，传入一行文本，返回是否具备文首职务姓名形态。"""
    return _seg_is_role_name_line(text)


def _is_header_role_date_pair(role_line: str, date_line: str) -> bool:
    """兼容旧私有入口，传入相邻两行文本，返回是否是职务姓名加日期。"""
    return _seg_is_header_role_date_pair(role_line, date_line)


def _opening_speech_title_text(text: str, ctx) -> str | None:
    """兼容旧私有入口，传入文本和上下文，返回文首讲话主标题文本。"""
    return _recognition_opening_speech_title_text(
        text,
        has_seen_body=ctx.has_seen_body,
        previous_type_id=ctx.prev_type_id,
        contains_colon_func=_contains_colon,
        match_numbering_func=_match_numbering,
    )


def _strip_inferred_speech_numbering(text: str) -> str:
    """兼容旧私有入口，传入文首讲话标题，返回去除误推断编号后的文本。"""
    return _recognition_strip_inferred_speech_numbering(
        text,
        match_numbering_func=_match_numbering,
    )


def _should_split_structural_line_breaks(parts: list[str], next_text: str) -> bool:
    """兼容旧私有入口，传入软换行行文本，返回是否需要拆成逻辑段。"""
    return _seg_should_split_structural_line_breaks(
        parts,
        next_text,
        detect_numbering_prefix_func=_detect_numbering_prefix,
        is_dispatch_number_line_func=_is_dispatch_number_line,
        is_key_value_line_func=_is_structural_key_value_line,
        is_sign_date_func=_normalization_is_sign_date_text,
        is_attachment_boundary_func=_is_attachment_boundary,
        is_tail_signature_org_func=_is_tail_signature_org_text,
        is_role_name_line_func=_is_role_name_line,
        is_header_role_date_pair_func=_is_header_role_date_pair,
    )


def _is_dispatch_number_line(text: str) -> bool:
    """兼容旧私有入口，传入一行文本，返回是否为结构化发文字号。"""
    return _seg_is_dispatch_number_line(text)


def _normalize_responsibility_line(text: str) -> str:
    """兼容旧私有入口，传入责任单位文本，返回规范化后的显示文本。"""
    return _normalization_normalize_responsibility_line(text)


def detect_structural_type(line: str, next_line: str, ctx,
                           feats: Optional[ParagraphFeatures] = None,
                           next_feats: Optional[ParagraphFeatures] = None):
    """兼容旧私有入口，传入当前/下一行和上下文，返回尾部结构类型结果。"""
    return _recognition_detect_legacy_tail_structural_type(
        line,
        next_line,
        ctx,
        feats,
        next_feats,
        is_responsibility_line_func=_normalization_is_responsibility_line,
        normalize_responsibility_line_func=_normalize_responsibility_line,
        match_attachment_note_func=_recognition_match_attachment_note,
        can_start_attachment_note_func=_can_start_attachment_note,
        match_attachment_item_func=_recognition_match_attachment_item,
        is_auto_numbered_item_func=_is_auto_numbered_item,
        looks_like_sign_org_func=_looks_like_sign_org,
        normalize_sign_org_func=_norm_sign_org,
        is_sign_date_func=_normalization_is_sign_date_text,
        normalize_sign_date_func=_norm_sign_date,
        is_attachment_boundary_func=_is_attachment_boundary,
        blocks_independent_sign_date_func=_blocks_independent_sign_date,
        is_attachment_page_mark_func=_is_attachment_page_mark,
        normalize_attachment_page_mark_func=_norm_attach_mark,
        contains_colon_func=_contains_colon,
        match_numbering_func=_match_numbering,
        record_structural_func=_record_structural,
    )


def _is_tail_signature_org_text(text: str) -> bool:
    """兼容旧私有入口，传入尾部短行文本，返回是否具备落款单位形态。"""
    value = (text or "").strip()
    return _recognition_has_signature_org_shape(value, max_length=40)


def _normalize_text(text: str) -> str:
    """兼容旧私有入口，传入文本并返回基础清理后的文本。"""
    return _normalization_normalize_basic_text(text)


def _to_chinese_punctuation(text: str) -> str:
    """兼容旧私有入口，传入文本并返回中文语境标点转换结果。"""
    return _normalization_to_chinese_punctuation(text)


def _normalize_quotes(text: str) -> str:
    """兼容旧私有入口，传入文本并返回中文语境引号转换结果。"""
    return _normalization_normalize_quotes(text)


def _feature_bool(value, default: bool = False) -> bool:
    """兼容旧私有入口，传入任意开关值，返回布尔配置结果。"""
    raw = str(value).strip().lower()
    if raw in {"1", "true", "yes", "on", "启用", "是"}:
        return True
    if raw in {"0", "false", "no", "off", "禁用", "否"}:
        return False
    return default


def _contains_colon(text: str) -> bool:
    """兼容旧私有入口，传入文本，返回是否包含中英文冒号。"""
    return _recognition_contains_colon(text)


def _colon_bold_match(text: str):
    """兼容旧私有入口，传入文本，返回应加粗标签的冒号位置或 -1。"""
    return _recognition_colon_bold_match(text)


def _find_numbered_bold_pos(text: str) -> int:
    """兼容旧私有入口，传入文本，返回“一是/一要”等强调位置或 -1。"""
    return _recognition_find_numbered_bold_pos(text, normalize_text=_normalize_text)


def _looks_like_heading(text: str) -> bool:
    """兼容旧私有入口，传入文本，返回是否像损坏的中文编号标题。"""
    return _recognition_looks_like_damaged_heading(text)


# ── 提取编号前缀（供 strip_numbering）──

def _match_numbering(text: str):
    """兼容旧私有入口，传入文本，返回 `(type_id, prefix)`。"""
    return _recognition_match_numbering(
        text,
        normalize_text=_normalize_text,
        contains_colon=_contains_colon,
    )


def _match_style_or_lvl(text: str, feats):
    """兼容旧私有入口，传入文本和特征，返回 Word 结构推导的标题类型。"""
    return _recognition_match_style_or_level(text, feats, normalize_text=_normalize_text)


def _detect_doc_type(ctx) -> str:
    """从头部标题文字检测文种。仅在 has_seen_body 首次变为 True 时调用一次。"""
    return _recognition_detect_legacy_doc_type(ctx.title_texts)


_STRUCTURE_SCORERS, _MODE_SCORERS, _FALLBACK_SCORERS = _recognition_build_legacy_scorer_registry(
    match_numbering_func=_match_numbering,
    contains_colon_func=_contains_colon,
    detect_doc_type_func=_detect_doc_type,
)


# ═══════════════════════════════════════════════════════════════
# Flow 层：显式状态机
# ═══════════════════════════════════════════════════════════════

def _flow_allows(candidate: str, ctx) -> bool:
    """候选类型是否被当前上下文允许。"""
    return _recognition_legacy_flow_allows(candidate, ctx.prev_type_id)


# ═══════════════════════════════════════════════════════════════
# Repair 层
# ═══════════════════════════════════════════════════════════════

def _repair_level(type_id: str, feats, ctx) -> str:
    """跳级修复：heading1→heading4 不允许，提级为 heading2。"""
    repaired = _recognition_legacy_repair_heading_level(type_id, ctx.current_level)
    if repaired != type_id:
        logger.debug(f"[修复] 跳级 {type_id}→{repaired}")
    return repaired


def _repair_heading4_colon(type_id: str, text: str, feats, ctx) -> str:
    """heading4 + 含冒号 → 回退为 body。"""
    repaired = _recognition_legacy_repair_heading4_colon(type_id, contains_colon=_contains_colon(text))
    if repaired != type_id:
        logger.debug("[修复] heading4 含冒号→body chars=%s", len(text))
    return repaired


# ═══════════════════════════════════════════════════════════════
# 主入口：classify（替代 detect_paragraph_type）
# ═══════════════════════════════════════════════════════════════

def detect_paragraph_type(text: str, feats: ParagraphFeatures,
                          ctx: DetectionContext,
                          rules: List[StyleRule]) -> Tuple[str, dict, str]:
    """v3 统一分类器：按 19 步优先级遍历 scorer → Flow 约束 → Repair。

    返回 (type_id, meta_patch, prefix)，与原接口完全兼容。
    """
    ctx.para_index = feats.paragraph_index
    meta: dict = {}
    prefix: str = ""
    from_word_structure = False
    score_log = []  # 收集各 scorer 得分
    unbound_object_label = _is_object_caption_text(text)

    # 未紧邻对象的“表1/图2...”不再具备题注语义，应作为正文清理
    # 直接格式；真正题注已在 raw_blocks 阶段转换为受保护对象。
    if unbound_object_label:
        type_id = "body"
        score_log.append("unbound_object_label:100")
    elif opening_speech_title := _opening_speech_title_text(text, ctx):
        # Word's built-in Heading 1 is often pasted onto the first line of a
        # speech manuscript.  A first-line “在……上的讲话” is a main title,
        # not the first numbered section.
        type_id = "title"
        meta["is_title"] = True
        if opening_speech_title != text.strip():
            meta["strip_inferred_speech_numbering"] = True
        score_log.append("opening_speech_title:100")
    else:
        # ── 先检查 Word 样式/多级列表 ──
        style_tid, style_prefix = _match_style_or_lvl(text, feats)
        if style_tid:
            type_id = style_tid
            prefix = style_prefix
            from_word_structure = True
        else:
            type_id, meta, prefix, score_log = _recognition_select_legacy_scored_type(
                text,
                feats,
                ctx,
                structure_scorers=_STRUCTURE_SCORERS,
                mode_scorers=_MODE_SCORERS,
                fallback_scorers=_FALLBACK_SCORERS,
                detect_doc_type_func=_detect_doc_type,
                flow_allows_func=_flow_allows,
            )

    # ── Repair ──
    type_id = _repair_heading4_colon(type_id, text, feats, ctx)
    if not from_word_structure:
        type_id = _repair_level(type_id, feats, ctx)

    repaired_type_id = _recognition_legacy_repair_ocr_heading(
        type_id,
        text,
        has_seen_body=ctx.has_seen_body,
        unbound_object_label=unbound_object_label,
        looks_like_heading_func=_looks_like_heading,
    )
    if repaired_type_id != type_id:
        type_id = repaired_type_id
        logger.debug("[修复] OCR 标题升级 chars=%s", len(text))

    # ── 打分日志 ──
    scores_str = ' → '.join(score_log) if score_log else 'by_style'
    logger.info("[打分] chars=%s text_sha256=%s | %s → %s", len(text), hashlib.sha256(text.encode("utf-8")).hexdigest()[:12], scores_str, type_id)

    repaired_type_id = _recognition_legacy_repair_heading2_continuation(
        type_id,
        text,
        ctx.prev_type_id,
        meta,
    )
    if repaired_type_id != type_id:
        type_id = repaired_type_id
        logger.debug("[修复] heading2 续行 chars=%s", len(text))

    meta = _recognition_enrich_legacy_type_metadata(
        text,
        type_id,
        feats,
        ctx,
        meta,
        heading_has_inline_body_func=_heading_has_inline_body,
        find_numbered_bold_pos_func=_find_numbered_bold_pos,
        colon_bold_match_func=_colon_bold_match,
        starts_report_heading_or_addressing_func=_recognition_starts_report_heading_or_addressing,
    )

    _recognition_legacy_update_context_after_type(
        ctx,
        type_id,
        text,
        meta,
        detect_doc_type_func=_detect_doc_type,
    )

    logger.debug(f"[决策] para={ctx.para_index} → {type_id} meta={meta}")
    return type_id, meta, prefix


# ═══════════════════════════════════════════════════════════════
# 编号剥离
# ═══════════════════════════════════════════════════════════════

def strip_numbering(text: str, prefix: Optional[str] = None) -> str:
    """兼容旧入口，传入标题文本和可选前缀，返回剥离编号后的正文。"""
    return _normalization_strip_numbering_prefix(
        text,
        prefix,
        numbering_patterns=_NUMBERING_PATTERNS,
    )


def _key_to_row(key: str) -> int:
    """兼容旧私有入口，传入编号层级 key，返回样式规则行号。"""
    return _normalization_style_key_to_rule_row(key)


# ═══════════════════════════════════════════════════════════════
# 导入器
# ═══════════════════════════════════════════════════════════════

class DocxImporter:
    """.docx 文件导入器。"""

    def load(
        self,
        filepath: str,
        rules: List[StyleRule],
        features: dict = None,
        *,
        strict_preservation: bool = True,
        recognition_mode: str = "authoritative",
    ) -> DocumentData:
        """加载 .docx，识别段落类型，返回 DocumentData。"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        features = features or {}
        processing_options = features.get("processing", {}) if isinstance(features.get("processing", {}), dict) else {}
        requested_strategy = str(
            processing_options.get("strategy")
            or processing_options.get("mode")
            or ""
        ).strip().lower()
        requested_strategy = {"smart": "structural"}.get(requested_strategy, requested_strategy)
        if requested_strategy:
            if requested_strategy not in {"strict", "structural", "normalize"}:
                raise ImportError("处理模式必须为 strict、smart、structural 或 normalize")
            processing_strategy = requested_strategy
        elif "strict_preservation" in processing_options:
            processing_strategy = "strict" if _feature_bool(
                processing_options.get("strict_preservation"), strict_preservation
            ) else "normalize"
        else:
            # Preserve the public importer default for library callers.  The
            # web configuration explicitly selects smart/structural mode.
            processing_strategy = "strict" if strict_preservation else "normalize"
        strict_preservation = processing_strategy == "strict"
        structural_preservation = processing_strategy == "structural"
        # A numbered heading followed by prose is rendered as title/body
        # paragraphs. The splitter only permits its first heading boundary;
        # later font changes inside the body cannot create extra paragraphs.
        split_inline_heading_body = structural_preservation and _feature_bool(
            processing_options.get("split_inline_heading_body", True), True
        )
        recognition_options = features.get("recognition", {}) if isinstance(features.get("recognition", {}), dict) else {}
        recognition_mode = str(recognition_options.get("mode", recognition_mode) or recognition_mode).lower()
        if recognition_mode not in {"legacy", "shadow", "authoritative"}:
            raise ImportError("识别模式必须为 legacy、shadow 或 authoritative")
        punctuation_options = features.get("punctuation", {}) if isinstance(features.get("punctuation", {}), dict) else {}
        numbering_options = features.get("numbering", {}) if isinstance(features.get("numbering", {}), dict) else {}
        numbering_enabled = _feature_bool(numbering_options.get("enabled", False), False)
        new_punctuation_enabled = _feature_bool(punctuation_options.get("enabled", False), False)
        punctuation_mode = str(punctuation_options.get("mode", "safe") or "safe")
        punctuation_enabled = _feature_bool(features.get("punctuation_enabled", True), True)
        punctuation_requested = new_punctuation_enabled or punctuation_enabled

        def normalize_text(text: str) -> str:
            if not text:
                return text
            if strict_preservation:
                return text
            if structural_preservation and not punctuation_requested:
                return text
            if new_punctuation_enabled:
                from docxtool.document.engine.punctuation import normalize_punctuation_text

                return normalize_punctuation_text(_normalize_text(text), mode=punctuation_mode)
            if punctuation_enabled:
                return _to_chinese_punctuation(_normalize_quotes(_normalize_text(text)))
            return text

        def normalize_tokens(tokens: List[InlineToken]) -> List[InlineToken]:
            if strict_preservation:
                return list(tokens or [])
            if structural_preservation and not punctuation_requested:
                return list(tokens or [])
            normalized = _normalize_inline_tokens(tokens, punctuation_enabled and not new_punctuation_enabled)
            if not new_punctuation_enabled:
                if not punctuation_enabled:
                    return normalized
                return [
                    InlineToken(token.kind, _normalize_text(token.text)) if token.kind == "text" else token
                    for token in normalized
                ]
            return [
                InlineToken(token.kind, normalize_text(token.text)) if token.kind == "text" else token
                for token in normalized
            ]

        # 物理文件打开与关系修复副本清理由 importing.reader 负责；传入
        # 旧回调，保留 importer 私有入口可被既有调用方和回归测试替换。
        original_filepath = filepath
        doc = _open_docx_document(
            filepath,
            document_factory=DocxDocument,
            repair_broken_rels_func=_repair_broken_rels,
            import_error_type=ImportError,
            cleanup_warning=logger.warning,
        )

        data = DocumentData(
            filepath=original_filepath,
            strict_preservation=strict_preservation,
            processing_strategy=processing_strategy,
            recognition_mode=recognition_mode,
        )
        source_visible_texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        from docxtool.document.engine.letterhead import detect_letterhead

        data.letterhead_detection = detect_letterhead(doc)
        protected_letterhead_indexes = set(data.letterhead_detection.protected_body_indexes)

        # 第一步：只提取物理 body 事实；Reader 不处理逻辑边界或类型。
        raw_blocks = _read_body_blocks(
            doc,
            data,
            strict_preservation=strict_preservation,
            protected_letterhead_indexes=protected_letterhead_indexes,
            extract_features_func=extract_features,
        )

        # 第二步：只按原始 source span 形成逻辑行。结构条件与顺序仍由
        # 原有 segmentation 规则决定，Importer 随后才开始旧链路和 Recognition。
        flat_lines = _build_logical_lines(
            raw_blocks,
            strict_preservation=strict_preservation,
            structural_preservation=structural_preservation,
            split_inline_heading_body_enabled=split_inline_heading_body,
            normalize_text_func=normalize_text,
            source_starts_body_region_func=_source_starts_body_region,
            split_inline_heading_body_spans_func=_split_inline_heading_body_spans,
            validate_numbered_heading_body_split_func=_validate_numbered_heading_body_split,
            should_split_structural_line_breaks_func=_should_split_structural_line_breaks,
            split_structural_tail_after_numbered_heading_func=(
                _split_structural_tail_after_numbered_heading
            ),
            validate_source_span_partition_func=_validate_source_span_partition,
            detect_numbering_prefix_func=_detect_numbering_prefix,
            inline_lead_bold_func=_has_inline_lead_bold_transition,
        )

        ctx = DetectionContext()
        # 预扫描只生成尾部边界事实，后续最终类型仍由识别链路决定。
        last_body_idx = _find_last_body_candidate_index(
            [item[1] if item[0] == "text" else "" for item in flat_lines],
            is_attachment_start_func=lambda text: bool(re.match(r'^附件', text)),
            is_sign_date_func=_normalization_is_sign_date_text,
            is_attachment_item_func=lambda text: bool(re.match(r'^\d+[.．、]', text)),
            is_attachment_page_mark_func=_is_attachment_page_mark,
        )
        for i, item in enumerate(flat_lines):
            ctx._remaining_has_no_body = (i >= last_body_idx)
            if item[0] == "table":
                pd = ParagraphData(text="", type_id="__table__",
                                   original_text="", features=None,
                                   meta={"table": item[1]})
                data.paragraphs.append(pd)
                continue
            if item[0] == "paragraph_xml":
                pd = ParagraphData(text="", type_id="__image__",
                                   original_text="", features=None,
                                   meta={"image_xml": item[1]})
                data.paragraphs.append(pd)
                continue
            if item[0] == "protected_paragraph_xml":
                caption_text = item[1].text
                pd = ParagraphData(text=caption_text, type_id="__object_caption__",
                                   original_text=caption_text, features=item[2],
                                   meta={"paragraph_xml": item[1]})
                data.paragraphs.append(pd)
                continue
            if item[0] == "letterhead_paragraph_xml":
                pd = ParagraphData(
                    text="",
                    type_id="__letterhead__",
                    original_text="",
                    features=None,
                    meta={"paragraph_xml": item[1]},
                )
                data.paragraphs.append(pd)
                continue

            _, line, sub_pf, inline_tokens, sectPr = item
            next_line = ""
            next_pf = None
            for next_item in flat_lines[i + 1:]:
                if next_item[0] == "text":
                    next_line = next_item[1]
                    next_pf = next_item[2]
                    break

            # 受管版头输出重新处理时，固定主标题样式优先于普通物理特征打分。
            managed_title = (
                data.letterhead_detection.status == "managed"
                and sub_pf.style_name == "Docxtool Title"
                and not ctx.has_seen_real_body
            )
            if managed_title:
                type_id = "title" if not ctx.title_texts else "title_cont"
                meta_patch = {"is_title": True} if type_id == "title" else {}
                prefix = ""
                clean_text = line
            else:
                # 结构检测优先
                st, sm, sp, ft = detect_structural_type(line, next_line, ctx, sub_pf, next_pf)
                if st:
                    sm.pop("numbering", None)
                    type_id = st
                    meta_patch = sm
                    prefix = sp
                    clean_text = ft
                    ctx.prev_type_id = st
                else:
                    type_id, meta_patch, prefix = detect_paragraph_type(line, sub_pf, ctx, rules)
                    clean_text = strip_numbering(line, prefix)

            if strict_preservation or structural_preservation:
                clean_text = line
                meta_patch = dict(meta_patch or {})
                meta_patch.pop("numbering", None)
                if meta_patch.get("strip_inferred_speech_numbering"):
                    clean_text = _strip_inferred_speech_numbering(line)

            if ctx.attachment_page_mode and type_id == "body":
                type_id = "attachment_body"

            # 跳级修正
            if type_id.startswith("heading") and not type_id == "heading1_report":
                lvl = int(type_id[-1])
                prev_lvl = ctx.current_level
                if lvl == getattr(ctx, '_last_detected_lvl', 0):
                    capped = prev_lvl
                else:
                    capped = min(lvl, prev_lvl + 1)
                if capped != lvl:
                    type_id = f"heading{capped}"
                ctx.current_level = capped
                ctx._last_detected_lvl = lvl
            elif type_id == "heading1_report":
                ctx.current_level = 1
            ctx.prev_type_id = type_id

            # 结构状态跟踪
            if type_id in ("body", "addressing", "responsibility_line"):
                ctx.has_seen_real_body = True
                _record_structural(ctx, "body", clean_text)
            elif type_id.startswith("heading") or type_id in ("title", "title2"):
                if meta_patch.get("heading_inline_body"):
                    ctx.has_seen_real_body = True
                    _record_structural(ctx, "body", clean_text)
                else:
                    _record_structural(ctx, type_id, clean_text)
            else:
                _record_structural(ctx, type_id, clean_text)
            if type_id == "sign_date":
                ctx.signature_complete = True
                ctx.attachment_page_mode = False

            if sectPr is not None:
                meta_patch = dict(meta_patch or {})
                meta_patch["sectPr"] = sectPr
            meta_patch = dict(meta_patch or {})
            meta_patch["legacy_type_id"] = {
                "value": type_id,
                "source": "legacy_importer",
                "recognition_version": RECOGNITION_VERSION_TAG,
            }
            pd = ParagraphData(
                text=clean_text, type_id=type_id,
                original_text=line, features=sub_pf, meta=meta_patch,
                inline_tokens=inline_tokens if strict_preservation or clean_text == line else [],
            )
            data.paragraphs.append(pd)
            text_hash = hashlib.sha256(clean_text.encode("utf-8")).hexdigest()[:12]
            logger.info(
                "[识别] #%s type=%s chars=%s text_sha256=%s",
                len(data.paragraphs) - 1,
                type_id,
                len(clean_text),
                text_hash,
            )
            # (body_blocks removed — tables/images now use paragraph stream placeholders)

        data.doc_mode = ctx.doc_mode
        before_normalization = _normalization_capture_pre_normalization_snapshot(
            data,
            source_visible_texts,
        )
        if strict_preservation:
            self._record_strict_normalization_suggestions(data)
        self._apply_core_classification(data, features)
        from dataclasses import replace
        from docxtool.document.recognition.config import DEFAULT_CONFIG

        apply_recognition(data, replace(DEFAULT_CONFIG, mode=recognition_mode))
        _normalization_apply_post_recognition_normalization(
            data,
            rules,
            doc.paragraphs,
            strict_preservation=strict_preservation,
            structural_preservation=structural_preservation,
            processing_strategy=processing_strategy,
            numbering_enabled=numbering_enabled,
            before_normalization=before_normalization,
            normalize_tail_structures_func=_normalize_tail_structures,
            reorder_attachment_note_before_signature_func=(
                self._reorder_attachment_note_before_signature
            ),
            assign_numbering_func=self._assign_numbering,
            merge_siblings_func=self._merge_siblings,
            record_applied_normalization_changes_func=(
                self._record_applied_normalization_changes
            ),
            fix_numbering_gaps_func=self._fix_numbering_gaps,
            strip_auto_numbering_func=self._strip_auto_numbering,
            sync_recognition_consistency_func=_sync_recognition_consistency,
        )
        # (old classification loop removed — replaced by flat_lines single pass above)

        logger.info(
            "[导入] file_sha256=%s paragraphs=%s tables=%s strategy=%s recognition=%s",
            hashlib.sha256(str(original_filepath).encode("utf-8")).hexdigest()[:12],
            len(data.paragraphs),
            len(data.tables),
            processing_strategy,
            recognition_mode,
        )
        return data

    def _record_strict_normalization_suggestions(self, data: DocumentData) -> None:
        """兼容旧私有入口，传入文档数据，记录 strict 模式规范化建议。"""
        from docxtool.document.engine.punctuation import normalize_punctuation_text

        _normalization_record_strict_normalization_suggestions(
            data,
            normalize_punctuation_func=lambda text: normalize_punctuation_text(text, mode="safe"),
        )

    def _record_applied_normalization_changes(
        self,
        data: DocumentData,
        before: list[tuple[str, str, str]],
    ) -> None:
        """兼容旧私有入口，传入文档数据和旧快照，记录已应用变化。"""
        _normalization_record_applied_normalization_changes(data, before)

    def _apply_core_classification(self, data: DocumentData, features: dict) -> None:
        classification_options = features.get("classification", {}) if isinstance(features.get("classification", {}), dict) else {}
        if not _feature_bool(classification_options.get("enabled", True), True):
            return
        threshold = classification_options.get("minimum_auto_format_confidence", 0.85)
        try:
            threshold = float(threshold)
        except (TypeError, ValueError):
            threshold = 0.85
        candidates = []
        indexes = []
        for index, paragraph in enumerate(data.paragraphs):
            if paragraph.type_id.startswith("__"):
                continue
            pf = paragraph.features or ParagraphFeatures()
            candidates.append(
                SimpleNamespace(
                    text=paragraph.original_text or paragraph.text,
                    style_name=pf.style_name,
                    alignment=pf.alignment,
                    first_line_indent=pf.first_line_indent,
                    font_size_pt=pf.font_size_pt,
                    bold=pf.bold,
                    native_numbering=bool(pf.numbering_prefix),
                )
            )
            indexes.append(index)
        if not candidates:
            return
        results = classify_paragraphs(candidates, ClassificationOptions(auto_format_threshold=threshold))
        for paragraph_index, result in zip(indexes, results):
            meta = dict(data.paragraphs[paragraph_index].meta or {})
            meta["classification_kind"] = result.kind.value
            meta["classification_confidence"] = round(result.confidence, 3)
            meta["classification_auto_format"] = bool(result.auto_format)
            data.paragraphs[paragraph_index].meta = meta

    def _reorder_attachment_note_before_signature(self, paragraphs: list) -> None:
        """兼容旧私有入口，传入段落列表，按正式尾部顺序原地重排。"""
        _normalization_reorder_attachment_note_before_signature(paragraphs)

    def _assign_numbering(self, paragraphs: list, rules: list, reset_on_attach: bool = True) -> None:
        """兼容旧私有入口，传入段落和样式规则，写入标题编号 meta。"""
        _normalization_assign_heading_numbering(
            paragraphs,
            rules,
            reset_on_attach=reset_on_attach,
            log_warning=logger.warning,
            log_debug=logger.debug,
        )

    def _strip_auto_numbering(self, paragraph) -> None:
        """删除段落中的 Word 自动编号标记 <w:numPr>。"""
        _normalization_strip_word_auto_numbering(
            paragraph,
            log_debug=logger.debug,
        )

    def _merge_siblings(self, paragraphs: list) -> None:
        """A. 同级合并：父标题下全是同模式子项 → 提升为父+1级。"""
        _normalization_merge_uniform_heading_siblings(
            paragraphs,
            log_debug=logger.debug,
            log_info=logger.info,
        )

    def _fix_numbering_gaps(self, paragraphs: list) -> None:
        """兼容旧私有入口，传入段落列表，修复标题编号连续性。"""
        _normalization_fix_heading_numbering_gaps(paragraphs, log_warning=logger.warning)


if __name__ == "__main__":
    ctx = DetectionContext()
    ctx.doc_mode = "NORMAL"
    pf = ParagraphFeatures()
    tid, _, _ = detect_paragraph_type("一、加强政治建设", pf, ctx, [])
    assert tid == "heading1", f"一、→ heading1: got {tid}"
    tid, _, _ = detect_paragraph_type("（一）坚持党的领导", pf, ctx, [])
    assert tid == "heading2", f"（一）→ heading2: got {tid}"
    tid, _, _ = detect_paragraph_type("1.完善组织体系", pf, ctx, [])
    assert tid == "heading3", f"1.→ heading3: got {tid}"
    pf_lvl = ParagraphFeatures(numbering_prefix="@lvl_0")
    tid, _, _ = detect_paragraph_type("带头固本培元、增强党性方面。", pf_lvl, DetectionContext(has_seen_body=True), [])
    assert tid == "heading2", f"@lvl_0 应识别为 heading2: got {tid}"
    tid, meta, _ = detect_paragraph_type("一要党内组织生活实效还有待提升。", pf_lvl, DetectionContext(has_seen_body=True), [])
    assert tid == "body" and meta.get("numbered_bold"), f"一要正文不应被自动编号误判为标题: {tid}, {meta}"
    assert strip_numbering("一、改革", "一、") == "改革"
    assert _looks_like_heading("一，加强领导")

    # 附件结构识别回归：正文后进入附件说明 → 落款 → 附件页 → 标题 → 正文
    ctx2 = DetectionContext(
        has_seen_body=True,
        has_seen_real_body=True,
        prev_type_id="body",
        last_structural_type="body",
    )
    cases = [
        ("附件：1.基本情况", "1. 具体情况", "attachment_note"),
        ("1. 具体情况", "2. 超级情况", "attachment_note_item"),
        ("2. 超级情况", "区政府人才保障工作组", "attachment_note_item"),
        ("区政府人才保障工作组", "2025年十月15日", "sign_org"),
        ("2025年十月15日", "附件1", "sign_date"),
        ("附件1", "标题", "attachment_page_mark"),
        ("标题", "测试正文。", "attachment_title"),
        ("测试正文。", "", "attachment_body"),
    ]
    for idx, (line, next_line, expected) in enumerate(cases):
        actual, _, _, _ = detect_structural_type(line, next_line, ctx2)
        assert actual == expected, f"附件结构第{idx}行识别失败: {line} -> {actual}, 期望 {expected}"

    # 附件结构识别回归：附件续项只有 Word 自动编号时，也应进入 attachment_note_item
    ctx3 = DetectionContext(
        has_seen_body=True,
        has_seen_real_body=True,
        prev_type_id="body",
        last_structural_type="body",
    )
    note_type, _, _, _ = detect_structural_type(
        "附件：基本情况", "具体情况", ctx3,
        ParagraphFeatures(paragraph_index=0),
        ParagraphFeatures(numbering_prefix="@lvl_0", paragraph_index=1),
    )
    assert note_type == "attachment_note", f"自动编号附件首行识别失败: {note_type}"
    item_type, _, _, fixed_text = detect_structural_type(
        "具体情况", "区政府人才保障工作组", ctx3,
        ParagraphFeatures(numbering_prefix="@lvl_0", paragraph_index=1),
        ParagraphFeatures(paragraph_index=2),
    )
    assert item_type == "attachment_note_item", f"自动编号附件续项识别失败: {item_type}"
    assert fixed_text.startswith("2. "), f"自动编号附件续项补号失败: {fixed_text}"
    print("✅ DOCX 导入器验证全部通过")
