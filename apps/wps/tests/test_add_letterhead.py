import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from apps.wps.control.add_letterhead import (
    LetterheadCommandError,
    add_letterhead_to_document,
    inspect_letterhead,
    normalize_letterhead_request,
)
from apps.wps.control.document_transaction import DocumentTransactionManager
from docxtool.security import validate_docx_integrity


def form(**changes):
    value = {
        "mark_text": "测试机关文件",
        "document_number": "测发〔2026〕1号",
        "signer": "",
        "separator_style": "straight",
        "replace_existing": False,
    }
    value.update(changes)
    return value


def test_add_letterhead_changes_only_the_output_and_can_be_inspected(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("公文标题")
    document.add_paragraph("正文内容。")
    document.save(source)
    before = source.read_bytes()

    result = add_letterhead_to_document(
        str(source),
        str(output),
        form(separator_style="star"),
        operation_id="operation-test",
        log_dir=tmp_path / "logs",
        request_id="request-test",
    )

    assert result.action == "generated"
    assert source.read_bytes() == before
    assert validate_docx_integrity(output).ok is True
    inspected = inspect_letterhead(str(output))
    assert inspected["status"] == "managed"
    assert inspected["fields"] == {
        "mark_text": "测试机关文件",
        "document_number": "测发〔2026〕1号",
        "signer": "",
        "separator_style": "star",
    }


def test_existing_letterhead_requires_explicit_replace(tmp_path):
    source = tmp_path / "source.docx"
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    document = Document()
    document.add_paragraph("公文标题")
    document.save(source)
    add_letterhead_to_document(
        str(source), str(first), form(),
        operation_id="first", log_dir=tmp_path / "logs",
    )

    with pytest.raises(LetterheadCommandError, match="WPS_LETTERHEAD_ALREADY_EXISTS"):
        add_letterhead_to_document(
            str(first), str(second), form(),
            operation_id="second", log_dir=tmp_path / "logs",
        )
    assert not second.exists()


@pytest.mark.parametrize(
    ("changes", "code"),
    [
        ({"mark_text": ""}, "WPS_LETTERHEAD_MARK_REQUIRED"),
        ({"document_number": "测发[2026]1号"}, "WPS_LETTERHEAD_DOCUMENT_NUMBER_INVALID"),
        ({"separator_style": "dots"}, "WPS_LETTERHEAD_SEPARATOR_INVALID"),
    ],
)
def test_form_validation_uses_stable_codes(changes, code):
    with pytest.raises(LetterheadCommandError, match=code):
        normalize_letterhead_request(form(**changes))


def test_overlong_mark_uses_the_wps_error_code_and_leaves_no_output(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("公文标题")
    document.save(source)

    with pytest.raises(LetterheadCommandError) as exc_info:
        add_letterhead_to_document(
            str(source),
            str(output),
            form(mark_text="机" * 80),
            operation_id="too-long",
            log_dir=tmp_path / "logs",
        )

    assert exc_info.value.code == "WPS_LETTERHEAD_MARK_TOO_LONG"
    assert not output.exists()


def test_joint_source_is_rejected_by_the_single_agency_wps_flow(tmp_path):
    source = tmp_path / "joint.docx"
    document = Document()
    first = document.add_paragraph("测试机关甲")
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    first.runs[0].font.size = Pt(32)
    second = document.add_paragraph("测试机关乙文件")
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    second.runs[0].font.size = Pt(32)
    document.add_paragraph("测发〔2026〕1号")
    separator = document.add_paragraph()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "FF0000")
    borders.append(bottom)
    separator._p.get_or_add_pPr().append(borders)
    document.add_paragraph("关于测试工作的通知")
    document.save(source)

    with pytest.raises(LetterheadCommandError) as exc_info:
        inspect_letterhead(str(source))

    assert exc_info.value.code == "WPS_LETTERHEAD_JOINT_SOURCE_UNSUPPORTED"


def _source_document(path):
    document = Document()
    document.styles["Normal"].paragraph_format.line_spacing = Pt(28.8)
    document.add_paragraph("公文标题")
    document.add_paragraph("正文内容。")
    document.save(path)


def test_letterhead_transaction_commit_and_finalize_updates_the_source(tmp_path):
    source = tmp_path / "source.docx"
    _source_document(source)
    before = source.read_bytes()
    manager = DocumentTransactionManager(tmp_path / "logs")

    operation = manager.prepare_letterhead(
        str(source), form(separator_style="star"), request_id="request-finalize"
    )

    assert source.read_bytes() == before
    assert operation.temporary_path.is_file()
    manager.commit(operation.operation_id, request_id="request-finalize")
    assert source.read_bytes() != before
    assert inspect_letterhead(str(source))["fields"]["separator_style"] == "star"
    manager.finalize(operation.operation_id, request_id="request-finalize")
    assert not operation.backup_path.exists()
    assert not operation.temporary_path.exists()
    assert not manager.journal_path.exists()


def test_letterhead_transaction_rollback_after_commit_restores_the_source(tmp_path):
    source = tmp_path / "source.docx"
    _source_document(source)
    before = source.read_bytes()
    manager = DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare_letterhead(
        str(source), form(), request_id="request-rollback"
    )
    manager.commit(operation.operation_id, request_id="request-rollback")

    manager.rollback(operation.operation_id, request_id="request-rollback")

    assert source.read_bytes() == before
    assert not operation.backup_path.exists()
    assert not operation.temporary_path.exists()
    assert not manager.journal_path.exists()


def test_failed_letterhead_prepare_does_not_publish_a_transaction(tmp_path):
    source = tmp_path / "source.docx"
    _source_document(source)
    manager = DocumentTransactionManager(tmp_path / "logs")

    with pytest.raises(LetterheadCommandError, match="WPS_LETTERHEAD_MARK_TOO_LONG"):
        manager.prepare_letterhead(
            str(source), form(mark_text="机" * 80), request_id="request-failed"
        )

    assert not manager.journal_path.exists()
    assert not list(tmp_path.glob(".source.docxtool-*.docx"))
    operation = manager.prepare_letterhead(
        str(source), form(), request_id="request-next"
    )
    manager.rollback(operation.operation_id, request_id="request-next")


def test_letterhead_gap_uses_the_source_body_line_pitch(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.styles["Normal"].paragraph_format.line_spacing = Pt(28.8)
    document.add_paragraph("公文标题")
    document.save(source)

    add_letterhead_to_document(
        str(source),
        str(output),
        form(),
        operation_id="grid-pitch",
        log_dir=tmp_path / "logs",
    )

    generated = Document(output)
    separator = next(
        paragraph
        for paragraph in generated.paragraphs
        if paragraph.style.style_id == "DCT-LetterheadSeparator"
    )
    spacing = separator._p.get_or_add_pPr().get_or_add_spacing()
    assert spacing.get(qn("w:after")) == "1152"
