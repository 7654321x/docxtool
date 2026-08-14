from __future__ import annotations

import errno
import http.client
import json
from logging.handlers import RotatingFileHandler
from pathlib import Path
import re
import threading
from types import SimpleNamespace
from xml.etree import ElementTree

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import apps.wps.main as wps_main
from apps.wps.control import logging_adapter
from apps.wps.control import server as server_module
from apps.wps.control import document_transaction as transaction_module
from apps.wps.control import format_current_document as format_module
from apps.wps.control import recognize_document as recognition_module
from apps.wps.control.format_current_document import FormatResult
from apps.wps.control.logging_adapter import (
    document_log_context,
    file_identity,
    sanitize_wps_log_fields,
)
from apps.wps.control.recognize_document import bind_preview
from apps.wps.control.server import _safe_warnings
from docxtool.sdk import recognize_docx
from docxtool.wps_server.format_config import load_active_format_profile


def _fake_result(output_path: Path, log_dir: Path) -> FormatResult:
    log_path = log_dir / "test.log"
    log_path.parent.mkdir(parents=True, exist_ok=True)
    log_path.write_text("", encoding="utf-8")
    return FormatResult(
        output_path=output_path,
        log_path=log_path,
        document_mode="NORMAL",
        paragraph_count=3,
        heading_count=1,
        body_count=2,
        export_stats={},
    )


def _add_native_heading2_numbering(document: Document, text: str) -> None:
    numbering = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), "91")
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "chineseCounting")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "（%1）")
    level.extend((start, number_format, level_text))
    abstract.append(level)
    numbering.find(qn("w:num")).addprevious(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), "91")
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), "91")
    number.append(abstract_ref)
    numbering.append(number)

    paragraph = document.add_paragraph(text)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "91")
    num_pr.extend((ilvl, num_id))
    paragraph._p.get_or_add_pPr().append(num_pr)


def _install_fake_formatter(monkeypatch):
    def fake_format(
        source_path,
        output_path,
        *,
        operation_id,
        log_dir,
        format_config=None,
        request_id="",
    ):
        target = Path(output_path)
        target.write_bytes(b"formatted")
        return _fake_result(target, Path(log_dir))

    monkeypatch.setattr(transaction_module, "format_current_document", fake_format)


def _transaction_journal_payload(source: Path) -> dict:
    operation_id = "a" * 32
    return {
        "version": 2,
        "operation_id": operation_id,
        "state": "prepared",
        "source_path": str(source),
        "temporary_path": str(
            source.with_name(f".{source.stem}.docxtool-{operation_id[:12]}.docx")
        ),
        "backup_path": str(
            source.with_name(f".{source.stem}.docxtool-backup-{operation_id[:12]}.docx")
        ),
        "original_source_sha256": "0" * 64,
        "temporary_sha256": "1" * 64,
        "backup_sha256": None,
        "formatted_source_sha256": None,
    }


def _snapshot(raw_text: str, *, snapshot_id: str = "snap-wps") -> dict:
    return {
        "schema_version": "host-snapshot-v1",
        "integration_contract_version": "integration-contract-v1",
        "snapshot_id": snapshot_id,
        "document_identity": "doc-wps",
        "document_revision": "rev-wps",
        "host": {"kind": "wps", "platform": "windows"},
        "host_type": "wps",
        "text_contract_version": "host-text-v1",
        "offset_encoding": "utf16_code_unit",
        "paragraphs": [
            {
                "host_paragraph_id": "main:000000",
                "host_paragraph_index": 0,
                "story_id": "main",
                "story_type": "main",
                "story_paragraph_index": 0,
                "section_index": None,
                "is_in_table": False,
                "raw_text": raw_text,
            }
        ],
    }


def _multi_paragraph_snapshot(texts: list[str]) -> dict:
    snapshot = _snapshot(texts[0])
    snapshot["paragraphs"] = [
        {
            "host_paragraph_id": f"main:{index:06d}",
            "host_paragraph_index": index,
            "story_id": "main",
            "story_type": "main",
            "story_paragraph_index": index,
            "section_index": None,
            "is_in_table": False,
            "raw_text": text,
        }
        for index, text in enumerate(texts)
    ]
    return snapshot


def test_transaction_commit_then_rollback_restores_original(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.prepare(str(source))
    assert source.read_bytes() == b"original"
    assert operation.temporary_path.read_bytes() == b"formatted"
    assert manager.journal_path.is_file()

    manager.commit(operation.operation_id)
    assert source.read_bytes() == b"formatted"
    assert operation.backup_path.read_bytes() == b"original"

    manager.rollback(operation.operation_id)
    assert source.read_bytes() == b"original"
    assert not manager.journal_path.exists()


def test_transaction_finalize_keeps_formatted_document(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.prepare(str(source))
    manager.commit(operation.operation_id)
    manager.finalize(operation.operation_id)

    assert source.read_bytes() == b"formatted"
    assert not operation.backup_path.exists()
    assert not manager.journal_path.exists()


def test_legacy_upgrade_commit_then_rollback_restores_original(tmp_path, monkeypatch):
    source = tmp_path / "sample.doc"
    source.write_bytes(b"legacy-original")
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.reserve_upgrade(str(source), command="apply")
    assert operation.state == "conversion_pending"
    assert operation.target_path == tmp_path / "sample.docx"
    operation.conversion_path.write_bytes(b"converted-docx")

    operation = manager.prepare_upgrade(operation.operation_id)
    assert operation.state == "prepared"
    assert operation.temporary_path.read_bytes() == b"formatted"

    manager.commit(operation.operation_id)
    assert not source.exists()
    assert operation.target_path.read_bytes() == b"formatted"
    assert operation.backup_path.read_bytes() == b"legacy-original"

    manager.rollback(operation.operation_id)
    assert source.read_bytes() == b"legacy-original"
    assert not operation.target_path.exists()
    assert not operation.conversion_path.exists()


def test_legacy_upgrade_finalize_keeps_only_docx(tmp_path, monkeypatch):
    source = tmp_path / "sample.wps"
    source.write_bytes(b"legacy-original")
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.reserve_upgrade(str(source), command="apply")
    operation.conversion_path.write_bytes(b"converted-docx")
    manager.prepare_upgrade(operation.operation_id)
    manager.commit(operation.operation_id)
    manager.finalize(operation.operation_id)

    assert not source.exists()
    assert operation.target_path.read_bytes() == b"formatted"
    assert not operation.backup_path.exists()
    assert not operation.conversion_path.exists()
    assert not manager.journal_path.exists()


def test_legacy_upgrade_rejects_existing_docx_before_reservation(tmp_path):
    source = tmp_path / "sample.doc"
    target = tmp_path / "sample.docx"
    source.write_bytes(b"legacy-original")
    target.write_bytes(b"existing-docx")

    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        manager.reserve_upgrade(str(source), command="preview")

    assert exc_info.value.code == "WPS_LEGACY_UPGRADE_TARGET_EXISTS"
    assert source.read_bytes() == b"legacy-original"
    assert target.read_bytes() == b"existing-docx"
    assert not manager.journal_path.exists()


def test_legacy_upgrade_prepare_converted_preserves_docx_bytes(tmp_path):
    source = tmp_path / "sample.doc"
    source.write_bytes(b"legacy-original")
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.reserve_upgrade(str(source), command="preview")
    operation.conversion_path.write_bytes(b"converted-docx")

    prepared = manager.prepare_converted_upgrade(operation.operation_id)

    assert prepared.state == "prepared"
    assert prepared.format_result is None
    assert prepared.temporary_path.read_bytes() == b"converted-docx"
    assert prepared.conversion_path.read_bytes() == b"converted-docx"
    manager.rollback(operation.operation_id)


@pytest.mark.parametrize("command", ["preview", "apply"])
def test_legacy_upgrade_prepare_rejects_the_wrong_transaction_command(
    tmp_path, monkeypatch, command
):
    source = tmp_path / "sample.doc"
    source.write_bytes(b"legacy-original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.reserve_upgrade(
        str(source), command=command, request_id="request-owner"
    )
    operation.conversion_path.write_bytes(b"converted-docx")

    with pytest.raises(
        transaction_module.DocumentTransactionError,
        match="WPS_TRANSACTION_COMMAND_MISMATCH",
    ):
        if command == "preview":
            manager.prepare_upgrade(
                operation.operation_id, request_id="request-owner"
            )
        else:
            manager.prepare_converted_upgrade(
                operation.operation_id, request_id="request-owner"
            )

    manager.rollback(operation.operation_id, request_id="request-owner")


@pytest.mark.parametrize("action", ["commit", "finalize", "rollback"])
def test_transaction_lifecycle_rejects_a_different_request_id(
    tmp_path, monkeypatch, action
):
    source = tmp_path / f"{action}.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / f"logs-{action}")
    operation = manager.prepare(str(source), request_id="request-owner")
    if action == "finalize":
        manager.commit(operation.operation_id, request_id="request-owner")

    with pytest.raises(
        transaction_module.DocumentTransactionError,
        match="WPS_TRANSACTION_REQUEST_MISMATCH",
    ):
        getattr(manager, action)(
            operation.operation_id, request_id="request-other"
        )

    manager.rollback(operation.operation_id, request_id="request-owner")


def test_legacy_upgrade_recovery_cleans_uncommitted_conversion(tmp_path):
    source = tmp_path / "sample.doc"
    source.write_bytes(b"legacy-original")
    log_dir = tmp_path / "logs"

    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.reserve_upgrade(str(source), command="preview")
    operation.conversion_path.write_bytes(b"converted-docx")

    recovered = transaction_module.DocumentTransactionManager(log_dir)

    assert source.read_bytes() == b"legacy-original"
    assert not operation.target_path.exists()
    assert not operation.conversion_path.exists()
    assert not recovered.journal_path.exists()


def test_legacy_upgrade_recovery_cleans_unjournaled_publish_copy(tmp_path):
    source = tmp_path / "sample.doc"
    source.write_bytes(b"legacy-original")
    log_dir = tmp_path / "logs"
    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.reserve_upgrade(str(source), command="preview")
    operation.conversion_path.write_bytes(b"converted-docx")
    operation.temporary_path.write_bytes(b"converted-docx")

    recovered = transaction_module.DocumentTransactionManager(log_dir)

    assert source.read_bytes() == b"legacy-original"
    assert not operation.target_path.exists()
    assert not operation.conversion_path.exists()
    assert not operation.temporary_path.exists()
    assert not recovered.journal_path.exists()


def test_legacy_upgrade_recovery_restores_committed_source(tmp_path, monkeypatch):
    source = tmp_path / "sample.wps"
    source.write_bytes(b"legacy-original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.reserve_upgrade(str(source), command="apply")
    operation.conversion_path.write_bytes(b"converted-docx")
    manager.prepare_upgrade(operation.operation_id)
    manager.commit(operation.operation_id)

    recovered = transaction_module.DocumentTransactionManager(log_dir)

    assert source.read_bytes() == b"legacy-original"
    assert not operation.target_path.exists()
    assert not operation.backup_path.exists()
    assert not operation.conversion_path.exists()
    assert not recovered.journal_path.exists()


def test_control_legacy_upgrade_routes_share_one_transaction(tmp_path, monkeypatch):
    source = tmp_path / "sample.doc"
    source.write_bytes(b"legacy-original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)
    application = object.__new__(server_module.WpsControlApplication)
    application.log_dir = log_dir
    application.transactions = transaction_module.DocumentTransactionManager(log_dir)
    application._authorization_lock = threading.RLock()
    application._authorized_requests = {
        "request-upgrade": {
            "started_at": 0.0,
            "config_version": "config-1",
            "format_config": {"features": {}},
            "host_generation": 1,
            "state": "authorized",
            "operation_id": "",
        }
    }

    reserved = application.dispatch(
        "/v1/format/upgrade/reserve",
        {"source_path": str(source), "command": "apply"},
        request_id="request-upgrade",
    )
    conversion_path = Path(reserved["conversion_path"])
    conversion_path.write_bytes(b"converted-docx")

    prepared = application.dispatch(
        "/v1/format/upgrade/prepare",
        {"operation_id": reserved["operation_id"]},
        request_id="request-upgrade",
    )
    rolled_back = application.dispatch(
        "/v1/format/rollback",
        {"operation_id": reserved["operation_id"]},
        request_id="request-upgrade",
    )

    assert reserved["state"] == "conversion_pending"
    assert reserved["source_format"] == "doc"
    assert Path(reserved["target_path"]) == source.with_suffix(".docx")
    assert prepared["state"] == "prepared"
    assert rolled_back["state"] == "rolled_back"
    assert source.read_bytes() == b"legacy-original"
    assert not conversion_path.exists()


def test_control_prepare_converted_upgrade_route(tmp_path):
    source = tmp_path / "sample.wps"
    source.write_bytes(b"legacy-original")
    log_dir = tmp_path / "logs"
    application = object.__new__(server_module.WpsControlApplication)
    application.log_dir = log_dir
    application.transactions = transaction_module.DocumentTransactionManager(log_dir)
    reserved = application.dispatch(
        "/v1/format/upgrade/reserve",
        {"source_path": str(source), "command": "preview"},
        request_id="request-preview-upgrade",
    )
    Path(reserved["conversion_path"]).write_bytes(b"converted-docx")

    prepared = application.dispatch(
        "/v1/format/upgrade/prepare-converted",
        {"operation_id": reserved["operation_id"]},
        request_id="request-preview-upgrade",
    )

    assert prepared == {
        "operation_id": reserved["operation_id"],
        "state": "prepared",
    }
    application.dispatch(
        "/v1/format/rollback",
        {"operation_id": reserved["operation_id"]},
        request_id="request-preview-upgrade",
    )


def test_control_format_prepare_requires_public_authorization(tmp_path):
    application = object.__new__(server_module.WpsControlApplication)
    application.log_dir = tmp_path / "logs"
    application.transactions = transaction_module.DocumentTransactionManager(
        application.log_dir
    )
    application._authorization_lock = threading.RLock()
    application._authorized_requests = {}

    with pytest.raises(
        transaction_module.DocumentTransactionError,
        match="WPS_APPLY_AUTHORIZATION_REQUIRED",
    ):
        application.dispatch(
            "/v1/format/prepare",
            {"source_path": str(tmp_path / "sample.docx")},
            request_id="request-without-authorization",
        )


def test_control_uses_authorized_config_once_and_ignores_request_body_config(
    tmp_path, monkeypatch
):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    observed_configs = []

    def fake_format(
        source_path,
        output_path,
        *,
        operation_id,
        log_dir,
        format_config=None,
        request_id="",
    ):
        target = Path(output_path)
        target.write_bytes(b"formatted")
        observed_configs.append(format_config)
        return _fake_result(target, Path(log_dir))

    monkeypatch.setattr(transaction_module, "format_current_document", fake_format)
    application = object.__new__(server_module.WpsControlApplication)
    application.log_dir = log_dir
    application.transactions = transaction_module.DocumentTransactionManager(log_dir)
    application._authorization_lock = threading.RLock()
    application._authorized_requests = {
        "request-authorized": {
            "started_at": 0.0,
            "config_version": "config-1",
            "format_config": {"features": {"numbering": {"enabled": True}}},
            "host_generation": 1,
            "state": "authorized",
            "operation_id": "",
        }
    }

    prepared = application.dispatch(
        "/v1/format/prepare",
        {
            "source_path": str(source),
            "format_config": {"features": {"numbering": {"enabled": False}}},
        },
        request_id="request-authorized",
    )

    assert observed_configs == [
        {"features": {"numbering": {"enabled": True}}}
    ]
    with pytest.raises(
        transaction_module.DocumentTransactionError,
        match="WPS_APPLY_AUTHORIZATION_CONSUMED",
    ):
        application.dispatch(
            "/v1/format/prepare",
            {"source_path": str(source)},
            request_id="request-authorized",
        )
    application.dispatch(
        "/v1/format/rollback",
        {"operation_id": prepared["operation_id"]},
        request_id="request-authorized",
    )


def test_prepare_journal_failure_does_not_publish_operation(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    original_write = Path.write_text

    def fail_journal_write(path, *args, **kwargs):
        if path == manager.journal_path.with_suffix(".tmp"):
            raise OSError("journal unavailable")
        return original_write(path, *args, **kwargs)

    monkeypatch.setattr(Path, "write_text", fail_journal_write)
    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        manager.prepare(str(source))
    assert exc_info.value.code == "WPS_TRANSACTION_JOURNAL_WRITE_FAILED"
    assert not manager._operations
    assert not list(tmp_path.glob(".sample.docxtool-*.docx"))


def test_commit_started_journal_failure_keeps_consistent_state(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(source))
    original_write = Path.write_text

    def fail_journal_write(path, *args, **kwargs):
        if path == manager.journal_path.with_suffix(".tmp"):
            raise OSError("journal unavailable")
        return original_write(path, *args, **kwargs)

    monkeypatch.setattr(Path, "write_text", fail_journal_write)
    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        manager.commit(operation.operation_id)
    assert exc_info.value.code == "WPS_TRANSACTION_JOURNAL_WRITE_FAILED"
    assert operation.state == "prepared"
    assert source.read_bytes() == b"original"
    assert operation.temporary_path.read_bytes() == b"formatted"
    assert not operation.backup_path.exists()


def test_commit_replace_failure_keeps_recoverable_state(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(source))
    original_replace = transaction_module.os.replace

    def fail_document_replace(source_path, destination_path):
        if Path(source_path) == operation.temporary_path and Path(destination_path) == operation.source_path:
            raise OSError("document locked")
        return original_replace(source_path, destination_path)

    monkeypatch.setattr(transaction_module.os, "replace", fail_document_replace)
    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        manager.commit(operation.operation_id)
    assert exc_info.value.code == "WPS_TRANSACTION_REPLACE_FAILED"
    assert operation.state == "commit_started"
    assert source.read_bytes() == b"original"
    assert operation.temporary_path.read_bytes() == b"formatted"
    assert operation.backup_path.read_bytes() == b"original"


def test_commit_backup_copy_failure_keeps_recoverable_state(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(source))

    def fail_backup_copy(*_args, **_kwargs):
        raise OSError("backup unavailable")

    monkeypatch.setattr(transaction_module.shutil, "copy2", fail_backup_copy)
    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        manager.commit(operation.operation_id)
    assert exc_info.value.code == "WPS_TRANSACTION_BACKUP_FAILED"
    assert operation.state == "commit_started"
    assert source.read_bytes() == b"original"
    assert operation.temporary_path.read_bytes() == b"formatted"
    assert not operation.backup_path.exists()


def test_committed_journal_failure_is_recoverable(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(source))
    original_write = Path.write_text
    write_count = 0

    def fail_committed_write(path, *args, **kwargs):
        nonlocal write_count
        if path == manager.journal_path.with_suffix(".tmp"):
            write_count += 1
            if write_count == 3:
                raise OSError("journal unavailable")
        return original_write(path, *args, **kwargs)

    monkeypatch.setattr(Path, "write_text", fail_committed_write)
    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        manager.commit(operation.operation_id)
    assert exc_info.value.code == "WPS_TRANSACTION_JOURNAL_WRITE_FAILED"
    assert operation.state == "commit_started"
    assert source.read_bytes() == b"formatted"
    monkeypatch.undo()
    transaction_module.DocumentTransactionManager(tmp_path / "logs")
    assert source.read_bytes() == b"original"
    assert not manager.journal_path.exists()


def test_recovery_refuses_changed_source_and_preserves_artifacts(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(source))
    manager.commit(operation.operation_id)
    source.write_bytes(b"changed outside transaction")

    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        transaction_module.DocumentTransactionManager(tmp_path / "logs")
    assert exc_info.value.code == "WPS_TRANSACTION_RECOVERY_REQUIRED"
    assert source.read_bytes() == b"changed outside transaction"
    assert operation.backup_path.exists()
    assert manager.journal_path.exists()


def test_recovery_restores_only_verified_formatted_source(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(source))
    manager.commit(operation.operation_id)

    transaction_module.DocumentTransactionManager(tmp_path / "logs")
    assert source.read_bytes() == b"original"
    assert not operation.backup_path.exists()
    assert not manager.journal_path.exists()


def test_recovery_refuses_wrong_backup(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(source))
    manager.commit(operation.operation_id)
    operation.backup_path.write_bytes(b"wrong backup")

    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        transaction_module.DocumentTransactionManager(tmp_path / "logs")
    assert exc_info.value.code == "WPS_TRANSACTION_RECOVERY_REQUIRED"
    assert source.read_bytes() == b"formatted"
    assert operation.backup_path.read_bytes() == b"wrong backup"


def test_restart_discards_committed_transaction_when_backup_is_missing(
    tmp_path, monkeypatch
):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.prepare(str(source))
    manager.commit(operation.operation_id)
    operation.backup_path.unlink()

    restarted = transaction_module.DocumentTransactionManager(log_dir)

    assert source.read_bytes() == b"formatted"
    assert not restarted.journal_path.exists()


def test_second_format_transaction_is_rejected_until_first_finishes(tmp_path, monkeypatch):
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    first.write_bytes(b"first")
    second.write_bytes(b"second")
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(tmp_path / "logs")
    operation = manager.prepare(str(first))
    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        manager.prepare(str(second))
    assert exc_info.value.code == "WPS_FORMAT_BUSY"

    manager.rollback(operation.operation_id)
    replacement = manager.prepare(str(second))
    manager.rollback(replacement.operation_id)


def test_restart_cleans_prepared_transaction(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.prepare(str(source))
    assert operation.temporary_path.exists()

    recovered = transaction_module.DocumentTransactionManager(log_dir)
    assert source.read_bytes() == b"original"
    assert not operation.temporary_path.exists()
    assert not recovered.journal_path.exists()


def test_restart_restores_committed_transaction(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)

    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.prepare(str(source))
    manager.commit(operation.operation_id)
    assert source.read_bytes() == b"formatted"
    assert operation.backup_path.exists()

    recovered = transaction_module.DocumentTransactionManager(log_dir)
    assert source.read_bytes() == b"original"
    assert not operation.backup_path.exists()
    assert not recovered.journal_path.exists()


def test_stale_transaction_recovery_logs_lifecycle(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.prepare(str(source))
    events = []
    monkeypatch.setattr(
        transaction_module,
        "log_event",
        lambda _level, _component, event, _message, _fields=None: events.append(event),
    )

    transaction_module.DocumentTransactionManager(log_dir)

    assert events == [
        "transaction.recovery.start",
        "transaction.recovery.temporary_cleanup.start",
        "transaction.recovery.temporary_cleanup.completed",
        "transaction.recovery.journal_clear.start",
        "transaction.recovery.journal_clear.completed",
        "transaction.recovery.completed",
    ]
    assert not operation.temporary_path.exists()


def test_stale_transaction_source_restore_failure_has_own_event(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(log_dir)
    operation = manager.prepare(str(source))
    manager.commit(operation.operation_id)
    original_replace = transaction_module.os.replace

    def fail_backup_restore(source_path, destination_path):
        if (
            Path(source_path) == operation.backup_path
            and Path(destination_path) == operation.source_path
        ):
            raise OSError("restore failed")
        return original_replace(source_path, destination_path)

    events = []
    monkeypatch.setattr(transaction_module.os, "replace", fail_backup_restore)
    monkeypatch.setattr(
        transaction_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields)
        ),
    )

    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        transaction_module.DocumentTransactionManager(log_dir)

    assert exc_info.value.code == "WPS_TRANSACTION_RECOVERY_SOURCE_RESTORE_FAILED"
    assert [event for event, _fields in events][-3:] == [
        "transaction.recovery.source_restore.start",
        "transaction.recovery.source_restore.failed",
        "transaction.recovery.failed",
    ]


def test_stale_transaction_journal_clear_failure_has_own_event(tmp_path, monkeypatch):
    source = tmp_path / "sample.docx"
    source.write_bytes(b"original")
    log_dir = tmp_path / "logs"
    _install_fake_formatter(monkeypatch)
    manager = transaction_module.DocumentTransactionManager(log_dir)
    manager.prepare(str(source))
    events = []

    def fail_journal_clear(_manager):
        raise OSError("journal clear failed")

    monkeypatch.setattr(
        transaction_module.DocumentTransactionManager,
        "_clear_journal",
        fail_journal_clear,
    )
    monkeypatch.setattr(
        transaction_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields)
        ),
    )

    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        transaction_module.DocumentTransactionManager(log_dir)

    assert exc_info.value.code == "WPS_TRANSACTION_RECOVERY_JOURNAL_CLEAR_FAILED"
    assert [event for event, _fields in events][-3:] == [
        "transaction.recovery.journal_clear.start",
        "transaction.recovery.journal_clear.failed",
        "transaction.recovery.failed",
    ]


@pytest.mark.parametrize(
    ("case", "expected_event", "expected_code"),
    [
        (
            "json",
            "transaction.journal.parse.failed",
            "WPS_TRANSACTION_JOURNAL_JSON_INVALID",
        ),
        (
            "schema",
            "transaction.journal.schema.invalid",
            "WPS_TRANSACTION_JOURNAL_SCHEMA_INVALID",
        ),
        (
            "path",
            "transaction.journal.paths.invalid",
            "WPS_TRANSACTION_JOURNAL_PATH_INVALID",
        ),
        (
            "hash",
            "transaction.journal.hashes.invalid",
            "WPS_TRANSACTION_JOURNAL_HASH_INVALID",
        ),
    ],
)
def test_invalid_stale_transaction_logs_exact_validation_failure(
    tmp_path, monkeypatch, case, expected_event, expected_code
):
    runtime_dir = tmp_path / "runtime"
    runtime_dir.mkdir()
    journal_path = runtime_dir / "transaction-state.json"
    source = tmp_path / "sample.docx"
    payload = _transaction_journal_payload(source)
    if case == "json":
        journal_path.write_text("{", encoding="utf-8")
    elif case == "schema":
        journal_path.write_text("{}", encoding="utf-8")
    else:
        if case == "path":
            payload["source_path"] = str(source.with_suffix(".txt"))
        elif case == "hash":
            payload["original_source_sha256"] = "invalid"
        journal_path.write_text(json.dumps(payload), encoding="utf-8")
    events = []
    monkeypatch.setattr(
        transaction_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append((event, fields)),
    )

    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        transaction_module.DocumentTransactionManager(tmp_path / "logs")

    assert [event for event, _fields in events] == [
        "transaction.recovery.start",
        expected_event,
        "transaction.recovery.failed",
    ]
    assert exc_info.value.code == expected_code
    assert events[-1][1]["error_code"] == expected_code


def test_stale_transaction_journal_read_failure_has_own_event(tmp_path, monkeypatch):
    runtime_dir = tmp_path / "runtime"
    runtime_dir.mkdir()
    journal_path = runtime_dir / "transaction-state.json"
    journal_path.write_text("{}", encoding="utf-8")
    original_read_text = Path.read_text

    def fail_journal_read(path, *args, **kwargs):
        if path == journal_path:
            raise OSError("journal read failed")
        return original_read_text(path, *args, **kwargs)

    events = []
    monkeypatch.setattr(Path, "read_text", fail_journal_read)
    monkeypatch.setattr(
        transaction_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields)
        ),
    )

    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        transaction_module.DocumentTransactionManager(tmp_path / "logs")

    assert [event for event, _fields in events] == [
        "transaction.recovery.start",
        "transaction.journal.read.failed",
        "transaction.recovery.failed",
    ]
    assert exc_info.value.code == "WPS_TRANSACTION_JOURNAL_READ_FAILED"


@pytest.mark.parametrize(
    ("role", "expected_event", "expected_code"),
    [
        (
            "source",
            "transaction.recovery.source_state.failed",
            "WPS_TRANSACTION_SOURCE_STATE_READ_FAILED",
        ),
        (
            "temporary",
            "transaction.recovery.temporary_state.failed",
            "WPS_TRANSACTION_TEMPORARY_STATE_READ_FAILED",
        ),
        (
            "backup",
            "transaction.recovery.backup_state.failed",
            "WPS_TRANSACTION_BACKUP_STATE_READ_FAILED",
        ),
    ],
)
def test_stale_transaction_file_state_failure_has_own_event(
    tmp_path, monkeypatch, role, expected_event, expected_code
):
    runtime_dir = tmp_path / "runtime"
    runtime_dir.mkdir()
    source = tmp_path / "sample.docx"
    payload = _transaction_journal_payload(source)
    paths = {
        "source": source,
        "temporary": Path(payload["temporary_path"]),
        "backup": Path(payload["backup_path"]),
    }
    paths["source"].write_bytes(b"original")
    paths["temporary"].write_bytes(b"formatted")
    paths["backup"].write_bytes(b"original")
    payload["original_source_sha256"] = transaction_module.sha256_file(paths["source"])
    payload["temporary_sha256"] = transaction_module.sha256_file(paths["temporary"])
    payload["backup_sha256"] = payload["original_source_sha256"]
    (runtime_dir / "transaction-state.json").write_text(
        json.dumps(payload), encoding="utf-8"
    )
    original_sha256_file = transaction_module.sha256_file

    def fail_selected_file_state(path):
        if Path(path) == paths[role]:
            raise OSError("file state read failed")
        return original_sha256_file(path)

    events = []
    monkeypatch.setattr(transaction_module, "sha256_file", fail_selected_file_state)
    monkeypatch.setattr(
        transaction_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields)
        ),
    )

    with pytest.raises(transaction_module.DocumentTransactionError) as exc_info:
        transaction_module.DocumentTransactionManager(tmp_path / "logs")

    assert [event for event, _fields in events] == [
        "transaction.recovery.start",
        expected_event,
        "transaction.recovery.failed",
    ]
    assert exc_info.value.code == expected_code


def test_preview_binding_uses_sdk_confirmed_host_range(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("普通正文内容")
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative")
    result = bind_preview(plan, _snapshot("普通正文内容"))

    assert result["confirmed_count"] >= 1
    eligible = [item for item in result["items"] if item["preview_eligible"]]
    assert eligible
    assert all(item["binding_status"] == "confirmed" for item in eligible)
    assert all(item["host_paragraph_index"] == 0 for item in eligible)
    assert all(item["raw_fragment_sha256"] for item in eligible)


def test_preview_binding_marks_canonical_review_range_as_preview_eligible(
    tmp_path, monkeypatch
):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("正文\u00a0内容")
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative")
    events = []
    monkeypatch.setattr(
        recognition_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields)
        ),
    )
    result = bind_preview(plan, _snapshot("正文 内容"))

    assert result["binding_review_count"] >= 1
    eligible = [item for item in result["items"] if item["preview_eligible"]]
    assert eligible
    assert result["confirmed_count"] == 0
    assert result["preview_eligible_count"] == len(eligible)
    assert all(item["binding_status"] == "review" for item in eligible)
    assert all(item["recommended_action"] == "preview_only" for item in eligible)
    warning_events = [
        fields for event, fields in events if event == "binding.item.warning"
    ]
    assert warning_events
    assert all(fields["warning_code"] == "RAW_TEXT_NORMALIZED" for fields in warning_events)
    assert all(fields["physical_paragraph_index"] == 0 for fields in warning_events)
    assert all(fields["physical_occurrence_index"] == 0 for fields in warning_events)
    assert all(fields["physical_text_length_utf16"] == 5 for fields in warning_events)
    assert all(fields["segment_index"] == 0 for fields in warning_events)
    assert all(fields["segment_count"] == 1 for fields in warning_events)
    assert all(fields["locator_verified"] is True for fields in warning_events)
    assert all(fields["locator_status"] == "confirmed" for fields in warning_events)


def test_preview_binding_keeps_ambiguous_range_unresolved_and_ineligible(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("重复内容")
    document.save(source)

    plan = recognize_docx(source, recognition_mode="authoritative")
    snapshot = _snapshot("重复内容")
    duplicate = dict(snapshot["paragraphs"][0])
    duplicate.update(
        host_paragraph_id="main:000001",
        host_paragraph_index=1,
        story_paragraph_index=1,
    )
    snapshot["paragraphs"].append(duplicate)
    result = bind_preview(plan, snapshot)

    assert result["unresolved_count"] >= 1
    assert result["preview_eligible_count"] == 0
    assert not any(item["preview_eligible"] for item in result["items"])


def test_preview_binding_sdk_failure_logs_exact_boundary(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("脱敏正文")
    document.save(source)
    plan = recognize_docx(source)
    events = []
    monkeypatch.setattr(
        recognition_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields)
        ),
    )
    monkeypatch.setattr(
        recognition_module,
        "bind_recognition_plan",
        lambda _plan, _snapshot: (_ for _ in ()).throw(
            recognition_module.DocxToolSdkError(
                "binder failed",
                code="INVALID_RECOGNITION_PLAN",
                details={"path": "$.blocks[1].block_id", "reason": "duplicate_id"},
            )
        ),
    )

    with pytest.raises(RuntimeError, match="WPS_BINDING_SDK_FAILED"):
        recognition_module.bind_preview(plan, {}, request_id="request-binding-fail")

    assert [event for event, _fields in events] == [
        "binding.start",
        "binding.sdk.failed",
    ]
    assert events[-1][1]["request_id"] == "request-binding-fail"
    assert events[-1][1]["error_code"] == "WPS_BINDING_SDK_FAILED"
    assert events[-1][1]["sdk_error_code"] == "INVALID_RECOGNITION_PLAN"
    assert events[-1][1]["sdk_error_path"] == "$.blocks[1].block_id"
    assert events[-1][1]["sdk_error_reason"] == "duplicate_id"


@pytest.mark.parametrize(
    ("case", "expected_event", "expected_code"),
    [
        ("config", "config.load.failed", "WPS_FORMAT_CONFIG_FAILED"),
        ("import", "import.failed", "WPS_FORMAT_IMPORT_FAILED"),
        ("export", "engine.export.failed", "WPS_FORMAT_EXPORT_FAILED"),
        ("integrity", "integrity.validate.failed", "WPS_FORMAT_INTEGRITY_FAILED"),
    ],
)
def test_format_pipeline_failure_logs_exact_stage(
    tmp_path, monkeypatch, case, expected_event, expected_code
):
    source = tmp_path / "source.docx"
    target = tmp_path / "output.docx"
    source.write_bytes(b"source")
    events = []

    class FakeImporter:
        def load(self, *_args, **_kwargs):
            return SimpleNamespace(
                doc_mode="NORMAL",
                paragraphs=[SimpleNamespace(type_id="body")],
            )

    def export_success(*_args, **_kwargs):
        target.write_bytes(b"output")
        return {}

    monkeypatch.setattr(
        format_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields)
        ),
    )
    monkeypatch.setattr(
        format_module,
        "load_rules_and_settings",
        lambda _config: (
            {},
            {},
            {
                "processing": {},
                "numbering": {"enabled": False},
                "punctuation": {"enabled": False},
            },
        ),
    )
    monkeypatch.setattr(format_module, "DocxImporter", FakeImporter)
    monkeypatch.setattr(format_module, "export_doc", export_success)
    monkeypatch.setattr(format_module, "validate_docx_integrity", lambda _path: None)

    if case == "config":
        monkeypatch.setattr(
            format_module,
            "load_rules_and_settings",
            lambda _config: (_ for _ in ()).throw(RuntimeError("config failed")),
        )
    elif case == "import":
        monkeypatch.setattr(
            FakeImporter,
            "load",
            lambda self, *_args, **_kwargs: (_ for _ in ()).throw(
                RuntimeError("import failed")
            ),
        )
    elif case == "export":
        monkeypatch.setattr(
            format_module,
            "export_doc",
            lambda *_args, **_kwargs: (_ for _ in ()).throw(
                RuntimeError("export failed")
            ),
        )
    elif case == "integrity":
        monkeypatch.setattr(
            format_module,
            "validate_docx_integrity",
            lambda _path: (_ for _ in ()).throw(RuntimeError("integrity failed")),
        )

    with pytest.raises(RuntimeError, match=expected_code):
        format_module.format_current_document(
            str(source),
            str(target),
            operation_id="operation-test",
            log_dir=tmp_path / "logs",
            request_id="request-format-fail",
        )

    failure = next(fields for event, fields in events if event == expected_event)
    assert failure["request_id"] == "request-format-fail"
    assert failure["error_code"] == expected_code


def test_wps_one_click_passes_builtin_style_profile_to_engine(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    target = tmp_path / "output.docx"
    source.write_bytes(b"source")
    captured = {}

    class FakeImporter:
        def load(self, *_args, **_kwargs):
            return SimpleNamespace(
                doc_mode="NORMAL",
                paragraphs=[SimpleNamespace(type_id="body")],
            )

    def fake_export(*_args, **kwargs):
        captured.update(kwargs)
        target.write_bytes(b"output")
        return {}

    monkeypatch.setattr(
        format_module,
        "load_rules_and_settings",
        lambda _config: (
            {},
            {},
            {
                "processing": {},
                "numbering": {"enabled": False},
                "punctuation": {"enabled": False},
            },
        ),
    )
    monkeypatch.setattr(format_module, "DocxImporter", FakeImporter)
    monkeypatch.setattr(format_module, "export_doc", fake_export)
    monkeypatch.setattr(format_module, "validate_docx_integrity", lambda _path: None)

    format_module.format_current_document(
        str(source),
        str(target),
        operation_id="operation-style-profile",
        log_dir=tmp_path / "logs",
        request_id="request-style-profile",
    )

    assert captured["style_profile"] == "wps_builtin"


def test_wps_one_click_format_rebuilds_heading_numbering_by_default(tmp_path):
    source = tmp_path / "source.docx"
    target = tmp_path / "output.docx"
    document = Document()
    for text in (
        "测试材料",
        "一、第一部分",
        "（六）第二层",
        "5.第三层",
        "（6）第四层",
        "正文内容正文内容正文内容。",
    ):
        document.add_paragraph(text)
    document.save(source)

    format_module.format_current_document(
        str(source),
        str(target),
        operation_id="operation-numbering",
        log_dir=tmp_path / "logs",
        request_id="request-numbering",
    )

    headings = [
        paragraph.text
        for paragraph in Document(target).paragraphs
        if paragraph.style.style_id
        in {"Heading1", "Heading2", "Heading3", "Heading4"}
    ]
    assert headings == [
        "一、第一部分",
        "（一）第二层",
        "1.第三层",
        "（1）第四层",
    ]


def test_wps_page_scope_recognizes_only_selected_source_paragraphs(tmp_path):
    source = tmp_path / "scoped-source.docx"
    target = tmp_path / "scoped-output.docx"
    document = Document()
    document.sections[0].top_margin = 720000
    first = document.add_paragraph("范围外首段")
    first.style = document.styles["Caption"]
    document.add_paragraph("一、范围内标题")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "范围外表格"
    last = document.add_paragraph("范围外末段")
    last.style = document.styles["Quote"]
    document.save(source)
    source_top_margin = Document(source).sections[0].top_margin

    result = format_module.format_current_document(
        str(source),
        str(target),
        operation_id="operation-page-scope",
        log_dir=tmp_path / "logs",
        request_id="request-page-scope",
        host_snapshot=_multi_paragraph_snapshot(
            ["范围外首段", "一、范围内标题", "范围外表格\r\x07", "范围外末段"]
        ),
        selected_host_paragraph_indexes=[1],
    )

    output = Document(target)
    assert result.paragraph_count == 1
    assert result.heading_count == 1
    assert [paragraph.text for paragraph in output.paragraphs] == [
        "范围外首段",
        "一、范围内标题",
        "范围外末段",
    ]
    assert output.paragraphs[0].style.name == "Caption"
    assert output.paragraphs[1].style.style_id == "Heading1"
    assert output.paragraphs[2].style.name == "Quote"
    assert output.tables[0].cell(0, 0).text == "范围外表格"
    assert output.sections[0].top_margin == source_top_margin


def test_wps_page_scope_rejects_unbound_nonempty_host_paragraph(tmp_path):
    source = tmp_path / "scope-bind-source.docx"
    target = tmp_path / "scope-bind-output.docx"
    document = Document()
    document.add_paragraph("源文档正文")
    document.save(source)

    with pytest.raises(ValueError, match="WPS_FORMAT_SCOPE_BIND_FAILED"):
        format_module.format_current_document(
            str(source),
            str(target),
            operation_id="operation-page-bind-failed",
            log_dir=tmp_path / "logs",
            request_id="request-page-bind-failed",
            host_snapshot=_multi_paragraph_snapshot(["宿主新增正文"]),
            selected_host_paragraph_indexes=[0],
        )

    assert not target.exists()


def test_wps_one_click_rebuilds_native_heading_numbering_from_server_config(tmp_path):
    source = tmp_path / "native-source.docx"
    target = tmp_path / "native-output.docx"
    document = Document()
    document.add_paragraph("测试材料", style="Title")
    document.add_paragraph("正文内容已经开始，并完整说明有关工作情况。")
    _add_native_heading2_numbering(document, "自动编号二级标题")
    document.add_paragraph("后续正文对该标题展开具体说明。")
    document.save(source)

    server_config = load_active_format_profile()["format_config"]
    assert server_config["numbering"]["enabled"] is False
    format_module.format_current_document(
        str(source),
        str(target),
        operation_id="operation-native-numbering",
        log_dir=tmp_path / "logs",
        format_config=server_config,
        request_id="request-native-numbering",
    )

    output = Document(target)
    heading = next(
        paragraph
        for paragraph in output.paragraphs
        if paragraph.style.style_id == "Heading2"
    )
    assert heading.text == "（一）自动编号二级标题"
    assert heading._p.get_or_add_pPr().find(qn("w:numPr")) is None
    assert heading.runs[0].text == "（一）"
    assert heading.runs[0].bold is True


def test_wps_one_click_format_uses_safe_punctuation_by_default(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    target = tmp_path / "output.docx"
    captured = {}
    load_config = format_module.load_rules_and_settings

    def capture_config(config):
        rules, settings, features = load_config(config)
        captured["features"] = features
        return rules, settings, features

    monkeypatch.setattr(format_module, "load_rules_and_settings", capture_config)
    document = Document()
    document.add_paragraph("测试材料")
    document.add_paragraph(
        "请访问 https://example.com/a,b?x=1.2, 并说明:可以吗?"
    )
    document.save(source)

    format_module.format_current_document(
        str(source),
        str(target),
        operation_id="operation-punctuation",
        log_dir=tmp_path / "logs",
        request_id="request-punctuation",
    )

    body_texts = [
        paragraph.text
        for paragraph in Document(target).paragraphs
        if paragraph.style.style_id == "Normal"
    ]
    assert captured["features"]["punctuation"]["enabled"] is True
    assert "请访问 https://example.com/a,b?x=1.2, 并说明：可以吗？" in body_texts


def test_document_log_name_does_not_expose_source_filename(tmp_path):
    source = tmp_path / "private-name.docx"
    source.write_bytes(b"x")
    log_dir = tmp_path / "logs"
    with document_log_context(source, log_dir, "1234567890abcdef") as log_path:
        assert "private-name" not in Path(log_path).name
        assert "document" in Path(log_path).name


def test_file_identity_does_not_expose_path(tmp_path):
    source = tmp_path / "private-name.docx"
    value = file_identity(source)
    assert len(value) == 12
    assert "private-name" not in value


def test_wps_log_rejects_document_name_and_document_path():
    fields = sanitize_wps_log_fields(
        {
            "document_name": "sample.docx",
            "source_path": r"C:\\fixtures\\sample.docx",
        }
    )
    assert fields == {}


def test_wps_log_accepts_style_profile_diagnostic():
    assert sanitize_wps_log_fields({"style_profile": "wps_builtin"}) == {
        "style_profile": "wps_builtin"
    }


def test_wps_log_accepts_taskpane_scroll_diagnostics():
    fields = sanitize_wps_log_fields(
        {
            "stage": "load_settled",
            "root_scroll_top": 80,
            "body_scroll_top": 80,
            "content_scroll_top": 80,
            "inner_width": 390,
            "inner_height": 720,
            "header_top": -64,
            "header_height": 64,
            "header_clipped_top": True,
            "document_has_focus": True,
            "active_element_tag": "BODY",
            "top_element_id": "taskpane_header",
            "scheduled_delay_ms": 100,
            "timer_drift_ms": 4,
            "state_wait_in_flight": True,
        }
    )
    assert fields == {
        "stage": "load_settled",
        "root_scroll_top": 80,
        "body_scroll_top": 80,
        "content_scroll_top": 80,
        "inner_width": 390,
        "inner_height": 720,
        "header_top": -64,
        "header_height": 64,
        "header_clipped_top": True,
        "document_has_focus": True,
        "active_element_tag": "BODY",
        "top_element_id": "taskpane_header",
        "scheduled_delay_ms": 100,
        "timer_drift_ms": 4,
        "state_wait_in_flight": True,
    }


def test_wps_log_accepts_taskpane_host_properties():
    fields = sanitize_wps_log_fields(
        {
            "page_version": "7",
            "pane_branch": "created",
            "pane_dock_position": 2,
            "pane_expected_dock_position": 2,
            "pane_found": True,
            "pane_id": "1",
            "pane_visible": True,
            "pane_width": 390,
        }
    )
    assert fields == {
        "page_version": "7",
        "pane_branch": "created",
        "pane_dock_position": 2,
        "pane_expected_dock_position": 2,
        "pane_found": True,
        "pane_id": "1",
        "pane_visible": True,
        "pane_width": 390,
    }


def test_wps_log_accepts_taskpane_transition_diagnostics():
    fields = sanitize_wps_log_fields(
        {
            "checkpoint": "after_target_activated",
            "observed_delay_ms": 500,
            "active_document_present": True,
            "active_window_present": True,
            "document_matches_expected": True,
            "pane_width_before": 640,
            "pane_width_requested": 390,
            "pane_width_after": 325,
            "pane_width_effective": False,
            "window_screen_x": 80,
            "window_screen_y": 120,
            "screen_width": 1920,
            "screen_avail_height": 1040,
            "physical_header_height": 104,
            "window_top_is_self": True,
            "frame_element_present": False,
            "header_transform": "none",
            "source_path": r"C:\\fixtures\\sample.docx",
        }
    )
    assert fields == {
        "checkpoint": "after_target_activated",
        "observed_delay_ms": 500,
        "active_document_present": True,
        "active_window_present": True,
        "document_matches_expected": True,
        "pane_width_before": 640,
        "pane_width_requested": 390,
        "pane_width_after": 325,
        "pane_width_effective": False,
        "window_screen_x": 80,
        "window_screen_y": 120,
        "screen_width": 1920,
        "screen_avail_height": 1040,
        "physical_header_height": 104,
        "window_top_is_self": True,
        "frame_element_present": False,
        "header_transform": "none",
    }


def test_wps_log_accepts_bridge_diagnostics():
    fields = sanitize_wps_log_fields(
        {
            "bridge_ready": True,
            "command_sequence": 3,
            "generation_changed": False,
            "host_generation": 2,
            "replaced": True,
            "state_revision": 9,
            "wait_timed_out": False,
        }
    )
    assert fields == {
        "bridge_ready": True,
        "command_sequence": 3,
        "generation_changed": False,
        "host_generation": 2,
        "replaced": True,
        "state_revision": 9,
        "wait_timed_out": False,
    }


def test_ribbon_only_keeps_the_taskpane_entry():
    ribbon_path = Path(__file__).resolve().parents[1] / "ribbon.xml"
    root = ElementTree.parse(ribbon_path).getroot()
    namespace = {"ui": "http://schemas.microsoft.com/office/2006/01/customui"}
    buttons = root.findall(".//ui:button", namespace)
    assert [(button.get("id"), button.get("label")) for button in buttons] == [
        ("panel", "打开侧边栏")
    ]
    assert buttons[0].get("getImage") == "GetImage"
    icon_path = ribbon_path.parent / "images" / "taskpane.svg"
    icon_root = ElementTree.parse(icon_path).getroot()
    assert icon_root.tag == "{http://www.w3.org/2000/svg}svg"
    assert icon_root.get("viewBox") == "0 0 24 24"


def test_compatibility_warnings_are_json_safe_and_bounded():
    warnings = _safe_warnings([
        "plain warning",
        {"code": "TEST", "count": 2, "nested": {"hidden": True}},
        object(),
    ])
    assert warnings == ["plain warning", {"code": "TEST", "count": 2}]


def test_wps_production_code_does_not_contain_old_second_format_engine():
    root = Path(__file__).resolve().parents[1]
    forbidden = (
        "LocalFormatCommandGenerator",
        "WpsApiDocumentExecutor",
        "paragraph.set_font",
        "paragraph.set_spacing",
        "section.set_page_setup",
    )
    for path in root.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in {".py", ".js"}:
            continue
        if "tests" in path.parts:
            continue
        text = path.read_text(encoding="utf-8")
        for token in forbidden:
            assert token not in text, f"old formatting engine token {token!r} found in {path}"


def test_wps_bootstrap_structure_is_explicit():
    root = Path(__file__).resolve().parents[1]
    index_source = (root / "index.html").read_text(encoding="utf-8")
    assert './main.js' in index_source
    for token in (
        "runtime/runtime-config.js",
        "js/bootstrap-log.js",
        "js/ribbon.js",
        "host-runtime.js",
        "taskpane.js",
        "PluginStorage",
        "Application",
        "fetch(",
        "runPreview",
        "runFormat",
        "buildHostSnapshot",
        "OnAddinLoad",
        "OnAction",
    ):
        assert token not in index_source

    main_source = (root / "main.js").read_text(encoding="utf-8")
    for token in (
        "runPreview",
        "runFormat",
        "buildHostSnapshot",
        "applyPreviewComments",
        "pollTaskpaneRequests",
        "DocxImporter",
    ):
        assert token not in main_source
    expected_order = (
        "runtime/config",
        "js/bootstrap-log.js",
        "host-runtime.js",
        "js/ribbon.js",
        "js/bootstrap-complete.js",
    )
    for source in expected_order:
        assert source in main_source
    positions = [main_source.index(source) for source in expected_order]
    assert positions == sorted(positions)
    assert "document.write(" not in main_source
    assert "response.json()" in main_source
    assert "document.createElement(\"script\")" in main_source
    assert "?v=" not in main_source
    assert "async " in main_source

    ribbon_source = (root / "js" / "ribbon.js").read_text(encoding="utf-8")
    for callback in ("onAddinLoad", "onAction", "getActionEnabled", "getImage"):
        assert re.search(rf"^function {callback}\(", ribbon_source, re.MULTILINE)
    assert 'return "images/taskpane.svg"' in ribbon_source
    assert "window.DocxToolHostRuntime.start()" in ribbon_source


def test_wps_logging_uses_ten_mib_rotating_file(tmp_path):
    logging_adapter.configure_wps_logging(tmp_path)
    handler = logging_adapter._WPS_FILE_HANDLER
    try:
        assert isinstance(handler, RotatingFileHandler)
        assert handler.maxBytes == 10 * 1024 * 1024
        assert handler.backupCount == 5
        assert handler.encoding.lower().replace("-", "") == "utf8"
        logging_adapter.log_event("INFO", "test", "log.info", "normal-message")
        logging_adapter.log_event("WARNING", "test", "log.warning", "warning-message")
        logging_adapter.log_event("ERROR", "test", "log.error", "error-message")
        handler.flush()
        content = Path(handler.baseFilename).read_text(encoding="utf-8")
        assert "normal-message" in content
        assert "warning-message" in content
        assert "error-message" in content
        assert "\x1b[" not in content
    finally:
        if handler is not None:
            logging_adapter.LOGGER.removeHandler(handler)
            handler.close()
            logging_adapter._WPS_FILE_HANDLER = None


def test_control_request_id_header_reaches_dispatch(monkeypatch, tmp_path):
    observed = []

    def fake_dispatch(self, path, body, request_id=""):
        observed.append((path, body, request_id))
        return {"status": "ok"}

    monkeypatch.setattr(server_module.WpsControlApplication, "dispatch", fake_dispatch)
    server = server_module.create_server(tmp_path, "test-token", 0)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    connection = http.client.HTTPConnection("127.0.0.1", server.server_address[1], timeout=5)
    try:
        connection.request(
            "POST",
            "/v1/recognize",
            body=json.dumps({"source_path": "ignored"}),
            headers={
                "Authorization": "Bearer test-token",
                "Content-Type": "application/json",
                "X-DocxTool-Request-Id": "pane-request-123",
            },
        )
        response = connection.getresponse()
        response.read()
        assert response.status == 200
    finally:
        connection.close()
        server.shutdown()
        server.server_close()
        thread.join(timeout=3)
    assert observed == [
        ("/v1/recognize", {"source_path": "ignored"}, "pane-request-123")
    ]


def _control_post(server, path, *, body=b"{}", token="test-token", headers=None):
    connection = http.client.HTTPConnection(
        "127.0.0.1", server.server_address[1], timeout=5
    )
    request_headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "X-DocxTool-Request-Id": "boundary-request",
    }
    request_headers.update(headers or {})
    try:
        connection.request("POST", path, body=body, headers=request_headers)
        response = connection.getresponse()
        payload = json.loads(response.read().decode("utf-8"))
        return response.status, payload
    finally:
        connection.close()


@pytest.mark.parametrize(
    ("case", "path", "body", "token", "headers", "expected_status", "expected_code", "expected_event"),
    [
        (
            "unauthorized",
            "/v1/recognize",
            b"{}",
            "wrong-token",
            {},
            401,
            "WPS_CONTROL_UNAUTHORIZED",
            "control.auth.rejected",
        ),
        (
            "invalid-content-length",
            "/v1/recognize",
            b"",
            "test-token",
            {"Content-Length": "invalid"},
            400,
            "WPS_CONTROL_INVALID_CONTENT_LENGTH",
            "control.body.length_invalid",
        ),
        (
            "negative-content-length",
            "/v1/recognize",
            b"",
            "test-token",
            {"Content-Length": "-1"},
            400,
            "WPS_CONTROL_NEGATIVE_CONTENT_LENGTH",
            "control.body.length_negative",
        ),
        (
            "body-too-large",
            "/v1/recognize",
            b"",
            "test-token",
            {"Content-Length": str(server_module.MAX_BODY_BYTES + 1)},
            400,
            "WPS_CONTROL_REQUEST_TOO_LARGE",
            "control.body.too_large",
        ),
        (
            "invalid-json",
            "/v1/recognize",
            b"{",
            "test-token",
            {},
            400,
            "WPS_CONTROL_JSON_INVALID",
            "control.body.json_invalid",
        ),
        (
            "json-object-required",
            "/v1/recognize",
            b"[]",
            "test-token",
            {},
            400,
            "WPS_CONTROL_JSON_OBJECT_REQUIRED",
            "control.body.object_required",
        ),
        (
            "route-not-found",
            "/v1/missing",
            b"{}",
            "test-token",
            {},
            404,
            "WPS_CONTROL_ROUTE_NOT_FOUND",
            "control.route.not_found",
        ),
    ],
)
def test_control_boundaries_have_distinct_events(
    monkeypatch,
    tmp_path,
    case,
    path,
    body,
    token,
    headers,
    expected_status,
    expected_code,
    expected_event,
):
    events = []
    monkeypatch.setattr(
        server_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields or {})
        ),
    )
    server = server_module.create_server(tmp_path / case, "test-token", 0)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        status, payload = _control_post(
            server, path, body=body, token=token, headers=headers
        )
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=3)

    assert status == expected_status
    assert payload == {"ok": False, "error_code": expected_code}
    matching = [fields for event, fields in events if event == expected_event]
    assert len(matching) == 1
    assert matching[0]["error_code"] == expected_code
    assert not any(
        event == "control.request.execution_failed" for event, _fields in events
    )


@pytest.mark.parametrize(
    ("monitor_state", "expected_code", "expected_event"),
    [
        ("busy", "WPS_COMMAND_BUSY", "control.command.busy"),
        ("stopped", "WPS_MONITOR_NOT_RUNNING", "control.monitor.unavailable"),
    ],
)
def test_control_monitor_boundaries_have_distinct_events(
    monkeypatch, tmp_path, monitor_state, expected_code, expected_event
):
    events = []
    monkeypatch.setattr(
        server_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append(
            (event, fields or {})
        ),
    )
    server = server_module.create_server(tmp_path / monitor_state, "test-token", 0)
    if monitor_state == "busy":
        with server.command_monitor._lock:
            server.command_monitor._occupied = True
    else:
        server.command_monitor.stop()
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        status, payload = _control_post(server, "/v1/recognize")
    finally:
        if monitor_state == "busy":
            with server.command_monitor._lock:
                server.command_monitor._occupied = False
        server.shutdown()
        server.server_close()
        thread.join(timeout=3)

    assert status == 400
    assert payload == {"ok": False, "error_code": expected_code}
    matching = [fields for event, fields in events if event == expected_event]
    assert len(matching) == 1
    assert matching[0]["error_code"] == expected_code


def test_log_route_emits_only_the_submitted_diagnostic_event(monkeypatch, tmp_path):
    events = []
    monkeypatch.setattr(
        server_module,
        "log_event",
        lambda _level, component, event, _message, fields=None: events.append(
            (component, event, fields or {})
        ),
    )
    server = server_module.create_server(tmp_path, "test-token", 0)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    connection = http.client.HTTPConnection(
        "127.0.0.1", server.server_address[1], timeout=5
    )
    try:
        body = json.dumps(
            {
                "level": "WARNING",
                "component": "taskpane",
                "event": "taskpane.request.blocked",
                "message": "blocked",
                "details": {
                    "request_id": "request-42",
                    "reason": "context_not_ready",
                },
            }
        )
        connection.request(
            "POST",
            "/v1/log",
            body=body,
            headers={
                "Authorization": "Bearer test-token",
                "Content-Type": "application/json",
            },
        )
        response = connection.getresponse()
        response.read()
        assert response.status == 200
        submitted_events = list(events)
    finally:
        connection.close()
        server.shutdown()
        server.server_close()
        thread.join(timeout=3)

    assert submitted_events == [
        (
            "taskpane",
            "taskpane.request.blocked",
            {"request_id": "request-42", "reason": "context_not_ready"},
        )
    ]


def test_log_detail_contract_keeps_diagnostics_and_rejects_sensitive_fields():
    details = server_module._safe_log_details(
        {
            "request_id": "request-42",
            "event_sequence": 9,
            "reason": "context_not_ready",
            "request_status": "BLOCKED",
            "host_ready": True,
            "host_instance_id_short": "host-1234abcd",
            "physical_paragraph_index": 12,
            "physical_occurrence_index": 1,
            "physical_text_length_utf16": 48,
            "segment_index": 0,
            "segment_count": 2,
            "locator_verified": True,
            "locator_status": "confirmed",
            "document_id_short": "doc-1234",
            "operation_id_short": "operation-12",
            "plan_id_short": "plan-1234",
            "body": "private body",
            "source_path": "C:/private/document.docx",
            "raw_text": "private body",
            "session_token": "secret",
            "sha256": "f" * 64,
            "traceback": "private traceback",
        }
    )

    assert details == {
        "request_id": "request-42",
        "event_sequence": 9,
        "reason": "context_not_ready",
        "request_status": "BLOCKED",
        "host_ready": True,
        "host_instance_id_short": "host-1234abcd",
        "physical_paragraph_index": 12,
        "physical_occurrence_index": 1,
        "physical_text_length_utf16": 48,
        "segment_index": 0,
        "segment_count": 2,
        "locator_verified": True,
        "locator_status": "confirmed",
        "document_id_short": "doc-1234",
        "operation_id_short": "operation-12",
        "plan_id_short": "plan-1234",
    }
    assert server_module._safe_log_details({"path": "/v1/health"}) == {
        "path": "/v1/health"
    }
    assert server_module._safe_log_details(
        {"path": "C:/private/document.docx"}
    ) == {}


def test_python_log_field_contract_rejects_sensitive_values():
    fields = logging_adapter._fields_text(
        {
            "request_id": "request-42",
            "error_code": "WPS_TEST_FAILED",
            "body": "private body",
            "source_path": "C:/private/document.docx",
            "raw_text": "private body",
            "session_token": "secret",
            "sha256": "f" * 64,
            "traceback": "private traceback",
        }
    )

    assert "request_id=request-42" in fields
    assert "error_code=WPS_TEST_FAILED" in fields
    assert "private" not in fields
    assert "secret" not in fields
    assert "traceback" not in fields


def test_python_log_field_contract_keeps_safe_sdk_validation_details():
    fields = sanitize_wps_log_fields(
        {
            "sdk_error_code": "INVALID_RECOGNITION_PLAN",
            "sdk_error_path": "$.blocks[1].block_id",
            "sdk_error_reason": "duplicate_id",
            "raw_text": "private body",
        }
    )

    assert fields == {
        "sdk_error_code": "INVALID_RECOGNITION_PLAN",
        "sdk_error_path": "$.blocks[1].block_id",
        "sdk_error_reason": "duplicate_id",
    }
    assert "C:/private" not in logging_adapter._fields_text(
        {"path": "C:/private/document.docx"}
    )
    assert "path=/v1/health" in logging_adapter._fields_text({"path": "/v1/health"})


def test_control_recognize_and_bind_logs_keep_request_id(monkeypatch, tmp_path):
    application = object.__new__(server_module.WpsControlApplication)
    application.log_dir = tmp_path
    application._remember_plan = lambda _plan: None
    application._get_plan = lambda _plan_id: object()
    events = []
    monkeypatch.setattr(
        server_module,
        "log_event",
        lambda _level, _component, event, _message, fields=None: events.append((event, fields)),
    )
    monkeypatch.setattr(
        server_module,
        "recognize_document",
        lambda *_args, **_kwargs: SimpleNamespace(
            plan=object(), public_result={"block_count": 3}
        ),
    )
    monkeypatch.setattr(
        server_module,
        "bind_preview",
        lambda _plan, _snapshot, request_id="": {
            "confirmed_count": 2,
            "unresolved_count": 1,
        },
    )

    application.dispatch("/v1/recognize", {}, request_id="request-42")
    application.dispatch(
        "/v1/recognize/bind",
        {"plan_id": "plan", "host_snapshot": {}},
        request_id="request-42",
    )

    assert [event for event, _fields in events] == [
        "recognize.request.start",
        "recognize.request.completed",
        "bind.request.start",
        "bind.request.completed",
    ]
    assert all(fields["request_id"] == "request-42" for _event, fields in events)


def test_verify_files_requires_new_bootstrap_files(monkeypatch, tmp_path):
    monkeypatch.setattr(wps_main, "APP_ROOT", tmp_path)
    for relative in (
        "package.json",
        "manifest.xml",
        "ribbon.xml",
        "index.html",
        "main.js",
        "js/bootstrap-log.js",
            "js/bootstrap-complete.js",
            "js/ribbon.js",
            "images/taskpane.svg",
            "images/check.svg",
            "images/eye.svg",
            "images/eye-off.svg",
            "images/taskpane-icons.svg",
            "images/login-window.png",
            "images/user.svg",
            "host-runtime.js",
        "taskpane.html",
        "taskpane.js",
        "client-config.json",
        "reader/reader-client.js",
        "reader/reader-ui.js",
        "reader/reader.css",
        "account_store.py",
        "account_runtime.py",
        "public_api.py",
        "login_window.py",
        "desktop_runtime.py",
        "windows_startup.py",
            "control/server.py",
            "control/host_bridge.py",
            "control/format_current_document.py",
            "control/add_letterhead.py",
        "control/reader_routes.py",
        "control/document_transaction.py",
        "control/logging_adapter.py",
        "control/recognize_document.py",
        "control/monitor.py",
    ):
        path = tmp_path / relative
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text("", encoding="utf-8")
    (tmp_path / "package.json").write_text(
        '{"devDependencies":{"wpsjs":"2.2.3"},'
        '"overrides":{"wpsjs-rpc-sdk-new":"1.1.0"}}',
        encoding="utf-8",
    )

    wps_main.verify_files()


def test_taskpane_scrolls_content_without_moving_header():
    source = (wps_main.APP_ROOT / "taskpane.html").read_text(encoding="utf-8")

    assert "html,body{height:100%;max-width:100%;overflow:hidden}" in source
    assert "body{display:flex;flex-direction:column" in source
    assert "header{flex:0 0 auto" in source
    assert "main{flex:1 1 auto;min-height:0;overflow-y:auto" in source
    assert '<header id="taskpane_header">' in source
    assert 'id="focus_document"' not in source
    assert "返回文档" not in source
    assert '<div id="status" class="brand-status">连接中</div>' in source
    assert 'document.getElementById("status").textContent="错误"' in source
    assert 'document.getElementById("message").textContent="运行配置加载失败，请重新打开状态面板。"' in source
    assert 'document.getElementById("error").textContent="错误代码：WPS_RUNTIME_CONFIG_LOAD_FAILED"' in source
    assert '<main id="content">' in source
    assert 'fetch("./runtime/config"' in source
    assert 'load("./reader/reader-client.js?v=1")' in source
    assert 'load("./reader/reader-ui.js?v=2")' in source
    assert 'load("./taskpane.js?v=15")' in source


def test_start_handles_keyboard_interrupt_without_traceback(monkeypatch):
    events = []

    class FakeControlServer:
        def serve_forever(self, **_kwargs):
            pass

        def shutdown(self):
            events.append("control-shutdown")

        def server_close(self):
            events.append("control-close")

    class FakeWebServer:
        def serve_forever(self, **_kwargs):
            raise KeyboardInterrupt

        def server_close(self):
            events.append("web-close")

    class FakeThread:
        def __init__(self, **_kwargs):
            pass

        def start(self):
            events.append("thread-start")

        def join(self, timeout):
            events.append(("join", timeout))

        def is_alive(self):
            return False

    class FakeAccountRuntime:
        def start(self):
            events.append("account-start")

        def stop(self):
            events.append("account-stop")

    monkeypatch.setattr(wps_main, "verify_files", lambda: None)
    monkeypatch.setattr(wps_main, "configure_wps_logging", lambda _root: None)
    monkeypatch.setattr(
        wps_main, "_start_control", lambda _port, _runtime: (FakeControlServer(), 45678)
    )
    monkeypatch.setattr(
        wps_main, "_start_web_server", lambda _port: (FakeWebServer(), 3889)
    )
    monkeypatch.setattr(
        wps_main, "_publish_addin", lambda port: events.append(("publish", port))
    )
    monkeypatch.setattr(wps_main.threading, "Thread", FakeThread)
    monkeypatch.setattr(wps_main, "clear_runtime_config", lambda: events.append("clear"))
    monkeypatch.setattr(wps_main, "log_event", lambda _level, _component, event, _message, _fields=None: events.append(event))

    wps_main.start(0, FakeAccountRuntime())

    assert "launcher.interrupt.received" in events
    assert ("publish", 3889) in events
    assert events.index("account-start") < events.index(("publish", 3889))
    assert "web-close" in events
    assert "control-shutdown" in events
    assert events[-1] == "launcher.session.stop"


def test_start_requires_account_runtime_before_any_service_side_effect(monkeypatch):
    calls = []

    monkeypatch.setattr(wps_main, "verify_files", lambda: calls.append("verify"))
    monkeypatch.setattr(
        wps_main, "configure_wps_logging", lambda _root: calls.append("logging")
    )
    monkeypatch.setattr(
        wps_main, "_start_control", lambda *_args: calls.append("control")
    )
    monkeypatch.setattr(
        wps_main, "_start_web_server", lambda *_args: calls.append("web")
    )
    monkeypatch.setattr(
        wps_main, "_publish_addin", lambda *_args: calls.append("publish")
    )
    monkeypatch.setattr(
        wps_main, "clear_runtime_config", lambda: calls.append("runtime-config")
    )
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    with pytest.raises(RuntimeError, match="WPS_ACCOUNT_RUNTIME_REQUIRED"):
        wps_main.start(0)

    assert calls == []


@pytest.mark.parametrize(
    ("failed_stage", "expected_error", "expected_event", "expected_error_code"),
    [
        (
            "account-stop",
            "ACCOUNT_STOP_FAILED",
            "launcher.account_runtime.stop.failed",
            "WPS_ACCOUNT_RUNTIME_STOP_FAILED",
        ),
        (
            "web-close",
            "WEB_CLOSE_FAILED",
            "launcher.web.close.failed",
            "WPS_WEB_SERVER_CLOSE_FAILED",
        ),
        (
            "control-shutdown",
            "CONTROL_SHUTDOWN_FAILED",
            "launcher.control.shutdown.failed",
            "WPS_CONTROL_SERVER_SHUTDOWN_FAILED",
        ),
        (
            "control-close",
            "CONTROL_CLOSE_FAILED",
            "launcher.control.close.failed",
            "WPS_CONTROL_SERVER_CLOSE_FAILED",
        ),
        (
            "thread-join",
            "THREAD_JOIN_FAILED",
            "launcher.control.thread.join.failed",
            "WPS_CONTROL_THREAD_JOIN_FAILED",
        ),
        (
            "thread-state",
            "THREAD_STATE_FAILED",
            "launcher.control.thread.state_check.failed",
            "WPS_CONTROL_THREAD_STATE_CHECK_FAILED",
        ),
        (
            "thread-timeout",
            "WPS_CONTROL_THREAD_STOP_TIMEOUT",
            "launcher.control.thread.stop_timeout",
            "WPS_CONTROL_THREAD_STOP_TIMEOUT",
        ),
        (
            "runtime-config",
            "RUNTIME_CONFIG_FAILED",
            "launcher.runtime_config.cleanup.failed",
            "WPS_RUNTIME_CONFIG_CLEANUP_FAILED",
        ),
    ],
)
def test_start_attempts_all_cleanup_after_each_failure(
    monkeypatch,
    failed_stage,
    expected_error,
    expected_event,
    expected_error_code,
):
    events = []
    log_records = []

    def step(name):
        events.append(name)
        if name == failed_stage:
            raise RuntimeError(expected_error)

    class FakeAccountRuntime:
        def start(self):
            events.append("account-start")

        def stop(self):
            step("account-stop")

    class FakeControlServer:
        def serve_forever(self, **_kwargs):
            pass

        def shutdown(self):
            step("control-shutdown")

        def server_close(self):
            step("control-close")

    class FakeWebServer:
        def serve_forever(self, **_kwargs):
            raise KeyboardInterrupt

        def server_close(self):
            step("web-close")

    class FakeThread:
        def __init__(self, **_kwargs):
            pass

        def start(self):
            events.append("thread-start")

        def join(self, timeout):
            events.append(("thread-join-timeout", timeout))
            step("thread-join")

        def is_alive(self):
            if failed_stage == "thread-timeout":
                events.append("thread-state")
                return True
            step("thread-state")
            return False

    monkeypatch.setattr(wps_main, "verify_files", lambda: None)
    monkeypatch.setattr(wps_main, "configure_wps_logging", lambda _root: None)
    monkeypatch.setattr(
        wps_main,
        "_start_control",
        lambda _port, _runtime: (FakeControlServer(), 45678),
    )
    monkeypatch.setattr(
        wps_main, "_start_web_server", lambda _port: (FakeWebServer(), 3889)
    )
    monkeypatch.setattr(wps_main, "_publish_addin", lambda _port: None)
    monkeypatch.setattr(wps_main.threading, "Thread", FakeThread)
    monkeypatch.setattr(wps_main, "clear_runtime_config", lambda: step("runtime-config"))
    monkeypatch.setattr(
        wps_main,
        "log_event",
        lambda _level, _component, event, _message, fields=None: (
            events.append(event),
            log_records.append((event, fields or {})),
        ),
    )

    with pytest.raises(RuntimeError, match=expected_error):
        wps_main.start(0, FakeAccountRuntime())

    assert expected_event in events
    assert (expected_event, expected_error_code) in [
        (event, fields.get("error_code")) for event, fields in log_records
    ]
    for cleanup_stage in (
        "account-stop",
        "web-close",
        "control-shutdown",
        "control-close",
        "thread-join",
        "runtime-config",
    ):
        assert cleanup_stage in events


def test_start_raises_first_cleanup_failure_and_continues(monkeypatch):
    events = []

    class FakeControlServer:
        def serve_forever(self, **_kwargs):
            pass

        def shutdown(self):
            events.append("control-shutdown")
            raise RuntimeError("CONTROL_SHUTDOWN_SECOND")

        def server_close(self):
            events.append("control-close")
            raise RuntimeError("CONTROL_CLOSE_THIRD")

    class FakeWebServer:
        def serve_forever(self, **_kwargs):
            raise KeyboardInterrupt

        def server_close(self):
            events.append("web-close")
            raise RuntimeError("WEB_CLOSE_FIRST")

    class FakeThread:
        def __init__(self, **_kwargs):
            pass

        def start(self):
            pass

        def join(self, timeout):
            events.append(("thread-join", timeout))

        def is_alive(self):
            return False

    class FakeAccountRuntime:
        def start(self):
            return None

        def stop(self):
            return None

    monkeypatch.setattr(wps_main, "verify_files", lambda: None)
    monkeypatch.setattr(wps_main, "configure_wps_logging", lambda _root: None)
    monkeypatch.setattr(
        wps_main, "_start_control", lambda _port, _runtime: (FakeControlServer(), 45678)
    )
    monkeypatch.setattr(
        wps_main, "_start_web_server", lambda _port: (FakeWebServer(), 3889)
    )
    monkeypatch.setattr(wps_main, "_publish_addin", lambda _port: None)
    monkeypatch.setattr(wps_main.threading, "Thread", FakeThread)
    monkeypatch.setattr(
        wps_main, "clear_runtime_config", lambda: events.append("runtime-config")
    )
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    with pytest.raises(RuntimeError, match="WEB_CLOSE_FIRST"):
        wps_main.start(0, FakeAccountRuntime())

    assert events == [
        "web-close",
        "control-shutdown",
        "control-close",
        ("thread-join", 3),
        "runtime-config",
    ]


def test_start_preserves_business_error_when_cleanup_also_fails(monkeypatch):
    events = []

    class FakeControlServer:
        def serve_forever(self, **_kwargs):
            pass

        def shutdown(self):
            events.append("control-shutdown")
            raise RuntimeError("CLEANUP_FAILED")

        def server_close(self):
            events.append("control-close")

    class FakeWebServer:
        def serve_forever(self, **_kwargs):
            raise ValueError("BUSINESS_FAILED")

        def server_close(self):
            events.append("web-close")

    class FakeThread:
        def __init__(self, **_kwargs):
            pass

        def start(self):
            pass

        def join(self, timeout):
            events.append(("thread-join", timeout))

        def is_alive(self):
            return False

    class FakeAccountRuntime:
        def start(self):
            return None

        def stop(self):
            return None

    monkeypatch.setattr(wps_main, "verify_files", lambda: None)
    monkeypatch.setattr(wps_main, "configure_wps_logging", lambda _root: None)
    monkeypatch.setattr(
        wps_main, "_start_control", lambda _port, _runtime: (FakeControlServer(), 45678)
    )
    monkeypatch.setattr(
        wps_main, "_start_web_server", lambda _port: (FakeWebServer(), 3889)
    )
    monkeypatch.setattr(wps_main, "_publish_addin", lambda _port: None)
    monkeypatch.setattr(wps_main.threading, "Thread", FakeThread)
    monkeypatch.setattr(
        wps_main, "clear_runtime_config", lambda: events.append("runtime-config")
    )
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    with pytest.raises(ValueError, match="BUSINESS_FAILED"):
        wps_main.start(0, FakeAccountRuntime())

    assert events == [
        "web-close",
        "control-shutdown",
        "control-close",
        ("thread-join", 3),
        "runtime-config",
    ]


def test_runtime_config_is_kept_in_launcher_memory():
    wps_main.write_runtime_config(9527, "test-token")
    assert wps_main._RUNTIME_CONFIG == {
        "controlBaseUrl": "http://127.0.0.1:9527",
        "sessionToken": "test-token",
    }
    wps_main.clear_runtime_config()
    assert wps_main._RUNTIME_CONFIG == {}


def test_wps_static_server_serves_plugin_without_launching_wps(tmp_path, monkeypatch):
    (tmp_path / "index.html").write_text("WPS_BACKGROUND_READY", encoding="utf-8")
    monkeypatch.setattr(wps_main, "APP_ROOT", tmp_path)
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    wps_main.write_runtime_config(9527, "test-token")
    server, port = wps_main._start_web_server(0)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    connection = http.client.HTTPConnection("127.0.0.1", port, timeout=5)
    try:
        connection.request("GET", "/index.html")
        response = connection.getresponse()
        assert response.status == 200
        assert response.getheader("Cache-Control") == (
            "no-store, no-cache, must-revalidate, max-age=0"
        )
        assert response.getheader("Pragma") == "no-cache"
        assert response.getheader("Expires") == "0"
        assert response.read().decode("utf-8") == "WPS_BACKGROUND_READY"
        connection.request("GET", "/runtime/config")
        response = connection.getresponse()
        assert response.status == 200
        assert response.getheader("Content-Type").startswith("application/json")
        assert response.getheader("X-Content-Type-Options") == "nosniff"
        assert response.getheader("Access-Control-Allow-Origin") is None
        assert json.loads(response.read()) == {
            "controlBaseUrl": "http://127.0.0.1:9527",
            "sessionToken": "test-token",
        }
        connection.request("GET", "/runtime/runtime-config.js")
        assert connection.getresponse().status == 404
    finally:
        connection.close()
        server.shutdown()
        server.server_close()
        thread.join(timeout=3)


def test_wps_web_server_uses_stable_addin_origin():
    assert wps_main.DEFAULT_WEB_PORT == 3889
    assert wps_main._WpsStaticHttpServer.allow_reuse_address is False


def test_wps_web_server_reports_the_fixed_port_conflict(monkeypatch):
    def address_in_use(*_args, **_kwargs):
        raise OSError(errno.EADDRINUSE, "address already in use")

    monkeypatch.setattr(wps_main, "_WpsStaticHttpServer", address_in_use)
    monkeypatch.setattr(wps_main, "_stop_previous_docxtool_service", lambda _port: False)
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    with pytest.raises(RuntimeError, match="WPS_WEB_SERVER_PORT_IN_USE") as exc_info:
        wps_main._start_web_server(wps_main.DEFAULT_WEB_PORT)

    assert isinstance(exc_info.value.__cause__, OSError)


def test_wps_web_server_retries_after_stopping_verified_previous_service(monkeypatch):
    attempts = []
    stopped = []

    class FakeServer:
        server_address = ("127.0.0.1", 3889)

    def create_server(*_args, **_kwargs):
        attempts.append(True)
        if len(attempts) == 1:
            raise OSError(errno.EADDRINUSE, "address already in use")
        return FakeServer()

    monkeypatch.setattr(wps_main, "_WpsStaticHttpServer", create_server)
    monkeypatch.setattr(
        wps_main,
        "_stop_previous_docxtool_service",
        lambda port: stopped.append(port) or True,
    )
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    _server, port = wps_main._start_web_server(wps_main.DEFAULT_WEB_PORT)

    assert port == 3889
    assert len(attempts) == 2
    assert stopped == [3889]


def test_publish_addin_updates_only_docxtool_entry(tmp_path, monkeypatch):
    publish_path = tmp_path / "publish.xml"
    publish_path.write_text(
        "<jsplugins>"
        '<jspluginonline name="another-addin" type="wps" url="http://example.test/" />'
        '<jspluginonline name="docxtool-wps-app" type="wps" url="http://127.0.0.1:3999/" />'
        '<jspluginonline name="docxtool-wps-app" type="wps" url="http://127.0.0.1:4000/" />'
        "</jsplugins>",
        encoding="utf-8",
    )
    monkeypatch.setattr(wps_main, "_publish_xml_path", lambda: publish_path)
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    wps_main._publish_addin(3889)

    root = ElementTree.parse(publish_path).getroot()
    entries = list(root)
    assert any(node.get("name") == "another-addin" for node in entries)
    docxtool = [node for node in entries if node.get("name") == "docxtool-wps-app"]
    assert len(docxtool) == 1
    assert docxtool[0].attrib == {
        "name": "docxtool-wps-app",
        "type": "wps",
        "url": "http://127.0.0.1:3889/",
        "debug": "",
        "enable": "enable_dev",
        "install": "null",
    }


def test_unpublish_addin_removes_only_docxtool_entries(tmp_path, monkeypatch):
    publish_path = tmp_path / "publish.xml"
    publish_path.write_text(
        "<jsplugins>"
        '<jspluginonline name="another-addin" type="wps" url="http://example.test/" />'
        '<jspluginonline name="docxtool-wps-app" type="wps" url="http://127.0.0.1:3889/" />'
        '<jspluginonline name="docxtool-wps-app" type="wps" url="http://127.0.0.1:3999/" />'
        "</jsplugins>",
        encoding="utf-8",
    )
    monkeypatch.setattr(wps_main, "_publish_xml_path", lambda: publish_path)
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    wps_main._unpublish_addin()

    entries = list(ElementTree.parse(publish_path).getroot())
    assert [node.get("name") for node in entries] == ["another-addin"]


def test_unpublish_addin_reports_atomic_write_failure(tmp_path, monkeypatch):
    publish_path = tmp_path / "publish.xml"
    publish_path.write_text(
        "<jsplugins>"
        '<jspluginonline name="docxtool-wps-app" type="wps" url="http://127.0.0.1:3889/" />'
        "</jsplugins>",
        encoding="utf-8",
    )
    monkeypatch.setattr(wps_main, "_publish_xml_path", lambda: publish_path)
    monkeypatch.setattr(wps_main, "log_event", lambda *_args, **_kwargs: None)

    def fail_write(self, *_args, **_kwargs):
        raise OSError("write failed")

    monkeypatch.setattr(ElementTree.ElementTree, "write", fail_write)

    with pytest.raises(RuntimeError, match="WPS_UNPUBLISH_WRITE_FAILED") as exc_info:
        wps_main._unpublish_addin()

    assert isinstance(exc_info.value.__cause__, OSError)


def test_wps_request_context_and_preview_safety_contracts_are_present():
    root = Path(__file__).resolve().parents[1]
    host = (root / "host-runtime.js").read_text(encoding="utf-8")
    taskpane = (root / "taskpane.js").read_text(encoding="utf-8")

    assert "currentRequestId" not in host
    for token in (
        'schema_version !== "wps-command-v2"',
        "/v1/bridge/host/register",
        "/v1/bridge/host/wait",
        "/v1/bridge/state",
        "hostContextId",
    ):
        assert token in host
    for token in (
        "/v1/bridge/command",
        "/v1/bridge/state/wait",
        "hostGeneration",
        "stateRevision",
        "pendingRequestId",
        "REQUEST_ACK_TIMEOUT",
    ):
        assert token in taskpane
    for token in (
        "host.bridge.command.schema_invalid",
        "host.bridge.command.request_id_missing",
        "host.bridge.command.command_missing",
        "taskpane.request.completed",
        "host.start.rollback",
        "PREVIEW_DOCUMENT_CHANGED",
        "preview.range.revalidate.failed",
        "document.context.changed",
        "taskpane.rebuild.failed",
    ):
        assert token in host
    assert "config.sessionId" not in host
    assert "config.sessionId" not in taskpane
    for obsolete in (
        "wps-request-v2",
        "docxtool_wps_state_v1",
        "docxtool_wps_request_v1",
        "setInterval(",
    ):
        assert obsolete not in host
        assert obsolete not in taskpane
    assert host.index("range = await previewRange(document, item)") < host.index("comments.Add(range")


def test_wps_entry_registers_ribbon_bridges_before_loading_children():
    root = Path(__file__).resolve().parents[1]
    main_source = (root / "main.js").read_text(encoding="utf-8")
    ribbon_source = (root / "js" / "ribbon.js").read_text(encoding="utf-8")

    first_child_load = main_source.index('loadScript("js/bootstrap-log.js"')
    for callback in ("OnAddinLoad", "OnAction", "GetActionEnabled", "GetImage"):
        declaration = main_source.index(f"function {callback}(")
        assert declaration < first_child_load
    assert "DocxToolRibbonCallbacks" in main_source
    assert "window.DocxToolRibbonCallbacks" in ribbon_source


def test_wps_diagnostic_event_contract_is_present():
    root = Path(__file__).resolve().parents[1]
    sources = {
        "main": (root / "main.js").read_text(encoding="utf-8"),
        "ribbon": (root / "js" / "ribbon.js").read_text(encoding="utf-8"),
        "host": (root / "host-runtime.js").read_text(encoding="utf-8"),
        "taskpane": (root / "taskpane.js").read_text(encoding="utf-8"),
    }
    expected = {
        "main": (
            "bootstrap.main.loaded",
            '"bootstrap." + label + ".load.start"',
            '"bootstrap." + label + ".loaded"',
            '"bootstrap." + label + ".failed"',
            '"bootstrap_log"',
                "bootstrap.runtime_config.load.start",
                "bootstrap.runtime_config.loaded",
            '"ribbon"',
            '"host_runtime"',
        ),
        "ribbon": (
            "ribbon.addin.load.enter",
            "ribbon.addin.host_start.call",
            "ribbon.addin.host_start.returned",
            "ribbon.addin.load.completed",
            "ribbon.addin.load.failed",
        ),
        "host": (
            "host.start.enter",
            "host.start.lazy.enter",
            "host.start.lazy.scheduled",
            "host.start.lazy.failed",
            "host.start.failed",
            "host.bridge.register.start",
            "host.bridge.register.completed",
            "host.bridge.wait.started",
            "host.bridge.wait.failed",
            "host.bridge.command.received",
            "host.bridge.state.published",
            "host.bridge.state.publish_failed",
            "host.bridge.state.flush_failed",
            "taskpane.storage_id.read_failed",
            "taskpane.storage_id.write_failed",
            "taskpane.create_call.failed",
            "taskpane.show.failed",
            "taskpane.width.failed",
            "document.format.detected",
            "document.upgrade.start",
            "document.upgrade.save_as.failed",
            "document.upgrade.verify.failed",
            "document.upgrade.publish.start",
            "document.upgrade.reopen.failed",
            "document.upgrade.completed",
            "document.upgrade.rollback.failed",
            "document.upgrade.rollback.completed",
            "preview.start",
            "preview.document_path.wait.start",
            "preview.document_path.wait.completed",
            "preview.document_path.wait.failed",
            "preview.range_selection.start",
            "preview.range_selection.completed",
            "preview.range.binding_unconfirmed",
            "preview.range.paragraph_unresolved",
            "preview.range.offset_invalid",
            "preview.range.paragraph_lookup_failed",
            "preview.range.paragraph_missing",
            "preview.range.paragraph_text_failed",
            "preview.range.paragraph_changed",
            "preview.range.fragment_mismatch",
            "preview.range.characters_unsupported",
            "preview.range.utf16_boundary_invalid",
            "preview.range.character_lookup_failed",
            "preview.range.boundary_invalid",
            "preview.range.set_failed",
            "preview.range.readback_failed",
            "preview.range.readback_mismatch",
            "preview.comment.create_call.failed",
            "preview.comment.create_result.empty",
            "preview.comment.metadata.failed",
            "preview.session.write_failed",
            "preview.comment_cleanup.item.failed",
            "preview.comments.rollback.completed",
            "preview.comments.rollback.failed",
            "preview.completed",
            "preview.failed",
            "format.start",
            "format.completed",
            "format.failed",
            "transaction.recovery.start",
            "transaction.recovery.completed",
            "transaction.recovery.failed",
            "host.command.failure_state.failed",
            "host.command.failure_panel_open.failed",
            "ribbon.invalidate.failed",
            "log.transport.unavailable",
        ),
        "taskpane": (
            "taskpane.request.prepare",
            "taskpane.request.blocked.busy",
            "taskpane.request.blocked.host_not_ready",
            "taskpane.request.claimed",
            "taskpane.request.completed",
            "taskpane.bridge.command.submit.start",
            "taskpane.bridge.command.submit.completed",
            "taskpane.bridge.command.submit.failed",
            "taskpane.bridge.state.wait.started",
            "taskpane.bridge.state.received",
            "taskpane.bridge.state.wait.failed",
            "taskpane.bridge.state.wait.stopped",
            "taskpane.bridge.host_generation.changed",
            "taskpane.load.failed",
            "log.transport.unavailable",
        ),
    }
    for source_name, events in expected.items():
        for event in events:
            assert event in sources[source_name], f"missing {event} in {source_name}"


def test_wps_python_diagnostic_event_contract_is_present():
    root = Path(__file__).resolve().parents[1]
    sources = {
        "monitor": (root / "control" / "monitor.py").read_text(encoding="utf-8"),
        "server": (
            (root / "control" / "server.py").read_text(encoding="utf-8")
            + (root / "control" / "transport" / "protocol.py").read_text(
                encoding="utf-8"
            )
        ),
        "recognition": (root / "control" / "recognize_document.py").read_text(
            encoding="utf-8"
        ),
        "format": (root / "control" / "format_current_document.py").read_text(
            encoding="utf-8"
        ),
        "transaction": (root / "control" / "document_transaction.py").read_text(
            encoding="utf-8"
        ),
        "launcher": (root / "main.py").read_text(encoding="utf-8"),
    }
    expected = {
        "monitor": (
            "monitor.thread.started",
            "monitor.thread.crashed",
            "monitor.command.received",
            "monitor.command.queued",
            "monitor.command.started",
            "monitor.command.completed",
            "monitor.command.failed",
            "monitor.command.busy",
            "monitor.command.unavailable",
        ),
        "server": (
            "bridge.host.registered",
            "bridge.host.replaced",
            "bridge.command.enqueued",
            "bridge.command.delivered",
            "bridge.state.published",
            "bridge.waiters.closed",
            "control.auth.rejected",
            "control.body.length_invalid",
            "control.body.length_negative",
            "control.body.too_large",
            "control.body.truncated",
            "control.body.json_invalid",
            "control.body.object_required",
            "control.route.not_found",
            "control.response.write_failed",
            "format.upgrade.prepare_converted.request.start",
            "format.upgrade.prepare_converted.request.failed",
            "format.upgrade.prepare_converted.request.completed",
        ),
        "recognition": (
            "recognition.input.invalid",
            "recognition.start",
            "recognition.pipeline.failed",
            "recognition.completed",
            "binding.start",
            "binding.sdk.failed",
            "binding.block_missing",
            "binding.item",
            "binding.item.warning",
            "binding.completed",
        ),
        "format": (
            "input.source.invalid",
            "input.output.invalid",
            "input.path_collision",
            "config.load.failed",
            "config.processing.invalid",
            "import.failed",
            "engine.export.failed",
            "integrity.validate.failed",
            "pipeline.completed",
        ),
        "transaction": (
            "prepare.rejected.busy",
            "prepare.source_hash.failed",
            "prepare.output_hash.failed",
            "prepare.source_changed",
            "upgrade.prepare_converted.start",
            "upgrade.prepare_converted.copy.failed",
            "upgrade.prepare_converted.completed",
            "commit.source.missing",
            "commit.source.changed",
            "commit.output.missing",
            "backup.copy.failed",
            "backup.verify.failed",
            "source.replace.failed",
            "source.replace.verify_failed",
            "finalize.cleanup.failed",
            "rollback.backup.missing",
            "rollback.backup.mismatch",
            "rollback.source.ambiguous",
            "rollback.temporary_cleanup.failed",
            "rollback.source_replace.failed",
            "rollback.backup_cleanup.failed",
            "rollback.journal_clear.failed",
            "transaction.journal.read.failed",
            "transaction.journal.parse.failed",
            "transaction.journal.schema.invalid",
            "transaction.journal.paths.invalid",
            "transaction.journal.hashes.invalid",
            "transaction.recovery.source_state.failed",
            "transaction.recovery.temporary_state.failed",
            "transaction.recovery.backup_state.failed",
            "transaction.recovery.prepared_state.invalid",
            "transaction.recovery.committed_state.invalid",
            "transaction.recovery.state.invalid",
            "transaction.recovery.temporary_cleanup.failed",
            "transaction.recovery.source_restore.failed",
            "transaction.recovery.backup_cleanup.failed",
            "transaction.recovery.journal_clear.failed",
            "transaction.recovery.failed",
        ),
        "launcher": (
            "launcher.control.create.failed",
            "launcher.runtime_config.publish.start",
            "launcher.runtime_config.publish.completed",
            "launcher.publish.parse.failed",
            "launcher.publish.schema.failed",
            "launcher.publish.write.failed",
            "launcher.web.create.failed",
            "launcher.web.previous_service.detected",
            "launcher.web.previous_service.stop.failed",
            "launcher.web.previous_service.stop.completed",
            "launcher.web.serve.failed",
            "launcher.control.thread.failed",
            "launcher.account_runtime.stop.failed",
            "launcher.web.close.failed",
            "launcher.control.shutdown.failed",
            "launcher.control.close.failed",
            "launcher.control.thread.join.failed",
            "launcher.control.thread.state_check.failed",
            "launcher.control.thread.stop_timeout",
            "launcher.runtime_config.cleanup.failed",
        ),
    }
    for source_name, events in expected.items():
        for event in events:
            assert event in sources[source_name], f"missing {event} in {source_name}"


@pytest.mark.parametrize(
    ("argv", "expected"),
    [
        (["main.py"], "start"),
        (["main.py", "start"], "start"),
        (["main.py", "control"], "control"),
        (["main.py", "verify"], "verify"),
    ],
)
def test_main_action_defaults_and_routes(monkeypatch, argv, expected):
    calls = []

    def fake_desktop(port, *, force_login=False):
        calls.append(("start", port, force_login))
        return 0

    def fake_control(port):
        calls.append(("control", port))

    def fake_verify():
        calls.append(("verify",))

    monkeypatch.setattr(wps_main, "run_desktop", fake_desktop)
    monkeypatch.setattr(wps_main, "control_only", fake_control)
    monkeypatch.setattr(wps_main, "verify_files", fake_verify)
    monkeypatch.setattr(wps_main.sys, "argv", argv)
    assert wps_main.main() == 0
    if expected == "start":
        assert calls == [("start", wps_main.DEFAULT_PORT, False)]
    elif expected == "control":
        assert calls == [("control", wps_main.DEFAULT_PORT)]
    else:
        assert calls == [("verify",)]


def test_main_closing_login_window_stops_before_start(monkeypatch):
    application = type("Application", (), {})()
    instance = type("Instance", (), {"acquire": lambda self: True})()
    calls = []
    monkeypatch.setattr("apps.wps.desktop_runtime.ensure_application", lambda: application)
    monkeypatch.setattr("apps.wps.desktop_runtime.SingleInstance", lambda: instance)
    monkeypatch.setattr("apps.wps.desktop_runtime.shift_pressed", lambda: False)
    monkeypatch.setattr(wps_main, "WpsPublicApi", lambda: "api")
    monkeypatch.setattr(wps_main, "_unpublish_addin", lambda: calls.append("unpublish"))
    monkeypatch.setattr(
        "apps.wps.desktop_runtime.DesktopController",
        lambda **_kwargs: pytest.fail("DesktopController must not be created"),
    )
    monkeypatch.setattr(
        wps_main,
        "resolve_startup_account",
        lambda _api, force_login=False: {},
    )

    assert wps_main.run_desktop(wps_main.DEFAULT_PORT) == 0
    assert calls == ["unpublish"]


def test_main_starts_services_only_after_login_returns_an_account(monkeypatch):
    calls = []

    class Signal:
        def connect(self, callback):
            calls.append(("connect", callback))

    class Instance:
        show_requested = Signal()

        def acquire(self):
            return True

    class Application:
        def exec_(self):
            calls.append("exec")
            return 0

    class Controller:
        restart_login_requested = False

        def __init__(self, **kwargs):
            calls.append(("controller", kwargs["account_runtime"].summary()["username"]))

        def show_settings(self):
            return None

        def start(self):
            calls.append("start")

        def shutdown(self):
            calls.append("shutdown")

    monkeypatch.setattr("apps.wps.desktop_runtime.ensure_application", Application)
    monkeypatch.setattr("apps.wps.desktop_runtime.SingleInstance", Instance)
    monkeypatch.setattr("apps.wps.desktop_runtime.shift_pressed", lambda: False)
    monkeypatch.setattr("apps.wps.desktop_runtime.DesktopController", Controller)
    monkeypatch.setattr(wps_main, "WpsPublicApi", lambda: "api")
    monkeypatch.setattr(
        wps_main,
        "resolve_startup_account",
        lambda _api, force_login=False: {"username": "User01"},
    )
    monkeypatch.setattr(wps_main, "_unpublish_addin", lambda: calls.append("unpublish"))

    assert wps_main.run_desktop(wps_main.DEFAULT_PORT) == 0
    assert ("controller", "User01") in calls
    assert "start" in calls
    assert "shutdown" in calls
    assert calls.count("unpublish") == 1


def test_logout_stops_services_unpublishes_then_reopens_login(monkeypatch):
    calls = []

    class Signal:
        def connect(self, _callback):
            calls.append("connect")

    class Instance:
        show_requested = Signal()

        def acquire(self):
            return True

        def close(self):
            calls.append("instance-close")

    class Application:
        def exec_(self):
            calls.append("exec")
            return 0

    class Controller:
        restart_login_requested = True

        def __init__(self, **_kwargs):
            calls.append("controller")

        def show_settings(self):
            return None

        def start(self):
            calls.append("start")

        def shutdown(self):
            calls.append("shutdown")

    monkeypatch.setattr("apps.wps.desktop_runtime.ensure_application", Application)
    monkeypatch.setattr("apps.wps.desktop_runtime.SingleInstance", Instance)
    monkeypatch.setattr("apps.wps.desktop_runtime.shift_pressed", lambda: False)
    monkeypatch.setattr("apps.wps.desktop_runtime.DesktopController", Controller)
    monkeypatch.setattr(wps_main, "WpsPublicApi", lambda: "api")
    monkeypatch.setattr(
        wps_main,
        "resolve_startup_account",
        lambda _api, force_login=False: {"username": "User01"},
    )
    monkeypatch.setattr(wps_main, "_unpublish_addin", lambda: calls.append("unpublish"))
    monkeypatch.setattr(
        wps_main.windows_startup, "launch", lambda argument: calls.append(("launch", argument))
    )

    assert wps_main.run_desktop(wps_main.DEFAULT_PORT) == 0
    unpublish_positions = [index for index, value in enumerate(calls) if value == "unpublish"]
    assert len(unpublish_positions) == 2
    assert calls.index("shutdown") < unpublish_positions[1] < calls.index("instance-close")
    assert ("launch", "--force-login") in calls
